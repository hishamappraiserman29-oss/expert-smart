"""
asset_type_reports.py
======================
Asset-type-specific Word report sections.

Each function appends a self-contained section to an open `python-docx` Document,
making it easy to compose with purpose_specific_reports.py to produce a single
report that combines (asset_type × purpose) framing.

Public API:
    add_asset_type_section(doc, asset_type_or_result) -> bool
    add_intangible_section(doc, result)
    add_partial_interest_section(doc, result)
    add_under_construction_section(doc, result)
    add_quarry_section(doc, result)

This module is ADDITIVE — independent of master_word_report.py.
"""

from __future__ import annotations
from typing import Any, Dict, Optional

try:
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _DOCX_OK = True
except Exception:
    _DOCX_OK = False


# Re-export shared style helpers from purpose_specific_reports — but keep local
# fallbacks in case import order varies.

def _para_rtl(p):
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


NAVY  = RGBColor(0x1F, 0x38, 0x64) if _DOCX_OK else None
GOLD  = RGBColor(0xC9, 0xA2, 0x27) if _DOCX_OK else None
GREEN = RGBColor(0x1E, 0x84, 0x49) if _DOCX_OK else None
GREY  = RGBColor(0x55, 0x55, 0x55) if _DOCX_OK else None


def _h2(doc, text: str, color=None):
    p = doc.add_paragraph()
    _para_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(text)
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = color or GOLD
    return p


def _para(doc, text: str, *, bold: bool = False, color=None, size: int = 12):
    p = doc.add_paragraph()
    _para_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    if color is not None:
        r.font.color.rgb = color
    return p


def _disclosure(doc, text: str):
    p = doc.add_paragraph()
    _para_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = GREY


def _kv(doc, rows):
    if not rows:
        return None
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = "Light Grid Accent 1"
    tblPr = tbl._tbl.tblPr
    bidi = OxmlElement("w:bidiVisual")
    tblPr.append(bidi)
    for i, (k, v) in enumerate(rows):
        rc = tbl.rows[i].cells
        rc[0].text = str(k)
        rc[1].text = str(v)
        for c in rc:
            for p in c.paragraphs:
                _para_rtl(p)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(11)
        rc[0].paragraphs[0].runs[0].font.bold = True
    return tbl


def _money(x: Any) -> str:
    try:
        return f"{float(x):,.0f} ج.م"
    except Exception:
        return "—"


def _pct(x: Any) -> str:
    try:
        return f"{float(x) * 100:.2f}%" if abs(float(x)) <= 2 else f"{float(x):.2f}%"
    except Exception:
        return "—"


# ════════════════════════════════════════════════════════════════════════════
#  INTANGIBLE ASSETS
# ════════════════════════════════════════════════════════════════════════════

def add_intangible_section(doc, r: Dict[str, Any]) -> None:
    _h2(doc, "تقييم الأصل المعنوي — MPEEM", color=NAVY)
    _para(doc,
        "طُبِّق نموذج Multi-Period Excess Earnings Method (MPEEM) وفق IFRS 13 "
        "و IVS 210، حيث يُعزى جزء من العائد السنوي للأصل المعنوي بعد خصم "
        "Contributory Asset Charge للأصول الأخرى المساهمة.")
    _kv(doc, [
        ("الإيراد السنوي",                         _money(r.get("annual_revenue"))),
        ("نسبة العائد المنسوبة للأصل المعنوي",     _pct(r.get("intangible_attribution_pct"))),
        ("Contributory Asset Charge",              _pct(r.get("contributory_asset_charge_pct"))),
        ("العائد الإضافي السنوي",                  _money(r.get("annual_excess_earnings"))),
        ("عمر الحق المعنوي (سنوات)",              str(r.get("license_term_years"))),
        ("معدل الخصم",                              _pct(r.get("discount_rate"))),
        ("القيمة الحالية للعائدات",                _money(r.get("pv_excess_earnings"))),
        ("ميزة الإطفاء الضريبي (TAB)",             _pct(r.get("tax_amortization_benefit_pct"))),
        ("القيمة المعتمدة للأصل المعنوي",          _money(r.get("reconciled_value"))),
    ])
    _disclosure(doc,
        "قيمة الأصل المعنوي تخضع لفحص دوري للهبوط (Impairment Test) وفق IAS 36. "
        "أي تغيير جوهري في الإيراد المنسوب أو معدل الخصم يستوجب إعادة التقييم.")


# ════════════════════════════════════════════════════════════════════════════
#  PARTIAL INTERESTS
# ════════════════════════════════════════════════════════════════════════════

def add_partial_interest_section(doc, r: Dict[str, Any]) -> None:
    _h2(doc, "تقييم الملكية الجزئية — Pro-rata × DLOC × DLOM", color=NAVY)
    _para(doc,
        "طُبِّق نهج IVS 200 للحصص الجزئية في الملكية: حساب الحصة على أساس "
        "نسبي من القيمة السوقية للملكية الكاملة، ثم تطبيق خصم عدم السيطرة "
        "(للحصص الأقلية) ثم خصم عدم القابلية للتسويق.")
    is_ctrl = r.get("is_controlling", False)
    _kv(doc, [
        ("القيمة السوقية للملكية الكاملة",     _money(r.get("total_property_value"))),
        ("نسبة الملكية الجزئية",                r.get("ownership_pct_label", "—")),
        ("صفة الحصة",                           "حصة مسيطرة" if is_ctrl else "حصة أقلية"),
        ("القيمة النسبية (Pro-rata)",          _money(r.get("pro_rata_value"))),
        ("DLOC — خصم عدم السيطرة",             _pct(r.get("dloc_pct"))),
        ("قيمة بعد DLOC",                       _money(r.get("value_after_dloc"))),
        ("DLOM — خصم عدم القابلية للتسويق",    _pct(r.get("dlom_pct"))),
        ("القيمة المعتمدة للحصة الجزئية",      _money(r.get("reconciled_value"))),
        ("النسبة الفعلية للقيمة الكاملة",       _pct(r.get("implied_per_unit_pct"))),
    ])
    _disclosure(doc,
        "وفق المادة 825–م. 850 من القانون المدني المصري للشيوع، حق الشريك في "
        "بيع حصته دون موافقة الشركاء الآخرين (مع حق الشفعة لهم). الخصومات "
        "المطبقة تعكس صعوبة التسويق وعدم السيطرة الفعلية على إدارة العقار.")


# ════════════════════════════════════════════════════════════════════════════
#  UNDER-CONSTRUCTION
# ════════════════════════════════════════════════════════════════════════════

def add_under_construction_section(doc, r: Dict[str, Any]) -> None:
    _h2(doc, "تقييم استثمار تحت الإنشاء — Cost-to-date + Risk", color=NAVY)
    _para(doc,
        "طُبِّقت صيغة: التكلفة المنفقة × نسبة الإنجاز − نسبة المخاطرة، "
        "مع تكامل مع نهج Discounted Residual Method عند توفر القيمة السوقية "
        "المخططة بعد الاكتمال (وفق IAS 16 و IFRS 13 و IVS 230).")

    _h2(doc, "Method A — التكلفة المنفقة المخصومة بالمخاطرة")
    _kv(doc, [
        ("التكلفة الكاملة المخططة",            _money(r.get("planned_total_cost"))),
        ("نسبة الإنجاز",                       r.get("completion_pct_label", "—")),
        ("التكلفة المنفقة حتى الآن",          _money(r.get("cost_incurred"))),
        ("نسبة مخاطرة الإنشاء",                _pct(r.get("construction_risk_pct"))),
        ("قيمة المخاطرة المُخصَمة",            _money(r.get("construction_risk_amount"))),
        ("قيمة Method A",                      _money(r.get("value_method_a"))),
    ])

    if r.get("value_method_b") is not None:
        _h2(doc, "Method B — Discounted Residual")
        _kv(doc, [
            ("القيمة السوقية المخططة بعد الاكتمال", _money(r.get("planned_market_value"))),
            ("تكلفة الإكمال المتبقية",              _money(r.get("remaining_cost_to_complete"))),
            ("هامش ربح المطور",                     _pct(r.get("developer_profit_pct"))),
            ("قيمة هامش الربح",                     _money(r.get("developer_profit_amount"))),
            ("الأشهر المتبقية للاكتمال",            str(r.get("months_to_completion"))),
            ("معدل الخصم",                           _pct(r.get("discount_rate"))),
            ("معامل القيمة الحالية",                f"{r.get('pv_factor', 0):.4f}"),
            ("قيمة Method B",                       _money(r.get("value_method_b"))),
        ])

    _h2(doc, "النتيجة المُوفَّقة")
    _para(doc, f"المرجحة: {r.get('weighting', '—')}", bold=True)
    _para(doc, f"القيمة المعتمدة لاستثمار تحت الإنشاء = {_money(r.get('reconciled_value'))}",
          bold=True, color=GREEN, size=14)

    _disclosure(doc,
        "وفقاً لـ IAS 16 §22، تُسجَّل الأصول قيد الإنشاء بالتكلفة (Cost) إلى أن "
        "تصبح صالحة للاستخدام، ثم تُحوَّل إلى Investment Property أو PPE. "
        "نسبة المخاطرة تعكس مخاطر التأخير وتجاوز الكلفة وفشل الترخيص.")


# ════════════════════════════════════════════════════════════════════════════
#  QUARRIES
# ════════════════════════════════════════════════════════════════════════════

def add_quarry_section(doc, r: Dict[str, Any]) -> None:
    _h2(doc, "تقييم منجم — DCF على الاحتياطي المؤكد", color=NAVY)
    _para(doc,
        "طُبِّق نهج خصم التدفقات النقدية على الاحتياطي المؤكد للمنجم وفق IVS 220 "
        "ومعايير SAMREC / JORC، مع خصم تكلفة إعادة التأهيل وإضافة قيمة الأرض "
        "المتبقية بعد نضوب الاحتياطي.")
    _kv(doc, [
        ("الاحتياطي المؤكد (طن)",        f"{float(r.get('reserve_tons', 0)):,.0f}"),
        ("الاستخراج السنوي (طن)",        f"{float(r.get('annual_extraction', 0)):,.0f}"),
        ("سعر الطن الصافي",              _money(r.get("price_per_ton"))),
        ("العمر الإنتاجي المتوقع (سنة)", f"{float(r.get('life_years', 0)):.1f}"),
        ("الإيراد السنوي",               _money(r.get("annual_revenue"))),
        ("صافي التدفق النقدي السنوي",    _money(r.get("annual_ncf"))),
        ("القيمة الحالية للتدفقات",      _money(r.get("pv_cash_flows"))),
        ("القيمة الحالية لإعادة التأهيل", _money(r.get("pv_rehab_cost"))),
        ("القيمة الحالية للأرض المتبقية", _money(r.get("pv_land_residual"))),
        ("معدل الخصم",                    _pct(r.get("discount_rate"))),
        ("قيمة المنجم المعتمدة",          _money(r.get("reconciled_value"))),
        ("القيمة لكل طن احتياطي",         _money(r.get("value_per_ton_reserve"))),
    ])
    _disclosure(doc,
        "تقدير الاحتياطي يعتمد على مسح جيولوجي معتمد. أي إعادة تصنيف للاحتياطي "
        "(من Possible إلى Probable إلى Proven) تستوجب إعادة التقييم. تكلفة إعادة "
        "التأهيل البيئي إلزامية وفق قانون البيئة المصري رقم 4 لسنة 1994.")


# ════════════════════════════════════════════════════════════════════════════
#  PUBLIC DISPATCHER
# ════════════════════════════════════════════════════════════════════════════

_TYPE_TO_BUILDER = {
    "intangible":          add_intangible_section,
    "partial_interest":    add_partial_interest_section,
    "under_construction":  add_under_construction_section,
    "quarry":              add_quarry_section,
}


def add_asset_type_section(doc, asset_result: Dict[str, Any]) -> bool:
    """
    يُضيف قسم Word مخصصاً لنوع الأصل بناءً على asset_result['asset_type'].
    يُرجع True عند الإضافة، False إذا لم يُعرَف النوع.
    """
    if not _DOCX_OK or not isinstance(asset_result, dict):
        return False
    at = (asset_result.get("asset_type") or "").strip().lower()
    builder = _TYPE_TO_BUILDER.get(at)
    if not builder:
        return False
    try:
        builder(doc, asset_result)
        return True
    except Exception as e:
        print(f"[asset_type_reports] {at} section failed: {e}")
        return False
