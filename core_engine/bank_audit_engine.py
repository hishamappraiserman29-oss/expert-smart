"""
bank_audit_engine.py
====================
Expert_Smart PropTech — Central Bank / Basel III/IV Compliance Module
Real Estate Collateral Auditing Only.

Functions:
    run_bank_audit(payload)                          → dict
    generate_bank_audit_excel(result, output_dir)   → str
    generate_bank_audit_word_section(doc, result)
"""
from __future__ import annotations
import os
from datetime import datetime, date
from typing import Dict

_BASE_DIR = os.path.dirname(os.path.abspath(__file__))
_OUT_DIR = os.path.join(_BASE_DIR, "outputs", "reports")

try:
    import xlsxwriter
    _XLSXWRITER_AVAILABLE = True
except ImportError:
    _XLSXWRITER_AVAILABLE = False

try:
    from docx import Document as _DocxDocument
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import lxml.etree as _lxml_etree
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

# ── Non-real-estate asset types (strictly excluded) ──────────────────────────
_EXCLUDED_TYPES = {"machinery", "business", "equipment", "vehicle", "aircraft"}

# ── Basel III LTV thresholds by loan type ────────────────────────────────────
_LTV_THRESHOLDS: Dict[str, float] = {
    "residential": 0.80,
    "commercial":  0.60,
    "investment":  0.65,
}

# ── Risk band labels ──────────────────────────────────────────────────────────
_RISK_LOW      = "منخفض (مقبول)"
_RISK_MEDIUM   = "متوسط (مراقبة)"
_RISK_HIGH     = "مرتفع (تحذير)"
_RISK_CRITICAL = "حرج (تدخل فوري)"

_ACTION_LOW      = "لا إجراء مطلوب — الضمان سليم"
_ACTION_MEDIUM   = "مراقبة دورية — إعادة تقييم خلال 12 شهراً"
_ACTION_HIGH     = "مطلوب زيادة الضمان أو تخفيض القرض"
_ACTION_CRITICAL = "إشعار فوري للجهة المقرضة — Basel III Article 124"


# ─────────────────────────────────────────────────────────────────────────────
# Core audit function
# ─────────────────────────────────────────────────────────────────────────────

def run_bank_audit(payload: dict) -> dict:
    """
    Run a Basel III/IV compliant bank audit on real-estate collateral.

    Parameters
    ----------
    payload : dict
        Required keys: property_type, location, area, current_market_value,
        original_valuation_value, original_valuation_date, loan_amount, loan_type.
        Optional keys: client_name, bank_name.

    Returns
    -------
    dict
        Full audit result with LTV metrics, risk classification, and audit trail.

    Raises
    ------
    ValueError
        If property_type is a non-real-estate asset class.
    """
    # ── 0. Input validation ───────────────────────────────────────────────────
    property_type: str = str(payload.get("property_type", "")).strip().lower()
    if property_type in _EXCLUDED_TYPES:
        raise ValueError(
            f"property_type='{payload['property_type']}' is not a real-estate asset. "
            f"bank_audit_engine only handles real estate collateral. "
            f"Excluded types: {sorted(_EXCLUDED_TYPES)}"
        )

    location:                str   = str(payload.get("location", "غير محدد"))
    area:                    float = float(payload.get("area", 0.0))
    current_market_value:    float = float(payload.get("current_market_value", 0.0))
    original_valuation_value: float = float(payload.get("original_valuation_value", 0.0))
    original_valuation_date:  str  = str(payload.get("original_valuation_date", ""))
    loan_amount:             float = float(payload.get("loan_amount", 0.0))
    loan_type:               str   = str(payload.get("loan_type", "residential")).strip().lower()
    client_name:             str   = str(payload.get("client_name", "غير مذكور"))
    bank_name:               str   = str(payload.get("bank_name", "البنك الأهلي المصري"))

    if current_market_value <= 0:
        raise ValueError("current_market_value must be > 0")
    if original_valuation_value <= 0:
        raise ValueError("original_valuation_value must be > 0")
    if loan_amount < 0:
        raise ValueError("loan_amount must be >= 0")
    if loan_type not in _LTV_THRESHOLDS:
        raise ValueError(
            f"loan_type='{loan_type}' is not recognised. "
            f"Valid values: {list(_LTV_THRESHOLDS.keys())}"
        )

    # ── 1. LTV Calculation ────────────────────────────────────────────────────
    ltv_ratio:     float = loan_amount / current_market_value
    ltv_threshold: float = _LTV_THRESHOLDS[loan_type]
    ltv_breach:    bool  = ltv_ratio > ltv_threshold

    # ── 2. Collateral Decay ───────────────────────────────────────────────────
    collateral_decay_pct: float = (
        (original_valuation_value - current_market_value) / original_valuation_value * 100
    )

    # Parse original_valuation_date
    try:
        if isinstance(original_valuation_date, date):
            orig_date = original_valuation_date
        else:
            orig_date = date.fromisoformat(original_valuation_date)
    except (ValueError, TypeError):
        orig_date = date.today()

    today = date.today()
    days_since = (today - orig_date).days
    years_since_valuation: float = days_since / 365.25

    annual_decay_rate: float = (
        collateral_decay_pct / years_since_valuation if years_since_valuation > 0 else 0.0
    )
    revaluation_required: bool = years_since_valuation > 1.0

    # ── 3. Risk Classification ────────────────────────────────────────────────
    t  = ltv_threshold
    t85 = t * 0.85
    t110 = t * 1.10

    if ltv_ratio <= t85:
        risk_level      = _RISK_LOW
        required_action = _ACTION_LOW
    elif ltv_ratio <= t:
        risk_level      = _RISK_MEDIUM
        required_action = _ACTION_MEDIUM
    elif ltv_ratio <= t110:
        risk_level      = _RISK_HIGH
        required_action = _ACTION_HIGH
    else:
        risk_level      = _RISK_CRITICAL
        required_action = _ACTION_CRITICAL

    # ── 4. Audit Trail ────────────────────────────────────────────────────────
    now = datetime.now()
    audit_id             = "AUDIT-" + now.strftime("%Y%m%d-%H%M%S")
    audit_timestamp      = now.isoformat()
    auditor_signature    = "Expert_Smart AI v2 | IVS 2022 | Basel III/IV Compliant"

    # ── 5. Build result ───────────────────────────────────────────────────────
    result: dict = {
        # Echo inputs
        "property_type":            payload.get("property_type"),
        "location":                 location,
        "area":                     area,
        "current_market_value":     current_market_value,
        "original_valuation_value": original_valuation_value,
        "original_valuation_date":  original_valuation_date,
        "loan_amount":              loan_amount,
        "loan_type":                loan_type,
        "client_name":              client_name,
        "bank_name":                bank_name,
        # LTV
        "ltv_ratio":                round(ltv_ratio, 6),
        "ltv_ratio_pct":            round(ltv_ratio * 100, 4),
        "ltv_threshold":            ltv_threshold,
        "ltv_threshold_pct":        round(ltv_threshold * 100, 2),
        "ltv_breach":               ltv_breach,
        # Collateral Decay
        "collateral_decay_pct":     round(collateral_decay_pct, 4),
        "years_since_valuation":    round(years_since_valuation, 4),
        "annual_decay_rate":        round(annual_decay_rate, 4),
        "revaluation_required":     revaluation_required,
        # Risk
        "risk_level":               risk_level,
        "required_action":          required_action,
        # Audit Trail
        "audit_id":                 audit_id,
        "audit_timestamp":          audit_timestamp,
        "auditor_signature":        auditor_signature,
    }
    return result


# ─────────────────────────────────────────────────────────────────────────────
# Excel report
# ─────────────────────────────────────────────────────────────────────────────

def generate_bank_audit_excel(result: dict, output_dir: str = "") -> str:
    """
    Generate a color-coded Basel III/IV bank-audit Excel report.

    Parameters
    ----------
    result      : dict  — output of run_bank_audit()
    output_dir  : str   — override directory (defaults to _OUT_DIR)

    Returns
    -------
    str — absolute path of saved .xlsx file
    """
    if not _XLSXWRITER_AVAILABLE:
        raise ImportError(
            "xlsxwriter is not installed. Run: pip install xlsxwriter"
        )

    out_dir = output_dir.strip() if output_dir.strip() else _OUT_DIR
    os.makedirs(out_dir, exist_ok=True)

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"bank_audit_{ts}.xlsx"
    filepath = os.path.join(out_dir, filename)

    wb  = xlsxwriter.Workbook(filepath, {"strings_to_urls": False})
    ws  = wb.add_worksheet("تقرير المراقبة الائتمانية")
    ws.set_tab_color("#1F4E78")
    ws.right_to_left()

    # ── Shared formats ────────────────────────────────────────────────────────
    _FONT = "Simplified Arabic"

    def _fmt(props: dict):
        base = {"font_name": _FONT, "font_size": 11, "align": "right", "valign": "vcenter"}
        base.update(props)
        return wb.add_format(base)

    fmt_title = _fmt({
        "font_size": 16, "bold": True, "font_color": "#FFFFFF",
        "bg_color": "#1F4E78", "align": "center", "border": 1,
    })
    fmt_subtitle = _fmt({
        "font_size": 12, "bold": True, "font_color": "#1F4E78",
        "bg_color": "#D6E4F0", "border": 1,
    })
    fmt_header = _fmt({
        "font_size": 11, "bold": True, "font_color": "#FFFFFF",
        "bg_color": "#1F4E78", "border": 1,
    })
    fmt_label = _fmt({"bold": True, "border": 1, "bg_color": "#F2F2F2"})
    fmt_value = _fmt({"border": 1})
    fmt_value_num = _fmt({"border": 1, "num_format": "#,##0.00"})
    fmt_value_pct = _fmt({"border": 1, "num_format": "0.00%"})

    # Risk color formats
    risk_level = result.get("risk_level", "")
    ltv_ratio  = result.get("ltv_ratio", 0)
    threshold  = result.get("ltv_threshold", 0.80)

    if ltv_ratio <= threshold * 0.85:
        gauge_color = "#00B050"   # green
    elif ltv_ratio <= threshold * 1.10:
        gauge_color = "#FFC000"   # yellow/amber
    else:
        gauge_color = "#FF0000"   # red

    fmt_risk = _fmt({
        "font_size": 13, "bold": True, "font_color": "#FFFFFF",
        "bg_color": gauge_color, "border": 2, "align": "center",
    })
    fmt_audit_box = _fmt({
        "font_size": 10, "italic": True, "font_color": "#1F4E78",
        "bg_color": "#EBF5FB", "border": 1, "text_wrap": True, "align": "center",
    })
    fmt_section = _fmt({
        "font_size": 13, "bold": True, "font_color": "#7F6000",
        "bg_color": "#FFF2CC", "border": 1,
    })

    # ── Column widths ─────────────────────────────────────────────────────────
    ws.set_column("A:A", 6)
    ws.set_column("B:B", 38)
    ws.set_column("C:C", 32)
    ws.set_column("D:D", 20)
    ws.set_column("E:E", 20)
    ws.set_row(0, 40)
    ws.set_row(1, 26)

    row = 0

    # ── Title row ─────────────────────────────────────────────────────────────
    ws.merge_range(row, 0, row, 4,
                   "تقرير مراقبة الضمان العقاري — Basel III/IV",
                   fmt_title)
    row += 1

    ws.merge_range(row, 0, row, 4,
                   f"البنك: {result.get('bank_name')}    |    "
                   f"العميل: {result.get('client_name')}    |    "
                   f"رقم التدقيق: {result.get('audit_id')}",
                   fmt_subtitle)
    row += 1
    row += 1  # blank

    # ── Section 1 — Property Info ─────────────────────────────────────────────
    ws.merge_range(row, 0, row, 4, "أولاً: بيانات العقار والقرض", fmt_section)
    row += 1

    prop_rows = [
        ("نوع العقار",          result.get("property_type"),          False),
        ("الموقع",              result.get("location"),                False),
        ("المساحة (م²)",        result.get("area"),                    True),
        ("القيمة السوقية الحالية (جنيه)", result.get("current_market_value"), True),
        ("قيمة التقييم الأصلية (جنيه)",   result.get("original_valuation_value"), True),
        ("تاريخ التقييم الأصلي",           result.get("original_valuation_date"),  False),
        ("مبلغ القرض القائم (جنيه)",        result.get("loan_amount"),   True),
        ("نوع التمويل",          result.get("loan_type"),               False),
    ]

    ws.write(row, 1, "البيان",  fmt_header)
    ws.write(row, 2, "القيمة",  fmt_header)
    row += 1

    for label, val, is_num in prop_rows:
        ws.write(row, 1, label, fmt_label)
        if is_num:
            ws.write_number(row, 2, float(val) if val is not None else 0.0, fmt_value_num)
        else:
            ws.write(row, 2, str(val) if val is not None else "", fmt_value)
        row += 1

    row += 1

    # ── Section 2 — LTV Gauge ─────────────────────────────────────────────────
    ws.merge_range(row, 0, row, 4, "ثانياً: نسبة القرض إلى القيمة (LTV) — معيار بازل III", fmt_section)
    row += 1

    ws.write(row, 1, "المؤشر",          fmt_header)
    ws.write(row, 2, "القيمة",          fmt_header)
    ws.write(row, 3, "الحد المسموح",    fmt_header)
    ws.write(row, 4, "الحالة",          fmt_header)
    row += 1

    ltv_pct       = result.get("ltv_ratio_pct", 0)
    thresh_pct    = result.get("ltv_threshold_pct", 0)
    ltv_breach    = result.get("ltv_breach", False)

    ws.write(row, 1, "نسبة LTV", fmt_label)
    ws.write(row, 2, f"{ltv_pct:.2f}%", fmt_value)
    ws.write(row, 3, f"{thresh_pct:.2f}%", fmt_value)
    ws.write(row, 4, "خرق ✗" if ltv_breach else "ضمن الحد ✓", fmt_risk)
    row += 1

    ws.write(row, 1, "مستوى المخاطرة", fmt_label)
    ws.merge_range(row, 2, row, 4, risk_level, fmt_risk)
    row += 1

    row += 1

    # ── Section 3 — Collateral Decay ─────────────────────────────────────────
    ws.merge_range(row, 0, row, 4, "ثالثاً: تحليل تآكل قيمة الضمان", fmt_section)
    row += 1

    decay_rows = [
        ("نسبة تآكل القيمة (%)",         f"{result.get('collateral_decay_pct', 0):.4f}%"),
        ("المدة منذ آخر تقييم (سنة)",    f"{result.get('years_since_valuation', 0):.2f}"),
        ("معدل التآكل السنوي (%)",        f"{result.get('annual_decay_rate', 0):.4f}%"),
        ("إعادة التقييم مطلوبة (>1 سنة)", "نعم ✓" if result.get("revaluation_required") else "لا"),
    ]

    ws.write(row, 1, "المؤشر",  fmt_header)
    ws.write(row, 2, "القيمة",  fmt_header)
    row += 1

    for label, val in decay_rows:
        ws.write(row, 1, label, fmt_label)
        ws.write(row, 2, str(val), fmt_value)
        row += 1

    row += 1

    # ── Section 4 — Required Action ──────────────────────────────────────────
    ws.merge_range(row, 0, row, 4, "رابعاً: الإجراء المطلوب", fmt_section)
    row += 1

    fmt_action = _fmt({
        "font_size": 12, "bold": True, "font_color": "#FFFFFF",
        "bg_color": gauge_color, "border": 2, "text_wrap": True, "align": "center",
    })
    ws.set_row(row, 30)
    ws.merge_range(row, 0, row, 4, result.get("required_action", ""), fmt_action)
    row += 1
    row += 1

    # ── Audit Trail Box ───────────────────────────────────────────────────────
    ws.merge_range(row, 0, row, 4, "سجل التدقيق — Audit Trail", fmt_header)
    row += 1

    audit_lines = [
        f"رقم التدقيق: {result.get('audit_id')}",
        f"وقت التدقيق: {result.get('audit_timestamp')}",
        f"توقيع المدقق: {result.get('auditor_signature')}",
        f"البنك: {result.get('bank_name')}    |    العميل: {result.get('client_name')}",
    ]
    for line in audit_lines:
        ws.set_row(row, 20)
        ws.merge_range(row, 0, row, 4, line, fmt_audit_box)
        row += 1

    wb.close()
    return os.path.abspath(filepath)


# ─────────────────────────────────────────────────────────────────────────────
# Word section generator
# ─────────────────────────────────────────────────────────────────────────────

def _set_rtl_paragraph(paragraph) -> None:
    """Force paragraph to RTL via XML bidi element."""
    if not _DOCX_AVAILABLE:
        return
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


def _para_font(paragraph, font_name: str = "Simplified Arabic",
               size_pt: int = 11, bold: bool = False,
               color_rgb: tuple | None = None) -> None:
    """Apply font settings to all runs in a paragraph, or add a blank run."""
    if not _DOCX_AVAILABLE:
        return
    if not paragraph.runs:
        paragraph.add_run()
    for run in paragraph.runs:
        run.font.name        = font_name
        run.font.size        = Pt(size_pt)
        run.font.bold        = bold
        if color_rgb:
            run.font.color.rgb = RGBColor(*color_rgb)
        # Arabic font hints
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:cs"), font_name)
        rPr.insert(0, rFonts)


def _set_cell_bg(cell, hex_color: str) -> None:
    """Set table cell background fill colour."""
    if not _DOCX_AVAILABLE:
        return
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    tcPr.append(shd)


def _risk_color(risk_level: str) -> tuple:
    """Return (hex_bg, RGBColor) for a given risk level string."""
    if "منخفض" in risk_level:
        return "#00B050", (0, 176, 80)
    elif "متوسط" in risk_level:
        return "#FFC000", (255, 192, 0)
    elif "مرتفع" in risk_level:
        return "#FF6600", (255, 102, 0)
    else:  # حرج
        return "#FF0000", (255, 0, 0)


def generate_bank_audit_word_section(doc, result: dict) -> None:
    """
    Append a Basel III/IV audit section to an existing python-docx Document.

    Parameters
    ----------
    doc    : docx.Document — existing document to append to
    result : dict          — output of run_bank_audit()
    """
    if not _DOCX_AVAILABLE:
        raise ImportError(
            "python-docx is not installed. Run: pip install python-docx"
        )

    _NAVY   = (31, 78, 120)
    _GOLD   = (191, 144, 0)
    _WHITE  = (255, 255, 255)
    _FONT   = "Simplified Arabic"

    risk_level  = result.get("risk_level", "")
    risk_hex, risk_rgb = _risk_color(risk_level)
    ltv_pct    = result.get("ltv_ratio_pct", 0)
    thresh_pct = result.get("ltv_threshold_pct", 0)

    # ── Section heading ───────────────────────────────────────────────────────
    h = doc.add_heading("", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_rtl_paragraph(h)
    run_h = h.add_run("تقرير الائتمان والضمان العقاري — بازل III/IV")
    run_h.font.name  = _FONT
    run_h.font.size  = Pt(16)
    run_h.font.bold  = True
    run_h.font.color.rgb = RGBColor(*_NAVY)
    rPr = run_h._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:cs"), _FONT)
    rPr.insert(0, rFonts)

    doc.add_paragraph()  # spacer

    # ── Risk Status badge ─────────────────────────────────────────────────────
    risk_para = doc.add_paragraph()
    risk_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_rtl_paragraph(risk_para)
    risk_run = risk_para.add_run(f"  مستوى المخاطرة: {risk_level}  ")
    risk_run.font.name  = _FONT
    risk_run.font.size  = Pt(14)
    risk_run.font.bold  = True
    risk_run.font.color.rgb = RGBColor(255, 255, 255)
    # Highlight via paragraph shading
    pPr = risk_para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  risk_hex.lstrip("#"))
    pPr.append(shd)
    rPr2 = risk_run._r.get_or_add_rPr()
    rFonts2 = OxmlElement("w:rFonts")
    rFonts2.set(qn("w:cs"), _FONT)
    rPr2.insert(0, rFonts2)

    doc.add_paragraph()  # spacer

    # ── LTV Table ─────────────────────────────────────────────────────────────
    ltv_label_p = doc.add_paragraph()
    ltv_label_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_rtl_paragraph(ltv_label_p)
    ltv_label_run = ltv_label_p.add_run("نسبة القرض إلى القيمة (LTV) — Basel III")
    ltv_label_run.font.name  = _FONT
    ltv_label_run.font.size  = Pt(13)
    ltv_label_run.font.bold  = True
    ltv_label_run.font.color.rgb = RGBColor(*_GOLD)
    rPr3 = ltv_label_run._r.get_or_add_rPr()
    rFonts3 = OxmlElement("w:rFonts")
    rFonts3.set(qn("w:cs"), _FONT)
    rPr3.insert(0, rFonts3)

    ltv_table = doc.add_table(rows=1, cols=4)
    ltv_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    ltv_table.style = "Table Grid"

    # Header row
    headers = ["المؤشر", "القيمة الفعلية", "الحد المسموح به", "الحالة"]
    hdr_cells = ltv_table.rows[0].cells
    for i, h_text in enumerate(headers):
        hdr_cells[i].text = h_text
        _set_cell_bg(hdr_cells[i], "#1F4E78")
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name  = _FONT
            run.font.size  = Pt(11)
            run.font.bold  = True
            run.font.color.rgb = RGBColor(*_WHITE)
            rPr_ = run._r.get_or_add_rPr()
            rF_ = OxmlElement("w:rFonts")
            rF_.set(qn("w:cs"), _FONT)
            rPr_.insert(0, rF_)

    ltv_data = [
        ("نسبة LTV",          f"{ltv_pct:.2f}%",   f"{thresh_pct:.2f}%",   "خرق ✗" if result.get("ltv_breach") else "ضمن الحد ✓"),
        ("مبلغ القرض (جنيه)",  f"{result.get('loan_amount', 0):,.0f}", "—", "—"),
        ("القيمة السوقية",     f"{result.get('current_market_value', 0):,.0f}", "—", "—"),
    ]

    for row_data in ltv_data:
        cells = ltv_table.add_row().cells
        for j, cell_text in enumerate(row_data):
            cells[j].text = cell_text
            bg = risk_hex if (j == 3 and row_data[0] == "نسبة LTV") else "#F2F2F2" if j == 0 else "#FFFFFF"
            _set_cell_bg(cells[j], bg)
            p = cells[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = _FONT
                run.font.size = Pt(10)
                rPr_ = run._r.get_or_add_rPr()
                rF_ = OxmlElement("w:rFonts")
                rF_.set(qn("w:cs"), _FONT)
                rPr_.insert(0, rF_)

    doc.add_paragraph()  # spacer

    # ── Collateral Decay Table ────────────────────────────────────────────────
    decay_label_p = doc.add_paragraph()
    decay_label_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_rtl_paragraph(decay_label_p)
    decay_label_run = decay_label_p.add_run("مؤشرات تآكل قيمة الضمان")
    decay_label_run.font.name  = _FONT
    decay_label_run.font.size  = Pt(13)
    decay_label_run.font.bold  = True
    decay_label_run.font.color.rgb = RGBColor(*_GOLD)
    rPr4 = decay_label_run._r.get_or_add_rPr()
    rFonts4 = OxmlElement("w:rFonts")
    rFonts4.set(qn("w:cs"), _FONT)
    rPr4.insert(0, rFonts4)

    decay_table = doc.add_table(rows=1, cols=2)
    decay_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    decay_table.style = "Table Grid"

    # Header
    d_hdr = decay_table.rows[0].cells
    for i, txt in enumerate(["المؤشر", "القيمة"]):
        d_hdr[i].text = txt
        _set_cell_bg(d_hdr[i], "#1F4E78")
        p = d_hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name  = _FONT
            run.font.size  = Pt(11)
            run.font.bold  = True
            run.font.color.rgb = RGBColor(*_WHITE)
            rPr_ = run._r.get_or_add_rPr()
            rF_ = OxmlElement("w:rFonts")
            rF_.set(qn("w:cs"), _FONT)
            rPr_.insert(0, rF_)

    decay_data = [
        ("القيمة الأصلية عند منح القرض (جنيه)", f"{result.get('original_valuation_value', 0):,.0f}"),
        ("القيمة السوقية الحالية (جنيه)",         f"{result.get('current_market_value', 0):,.0f}"),
        ("نسبة تآكل القيمة (%)",                  f"{result.get('collateral_decay_pct', 0):.4f}%"),
        ("المدة منذ آخر تقييم (سنة)",              f"{result.get('years_since_valuation', 0):.2f}"),
        ("معدل التآكل السنوي (%)",                 f"{result.get('annual_decay_rate', 0):.4f}%"),
        ("إعادة التقييم مطلوبة",                   "نعم ✓" if result.get("revaluation_required") else "لا"),
    ]

    for label, val in decay_data:
        cells = decay_table.add_row().cells
        cells[0].text = label
        cells[1].text = val
        _set_cell_bg(cells[0], "#F2F2F2")
        for j in range(2):
            p = cells[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = _FONT
                run.font.size = Pt(10)
                rPr_ = run._r.get_or_add_rPr()
                rF_ = OxmlElement("w:rFonts")
                rF_.set(qn("w:cs"), _FONT)
                rPr_.insert(0, rF_)

    doc.add_paragraph()  # spacer

    # ── Required Action ───────────────────────────────────────────────────────
    action_p = doc.add_paragraph()
    action_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_rtl_paragraph(action_p)
    action_run = action_p.add_run(f"الإجراء المطلوب: {result.get('required_action', '')}")
    action_run.font.name  = _FONT
    action_run.font.size  = Pt(12)
    action_run.font.bold  = True
    action_run.font.color.rgb = RGBColor(*_WHITE)
    pPr_a = action_p._p.get_or_add_pPr()
    shd_a = OxmlElement("w:shd")
    shd_a.set(qn("w:val"),   "clear")
    shd_a.set(qn("w:color"), "auto")
    shd_a.set(qn("w:fill"),  risk_hex.lstrip("#"))
    pPr_a.append(shd_a)
    rPr5 = action_run._r.get_or_add_rPr()
    rFonts5 = OxmlElement("w:rFonts")
    rFonts5.set(qn("w:cs"), _FONT)
    rPr5.insert(0, rFonts5)

    doc.add_paragraph()  # spacer

    # ── Audit Trail Certification ─────────────────────────────────────────────
    cert_p = doc.add_paragraph()
    cert_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_rtl_paragraph(cert_p)
    pPr_c = cert_p._p.get_or_add_pPr()
    shd_c = OxmlElement("w:shd")
    shd_c.set(qn("w:val"),   "clear")
    shd_c.set(qn("w:color"), "auto")
    shd_c.set(qn("w:fill"),  "EBF5FB")
    pPr_c.append(shd_c)

    cert_lines = [
        f"رقم التدقيق: {result.get('audit_id')}",
        f"وقت التدقيق: {result.get('audit_timestamp')}",
        f"البنك: {result.get('bank_name')}    |    العميل: {result.get('client_name')}",
        f"توقيع المدقق: {result.get('auditor_signature')}",
    ]
    cert_run = cert_p.add_run("\n".join(cert_lines))
    cert_run.font.name   = _FONT
    cert_run.font.size   = Pt(10)
    cert_run.font.italic = True
    cert_run.font.color.rgb = RGBColor(*_NAVY)
    rPr6 = cert_run._r.get_or_add_rPr()
    rFonts6 = OxmlElement("w:rFonts")
    rFonts6.set(qn("w:cs"), _FONT)
    rPr6.insert(0, rFonts6)

    doc.add_paragraph()  # trailing spacer
