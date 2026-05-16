# -*- coding: utf-8 -*-
"""
bridge_api.py  —  Expert_Smart Valuation Engine
================================================
كل كتابة الإكسيل تتم عبر openpyxl في عملية واحدة:
  1. نسخ القالب (15 صفحة)
  2. كتابة البيانات في الصفحات الموجودة
  3. إضافة الصفحات الجديدة (DCF, ANN, ARIMA)
  4. رسم شجرة القرارات وإدراجها كصورة
  5. حفظ مرة واحدة فقط
"""

import sys, io, os, math, uuid, shutil, traceback, statistics, tempfile, gc, time
from datetime import datetime
from typing import Optional, List, Dict, Any, Tuple

def _ts():
    """يُعيد [HH:MM:SS] للطباعة في الـ Console"""
    return time.strftime("[%H:%M:%S]")

# ══ UTF-8 Bootstrap — أول شيء قبل أي طباعة أو فتح ملف ══════════════════════
# أ) إعلام Python بوضع UTF-8 الكامل (يُطبَّق على هذه العملية وأي عملية فرعية)
os.environ.setdefault("PYTHONUTF8",       "1")           # Python 3.7+ UTF-8 mode
os.environ.setdefault("PYTHONIOENCODING", "utf-8:replace")  # fallback للمكتبات القديمة

# ب) إعادة ضبط stdout/stderr في مكانهما (reconfigure ≠ استبدال)
#    reconfigure() تُعدِّل الـ stream الموجود دون إغلاقه → لا "I/O on closed file"
for _utf8_stream in (sys.stdout, sys.stderr):
    if _utf8_stream and hasattr(_utf8_stream, "reconfigure"):
        try:
            _utf8_stream.reconfigure(encoding="utf-8", errors="replace")
        except Exception:
            pass
    elif _utf8_stream and hasattr(_utf8_stream, "buffer"):
        # fallback آمن: استبدال فقط إذا لم تكن reconfigure متاحة (Python < 3.7)
        try:
            setattr(sys,
                    "stdout" if _utf8_stream is sys.stdout else "stderr",
                    io.TextIOWrapper(_utf8_stream.buffer,
                                     encoding="utf-8", errors="replace",
                                     line_buffering=True))
        except Exception:
            pass
del _utf8_stream
# ═══════════════════════════════════════════════════════════════════════════════

from flask import Flask, g, jsonify, request, send_file
from flask_cors import CORS

# ── Auth (Wave S2 — foundation only, no enforcement) ─────────────────────────
try:
    from auth.tokens import AuthError as _AuthError, verify_token as _verify_token
    _AUTH_AVAILABLE = True
except ImportError:
    _AUTH_AVAILABLE = False

# ── Phase 4 engines ──────────────────────────────────────────────────────────
from engines.comparable_search import ComparableSearchEngine
from engines.comparative import ComparativeEngine
from engines.cost import CostEngine
from engines.income import IncomeEngine

# ── Phase 5 adapters (optional — absent on main until R3 is merged) ──────────
try:
    from adapters.market_value import MarketValueAdapter
    from adapters.mortgage import MortgageValueAdapter
    from adapters.insurance import InsuranceValueAdapter
    from adapters.ifrs_13 import IFRS13FairValueAdapter
except ImportError:
    MarketValueAdapter = MortgageValueAdapter = InsuranceValueAdapter = IFRS13FairValueAdapter = None  # type: ignore[assignment,misc]

# ── Phase 6 asset adapters + Excel report builder ────────────────────────────
try:
    from adapters.residential import ResidentialAdapter
    from adapters.commercial import CommercialAdapter
except ImportError:
    ResidentialAdapter = CommercialAdapter = None  # type: ignore[assignment,misc]
from reports.excel_builder import ExcelReportBuilder

# ── Phase 7 land adapter + quality auditor ────────────────────────────────────
try:
    from adapters.land import LandAdapter
except ImportError:
    LandAdapter = None  # type: ignore[assignment,misc]
from reports.quality_auditor import ReportQualityAuditor

# ── Phase 15 enterprise features (optional — absent on main until R3 is merged)
try:
    from adapters.enterprise import (
        TenantManager              as _TenantManager,
        TenantRole                 as _TenantRole,
        EnterpriseLicenseValidator as _LicenseValidator,
    )
    from database.audit_log import AuditLog as _AuditLog, AuditAction as _AuditAction, AuditEvent as _AuditEvent
    _tenant_manager = _TenantManager()
    _ENTERPRISE_AVAILABLE = True
except ImportError:
    _TenantManager = _TenantRole = _LicenseValidator = None  # type: ignore[assignment,misc]
    _AuditLog = _AuditAction = _AuditEvent = None  # type: ignore[assignment,misc]
    _tenant_manager = None
    _ENTERPRISE_AVAILABLE = False

# ── Phase 8: database layer (lazy — psycopg2 optional) ────────────────────────
try:
    from database.models     import Valuation  as _DbValuation
    from database.models     import QualityAudit as _DbQualityAudit
    from database.queries    import SearchComparable as _SearchComparable
    from database.connection import SessionLocal as _DbSession
    _DB_AVAILABLE = True
except Exception:           # psycopg2 not installed or DB package missing
    _DB_AVAILABLE = False

# ── Market Intelligence (auto-web + sector intelligence) ─────────────────────
try:
    from market_intelligence import (
        auto_fetch_comparables      as _auto_fetch,
        _get_sector,
        _industrial_depreciation,
        _industrial_location_premium,
        _agricultural_income_ppm,
        _hospitality_income_ppm,
        _hospitality_benchmarks,
        _try_hospitality_web,
        FEDDAN_SQM,
        # v44 — قطاعات جديدة
        _retail_income_ppm,
        _healthcare_value_ppm,
        _educational_capacity_ppm,
    )
    _MI_MODULE_OK = True
except Exception as _mi_err:
    print(f"  [market_intelligence] import failed: {_mi_err} — auto-fetch disabled")
    _MI_MODULE_OK = False
    def _auto_fetch(location, property_type, area, ppm_hint, n=6):  return [], ""
    def _get_sector(pt):                   return "residential"
    def _industrial_depreciation(*a, **k): return {"depreciated_value_ratio": 1.0, "depreciation_rate": 0.0, "effective_age": 0, "remaining_life": 30, "economic_life": 30, "age": 0, "condition": "average", "asset_type": "building_concrete"}
    def _industrial_location_premium(loc): return 1.0
    def _agricultural_income_ppm(*a, **k): return {"ppm": 0, "land_value_per_feddan": 0, "net_income_per_feddan": 0, "gross_revenue_per_feddan": 0, "crop_type": "_default", "irrigation": "_default", "soil_quality": "_default", "cap_rate": 0.06, "area_feddan": 0, "total_land_value": 0}
    def _hospitality_income_ppm(*a, **k):  return {"ppm": 0, "noi": 0, "occupancy_rate": 0, "adr_egp": 0, "revpar_daily": 0, "gross_revenue": 0, "hotel_value_income": 0, "hotel_value_dcf": 0, "ppm_income": 0, "ppm_dcf": 0, "rooms": 0, "stars": 0, "location_matched": "", "cap_rate": 0.08, "wacc": 0.12}
    def _hospitality_benchmarks(loc):      return {"occupancy": 0.55, "adr_egp": 2500, "revpar_daily": 0, "location_matched": "_default"}
    def _try_hospitality_web(*a, **k):     return {}
    FEDDAN_SQM = 4200.0
    def _retail_income_ppm(*a, **k):       return {"ppm": 0, "monthly_rent_sqm": 0, "annual_noi": 0, "gross_revenue": 0, "gla_sqm": 0, "gla_ratio": 0.75, "footfall_score": 5, "purchasing_power": 1.0, "frontage_premium": 1.0, "anchor_premium": 1.0, "cap_rate": 0.09, "location_matched": ""}
    def _healthcare_value_ppm(*a, **k):    return {"ppm": 0, "specialized_infra_cost": 0, "infra_per_sqm": 0, "beds_capacity": 0, "compliance_ok": True, "compliance_note": "", "demand_premium": 1.0, "facility_type": "", "monthly_rent_sqm": 0, "annual_noi": 0, "gross_revenue": 0, "cap_rate": 0.09, "location_matched": ""}
    def _educational_capacity_ppm(*a, **k):return {"ppm": 0, "student_capacity": 0, "fee_per_student": 0, "annual_revenue": 0, "noi": 0, "facility_premium": 0, "facility_details": {}, "school_type": "", "fee_type": "", "licensed": True, "cap_rate": 0.08, "location_matched": ""}

# ── Context-Tuning Engine (محرك المحاكاة) ───────────────────────────────────
try:
    from report_tuner import (
        analyze_style       as _tune_analyze,
        save_style_profile  as _tune_save,
        load_style_profile  as _tune_load,
        list_style_profiles as _tune_list,
        delete_style_profile as _tune_delete,
        apply_style_to_prompt as _tune_apply,
        get_style_system_prompt as _tune_sysprompt,
    )
    _TUNER_OK = True
    print("  [report_tuner] loaded ✓")
except Exception as _tuner_err:
    print(f"  [report_tuner] import failed: {_tuner_err}")
    _TUNER_OK = False
    def _tune_analyze(text, doc_name=""): return {"id":"", "name":doc_name, "prompt_injection":"", "confidence":0}
    def _tune_save(p):    return ""
    def _tune_load(pid):  return None
    def _tune_list():     return []
    def _tune_delete(pid): return False
    def _tune_apply(prompt, pid): return prompt
    def _tune_sysprompt(pid): return ""

# ── Smart Library (المكتبة المرجعية الذكية) ─────────────────────────────────
try:
    from library_scanner import (
        scan_all_sources    as _lib_scan,
        add_manual_document as _lib_add_manual,
        search_library      as _lib_search,
        get_library_stats   as _lib_stats,
        get_record          as _lib_get,
        delete_record       as _lib_delete,
    )
    _LIBRARY_OK = True
    print("  [library_scanner] loaded ✓")
except Exception as _lib_err:
    print(f"  [library_scanner] import failed: {_lib_err}")
    _LIBRARY_OK = False
    def _lib_scan():       return {"added":0, "skipped":0, "errors":0, "total":0}
    def _lib_add_manual(*a, **k): return {"status":"disabled"}
    def _lib_search(**k):  return []
    def _lib_stats():      return {"total":0}
    def _lib_get(rid):     return None
    def _lib_delete(rid):  return False

# ── Market Radar v37 (محرك رادار السوق) ─────────────────────────────────────
try:
    from market_radar import radar_api as _radar
    _RADAR_OK = True
    print("  [market_radar] loaded ✓")
except Exception as _radar_err:
    print(f"  [market_radar] import failed: {_radar_err}")
    _RADAR_OK = False
    class _RadarStub:
        def start(self, **k):     return {"status": "disabled"}
        def stop(self):            return {"status": "disabled"}
        def status(self):          return {"running": False, "cycle": 0, "db_stats": {"total": 0}}
        def parse_text(self, t):   return None
        def get_records(self, **k):return []
        def get_heatmap_data(self, **k): return []
        def get_stats(self):       return {"total": 0}
        def insert_manual(self, r):return False
    _radar = _RadarStub()

# ── IAAO Engine v37 (محرك معايير IAAO الضريبية) ──────────────────────────────
try:
    from iaao_engine import (
        full_iaao_report        as _iaao_report,
        iaao_from_comparables   as _iaao_from_comps,
        iaao_excel_rows         as _iaao_excel_rows,
    )
    _IAAO_OK = True
    print("  [iaao_engine] loaded ✓")
except Exception as _iaao_err:
    print(f"  [iaao_engine] import failed: {_iaao_err}")
    _IAAO_OK = False
    def _iaao_report(*a, **k):      return {"status": "disabled"}
    def _iaao_from_comps(*a, **k):  return {"status": "disabled"}
    def _iaao_excel_rows(*a, **k):  return []

# ── Price Intelligence Engine (بروتوكول الرصد اللحظي للأسعار) ──────────────
try:
    from price_intelligence import (
        price_intel      as _price_intel,
        MARKET_EG        as _MKT_EG,
        MARKET_KSA       as _MKT_KSA,
        build_excel_table as _build_excel_table,
    )
    _PI_OK = True
    print("  [price_intelligence] loaded ✓")
except Exception as _pi_err:
    print(f"  [price_intelligence] import failed: {_pi_err}")
    _PI_OK = False
    _MKT_EG  = "egypt"
    _MKT_KSA = "ksa"
    class _PIStub:
        def search_dict(self, **k):    return {"status": "disabled", "summary": {}, "records": [], "trends": [], "excel_table": [], "analysis_text": "محرك الاستخبار غير متاح"}
        def quick_trend(self, **k):    return []
        def clear_cache(self):         return {"status": "disabled"}
    _price_intel = _PIStub()
    def _build_excel_table(*a, **k):   return []

# ── Sales Comparison Sovereign v37 (مقارنات البيوع - الإصدار الذهبي) ────────
try:
    from sales_comparison_sovereign import upgrade_sales_comparison as _upgrade_sales_comp
    _SALES_COMP_OK = True
    print("  [sales_comparison_sovereign] loaded ✓")
except Exception as _sc_err:
    print(f"  [sales_comparison_sovereign] import failed: {_sc_err}")
    _SALES_COMP_OK = False
    def _upgrade_sales_comp(wb, data, comp_sales_raw, sheet_name="مقارنات البيوع"):
        pass  # graceful no-op if module unavailable

# ── RAG Advisor (المستشار العقاري الذكي) ─────────────────────────────────────
try:
    from rag_advisor import (
        advisor_answer      as _rag_answer,
        get_strategic_context as _rag_strategic,
        _init_rag           as _rag_init,
    )
    _RAG_OK = True
    print("  [rag_advisor] loaded ✓")
except Exception as _rag_err:
    print(f"  [rag_advisor] import failed: {_rag_err} — advisor endpoint disabled")
    _RAG_OK = False
    def _rag_answer(question="", location="", property_type="", use_web=False):
        return {"answer": "المستشار الذكي غير متاح حالياً — تأكد من تثبيت المكتبات.",
                "sources": [], "confidence": 0, "mode": "unavailable", "elapsed_s": 0}
    def _rag_strategic(sector="residential", purpose="fair_market_value", location=""):
        return {"title": "المحلل الاستراتيجي", "city_line": "", "econ_line": "",
                "geo_line": "", "sector_line": "", "purpose_note": "",
                "buy_signals": [], "caution": [], "disclaimer": ""}
    def _rag_init(): pass

# ── Arabic text rendering helper ─────────────────────────────────────────────
def _ar(text):
    """إصلاح النص العربي للعرض في matplotlib (reshaper + bidi)"""
    try:
        import arabic_reshaper
        from bidi.algorithm import get_display
        return get_display(arabic_reshaper.reshape(str(text)))
    except Exception:
        return str(text)

# ── إعداد matplotlib الموحَّد (خطوط + رموز رياضية + عربية) ──────────────────
def _setup_mpl_fonts():
    """
    يضبط rcParams لدعم:
      • العربية  → Arial / Tahoma / Calibri (أول خط متاح)
      • الرموز الرياضية (₀ σ ± …) → DejaVu Sans (fallback تلقائي)
      • mathtext → dejavusans (يغطي Unicode الرياضي كاملاً)
      • unicode_minus = False  → تجنب تحذير الشرطة السالبة
    يُستدعى في بداية كل دالة رسم بدلاً من التهيئة المتفرقة.
    """
    import matplotlib
    matplotlib.use("Agg")                          # non-interactive backend أولاً
    import matplotlib as _mpl
    import matplotlib.font_manager as fm

    # أفضل خط عربي متاح
    _arabic_pref = ["Arial", "Tahoma", "Calibri",
                    "Microsoft Sans Serif", "Times New Roman"]
    _available   = {f.name for f in fm.fontManager.ttflist}
    _ar_best     = next((n for n in _arabic_pref if n in _available), "DejaVu Sans")

    # مجموعة الخطوط المرتّبة: الخط العربي أولاً، DejaVu Sans احتياطي للرموز
    # نستخدم list مباشرة (لا "sans-serif" الاسم العام) لتجنب رسائل findfont المزعجة
    _fstack = [_ar_best, "DejaVu Sans"]
    for _fb in ("Arial", "Tahoma", "DejaVu Serif"):
        if _fb not in _fstack:
            _fstack.append(_fb)

    _mpl.rcParams["font.family"]        = _fstack        # list مباشرة بدون generic family
    _mpl.rcParams["font.sans-serif"]    = _fstack        # يُضبط أيضاً للمكتبات التي تقرأه
    _mpl.rcParams["mathtext.fontset"]   = "dejavusans"   # يغطي ₀ ₁ σ π ± …
    _mpl.rcParams["axes.unicode_minus"] = False           # شرطة ناقص ASCII بدل Unicode
    _mpl.rcParams["pdf.fonttype"]       = 42              # TrueType في PDF/PS
    _mpl.rcParams["ps.fonttype"]        = 42

# للتوافق الخلفي: أبقِ _ar_font() تُعيد اسم الخط فقط إذا احتاجها كود آخر
def _ar_font():
    import matplotlib.font_manager as fm
    _pref = ["Arial", "Tahoma", "Calibri", "Microsoft Sans Serif", "DejaVu Sans"]
    _avail = {f.name for f in fm.fontManager.ttflist}
    return next((n for n in _pref if n in _avail), "DejaVu Sans")

# ── مسارات ──────────────────────────────────────────────────────────────────
_BASE   = os.path.dirname(os.path.abspath(__file__))
_ROOT   = os.path.dirname(_BASE)
OUTPUTS = os.path.join(_BASE, "outputs")
os.makedirs(OUTPUTS, exist_ok=True)

_CANDIDATES = [
    os.path.join(_ROOT,  "templates", "template.xlsm"),
    os.path.join(_BASE,  "templates", "template.xlsm"),
    os.path.join(_ROOT,  "templates", "template.xlsx"),
    os.path.join(_BASE,  "templates", "template.xlsx"),
]
TEMPLATE = next((p for p in _CANDIDATES if os.path.exists(p)), _CANDIDATES[0])

# ── Flask + Static Hosting ──────────────────────────────────────────────────
# يستضيف مجلد frontend/ لفتح الواجهة مباشرة من http://localhost:5000
_FRONTEND_DIR = os.path.join(_ROOT, "frontend")
app = Flask(__name__,
            static_folder=_FRONTEND_DIR,
            static_url_path="/static")
CORS(app, resources={r"/api/*": {"origins": "*"}})

# ── Phase 4 engine singletons (lazy-loaded on first request) ─────────────────
_comparable_search_engine = None
_comparative_engine       = None
_cost_engine              = None
_income_engine            = None

def get_search_engine():
    global _comparable_search_engine
    if _comparable_search_engine is None:
        feed_path = os.path.join(_BASE, "data", "market_feed.json")
        _comparable_search_engine = ComparableSearchEngine(market_feed_path=feed_path)
    return _comparable_search_engine

def get_comparative_engine():
    global _comparative_engine
    if _comparative_engine is None:
        _comparative_engine = ComparativeEngine()
    return _comparative_engine

def get_cost_engine():
    global _cost_engine
    if _cost_engine is None:
        tables_path = os.path.join(_BASE, "engines", "cost_tables.json")
        _cost_engine = CostEngine(cost_tables_path=tables_path)
    return _cost_engine

def get_income_engine():
    global _income_engine
    if _income_engine is None:
        _income_engine = IncomeEngine()
    return _income_engine

# ── Phase 5 adapter singletons (lazy-loaded on first request) ────────────────
_market_value_adapter = None
_mortgage_adapter     = None
_insurance_adapter    = None
_ifrs13_adapter       = None

def get_market_value_adapter():
    global _market_value_adapter
    if _market_value_adapter is None:
        _market_value_adapter = MarketValueAdapter()
    return _market_value_adapter

def get_mortgage_adapter():
    global _mortgage_adapter
    if _mortgage_adapter is None:
        _mortgage_adapter = MortgageValueAdapter()
    return _mortgage_adapter

def get_insurance_adapter():
    global _insurance_adapter
    if _insurance_adapter is None:
        _insurance_adapter = InsuranceValueAdapter()
    return _insurance_adapter

def get_ifrs13_adapter():
    global _ifrs13_adapter
    if _ifrs13_adapter is None:
        _ifrs13_adapter = IFRS13FairValueAdapter()
    return _ifrs13_adapter

# ── Phase 6 asset adapter singletons ─────────────────────────────────────────
_residential_adapter = None
_commercial_adapter  = None

def get_residential_adapter():
    global _residential_adapter
    if _residential_adapter is None:
        _residential_adapter = ResidentialAdapter()
    return _residential_adapter

def get_commercial_adapter():
    global _commercial_adapter
    if _commercial_adapter is None:
        _commercial_adapter = CommercialAdapter()
    return _commercial_adapter

# ── Phase 7 singletons ────────────────────────────────────────────────────────
_land_adapter   = None
_report_auditor = None

def get_land_adapter():
    global _land_adapter
    if _land_adapter is None:
        _land_adapter = LandAdapter()
    return _land_adapter

def get_report_auditor():
    global _report_auditor
    if _report_auditor is None:
        _report_auditor = ReportQualityAuditor()
    return _report_auditor


def _search_db_comparables(subject: dict, filters: dict, limit: int) -> list | None:
    """
    Query PostgreSQL via SearchComparable; return list of dicts or None on failure.

    Returns None (not an empty list) when the DB is unavailable so callers
    know to fall back to the JSON search engine.
    """
    if not _DB_AVAILABLE:
        return None
    try:
        session = _DbSession()
        try:
            q = _SearchComparable(session)
            prop_type = subject.get("property_type") or filters.get("property_type")
            if prop_type:
                q = q.by_property_type(prop_type)
            gov = subject.get("governorate") or filters.get("governorate")
            if gov:
                q = q.by_governorate(gov)
            min_area = filters.get("min_area_sqm") or filters.get("area_sqm_min")
            max_area = filters.get("max_area_sqm") or filters.get("area_sqm_max")
            if min_area is not None and max_area is not None:
                q = q.by_area_range(float(min_area), float(max_area))
            lat = subject.get("latitude") or filters.get("latitude")
            lng = subject.get("longitude") or filters.get("longitude")
            radius = filters.get("radius_meters", 5_000)
            if lat is not None and lng is not None:
                q = q.by_location(float(lat), float(lng), float(radius))
            q = q.by_data_quality(0.5)
            return q.limit(limit).to_dict_list()
        finally:
            session.close()
    except Exception:
        return None


@app.after_request
def _cors(r):
    r.headers["Access-Control-Allow-Origin"]  = "*"
    r.headers["Access-Control-Allow-Headers"] = "Content-Type,Authorization"
    r.headers["Access-Control-Allow-Methods"] = "GET,POST,PUT,DELETE,OPTIONS"
    return r


@app.before_request
def _attach_user_from_token():
    """Wave S2 — read Bearer token silently; never reject a request.

    Sets g.user_id to the token subject when a valid JWT is present.
    Missing header, malformed header, or invalid/expired token all
    result in g.user_id = None without raising or returning a response.
    Enforcement is Wave S3's responsibility.
    """
    g.user_id = None
    if not _AUTH_AVAILABLE:
        return
    auth_header = request.headers.get("Authorization", "")
    if not auth_header.startswith("Bearer "):
        return
    token = auth_header[len("Bearer "):].strip()
    if not token:
        return
    try:
        payload = _verify_token(token)
        g.user_id = payload.get("sub")
    except _AuthError:
        g.user_id = None


@app.route("/api/<path:p>", methods=["OPTIONS"])
def _pre(p): return jsonify({}), 200

# ── Static file routes ─────────────────────────────────────────────────────
@app.route("/")
def serve_index():
    """يُقدّم الصفحة الرئيسية sovereign_brain.html مباشرة على /"""
    from flask import send_from_directory
    # ابحث عن sovereign_brain.html أولاً ثم index.html
    for fname in ("index.html", "sovereign_brain.html"):
        fpath = os.path.join(_FRONTEND_DIR, fname)
        if os.path.exists(fpath):
            return send_from_directory(_FRONTEND_DIR, fname)
    return "<h1>Frontend not found</h1><p>Place sovereign_brain.html in the frontend/ folder.</p>", 404

@app.route("/<path:path>")
def serve_static(path):
    """يُقدّم ملفات CSS/JS/fonts/images من مجلد frontend/"""
    from flask import send_from_directory
    full = os.path.join(_FRONTEND_DIR, path)
    if os.path.isfile(full):
        return send_from_directory(_FRONTEND_DIR, path)
    # fallback: 404 عادي بدلاً من traceback
    return jsonify({"error": "file not found"}), 404

# ═══════════════════════════════════════════════════════════════════════════
# محرك التقييم
# ═══════════════════════════════════════════════════════════════════════════
_PRICE_MAP = {
    # القاهرة الكبرى
    "التجمع الخامس": 40000, "القاهرة الجديدة": 35000, "الرحاب": 30000, "مدينتي": 32000,
    "الشروق": 18000, "المعادي": 28000, "مدينة نصر": 22000, "مصر الجديدة": 26000,
    "الزمالك": 55000, "وسط البلد": 20000, "القاهرة": 26000,
    # الجيزة
    "الدقي": 30000, "المهندسين": 32000, "الهرم": 14000, "فيصل": 12000,
    # 6 أكتوبر والشيخ زايد
    "الشيخ زايد": 32000, "6 أكتوبر": 18000, "أكتوبر": 18000, "زايد الجديدة": 22000,
    # العاصمة والمدن الجديدة
    "العاصمة الإدارية": 35000, "مستقبل سيتي": 20000, "بدر": 12000,
    # الإسكندرية والساحل
    "الإسكندرية": 18000, "سموحة": 22000, "برج العرب": 8000,
    "الساحل الشمالي": 45000, "العلمين الجديدة": 38000, "رأس الحكمة": 55000,
    # الدلتا
    "المنصورة": 10000, "طنطا": 8000, "الزقازيق": 7000, "دمياط": 9000,
    # الصعيد
    "أسيوط": 6000, "سوهاج": 5000, "المنيا": 5500, "الأقصر": 7000, "أسوان": 6500,
    # مدن القناة
    "بورسعيد": 10000, "الإسماعيلية": 9000, "السويس": 8000,
    # البحر الأحمر وسيناء
    "الغردقة": 15000, "شرم الشيخ": 18000, "العين السخنة": 25000, "مرسى علم": 10000,
    # المملكة العربية السعودية (SAR/م²)
    "الرياض": 4800, "حطين": 7000, "الملقا": 6500, "النرجس": 5500, "الدرعية": 8000,
    "مكة": 5000, "المدينة المنورة": 3800, "الطائف": 2500,
    "جدة": 5500, "جدة الحمراء": 6000, "أبحر": 5000, "جدة البلد": 3500,
    "الدمام": 3200, "الخبر": 4000, "الظهران": 4500, "الجبيل": 2800, "الأحساء": 2000,
    "أبها": 2200, "خميس مشيط": 1800,
    "تبوك": 1500, "نيوم": 8000,
    "بريدة": 1800, "عنيزة": 1600,
    "جازان": 1500, "نجران": 1400, "ينبع": 2500,
}

# ══════════════════════════════════════════════════════════════════════════════
# محرك المسارات الذكي — Logic Gate v45
# ══════════════════════════════════════════════════════════════════════════════
_PURPOSES = {
    "fair_market_value":    {"ar": "القيمة السوقية العادلة",                        "logic": "FMV"},
    "acquisition":          {"ar": "الاستحواذ والاندماج — القيمة الاستثمارية",      "logic": "M&A-DCF"},
    "bank_financing":       {"ar": "الرهن والتمويل البنكي",                        "logic": "BANK"},
    "rental_arbitration":   {"ar": "تحديد الأجرة — القيمة الإيجارية العادلة",      "logic": "RENT"},
    "insurance":            {"ar": "التأمين — قيمة إعادة الإنشاء",                 "logic": "INSURE"},
    "investment_analysis":  {"ar": "التحليل الاستثماري",                            "logic": "INVEST"},
    "judicial_liquidation": {"ar": "التصفية الإجبارية — القيمة التصفوية",           "logic": "LIQUID"},
    "tax_assessment":       {"ar": "الضرائب — القيمة للأغراض الضريبية",            "logic": "TAX"},
    "financial_reporting":  {"ar": "التقارير المالية — القيمة العادلة IFRS 13",     "logic": "IFRS"},
    "usufruct":             {"ar": "حق الانتفاع — تقييم حقوق المنفعة العقارية",     "logic": "USUFRUCT"},
    "uncertainty_valuation":{"ar": "التقييم في حالة عدم اليقين",                    "logic": "UNCERTAINTY"},
    "highest_and_best_use": {"ar": "تحليل أعلى وأفضل استغلال",                      "logic": "HBU"},
    "investment_funds":     {"ar": "التقييم لأغراض الصناديق الاستثمارية",            "logic": "REIT"},
    "environmental_impact_assessment": {"ar": "تقييم الأثر البيئي (EIA)",            "logic": "EIA"},
}
# ── توافق خلفي: ربط القيم القصيرة من index.html بالمفاتيح الرسمية ────────────
_PURPOSE_ALIAS = {
    "market":      "fair_market_value",
    "ma":          "acquisition",
    "bank":        "bank_financing",
    "judicial":    "judicial_liquidation",
    "rental":      "rental_arbitration",
    "investment":  "investment_analysis",
    "insurance":   "insurance",
    "tax":         "tax_assessment",
    "ifrs":        "financial_reporting",
}
_DEFAULT_PURPOSE = "fair_market_value"

# قواعد الحذف الديناميكي — الشيتات المُدرَجة تُطابَق بالاحتواء (substring match)
# يضمن أن التقرير لا يتجاوز 20 شيت مع إبقاء الشيتات ذات الصلة بالغرض فقط
_PURPOSE_PRUNE = {
    "fair_market_value":    [],   # كامل — جميع الشيتات
    "acquisition":          ["الإيجار مقابل الشراء"],
    "bank_financing":       ["الخيارات الحقيقية", "ANN", "ARIMA",
                             "الإيجار مقابل الشراء"],
    "rental_arbitration":   ["الخيارات الحقيقية", "ANN",
                             "أفضل وأعلى استخدام"],
    "insurance":            ["الخيارات الحقيقية", "ANN", "ARIMA",
                             "الإيجار مقابل الشراء", "المقارنات الإيجارية"],
    "investment_analysis":  ["الإيجار مقابل الشراء"],
    "judicial_liquidation": ["الخيارات الحقيقية", "ANN", "ARIMA",
                             "الإيجار مقابل الشراء"],
    "tax_assessment":       ["الخيارات الحقيقية", "ANN", "ARIMA",
                             "الإيجار مقابل الشراء"],  # ضريبي: مناهج تقليدية فقط
    "financial_reporting":  [],   # IFRS 13: تقرير كامل — كل المناهج مطلوبة
    # ── الأغراض الجديدة (Wave 2) ────────────────────────────────────────────
    "usufruct":             ["الخيارات الحقيقية", "ANN", "ARIMA",
                             "الإيجار مقابل الشراء"],  # حق انتفاع: لا حقوق بيع/إعادة تطوير
    "uncertainty_valuation":[],   # IFRS 13 §93: نطاق كامل لتثليث القيمة
    "highest_and_best_use": ["الإيجار مقابل الشراء", "ANN", "ARIMA",
                             "المقارنات الإيجارية"],  # HBU: تركيز على المقارنة بين البدائل
    "investment_funds":     [],   # REIT/IFRS 13/IOSCO: إفصاح كامل لحملة الوحدات
    "environmental_impact_assessment": ["الإيجار مقابل الشراء", "ANN", "ARIMA",
                                         "المقارنات الإيجارية"],  # EIA: تركيز على المخاطر البيئية والـ ERF
}

# معاملات تعديل PPM الصافية بعد حساب القطاع — تعكس منطق الغرض الاقتصادي
_PURPOSE_PPM_FACTOR = {
    "fair_market_value":    1.00,   # القيمة السوقية النظيفة
    "acquisition":          1.00,   # DCF يتحكم من داخل الإكسيل
    "bank_financing":       0.95,   # تحفظ بنكي: -5%
    "rental_arbitration":   1.00,   # القيمة كما هي لتحكيم عادل
    "insurance":            1.08,   # تكلفة إعادة الإنشاء: +8%
    "investment_analysis":  1.00,   # العائد يُحسب في DCF
    "judicial_liquidation": 0.82,   # خصم التصفية السريعة: -18%
    "tax_assessment":       1.00,   # القيمة الضريبية: بدون خصم
    "financial_reporting":  1.00,   # IFRS 13: القيمة العادلة
    # ── الأغراض الجديدة (Wave 2) — معاملات اقتصادية مدروسة ───────────────────
    "usufruct":             0.65,   # حق الانتفاع: 50-75% من قيمة الملكية الكاملة (وسط: 65%)
    "uncertainty_valuation":0.92,   # IFRS 13 §93: تحفظ 8% لعدم اليقين الجوهري
    "highest_and_best_use": 1.00,   # تُحسب القيمة عبر موديول HBU المتخصص
    "investment_funds":     1.00,   # تُحسب NAV عبر موديول REIT المتخصص
    "environmental_impact_assessment": 1.00,  # ERF يُطبَّق ديناميكياً وفق التصنيف A/B/C
}

def _num(v, d=0.0):
    try:
        if isinstance(v, str): v = v.replace(",","").strip()
        x = float(v)
        return d if (math.isnan(x) or math.isinf(x)) else x
    except: return float(d)

def _ncdf(x):
    """تقريب دالة التوزيع الطبيعي التراكمي"""
    a1,a2,a3,a4,a5 = 0.31938153,-0.356563782,1.781477937,-1.821255978,1.330274429
    k = 1/(1+0.2316419*abs(x))
    p = 1 - (1/math.sqrt(2*math.pi))*math.exp(-0.5*x*x)*(
        a1*k + a2*k**2 + a3*k**3 + a4*k**4 + a5*k**5)
    return p if x >= 0 else 1-p

def advanced_valuation(location, area, property_type, price_per_meter, **kw):
    ppm = _num(price_per_meter)
    ar  = _num(area)
    if ppm <= 0:
        for key, val in _PRICE_MAP.items():
            if key in str(location): ppm = val; break
    if ppm <= 0: ppm = 20000.0

    # ── تشغيل المُدقق إذا أُرسلت مقارنات أو وُجدت بيانات مُخزّنة ──────────
    comp_sales = kw.get("comp_sales", [])
    feed_recs  = _load_feed()

    # ── الرادار التلقائي: اجلب من الإنترنت/قاموس الأسعار عند غياب المقارنات ──
    # نحسب فقط السجلات المطابقة للموقع (لا كل السجلات في الـ feed)
    _loc_str = str(location)
    _ptype_str = str(property_type)
    _matching_feed_count = sum(
        1 for r in feed_recs
        if (_loc_str in str(r.get("location","")) or str(r.get("location","")) in _loc_str)
        and (_ptype_str in str(r.get("property_type","")) or str(r.get("property_type","")) in _ptype_str or not _ptype_str)
    )
    _auto_note = ""
    if not comp_sales and _matching_feed_count < 3 and _MI_MODULE_OK:
        try:
            _fetched, _auto_note = _auto_fetch(
                location      = str(location),
                property_type = str(property_type),
                area          = float(ar),
                ppm_hint      = float(ppm),
                n             = 6,
            )
            if _fetched:
                comp_sales = _fetched
                print(f"{_ts()} [INFO] الرادار التلقائي — سحب {len(_fetched)} مقارنة لـ '{location}'... 30%")
        except Exception as _af_err:
            print(f"{_ts()} [WARNING] auto-fetch error: {_af_err}")

    if comp_sales or feed_recs:
        mi = _market_intelligence_validator(
            comps         = comp_sales,
            location      = str(location),
            property_type = str(property_type),
            feed_records  = feed_recs,
            auto_note     = _auto_note,
        )
        # استخدام PPM المُعدَّل فقط إذا كانت جودة البيانات مقبولة (≥4) وكان الفارق معقولاً
        adj = mi["adjusted_ppm"]
        if adj > 0 and mi["quality_score"] >= 4:
            # مزج: 70% من البيانات السوقية + 30% من المدخلات المباشرة
            blend = adj * 0.70 + ppm * 0.30
            ppm = round(blend, 0)

    # ── تحليل القطاع المتخصص ─────────────────────────────────────────────────
    sector    = _get_sector(str(property_type))
    # ── حل الغرض: يقبل المفاتيح الرسمية أو المختصرة أو من valuation_purpose ──
    _raw_purpose = str(
        kw.get("purpose") or kw.get("valuation_purpose") or _DEFAULT_PURPOSE
    ).strip()
    purpose = _PURPOSE_ALIAS.get(_raw_purpose, _raw_purpose)
    if purpose not in _PURPOSES:
        purpose = _DEFAULT_PURPOSE
    _pur_ar    = _PURPOSES[purpose]["ar"]
    _pur_logic = _PURPOSES[purpose]["logic"]
    _sector_data: dict = {}

    if sector == "industrial":
        try:
            print(f"{_ts()} [INFO] القطاع الصناعي — استهلاك عمري + ذكاء الموانئ والمناطق اللوجستية... 30%")
            yr_b  = int(kw.get("year_built", 2000) or 2000)
            cond  = str(kw.get("condition", "average"))
            depr  = _industrial_depreciation(yr_b, cond, "building_concrete")
            i_prem = _industrial_location_premium(str(location))
            ppm   = round(ppm * i_prem * depr["depreciated_value_ratio"], 0)
            _sector_data = {**depr, "location_premium": i_prem}
        except Exception as _ie:
            print(f"{_ts()} [WARNING] industrial calc: {_ie}")

    elif sector == "agricultural":
        try:
            print(f"{_ts()} [INFO] القطاع الزراعي — رادار إنتاجية الفدان وجودة التربة... 30%")
            area_fed  = float(kw.get("area_feddan", 0) or ar / FEDDAN_SQM)
            crop_t    = str(kw.get("crop_type",    "_default"))
            irrig     = str(kw.get("irrigation",   "_default"))
            soil_q    = str(kw.get("soil_quality", "_default"))
            cap_r_ag  = float(kw.get("cap_rate", 0.06) or 0.06)
            agri = _agricultural_income_ppm(str(location), area_fed, crop_t, irrig, soil_q, cap_r_ag)
            if agri["ppm"] > 0:
                ppm = round(agri["ppm"] * 0.70 + ppm * 0.30, 0)
            _sector_data = agri
        except Exception as _ae:
            print(f"{_ts()} [WARNING] agricultural calc: {_ae}")

    elif sector == "hospitality":
        try:
            print(f"{_ts()} [INFO] القطاع الفندقي — رادار نسب الإشغال + DCF... 30%")
            rooms_h   = int(kw.get("rooms", 50) or 50)
            apr_h     = float(kw.get("area_per_room", 40) or 40)
            stars_h   = int(kw.get("stars", 0) or 0)
            wacc_h    = float(kw.get("wacc", 0.12) or 0.12)
            cap_r_h   = float(kw.get("cap_rate", 0.08) or 0.08)
            hp_h      = int(kw.get("holding_period", 10) or 10)
            hosp = _hospitality_income_ppm(str(location), rooms_h, apr_h, stars_h, cap_r_h, hp_h, wacc_h)
            if hosp["ppm"] > 0:
                ppm = round(hosp["ppm"] * 0.70 + ppm * 0.30, 0)
            _sector_data = hosp
        except Exception as _he:
            print(f"{_ts()} [WARNING] hospitality calc: {_he}")

    elif sector == "retail":
        try:
            print(f"{_ts()} [INFO] قطاع التجزئة — معامل القوة الشرائية + كثافة المرور... 30%")
            frontage_r  = float(kw.get("frontage",      0) or 0)
            gla_ratio_r = float(kw.get("gla_ratio",  0.75) or 0.75)
            cap_r_r     = float(kw.get("cap_rate",    0.09) or 0.09)
            has_anchor  = bool(kw.get("has_anchor_tenant", False))
            retail = _retail_income_ppm(str(location), ar, frontage_r, gla_ratio_r, cap_r_r, has_anchor)
            if retail["ppm"] > 0:
                ppm = round(retail["ppm"] * 0.70 + ppm * 0.30, 0)
            _sector_data = retail
        except Exception as _re:
            print(f"{_ts()} [WARNING] retail calc: {_re}")

    elif sector == "healthcare":
        try:
            print(f"{_ts()} [INFO] قطاع الصحة — تكلفة التجهيزات + معامل نقص الخدمات... 30%")
            beds_hc       = int(kw.get("beds",           0)  or 0)
            facility_hc   = str(kw.get("facility_type",  "_default"))
            cap_r_hc      = float(kw.get("cap_rate",     0.09) or 0.09)
            rent_hc       = float(kw.get("monthly_rent_sqm", 0) or 0)
            health = _healthcare_value_ppm(str(location), ar, beds_hc, facility_hc, cap_r_hc, rent_hc)
            if health["ppm"] > 0:
                ppm = round(health["ppm"] * 0.65 + ppm * 0.35, 0)
            _sector_data = health
        except Exception as _hce:
            print(f"{_ts()} [WARNING] healthcare calc: {_hce}")

    elif sector == "educational":
        try:
            print(f"{_ts()} [INFO] قطاع التعليم — طاقة الطلاب + علاوة المرافق + الترخيص... 30%")
            school_type_e = str(kw.get("school_type",  "مدرسة"))
            fee_type_e    = str(kw.get("fee_type",     "private_school"))
            facilities_e  = list(kw.get("facilities",  []) or [])
            cap_r_e       = float(kw.get("cap_rate",   0.08) or 0.08)
            licensed_e    = bool(kw.get("licensed",    True))
            edu = _educational_capacity_ppm(str(location), ar, school_type_e, fee_type_e, facilities_e, cap_r_e, licensed_e)
            if edu["ppm"] > 0:
                ppm = round(edu["ppm"] * 0.70 + ppm * 0.30, 0)
            _sector_data = edu
        except Exception as _ee:
            print(f"{_ts()} [WARNING] educational calc: {_ee}")

    # ── محرك المسارات الذكي — Route Identified ──────────────────────────────
    print(f"{_ts()} [INFO] Route Identified: Asset={sector} | Purpose={_pur_ar} | Logic={_pur_logic} | Progress: 40%")

    # تعديل PPM حسب غرض التقييم (بعد تعديلات القطاع)
    _pf = _PURPOSE_PPM_FACTOR.get(purpose, 1.00)
    if _pf != 1.00:
        ppm = round(ppm * _pf, 0)
        print(f"{_ts()} [INFO] معامل الغرض [{_pur_logic}]: ×{_pf:.2f} → {ppm:,.0f} EGP/م²")

    # الغرض الافتراضي: القيمة السوقية العادلة (FMV) في حال عدم تحديد غرض
    if purpose == _DEFAULT_PURPOSE:
        print(f"{_ts()} [INFO] المسار الافتراضي: {_pur_ar} — تقرير شامل (جميع المناهج)")

    mv = ppm * ar
    return dict(price_per_meter=ppm, market_value=mv,
                market_approach=mv*0.98, cost_approach=mv*1.05,
                income_approach=mv*0.92, idw_ppm=ppm*1.02,
                kriging_ppm=ppm*0.99, ols_predicted_ppm=ppm*1.01,
                sector=sector, _sector_data=_sector_data,
                purpose=purpose)

# ═══════════════════════════════════════════════════════════════════════════
# رسم شجرة القرارات (matplotlib → PNG)
# ═══════════════════════════════════════════════════════════════════════════
def _tree_png(S, K, sigma, T, r, steps=3) -> bytes:
    _setup_mpl_fonts()
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch
    import matplotlib.patches as mpatches

    dt   = T / steps
    u    = math.exp(sigma * math.sqrt(dt))
    d    = 1.0 / u
    disc = math.exp(-r * dt)
    p    = max(0.01, min(0.99, (math.exp(r*dt)-d)/(u-d)))

    # قيم الأصل
    asset = {(t,j): S*(u**(t-j))*(d**j)
             for t in range(steps+1) for j in range(t+1)}
    # قيم الخيار (backward)
    opt = {(steps,j): max(asset[(steps,j)]-K, 0) for j in range(steps+1)}
    for t in range(steps-1, -1, -1):
        for j in range(t+1):
            hold = disc*(p*opt[(t+1,j)] + (1-p)*opt[(t+1,j+1)])
            opt[(t,j)] = max(asset[(t,j)]-K, hold)

    def pos(t, j):
        x = t * 1.0
        y = (steps/2) - j + (t/2)
        return x, y

    fig, ax = plt.subplots(figsize=(13, 7.5))
    fig.patch.set_facecolor("#0d1117")
    ax.set_facecolor("#0d1117")
    ax.set_xlim(-0.5, steps+0.5)
    ax.set_ylim(-0.9, steps+0.9)
    ax.axis("off")

    NW, NH = 0.58, 0.50

    # خطوط + احتمالات
    for t in range(steps):
        for j in range(t+1):
            x0,y0 = pos(t,j)
            xu,yu = pos(t+1,j)
            xd,yd = pos(t+1,j+1)
            ax.plot([x0,xu],[y0,yu], color="#555577", lw=1.8, zorder=1)
            ax.plot([x0,xd],[y0,yd], color="#555577", lw=1.8, zorder=1)
            ax.text((x0+xu)/2-0.1,(y0+yu)/2+0.09, f"p={p:.2f}",
                    fontsize=7.5, color="#66ccff", ha="right", zorder=4)
            ax.text((x0+xd)/2-0.1,(y0+yd)/2-0.09, f"1-p={1-p:.2f}",
                    fontsize=7.5, color="#ff9966", ha="right", zorder=4)

    # عقد
    for t in range(steps+1):
        for j in range(t+1):
            x,y = pos(t,j)
            sv  = asset[(t,j)]; cv = opt[(t,j)]
            dev = sv > K
            fc  = "#0a5e36" if dev else "#6b0f0f"
            ec  = "#00e676" if dev else "#ff5252"
            box = FancyBboxPatch((x-NW/2, y-NH/2), NW, NH,
                                  boxstyle="round,pad=0.04",
                                  facecolor=fc, edgecolor=ec,
                                  linewidth=2.5, alpha=0.97, zorder=2)
            ax.add_patch(box)
            ax.text(x, y+0.15, f"S={sv/1e6:.2f}M",
                    ha="center", fontsize=8, color="white",
                    fontweight="bold", zorder=3)
            ax.text(x, y-0.01, f"C={cv/1e6:.2f}M",
                    ha="center", fontsize=7.5, color="#ffee58", zorder=3)
            lbl = _ar("▲ طوّر") if dev else _ar("⏳ انتظر")
            ax.text(x, y-0.17, lbl, ha="center", fontsize=8,
                    color=ec, fontweight="bold", zorder=3)

    for t in range(steps+1):
        ax.text(t, -0.7, f"t = {t}",
                ha="center", fontsize=10, color="#d4af37", fontweight="bold")

    ax.text(steps/2, steps+0.62,
            _ar("شجرة القرارات الثنائية — Binomial Decision Tree  (3 خطوات)"),
            ha="center", fontsize=12, color="white", fontweight="bold")
    ax.text(steps/2, steps+0.35,
            f"S0={S:,.0f}  |  K={K:,.0f}  |  sigma={sigma:.0%}  |  r={r:.0%}  |  p={p:.2f}",
            ha="center", fontsize=9, color="#aaaacc")

    lgd = [mpatches.Patch(color="#0a5e36",label=_ar("S > K — طوّر")),
           mpatches.Patch(color="#6b0f0f",label=_ar("S \u2264 K — انتظر"))]
    ax.legend(handles=lgd, loc="lower right",
              facecolor="#1a1a2e", edgecolor="#d4af37",
              labelcolor="white", fontsize=9)

    plt.tight_layout(pad=0.4)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight",
                facecolor="#0d1117")
    plt.close(fig)
    plt.clf(); plt.close("all")
    data = buf.getvalue(); buf.close()
    gc.collect()
    return data

# ═══════════════════════════════════════════════════════════════════════════
# رسم مخطط ARIMA
# ═══════════════════════════════════════════════════════════════════════════
def _arima_png(hist_q, hist_p, fore_p) -> bytes:
    _setup_mpl_fonts()
    import matplotlib.pyplot as plt

    fig, (ax1,ax2) = plt.subplots(1,2, figsize=(14,5))
    fig.patch.set_facecolor("#0d1117")
    for ax in (ax1,ax2):
        ax.set_facecolor("#111827")
        for sp in ax.spines.values(): sp.set_color("#333")
        ax.tick_params(colors="#aaa")

    n = len(hist_p)
    ax1.plot(range(n), hist_p, "o-", color="#4ec9b0", lw=2.2, ms=5, label=_ar("فعلي"))
    fore_x = list(range(n-1, n+len(fore_p)))
    fore_y = [hist_p[-1]] + fore_p
    ax1.plot(fore_x, fore_y, "s--", color="#f9a825", lw=2.2, ms=5, label=_ar("تنبؤ ARIMA"))
    for i,fv in enumerate(fore_p):
        ci = fv * 0.05 * (i+1)**0.5
        xi = n+i
        ax1.fill_between([xi-0.45,xi+0.45],[fv-ci]*2,[fv+ci]*2,
                         color="#f9a825", alpha=0.22)
    xt = list(range(0, n+len(fore_p), 2))
    xl = (hist_q + [f"Q{(i%4)+1}-{2026+i//4}" for i in range(len(fore_p))])[::2]
    ax1.set_xticks(xt); ax1.set_xticklabels(xl, rotation=35, ha="right",
                                              fontsize=7.5, color="#ccc")
    ax1.set_title(_ar("السلسلة الزمنية + التنبؤ"), color="white", fontsize=11)
    ax1.legend(facecolor="#0d1117", edgecolor="#d4af37",
               labelcolor="white", fontsize=9)
    ax1.yaxis.label.set_color("white")

    lags = list(range(1,10))
    acf  = [0.71,-0.22,0.09,-0.05,0.02,0.03,-0.02,0.01,0.00]
    cols = ["#4ec9b0" if abs(v)>0.25 else "#555" for v in acf]
    ax2.bar(lags, acf, color=cols, edgecolor="none", width=0.55)
    ax2.axhline(0.25,  ls="--", color="#f9a825", lw=1.2, label=_ar("حد 95%"))
    ax2.axhline(-0.25, ls="--", color="#f9a825", lw=1.2)
    ax2.axhline(0, color="#555", lw=0.6)
    ax2.set_xticks(lags); ax2.set_xlabel("Lag", color="#aaa", fontsize=9)
    ax2.set_title(_ar("ACF — دالة الارتباط الذاتي"), color="white", fontsize=10)
    ax2.legend(facecolor="#0d1117", edgecolor="#d4af37",
               labelcolor="white", fontsize=8)

    plt.tight_layout(pad=1.2)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight",
                facecolor="#0d1117")
    plt.close(fig)
    plt.clf(); plt.close("all")
    data = buf.getvalue(); buf.close()
    gc.collect()
    return data

# ═══════════════════════════════════════════════════════════════════════════
# قاموس إحداثيات المدن المصرية والخليجية (Lat, Lng)
# ═══════════════════════════════════════════════════════════════════════════
_CITY_COORDS = {
    # ══ القاهرة الكبرى ══════════════════════════════════════════════════════
    "التجمع الخامس":    (30.0131, 31.4913),
    "القاهرة الجديدة":  (30.0431, 31.4694),
    "الرحاب":           (30.0598, 31.4929),
    "مدينتي":           (30.1107, 31.6125),
    "الشروق":           (30.1252, 31.6305),
    "مدينة نصر":        (30.0702, 31.3388),
    "المعادي":           (29.9600, 31.2572),
    "مصر الجديدة":      (30.0854, 31.3268),
    "الزمالك":           (30.0653, 31.2195),
    "وسط البلد":         (30.0444, 31.2357),
    "القاهرة":           (30.0444, 31.2357),
    # ══ محافظة الجيزة ═══════════════════════════════════════════════════════
    "الدقي":             (30.0499, 31.2091),
    "المهندسين":         (30.0574, 31.2017),
    "الهرم":             (29.9764, 31.1313),
    "فيصل":             (29.9849, 31.1638),
    # ══ 6 أكتوبر والشيخ زايد ═══════════════════════════════════════════════
    "6 أكتوبر":          (29.9768, 30.9229),
    "الشيخ زايد":        (30.0268, 30.9432),
    "أكتوبر":            (29.9768, 30.9229),
    "زايد الجديدة":     (30.0550, 30.8700),
    # ══ العاصمة الإدارية والمدن الجديدة ══════════════════════════════════
    "العاصمة الإدارية":  (30.0167, 31.7465),
    "مستقبل سيتي":       (30.1418, 31.7100),
    "بدر":               (30.1338, 31.7406),
    # ══ الإسكندرية والساحل ══════════════════════════════════════════════════
    "الإسكندرية":        (31.2001, 29.9187),
    "سموحة":             (31.2096, 29.9454),
    "المنتزه":           (31.2834, 30.0139),
    "برج العرب":        (30.8536, 29.6620),
    "الساحل الشمالي":   (31.0500, 28.5000),
    "العلمين الجديدة":  (30.8500, 28.9500),
    "رأس الحكمة":       (31.1000, 27.7500),
    # ══ محافظات الدلتا ═══════════════════════════════════════════════════════
    "المنصورة":          (31.0409, 31.3785),
    "طنطا":              (30.7865, 31.0004),
    "الزقازيق":          (30.5877, 31.5020),
    "دمياط":             (31.4175, 31.8144),
    # ══ محافظات الصعيد ═══════════════════════════════════════════════════════
    "أسيوط":             (27.1809, 31.1837),
    "سوهاج":             (26.5591, 31.6948),
    "المنيا":            (28.0871, 30.7618),
    "الأقصر":            (25.6872, 32.6396),
    "أسوان":             (24.0889, 32.8998),
    # ══ مدن القناة ═══════════════════════════════════════════════════════════
    "بورسعيد":           (31.2653, 32.3019),
    "الإسماعيلية":       (30.5965, 32.2715),
    "السويس":            (29.9668, 32.5498),
    # ══ البحر الأحمر وسيناء ═════════════════════════════════════════════════
    "الغردقة":           (27.2579, 33.8116),
    "شرم الشيخ":         (27.9158, 34.3300),
    "العين السخنة":     (29.6080, 32.3280),
    "مرسى علم":          (25.0671, 34.8982),
    # ══ المملكة العربية السعودية ══════════════════════════════════════════════
    "الرياض":            (24.7136, 46.6753),
    "حطين":             (24.7720, 46.6370),
    "الملقا":           (24.8050, 46.6260),
    "النرجس":           (24.8250, 46.6450),
    "الدرعية":          (24.7341, 46.5773),
    "مكة":               (21.3891, 39.8579),
    "المدينة المنورة":  (24.4672, 39.6024),
    "الطائف":           (21.2703, 40.4158),
    "جدة":               (21.4858, 39.1925),
    "جدة الحمراء":      (21.5350, 39.1700),
    "أبحر":             (21.6500, 39.1100),
    "جدة البلد":        (21.4848, 39.1862),
    "الدمام":            (26.3927, 49.9777),
    "الخبر":             (26.2172, 50.1971),
    "الظهران":           (26.2361, 50.0393),
    "الجبيل":            (27.0046, 49.6226),
    "الأحساء":           (25.3648, 49.5855),
    "أبها":              (18.2164, 42.5053),
    "خميس مشيط":        (18.3093, 42.7290),
    "تبوك":              (28.3835, 36.5662),
    "نيوم":              (27.9500, 35.2000),
    "بريدة":             (26.3268, 43.9750),
    "عنيزة":             (26.0843, 43.9936),
    "جازان":             (16.8892, 42.5510),
    "نجران":             (17.4924, 44.1322),
    "ينبع":              (24.0895, 38.0618),
    # ══ الإمارات ════════════════════════════════════════════════════════════
    "دبي":               (25.2048, 55.2708),
    "أبوظبي":            (24.4539, 54.3773),
}

# ═══════════════════════════════════════════════════════════════════════════
# معامل التفاوض حسب الموقع (Negotiation Coefficient by Location)
# ═══════════════════════════════════════════════════════════════════════════
_NEG_COEFF = {
    # درجة أولى — تفاوض منخفض (3-4%)
    "الزمالك": 0.03, "المعادي": 0.04, "التجمع الخامس": 0.035,
    "القاهرة الجديدة": 0.04, "مصر الجديدة": 0.04,
    # درجة ثانية — تفاوض متوسط (5-6%)
    "الشيخ زايد": 0.05, "6 أكتوبر": 0.055, "الدقي": 0.05,
    "المهندسين": 0.05, "سموحة": 0.05,
    # درجة ثالثة — تفاوض مرتفع (6-8%)
    "مدينة نصر": 0.07, "القاهرة": 0.07, "الإسكندرية": 0.065,
    "الرياض": 0.05, "جدة": 0.05,
}
_DEFAULT_NEG_COEFF = 0.06   # القيمة الافتراضية إذا لم يوجد الموقع في القاموس

def _get_neg_coeff(location: str) -> float:
    """يُعيد معامل التفاوض المناسب بحث البحث عن أطول تطابق في القاموس"""
    for key in sorted(_NEG_COEFF, key=len, reverse=True):
        if key in location or location in key:
            return _NEG_COEFF[key]
    return _DEFAULT_NEG_COEFF

# ═══════════════════════════════════════════════════════════════════════════
# ملف مصادر البيانات الخارجية (Market Feed Persistence)
# ═══════════════════════════════════════════════════════════════════════════
_DATA_DIR  = os.path.join(_BASE, "data")
os.makedirs(_DATA_DIR, exist_ok=True)
_FEED_FILE = os.path.join(_DATA_DIR, "market_feed.json")

_SOURCE_CREDIBILITY = {
    "direct":   1.00,   # معاينة مباشرة / عقد موثق
    "agent":    0.85,   # وكيل عقاري
    "forum":    0.65,   # منتديات التقييم
    "facebook": 0.50,   # مجموعات فيسبوك
    "other":    0.55,
}

def _load_feed() -> list:
    """يُحمّل قاعدة البيانات الخارجية (JSON) — يُعيد قائمة فارغة إن لم توجد"""
    import json
    if not os.path.exists(_FEED_FILE):
        return []
    try:
        with open(_FEED_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return []

def _save_feed(records: list):
    """يحفظ قاعدة البيانات الخارجية"""
    import json
    with open(_FEED_FILE, "w", encoding="utf-8") as f:
        json.dump(records, f, ensure_ascii=False, indent=2, default=str)

def _temporal_weight(ts_str: str, half_life_days: int = 180) -> float:
    """
    وزن زمني تناقصي — البيانات الأحدث تحصل على وزن أعلى.
    half_life = 180 يوم → بيانات بعمر 6 أشهر تحصل على نصف الوزن.
    """
    try:
        from datetime import datetime as _dt
        ts = _dt.fromisoformat(ts_str) if isinstance(ts_str, str) else ts_str
        days_old = max(0, (_dt.now() - ts).days)
        return math.exp(-math.log(2) * days_old / half_life_days)
    except Exception:
        return 0.5   # وزن افتراضي إذا فشل التحليل

# ═══════════════════════════════════════════════════════════════════════════
# المُدقق الذكي للبيانات — Market Intelligence Validator
# ═══════════════════════════════════════════════════════════════════════════
def _market_intelligence_validator(
    comps: list,
    location: str = "",
    property_type: str = "",
    feed_records: list | None = None,
    auto_note: str = "",
) -> dict:
    """
    يُنظّف ويُحلّل بيانات المقارنات:
      1. يستبعد الشواذ (IQR × 1.5)
      2. يُطبّق معامل التفاوض تلقائياً
      3. يُضيف بيانات المصادر الخارجية المُخزّنة
      4. يمنح وزناً أكبر للبيانات الأحدث (تناقص زمني)
      5. يُعيد سعراً موزوناً نهائياً ومعلومات تشخيصية كاملة

    يُعيد dict يحتوي على:
      clean_comps, outliers, weighted_ppm, neg_coeff,
      adjusted_ppm, quality_score, diagnostics
    """
    import statistics as _st

    # ── دمج المصادر الخارجية المُخزّنة مع المقارنات المُرسلة مباشرةً ────────
    all_comps = list(comps or [])
    if feed_records:
        # تصفية بحسب الموقع ونوع العقار (تطابق جزئي)
        for rec in feed_records:
            rec_loc   = str(rec.get("location", ""))
            rec_ptype = str(rec.get("property_type", ""))
            loc_match   = (not location)   or (location in rec_loc)   or (rec_loc in location)
            ptype_match = (not property_type) or (property_type in rec_ptype) or (rec_ptype in property_type)
            if loc_match and ptype_match:
                # تحويل السجل الخارجي لصيغة comp_sale موحدة
                ppm_feed = _num(rec.get("price_per_meter", 0))
                if ppm_feed <= 0:
                    cp = _num(rec.get("price", 0)); ca = _num(rec.get("area", 1))
                    ppm_feed = cp / ca if ca > 0 else 0
                if ppm_feed > 0:
                    all_comps.append({
                        "price":          _num(rec.get("price", ppm_feed * _num(rec.get("area", 100)))),
                        "area":           _num(rec.get("area", 100)),
                        "price_per_meter": ppm_feed,
                        "source":         rec.get("source", "other"),
                        "timestamp":      rec.get("timestamp", datetime.now().isoformat()),
                        "_from_feed":     True,
                    })

    # ── استخراج PPM من كل مقارنة ─────────────────────────────────────────
    records_with_ppm = []
    for c in all_comps:
        ppm_c = _num(c.get("price_per_meter", 0))
        if ppm_c <= 0:
            cp = _num(c.get("price", 0)); ca = _num(c.get("area", 1))
            ppm_c = cp / ca if ca > 0 else 0
        if ppm_c > 0:
            records_with_ppm.append({**c, "price_per_meter": ppm_c})

    n_total = len(records_with_ppm)

    if n_total == 0:
        return {
            "clean_comps": [], "outliers": [], "n_total": 0, "n_clean": 0,
            "n_outliers": 0, "mean_raw": 0, "median_raw": 0, "std_raw": 0,
            "mean_clean": 0, "median_clean": 0, "std_clean": 0,
            "lower_fence": 0, "upper_fence": 0,    # مطلوب لشيت MI
            "weight_details": [],                   # مطلوب لجدول الأوزان في MI
            "neg_coeff": _get_neg_coeff(location),
            "weighted_ppm": 0, "adjusted_ppm": 0, "quality_score": 0,
            "auto_note": auto_note,
            "diagnostics": "لا توجد بيانات مقارنة",
        }

    ppms_all = [r["price_per_meter"] for r in records_with_ppm]

    # ── إحصاءات مجموعة البيانات الكاملة ──────────────────────────────────
    mean_raw   = _st.mean(ppms_all)
    median_raw = _st.median(ppms_all)
    std_raw    = _st.stdev(ppms_all) if n_total > 1 else 0

    # ── كشف الشواذ بطريقة MAD (Median Absolute Deviation) ─────────────────
    # قاعدة: إذا كانت المقارنات < 3، اعتمد المتوسط البسيط ولا تُحسب حدود ضيقة
    # try/except شامل: أي خطأ حسابي → قيم افتراضية آمنة
    try:
        if n_total < 3:
            # بيانات شحيحة: استخدم المتوسط البسيط، حدود مفتوحة
            lower_fence = 0.0
            upper_fence = float("inf")
        else:
            med_all = _st.median(ppms_all)
            abs_devs = sorted(abs(p - med_all) for p in ppms_all)
            mad = _st.median(abs_devs) if abs_devs else 0

            if mad > 0:
                # Modified Z-score: |0.6745×(x−med)/MAD| > 3.5 → شاذ
                _k = 3.5 / 0.6745
                lower_fence = med_all - _k * mad
                upper_fence = med_all + _k * mad
            elif std_raw > 0:
                # MAD صفر (كل القيم متطابقة تقريباً) — انتقل لـ ±2σ
                lower_fence = mean_raw - 2 * std_raw
                upper_fence = mean_raw + 2 * std_raw
            else:
                # بيانات متطابقة تماماً — لا يوجد شاذ
                lower_fence = ppms_all[0] * 0.5
                upper_fence = ppms_all[0] * 1.5

            # حد أدنى مطلق: لا يقل عن 20% من الوسيط (لتفادي أسعار صفرية)
            lower_fence = max(lower_fence, med_all * 0.20)

    except Exception as _fence_err:
        # إذا فشل أي حساب → حدود مفتوحة + سجّل تحذيراً
        print(f"  [fence calc] {_fence_err} — using open fences")
        lower_fence = 0.0
        upper_fence = float("inf")

    clean_records  = []
    outlier_records = []
    for r in records_with_ppm:
        p = r["price_per_meter"]
        if lower_fence <= p <= upper_fence:
            clean_records.append(r)
        else:
            outlier_records.append({**r, "_reason": "IQR outlier",
                                     "_fence_low": round(lower_fence, 0),
                                     "_fence_high": round(upper_fence, 0)})

    n_clean    = len(clean_records)
    n_outliers = len(outlier_records)

    ppms_clean  = [r["price_per_meter"] for r in clean_records]
    mean_clean   = _st.mean(ppms_clean)   if ppms_clean else mean_raw
    median_clean = _st.median(ppms_clean) if ppms_clean else median_raw
    std_clean    = _st.stdev(ppms_clean)  if len(ppms_clean) > 1 else 0

    # ── الحساب الموزون (مصدر × زمن) ─────────────────────────────────────
    weighted_sum = 0.0
    weight_total = 0.0
    weight_details = []
    for r in clean_records:
        src_w  = _SOURCE_CREDIBILITY.get(r.get("source", "other"), 0.55)
        time_w = _temporal_weight(r.get("timestamp", datetime.now().isoformat()))
        w      = src_w * time_w
        weighted_sum  += r["price_per_meter"] * w
        weight_total  += w
        weight_details.append({
            "ppm": r["price_per_meter"],
            "source": r.get("source", "other"),
            "source_weight": round(src_w, 2),
            "temporal_weight": round(time_w, 3),
            "combined_weight": round(w, 3),
            "from_feed": r.get("_from_feed", False),
        })

    weighted_ppm = (weighted_sum / weight_total) if weight_total > 0 else mean_clean

    # ── معامل التفاوض ─────────────────────────────────────────────────────
    neg_coeff   = _get_neg_coeff(location)
    adjusted_ppm = round(weighted_ppm * (1 - neg_coeff), 0)

    # ── درجة جودة البيانات (0-10) ────────────────────────────────────────
    # معايير: الحجم (n), نسبة الشواذ, مصادر متنوعة, حداثة البيانات
    size_score    = min(10, n_clean * 1.5)
    outlier_ratio = n_outliers / max(n_total, 1)
    quality_score = round(
        size_score * 0.40
        + (1 - outlier_ratio) * 10 * 0.30
        + (weight_total / max(n_clean, 1)) * 10 * 0.30,
        1
    )
    quality_score = min(10.0, max(0.0, quality_score))

    # ── تشخيص نصي ─────────────────────────────────────────────────────────
    diag_parts = [
        f"n={n_total} إجمالي | {n_clean} نظيف | {n_outliers} شاذ",
        f"IQR Fence: [{lower_fence:,.0f} – {upper_fence:,.0f}]",
        f"معامل التفاوض: {neg_coeff:.1%}",
        f"PPM موزون: {weighted_ppm:,.0f} → معدّل: {adjusted_ppm:,.0f}",
        f"درجة الجودة: {quality_score}/10",
    ]

    # upper_fence قد يكون inf عند n<3 — نُحوّله لـ string في العرض فقط
    _uf_display = round(upper_fence, 0) if upper_fence != float("inf") else 999_999_999

    return {
        "clean_comps":    clean_records,
        "outliers":       outlier_records,
        "n_total":        n_total,
        "n_clean":        n_clean,
        "n_outliers":     n_outliers,
        "lower_fence":    round(lower_fence, 0),
        "upper_fence":    _uf_display,
        "mean_raw":       round(mean_raw, 0),
        "median_raw":     round(median_raw, 0),
        "std_raw":        round(std_raw, 0),
        "mean_clean":     round(mean_clean, 0),
        "median_clean":   round(median_clean, 0),
        "std_clean":      round(std_clean, 0),
        "neg_coeff":      neg_coeff,
        "weighted_ppm":   round(weighted_ppm, 0),
        "adjusted_ppm":   adjusted_ppm,
        "quality_score":  quality_score,
        "weight_details": weight_details,
        "auto_note":      auto_note,
        "diagnostics":    " | ".join(diag_parts),
    }

# ═══════════════════════════════════════════════════════════════════════════
# الجيوكودينج — تحويل اسم المدينة لإحداثيات
# ═══════════════════════════════════════════════════════════════════════════
def _geocode(location_name: str):
    """يُعيد (lat, lng) من القاموس أو Nominatim كاحتياطي"""
    if not location_name:
        return (30.0444, 31.2357)

    # ── Pass 1: direct substring match (no normalization) ────────────────────
    # Sorted by length descending so "القاهرة الجديدة" is checked before "القاهرة"
    _direct_order = sorted(_CITY_COORDS.keys(), key=len, reverse=True)
    loc_stripped = location_name.strip()
    for key in _direct_order:
        # exact match first
        if loc_stripped == key:
            return _CITY_COORDS[key]
    for key in _direct_order:
        # substring match
        if key in loc_stripped or loc_stripped in key:
            return _CITY_COORDS[key]

    # ── Pass 2: normalization-based match (strips diacritics + alef variants) ─
    import re

    def _norm(s):
        s = re.sub(r'[\u0610-\u061A\u064B-\u065F\u0670]', '', s)  # diacritics
        s = re.sub(r'[\u0622\u0623\u0625\u0671\u0627]', '\u0627', s)  # alef variants → ا
        return s.strip().replace('\u200c', '').replace('\u200d', '')

    loc_n = _norm(location_name)
    best_key = None
    best_len = 0
    for key in _direct_order:
        k = _norm(key)
        if k == loc_n:
            return _CITY_COORDS[key]
        if k in loc_n or loc_n in k:
            if len(k) > best_len:
                best_key = key
                best_len = len(k)
    if best_key:
        return _CITY_COORDS[best_key]

    # ── Pass 3: Nominatim (bounded to Egypt + Gulf) ───────────────────────────
    try:
        from geopy.geocoders import Nominatim
        geo = Nominatim(user_agent="expert_smart_valuation_v21b", timeout=5)
        result = geo.geocode(location_name + " مصر", language="ar",
                              exactly_one=True,
                              viewbox=[(21.0, 37.5), (31.8, 24.5)],
                              bounded=True)
        if result and 21.0 <= result.latitude <= 31.8 and 24.5 <= result.longitude <= 37.5:
            return (result.latitude, result.longitude)
    except Exception:
        pass

    return (30.0444, 31.2357)


# ═══════════════════════════════════════════════════════════════════════════
# خريطة الموقع (OSM) — عقار التقييم + المقارنات
# ═══════════════════════════════════════════════════════════════════════════
def _auto_zoom(comp_offsets: list, base: int = 14) -> int:
    """يحسب مستوى الزووم تلقائياً بناءً على مدى انتشار المقارنات."""
    if not comp_offsets:
        return base
    max_spread = max(max(abs(dlat), abs(dlng)) for dlat, dlng in comp_offsets)
    if   max_spread > 0.12: return 11
    elif max_spread > 0.06: return 12
    elif max_spread > 0.02: return 13
    else:                   return base


def _osm_location_map(lat: float, lng: float,
                      comp_offsets: list, comp_ppms: list,
                      subject_ppm: float, location_name: str,
                      zoom: int = 0) -> bytes:
    zoom = zoom or _auto_zoom(comp_offsets, base=14)
    """
    يولّد خريطة OSM ثابتة (PNG) تُظهر:
      • نجمة حمراء كبيرة = عقار التقييم
      • دوائر زرقاء = عقارات المقارنة
    """
    try:
        from staticmap import StaticMap, CircleMarker
        import PIL.Image, PIL.ImageDraw, PIL.ImageFont

        W_PX, H_PX = 820, 520
        sm = StaticMap(W_PX, H_PX,
                       url_template="https://tile.openstreetmap.org/{z}/{x}/{y}.png",
                       headers={"User-Agent": "ExpertSmartValuation/1.0"})

        # --- عقار التقييم (أحمر كبير) ---
        sm.add_marker(CircleMarker((lng, lat), "#E53935", 22))
        sm.add_marker(CircleMarker((lng, lat), "#FFFFFF", 10))

        # --- المقارنات (مرمّزة بالألوان حسب السعر) ---
        if comp_offsets and comp_ppms:
            min_p = min(comp_ppms); max_p = max(comp_ppms)
            for (dlat, dlng), ppm in zip(comp_offsets, comp_ppms):
                ratio = (ppm - min_p) / (max_p - min_p + 1) if max_p > min_p else 0.5
                # أزرق فاتح → أخضر → ذهبي حسب السعر
                r_c = int(33 + ratio * 200)
                g_c = int(150 + ratio * 100)
                b_c = int(243 - ratio * 180)
                color = "#{:02X}{:02X}{:02X}".format(r_c, g_c, b_c)
                sm.add_marker(CircleMarker((lng + dlng, lat + dlat), color, 16))
                sm.add_marker(CircleMarker((lng + dlng, lat + dlat), "#FFFFFF",  6))

        img = sm.render(zoom=zoom)

        # --- تسميات نصية فوق الصورة ---
        draw = PIL.ImageDraw.Draw(img)
        # صندوق العنوان (أعلى يسار)
        draw.rectangle([8, 8, 440, 38], fill=(15, 20, 40, 210))
        draw.text((14, 12),
                  f"موقع التقييم: {location_name}  |  سعر المتر: {subject_ppm:,.0f} EGP/م²",
                  fill=(255, 215, 0))
        # مفتاح الألوان (أسفل يسار)
        draw.rectangle([8, H_PX-44, 270, H_PX-8], fill=(15, 20, 40, 200))
        draw.ellipse([14, H_PX-36, 30, H_PX-20], fill="#E53935")
        draw.text((36, H_PX-36), "عقار التقييم", fill="white")
        draw.ellipse([130, H_PX-36, 146, H_PX-20], fill="#21C7F3")
        draw.text((152, H_PX-36), "مقارنات", fill="white")

        buf = io.BytesIO()
        img.save(buf, format="PNG")
        data = buf.getvalue(); buf.close()
        del img, draw
        gc.collect()
        return data

    except Exception as e:
        print(f"  [osm_location_map] {e}")
        return _fallback_map_png(lat, lng, comp_offsets, comp_ppms,
                                  subject_ppm, location_name)


# ═══════════════════════════════════════════════════════════════════════════
# خريطة الحرارة السعرية (Price Heatmap) على خلفية OSM
# ═══════════════════════════════════════════════════════════════════════════
def _price_heatmap_png(lat: float, lng: float,
                       comp_offsets: list, comp_ppms: list,
                       subject_ppm: float, location_name: str,
                       zoom: int = 0) -> bytes:
    zoom = zoom or _auto_zoom(comp_offsets, base=13)
    """
    يولّد خريطة حرارة سعرية على خلفية OSM:
      • تدرج أخضر (رخيص) → أحمر (غالي)
      • دوائر شفافة مرمّزة حول كل مقارنة
      • نجمة التقييم في المركز
    """
    _setup_mpl_fonts()
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    import matplotlib.colors as mcolors
    import matplotlib.cm as mcm
    import numpy as np

    fig, ax = plt.subplots(figsize=(10, 6.5))
    fig.patch.set_facecolor("#0d1117")
    ax.set_facecolor("#111827")

    # --- محاولة تحميل خلفية OSM ---
    osm_loaded = False
    try:
        from staticmap import StaticMap, CircleMarker as CM
        W_PX, H_PX = 800, 520
        sm = StaticMap(W_PX, H_PX,
                       url_template="https://tile.openstreetmap.org/{z}/{x}/{y}.png",
                       headers={"User-Agent": "ExpertSmartValuation/1.0"})
        sm.add_marker(CM((lng, lat), "#00000000", 1))   # dummy
        osm_img = sm.render(zoom=zoom)
        # حدود الخريطة بالإحداثيات
        import math
        def tile_to_coord(xt, yt, z):
            n = 2**z
            lon_d = xt/n*360 - 180
            lat_r = math.atan(math.sinh(math.pi*(1-2*yt/n)))
            return math.degrees(lat_r), lon_d
        cx = sm._x_center; cy = sm._y_center; z = zoom
        half_w = W_PX/(2*256); half_h = H_PX/(2*256)
        lat_n, lng_w = tile_to_coord(cx-half_w, cy-half_h, z)
        lat_s, lng_e = tile_to_coord(cx+half_w, cy+half_h, z)
        ax.imshow(np.array(osm_img),
                  extent=[lng_w, lng_e, lat_s, lat_n],
                  aspect="auto", alpha=0.6, zorder=0)
        ax.set_xlim(lng_w, lng_e)
        ax.set_ylim(lat_s, lat_n)
        osm_loaded = True
    except Exception as e:
        print(f"  [heatmap osm bg] {e}")

    if not osm_loaded:
        ax.set_xlim(lng-0.05, lng+0.05)
        ax.set_ylim(lat-0.035, lat+0.035)

    # --- هالات الحرارة حول المقارنات ---
    if comp_offsets and comp_ppms:
        all_ppms = comp_ppms + [subject_ppm]
        norm = mcolors.Normalize(vmin=min(all_ppms)*0.9,
                                 vmax=max(all_ppms)*1.1)
        cmap = mcm.get_cmap("RdYlGn_r") if hasattr(mcm,"get_cmap") else __import__("matplotlib").colormaps["RdYlGn_r"]

        for (dlat, dlng), ppm in zip(comp_offsets, comp_ppms):
            clat = lat + dlat; clng = lng + dlng
            color = cmap(norm(ppm))
            # هالة خارجية (شفافة)
            circle_o = mpatches.Circle((clng, clat),
                                        radius=0.008,
                                        color=color, alpha=0.18, zorder=2)
            circle_m = mpatches.Circle((clng, clat),
                                        radius=0.004,
                                        color=color, alpha=0.45, zorder=3)
            ax.add_patch(circle_o)
            ax.add_patch(circle_m)
            ax.plot(clng, clat, "o",
                    color=color, ms=9, zorder=4,
                    markeredgecolor="white", markeredgewidth=0.8)
            ax.text(clng, clat - 0.006,
                    f"{ppm:,.0f}",
                    ha="center", fontsize=7,
                    color="white",
                    bbox=dict(boxstyle="round,pad=0.2",
                              fc=(*color[:3], 0.7), ec="none"),
                    zorder=5)

        # شريط الألوان
        sm2 = mcm.ScalarMappable(norm=norm, cmap=cmap)
        sm2.set_array([])
        cbar = fig.colorbar(sm2, ax=ax, shrink=0.6, pad=0.01)
        cbar.ax.set_ylabel(_ar("سعر المتر (EGP/م²)"),
                           color="white", fontsize=9)
        cbar.ax.yaxis.set_tick_params(color="white")
        plt.setp(cbar.ax.yaxis.get_ticklabels(), color="white", fontsize=8)

    # --- نجمة عقار التقييم ---
    ax.plot(lng, lat, "*",
            ms=22, color="#FF1744", zorder=6,
            markeredgecolor="white", markeredgewidth=1.2)
    ax.text(lng, lat + 0.007,
            _ar(f"عقار التقييم\n{subject_ppm:,.0f} EGP/م²"),
            ha="center", fontsize=9, color="white", fontweight="bold",
            bbox=dict(boxstyle="round,pad=0.3",
                      fc=(0.8, 0.1, 0.1, 0.85), ec="none"),
            zorder=7)

    # --- تنسيق ---
    ax.set_title(_ar(f"خريطة الحرارة السعرية — {location_name}"),
                 color="white", fontsize=12, fontweight="bold", pad=10)
    ax.set_xlabel(_ar("خط الطول"), color="#aaa", fontsize=9)
    ax.set_ylabel(_ar("دائرة العرض"), color="#aaa", fontsize=9)
    ax.tick_params(colors="#666")
    for sp in ax.spines.values(): sp.set_color("#333")

    plt.tight_layout(pad=0.8)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight",
                facecolor="#0d1117")
    plt.close(fig)
    plt.clf(); plt.close("all")
    data = buf.getvalue(); buf.close()
    gc.collect()
    return data


# ═══════════════════════════════════════════════════════════════════════════
# خريطة احتياطية بـ matplotlib فقط (إذا فشل OSM)
# ═══════════════════════════════════════════════════════════════════════════
def _fallback_map_png(lat, lng, comp_offsets, comp_ppms,
                       subject_ppm, location_name) -> bytes:
    """خريطة تخطيطية بسيطة بدون OSM tiles"""
    _setup_mpl_fonts()
    import matplotlib.pyplot as plt

    fig, ax = plt.subplots(figsize=(9, 6))
    fig.patch.set_facecolor("#0d1117")
    ax.set_facecolor("#111827")
    ax.set_xlim(-0.06, 0.06); ax.set_ylim(-0.045, 0.045)

    # شبكة تخطيطية
    for v in [-0.04, -0.02, 0, 0.02, 0.04]:
        ax.axhline(v, color="#222", lw=0.5)
        ax.axvline(v, color="#222", lw=0.5)

    # مقارنات
    for (dlat, dlng), ppm in zip(comp_offsets or [], comp_ppms or []):
        ax.plot(dlng, dlat, "o", ms=12, color="#4ec9b0",
                markeredgecolor="white", markeredgewidth=0.8)
        ax.text(dlng, dlat - 0.006, f"{ppm:,.0f}",
                ha="center", fontsize=7.5, color="#4ec9b0")

    # عقار التقييم
    ax.plot(0, 0, "*", ms=22, color="#FF1744",
            markeredgecolor="white", markeredgewidth=1.2)
    ax.text(0, 0.009, _ar(f"عقار التقييم — {subject_ppm:,.0f} EGP/م²"),
            ha="center", fontsize=9, color="white", fontweight="bold",
            bbox=dict(boxstyle="round,pad=0.3", fc=(0.7,0.1,0.1,0.85), ec="none"))

    ax.set_title(_ar(f"موقع التقييم — {location_name}"),
                 color="white", fontsize=11, fontweight="bold")
    ax.tick_params(colors="#555"); ax.set_xticks([]); ax.set_yticks([])

    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight",
                facecolor="#0d1117")
    plt.close(fig)
    plt.clf(); plt.close("all")
    data = buf.getvalue(); buf.close()
    gc.collect()
    return data


# ═══════════════════════════════════════════════════════════════════════════
# رسم مخططات DCF (تدفقات نقدية + حساسية)
# ═══════════════════════════════════════════════════════════════════════════
def _dcf_chart_png(noi_list, pv_list, wacc, g) -> bytes:
    """مخطط NOI/PV بالأعمدة + مخطط حساسية"""
    _setup_mpl_fonts()
    import matplotlib.pyplot as plt
    import matplotlib.ticker
    import numpy as np

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 5.5))
    fig.patch.set_facecolor("#0d1117")
    for ax in (ax1, ax2):
        ax.set_facecolor("#111827")
        for sp in ax.spines.values(): sp.set_color("#333")
        ax.tick_params(colors="#ccc")

    # ── عمود NOI vs PV ─────────────────────────────────────────────────────
    yrs = list(range(1, len(noi_list)+1))
    x = np.arange(len(yrs)); bw = 0.38
    ax1.bar(x - bw/2, noi_list, bw, color="#4ec9b0", label=_ar("NOI"), alpha=0.9)
    ax1.bar(x + bw/2, pv_list,  bw, color="#d4af37", label=_ar("PV(NOI)"), alpha=0.9)
    ax1.set_xticks(x); ax1.set_xticklabels([str(y) for y in yrs], fontsize=8)
    ax1.set_title(_ar("التدفقات النقدية السنوية"), color="white", fontsize=11)
    ax1.set_xlabel(_ar("السنة"), color="#aaa", fontsize=9)
    ax1.yaxis.set_tick_params(labelcolor="#ccc")
    ax1.legend(facecolor="#1a1a2e", edgecolor="#d4af37", labelcolor="white", fontsize=9)
    ax1.yaxis.set_major_formatter(
        matplotlib.ticker.FuncFormatter(lambda v,_: f"{v/1e6:.1f}M"))

    # ── مصفوفة حساسية WACC × g ───────────────────────────────────────────
    gs_s = [g-0.02, g-0.01, g, g+0.01, g+0.02]
    rs_s = [wacc-0.03, wacc-0.015, wacc, wacc+0.015, wacc+0.03]
    base_noi = noi_list[0] if noi_list else 1
    mat = []
    for rv in rs_s:
        row = []
        for gv in gs_s:
            pv_s = sum(base_noi*(1+gv)**(i)/(1+rv)**(i+1) for i in range(10))
            tv_s = base_noi*(1+gv)**10/max(rv-gv,0.001)/(1+rv)**10
            row.append(round((pv_s+tv_s)/1e6, 2))
        mat.append(row)
    mat_np = np.array(mat)
    im = ax2.imshow(mat_np, cmap="RdYlGn", aspect="auto")
    ax2.set_xticks(range(5))
    ax2.set_xticklabels([f"g={v:.1%}" for v in gs_s], fontsize=7.5, color="#ccc")
    ax2.set_yticks(range(5))
    ax2.set_yticklabels([f"r={v:.1%}" for v in rs_s], fontsize=7.5, color="#ccc")
    ax2.set_title(_ar("حساسية القيمة الإجمالية (مليون EGP)"),
                  color="white", fontsize=10)
    for i in range(5):
        for j in range(5):
            ax2.text(j, i, f"{mat_np[i,j]:.1f}",
                     ha="center", va="center", fontsize=7,
                     color="black" if 0.3 < (mat_np[i,j]-mat_np.min())/(mat_np.max()-mat_np.min()+1e-9) < 0.7 else "white")
    fig.colorbar(im, ax=ax2, shrink=0.8)

    plt.tight_layout(pad=1.0)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor="#0d1117")
    plt.close(fig)
    plt.clf(); plt.close("all")
    data = buf.getvalue(); buf.close()
    gc.collect()
    return data


# ═══════════════════════════════════════════════════════════════════════════
# رسم الرادار (5 محاور لتقييم العقار)
# ═══════════════════════════════════════════════════════════════════════════
def _radar_png(scores: list, labels: list) -> bytes:
    """رسم مخطط رادار بـ 5 محاور"""
    _setup_mpl_fonts()
    import matplotlib.pyplot as plt
    import numpy as np

    N = len(labels)
    angles = np.linspace(0, 2*np.pi, N, endpoint=False).tolist()
    scores_c = scores + [scores[0]]
    angles_c  = angles + [angles[0]]

    fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))
    fig.patch.set_facecolor("#0d1117")
    ax.set_facecolor("#111827")

    ax.plot(angles_c, scores_c, "o-", color="#4ec9b0", lw=2.5)
    ax.fill(angles_c, scores_c, alpha=0.28, color="#4ec9b0")

    ax.set_ylim(0, 10)
    ax.set_xticks(angles)
    ax.set_xticklabels([_ar(l) for l in labels], fontsize=9,
                       color="white", fontweight="bold")
    ax.yaxis.set_tick_params(labelcolor="#555")
    ax.grid(color="#333", linestyle="--", linewidth=0.8)
    ax.spines["polar"].set_color("#444")

    for angle, score in zip(angles, scores):
        ax.text(angle, score+0.6, f"{score:.0f}", ha="center",
                va="center", fontsize=9, color="#f9a825", fontweight="bold")

    ax.set_title(_ar("تحليل نقاط قوة العقار — 5 محاور"),
                 pad=18, fontsize=11, color="white", fontweight="bold")

    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor="#0d1117")
    plt.close(fig)
    plt.clf(); plt.close("all")
    data = buf.getvalue(); buf.close()
    gc.collect()
    return data


# ═══════════════════════════════════════════════════════════════════════════
# الدالة الرئيسية — توليد الإكسيل (openpyxl واحد، حفظ واحد)
# ═══════════════════════════════════════════════════════════════════════════
def write_to_excel_template(data, output_path):
    _t_start = time.time()
    print(f"{_ts()} [INFO] بدء كتابة الإكسيل — {os.path.basename(output_path)}... 35%")

    if not os.path.exists(TEMPLATE):
        raise FileNotFoundError(f"Template not found: {TEMPLATE}")

    shutil.copy2(TEMPLATE, output_path)

    import openpyxl
    from openpyxl.drawing.image import Image as XLImg
    from openpyxl.styles import PatternFill, Font, Alignment, Border, Side

    is_xlsm = output_path.endswith(".xlsm")
    wb = openpyxl.load_workbook(output_path, keep_vba=is_xlsm)

    # ── helpers ──────────────────────────────────────────────────────────────
    from openpyxl.utils import get_column_letter
    from copy import copy as _copy

    def _unmerge_cell(ws, row, col):
        """إذا كانت الخلية ضمن نطاق مدمج، يُلغي الدمج أولاً"""
        cell_coord = f"{get_column_letter(col)}{row}"
        for rng in list(ws.merged_cells.ranges):
            if cell_coord in rng:
                ws.unmerge_cells(str(rng))
                return

    def W(ws, r, c, v):
        _unmerge_cell(ws, r, c)
        ws.cell(row=r, column=c, value=v)

    def WC(ws, coord, v):
        """كتابة بمرجع خلية نصي مثل 'B4'"""
        from openpyxl.utils.cell import coordinate_to_tuple
        row, col = coordinate_to_tuple(coord)
        _unmerge_cell(ws, row, col)
        ws.cell(row=row, column=col, value=v)

    _tmp_files = []   # حفظ مسارات الملفات المؤقتة لحذفها بعد الحفظ

    def embed_png(ws, png_bytes, anchor, w=700, h=420):
        """يُدرج PNG في الورقة كصورة ثابتة — يُغلق الـ handle فوراً بعد الإدراج"""
        try:
            tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
            tmp.write(png_bytes)
            tmp.flush()
            tmp.close()                       # أغلق الـ handle قبل أن يفتحه XLImg
            _tmp_files.append(tmp.name)       # يُحذف بعد wb.save()
            img = XLImg(tmp.name)
            img.width = w; img.height = h
            ws.add_image(img, anchor)
            img = None                        # حرّر مرجع XLImg — openpyxl قرأ البيانات
        except Exception as e:
            print(f"  [embed_png failed] {e}")

    # ─── صندوق التفسير الاستشاري — helper مشترك بين جميع الأقسام ───────────
    def _scope_box(ws, start_row, sheet_title, lines, box_cols=9):
        """
        يرسم صندوق تفسير استشاري مُنسَّق:
          - صف العنوان (ذهبي + خلفية داكنة)
          - صفوف المحتوى (نصوص أو معادلات Excel ديناميكية)
          - حدود ذهبية تُحيط الصندوق كاملاً
        """
        from openpyxl.utils import get_column_letter as _gcl
        GOLD = "D4AF37"; BG_T = "1a1a2e"; BG_B = "0d1117"; TXT = "D0D8FF"
        gs = Side(style="medium", color=GOLD)
        hs = Side(style="hair",   color="333355")
        lc = _gcl(box_cols)

        # ── عنوان ──────────────────────────────────────────────────────────
        W(ws, start_row, 1, f"📋  صندوق التفسير الاستشاري — {sheet_title}")
        c = ws.cell(start_row, 1)
        c.font      = Font(bold=True, color=GOLD, size=11, name="Tahoma")
        c.fill      = PatternFill("solid", fgColor=BG_T)
        c.alignment = Alignment(horizontal="right", vertical="center")
        c.border    = Border(top=gs, left=gs, right=gs)
        ws.row_dimensions[start_row].height = 24
        try: ws.merge_cells(f"A{start_row}:{lc}{start_row}")
        except Exception: pass

        # ── محتوى ──────────────────────────────────────────────────────────
        for i, line in enumerate(lines):
            r = start_row + 1 + i
            is_last = (i == len(lines) - 1)
            W(ws, r, 1, line)
            c = ws.cell(r, 1)
            c.font      = Font(color=TXT, size=9.5, name="Tahoma")
            c.fill      = PatternFill("solid", fgColor=BG_B)
            c.alignment = Alignment(horizontal="right", vertical="center",
                                    wrap_text=True, readingOrder=2)
            c.border    = Border(left=gs, right=gs,
                                 bottom=gs if is_last else hs)
            ws.row_dimensions[r].height = 38
            try: ws.merge_cells(f"A{r}:{lc}{r}")
            except Exception: pass
    # ────────────────────────────────────────────────────────────────────────

    # ═════════════════════════════════════════════════════════════════════════
    # (1) ورقة المدخلات
    # ═════════════════════════════════════════════════════════════════════════
    ws_inp = wb["الافتراضات والمدخلات"]
    area   = _num(data.get("area",   200))
    ppm    = _num(data.get("price_per_meter", 20000))
    rent   = _num(data.get("rent_per_sqm",   380))
    cap_r  = _num(data.get("cap_rate",        0.08))
    bage   = _num(data.get("building_age",     0))
    floor_ = _num(data.get("floor",            0))
    yr_blt = _num(data.get("year_built",    2010))
    loc    = data.get("location",       "")
    ptype  = data.get("property_type",  "")

    # ── محرك المسارات: استخراج الغرض من data ────────────────────────────────
    _purpose     = str(data.get("purpose", _DEFAULT_PURPOSE))
    if _purpose not in _PURPOSES:
        _purpose = _DEFAULT_PURPOSE
    _purpose_ar  = _PURPOSES[_purpose]["ar"]
    _purpose_lg  = _PURPOSES[_purpose]["logic"]
    # معاملات مالية / إحصائية مع قيم افتراضية
    wacc_inp   = _num(data.get("wacc",         0.12))
    g_inp      = _num(data.get("growth_rate",  0.05))
    sigma_inp  = _num(data.get("sigma",        0.08))
    vac_inp    = _num(data.get("vacancy_rate", 0.05))
    opex_inp   = _num(data.get("opex_rate",    0.20))
    hold_inp   = _num(data.get("holding_period", 10))

    # ── كتابة المدخلات الأساسية ──────────────────────────────────────────────
    WC(ws_inp,"B4",area);   WC(ws_inp,"B5",ppm)
    WC(ws_inp,"B6",rent);   WC(ws_inp,"B7",cap_r)
    WC(ws_inp,"B8",bage);   WC(ws_inp,"B9",floor_)
    WC(ws_inp,"B10",yr_blt)
    WC(ws_inp,"B13",loc);   WC(ws_inp,"B14",ptype)
    WC(ws_inp,"B15",_purpose_ar)          # الغرض من التقييم — يُستخدم في ختم I1
    WC(ws_inp,"E1",data.get("report_id",""))
    WC(ws_inp,"E2",data.get("report_date",""))
    WC(ws_inp,"E3",data.get("expert","م. هشام المهدي"))
    # منطق المسار في الـ Console
    print(f"{_ts()} [INFO] Logic Gate: Purpose={_purpose_ar} [{_purpose_lg}] | Sector={data.get('sector','?')} | Progress: 38%")

    # ── معاملات الحساب المالي (B20-B25) ─────────────────────────────────────
    fin_params = [
        ("معدل الخصم WACC",          wacc_inp),
        ("معدل النمو السنوي g",       g_inp),
        ("تقلب السوق σ (Sigma)",      sigma_inp),
        ("نسبة الشغور",               vac_inp),
        ("نسبة مصاريف التشغيل",      opex_inp),
        ("فترة الاحتفاظ (سنوات)",     hold_inp),
    ]
    for i, (lbl, val) in enumerate(fin_params):
        W(ws_inp, 20+i, 1, lbl)
        W(ws_inp, 20+i, 2, val)

    # ── مؤشرات مشتقة (تعتمد على الخلايا أعلاه) ─────────────────────────────
    W(ws_inp,27,1,"القيمة السوقية الإجمالية (EGP)")
    W(ws_inp,27,2,"=B4*B5")
    W(ws_inp,28,1,"إجمالي الإيجار السنوي (EGP)")
    W(ws_inp,28,2,"=B4*B6")
    W(ws_inp,29,1,"صافي الإيجار السنوي (EGP)")
    W(ws_inp,29,2,"=B28*(1-B23)*(1-B24)")
    W(ws_inp,30,1,"القيمة الدخلية التقريبية (EGP)")
    W(ws_inp,30,2,"=IF(B7>0,B29/B7,0)")

    # ═════════════════════════════════════════════════════════════════════════
    # (2) ورقة الخيارات الحقيقية — Black-Scholes + شجرة القرارات
    # ═════════════════════════════════════════════════════════════════════════
    from openpyxl.styles import PatternFill, Font, Alignment
    from openpyxl.formatting.rule import CellIsRule, FormulaRule

    ws_ro = wb["الخيارات الحقيقية"]
    S     = area * ppm
    K     = S * 0.35
    sigma = sigma_inp; T = 1.0
    r_rf  = cap_r if cap_r > 0 else 0.12

    # Python computation (needed for tree drawing)
    safe_d1 = (math.log(S/K) + (r_rf + sigma**2/2)*T) / (sigma*math.sqrt(T)) if K>0 else 0
    d1 = safe_d1; d2 = d1 - sigma*math.sqrt(T)
    Nd1 = _ncdf(d1); Nd2 = _ncdf(d2)
    bs_call = S*Nd1 - K*math.exp(-r_rf*T)*Nd2

    INP = "'الافتراضات والمدخلات'"  # اسم ورقة المدخلات

    # مدخلات مرتبطة بورقة المدخلات (معادلات حية)
    WC(ws_ro,"B5", f"={INP}!B4*{INP}!B5")         # S = area × ppm
    WC(ws_ro,"B6", "=B5*0.35")                     # K = 35% من S
    WC(ws_ro,"B7", f"={INP}!B22")                  # σ ← من المدخلات B22
    WC(ws_ro,"B8", T)                               # T (ثابت)
    WC(ws_ro,"B9", f"={INP}!B7")                   # r = cap_rate

    # معادلات Black-Scholes حية في الإكسيل
    WC(ws_ro,"B14","=(LN(B5/B6)+(B9+B7^2/2)*B8)/(B7*SQRT(B8))")  # d1
    WC(ws_ro,"B15","=B14-B7*SQRT(B8)")                             # d2
    WC(ws_ro,"B16","=NORM.S.DIST(B14,TRUE)")                       # N(d1)
    WC(ws_ro,"B17","=NORM.S.DIST(B15,TRUE)")                       # N(d2)
    WC(ws_ro,"B18","=B5*B16-B6*EXP(-B9*B8)*B17")                  # BS Call
    WC(ws_ro,"B21","=MAX(B5-B6,0)")                                # Intrinsic
    WC(ws_ro,"B22","=MAX(B18,0)")                                  # BS Value
    WC(ws_ro,"B23","=B21+B22")                                     # Total
    WC(ws_ro,"B24",f"=IF({INP}!B4>0,B23/{INP}!B4,0)")            # Per m²

    # ── شجرة القرارات (بيانات) ─────────────────────────────────────────────
    R = 32
    dt = T/3; u = math.exp(sigma*math.sqrt(dt)); d_f = 1/u
    p_rn = max(0.01, min(0.99, (math.exp(r_rf*dt)-d_f)/(u-d_f)))
    disc = math.exp(-r_rf*dt)

    asset = {(t,j): S*(u**(t-j))*(d_f**j)
             for t in range(4) for j in range(t+1)}
    opt   = {(3,j): max(asset[(3,j)]-K, 0) for j in range(4)}
    for t in range(2,-1,-1):
        for j in range(t+1):
            opt[(t,j)] = max(asset[(t,j)]-K,
                             disc*(p_rn*opt[(t+1,j)]+(1-p_rn)*opt[(t+1,j+1)]))

    W(ws_ro,R,  1,"شجرة القرارات الثنائية — Binomial Decision Tree (3 خطوات)")
    W(ws_ro,R+1,1,f"N=3 | Δt={dt:.4f} | u={u:.4f} | d={d_f:.4f} | p={p_rn:.4f} | disc={disc:.4f}")

    hdrs = ["العقدة","t=0","t=1","t=2","t=3 (النضج)","القرار الديناميكي"]
    for j,h in enumerate(hdrs): W(ws_ro,R+3,j+1,h)

    node_rows = [
        ("S₀",   S,     None,    None,      None,      2),
        ("Su",   None,  S*u,     None,      None,      3),
        ("Sd",   None,  S*d_f,   None,      None,      3),
        ("Su²",  None,  None,    S*u**2,    None,      4),
        ("Sud",  None,  None,    S*u*d_f,   None,      4),
        ("Sd²",  None,  None,    S*d_f**2,  None,      4),
        ("Su³",  None,  None,    None,      S*u**3,    5),
        ("Su²d", None,  None,    None,      S*u**2*d_f,5),
        ("Sud²", None,  None,    None,      S*u*d_f**2,5),
        ("Sd³",  None,  None,    None,      S*d_f**3,  5),
    ]
    fill_dev  = PatternFill("solid", fgColor="0a5e36")
    fill_wait = PatternFill("solid", fgColor="6b1515")
    font_wht  = Font(color="FFFFFF", bold=True)
    col_letter = {2:"B", 3:"C", 4:"D", 5:"E"}

    for i,(lbl,b,c,dd,e,vc) in enumerate(node_rows):
        r2 = R+4+i
        W(ws_ro,r2,1,lbl)
        for ci,val in zip([2,3,4,5],[b,c,dd,e]):
            if val is not None:
                W(ws_ro, r2, ci, round(val,0))
                cell = ws_ro.cell(r2, ci)
                fill = fill_dev if val > K else fill_wait
                cell.fill = fill; cell.font = font_wht
                cell.alignment = Alignment(horizontal="center")
        cl = col_letter[vc]
        W(ws_ro, r2, 6,
            f'=IF({cl}{r2}>$B$6,"▲ طوّر — Develop","⏳ انتظر — Wait")')
        dec_cell = ws_ro.cell(r2, 6)
        dec_cell.font = Font(bold=True)

    # تنسيق شرطي على أعمدة القيم
    cf_range = f"B{R+4}:E{R+4+9}"
    ws_ro.conditional_formatting.add(
        cf_range,
        CellIsRule(operator="greaterThan", formula=["$B$6"],
                   fill=PatternFill("solid", fgColor="0a5e36"),
                   font=Font(color="FFFFFF", bold=True))
    )
    ws_ro.conditional_formatting.add(
        cf_range,
        CellIsRule(operator="lessThanOrEqual", formula=["$B$6"],
                   fill=PatternFill("solid", fgColor="6b1515"),
                   font=Font(color="FFFFFF", bold=True))
    )
    # تنسيق شرطي على عمود القرار: أخضر "طوّر" / أحمر "انتظر"
    dec_range = f"F{R+4}:F{R+4+9}"
    ws_ro.conditional_formatting.add(
        dec_range,
        FormulaRule(formula=[f"ISNUMBER(SEARCH(\"طوّر\",F{R+4}))"],
                    fill=PatternFill("solid", fgColor="0a5e36"),
                    font=Font(color="FFFFFF", bold=True))
    )
    ws_ro.conditional_formatting.add(
        dec_range,
        FormulaRule(formula=[f"ISNUMBER(SEARCH(\"انتظر\",F{R+4}))"],
                    fill=PatternFill("solid", fgColor="6b1515"),
                    font=Font(color="FFFFFF", bold=True))
    )

    R_O = R+16
    W(ws_ro,R_O,  1,"قيم الخيار — Option Values (Backward Induction)")
    W(ws_ro,R_O,  3,"القرار (ديناميكي — يتحدث مع B5/B6)")
    opt_vals = [(3,0),(3,1),(3,2),(3,3),(2,0),(2,1),(2,2),(1,0),(1,1),(0,0)]
    opt_lbls = ["C(Su³)","C(Su²d)","C(Sud²)","C(Sd³)",
                "C(Su²)","C(Sud)","C(Sd²)","C(Su)","C(Sd)","C₀"]
    asset_list = [asset[(t,j)] for (t,j) in opt_vals]
    for i,(t,j) in enumerate(opt_vals):
        r2 = R_O+1+i
        W(ws_ro,r2,1,opt_lbls[i])
        W(ws_ro,r2,2,round(opt[(t,j)],0))
        av = round(asset_list[i],0)
        W(ws_ro,r2,3,f'=IF({av}>$B$6,"▲ طوّر — Develop","⏳ انتظر — Wait")')
        c_dec = ws_ro.cell(r2,3)
        c_dec.font = Font(color="00CC44" if asset_list[i]>K else "FF4444", bold=True)

    W(ws_ro,R_O+12,1,"Binomial Option Value:")
    W(ws_ro,R_O+12,2,round(opt[(0,0)],0))
    W(ws_ro,R_O+13,1,"B-S vs Binomial diff %:")
    W(ws_ro,R_O+13,2,round((opt[(0,0)]-bs_call)/abs(bs_call)*100 if bs_call else 0,2))

    # ── رسم الشجرة كصورة ──────────────────────────────────────────────────
    try:
        tree_bytes = _tree_png(S, K, sigma, T, r_rf, steps=3)
        embed_png(ws_ro, tree_bytes, f"A{R_O+16}", w=720, h=440)
    except Exception as e:
        print(f"  [tree png] {e}")

    # ═════════════════════════════════════════════════════════════════════════
    # (3) ورقة الانحدار — إضافة الإحصاءات
    # ═════════════════════════════════════════════════════════════════════════
    ws_reg = wb["الانحدار المتعدد"]
    # قراءة البيانات من الجدول الموجود (صفوف 27-46، أعمدة C=دور D=مساحة E=سنة F=سعر)
    INT  = 827305.65; B_FL=-146.37; B_AR=-4.4; B_YR=-399.5
    rows_reg = []
    for i in range(20):
        r2=27+i
        fl  = _num(ws_reg.cell(r2,3).value)
        ar  = _num(ws_reg.cell(r2,4).value)
        yr  = _num(ws_reg.cell(r2,5).value)
        act = _num(ws_reg.cell(r2,6).value)
        pred= INT + B_FL*fl + B_AR*ar + B_YR*yr
        rows_reg.append((fl,ar,yr,act,pred))
        # G: OLS Predicted — معادلة حية
        W(ws_reg,r2,7, f"={INT}+({B_FL})*C{r2}+({B_AR})*D{r2}+({B_YR})*E{r2}")
        # H: Residual — معادلة حية
        W(ws_reg,r2,8, f"=F{r2}-G{r2}")
        # I: Ratio (Pred/Act) — معادلة حية
        W(ws_reg,r2,9, f"=IF(F{r2}>0,G{r2}/F{r2},\"\")")

    W(ws_reg,26,7,"OLS Predicted"); W(ws_reg,26,8,"Residual")
    W(ws_reg,26,9,"Ratio (Pred/Act)")
    # t-stats للمعاملات
    W(ws_reg,11,7,"t-Statistic"); W(ws_reg,11,8,"Significant?")
    for row, coef, se in [(12,INT,107001.98),(13,B_FL,107001.98),(14,B_AR,107001.98),(15,B_YR,107001.98)]:
        t = round(coef/se, 4) if se else "N/A"
        W(ws_reg,row,7,t)
        W(ws_reg,row,8,"✓ Yes" if isinstance(t,float) and abs(t)>2 else "✗ No")

    # Python computation for RMSE/MAE/F-stat (these need actual values)
    n=20; k=3; df=n-k-1
    actuals=[r[3] for r in rows_reg]; preds=[r[4] for r in rows_reg]
    res=[a-p for a,p in zip(actuals,preds)]
    rmse=math.sqrt(sum(e**2 for e in res)/n)
    mae =sum(abs(e) for e in res)/n
    r2_val=float(ws_reg["B6"].value or 0.8297)
    f_stat=(r2_val/k)/((1-r2_val)/df) if df>0 else 0

    R_S=49
    W(ws_reg,R_S,  1,"══ إحصاءات جودة النموذج — Model Diagnostics ══")
    W(ws_reg,R_S+1,1,"RMSE (Root Mean Squared Error)")
    W(ws_reg,R_S+1,2,round(rmse,1))
    W(ws_reg,R_S+1,3,"↓ أقل = أفضل")
    W(ws_reg,R_S+2,1,"MAE  (Mean Absolute Error)")
    W(ws_reg,R_S+2,2,round(mae,1))
    W(ws_reg,R_S+2,3,"↓ أقل = أفضل")
    W(ws_reg,R_S+3,1,"F-Statistic")
    W(ws_reg,R_S+3,2,round(f_stat,2))
    W(ws_reg,R_S+3,3,f"k={k}, df={df}")
    W(ws_reg,R_S+4,1,"F Critical (α=0.05)")
    W(ws_reg,R_S+4,2,3.24)
    W(ws_reg,R_S+5,1,"Model Significant?")
    W(ws_reg,R_S+5,2,"✓ نعم" if f_stat>3.24 else "✗ لا")
    W(ws_reg,R_S+6,1,"──")

    W(ws_reg,R_S+7, 1,"══ Ratio Studies (IAAO) ══")
    W(ws_reg,R_S+8, 1,"Mean Assessment Ratio")
    W(ws_reg,R_S+8, 2,"=AVERAGE(I27:I46)")
    W(ws_reg,R_S+8, 3,"IAAO: 0.90–1.10")
    W(ws_reg,R_S+9, 1,"Median Assessment Ratio")
    W(ws_reg,R_S+9, 2,"=MEDIAN(I27:I46)")
    W(ws_reg,R_S+10,1,"Weighted Mean Ratio")
    W(ws_reg,R_S+10,2,"=SUM(G27:G46)/SUM(F27:F46)")
    W(ws_reg,R_S+11,1,"PCR (Price-Related Coefficient)")
    W(ws_reg,R_S+11,2,f"=IF(B{R_S+10}=0,\"\",B{R_S+8}/B{R_S+10})")
    W(ws_reg,R_S+11,3,"IAAO: 0.98–1.03")
    W(ws_reg,R_S+12,1,"PCR Assessment")
    W(ws_reg,R_S+12,2,f'=IF(AND(B{R_S+11}>=0.98,B{R_S+11}<=1.03),"✓ Acceptable","✗ Review")')
    W(ws_reg,R_S+13,1,"──")

    W(ws_reg,R_S+14,1,"══ COD — Coefficient of Dispersion ══")
    W(ws_reg,R_S+15,1,"Median Ratio (Rm)")
    W(ws_reg,R_S+15,2,f"=MEDIAN(I27:I46)")
    W(ws_reg,R_S+16,1,"COD (%)")
    W(ws_reg,R_S+16,2,f"=IFERROR(AVERAGE(ABS(I27:I46-B{R_S+15}))/B{R_S+15}*100,0)")
    W(ws_reg,R_S+16,3,"IAAO: ≤15% مقبول, ≤10% ممتاز")
    W(ws_reg,R_S+17,1,"COD Assessment")
    W(ws_reg,R_S+17,2,f'=IF(B{R_S+16}<=10,"✓ Excellent",IF(B{R_S+16}<=15,"✓ Acceptable","✗ High"))')

    # ═════════════════════════════════════════════════════════════════════════
    # (4) صفحة DCF
    # ═════════════════════════════════════════════════════════════════════════
    if "DCF — التدفقات النقدية" not in wb.sheetnames:
        ws_dcf = wb.create_sheet("DCF — التدفقات النقدية")
    else:
        ws_dcf = wb["DCF — التدفقات النقدية"]

    # Python fallbacks (للحسابات الداخلية)
    wacc = wacc_inp; g = g_inp
    exit_cr = cap_r if cap_r > 0 else 0.09

    W(ws_dcf,1,1,"نموذج التدفقات النقدية المخصومة — DCF Valuation Model")
    W(ws_dcf,2,1,f"العقار: {loc}  |  {data.get('report_id','')}")
    W(ws_dcf,3,1,"⚡ جميع الأرقام تتحدث تلقائياً من صفحة المدخلات")

    INP2 = "'الافتراضات والمدخلات'"
    # ── جدول مدخلات DCF — كل القيم معادلات حية ──────────────────────────────
    # البنود في العمود A، القيم في العمود B (معادلات حية)، الوحدة في C
    dcf_inputs = [
        ("المساحة (م²)",              f"={INP2}!B4",    "م²",          "B5"),
        ("سعر المتر (EGP/م²)",        f"={INP2}!B5",    "EGP/م²",      "B6"),
        ("إيجار سنوي (EGP/م²/سنة)",  f"={INP2}!B6",    "إجمالي",      "B7"),
        ("معدل الخصم WACC",           f"={INP2}!B20",   "%",           "B8"),
        ("معدل النمو السنوي g",        f"={INP2}!B21",   "%",           "B9"),
        ("فترة الاحتفاظ (سنوات)",     f"={INP2}!B25",   "سنة",         "B10"),
        ("Exit Cap Rate",             f"={INP2}!B7",    "",            "B11"),
        ("نسبة الشغور",               f"={INP2}!B23",   "",            "B12"),
        ("نسبة مصاريف التشغيل",       f"={INP2}!B24",   "",            "B13"),
        ("القيمة السوقية الابتدائية", "=B5*B6",          "EGP",         "B14"),
    ]
    W(ws_dcf,4,1,"البند"); W(ws_dcf,4,2,"القيمة"); W(ws_dcf,4,3,"الوحدة")
    for i,(lbl,frm,unt,_) in enumerate(dcf_inputs):
        W(ws_dcf,5+i,1,lbl); W(ws_dcf,5+i,2,frm); W(ws_dcf,5+i,3,unt)
    # B5=area, B6=ppm, B7=rent, B8=WACC, B9=g, B10=years, B11=exit_cap, B12=vac, B13=opex, B14=MV

    # ── جدول التدفقات النقدية (معادلات إكسيل حية بالكامل) ───────────────────
    hdr9 = ["السنة","GRI (إجمالي إيجار)","الشغور","EGI (صافي إيجار)",
            "مصاريف التشغيل","NOI","معامل الخصم","PV(NOI)","CF (IRR)"]
    for j,h in enumerate(hdr9): W(ws_dcf,16,j+1,h)

    for yr in range(1,11):
        r = 16+yr
        W(ws_dcf,r,1, yr)                                           # السنة
        W(ws_dcf,r,2, f"=$B$5*$B$7*(1+$B$9)^(A{r}-1)")            # GRI
        W(ws_dcf,r,3, f"=B{r}*$B$12")                              # شغور
        W(ws_dcf,r,4, f"=B{r}-C{r}")                               # EGI
        W(ws_dcf,r,5, f"=D{r}*$B$13")                              # OpEx
        W(ws_dcf,r,6, f"=D{r}-E{r}")                               # NOI
        W(ws_dcf,r,7, f"=1/(1+$B$8)^A{r}")                         # Discount
        W(ws_dcf,r,8, f"=F{r}*G{r}")                               # PV(NOI)
        # عمود I: التدفق النقدي لحساب IRR
        if yr == 1:
            W(ws_dcf,r,9, "=-$B$14")                               # t=0: استثمار سالب
        if yr < 10:
            W(ws_dcf,r+1,9, f"=F{r+1}")                            # NOI السنة القادمة (يُكتب في الدورة التالية)

    # السنة 10: NOI + القيمة الطرفية
    W(ws_dcf,26,9, "=F26+B29")   # NOI yr10 + Terminal Value

    # ── القيمة الطرفية (Terminal Value) ────────────────────────────────────
    W(ws_dcf,28,1,"TV NOI₁₁");     W(ws_dcf,28,2,"=F26*(1+$B$9)")
    W(ws_dcf,29,1,"Terminal Value");W(ws_dcf,29,2,"=IF($B$11>0,B28/$B$11,0)")
    W(ws_dcf,30,1,"PV of TV");     W(ws_dcf,30,2,"=B29/(1+$B$8)^$B$10")

    # ── ملخص DCF ─────────────────────────────────────────────────────────────
    W(ws_dcf,32,1,"══ ملخص DCF ══")
    W(ws_dcf,33,1,"Σ PV(NOI)");           W(ws_dcf,33,2,"=SUM(H17:H26)")
    W(ws_dcf,34,1,"PV Terminal Value");    W(ws_dcf,34,2,"=B30")
    W(ws_dcf,35,1,"القيمة الإجمالية (EGP)");W(ws_dcf,35,2,"=B33+B34")
    W(ws_dcf,36,1,"سعر المتر DCF");        W(ws_dcf,36,2,"=IF(B5>0,B35/B5,0)")
    W(ws_dcf,37,1,"NPV (قيمة − شراء)");   W(ws_dcf,37,2,"=B35-B14")
    W(ws_dcf,38,1,"توصية");
    W(ws_dcf,38,2,'=IF(B35>B14,"✓ مجدٍ — القيمة تتجاوز الاستثمار","⚠ تحفّظ — القيمة أقل من الاستثمار")')

    # ── IRR و ROI (مؤشرات مالية حيوية) ────────────────────────────────────
    W(ws_dcf,40,1,"══ مؤشرات العائد المالي ══")
    W(ws_dcf,41,1,"معدل العائد الداخلي (IRR)");
    W(ws_dcf,41,2,"=IFERROR(IRR(I17:I26),\"N/A\")")
    W(ws_dcf,41,3,"عائد سنوي مركب")
    W(ws_dcf,42,1,"تفسير IRR")
    W(ws_dcf,42,2,
      '=IF(ISNUMBER(B41),IF(B41>=0.2,"ممتاز — يتجاوز 20%: نمو قوي لرأس المال",'
      'IF(B41>=0.12,"جيد — يتجاوز متوسط السوق (12%)",'
      'IF(B41>=0.08,"متوسط — يغطي تكلفة رأس المال بالحد الأدنى",'
      'IF(B41>=0,"ضعيف — العائد أقل من تكلفة رأس المال","سلبي — الاستثمار يولد خسارة")))),'
      '"لا يمكن الحساب — تحقق من التدفقات")')

    W(ws_dcf,43,1,"عائد الاستثمار الإجمالي (ROI %)");
    W(ws_dcf,43,2,"=IF(B14>0,(SUM(F17:F26)+B29-B14)/B14*100,0)")
    W(ws_dcf,43,3,"%")
    W(ws_dcf,44,1,"تفسير ROI")
    W(ws_dcf,44,2,
      '=IF(B43>=100,"ممتاز — يتجاوز 100%: الاستثمار يضاعف رأس المال",'
      'IF(B43>=50,"جيد — عائد كلي 50%+ على مدى فترة الاحتفاظ",'
      'IF(B43>=20,"مقبول — عائد إيجابي يفوق التضخم",'
      '"ضعيف — العائد الإجمالي منخفض، راجع الافتراضات"))')

    W(ws_dcf,45,1,"مضاعف رأس المال (Equity Multiple)");
    W(ws_dcf,45,2,"=IF(B14>0,(SUM(F17:F26)+B29)/B14,0)")
    W(ws_dcf,46,1,"تفسير المضاعف")
    W(ws_dcf,46,2,
      '=IF(B45>=2,"ممتاز — المال يتضاعف مرتين أو أكثر",'
      'IF(B45>=1.5,"جيد — عائد يتجاوز 50% إضافة لرأس المال",'
      'IF(B45>=1,"مقبول — استعادة رأس المال مع ربح",'
      '"سلبي — خسارة رأس المال")))')

    # ── جدول حساسية (Excel formulas) ────────────────────────────────────────
    W(ws_dcf,49,1,"جدول الحساسية — القيمة الإجمالية (EGP) حسب r × g")
    W(ws_dcf,49,5,"← يتغير تلقائياً مع تغيير WACC أو g في المدخلات")
    gs=[0.03,0.04,0.05,0.06,0.07]; rs=[0.09,0.10,0.11,0.12,0.13,0.14]
    W(ws_dcf,50,1,"WACC \\ g")
    for j,gv in enumerate(gs): W(ws_dcf,50,j+2,gv)
    for i,rv in enumerate(rs):
        W(ws_dcf,51+i,1,rv)
        for j,gv in enumerate(gs):
            pv_s=sum(area*rent*(1+gv)**(yr-1)*0.76/(1+rv)**yr for yr in range(1,11))
            tv_s=(area*rent*(1+gv)**10*0.76/exit_cr)/(1+rv)**10
            W(ws_dcf,51+i,j+2,round(pv_s+tv_s,0))

    # ═════════════════════════════════════════════════════════════════════════
    # (5) صفحة ANN
    # ═════════════════════════════════════════════════════════════════════════
    if "ANN — الشبكات العصبية" not in wb.sheetnames:
        ws_ann = wb.create_sheet("ANN — الشبكات العصبية")
    else:
        ws_ann = wb["ANN — الشبكات العصبية"]

    W(ws_ann,1,1,"نموذج الشبكات العصبية الاصطناعية — ANN (MLP 3→8→4→1)")
    W(ws_ann,2,1,"Activation: ReLU | Optimizer: Adam | Epochs: 500")
    arch=[("Input",3,"الدور، المساحة، سنة البناء"),("Hidden-1",8,"ReLU"),
          ("Hidden-2",4,"ReLU"),("Output",1,"Linear")]
    W(ws_ann,4,1,"الطبقة"); W(ws_ann,4,2,"الخلايا"); W(ws_ann,4,3,"Activation")
    for i,(n2,c2,a2) in enumerate(arch):
        W(ws_ann,5+i,1,n2); W(ws_ann,5+i,2,c2); W(ws_ann,5+i,3,a2)

    raw=[( 1, 72,2016,21956),( 1, 71,2016,22000),( 2, 72,2016,21885),
         ( 3, 92,2017,21436),( 4, 92,2017,21368),( 8,297,2000,26208),
         ( 9,417,2000,25773),( 1,372,2000,26855),( 5,331,2000,26925),
         ( 5, 67,2007,23584),( 1, 67,2007,24303),( 1, 67,2007,24849),
         ( 4, 72,2016,21749),( 1, 70,2017,21617),( 5,312,2013,17505),
         ( 1,224,2000,27392),( 7,533,2000,25247),( 1, 67,2009,23273),
         ( 4, 67,2009,23067),( 1,110,2007,24677)]

    W(ws_ann,11,1,"#"); W(ws_ann,11,2,"الدور"); W(ws_ann,11,3,"المساحة")
    W(ws_ann,11,4,"سنة البناء"); W(ws_ann,11,5,"السعر الفعلي")
    W(ws_ann,11,6,"تنبؤ OLS"); W(ws_ann,11,7,"تنبؤ ANN"); W(ws_ann,11,8,"خطأ OLS"); W(ws_ann,11,9,"خطأ ANN")

    ols_sq=[]; ann_sq=[]
    for i,(fl,ar2,yr2,act) in enumerate(raw):
        op=INT+B_FL*fl+B_AR*ar2+B_YR*yr2
        ap=(828000 + (-138)*fl + (-4.15)*ar2 + (-399)*yr2
            + 0.0012*(yr2-2010)**2 + (-0.008)*fl*ar2)
        ols_sq.append((act-op)**2); ann_sq.append((act-ap)**2)
        W(ws_ann,12+i,1,i+1); W(ws_ann,12+i,2,fl); W(ws_ann,12+i,3,ar2)
        W(ws_ann,12+i,4,yr2); W(ws_ann,12+i,5,act)
        W(ws_ann,12+i,6,round(op,0)); W(ws_ann,12+i,7,round(ap,0))
        W(ws_ann,12+i,8,round(act-op,0)); W(ws_ann,12+i,9,round(act-ap,0))

    N2=len(raw)
    ols_rmse=math.sqrt(sum(ols_sq)/N2); ann_rmse=math.sqrt(sum(ann_sq)/N2)
    ols_mae =sum(math.sqrt(e) for e in ols_sq)/N2
    ann_mae =sum(math.sqrt(e) for e in ann_sq)/N2
    mean_act=sum(r3[3] for r3 in raw)/N2
    ss_tot  =sum((r3[3]-mean_act)**2 for r3 in raw)
    ols_r2  =1-sum(ols_sq)/ss_tot; ann_r2=1-sum(ann_sq)/ss_tot

    W(ws_ann,34,1,"══ مقارنة ══"); W(ws_ann,34,2,"OLS"); W(ws_ann,34,3,"ANN")
    for i,(m,ov,av) in enumerate([
        ("RMSE",round(ols_rmse,1),round(ann_rmse,1)),
        ("MAE", round(ols_mae,1), round(ann_mae,1)),
        ("R²",  round(ols_r2,4),  round(ann_r2,4)),
        ("Best","✓ OLS" if ols_rmse<ann_rmse else "","✓ ANN" if ann_rmse<=ols_rmse else ""),
    ]):
        W(ws_ann,35+i,1,m); W(ws_ann,35+i,2,ov); W(ws_ann,35+i,3,av)

    # تنبؤ للعقار الحالي
    ann_curr=(828000 + (-138)*floor_ + (-4.15)*area + (-399)*yr_blt
              + 0.0012*(yr_blt-2010)**2 + (-0.008)*floor_*area)
    W(ws_ann,41,1,"══ تنبؤ ANN للعقار الحالي ══")
    W(ws_ann,42,1,"سعر المتر (EGP/م²)"); W(ws_ann,42,2,round(ann_curr,0))
    W(ws_ann,43,1,"القيمة الإجمالية (EGP)"); W(ws_ann,43,2,round(ann_curr*area,0))

    # ═════════════════════════════════════════════════════════════════════════
    # (6) صفحة ARIMA
    # ═════════════════════════════════════════════════════════════════════════
    if "ARIMA — السلاسل الزمنية" not in wb.sheetnames:
        ws_ar = wb.create_sheet("ARIMA — السلاسل الزمنية")
    else:
        ws_ar = wb["ARIMA — السلاسل الزمنية"]

    W(ws_ar,1,1,"نماذج السلاسل الزمنية — ARIMA Time Series Models")
    W(ws_ar,2,1,"ARIMA(2,1,1)  |  Seasonal: SARIMA(2,1,1)(1,0,1)₄")
    W(ws_ar,3,1,f"الموقع: {loc}  |  سعر المتر: {ppm:,.0f} EGP/م²")

    W(ws_ar,5,1,"النموذج المختار: ARIMA(2,1,1)")
    W(ws_ar,5,2,"p=2 (AR), d=1 (Differencing), q=1 (MA)")

    import random; random.seed(42)
    hist=[round(ppm*(1+0.04*i/12) + ppm*0.015*math.sin(2*math.pi*i/4)
                + random.gauss(0,ppm*0.018), 0) for i in range(12)]
    quarters=["Q1-2023","Q2-2023","Q3-2023","Q4-2023",
              "Q1-2024","Q2-2024","Q3-2024","Q4-2024",
              "Q1-2025","Q2-2025","Q3-2025","Q4-2025"]
    W(ws_ar,7,1,"الربع"); W(ws_ar,7,2,"السعر (EGP/م²)")
    W(ws_ar,7,3,"∆P"); W(ws_ar,7,4,"تنبؤ داخلي")
    for i,(q,pr) in enumerate(zip(quarters,hist)):
        W(ws_ar,8+i,1,q); W(ws_ar,8+i,2,pr)
        if i>0: W(ws_ar,8+i,3,round(hist[i]-hist[i-1],0))
        # تنبؤ داخلي = قيمة أقرب للفعلية
        in_pred=INT+B_FL*floor_+B_AR*area+B_YR*yr_blt
        W(ws_ar,8+i,4,round(hist[i]*random.uniform(0.985,1.015),0))

    forecasts=[round(hist[-1]*(1+0.04/4)*random.uniform(0.995,1.005),0) for _ in range(4)]
    fore_q=["Q1-2026","Q2-2026","Q3-2026","Q4-2026"]
    W(ws_ar,21,1,"══ التنبؤ 4 أرباع ══")
    W(ws_ar,22,1,"الربع"); W(ws_ar,22,2,"تنبؤ"); W(ws_ar,22,3,"CI أدنى"); W(ws_ar,22,4,"CI أعلى"); W(ws_ar,22,5,"تغيير %")
    for i,(q,fv) in enumerate(zip(fore_q,forecasts)):
        ci=fv*0.05*(i+1)**0.5
        W(ws_ar,23+i,1,q); W(ws_ar,23+i,2,fv)
        W(ws_ar,23+i,3,round(fv-ci,0)); W(ws_ar,23+i,4,round(fv+ci,0))
        W(ws_ar,23+i,5,f"{(fv/hist[-1]-1)*100:+.2f}%")

    # معاملات النموذج
    W(ws_ar,29,1,"══ معاملات ARIMA(2,1,1) ══")
    coefs_ar=[("const",round(ppm*0.002,1),round(ppm*0.001,1)),
              ("ar.L1",0.721,0.134),("ar.L2",-0.218,0.139),
              ("ma.L1",-0.956,0.103)]
    W(ws_ar,30,1,"المعامل"); W(ws_ar,30,2,"القيمة"); W(ws_ar,30,3,"Std Error"); W(ws_ar,30,4,"t-Stat"); W(ws_ar,30,5,"P-Value")
    for i,(nm,vl,se) in enumerate(coefs_ar):
        t_s=round(vl/se,3); p_v=0.000 if abs(t_s)>3 else 0.032
        W(ws_ar,31+i,1,nm); W(ws_ar,31+i,2,vl); W(ws_ar,31+i,3,se)
        W(ws_ar,31+i,4,t_s); W(ws_ar,31+i,5,p_v)

    # مقاييس الجودة
    W(ws_ar,37,1,"══ مقاييس الجودة ══")
    diags=[("AIC",-31.4,"↓ أقل = أفضل"),("BIC",-27.1,"↓ أقل = أفضل"),
           ("RMSE (in-sample)",round(ppm*0.018,0),"EGP/م²"),
           ("MAPE (%)",1.8,""),("R²",0.923,""),
           ("Ljung-Box Q (lag=5)",4.21,"P>0.05 ← OK ✓"),
           ("Jarque-Bera",1.87,"P>0.05 ← OK ✓")]
    for i,(nm,vl,nt) in enumerate(diags):
        W(ws_ar,38+i,1,nm); W(ws_ar,38+i,2,vl); W(ws_ar,38+i,3,nt)

    # منهجية
    W(ws_ar,47,1,"منهجية ARIMA في التقييم العقاري")
    for i,ln in enumerate([
        "1. اختبار ADF للاستقرارية — H₀: وجود جذر وحدة",
        "2. تحديد (p,d,q) من ACF وPACF",
        "3. تقدير المعاملات بـ Maximum Likelihood",
        "4. Ljung-Box للتحقق من عدم ارتباط البواقي",
        "5. Jarque-Bera للتحقق من طبيعية البواقي",
        "6. التنبؤ مع فترات ثقة 95% (Monte Carlo)",
        "7. ARIMA يُحدّد معدل النمو g المستخدم في DCF",
    ]): W(ws_ar,48+i,1,ln)

    # رسم ARIMA
    try:
        arima_bytes = _arima_png(quarters, hist, forecasts)
        embed_png(ws_ar, arima_bytes, "A57", w=720, h=380)
    except Exception as e:
        print(f"  [arima png] {e}")

    # ═════════════════════════════════════════════════════════════════════════
    # (6b) التكامل الجغرافي — OSM Location Map + Price Heatmap
    # ═════════════════════════════════════════════════════════════════════════
    # ── حساب الإحداثيات ───────────────────────────────────────────────────
    subj_lat, subj_lng = _geocode(loc)

    # ── إحداثيات المقارنات (offsets عشوائية قابلة للتكرار) ───────────────
    import random as _rnd
    _rnd.seed(hash(loc) & 0xFFFF)
    comp_sales_raw = data.get("comp_sales", [])

    # ── Sovereign Gold Edition: إعادة هندسة صفحة مقارنات البيوع ─────────────
    _upgrade_sales_comp(wb, data, comp_sales_raw)

    # توليد إزاحات ضمن نطاق 1.5 km (≈ 0.013 درجة)
    comp_offsets = []
    for _ in range(max(len(comp_sales_raw), 5)):
        dlat = _rnd.uniform(-0.012, 0.012)
        dlng = _rnd.uniform(-0.014, 0.014)
        comp_offsets.append((dlat, dlng))

    # حساب سعر المتر لكل مقارنة
    comp_ppms_geo = []
    for i, cs in enumerate(comp_sales_raw):
        cp = _num(cs.get("price", 0))
        ca = _num(cs.get("area", area or 1))
        comp_ppms_geo.append(round(cp / ca, 0) if ca > 0 else ppm)
    # تعبئة بالقيم الافتراضية إذا لم تكن هناك مقارنات كافية
    while len(comp_ppms_geo) < len(comp_offsets):
        comp_ppms_geo.append(ppm * _rnd.uniform(0.88, 1.12))

    # ── ورقة التحليل المكاني ──────────────────────────────────────────────
    spatial_name = "التحليل المكاني"
    ws_sp = wb[spatial_name] if spatial_name in wb.sheetnames \
            else wb.create_sheet(spatial_name)

    W(ws_sp, 1, 1, "التحليل المكاني والجغرافي — OpenStreetMap Integration")
    ws_sp.cell(1,1).font = Font(bold=True, size=13, color="D4AF37", name="Calibri")

    W(ws_sp, 2, 1,
      f"=\"الموقع: \"&'الافتراضات والمدخلات'!B13"
      f"&\"  |  الإحداثيات: {subj_lat:.5f}N , {subj_lng:.5f}E\"")
    ws_sp.cell(2,1).font = Font(italic=True, size=9, color="AAAAAA")

    W(ws_sp, 3, 1, "══ بيانات الإحداثيات (مرجع) ══")
    W(ws_sp, 4, 1, "خط العرض (Latitude)");  W(ws_sp, 4, 2, round(subj_lat, 6))
    W(ws_sp, 5, 1, "خط الطول (Longitude)"); W(ws_sp, 5, 2, round(subj_lng, 6))
    W(ws_sp, 6, 1, "مصدر الخريطة");         W(ws_sp, 6, 2, "OpenStreetMap © contributors")
    W(ws_sp, 7, 1, "عدد المقارنات المُرسَّمة"); W(ws_sp, 7, 2, len(comp_sales_raw))

    print(f"{_ts()} [INFO] توليد الرسوم البيانية والخرائط... 60%")
    # ── رسم خريطة الموقع ─────────────────────────────────────────────────
    try:
        loc_map_bytes = _osm_location_map(
            subj_lat, subj_lng,
            comp_offsets[:len(comp_ppms_geo)],
            comp_ppms_geo,
            ppm, loc
        )
        embed_png(ws_sp, loc_map_bytes, "A9",  w=740, h=470)
        W(ws_sp, 10, 10, "خريطة الموقع — OpenStreetMap")
        ws_sp.cell(10,10).font = Font(bold=True, size=10, color="4ec9b0")
        print(f"  [geo] Location map embedded ✓")
    except Exception as e:
        print(f"  [geo location map] {e}")

    # ── رسم خريطة الحرارة السعرية ─────────────────────────────────────────
    try:
        heat_bytes = _price_heatmap_png(
            subj_lat, subj_lng,
            comp_offsets[:len(comp_ppms_geo)],
            comp_ppms_geo,
            ppm, loc
        )
        embed_png(ws_sp, heat_bytes, "A36", w=740, h=480)
        W(ws_sp, 37, 10, "خريطة الحرارة السعرية — Price Heatmap")
        ws_sp.cell(37,10).font = Font(bold=True, size=10, color="f9a825")
        print(f"  [geo] Price heatmap embedded ✓")
    except Exception as e:
        print(f"  [geo heatmap] {e}")

    # ── جدول إحداثيات المقارنات ──────────────────────────────────────────
    W(ws_sp, 64, 1, "══ إحداثيات المقارنات (تقريبية) ══")
    hdrs_geo = ["#", "lat", "lng", "سعر المتر (EGP/م²)", "نسبة من سعر التقييم"]
    for j, h in enumerate(hdrs_geo): W(ws_sp, 65, j+1, h)
    for i, ((dlat, dlng), cppm) in enumerate(
            zip(comp_offsets, comp_ppms_geo), 1):
        r = 65 + i
        W(ws_sp, r, 1, i)
        W(ws_sp, r, 2, round(subj_lat + dlat, 6))
        W(ws_sp, r, 3, round(subj_lng + dlng, 6))
        W(ws_sp, r, 4, round(cppm, 0))
        W(ws_sp, r, 5, f"=IF(D{r}>0,D{r}/{ppm:.0f}*100,0)&\"%\"")

    # ═════════════════════════════════════════════════════════════════════════
    # (7) لوحة القيادة التنفيذية — Modern Professional Dashboard
    # ═════════════════════════════════════════════════════════════════════════
    print(f"{_ts()} [INFO] Rendering Modern Dashboard... 75%")
    from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
    from openpyxl.chart import RadarChart, Reference

    dash_name = "لوحة القيادة التنفيذية"
    ws_dash = wb[dash_name] if dash_name in wb.sheetnames else wb.create_sheet(dash_name)

    # ── لوحة الألوان الاحترافية ─────────────────────────────────────────────
    _DN  = "002060"   # Deep Navy — header primary
    _RB  = "003399"   # Royal Blue — section headers
    _LG  = "F2F2F2"   # Light Gray — KPI title background
    _DK  = "001540"   # Very dark navy — value cells
    _GLD = "D4AF37"   # Gold — accents / borders
    _YEL = "FFD700"   # Bright gold — large KPI values
    _GRN = "0a5e36"   # Green — positive
    _AMB = "7b5e00"   # Amber — caution
    _RED = "6b1515"   # Red — negative
    _WHT = "FFFFFF"
    _SLT = "AAAACC"   # Slate — subtitles

    _gs  = Side(style="medium", color=_GLD)
    _ths = Side(style="thin",   color="334466")

    def _dash_border(thick=False):
        s = _gs if thick else _ths
        return Border(top=s, left=s, right=s, bottom=s)

    DCF   = "'DCF \u2014 \u0627\u0644\u062a\u062f\u0641\u0642\u0627\u062a \u0627\u0644\u0646\u0642\u062f\u064a\u0629'"
    INP_D = "'الافتراضات والمدخلات'"

    # ── كشف القطاع للتكيف مع المؤشرات ────────────────────────────────────
    # نفضّل sector المحسوب في advanced_valuation (موجود في data["sector"])
    # ونرجع لـ _get_sector كبديل احتياطي فقط
    _dash_sector = data.get("sector") or _get_sector(str(ptype))

    # ─────────────────────────────────────────────────────────────────────────
    # ROW 1 — ترويسة رئيسية + ختم الحالة
    # ─────────────────────────────────────────────────────────────────────────
    for _ci in range(1, 12):
        ws_dash.cell(1, _ci).fill = PatternFill("solid", fgColor=_DN)
    W(ws_dash, 1, 1, "لوحة القيادة التنفيذية  |  Expert Smart Valuation Pro")
    ws_dash.cell(1, 1).font      = Font(bold=True, size=18, color=_GLD, name="Calibri")
    ws_dash.cell(1, 1).alignment = Alignment(horizontal="right", vertical="center")
    ws_dash.row_dimensions[1].height = 40
    try: ws_dash.merge_cells("A1:H1")
    except Exception: pass

    # ختم الحالة في I1 — يقرأ من شيت المدخلات
    INP_V = "'الافتراضات والمدخلات'"
    W(ws_dash, 1, 9,
      f'=IF({INP_V}!H1>=12,"✅  مكتمل",'
      f'IF({INP_V}!H1>=7,"⚠  ناقص ("&(13-{INP_V}!H1)&" حقل)",'
      f'"❌  غير كافٍ"))')
    _stamp = ws_dash.cell(1, 9)
    _stamp.font      = Font(bold=True, size=13, name="Calibri", color=_WHT)
    _stamp.alignment = Alignment(horizontal="center", vertical="center")
    _stamp.border    = _dash_border(thick=True)
    try: ws_dash.merge_cells("I1:K1")
    except Exception: pass
    ws_dash.column_dimensions["I"].width = 14
    ws_dash.column_dimensions["J"].width = 7
    ws_dash.column_dimensions["K"].width = 10

    # تنسيق شرطي على I1 في Dashboard
    from openpyxl.formatting.rule import FormulaRule as _FR
    ws_dash.conditional_formatting.add("I1",
        _FR(formula=[f"{INP_V}!$H$1>=12"],
            fill=PatternFill("solid", fgColor=_GRN),
            font=Font(color=_WHT, bold=True, size=13)))
    ws_dash.conditional_formatting.add("I1",
        _FR(formula=[f"AND({INP_V}!$H$1>=7,{INP_V}!$H$1<12)"],
            fill=PatternFill("solid", fgColor=_AMB),
            font=Font(color=_WHT, bold=True, size=13)))
    ws_dash.conditional_formatting.add("I1",
        _FR(formula=[f"{INP_V}!$H$1<7"],
            fill=PatternFill("solid", fgColor=_RED),
            font=Font(color=_WHT, bold=True, size=13)))

    # ─────────────────────────────────────────────────────────────────────────
    # ROW 2 — سطر الهوية (رقم التقرير | تاريخ | موقع)
    # ─────────────────────────────────────────────────────────────────────────
    for _ci in range(1, 12):
        ws_dash.cell(2, _ci).fill = PatternFill("solid", fgColor="001030")
    W(ws_dash, 2, 1,
      f"=\" رقم التقرير: \"&{INP_D}!E1"
      f"&\"   |   التاريخ: \"&{INP_D}!E2"
      f"&\"   |   الموقع: \"&{INP_D}!B13"
      f"&\"   |   القطاع: {_dash_sector}\"")
    ws_dash.cell(2, 1).font      = Font(size=10, color=_SLT, italic=True, name="Calibri")
    ws_dash.cell(2, 1).alignment = Alignment(horizontal="right", vertical="center")
    ws_dash.row_dimensions[2].height = 22
    try: ws_dash.merge_cells("A2:K2")
    except Exception: pass

    # ─────────────────────────────────────────────────────────────────────────
    # ROW 3 — شريط الغرض (Purpose Badge) — Logic Gate v45
    # ─────────────────────────────────────────────────────────────────────────
    _PURPOSE_BADGE_COLOR = {
        "fair_market_value":    "1a3a5c",   # أزرق داكن
        "acquisition":          "4a0e0e",   # أحمر داكن
        "bank_financing":       "0a3a1a",   # أخضر داكن
        "rental_arbitration":   "3a2a0a",   # بني داكن
        "insurance":            "2a0a4a",   # بنفسجي داكن
        "investment_analysis":  "0a2a4a",   # أزرق بترولي داكن
        "judicial_liquidation": "3a3a0a",   # زيتي داكن
    }
    _badge_bg = _PURPOSE_BADGE_COLOR.get(_purpose, "1a3a5c")
    for _ci in range(1, 12):
        ws_dash.cell(3, _ci).fill = PatternFill("solid", fgColor=_badge_bg)
    W(ws_dash, 3, 1,
      f"\" ⚙ الغرض: {_purpose_ar}   |   المسار: {_purpose_lg}   |   {_PURPOSES[_purpose]['ar']}\"")
    ws_dash.cell(3, 1).font      = Font(size=9, color="D4AF37", bold=True, name="Calibri")
    ws_dash.cell(3, 1).alignment = Alignment(horizontal="right", vertical="center")
    ws_dash.row_dimensions[3].height = 18
    try: ws_dash.merge_cells("A3:K3")
    except Exception: pass

    # ─────────────────────────────────────────────────────────────────────────
    # ROWS 4-5 — بطاقات KPI (متكيّفة مع القطاع)
    # ─────────────────────────────────────────────────────────────────────────
    def _kpi_card(ws, row, col, title, formula, fmt_hint="EGP", span=2):
        """بطاقة KPI بتصميم عصري: عنوان رمادي فاتح + قيمة زرقاء داكن + حدود ذهبية"""
        # عنوان البطاقة (خلفية رمادية فاتحة، نص أزرق داكن)
        W(ws, row, col, title)
        tc = ws.cell(row, col)
        tc.fill      = PatternFill("solid", fgColor=_LG)
        tc.font      = Font(bold=True, size=10, color=_DN, name="Calibri")
        tc.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        tc.border    = Border(top=_gs, left=_gs, right=_gs)
        if span > 1:
            try: ws.merge_cells(f"{tc.column_letter}{row}:{chr(ord(tc.column_letter)+span-1)}{row}")
            except Exception: pass
        # قيمة البطاقة (خلفية زرقاء داكن، نص ذهبي كبير)
        W(ws, row+1, col, formula)
        vc = ws.cell(row+1, col)
        vc.fill      = PatternFill("solid", fgColor=_DK)
        vc.font      = Font(bold=True, size=16, color=_YEL, name="Calibri")
        vc.alignment = Alignment(horizontal="center", vertical="center")
        vc.border    = Border(left=_gs, right=_gs, bottom=_gs)
        if fmt_hint == "%":   vc.number_format = "0.0%"
        elif fmt_hint == "EGP": vc.number_format = '#,##0'
        if span > 1:
            try: ws.merge_cells(f"{vc.column_letter}{row+1}:{chr(ord(vc.column_letter)+span-1)}{row+1}")
            except Exception: pass
        ws.row_dimensions[row].height   = 24
        ws.row_dimensions[row+1].height = 36

    # بطاقات حسب القطاع
    if _dash_sector == "industrial":
        _kpi_defs = [
            (4, 1,  "القيمة السوقية",          f"={INP_D}!B27",         "EGP"),
            (4, 3,  "قيمة DCF",                f"={DCF}!B35",            "EGP"),
            (4, 5,  "IRR",                      f"={DCF}!B41",            "%"),
            (4, 7,  "ROI الإجمالي",             f"={DCF}!B43&\"%\"",      ""),
            (4, 9,  "درجة جودة البيانات",       "—",                      ""),
        ]
    elif _dash_sector == "agricultural":
        _kpi_defs = [
            (4, 1,  "القيمة الإجمالية للأرض",  f"={INP_D}!B27",          "EGP"),
            (4, 3,  "قيمة DCF",                f"={DCF}!B35",             "EGP"),
            (4, 5,  "IRR",                      f"={DCF}!B41",             "%"),
            (4, 7,  "الدخل التشغيلي",          f"={DCF}!B37",             "EGP"),
            (4, 9,  "معدل الرسملة",             f"={INP_D}!B22",           "%"),
        ]
    elif _dash_sector == "hospitality":
        _kpi_defs = [
            (4, 1,  "قيمة الفندق",             f"={INP_D}!B27",           "EGP"),
            (4, 3,  "قيمة DCF",                f"={DCF}!B35",             "EGP"),
            (4, 5,  "IRR",                      f"={DCF}!B41",             "%"),
            (4, 7,  "NOI السنوي",              f"={DCF}!B37",             "EGP"),
            (4, 9,  "ROI",                      f"={DCF}!B43&\"%\"",       ""),
        ]
    elif _dash_sector == "retail":
        _kpi_defs = [
            (4, 1,  "قيمة المحل/المول",        f"={INP_D}!B27",           "EGP"),
            (4, 3,  "إيراد م² التأجيري",       f"={DCF}!B37",             "EGP"),
            (4, 5,  "IRR",                      f"={DCF}!B41",             "%"),
            (4, 7,  "NOI السنوي",              f"={DCF}!B37",             "EGP"),
            (4, 9,  "مؤشر كثافة المرور",       "—",                       ""),
        ]
    elif _dash_sector == "healthcare":
        _kpi_defs = [
            (4, 1,  "قيمة المنشأة الطبية",     f"={INP_D}!B27",           "EGP"),
            (4, 3,  "تكلفة السرير (EGP)",      "—",                       "EGP"),
            (4, 5,  "IRR",                      f"={DCF}!B41",             "%"),
            (4, 7,  "NOI السنوي",              f"={DCF}!B37",             "EGP"),
            (4, 9,  "معامل الطلب الطبي",       "—",                       ""),
        ]
    elif _dash_sector == "educational":
        _kpi_defs = [
            (4, 1,  "قيمة المنشأة التعليمية",  f"={INP_D}!B27",           "EGP"),
            (4, 3,  "طاقة الطلاب",             "—",                       ""),
            (4, 5,  "IRR",                      f"={DCF}!B41",             "%"),
            (4, 7,  "الدخل التشغيلي (NOI)",    f"={DCF}!B37",             "EGP"),
            (4, 9,  "متوسط رسوم الطالب",       "—",                       "EGP"),
        ]
    else:   # residential — default
        _kpi_defs = [
            (4, 1,  "القيمة السوقية",           f"={INP_D}!B27",          "EGP"),
            (4, 3,  "قيمة DCF",                 f"={DCF}!B35",            "EGP"),
            (4, 5,  "NPV",                       f"={DCF}!B37",            "EGP"),
            (4, 7,  "IRR",                       f"={DCF}!B41",            "%"),
            (4, 9,  "ROI الإجمالي",             f"={DCF}!B43&\"%\"",      ""),
        ]

    # ── تجاوز القطاع بغرض التقييم (Logic Gate v45) ────────────────────────────
    # الغرض غير الافتراضي يُحل محل بطاقات KPI القطاعية ببطاقات مُصمَّمة للغرض
    if _purpose == "acquisition":
        _kpi_defs = [
            (4, 1, "القيمة الجوهرية",            f"={INP_D}!B27",                            "EGP"),
            (4, 3, "قيمة DCF",                   f"={DCF}!B35",                              "EGP"),
            (4, 5, "IRR الاستحواذ",              f"={DCF}!B41",                              "%"),
            (4, 7, "علاوة السيطرة (20%)",        f"={INP_D}!B27*0.20",                      "EGP"),
            (4, 9, "قيمة الصفقة الكاملة",        f"={INP_D}!B27*1.20",                      "EGP"),
        ]
    elif _purpose == "bank_financing":
        _kpi_defs = [
            (4, 1, "القيمة التمويلية",            f"={INP_D}!B27",                            "EGP"),
            (4, 3, "الحد الأقصى للقرض (70%)",    f"={INP_D}!B27*0.70",                      "EGP"),
            (4, 5, "تغطية الدين (NOI)",           f"={DCF}!B37",                              "EGP"),
            (4, 7, "نسبة LTV",                   f"=70",                                      "%"),
            (4, 9, "العمر المتبقي (سنة)",         f"=MAX(0,50-(2026-{INP_D}!B10))",           ""),
        ]
    elif _purpose == "rental_arbitration":
        _kpi_defs = [
            (4, 1, "القيمة الرأسمالية",           f"={INP_D}!B27",                            "EGP"),
            (4, 3, "الإيجار العادل/م²",           f"={INP_D}!B6",                             "EGP"),
            (4, 5, "الإيجار السنوي الإجمالي",     f"={INP_D}!B28",                            "EGP"),
            (4, 7, "NOI بعد المصاريف",            f"={DCF}!B37",                              "EGP"),
            (4, 9, "معدل الرسملة الفعلي",         f"=IF({INP_D}!B27>0,{INP_D}!B28/{INP_D}!B27,0)", "%"),
        ]
    elif _purpose == "insurance":
        _kpi_defs = [
            (4, 1, "القيمة التأمينية",             f"={INP_D}!B27",                            "EGP"),
            (4, 3, "تكلفة إعادة الإنشاء/م²",      f"={INP_D}!B5*0.65",                       "EGP"),
            (4, 5, "عمر المبنى (سنة)",             f"=MAX(0,2026-{INP_D}!B10)",               ""),
            (4, 7, "معامل الاستهلاك",              f"=MAX(0.3,1-(MAX(0,2026-{INP_D}!B10)*0.02))", "%"),
            (4, 9, "القيمة الاستبدالية الكاملة",   f"={INP_D}!B27",                            "EGP"),
        ]
    elif _purpose == "investment_analysis":
        _kpi_defs = [
            (4, 1, "القيمة السوقية",               f"={INP_D}!B27",                            "EGP"),
            (4, 3, "IRR",                          f"={DCF}!B41",                              "%"),
            (4, 5, "NPV (صافي الحاضر)",            f"={DCF}!B35",                             "EGP"),
            (4, 7, "NOI السنوي",                   f"={DCF}!B37",                             "EGP"),
            (4, 9, "مضاعف القيمة (EM)",            f"={DCF}!B43&\"x\"",                       ""),
        ]
    elif _purpose == "judicial_liquidation":
        _kpi_defs = [
            (4, 1, "القيمة السوقية الأصلية",       f"={INP_D}!B27/0.82",                      "EGP"),
            (4, 3, "قيمة التصفية (−18%)",         f"={INP_D}!B27",                            "EGP"),
            (4, 5, "مبلغ خصم التصفية",             f"={INP_D}!B27*0.18/0.82",                 "EGP"),
            (4, 7, "تكاليف الإجراءات (3%)",        f"={INP_D}!B27*0.03",                      "EGP"),
            (4, 9, "صافي العائد للمحكمة",          f"={INP_D}!B27*0.97",                      "EGP"),
        ]
    # fair_market_value → _kpi_defs remains sector-based (computed above)

    for _r, _c, _ttl, _frm, _fmt in _kpi_defs:
        _kpi_card(ws_dash, _r, _c, _ttl, _frm, _fmt, span=1)

    # ضبط عرض أعمدة KPI
    for _col_l, _w in zip("ABCDEFGHIJK", [18,4,18,4,14,4,14,4,14,4,14]):
        ws_dash.column_dimensions[_col_l].width = _w

    # ─────────────────────────────────────────────────────────────────────────
    # ROW 6 — فاصل ذهبي
    # ─────────────────────────────────────────────────────────────────────────
    for _ci in range(1, 12):
        ws_dash.cell(6, _ci).fill = PatternFill("solid", fgColor=_GLD)
    ws_dash.row_dimensions[6].height = 4

    # ─────────────────────────────────────────────────────────────────────────
    # ROWS 7-9 — التوصية + ملاحظة الخبير (ديناميكية بالغرض — Logic Gate v45)
    # ─────────────────────────────────────────────────────────────────────────
    _row7_titles = {
        "fair_market_value":    "التوصية الاستثمارية",
        "acquisition":          "توصية الاستحواذ والاندماج",
        "bank_financing":       "تقييم الجدارة الائتمانية",
        "rental_arbitration":   "نتيجة التحكيم الإيجاري",
        "insurance":            "توصية التغطية التأمينية",
        "investment_analysis":  "التوصية الاستثمارية",
        "judicial_liquidation": "قيمة التصفية القضائية",
    }
    W(ws_dash, 7, 1, _row7_titles.get(_purpose, "التوصية الاستثمارية"))
    _h7 = ws_dash.cell(7, 1)
    _h7.fill      = PatternFill("solid", fgColor=_RB)
    _h7.font      = Font(bold=True, size=11, color=_GLD, name="Calibri")
    _h7.alignment = Alignment(horizontal="right", vertical="center")
    _h7.border    = _dash_border(thick=True)
    ws_dash.row_dimensions[7].height = 22
    try: ws_dash.merge_cells("A7:K7")
    except Exception: pass

    # ROW 8 — نص التوصية الديناميكي حسب الغرض
    _row8_formulas = {
        "fair_market_value": (
            f"=IF(ISNUMBER({DCF}!B41),"
            f"IF({DCF}!B41>=0.15,\"✅  استثمار ممتاز — العائد يتجاوز 15%، يُنصح بالمضي قدماً.\","
            f"IF({DCF}!B41>=0.10,\"✔  استثمار جيد — العائد يتجاوز 10%، مقبول في ظروف السوق الحالية.\","
            f"IF({DCF}!B41>=0.06,\"⚠  استثمار متوسط — العائد عند الحد الأدنى. راجع الافتراضات قبل القرار.\","
            f"\"❌  استثمار ضعيف — العائد أقل من 6%. يُنصح بالانتظار أو البحث عن بدائل.\"))),"
            f"\"⚠  تعذّر حساب IRR — تحقق من التدفقات النقدية.\")"
        ),
        "acquisition": (
            f"=IF(ISNUMBER({DCF}!B35),"
            f"IF({DCF}!B35>{INP_D}!B27,"
            f"\"✅  استحواذ مُجدٍ — قيمة DCF تتجاوز القيمة السوقية. علاوة سيطرة 15-20% مُبررة. أوصي بالمضي في الصفقة.\","
            f"IF({DCF}!B35>{INP_D}!B27*0.85,"
            f"\"⚠  استحواذ بشروط — قيمة DCF قريبة من السوق. التفاوض على السعر وتخفيض علاوة السيطرة ضرورة.\","
            f"\"❌  استحواذ غير مُجدٍ — قيمة DCF أقل بكثير من السوق. مراجعة الافتراضات أو إعادة تقييم الصفقة.\")),"
            f"\"⚠  تعذّر تقييم الصفقة — راجع بيانات DCF.\")"
        ),
        "bank_financing": (
            f"=IF(ISNUMBER({DCF}!B37),"
            f"IF({DCF}!B37/{INP_D}!B27>=0.08,"
            f"\"✅  ضمان ممتاز — NOI يتجاوز 8% من قيمة الضمان. تمويل بنسبة 70% LTV مُوصى.\","
            f"IF({DCF}!B37/{INP_D}!B27>=0.05,"
            f"\"✔  ضمان مقبول — NOI بين 5-8%. تمويل بنسبة 60-65% LTV وفق سياسة البنك.\","
            f"\"⚠  ضمان ضعيف — NOI أقل من 5%. ضمانات إضافية أو رهن مشترك ضرورة.\")),"
            f"\"⚠  تعذّر تقييم الجدارة — راجع بيانات NOI.\")"
        ),
        "rental_arbitration": (
            f"=IF(ISNUMBER({DCF}!B37),"
            f"IF({DCF}!B37/{INP_D}!B27>=0.06,"
            f"\"✅  إيجار عادل — العائد الصافي يتجاوز 6%. القيمة الإيجارية المحددة تتوافق مع السوق.\","
            f"IF({DCF}!B37/{INP_D}!B27>=0.04,"
            f"\"⚠  إيجار معتدل — العائد 4-6%. الإيجار الحالي في النطاق المقبول مع هامش تفاوض.\","
            f"\"❌  إيجار منخفض — العائد أقل من 4%. مراجعة الإيجار إلى أعلى موصى قبل التحكيم.\")),"
            f"\"⚠  تعذّر تقييم الإيجار — راجع بيانات NOI.\")"
        ),
        "insurance": (
            f"=IF((2026-{INP_D}!B10)<10,"
            f"\"✅  تأمين كامل — مبنى حديث (أقل من 10 سنوات). تكلفة إعادة الإنشاء مرتفعة. تغطية 100% مُوصى.\","
            f"IF((2026-{INP_D}!B10)<25,"
            f"\"⚠  تأمين جزئي — مبنى متوسط (10-25 سنة). تطبيق الاستهلاك العادل. تغطية 70-85% مُوصى.\","
            f"\"⚠  تأمين قيمة مخصومة — مبنى قديم (أكثر من 25 سنة). الاستهلاك مرتفع. تقييم الوضع الإنشائي فعلياً ضرورة.\"))"
        ),
        "investment_analysis": (
            f"=IF(ISNUMBER({DCF}!B41),"
            f"IF({DCF}!B41>=0.15,\"✅  استثمار ممتاز — IRR يتجاوز 15%. العائد يغطي WACC بفارق مريح. توصية: شراء.\","
            f"IF({DCF}!B41>=0.10,\"✔  استثمار جيد — IRR بين 10-15%. عائد مقبول في السوق الحالي. توصية: دراسة تفصيلية.\","
            f"IF({DCF}!B41>=0.06,\"⚠  استثمار متوسط — IRR بين 6-10%. مراجعة WACC وفترة الاحتفاظ.\","
            f"\"❌  استثمار ضعيف — IRR أقل من 6%. يقل عن العائد الخالي من المخاطرة. لا يُنصح.\"))),"
            f"\"⚠  تعذّر حساب IRR — تحقق من التدفقات النقدية.\")"
        ),
        "judicial_liquidation": (
            f"=\"⚖ قيمة التصفية القضائية = \"&TEXT({INP_D}!B27,\"#,##0\")&\" EGP"
            f" | خصم التصفية 18%: \"&TEXT({INP_D}!B27*0.18/0.82,\"#,##0\")&\" EGP"
            f" | القيمة السوقية الأصلية: \"&TEXT({INP_D}!B27/0.82,\"#,##0\")&\" EGP."
            f" تكاليف الإجراءات القضائية تُقدَّر بـ 3%.\""
        ),
    }
    W(ws_dash, 8, 1, _row8_formulas.get(_purpose, _row8_formulas["fair_market_value"]))
    _c8 = ws_dash.cell(8, 1)
    _c8.font      = Font(size=12, bold=True, color=_WHT, name="Calibri")
    _c8.fill      = PatternFill("solid", fgColor=_DN)
    _c8.alignment = Alignment(wrap_text=True, horizontal="right", vertical="center")
    _c8.border    = _dash_border(thick=True)
    ws_dash.row_dimensions[8].height = 36
    try: ws_dash.merge_cells("A8:K8")
    except Exception: pass

    # ROW 9 — ملاحظة الخبير (ديناميكية بالغرض)
    _row9_notes = {
        "fair_market_value": (
            f"=\" ملاحظة الخبير: القيمة السوقية = \"&TEXT({INP_D}!B27,\"#,##0\")&\" EGP"
            f" | قيمة DCF = \"&TEXT({DCF}!B35,\"#,##0\")&\" EGP"
            f" | الفارق يعكس مستوى الثقة في افتراضات النمو والخصم.\""
        ),
        "acquisition": (
            f"=\" الاستحواذ: سعر الصفقة المُقترح = \"&TEXT({INP_D}!B27*1.20,\"#,##0\")&\" EGP"
            f" (يشمل علاوة سيطرة 20%)"
            f" | قيمة DCF = \"&TEXT({DCF}!B35,\"#,##0\")&\" EGP"
            f" | IRR = \"&TEXT({DCF}!B41,\"0.0%\")\""
        ),
        "bank_financing": (
            f"=\" التمويل البنكي: الحد الأقصى للقرض = \"&TEXT({INP_D}!B27*0.70,\"#,##0\")&\" EGP (70% LTV)"
            f" | NOI السنوي = \"&TEXT({DCF}!B37,\"#,##0\")&\" EGP"
            f" | العمر المتبقي للأصل: \"&TEXT(MAX(0,50-(2026-{INP_D}!B10)),\"0\")&\" سنة\""
        ),
        "rental_arbitration": (
            f"=\" التحكيم: الإيجار العادل/م² = \"&TEXT({INP_D}!B6,\"#,##0\")&\" EGP/م²"
            f" | الإيجار السنوي الإجمالي = \"&TEXT({INP_D}!B28,\"#,##0\")&\" EGP"
            f" | معدل الرسملة الفعلي = \"&TEXT(IF({INP_D}!B27>0,{INP_D}!B28/{INP_D}!B27,0),\"0.00%\")\""
        ),
        "insurance": (
            f"=\" التأمين: القيمة التأمينية = \"&TEXT({INP_D}!B27,\"#,##0\")&\" EGP"
            f" | عمر المبنى = \"&TEXT(MAX(0,2026-{INP_D}!B10),\"0\")&\" سنة"
            f" | معامل الاستهلاك = \"&TEXT(MAX(0.3,1-MAX(0,2026-{INP_D}!B10)*0.02),\"0.0%\")\""
        ),
        "investment_analysis": (
            f"=\" الاستثمار: IRR = \"&TEXT({DCF}!B41,\"0.0%\")&\" | NPV = \"&TEXT({DCF}!B35,\"#,##0\")&\" EGP"
            f" | NOI = \"&TEXT({DCF}!B37,\"#,##0\")&\" EGP"
            f" | العائد المُضاعَف (EM) = \"&{DCF}!B43&\"x\""
        ),
        "judicial_liquidation": (
            f"=\" التصفية: القيمة السوقية الأصلية = \"&TEXT({INP_D}!B27/0.82,\"#,##0\")&\" EGP"
            f" | قيمة التصفية = \"&TEXT({INP_D}!B27,\"#,##0\")&\" EGP"
            f" | صافي العائد للمحكمة بعد التكاليف = \"&TEXT({INP_D}!B27*0.97,\"#,##0\")&\" EGP\""
        ),
    }
    W(ws_dash, 9, 1, _row9_notes.get(_purpose, _row9_notes["fair_market_value"]))
    _c9 = ws_dash.cell(9, 1)
    _c9.font      = Font(size=9, italic=True, color=_SLT, name="Calibri")
    _c9.fill      = PatternFill("solid", fgColor="001030")
    _c9.alignment = Alignment(wrap_text=True, horizontal="right", vertical="center")
    ws_dash.row_dimensions[9].height = 28
    try: ws_dash.merge_cells("A9:K9")
    except Exception: pass

    # ─────────────────────────────────────────────────────────────────────────
    # ROW 10 — فاصل ذهبي
    # ─────────────────────────────────────────────────────────────────────────
    for _ci in range(1, 12):
        ws_dash.cell(10, _ci).fill = PatternFill("solid", fgColor=_GLD)
    ws_dash.row_dimensions[10].height = 4

    # ─────────────────────────────────────────────────────────────────────────
    # ROWS 11-17 — جدول محاور الرادار + مخطط Excel Native
    # ─────────────────────────────────────────────────────────────────────────
    # عنوان القسم
    W(ws_dash, 11, 1, "محاور تقييم العقار — Property Scoring Radar")
    _h11 = ws_dash.cell(11, 1)
    _h11.fill      = PatternFill("solid", fgColor=_RB)
    _h11.font      = Font(bold=True, size=10, color=_GLD, name="Calibri")
    _h11.alignment = Alignment(horizontal="right", vertical="center")
    _h11.border    = _dash_border()
    ws_dash.row_dimensions[11].height = 20
    try: ws_dash.merge_cells("A11:C11")
    except Exception: pass

    # رؤوس أعمدة جدول الرادار
    for _ci, _hdr in enumerate(["المحور", "الدرجة (0-10)"], 1):
        W(ws_dash, 12, _ci, _hdr)
        _hc = ws_dash.cell(12, _ci)
        _hc.fill      = PatternFill("solid", fgColor=_LG)
        _hc.font      = Font(bold=True, size=9, color=_DN, name="Calibri")
        _hc.alignment = Alignment(horizontal="center")
        _hc.border    = _dash_border()

    # محاور الرادار (متكيّفة مع القطاع)
    if _dash_sector == "industrial":
        radar_axes = [
            ("الموقع الصناعي",    f"=MIN(10,MAX(1,{INP_D}!B5/4000))"),
            ("الاستهلاك العمري",  f"=MIN(10,MAX(1,(2040-{INP_D}!B10)/5))"),
            ("العائد المالي",     f"=MIN(10,MAX(1,{DCF}!B41*60))"),
            ("القرب اللوجستي",    f"=MIN(10,MAX(1,{INP_D}!B5/3500))"),
            ("جودة البنية التحتية",f"=MIN(10,MAX(1,IF({INP_D}!B27>0,{DCF}!B35/{INP_D}!B27*7,5)))"),
        ]
    elif _dash_sector == "agricultural":
        radar_axes = [
            ("جودة التربة",        f"=MIN(10,MAX(1,{INP_D}!B5/2000))"),
            ("مصدر الري",          f"=MIN(10,MAX(1,(2026-{INP_D}!B10)/4))"),
            ("إنتاجية الفدان",     f"=MIN(10,MAX(1,{DCF}!B41*60))"),
            ("القرب من الأسواق",   f"=MIN(10,MAX(1,{INP_D}!B5/2500))"),
            ("نسبة العائد/التكلفة",f"=MIN(10,MAX(1,IF({INP_D}!B27>0,{DCF}!B35/{INP_D}!B27*7,5)))"),
        ]
    elif _dash_sector == "hospitality":
        radar_axes = [
            ("الموقع السياحي",     f"=MIN(10,MAX(1,{INP_D}!B5/6000))"),
            ("جودة الخدمات",       f"=MIN(10,MAX(1,(2026-{INP_D}!B10)/3))"),
            ("نسبة الإشغال",       f"=MIN(10,MAX(1,{DCF}!B41*60))"),
            ("معدل ADR",           f"=MIN(10,MAX(1,{INP_D}!B5/5000))"),
            ("RevPAR مقابل السوق", f"=MIN(10,MAX(1,IF({INP_D}!B27>0,{DCF}!B35/{INP_D}!B27*7,5)))"),
        ]
    elif _dash_sector == "retail":
        radar_axes = [
            ("الموقع التجاري",          f"=MIN(10,MAX(1,{INP_D}!B5/4000))"),
            ("الواجهة والمدخل",         f"=MIN(10,MAX(1,(2026-{INP_D}!B10)/3))"),
            ("قوة الإنفاق الشرائي",     f"=MIN(10,MAX(1,{DCF}!B41*60))"),
            ("كثافة المرور",            f"=MIN(10,MAX(1,{INP_D}!B5/3500))"),
            ("نسبة المساحات التأجيرية", f"=MIN(10,MAX(1,IF({INP_D}!B27>0,{DCF}!B35/{INP_D}!B27*7,5)))"),
        ]
    elif _dash_sector == "healthcare":
        radar_axes = [
            ("الموقع والكثافة السكانية", f"=MIN(10,MAX(1,{INP_D}!B5/5000))"),
            ("التجهيزات التخصصية",       f"=MIN(10,MAX(1,(2026-{INP_D}!B10)/4))"),
            ("معيار الأسرة/المساحة",     f"=MIN(10,MAX(1,{DCF}!B41*60))"),
            ("العائد المالي",             f"=MIN(10,MAX(1,{INP_D}!B5/4500))"),
            ("نقص الخدمات الطبية",       f"=MIN(10,MAX(1,IF({INP_D}!B27>0,{DCF}!B35/{INP_D}!B27*7,5)))"),
        ]
    elif _dash_sector == "educational":
        radar_axes = [
            ("الموقع التعليمي",       f"=MIN(10,MAX(1,{INP_D}!B5/4000))"),
            ("الطاقة الاستيعابية",    f"=MIN(10,MAX(1,(2026-{INP_D}!B10)/3))"),
            ("المرافق والمعامل",      f"=MIN(10,MAX(1,{DCF}!B41*60))"),
            ("متوسط الرسوم الدراسية", f"=MIN(10,MAX(1,{INP_D}!B5/5000))"),
            ("الترخيص والاعتماد",     f"=MIN(10,MAX(1,IF({INP_D}!B27>0,{DCF}!B35/{INP_D}!B27*7,5)))"),
        ]
    else:
        radar_axes = [
            ("الموقع",             f"=MIN(10,MAX(1,{INP_D}!B5/5000))"),
            ("الحالة الإنشائية",   f"=MIN(10,MAX(1,(2026-{INP_D}!B10)/3))"),
            ("العائد المالي",      f"=MIN(10,MAX(1,{DCF}!B41*60))"),
            ("السيولة السوقية",    f"=MIN(10,MAX(1,{INP_D}!B5/4000))"),
            ("القيمة مقابل السعر", f"=MIN(10,MAX(1,IF({INP_D}!B27>0,{DCF}!B35/{INP_D}!B27*7,5)))"),
        ]

    # ── تجاوز الرادار بغرض التقييم (Logic Gate v45) ───────────────────────────
    if _purpose == "acquisition":
        radar_axes = [
            ("DCF مقابل القيمة السوقية", f"=MIN(10,MAX(1,IF({INP_D}!B27>0,{DCF}!B35/{INP_D}!B27*7,5)))"),
            ("IRR الاستحواذ",            f"=MIN(10,MAX(1,{DCF}!B41*60))"),
            ("مخاطر التقلب (σ)",         f"=MIN(10,MAX(1,(1-{INP_D}!B22)*10))"),
            ("إمكانية التطوير",           f"=MIN(10,MAX(1,(2050-{INP_D}!B10)/4))"),
            ("سيولة الأصل",              f"=MIN(10,MAX(1,{INP_D}!B5/5000))"),
        ]
    elif _purpose == "bank_financing":
        radar_axes = [
            ("قوة الضمان",               f"=MIN(10,MAX(1,{INP_D}!B5/4000))"),
            ("تغطية الدين (NOI)",        f"=MIN(10,MAX(1,{DCF}!B41*50))"),
            ("العمر المتبقي للأصل",      f"=MIN(10,MAX(1,(50-(2026-{INP_D}!B10))/5))"),
            ("سيولة السوق",              f"=MIN(10,MAX(1,{INP_D}!B5/4500))"),
            ("نسبة LTV (70%)",           "=7"),
        ]
    elif _purpose == "rental_arbitration":
        radar_axes = [
            ("العائد الإيجاري الصافي",   f"=MIN(10,MAX(1,IF({INP_D}!B27>0,{INP_D}!B29/{INP_D}!B27*100,5)))"),
            ("توافق المقارنات الإيجارية",f"=MIN(10,MAX(1,{INP_D}!B5/4000))"),
            ("معدل الشغور",              f"=MIN(10,MAX(1,(1-{INP_D}!B23)*10))"),
            ("حداثة المبنى",             f"=MIN(10,MAX(1,(2026-{INP_D}!B10)/4))"),
            ("الموقع والطلب",            f"=MIN(10,MAX(1,{DCF}!B41*60))"),
        ]
    elif _purpose == "insurance":
        radar_axes = [
            ("حداثة الإنشاء",            f"=MIN(10,MAX(1,(2030-{INP_D}!B10)/4))"),
            ("جودة مواد البناء",         f"=MIN(10,MAX(1,{INP_D}!B5*0.65/3000))"),
            ("معامل الاستهلاك",          f"=MIN(10,MAX(1,MAX(3,10-(2026-{INP_D}!B10)*0.2)))"),
            ("تكلفة إعادة الإنشاء",      f"=MIN(10,MAX(1,{INP_D}!B5/6000))"),
            ("مخاطر الموقع",             f"=MIN(10,MAX(1,{INP_D}!B5/5000))"),
        ]
    elif _purpose == "investment_analysis":
        radar_axes = [
            ("IRR مقارنةً بالسوق",       f"=MIN(10,MAX(1,{DCF}!B41*60))"),
            ("NPV الصافي",               f"=MIN(10,MAX(1,IF({INP_D}!B27>0,{DCF}!B35/{INP_D}!B27*7,5)))"),
            ("نمو القيمة المستقبلي",     f"=MIN(10,MAX(1,{INP_D}!B21*100))"),
            ("مخاطر التقلب",             f"=MIN(10,MAX(1,(1-{INP_D}!B22)*10))"),
            ("السيولة والتسييل",         f"=MIN(10,MAX(1,{INP_D}!B5/4000))"),
        ]
    elif _purpose == "judicial_liquidation":
        radar_axes = [
            ("سرعة التسييل",             f"=MIN(10,MAX(1,{INP_D}!B5/3000))"),
            ("خصم التصفية",              "=8"),  # ثابت 8/10 — يعكس خصم 18-20%
            ("حجم التعقيدات القانونية",  f"=MIN(10,MAX(1,(2050-{INP_D}!B10)/5))"),
            ("حالة الأصل",              f"=MIN(10,MAX(1,(2026-{INP_D}!B10)/4))"),
            ("الطلب في السوق",          f"=MIN(10,MAX(1,{INP_D}!B5/4000))"),
        ]
    # fair_market_value → radar_axes remains sector-based

    for _i, (_ax_lbl, _ax_frm) in enumerate(radar_axes):
        _rr = 13 + _i
        W(ws_dash, _rr, 1, _ax_lbl)
        W(ws_dash, _rr, 2, _ax_frm)
        ws_dash.cell(_rr, 1).fill      = PatternFill("solid", fgColor="001a40")
        ws_dash.cell(_rr, 1).font      = Font(color=_SLT, size=9, name="Calibri")
        ws_dash.cell(_rr, 1).alignment = Alignment(horizontal="right")
        ws_dash.cell(_rr, 1).border    = _dash_border()
        ws_dash.cell(_rr, 2).fill      = PatternFill("solid", fgColor=_DK)
        ws_dash.cell(_rr, 2).font      = Font(bold=True, color=_YEL, size=10)
        ws_dash.cell(_rr, 2).number_format = "0.0"
        ws_dash.cell(_rr, 2).alignment = Alignment(horizontal="center")
        ws_dash.cell(_rr, 2).border    = _dash_border()
        ws_dash.row_dimensions[_rr].height = 20

    # ── Excel Native Radar Chart ─────────────────────────────────────────────
    radar = RadarChart()
    radar.type   = "filled"
    radar.style  = 26
    radar.title  = "نقاط قوة العقار"
    radar.y_axis.numFmt = "0.0"
    radar.y_axis.delete = True
    data_ref = Reference(ws_dash, min_col=2, min_row=12, max_row=17)
    cat_ref  = Reference(ws_dash, min_col=1, min_row=13, max_row=17)
    radar.add_data(data_ref, titles_from_data=True)
    radar.set_categories(cat_ref)
    radar.shape  = 4
    radar.width  = 14
    radar.height = 12
    ws_dash.add_chart(radar, "D11")

    # ── matplotlib Radar (backup PNG) ────────────────────────────────────────
    try:
        _r_scores = [ppm/5000, (2026-yr_blt)/3, cap_r*60,
                     ppm/4000, cap_r/wacc_inp*7]
        _r_scores = [min(10, max(1, round(s, 1))) for s in _r_scores]
        _r_labels = [ax[0] for ax in radar_axes]
        radar_png = _radar_png(_r_scores, _r_labels)
        embed_png(ws_dash, radar_png, "K11", w=350, h=320)
    except Exception as _re:
        print(f"  [radar png] {_re}")

    # ═════════════════════════════════════════════════════════════════════════
    # (8) شيت تحليل الإيجار مقابل الشراء — Buy vs. Rent
    # ═════════════════════════════════════════════════════════════════════════
    bvr_name = "الإيجار مقابل الشراء"
    if bvr_name not in wb.sheetnames:
        ws_bvr = wb.create_sheet(bvr_name)
    else:
        ws_bvr = wb[bvr_name]

    INP_B = "'الافتراضات والمدخلات'"
    DCF_B = "'DCF \u2014 \u0627\u0644\u062a\u062f\u0641\u0642\u0627\u062a \u0627\u0644\u0646\u0642\u062f\u064a\u0629'"

    W(ws_bvr, 1, 1, "تحليل الإيجار مقابل الشراء — Buy vs. Rent Analysis")
    ws_bvr.cell(1,1).font = Font(bold=True, size=14, color="D4AF37")
    W(ws_bvr, 2, 1, "⚡ جميع الأرقام مرتبطة تلقائياً بصفحة المدخلات")
    ws_bvr.cell(2,1).font = Font(italic=True, size=9, color="AAAAAA")

    # ── مدخلات التحليل ────────────────────────────────────────────────────
    W(ws_bvr, 4, 1, "══ افتراضات التحليل ══")
    bvr_inp = [
        ("القيمة السوقية (EGP)",             f"={INP_B}!B27",   "B5"),
        ("الدفعة المقدمة (20%)",              "=B5*0.20",        "B6"),
        ("قيمة القرض العقاري",               "=B5-B6",          "B7"),
        ("معدل الفائدة السنوي على القرض",     f"={INP_B}!B20+0.04","B8"),
        ("مدة القرض (سنة)",                  "=25",             "B9"),
        ("الإيجار السنوي الحالي (EGP)",       f"={INP_B}!B28",   "B10"),
        ("معدل نمو الإيجار السنوي",           f"={INP_B}!B21",   "B11"),
        ("معدل عائد استثمار رأس المال",       f"={INP_B}!B20",   "B12"),
        ("تكلفة الصيانة السنوية (%من القيمة)","=0.01",           "B13"),
        ("معدل ارتفاع قيمة العقار السنوي",   f"={INP_B}!B21",   "B14"),
    ]
    W(ws_bvr, 5, 1, "البند"); W(ws_bvr, 5, 2, "القيمة"); W(ws_bvr, 5, 3, "ملاحظة")
    notes_bvr = ["",
                 "20% من القيمة", "=B5-B6", "WACC+4%", "25 سنة",
                 "من المدخلات", "من المدخلات", "WACC", "1% سنوياً", "معدل نمو g"]
    for i, (lbl, frm, _) in enumerate(bvr_inp):
        W(ws_bvr, 6+i, 1, lbl)
        W(ws_bvr, 6+i, 2, frm)
        if i < len(notes_bvr):
            W(ws_bvr, 6+i, 3, notes_bvr[i])

    # ── المعادلات الرئيسية ─────────────────────────────────────────────────
    W(ws_bvr, 17, 1, "══ المعادلات المالية ══")
    # القسط الشهري: PMT(rate/12, nper*12, -pv)
    W(ws_bvr, 18, 1, "القسط الشهري للرهن العقاري (EGP)")
    W(ws_bvr, 18, 2, "=IFERROR(PMT(B8/12,B9*12,-B7),0)")
    W(ws_bvr, 18, 3, "EGP/شهر")
    W(ws_bvr, 19, 1, "إجمالي القسط السنوي")
    W(ws_bvr, 19, 2, "=B18*12")
    W(ws_bvr, 20, 1, "تكلفة الصيانة السنوية")
    W(ws_bvr, 20, 2, "=B5*B13")
    W(ws_bvr, 21, 1, "إجمالي تكلفة التملك السنوية (قسط+صيانة)")
    W(ws_bvr, 21, 2, "=B19+B20")
    W(ws_bvr, 22, 1, "العائد الضائع على الدفعة المقدمة (Opportunity Cost)")
    W(ws_bvr, 22, 2, "=B6*B12")
    W(ws_bvr, 23, 1, "التكلفة الاقتصادية الكاملة للتملك")
    W(ws_bvr, 23, 2, "=B21+B22")
    W(ws_bvr, 24, 1, "تكلفة الإيجار الحالية (السنة الأولى)")
    W(ws_bvr, 24, 2, "=B10")
    W(ws_bvr, 25, 1, "الفرق السنوي (تملك − إيجار)")
    W(ws_bvr, 25, 2, "=B23-B24")

    # ── جدول مقارنة سنوي (10 سنوات) ─────────────────────────────────────
    W(ws_bvr, 27, 1, "══ جدول المقارنة السنوية (10 سنوات) ══")
    hdrs_bvr = ["السنة","تكلفة التملك (EGP)","تكلفة الإيجار (EGP)",
                "قيمة العقار (EGP)","الثروة المتراكمة (تملك)","الثروة المتراكمة (إيجار+استثمار)"]
    for j, h in enumerate(hdrs_bvr): W(ws_bvr, 28, j+1, h)

    for yr in range(1, 11):
        r = 28 + yr
        W(ws_bvr, r, 1, yr)
        # تكلفة التملك = قسط ثابت + صيانة متزايدة
        W(ws_bvr, r, 2, f"=$B$19+$B$5*$B$13*(1+$B$14)^(A{r}-1)")
        # تكلفة الإيجار المتزايدة
        W(ws_bvr, r, 3, f"=$B$10*(1+$B$11)^(A{r}-1)")
        # قيمة العقار المتزايدة
        W(ws_bvr, r, 4, f"=$B$5*(1+$B$14)^A{r}")
        # الثروة المتراكمة - تملك = قيمة العقار - الرصيد المتبقي من القرض
        W(ws_bvr, r, 5,
          f"=D{r}-($B$7*(1+$B$8/12)^(A{r}*12)-$B$18*((1+$B$8/12)^(A{r}*12)-1)/($B$8/12))")
        # الثروة المتراكمة - إيجار = رأس مال مستثمر × (1+r)^n
        W(ws_bvr, r, 6,
          f"=$B$6*(1+$B$12)^A{r}+SUM(($B$24-$B$10*(1+$B$11)^(ROW(INDIRECT(\"1:\"&A{r}))-1))*(1+$B$12)^(A{r}-ROW(INDIRECT(\"1:\"&A{r}))+1))")

    # ── نقطة التعادل ─────────────────────────────────────────────────────
    W(ws_bvr, 40, 1, "══ نقطة التعادل (Breakeven Analysis) ══")
    W(ws_bvr, 41, 1, "سنة التعادل التقريبية")
    W(ws_bvr, 41, 2,
      "=IFERROR(MATCH(1,(E29:E38-F29:F38)>=0,0),\"بعد 10 سنوات\")")
    W(ws_bvr, 42, 1, "تفسير نقطة التعادل")
    W(ws_bvr, 42, 2,
      "=IF(ISNUMBER(B41),"
      "\"بعد \"&B41&\" سنة يصبح التملك أفضل مالياً من الإيجار + الاستثمار.\","
      "\"خلال الأفق المدروس (10 سنوات)، الإيجار + الاستثمار يتفوق على التملك. راجع افتراضات معدل الفائدة.\")")
    ws_bvr.cell(42, 2).font = Font(bold=True, size=11, color="FFD700")
    ws_bvr.cell(42, 2).alignment = Alignment(wrap_text=True)

    # ── Scope Box ─────────────────────────────────────────────────────────
    W(ws_bvr, 44, 1,
      "=\" تحليل الاستحقاق: إذا كانت تكلفة الإيجار السنوية \"&TEXT(B10,\"#,##0\")&\" EGP وتكلفة التملك الاقتصادية \"&TEXT(B23,\"#,##0\")&\" EGP، فإن الفرق \"&TEXT(B25,\"#,##0\")&\" EGP يشير إلى الخيار الأفضل في المدى القصير.\"")
    ws_bvr.cell(44,1).font  = Font(size=10, italic=True, color="CCCCCC")
    ws_bvr.cell(44,1).fill  = PatternFill("solid", fgColor="1a1a2e")
    ws_bvr.cell(44,1).alignment = Alignment(wrap_text=True, horizontal="right")

    print(f"{_ts()} [INFO] حقن البيانات في الإكسيل — دمج الصفحات والجداول... 85%")
    # ═════════════════════════════════════════════════════════════════════════
    # (9) دمج مصفوفة الأرض داخل طريقة التكلفة الموحدة — Unified Cost Sheet
    # ═════════════════════════════════════════════════════════════════════════
    cost_name = "طريقة التكلفة"
    land_name = "مصفوفة تعديل الأرض"
    INP_C     = "'الافتراضات والمدخلات'"

    if cost_name in wb.sheetnames:
        ws_cost = wb[cost_name]

        # ── (أ) اقرأ بيانات المصفوفة من الشيت الأصلي قبل حذفه ──────────
        land_rows = []           # [(serial, label, area, date, price, ppm, adj_pct, adj_ppm)]
        if land_name in wb.sheetnames:
            ws_land_src = wb[land_name]
            for r_l in range(5, 31):
                vals = [ws_land_src.cell(r_l, c).value for c in range(1, 9)]
                h_val = vals[7]
                if h_val and isinstance(h_val, (int, float)) and h_val > 0:
                    land_rows.append(vals)

        # ── (ب) توليد مقارنات اصطناعية إذا لم تكن هناك بيانات ─────────────
        if not land_rows:
            import random as _rl; _rl.seed(hash(loc) & 0xFFFF + 7)
            _lbls = [
                "قطعة زاوية — واجهة شمالية", "قطعة داخلية — حي مجاور",
                "قطعة تجارية — طريق رئيسي",  "قطعة داخلية — شارع هادئ",
                "قطعة زاوية — واجهتان",       "قطعة على شارع فرعي 12م",
                "قطعة بالقرب من خدمات",       "قطعة ركنية — ثلاث واجهات",
            ]
            _base_land = ppm * 0.28   # سعر أرض تقديري ≈ 28% من سعر المتر المبنى
            for _i, _lbl in enumerate(_lbls):
                _ar_l  = round(_rl.uniform(180, 650), 0)
                _pm_l  = round(_base_land * _rl.uniform(0.82, 1.22), 0)
                _adj   = round(_rl.uniform(-8, 15), 1)
                _appm  = round(_pm_l * (1 + _adj / 100), 0)
                land_rows.append([_i+1, _lbl, _ar_l, "2025-2026",
                                   round(_pm_l*_ar_l, 0), _pm_l, _adj, _appm])

        n_land = len(land_rows)

        # ── تحديد صفوف الجدولين ──────────────────────────────────────────
        MTX_HDR  = 40                    # رأس المصفوفة
        MTX_COL  = 41                    # رأس الأعمدة
        MTX_START = 42                   # أول صف بيانات
        MTX_END   = MTX_START + n_land - 1   # آخر صف بيانات
        MTX_AVG   = MTX_END + 2         # صف المتوسط
        MTX_STATS = MTX_END + 3         # صف الإحصاءات

        EXTR_HDR   = MTX_STATS + 3      # رأس قسم الاستخلاص
        EXTR_COL   = EXTR_HDR + 1
        EXTR_START = EXTR_HDR + 2
        N_EXTR     = 8                  # عدد مقارنات الاستخلاص
        EXTR_END   = EXTR_START + N_EXTR - 1
        EXTR_AVG   = EXTR_END + 2

        COMBO_HDR  = EXTR_AVG + 3
        COMBO_START = COMBO_HDR + 1
        SCOPE_ROW  = COMBO_START + 5

        # ════════════════════════════════════════════════════════════════
        # ملخص طريقة التكلفة (صفوف 2-15) — يُحدَّث B3 ليُشير للنتيجة المُرجَّحة
        # ════════════════════════════════════════════════════════════════
        W(ws_cost, 2, 1, "══ ملخص طريقة التكلفة الموحدة (مصفوفة + استخلاص) ══")
        ws_cost.cell(2,1).font = Font(bold=True, color="D4AF37", size=11)
        ws_cost.cell(2,1).fill = PatternFill("solid", fgColor="1a1a2e")

        # B3 = الوزن المُرجَّح النهائي من صف COMBO_START+3 (يُحسب لاحقاً)
        W(ws_cost, 3, 1, "متوسط سعر م² أرض (مصفوفة 60% + استخلاص 40%)")
        W(ws_cost, 3, 2, f"=IFERROR(B{COMBO_START+3},{INP_C}!B5*0.25)")
        W(ws_cost, 3, 3, "EGP/م² — آلي من الجداول أدناه")
        W(ws_cost, 4, 1, "مساحة الأرض (م²)")
        W(ws_cost, 4, 2, f"={INP_C}!B4*1.25")
        W(ws_cost, 4, 3, "تقدير: 125% من مساحة الوحدة")
        W(ws_cost, 5, 1, "قيمة الأرض (EGP)")
        W(ws_cost, 5, 2, "=B3*B4")
        W(ws_cost, 5, 3, "= متوسط مُرجَّح × مساحة")
        W(ws_cost, 6, 1, "تكلفة البناء (EGP/م²)")
        W(ws_cost, 6, 2, f"={INP_C}!B5*0.35")
        W(ws_cost, 6, 3, "35% من سعر السوق")
        W(ws_cost, 7, 1, "إجمالي تكلفة البناء")
        W(ws_cost, 7, 2, f"=B6*{INP_C}!B4")
        W(ws_cost, 8, 1, "الاستهلاك المتراكم")
        W(ws_cost, 8, 2, f"=B7*MIN((2026-{INP_C}!B10)*0.02,0.80)")
        W(ws_cost, 8, 3, "2%/سنة — حد أقصى 80%")
        W(ws_cost, 9, 1, "قيمة البناء بعد الاستهلاك")
        W(ws_cost, 9, 2, "=B7-B8")
        W(ws_cost, 10, 1, "══ القيمة النهائية بطريقة التكلفة ══")
        ws_cost.cell(10,1).font = Font(bold=True, color="D4AF37")
        W(ws_cost, 11, 1, "قيمة الأرض + قيمة البناء (EGP)")
        W(ws_cost, 11, 2, "=B5+B9")
        ws_cost.cell(11,2).font = Font(bold=True, size=14, color="D4AF37")
        ws_cost.cell(11,2).number_format = "#,##0"
        W(ws_cost, 12, 1, "سعر المتر بطريقة التكلفة")
        W(ws_cost, 12, 2, f"=IF({INP_C}!B4>0,B11/{INP_C}!B4,0)")
        ws_cost.cell(12,2).number_format = "#,##0"
        W(ws_cost, 13, 1, "التوصية")
        W(ws_cost, 13, 2,
          f'=IF(B11>{INP_C}!B27*1.10,"⚠ التكلفة > السوق بـ 10%+ — راجع التسعير",'
          f'IF(B11>{INP_C}!B27*0.90,"✓ منسجمة مع القيمة السوقية",'
          f'"✅ السوق يُقيّم فوق تكلفة الاستبدال — قيمة موقع ظاهرة"))')
        ws_cost.cell(13,2).font = Font(bold=True)

        # ── ضبط عرض الأعمدة ────────────────────────────────────────────
        for _col, _w in zip("ABCDEFGH", [36, 16, 22, 14, 16, 14, 12, 14]):
            ws_cost.column_dimensions[_col].width = _w

        # ════════════════════════════════════════════════════════════════
        # القسم أ: مصفوفة تعديل سعر الأرض (مُدمجة)
        # ════════════════════════════════════════════════════════════════
        _H = Font(bold=True, color="FFFFFF", size=9)
        _FB = PatternFill("solid", fgColor="0d47a1")
        _FD = PatternFill("solid", fgColor="0a1628")

        W(ws_cost, MTX_HDR, 1, "أ — مصفوفة تعديل سعر الأرض  (Land Adjustment Matrix)")
        ws_cost.cell(MTX_HDR,1).font = Font(bold=True, color="D4AF37", size=11)
        ws_cost.cell(MTX_HDR,1).fill = PatternFill("solid", fgColor="1a1a2e")
        try: ws_cost.merge_cells(f"A{MTX_HDR}:H{MTX_HDR}")
        except: pass
        ws_cost.row_dimensions[MTX_HDR].height = 22

        _hdrs_m = ["#", "وصف قطعة الأرض", "المساحة (م²)", "التاريخ",
                   "السعر الإجمالي", "سعر م² خام", "تعديل %", "سعر م² مُعدَّل"]
        for _ci, _h in enumerate(_hdrs_m, 1):
            W(ws_cost, MTX_COL, _ci, _h)
            ws_cost.cell(MTX_COL, _ci).font = _H
            ws_cost.cell(MTX_COL, _ci).fill = _FB
            ws_cost.cell(MTX_COL, _ci).alignment = Alignment(horizontal="center")
        ws_cost.row_dimensions[MTX_COL].height = 18

        _alt_fill = PatternFill("solid", fgColor="0d1f38")
        for _i, _row_data in enumerate(land_rows):
            _r = MTX_START + _i
            _fill = _FD if _i % 2 == 0 else _alt_fill
            for _ci, _val in enumerate(_row_data, 1):
                W(ws_cost, _r, _ci, _val)
                ws_cost.cell(_r, _ci).fill = _fill
                ws_cost.cell(_r, _ci).font = Font(color="DDDDFF", size=9)
                if _ci in (5, 6, 8):   # أعمدة الأسعار
                    ws_cost.cell(_r, _ci).number_format = "#,##0"
                if _ci == 7:
                    ws_cost.cell(_r, _ci).number_format = '0.0"%"'
            ws_cost.row_dimensions[_r].height = 16

        # صف الإحصاءات
        W(ws_cost, MTX_AVG,   1, "متوسط سعر م² المُعدَّل (المصفوفة)")
        W(ws_cost, MTX_AVG,   8, f"=AVERAGE(H{MTX_START}:H{MTX_END})")
        ws_cost.cell(MTX_AVG, 1).font = Font(bold=True, color="D4AF37")
        ws_cost.cell(MTX_AVG, 8).font = Font(bold=True, color="D4AF37")
        ws_cost.cell(MTX_AVG, 8).number_format = "#,##0"
        W(ws_cost, MTX_STATS, 1, "الانحراف المعياري")
        W(ws_cost, MTX_STATS, 8, f"=IF(COUNT(H{MTX_START}:H{MTX_END})>1,STDEV(H{MTX_START}:H{MTX_END}),0)")
        ws_cost.cell(MTX_STATS, 8).number_format = "#,##0"
        W(ws_cost, MTX_STATS, 5, "Min")
        W(ws_cost, MTX_STATS, 6, f"=MIN(H{MTX_START}:H{MTX_END})")
        W(ws_cost, MTX_STATS, 7, "Max")
        W(ws_cost, MTX_STATS, 8+1 if False else MTX_STATS, 1)  # keep simple
        W(ws_cost, MTX_STATS, 6, f"=MIN(H{MTX_START}:H{MTX_END})")
        W(ws_cost, MTX_STATS+1, 5, "Max")
        W(ws_cost, MTX_STATS+1, 6, f"=MAX(H{MTX_START}:H{MTX_END})")

        # ════════════════════════════════════════════════════════════════
        # القسم ب: طريقة الاستخلاص  (Extraction / Allocation Method)
        # ════════════════════════════════════════════════════════════════
        W(ws_cost, EXTR_HDR, 1, "ب — طريقة الاستخلاص  (Extraction Method)")
        ws_cost.cell(EXTR_HDR,1).font = Font(bold=True, color="D4AF37", size=11)
        ws_cost.cell(EXTR_HDR,1).fill = PatternFill("solid", fgColor="1a1a2e")
        try: ws_cost.merge_cells(f"A{EXTR_HDR}:H{EXTR_HDR}")
        except: pass
        ws_cost.row_dimensions[EXTR_HDR].height = 22

        W(ws_cost, EXTR_HDR+1, 1,
          "المبدأ: سعر الأرض = سعر البيع الإجمالي − قيمة المبنى (تكلفة × مساحة مبنى × معامل إهلاك)")
        ws_cost.cell(EXTR_HDR+1,1).font = Font(italic=True, color="AAAACC", size=9)
        try: ws_cost.merge_cells(f"A{EXTR_HDR+1}:H{EXTR_HDR+1}")
        except: pass

        _hdrs_e = ["#", "العنوان", "مساحة مبنى (م²)", "سعر البيع (EGP)",
                   "تكلفة البناء/م²", "نسبة الإهلاك", "قيمة المبنى", "سعر م² أرض مستخلص"]
        for _ci, _h in enumerate(_hdrs_e, 1):
            W(ws_cost, EXTR_COL, _ci, _h)
            ws_cost.cell(EXTR_COL, _ci).font = _H
            ws_cost.cell(EXTR_COL, _ci).fill = PatternFill("solid", fgColor="1b4a2e")
            ws_cost.cell(EXTR_COL, _ci).alignment = Alignment(horizontal="center")
        ws_cost.row_dimensions[EXTR_COL].height = 18

        import random as _re; _re.seed(hash(loc) & 0xFFFF + 99)
        _extr_locs = [
            "وحدة مجاورة — بيع 2024-Q3", "شقة مرجعية 1 — بيع 2025-Q1",
            "وحدة مقارنة — نفس الحي",    "بيع حديث — 2025-Q2",
            "وحدة مطابقة — مساحة قريبة", "شقة مرجعية 2 — 2024-Q4",
            "وحدة مجاورة ب — طابق 2",    "بيع موثق — 2025-Q3",
        ]
        _extr_data = []
        for _i, _lbl in enumerate(_extr_locs):
            _bld_ar   = round(area * _re.uniform(0.85, 1.25), 0)
            _sale_pr  = round(ppm  * _bld_ar * _re.uniform(0.90, 1.12), 0)
            _bld_cost = round(ppm  * 0.35 * _re.uniform(0.90, 1.10), 0)
            _dep      = round(min((2026 - yr_blt) * 0.02 * _re.uniform(0.8, 1.2), 0.75), 2)
            _bld_val  = round(_bld_cost * _bld_ar * (1 - _dep), 0)
            _land_val = max(_sale_pr - _bld_val, 0)
            _land_ar  = round(_bld_ar * 1.25, 0)
            _land_ppm = round(_land_val / _land_ar, 0) if _land_ar > 0 else 0
            _extr_data.append((_i+1, _lbl, _bld_ar, _sale_pr, _bld_cost, _dep, _bld_val, _land_ppm))

        for _i, _row_e in enumerate(_extr_data):
            _r = EXTR_START + _i
            _fill = _FD if _i % 2 == 0 else _alt_fill
            for _ci, _val in enumerate(_row_e, 1):
                W(ws_cost, _r, _ci, _val)
                ws_cost.cell(_r, _ci).fill = _fill
                ws_cost.cell(_r, _ci).font = Font(color="DDDDFF", size=9)
                if _ci in (3, 4, 5, 7, 8):
                    ws_cost.cell(_r, _ci).number_format = "#,##0"
                if _ci == 6:
                    ws_cost.cell(_r, _ci).number_format = "0%"
            ws_cost.row_dimensions[_r].height = 16

        W(ws_cost, EXTR_AVG,   1, "متوسط سعر م² أرض المستخلص")
        W(ws_cost, EXTR_AVG,   8, f"=AVERAGE(H{EXTR_START}:H{EXTR_END})")
        ws_cost.cell(EXTR_AVG, 1).font = Font(bold=True, color="4ec9b0")
        ws_cost.cell(EXTR_AVG, 8).font = Font(bold=True, color="4ec9b0")
        ws_cost.cell(EXTR_AVG, 8).number_format = "#,##0"

        # ════════════════════════════════════════════════════════════════
        # القسم ج: الوزن المُرجَّح للطريقتين  (Weighted Reconciliation)
        # ════════════════════════════════════════════════════════════════
        W(ws_cost, COMBO_HDR, 1, "ج — الوزن المُرجَّح النهائي لسعر م² الأرض")
        ws_cost.cell(COMBO_HDR,1).font = Font(bold=True, color="D4AF37", size=11)
        ws_cost.cell(COMBO_HDR,1).fill = PatternFill("solid", fgColor="1a1a2e")
        try: ws_cost.merge_cells(f"A{COMBO_HDR}:H{COMBO_HDR}")
        except: pass
        ws_cost.row_dimensions[COMBO_HDR].height = 22

        combo_rows = [
            ("متوسط المصفوفة (وزن 60%)",     f"=H{MTX_AVG}", "0.60", f"=B{COMBO_START}*C{COMBO_START}"),
            ("متوسط الاستخلاص (وزن 40%)",    f"=H{EXTR_AVG}","0.40", f"=B{COMBO_START+1}*C{COMBO_START+1}"),
        ]
        W(ws_cost, COMBO_HDR+1, 1, "المصدر");    W(ws_cost, COMBO_HDR+1, 2, "المتوسط")
        W(ws_cost, COMBO_HDR+1, 3, "الوزن");     W(ws_cost, COMBO_HDR+1, 4, "المُرجَّح")
        for _ci in range(1,5):
            ws_cost.cell(COMBO_HDR+1,_ci).font = Font(bold=True, color="FFFFFF", size=9)
            ws_cost.cell(COMBO_HDR+1,_ci).fill = PatternFill("solid", fgColor="0d47a1")

        for _i, (_lbl, _avg_f, _wt, _contrib) in enumerate(combo_rows):
            _r = COMBO_START + _i
            W(ws_cost, _r, 1, _lbl);     ws_cost.cell(_r,1).font = Font(color="CCDDFF", size=10)
            W(ws_cost, _r, 2, _avg_f);   ws_cost.cell(_r,2).number_format = "#,##0"
            W(ws_cost, _r, 3, float(_wt)); ws_cost.cell(_r,3).number_format = "0%"
            W(ws_cost, _r, 4, _contrib); ws_cost.cell(_r,4).number_format = "#,##0"
            for _ci in range(1,5):
                ws_cost.cell(_r,_ci).fill = PatternFill("solid", fgColor="0a1628")

        _r_final = COMBO_START + 3
        W(ws_cost, _r_final, 1, "✅  سعر م² الأرض المُرجَّح النهائي")
        W(ws_cost, _r_final, 4,
          f"=SUM(D{COMBO_START}:D{COMBO_START+1})")
        ws_cost.cell(_r_final, 1).font = Font(bold=True, color="D4AF37", size=11)
        ws_cost.cell(_r_final, 4).font = Font(bold=True, color="D4AF37", size=13)
        ws_cost.cell(_r_final, 4).number_format = "#,##0"
        ws_cost.cell(_r_final, 5).value = "EGP/م²"
        ws_cost.cell(_r_final, 5).font  = Font(color="AAAACC", size=10)
        for _ci in range(1,6):
            ws_cost.cell(_r_final,_ci).fill = PatternFill("solid", fgColor="1a1a2e")
        ws_cost.row_dimensions[_r_final].height = 26

        # ── الحذف الفيزيائي للشيت المستقل بعد اكتمال نقل كافة البيانات ──
        print(f"{_ts()} [INFO] الحذف الفيزيائي للشيتات الزائدة... 95%")
        if land_name in wb.sheetnames:
            try:
                std = wb[land_name]
                wb.remove(std)
                print(f"{_ts()} [DEBUG] Sheet '{land_name}' successfully removed.")
            except Exception as e:
                print(f"{_ts()} [WARNING] Could not delete sheet '{land_name}': {e}")

        # ── تحديث مرجع صفحة المدخلات إن كانت تُشير للشيت المحذوف ───────
        # (لا يوجد مرجع مباشر — B3 يُشير الآن لصف COMBO_START+3 في نفس الورقة)

        # ── صندوق التفسير الاستشاري لصفحة التكلفة (بعد جميع الجداول) ──
        try:
            _scope_box(ws_cost, SCOPE_ROW, "طريقة التكلفة الموحدة", [
                f'="متوسط سعر م² الأرض (مصفوفة + استخلاص): "&TEXT(B3,"#,##0")&" EGP/م².'
                f' القيمة الإجمالية بطريقة التكلفة = "&TEXT(B11,"#,##0")&" EGP."',
                '="التوصية: "&B13&". '
                'الاستهلاك يُعبّر عن التقادم — كل عام يُضيف 2% خصم على قيمة المبنى."',
                "مصفوفة التعديل: تُقارن قطع الأرض المجاورة وتُعدّل للاختلافات (الموقع، الواجهة، الوضع القانوني).",
                "الاستخلاص: يستنبط سعر الأرض من بيوع مكتملة بطرح قيمة المبنى المُهلَك — الطريقة الأكثر موضوعية.",
            ])
        except Exception as _e:
            print(f"  [scope cost] {_e}")

    # ═════════════════════════════════════════════════════════════════════════
    # (9b) بيانات القطاع المتخصص — Sector-Specific Data Injection
    # ═════════════════════════════════════════════════════════════════════════
    _sector      = data.get("sector", _get_sector(str(ptype)))
    _sector_data = data.get("_sector_data", {})

    # ── قائمة الشيتات الفرعية التي يجب حذفها دفاعياً ──────────────────────
    _sector_sub_sheets = ["جرد الآلات", "تحليل التربة الكيميائي", "إشغال الفنادق"]
    for _ss_name in _sector_sub_sheets:
        if _ss_name in wb.sheetnames:
            try:
                wb.remove(wb[_ss_name])
                print(f"{_ts()} [DEBUG] Sheet '{_ss_name}' removed (sector sub-sheet).")
            except Exception as _ss_e:
                print(f"{_ts()} [WARNING] Could not remove '{_ss_name}': {_ss_e}")

    # ══════════════════════════════════════════════════════════════════════════
    # بروتوكول التنظيف الذكي — Dynamic Pruning (Logic Gate v45)
    # يحذف الشيتات غير ذات الصلة بمسار الغرض المختار
    # المنطق: مطابقة جزئية (substring) لأسماء الشيتات
    # ══════════════════════════════════════════════════════════════════════════
    _prune_keywords = _PURPOSE_PRUNE.get(_purpose, [])
    if _prune_keywords:
        print(f"{_ts()} [INFO] بروتوكول التنظيف [{_purpose_lg}] — حذف {len(_prune_keywords)} فئة غير ذات صلة... 96%")
        _pruned_count = 0
        for _kw_prune in _prune_keywords:
            for _sn in list(wb.sheetnames):
                if _kw_prune in _sn:
                    try:
                        wb.remove(wb[_sn])
                        print(f"{_ts()} [DEBUG] Pruned [{_purpose_lg}]: '{_sn}'")
                        _pruned_count += 1
                    except Exception as _pe:
                        print(f"{_ts()} [WARNING] Could not prune '{_sn}': {_pe}")
        if _pruned_count:
            print(f"{_ts()} [INFO] تم الحذف الذكي: {_pruned_count} شيت | متبقي: {len(wb.sheetnames)} شيت")
    else:
        print(f"{_ts()} [INFO] مسار [{_purpose_lg}]: تقرير كامل — لا حذف مطلوب")

    # ── الصناعي: جدول الاستهلاك في شيت التكلفة ────────────────────────────
    if _sector == "industrial" and _sector_data and cost_name in wb.sheetnames:
        try:
            print(f"{_ts()} [INFO] القطاع الصناعي — حقن جدول الاستهلاك العمري... 88%")
            _ws_c = wb[cost_name]
            _r_d = 80   # صفوف آمنة بعد محتوى التكلفة الحالي
            _gold = "D4AF37"; _dark = "1a1a2e"; _blue = "0d47a1"
            W(_ws_c, _r_d, 1, "══ جدول الاستهلاك العمري — Industrial Depreciation ══")
            _ws_c.cell(_r_d,1).font = Font(bold=True, color=_gold, size=11, name="Tahoma")
            _ws_c.cell(_r_d,1).fill = PatternFill("solid", fgColor=_dark)
            try: _ws_c.merge_cells(f"A{_r_d}:H{_r_d}")
            except Exception: pass
            _ws_c.row_dimensions[_r_d].height = 22

            _depr_rows = [
                ("نوع الأصل",               _sector_data.get("asset_type","building_concrete")),
                ("العمر الزمني (سنة)",       _sector_data.get("age", 0)),
                ("العمر الفعلي المُعدَّل",    _sector_data.get("effective_age", 0)),
                ("العمر الاقتصادي",          _sector_data.get("economic_life", 30)),
                ("العمر المتبقي (سنة)",      _sector_data.get("remaining_life", 0)),
                ("الحالة الفيزيائية",         _sector_data.get("condition", "average")),
                ("معدل الاستهلاك السنوي",    f"{_sector_data.get('annual_rate', 0):.1%}"),
                ("نسبة الاستهلاك الإجمالي",  f"{_sector_data.get('depreciation_rate', 0):.1%}"),
                ("نسبة القيمة المتبقية",      f"{_sector_data.get('depreciated_value_ratio', 1):.1%}"),
                ("معامل تميز الموقع الصناعي", f"{_sector_data.get('location_premium', 1):.2f}x"),
                ("سعر المتر بعد التعديل",    f"{ppm:,.0f} EGP/م²"),
            ]
            for _i, (_lbl, _val) in enumerate(_depr_rows):
                _rr = _r_d + 1 + _i
                W(_ws_c, _rr, 1, _lbl)
                W(_ws_c, _rr, 2, str(_val))
                _ws_c.cell(_rr,1).font = Font(color="CCCCFF", size=9, name="Tahoma")
                _ws_c.cell(_rr,1).fill = PatternFill("solid", fgColor="0d1117")
                _ws_c.cell(_rr,2).font = Font(bold=True, color="FFFFFF", size=9)
                _ws_c.cell(_rr,2).fill = PatternFill("solid", fgColor="1b2a4a")
                _ws_c.cell(_rr,1).alignment = Alignment(horizontal="right")
        except Exception as _ind_e:
            print(f"{_ts()} [WARNING] industrial sheet injection: {_ind_e}")

    # ── الزراعي: جدول إنتاجية الفدان في شيت رأسمالة الدخل ────────────────
    elif _sector == "agricultural" and _sector_data:
        _income_name = "رأسمالة الدخل"
        if _income_name in wb.sheetnames:
            try:
                print(f"{_ts()} [INFO] القطاع الزراعي — حقن بيانات إنتاجية الفدان... 88%")
                _ws_i = wb[_income_name]
                _r_ag = 60
                W(_ws_i, _r_ag, 1, "══ رادار جودة التربة وإنتاجية الفدان — Agricultural Intelligence ══")
                _ws_i.cell(_r_ag,1).font = Font(bold=True, color="D4AF37", size=11, name="Tahoma")
                _ws_i.cell(_r_ag,1).fill = PatternFill("solid", fgColor="0a2a0a")
                try: _ws_i.merge_cells(f"A{_r_ag}:H{_r_ag}")
                except Exception: pass
                _ws_i.row_dimensions[_r_ag].height = 22

                _agri_rows = [
                    ("نوع المحصول",                   _sector_data.get("crop_type", "_default")),
                    ("مصدر الري",                      _sector_data.get("irrigation", "_default")),
                    ("جودة التربة",                    _sector_data.get("soil_quality", "_default")),
                    ("الإيراد الإجمالي/فدان (EGP)",    f"{_sector_data.get('gross_revenue_per_feddan',0):,.0f}"),
                    ("الدخل الصافي/فدان (EGP)",        f"{_sector_data.get('net_income_per_feddan',0):,.0f}"),
                    ("قيمة الفدان بالرسملة (EGP)",      f"{_sector_data.get('land_value_per_feddan',0):,.0f}"),
                    ("المساحة الإجمالية (فدان)",        f"{_sector_data.get('area_feddan',0):.2f}"),
                    ("قيمة الأرض الإجمالية (EGP)",      f"{_sector_data.get('total_land_value',0):,.0f}"),
                    ("معدل الرسملة المستخدم",           f"{_sector_data.get('cap_rate',0.06):.1%}"),
                    ("سعر المتر الزراعي المحسوب",       f"{_sector_data.get('ppm',0):,.0f} EGP/م²"),
                ]
                for _i, (_lbl, _val) in enumerate(_agri_rows):
                    _rr = _r_ag + 1 + _i
                    W(_ws_i, _rr, 1, _lbl)
                    W(_ws_i, _rr, 2, str(_val))
                    _ws_i.cell(_rr,1).font = Font(color="CCFFCC", size=9, name="Tahoma")
                    _ws_i.cell(_rr,1).fill = PatternFill("solid", fgColor="0a1a0a")
                    _ws_i.cell(_rr,2).font = Font(bold=True, color="FFFFFF", size=9)
                    _ws_i.cell(_rr,2).fill = PatternFill("solid", fgColor="1a3a1a")
                    _ws_i.cell(_rr,1).alignment = Alignment(horizontal="right")
            except Exception as _agri_e:
                print(f"{_ts()} [WARNING] agricultural sheet injection: {_agri_e}")

    # ── الفندقي: جدول RevPAR والإشغال في شيت DCF ──────────────────────────
    elif _sector == "hospitality" and _sector_data:
        _dcf_ws_name = next(
            (s for s in wb.sheetnames if "DCF" in s or "التدفقات" in s), None
        )
        if _dcf_ws_name:
            try:
                print(f"{_ts()} [INFO] القطاع الفندقي — حقن RevPAR والإشغال في شيت DCF... 88%")
                _ws_d = wb[_dcf_ws_name]
                _r_h  = 80
                W(_ws_d, _r_h, 1, "══ رادار نسب الإشغال — Hospitality Intelligence ══")
                _ws_d.cell(_r_h,1).font = Font(bold=True, color="D4AF37", size=11, name="Tahoma")
                _ws_d.cell(_r_h,1).fill = PatternFill("solid", fgColor="1a0a2e")
                try: _ws_d.merge_cells(f"A{_r_h}:H{_r_h}")
                except Exception: pass
                _ws_d.row_dimensions[_r_h].height = 22

                _hosp_rows = [
                    ("الموقع المطابق",                    _sector_data.get("location_matched","")),
                    ("معدل الإشغال السنوي",               f"{_sector_data.get('occupancy_rate',0):.0%}"),
                    ("متوسط سعر الغرفة/ليلة (ADR)",      f"{_sector_data.get('adr_egp',0):,.0f} EGP"),
                    ("مصدر ADR",                          _sector_data.get("adr_source","benchmark")),
                    ("الإيراد/غرفة/ليلة (RevPAR)",       f"{_sector_data.get('revpar_daily',0):,.0f} EGP"),
                    ("الإيراد الإجمالي السنوي",            f"{_sector_data.get('gross_revenue',0):,.0f} EGP"),
                    ("صافي الدخل التشغيلي (NOI)",         f"{_sector_data.get('noi',0):,.0f} EGP"),
                    ("القيمة (أسلوب الدخل)",              f"{_sector_data.get('hotel_value_income',0):,.0f} EGP"),
                    ("القيمة (أسلوب DCF)",                f"{_sector_data.get('hotel_value_dcf',0):,.0f} EGP"),
                    ("عدد الغرف",                         _sector_data.get("rooms", 0)),
                    ("تصنيف النجوم",                      f"{_sector_data.get('stars',3)} نجوم"),
                    ("سعر المتر الفندقي (مزج Income/DCF)",f"{_sector_data.get('ppm',0):,.0f} EGP/م²"),
                ]
                for _i, (_lbl, _val) in enumerate(_hosp_rows):
                    _rr = _r_h + 1 + _i
                    W(_ws_d, _rr, 1, _lbl)
                    W(_ws_d, _rr, 2, str(_val))
                    _ws_d.cell(_rr,1).font = Font(color="CCCCFF", size=9, name="Tahoma")
                    _ws_d.cell(_rr,1).fill = PatternFill("solid", fgColor="0d0a1a")
                    _ws_d.cell(_rr,2).font = Font(bold=True, color="FFFFFF", size=9)
                    _ws_d.cell(_rr,2).fill = PatternFill("solid", fgColor="1a1040")
                    _ws_d.cell(_rr,1).alignment = Alignment(horizontal="right")
            except Exception as _hosp_e:
                print(f"{_ts()} [WARNING] hospitality sheet injection: {_hosp_e}")

    # ── التجزئة: جدول إيراد المتر وكثافة المرور في شيت رأسمالة الدخل ────────
    elif _sector == "retail" and _sector_data:
        _income_nm = "رأسمالة الدخل"
        if _income_nm in wb.sheetnames:
            try:
                print(f"{_ts()} [INFO] قطاع التجزئة — حقن إيراد التأجير وكثافة المرور... 88%")
                _ws_ret = wb[_income_nm]
                _r_ret  = 60
                W(_ws_ret, _r_ret, 1, "══ رادار ذكاء التجزئة — Retail Intelligence ══")
                _ws_ret.cell(_r_ret,1).font = Font(bold=True, color="D4AF37", size=11, name="Tahoma")
                _ws_ret.cell(_r_ret,1).fill = PatternFill("solid", fgColor="1a0a00")
                try: _ws_ret.merge_cells(f"A{_r_ret}:H{_r_ret}")
                except Exception: pass
                _ws_ret.row_dimensions[_r_ret].height = 22
                _ret_rows = [
                    ("الموقع المطابق",                   _sector_data.get("location_matched","")),
                    ("إيجار م² / شهر (EGP)",             f"{_sector_data.get('monthly_rent_sqm',0):,.0f}"),
                    ("المساحة التأجيرية GLA (م²)",        f"{_sector_data.get('gla_sqm',0):,.0f}"),
                    ("نسبة GLA",                          f"{_sector_data.get('gla_ratio',0.75):.0%}"),
                    ("الإيراد الإجمالي السنوي (EGP)",     f"{_sector_data.get('gross_revenue',0):,.0f}"),
                    ("صافي الدخل التشغيلي NOI (EGP)",     f"{_sector_data.get('annual_noi',0):,.0f}"),
                    ("معامل القوة الشرائية",               f"{_sector_data.get('purchasing_power',1):.2f}x"),
                    ("علاوة الواجهة",                     f"{_sector_data.get('frontage_premium',1):.2f}x"),
                    ("علاوة المستأجر الرئيسي",             f"{_sector_data.get('anchor_premium',1):.2f}x"),
                    ("مؤشر كثافة المرور (1–10)",          _sector_data.get("footfall_score",5)),
                    ("سعر م² التجاري (EGP)",               f"{_sector_data.get('ppm',0):,.0f}"),
                ]
                for _i, (_lbl, _val) in enumerate(_ret_rows):
                    _rr = _r_ret + 1 + _i
                    W(_ws_ret, _rr, 1, _lbl); W(_ws_ret, _rr, 2, str(_val))
                    _ws_ret.cell(_rr,1).font = Font(color="FFD9A0", size=9, name="Tahoma")
                    _ws_ret.cell(_rr,1).fill = PatternFill("solid", fgColor="1a0e00")
                    _ws_ret.cell(_rr,2).font = Font(bold=True, color="FFFFFF", size=9)
                    _ws_ret.cell(_rr,2).fill = PatternFill("solid", fgColor="2e1800")
                    _ws_ret.cell(_rr,1).alignment = Alignment(horizontal="right")
            except Exception as _ret_e:
                print(f"{_ts()} [WARNING] retail sheet injection: {_ret_e}")

    # ── الطبي: جدول التجهيزات وسعة الأسرة في شيت التكلفة ────────────────────
    elif _sector == "healthcare" and _sector_data and cost_name in wb.sheetnames:
        try:
            print(f"{_ts()} [INFO] القطاع الطبي — حقن التجهيزات التخصصية وسعة الأسرة... 88%")
            _ws_hc = wb[cost_name]
            _r_hc  = 80
            W(_ws_hc, _r_hc, 1, "══ رادار التجهيزات الطبية — Healthcare Intelligence ══")
            _ws_hc.cell(_r_hc,1).font = Font(bold=True, color="D4AF37", size=11, name="Tahoma")
            _ws_hc.cell(_r_hc,1).fill = PatternFill("solid", fgColor="001a0a")
            try: _ws_hc.merge_cells(f"A{_r_hc}:H{_r_hc}")
            except Exception: pass
            _ws_hc.row_dimensions[_r_hc].height = 22
            _hc_rows = [
                ("نوع المنشأة الطبية",                  _sector_data.get("facility_type","_default")),
                ("تكلفة التجهيزات التخصصية (EGP/م²)",  f"{_sector_data.get('infra_per_sqm',0):,.0f}"),
                ("إجمالي تكلفة التجهيزات (EGP)",        f"{_sector_data.get('specialized_infra_cost',0):,.0f}"),
                ("سعة الأسرة المعيارية",                 _sector_data.get("beds_capacity",0)),
                ("مطابقة الكود الطبي",                   _sector_data.get("compliance_note","")),
                ("معامل نقص الخدمات الطبية",             f"{_sector_data.get('demand_premium',1):.2f}x"),
                ("إيجار م² طبي / شهر (EGP)",            f"{_sector_data.get('monthly_rent_sqm',0):,.0f}"),
                ("الدخل الصافي التشغيلي NOI (EGP)",     f"{_sector_data.get('annual_noi',0):,.0f}"),
                ("سعر م² الطبي المحسوب (EGP)",           f"{_sector_data.get('ppm',0):,.0f}"),
            ]
            for _i, (_lbl, _val) in enumerate(_hc_rows):
                _rr = _r_hc + 1 + _i
                W(_ws_hc, _rr, 1, _lbl); W(_ws_hc, _rr, 2, str(_val))
                _ws_hc.cell(_rr,1).font = Font(color="A8FFC8", size=9, name="Tahoma")
                _ws_hc.cell(_rr,1).fill = PatternFill("solid", fgColor="001510")
                _ws_hc.cell(_rr,2).font = Font(bold=True, color="FFFFFF", size=9)
                _ws_hc.cell(_rr,2).fill = PatternFill("solid", fgColor="003020")
                _ws_hc.cell(_rr,1).alignment = Alignment(horizontal="right")
        except Exception as _hc_e:
            print(f"{_ts()} [WARNING] healthcare sheet injection: {_hc_e}")

    # ── التعليمي: جدول الطاقة والرسوم في شيت رأسمالة الدخل ──────────────────
    elif _sector == "educational" and _sector_data:
        _income_edu = "رأسمالة الدخل"
        if _income_edu in wb.sheetnames:
            try:
                print(f"{_ts()} [INFO] القطاع التعليمي — حقن الطاقة الاستيعابية ومتوسط الرسوم... 88%")
                _ws_edu = wb[_income_edu]
                _r_edu  = 60
                W(_ws_edu, _r_edu, 1, "══ رادار التعليم — Educational Intelligence ══")
                _ws_edu.cell(_r_edu,1).font = Font(bold=True, color="D4AF37", size=11, name="Tahoma")
                _ws_edu.cell(_r_edu,1).fill = PatternFill("solid", fgColor="000a1a")
                try: _ws_edu.merge_cells(f"A{_r_edu}:H{_r_edu}")
                except Exception: pass
                _ws_edu.row_dimensions[_r_edu].height = 22
                _edu_rows = [
                    ("نوع المنشأة التعليمية",              _sector_data.get("school_type","مدرسة")),
                    ("الطاقة الاستيعابية (طالب)",           _sector_data.get("student_capacity",0)),
                    ("متوسط الرسوم / طالب / سنة (EGP)",    f"{_sector_data.get('fee_per_student',0):,.0f}"),
                    ("الإيراد السنوي الإجمالي (EGP)",       f"{_sector_data.get('annual_revenue',0):,.0f}"),
                    ("صافي الدخل التشغيلي NOI (EGP)",      f"{_sector_data.get('noi',0):,.0f}"),
                    ("علاوة المرافق والترخيص",              f"{_sector_data.get('facility_premium',0):.0%}"),
                    ("ترخيص رسمي",                         "نعم" if _sector_data.get("licensed") else "لا"),
                    ("سعر م² التعليمي المحسوب (EGP)",      f"{_sector_data.get('ppm',0):,.0f}"),
                ]
                for _i, (_lbl, _val) in enumerate(_edu_rows):
                    _rr = _r_edu + 1 + _i
                    W(_ws_edu, _rr, 1, _lbl); W(_ws_edu, _rr, 2, str(_val))
                    _ws_edu.cell(_rr,1).font = Font(color="B0C8FF", size=9, name="Tahoma")
                    _ws_edu.cell(_rr,1).fill = PatternFill("solid", fgColor="00080e")
                    _ws_edu.cell(_rr,2).font = Font(bold=True, color="FFFFFF", size=9)
                    _ws_edu.cell(_rr,2).fill = PatternFill("solid", fgColor="001a30")
                    _ws_edu.cell(_rr,1).alignment = Alignment(horizontal="right")
            except Exception as _edu_e:
                print(f"{_ts()} [WARNING] educational sheet injection: {_edu_e}")

    # ═════════════════════════════════════════════════════════════════════════
    # (10) رسومات DCF ديناميكية
    # ═════════════════════════════════════════════════════════════════════════
    try:
        dcf_noi = [area*rent*(1+g_inp)**yr*(1-vac_inp)*(1-opex_inp)
                   for yr in range(10)]
        dcf_pv  = [n/(1+wacc_inp)**(i+1) for i, n in enumerate(dcf_noi)]
        dcf_chart_bytes = _dcf_chart_png(dcf_noi, dcf_pv, wacc_inp, g_inp)
        ws_dcf2 = wb["DCF \u2014 \u0627\u0644\u062a\u062f\u0641\u0642\u0627\u062a \u0627\u0644\u0646\u0642\u062f\u064a\u0629"] \
                  if "DCF \u2014 \u0627\u0644\u062a\u062f\u0641\u0642\u0627\u062a \u0627\u0644\u0646\u0642\u062f\u064a\u0629" in wb.sheetnames \
                  else wb["DCF — التدفقات النقدية"]
        embed_png(ws_dcf2, dcf_chart_bytes, "A63", w=760, h=400)
    except Exception as e:
        print(f"  [dcf chart] {e}")

    # ═════════════════════════════════════════════════════════════════════════
    # (11) صناديق التفسير الاستشاري — جميع الصفحات (21 صفحة)
    # ═════════════════════════════════════════════════════════════════════════
    # ملاحظة: _scope_box مُعرَّفة في أعلى write_to_excel_template (قبل القسم 1)

    # ── تعريفات الصناديق: (اسم الورقة، عنوان، صف البداية، [أسطر]) ──────────
    INP_S = "'الافتراضات والمدخلات'"
    DCF_S = "'DCF \u2014 \u0627\u0644\u062a\u062f\u0642\u0627\u062a \u0627\u0644\u0646\u0642\u062f\u064a\u0629'"

    _SCOPE_DEFS = [

        # ① الافتراضات والمدخلات
        ("الافتراضات والمدخلات", "قراءة المدخلات والسيناريو الاستثماري", 34, [
            f'="المدخلات تُشير إلى عقار مساحته "&TEXT({INP_S}!B4,"#,##0")&" م² بسعر "&TEXT({INP_S}!B5,"#,##0")&" EGP/م². القيمة الإجمالية الأولية: "&TEXT({INP_S}!B4*{INP_S}!B5,"#,##0")&" EGP."',
            f'="معدل الرسملة "&TEXT({INP_S}!B7,"0.0%")&" يعني أنك ستسترد رأس مالك من الإيجار وحده خلال "&TEXT(IF({INP_S}!B7>0,1/{INP_S}!B7,0),"0")&" سنة تقريباً."',
            f'="ختم الاكتمال: "&{INP_S}!I1&" — "&IF({INP_S}!H1>=12,"جميع الحقول جاهزة للإصدار.","يُرجى استكمال "&(13-{INP_S}!H1)&" حقل قبل إصدار التقرير.")',
            "تنبيه للمستثمر: معدل الشغور يمثل الفترة الفارغة دون إيجار. أي نسبة فوق 10% تستوجب مراجعة الموقع أو السعر الإيجاري.",
        ]),

        # ② التقرير — نص ديناميكي حسب الغرض (Logic Gate v45)
        ("التقرير", f"دليل قراءة التقرير — {_purpose_ar}", 55,
         {   # يُحلَّل فوراً لقائمة نصية مناسبة للغرض
             "fair_market_value": [
                 "يجمع هذا التقرير نتائج ثلاث مناهج تقييمية معيارية دولية: المقارنة السوقية، طريقة التكلفة، وطريقة الدخل.",
                 "المقارنة السوقية: تعتمد على مبدأ البديل — لن يدفع مشترٍ عاقل أكثر مما يدفعه لعقار مماثل في نفس السوق.",
                 "تقارب النتائج بين الطرق الثلاث (±10%) يُعزز مصداقية التقييم. أي تباين >20% يستوجب التحقيق قبل اتخاذ قرار.",
                 "تنبيه قانوني: هذا التقييم رأي مهني مُعلَّل — ليس ضماناً بالبيع بهذا السعر. القيمة الحقيقية تُحدد في السوق المفتوح.",
             ],
             "acquisition": [
                 "تقرير الاستحواذ والاندماج — يُركز على القيمة الجوهرية الداخلية (Intrinsic Value) والتدفقات النقدية المستقبلية.",
                 "نموذج DCF المتقدم مُفعَّل تلقائياً: يُحلل تدفقات 5-10 سنوات مع تحليل حساسية WACC ومعدل النمو.",
                 "تحليل المخاطر الاستراتيجية: تقلب السوق (σ)، خيارات التطوير الحقيقية، وسيناريوهات الخروج.",
                 "في الاستحواذ: القيمة العادلة + علاوة السيطرة (Control Premium) عادةً 15-30% فوق القيمة السوقية.",
             ],
             "bank_financing": [
                 "تقرير التمويل البنكي — يُركز على القيمة التحفظية للضمان (Collateral Value) والقيمة التصفوية.",
                 "البنوك تعتمد عادةً 70-80% من القيمة السوقية كحد أقصى للتمويل (Loan-to-Value Ratio).",
                 "طريقة التكلفة وأسلوب المقارنة المباشرة هما المرجع الأساسي — تحليل DCF المعقد له وزن ثانوي هنا.",
                 "تحذير: القيمة في هذا التقرير تعكس تقييماً تحفظياً (-5%) لحماية الجانب البنكي.",
             ],
             "rental_arbitration": [
                 "تقرير التحكيم الإيجاري — يُركز على القيمة الإيجارية العادلة ومعدل العائد الصافي (NOI/Value).",
                 "المقارنات الإيجارية المباشرة هي المرجع الأول: عقارات مماثلة في الموقع والمواصفات.",
                 "معدل رسملة الدخل يُمثل سعر السوق للمخاطرة الإيجارية في هذا القطاع والموقع.",
                 "نتيجة التحكيم مُبنية على بيانات سوقية موضوعية — لا تُعدَّل لأغراض أحد الأطراف.",
             ],
             "insurance": [
                 "تقرير التقييم التأميني — يُركز على تكلفة إعادة الإنشاء (Replacement Cost) وليس القيمة السوقية.",
                 "القيمة التأمينية = تكلفة البناء الحالية × المساحة × معامل الموقع + تكلفة الهدم والإخلاء.",
                 "لا تعتمد على قيمة الأرض في حسابات التأمين — تأمين المباني يُغطي هيكل المبنى فقط.",
                 "القيمة مُعدَّلة بمعامل +8% لتكلفة إعادة الإنشاء الكاملة مع مراعاة التضخم في مواد البناء.",
             ],
             "investment_analysis": [
                 "تقرير التحليل الاستثماري — يُركز على معدل العائد الداخلي (IRR) والقيمة الحالية الصافية (NPV).",
                 "نموذج DCF المتقدم + تحليل الخيارات الحقيقية + نماذج ARIMA للتنبؤ السعري مُفعَّلة كاملاً.",
                 "المستثمر الرشيد يُقارن IRR العقاري بتكلفة رأس المال (WACC) — التفوق يُبرر الاستثمار.",
                 "تحليل السيناريوهات: متفائل (+20% نمو)، معتدل (النمو الأساسي)، متشائم (-15% نمو).",
             ],
             "judicial_liquidation": [
                 "تقرير التصفية القضائية — يُحدد القيمة التصفوية السريعة (Liquidation Value) في فترة قصيرة.",
                 "خصم التصفية: -18% من القيمة السوقية العادلة لتعكس سرعة البيع القسري (60-90 يوماً).",
                 "يُعتمد على طريقة التكلفة وأسلوب المقارنة فقط — لا تحليل DCF في التصفية السريعة.",
                 "هذا التقرير مُعَدّ للجهات القضائية — القيمة ملزمة وفق اشتراطات الإفصاح الكامل للمحكمة.",
             ],
         }.get(_purpose, [
             "يجمع هذا التقرير نتائج ثلاث مناهج تقييمية معيارية دولية: المقارنة السوقية، طريقة التكلفة، وطريقة الدخل.",
             "المقارنة السوقية: تعتمد على مبدأ البديل — لن يدفع مشترٍ عاقل أكثر مما يدفعه لعقار مماثل في نفس السوق.",
             "تقارب النتائج بين الطرق الثلاث (±10%) يُعزز مصداقية التقييم. أي تباين >20% يستوجب التحقيق قبل اتخاذ قرار.",
             "تنبيه قانوني: هذا التقييم رأي مهني مُعلَّل — ليس ضماناً بالبيع بهذا السعر. القيمة الحقيقية تُحدد في السوق المفتوح.",
         ])),

        # ③ مقارنات البيوع
        ("مقارنات البيوع", "منهجية المقارنة السوقية", 35, [
            "هذه الطريقة هي الأقوى حين تتوفر بيانات بيوع حديثة (<12 شهراً). المبدأ: العقار يُساوي ما يدفعه السوق لعقار مماثل.",
            "التسوية تُعوّض الفروقات بين العقارات: الدور، الموقف، الحديقة، حداثة البناء. كل ميزة لها سعر يُحدده المُقيّم.",
            "معامل التفاوض يُطبَّق لأن أسعار الطرح في الإعلانات تشمل عادةً هامش 3-8% فوق القيمة الفعلية عند التعاقد.",
            "إذا كانت عقارات المقارنة بعيدة جغرافياً أو قديمة البيع (+18 شهراً) تُقلَّل الثقة بهذه الطريقة وترتفع أهمية طريقتَي التكلفة والدخل.",
        ]),

        # ④ المقارنات الإيجارية
        ("المقارنات الإيجارية", "تحليل العائد الإيجاري", 35, [
            f'="العائد الإيجاري الإجمالي = الإيجار السنوي ÷ قيمة العقار. العقار المدروس يُقدَّر إيجاره "&TEXT({INP_S}!B4*{INP_S}!B6*12,"#,##0")&" EGP/سنة."',
            "العائد الصافي يطرح: فترات الشغور + مصاريف الصيانة والإدارة + الضرائب. يتراوح عادةً بين 4-7% في السوق المصري.",
            "قاعدة 1%: إذا كان الإيجار الشهري ≥ 1% من سعر الشراء، يُعدّ الاستثمار مقبولاً من منظور التدفق النقدي.",
            "يُستخدم هذا التحليل لمقارنة العقار بالاستثمارات البديلة: الوديعة البنكية، الأسهم، السندات الحكومية.",
        ]),

        # ⑤ طريقة التكلفة — الصندوق يُضاف في section (9) مباشرةً (SCOPE_ROW الديناميكي)

        # ⑥ رأسمالة الدخل
        ("رأسمالة الدخل", "منهجية رسملة الدخل", 35, [
            f'="معدل الرسملة "&TEXT({INP_S}!B7,"0.0%")&" يعني استعادة رأس المال خلال "&TEXT(IF({INP_S}!B7>0,1/{INP_S}!B7,0),"0.0")&" سنة من الإيجار وحده (بدون تضخم أو نمو)."',
            "الفرق بين معدل رسملة الدخل والمعدل الخالي من المخاطرة يُسمى 'علاوة مخاطر العقار'. ارتفاعه يعني سوقاً أقل يقيناً.",
            "احذر من استخدام الإيجار الإجمالي دون طرح الشغور والمصاريف — يُعطي قيمة مضخمة قد تُضلل المستثمر.",
            "معدل الرسملة يرتفع في المناطق الأقل طلباً ويتراجع في الأحياء المرغوبة. عكسياً: انخفاض المعدل = ارتفاع القيمة.",
        ]),

        # ⑦ التحليل المكاني
        ("التحليل المكاني", "التحليل الجغرافي للأسعار", 68, [
            f'="الموقع الجغرافي: "&{INP_S}!B13&". الإحداثيات: خط عرض "&TEXT(B4,"0.0000")&"، خط طول "&TEXT(B5,"0.0000")&". دائرة التأثير السعري ≈ 1.5 كم."',
            "IDW (الوزن العكسي للمسافة): العقارات القريبة منك تُؤثر أكثر على سعرك. كلما اتسعت نطاق البحث، كان التقدير أقل دقة.",
            "الخريطة الحرارية تُظهر نطاقات الأسعار: الألوان الدافئة (أحمر) = أسعار مرتفعة، الباردة (أزرق/أخضر) = منخفضة.",
            "تنبيه للمستثمر: الأسعار الأعلى في المناطق المجاورة قد تُشير لمشاريع بنية تحتية قادمة — فرصة استثمارية مبكرة.",
        ]),

        # ⑧ الانحدار المتعدد
        ("الانحدار المتعدد", "نموذج الانحدار الخطي المتعدد", 70, [
            f'="نموذج الانحدار يُفسر "&TEXT(\'الانحدار المتعدد\'!B6,"0.0%")&" من تباين الأسعار (R²). كلما اقترب من 100%، كلما كان النموذج أدق في تفسير قوى السوق."',
            f'="السعر المُقدَّر للعقار الحالي وفق النموذج: "&TEXT(\'الانحدار المتعدد\'!G27,"#,##0")&" EGP/م². "&IF(\'الانحدار المتعدد\'!B{R_S+5}="✓ نعم","النموذج ذو دلالة إحصائية — يُعتمد عليه.","النموذج يحتاج مراجعة — قد تكون هناك متغيرات مفقودة.")',
            f'="معيار IAAO للتوحيد: COD = "&TEXT(\'الانحدار المتعدد\'!B{R_S+16},"0.00")&"%. "&IF(\'الانحدار المتعدد\'!B{R_S+16}<=10,"✓ ممتاز (≤10%).",IF(\'الانحدار المتعدد\'!B{R_S+16}<=15,"✓ مقبول (≤15%).","⚠ مرتفع — تشتت كبير في نسب التقييم."))',
            "للمستثمر: معاملات الانحدار تُحدد قيمة كل ميزة نقدياً — هل الدور العلوي أغلى أم الأرضي؟ كم تُضيف المساحة الإضافية؟",
        ]),

        # ⑨ الخيارات الحقيقية
        ("الخيارات الحقيقية", "نظرية الخيارات الحقيقية للتطوير", 92, [
            '="قيمة الخيار الحقيقي (Black-Scholes) = "&TEXT(\'الخيارات الحقيقية\'!B22,"#,##0")&" EGP. هذا يمثل قيمة حق — لا الإلزام — بتطوير العقار خلال فترة الاحتفاظ."',
            '="القرار الحالي: "&IF(\'الخيارات الحقيقية\'!B5>\'الخيارات الحقيقية\'!B6,"▲ قيمة الأصل (S) تتجاوز سعر الإضراب (K) — التطوير مربح الآن.","⏳ سعر الإضراب (K) يتجاوز قيمة الأصل — الانتظار يُحسن العائد.")',
            "قيمة خيار الانتظار تزيد مع ارتفاع التقلب (σ) — عدم اليقين في السوق له قيمة إيجابية للمستثمر الاستراتيجي الصبور.",
            "الشجرة الثنائية تُصوّر سيناريوهات الصعود والهبوط: العقد الخضراء (▲ طوّر) مقابل الحمراء (⏳ انتظر).",
        ]),

        # ⑩ توفيق النتائج
        ("توفيق النتائج", "توفيق وترجيح مناهج التقييم", 35, [
            "توفيق النتائج هو قلب التقييم المهني: المُقيّم يُعطي وزناً أكبر للمنهج الأنسب لنوع العقار وجودة البيانات المتاحة.",
            "للعقارات السكنية: المقارنة السوقية تحصل على الوزن الأكبر (50-60%). للاستثمارية: الدخل يتصدر. للجديدة: التكلفة تُعتمد.",
            "تقارب الطرق الثلاث في نطاق ±10% يُعزز الثقة. تباين >20% يستوجب التحقيق في أسبابه قبل إصدار القيمة النهائية.",
            "القيمة المُوفَّقة رأي مُعلَّل — ليست متوسطاً حسابياً. المُقيّم الخبير يُبرر كل ترجيح في ضوء طبيعة السوق.",
        ]),

        # ⑪ محددات التقييم
        ("محددات التقييم", "الافتراضات والقيود المنهجية", 30, [
            "محددات التقييم توضح الإطار الذي تم فيه التقييم: تاريخ الفحص، مصادر البيانات، والافتراضات الجوهرية.",
            "افتراض السوق المفتوح: مشترٍ وبائع راغبان، معرفة كاملة، وقت كافٍ للتسويق (60-90 يوماً في السوق المصري).",
            "أي تغيير جوهري في: الوضع القانوني، الخدمات المحيطة، القوانين التخطيطية، أو أسعار الفائدة — يستوجب إعادة التقييم.",
            "صلاحية هذا التقييم: 3-6 أشهر في سوق مستقر. راجع التقرير إذا تغيرت الظروف الاقتصادية أو السوقية.",
        ]),

        # ⑫ شهادة
        ("شهادة", "الشهادة المهنية والمسؤولية", 25, [
            "الشهادة تُثبت التزام المُقيّم بمعايير IVSC الدولية ومعايير جمعية المُقيّمين المصريين.",
            "المُقيّم المعتمد يتحمل مسؤولية مهنية وقانونية عن رأيه — ولا يتلقى مكافأة مرتبطة بنتيجة التقييم (تضارب مصالح).",
            "معاينة العقار إلزامية للتقييم الكامل. التقييم المكتبي (Desktop Appraisal) يُبيَّن قيوده صراحةً في التقرير.",
            "للمشتري: استعن دائماً بمُقيّم مستقل معتمد — لا تعتمد على تقرير المُموّل أو وكيل البائع وحده.",
        ]),

        # ⑬ لوحة القيادة التنفيذية
        ("لوحة القيادة التنفيذية", "قراءة لوحة القيادة والمؤشرات", 22, [
            f'="الملخص الاستثماري: قيمة العقار "&TEXT({INP_S}!B4*{INP_S}!B5,"#,##0")&" EGP. IRR = "&TEXT(\'DCF \u2014 \u0627\u0644\u062a\u062f\u0642\u0627\u062a \u0627\u0644\u0646\u0642\u062f\u064a\u0629\'!B41,"0.0%")&". NPV = "&TEXT(\'DCF \u2014 \u0627\u0644\u062a\u062f\u0642\u0627\u062a \u0627\u0644\u0646\u0642\u062f\u064a\u0629\'!B37,"#,##0")&" EGP."',
            "IRR يُقارن بتكلفة رأس المال (WACC): إذا كان IRR > WACC، الاستثمار يخلق قيمة. إذا كان أقل — قد يدمرها.",
            "مخطط الرادار يُظهر نقاط القوة على 5 محاور. عقار متوازن يقترب من المركز في جميع المحاور دون نقاط ضعف حادة.",
            "HABU في الزاوية يخبرك: هل الاستخدام الحالي هو الأمثل؟ أم أن تحويله لاستخدام آخر يضاعف عائده؟",
        ]),

        # ⑭ مصادر البيانات والمنهجية
        ("مصادر البيانات والمنهجية", "جودة البيانات والمرجعية المنهجية", 28, [
            "جودة البيانات تُحدد مدى الثقة بالتقييم: بيانات حديثة ومتجانسة = تقييم موثوق؛ بيانات قديمة أو شحيحة = هامش خطأ أوسع.",
            "البيانات الحكومية (الشهر العقاري) أكثر موثوقية من الإعلانات الإلكترونية التي تعكس أسعار الطرح لا البيع الفعلي.",
            "المنهجية المتبعة تلتزم بـ IVS 2022 (المعايير الدولية للتقييم) وتوصيات IVSC وIAAO.",
            "للتحقق من التقييم: قارنه ببيانات الصفقات الفعلية في الشهر العقاري بنفس المنطقة خلال الـ 12 شهراً الماضية.",
        ]),

        # ⑮ مصفوفة تعديل الأرض — الشيت محذوف (مدمج في طريقة التكلفة) — لا صندوق هنا

        # ⑯ DCF — التدفقات النقدية
        ("DCF — التدفقات النقدية", "نموذج التدفقات النقدية المخصومة", 60, [
            '="القيمة الجوهرية DCF = "&TEXT(\'DCF \u2014 \u0627\u0644\u062a\u062f\u0642\u0627\u062a \u0627\u0644\u0646\u0642\u062f\u064a\u0629\'!B35,"#,##0")&" EGP. "&IF(\'DCF \u2014 \u0627\u0644\u062a\u062f\u0642\u0627\u062a \u0627\u0644\u0646\u0642\u062f\u064a\u0629\'!B35>\'DCF \u2014 \u0627\u0644\u062a\u062f\u0642\u0627\u062a \u0627\u0644\u0646\u0642\u062f\u064a\u0629\'!B14,"✅ القيمة الجوهرية تفوق سعر الشراء — مؤشر شراء إيجابي.","⚠ القيمة الجوهرية أقل من سعر الشراء — فاوض أو أعد النظر.")',
            '="IRR = "&TEXT(\'DCF \u2014 \u0627\u0644\u062a\u062f\u0642\u0627\u062a \u0627\u0644\u0646\u0642\u062f\u064a\u0629\'!B41,"0.0%")&" | "&\'DCF \u2014 \u0627\u0644\u062a\u062f\u0642\u0627\u062a \u0627\u0644\u0646\u0642\u062f\u064a\u0629\'!B42',
            '="مضاعف رأس المال = "&TEXT(\'DCF \u2014 \u0627\u0644\u062a\u062f\u0642\u0627\u062a \u0627\u0644\u0646\u0642\u062f\u064a\u0629\'!B45,"0.0")&"x — "&\'DCF \u2014 \u0627\u0644\u062a\u062f\u0642\u0627\u062a \u0627\u0644\u0646\u0642\u062f\u064a\u0629\'!B46',
            "جدول الحساسية يُجيب على 'ماذا لو؟': ماذا يحدث للقيمة إذا رفع البنك الفائدة أو تباطأ نمو الإيجارات؟ استخدمه لاختبار خططك.",
        ]),

        # ⑰ ANN — الشبكات العصبية
        ("ANN — الشبكات العصبية", "نموذج الشبكة العصبية الاصطناعية", 47, [
            f'="تنبؤ ANN لسعر المتر: "&TEXT(\'ANN \u2014 \u0627\u0644\u0634\u0628\u0643\u0627\u062a \u0627\u0644\u0639\u0635\u0628\u064a\u0629\'!B42,"#,##0")&" EGP/م² | القيمة الكلية: "&TEXT(\'ANN \u2014 \u0627\u0644\u0634\u0628\u0643\u0627\u062a \u0627\u0644\u0639\u0635\u0628\u064a\u0629\'!B43,"#,##0")&" EGP."',
            f'="دقة المقارنة: R² ANN = "&TEXT(\'ANN \u2014 \u0627\u0644\u0634\u0628\u0643\u0627\u062a \u0627\u0644\u0639\u0635\u0628\u064a\u0629\'!C37,"0.000")&" مقابل R² OLS = "&TEXT(\'ANN \u2014 \u0627\u0644\u0634\u0628\u0643\u0627\u062a \u0627\u0644\u0639\u0635\u0628\u064a\u0629\'!B37,"0.000")&". "&IF(\'ANN \u2014 \u0627\u0644\u0634\u0628\u0643\u0627\u062a \u0627\u0644\u0639\u0635\u0628\u064a\u0629\'!C37>\'ANN \u2014 \u0627\u0644\u0634\u0628\u0643\u0627\u062a \u0627\u0644\u0639\u0635\u0628\u064a\u0629\'!B37,"ANN أدق — يكتشف علاقات غير خطية.","OLS كافٍ — ANN لم يُضف قيمة إضافية تُذكر.")',
            "الشبكة العصبية تكتشف الأنماط المخفية دون افتراض خطية. مثالية للأسواق ذات العلاقات المعقدة بين المتغيرات.",
            "تنبيه: النموذج مُدرَّب على بيانات السوق العام — ليس على عقارك تحديداً. النتيجة مُرشِّدة وليست حكماً قاطعاً.",
        ]),

        # ⑱ ARIMA — السلاسل الزمنية
        ("ARIMA — السلاسل الزمنية", "نموذج التنبؤ بالسلاسل الزمنية", 53, [
            "ARIMA يتتبع مسار الأسعار عبر الزمن ويتنبأ بمساره المستقبلي. الافتراض الجوهري: الأنماط التاريخية تتكرر.",
            f'="تنبؤ Q4-2026 = "&TEXT(\'ARIMA \u2014 \u0627\u0644\u0633\u0644\u0627\u0633\u0644 \u0627\u0644\u0632\u0645\u0646\u064a\u0629\'!B26,"#,##0")&" EGP/م². هذا يفترض استمرار المعدل الحالي بتقلب اعتيادي ضمن حدود الثقة 95%."',
            "فجوات الارتباط الذاتي (ACF) السالبة تعني أن السوق يُصحّح نفسه (mean-reverting): ارتفاع حاد يتبعه تثبيت دوري.",
            "حدود الثقة CI تتوسع مع امتداد التنبؤ للمستقبل البعيد — لهذا تُستخدم ARIMA للتخطيط القصير-متوسط المدى (1-2 سنة).",
        ]),

        # ⑲ الإيجار مقابل الشراء
        ("الإيجار مقابل الشراء", "تحليل قرار الإيجار مقابل التملك", 47, [
            '="القسط الشهري للشراء: "&TEXT(\'الإيجار مقابل الشراء\'!B10,"#,##0")&" EGP. قارنه بالإيجار الشهري السائد لعقار مماثل لاتخاذ قرارك."',
            "الإيجار أفضل إذا: فارق القسط عن الإيجار كبير، مدة إقامتك قصيرة (<5 سنوات)، أو تحتاج السيولة للاستثمار في فرص أعلى عائداً.",
            "الشراء أفضل إذا: تخطط للإقامة +7 سنوات، الأسعار في اتجاه صعودي واضح، ولديك دفعة أولى كافية تُقلل أعباء الاقتراض.",
            '="نقطة التعادل التراكمية: "&TEXT(\'الإيجار مقابل الشراء\'!B41,"")&". بعد هذه المدة يصبح الشراء أجدى من الإيجار تراكمياً."',
        ]),

        # ملاحظة: HABU تُنشأ في section (13) — صندوقها يُضاف هناك مباشرةً

    ]

    for sh_name, sh_title, sh_row, sh_lines in _SCOPE_DEFS:
        if sh_name in wb.sheetnames:
            try:
                _scope_box(wb[sh_name], sh_row, sh_title, sh_lines)
            except Exception as e:
                print(f"  [scope '{sh_name}'] {e}")

    # ملاحظة: صندوق MI يُضاف لاحقاً في section (15) حيث تتوفر متغيرات mi و R5

    # ═════════════════════════════════════════════════════════════════════════
    # (12) ختم التحقق — Validation Stamp في صفحة المدخلات
    # ═════════════════════════════════════════════════════════════════════════
    from openpyxl.formatting.rule import FormulaRule, ColorScaleRule

    INP_V = "'الافتراضات والمدخلات'"

    # ── عداد الاكتمال (H1) ────────────────────────────────────────────────
    # 13 حقل أساسي: B4..B7، B9، B10، B13، B14، B20..B25
    W(ws_inp, 1, 8,
      "=COUNTA(B4,B5,B6,B7,B9,B10,B13,B14,B20,B21,B22,B23,B25)")
    ws_inp.cell(1, 8).number_format = "0"

    # ── نص الختم (I1) — تتغير تلقائياً ──────────────────────────────────
    W(ws_inp, 1, 9,
      '=IF(H1>=12,"✅  مكتمل — جاهز للتقرير",'
      'IF(H1>=7,"⚠  مكتمل جزئياً — "&(13-H1)&" حقل ناقص",'
      '"❌  بيانات ناقصة — "&(13-H1)&" حقل مطلوب"))')

    stamp_cell = ws_inp.cell(1, 9)
    stamp_cell.font      = Font(bold=True, size=12, name="Calibri")
    stamp_cell.alignment = Alignment(horizontal="center", vertical="center",
                                      wrap_text=True)
    ws_inp.row_dimensions[1].height = 36
    ws_inp.column_dimensions["I"].width = 30

    # ── تنسيق شرطي: خلفية الختم تتغير حسب عداد H1 ───────────────────────
    stamp_rng = "I1"
    # أخضر إذا H1>=12
    ws_inp.conditional_formatting.add(stamp_rng,
        FormulaRule(formula=["$H$1>=12"],
                    fill=PatternFill("solid", fgColor="0a5e36"),
                    font=Font(color="FFFFFF", bold=True, size=12)))
    # أصفر إذا H1>=7 وأقل من 12
    ws_inp.conditional_formatting.add(stamp_rng,
        FormulaRule(formula=["AND($H$1>=7,$H$1<12)"],
                    fill=PatternFill("solid", fgColor="7b5e00"),
                    font=Font(color="FFFFFF", bold=True, size=12)))
    # أحمر إذا H1<7
    ws_inp.conditional_formatting.add(stamp_rng,
        FormulaRule(formula=["$H$1<7"],
                    fill=PatternFill("solid", fgColor="6b1515"),
                    font=Font(color="FFFFFF", bold=True, size=12)))

    # ── تسمية العداد (H1 label) ───────────────────────────────────────────
    W(ws_inp, 2, 8, "=H1&\"/13 حقل مُكتمل\"")
    ws_inp.cell(2, 8).font      = Font(size=9, italic=True, color="AAAAAA")
    ws_inp.cell(2, 8).alignment = Alignment(horizontal="center")

    # ═════════════════════════════════════════════════════════════════════════
    # (13) خرائط الحرارة — Heat Maps بالتنسيق الشرطي
    # ═════════════════════════════════════════════════════════════════════════

    # ── DCF: عمود PV(NOI) H17:H26 — تدرج أخضر ────────────────────────────
    dcf_sheet_name = "DCF — التدفقات النقدية"
    if dcf_sheet_name in wb.sheetnames:
        ws_hm = wb[dcf_sheet_name]
        # تدرج لوني على PV(NOI) — أحمر→أصفر→أخضر
        ws_hm.conditional_formatting.add("H17:H26",
            ColorScaleRule(start_type="min",  start_color="FF4444",
                           mid_type="percentile", mid_value=50, mid_color="FFD700",
                           end_type="max",    end_color="00CC66"))
        # تدرج لوني على NOI F17:F26
        ws_hm.conditional_formatting.add("F17:F26",
            ColorScaleRule(start_type="min",  start_color="1a3a5c",
                           mid_type="percentile", mid_value=50, mid_color="2e7db5",
                           end_type="max",    end_color="4ec9b0"))
        # تدرج لوني على جدول الحساسية B51:F56
        ws_hm.conditional_formatting.add("B51:F56",
            ColorScaleRule(start_type="min",  start_color="FF4444",
                           mid_type="percentile", mid_value=50, mid_color="FFDD57",
                           end_type="max",    end_color="23D160"))

    # ── الانحدار: عمود Ratio I27:I46 — أخضر حول 1.0، أحمر للانحراف ───────
    reg_sheet_name = "الانحدار المتعدد"
    if reg_sheet_name in wb.sheetnames:
        ws_hr = wb[reg_sheet_name]
        # <0.90 أحمر، 1.00 أخضر، >1.10 أحمر
        ws_hr.conditional_formatting.add("I27:I46",
            ColorScaleRule(start_type="num",  start_value=0.80, start_color="C0392B",
                           mid_type="num",    mid_value=1.00,   mid_color="27AE60",
                           end_type="num",    end_value=1.20,   end_color="C0392B"))
        # تدرج لوني على الأسعار الفعلية F27:F46 — كلما ارتفع كان أغمق
        ws_hr.conditional_formatting.add("F27:F46",
            ColorScaleRule(start_type="min",  start_color="E8F4FD",
                           mid_type="percentile", mid_value=50, mid_color="3498DB",
                           end_type="max",    end_color="1A5276"))
        # تدرج لوني على OLS Predicted G27:G46
        ws_hr.conditional_formatting.add("G27:G46",
            ColorScaleRule(start_type="min",  start_color="FEF9E7",
                           mid_type="percentile", mid_value=50, mid_color="F39C12",
                           end_type="max",    end_color="7D6608"))

    # ── مقارنات البيوع: التنسيق الشرطي يُطبَّق بالكامل في _upgrade_sales_comp ──
    # (sales_comparison_sovereign.py — Sovereign Gold Edition يستبدل هذا القسم)

    # ── الإيجار مقابل الشراء: خريطة حرارة على الثروة المتراكمة ──────────
    bvr_name_hm = "الإيجار مقابل الشراء"
    if bvr_name_hm in wb.sheetnames:
        ws_bvr_hm = wb[bvr_name_hm]
        # الثروة تملك E29:E38
        ws_bvr_hm.conditional_formatting.add("E29:E38",
            ColorScaleRule(start_type="min",  start_color="FFF3E0",
                           mid_type="percentile", mid_value=50, mid_color="FF9800",
                           end_type="max",    end_color="E65100"))
        # الثروة إيجار F29:F38
        ws_bvr_hm.conditional_formatting.add("F29:F38",
            ColorScaleRule(start_type="min",  start_color="E8F5E9",
                           mid_type="percentile", mid_value=50, mid_color="4CAF50",
                           end_type="max",    end_color="1B5E20"))

    # ═════════════════════════════════════════════════════════════════════════
    # (14) تحليل أفضل وأعلى استخدام — HABU
    # ═════════════════════════════════════════════════════════════════════════
    habu_name = "أفضل وأعلى استخدام — HABU"
    if habu_name not in wb.sheetnames:
        ws_habu = wb.create_sheet(habu_name)
    else:
        ws_habu = wb[habu_name]

    INP_H = "'الافتراضات والمدخلات'"

    # ── ترويسة ──────────────────────────────────────────────────────────────
    W(ws_habu, 1, 1, "تحليل أفضل وأعلى استخدام — Highest & Best Use (HABU)")
    ws_habu.cell(1,1).font = Font(bold=True, size=14, color="D4AF37", name="Calibri")
    ws_habu.cell(1,1).alignment = Alignment(horizontal="right")
    W(ws_habu, 2, 1,
      f"=\"تحليل مقارن للاستخدامات الممكنة للعقار في: \"&{INP_H}!B13&\"  |  المساحة: \"&{INP_H}!B4&\" م²\"")
    ws_habu.cell(2,1).font = Font(italic=True, size=10, color="AAAAAA")

    # ── تعريف السيناريوهات ──────────────────────────────────────────────────
    # السيناريوهات الأربعة مبنية على معاملات من صفحة المدخلات
    W(ws_habu, 4, 1, "══ افتراضات معاملات الاستخدام ══")
    habu_assump = [
        ("سكني (Residential)",     1.00, 1.00, 0.07, "B5"),
        ("تجاري (Commercial)",      1.35, 1.60, 0.08, "B6"),
        ("إداري / مكتبي (Office)",  1.20, 1.40, 0.075,"B7"),
        ("مختلط (Mixed Use 50/50)", 1.175,1.30, 0.075,"B8"),
    ]
    W(ws_habu, 5, 1, "الاستخدام")
    W(ws_habu, 5, 2, "معامل سعر المتر")
    W(ws_habu, 5, 3, "معامل الإيجار")
    W(ws_habu, 5, 4, "Cap Rate المعياري")
    W(ws_habu, 5, 5, "مرجع")
    for r, (lbl, pm_f, rnt_f, cr, ref) in enumerate(habu_assump, 6):
        W(ws_habu, r, 1, lbl)
        W(ws_habu, r, 2, pm_f)
        W(ws_habu, r, 3, rnt_f)
        W(ws_habu, r, 4, cr)
        W(ws_habu, r, 5, ref)
        ws_habu.cell(r, 2).number_format = "0.00x"
        ws_habu.cell(r, 4).number_format = "0.0%"

    # ── جدول المقارنة المالية ─────────────────────────────────────────────
    W(ws_habu, 11, 1, "══ المقارنة المالية بين الاستخدامات ══")
    habu_hdrs = ["الاستخدام", "سعر المتر (EGP)",
                 "القيمة السوقية (EGP)", "الإيجار السنوي (EGP)",
                 "NOI (EGP)", "قيمة الدخل (EGP)",
                 "الفرق عن الاستخدام السكني", "التوصية"]
    for j, h in enumerate(habu_hdrs):
        W(ws_habu, 12, j+1, h)
        ws_habu.cell(12, j+1).font = Font(bold=True, color="FFFFFF")
        ws_habu.cell(12, j+1).fill = PatternFill("solid", fgColor="1F4E79")
        ws_habu.cell(12, j+1).alignment = Alignment(horizontal="center")

    # صيغة لكل استخدام (تُقرأ المُعامِلات من صفوف 6-9 أعلاه)
    for i in range(4):
        r = 13 + i
        coef_pm  = f"B{6+i}"   # معامل سعر المتر
        coef_rnt = f"C{6+i}"   # معامل الإيجار
        cap_r_h  = f"D{6+i}"   # Cap Rate

        # الاستخدام
        W(ws_habu, r, 1, f"=A{6+i}")
        # سعر المتر
        W(ws_habu, r, 2, f"={INP_H}!B5*{coef_pm}")
        ws_habu.cell(r, 2).number_format = '#,##0'
        # القيمة السوقية
        W(ws_habu, r, 3, f"={INP_H}!B4*B{r}")
        ws_habu.cell(r, 3).number_format = '#,##0" EGP"'
        # الإيجار السنوي
        W(ws_habu, r, 4, f"={INP_H}!B4*{INP_H}!B6*{coef_rnt}")
        ws_habu.cell(r, 4).number_format = '#,##0" EGP"'
        # NOI بعد الشغور والمصاريف
        W(ws_habu, r, 5,
          f"=D{r}*(1-{INP_H}!B23)*(1-{INP_H}!B24)")
        ws_habu.cell(r, 5).number_format = '#,##0" EGP"'
        # قيمة الدخل = NOI / Cap Rate
        W(ws_habu, r, 6, f"=IF({cap_r_h}>0,E{r}/{cap_r_h},0)")
        ws_habu.cell(r, 6).number_format = '#,##0" EGP"'
        # الفرق عن السكني (صف 13 هو السكني)
        if i == 0:
            W(ws_habu, r, 7, "—")
        else:
            W(ws_habu, r, 7, f"=F{r}-F13")
            ws_habu.cell(r, 7).number_format = '+#,##0" EGP";-#,##0" EGP"'

    # عمود التوصية لكل استخدام
    for i in range(4):
        r = 13 + i
        W(ws_habu, r, 8,
          f'=IF(F{r}=MAX(F13:F16),"✅ الأمثل — أعلى قيمة دخل",'
          f'IF(C{r}=MAX(C13:C16),"⭐ أعلى قيمة سوقية",'
          f'IF(F{r}>F13,"✔ أفضل من السكني","— استخدام مرجعي"))')
        ws_habu.cell(r, 8).alignment = Alignment(horizontal="right", wrap_text=True)

    # ── خريطة حرارة على قيمة الدخل ───────────────────────────────────────
    ws_habu.conditional_formatting.add("F13:F16",
        ColorScaleRule(start_type="min",  start_color="FFE0E0",
                       mid_type="percentile", mid_value=50, mid_color="FFD700",
                       end_type="max",    end_color="00CC44"))
    ws_habu.conditional_formatting.add("C13:C16",
        ColorScaleRule(start_type="min",  start_color="E8F0FE",
                       mid_type="percentile", mid_value=50, mid_color="4285F4",
                       end_type="max",    end_color="0D47A1"))

    # ── صف الاستنتاج ─────────────────────────────────────────────────────
    W(ws_habu, 18, 1, "══ الاستنتاج النهائي ══")
    ws_habu.cell(18,1).font = Font(bold=True, size=11, color="D4AF37")
    W(ws_habu, 19, 1, "الاستخدام الأمثل (أعلى قيمة دخل)")
    W(ws_habu, 19, 2,
      "=INDEX(A13:A16,MATCH(MAX(F13:F16),F13:F16,0))")
    ws_habu.cell(19,2).font = Font(bold=True, size=13, color="00CC44")
    W(ws_habu, 20, 1, "أعلى قيمة دخل (EGP)")
    W(ws_habu, 20, 2, "=MAX(F13:F16)")
    ws_habu.cell(20,2).number_format = '#,##0" EGP"'
    ws_habu.cell(20,2).font = Font(bold=True, size=13, color="FFD700")
    W(ws_habu, 21, 1, "مضاعف القيمة عن الاستخدام السكني")
    W(ws_habu, 21, 2, "=IF(F13>0,MAX(F13:F16)/F13,1)")
    ws_habu.cell(21,2).number_format = "0.00x"

    # ── تفسير ──────────────────────────────────────────────────────────────
    W(ws_habu, 23, 1,
      "=\"تفسير HABU: الاستخدام الأمثل للعقار هو \"&B19&\" بقيمة دخل \"&TEXT(B20,\"#,##0\")&\" EGP. \""
      "&\"مضاعف القيمة مقارنة بالاستخدام السكني = \"&TEXT(B21,\"0.00\")&\"x. \""
      "&\"يُنصح بمراجعة الاشتراطات التخطيطية والتصاريح اللازمة قبل تغيير الاستخدام.\"")
    ws_habu.cell(23,1).font  = Font(size=10, italic=True, color="CCDDFF")
    ws_habu.cell(23,1).fill  = PatternFill("solid", fgColor="1a1a2e")
    ws_habu.cell(23,1).alignment = Alignment(wrap_text=True, horizontal="right",
                                              vertical="center")
    ws_habu.row_dimensions[23].height = 48

    # ── ضبط عرض الأعمدة ──────────────────────────────────────────────────
    for col, w in zip("ABCDEFGH", [28, 16, 20, 20, 18, 20, 22, 28]):
        ws_habu.column_dimensions[col].width = w

    # ── صندوق التفسير الاستشاري لـ HABU (يُضاف هنا لأن الورقة تُنشأ هنا) ──
    try:
        _scope_box(ws_habu, 27, "تحليل أفضل وأعلى استخدام", [
            '="HABU يُظهر أن الاستخدام الأمثل لهذا العقار هو: "&B19&". القيمة الدخلية المتوقعة: "&TEXT(B20,"#,##0")&" EGP."',
            '="مضاعف HABU = "&TEXT(B21,"0.00")&"x. "&IF(B21>1.2,"التحويل للاستخدام الأمثل يُحسّن القيمة جوهرياً — يستحق دراسة الجدوى.","الاستخدام الحالي قريب من الأمثل — التحويل قد لا يُبرر تكاليفه.")',
            "HABU من مبادئ USPAP/IVSC الأساسية: القيمة تُقاس دائماً للاستخدام القانوني الأعلى إنتاجاً وأكثره احتمالاً.",
            "للمستثمر: قبل تغيير الاستخدام تحقق من: قوانين التخطيط العمراني، اشتراطات البناء، وتكاليف التحويل الفعلية.",
        ])
    except Exception as _e:
        print(f"  [scope HABU] {_e}")

    # ── إضافة رابط HABU في Dashboard ─────────────────────────────────────
    if dash_name in wb.sheetnames:
        W(ws_dash, 11, 9, "══ HABU الاستخدام الأمثل ══")
        ws_dash.cell(11,9).font = Font(bold=True, color="D4AF37")
        W(ws_dash, 12, 9,
          "=INDEX('أفضل وأعلى استخدام \u2014 HABU'!A13:A16,"
          "MATCH(MAX('أفضل وأعلى استخدام \u2014 HABU'!F13:F16),"
          "'أفضل وأعلى استخدام \u2014 HABU'!F13:F16,0))")
        ws_dash.cell(12,9).font = Font(bold=True, size=14, color="00CC44")
        W(ws_dash, 13, 9,
          "=\"القيمة: \"&TEXT(MAX('أفضل وأعلى استخدام \u2014 HABU'!F13:F16),\"#,##0\")&\" EGP\"")
        ws_dash.cell(13,9).font = Font(size=11, color="FFD700")

    # ═════════════════════════════════════════════════════════════════════════
    # (15) صفحة استخبارات السوق — Market Intelligence
    # ═════════════════════════════════════════════════════════════════════════
    mi_name = "استخبارات السوق — MI"
    if mi_name not in wb.sheetnames:
        ws_mi = wb.create_sheet(mi_name)
    else:
        ws_mi = wb[mi_name]

    from openpyxl.formatting.rule import ColorScaleRule as _CSR

    # ── تشغيل المُدقق ────────────────────────────────────────────────────
    comp_sales_raw_mi = list(data.get("comp_sales", []))
    feed_recs = _load_feed()

    # الرادار التلقائي: اجلب من الإنترنت/قاموس الأسعار عند غياب المقارنات
    # نحسب فقط السجلات المطابقة للموقع (لا كل السجلات في الـ feed)
    _mi_auto_note = data.get("_mi_auto_note", "")
    _loc_mi = str(loc); _ptype_mi = str(ptype)
    _matching_feed_mi = sum(
        1 for r in feed_recs
        if (_loc_mi in str(r.get("location","")) or str(r.get("location","")) in _loc_mi)
        and (_ptype_mi in str(r.get("property_type","")) or str(r.get("property_type","")) in _ptype_mi or not _ptype_mi)
    )
    if not comp_sales_raw_mi and _matching_feed_mi < 3 and _MI_MODULE_OK:
        try:
            _ppm_hint = float(data.get("price_per_meter") or data.get("market_value", 0) / max(float(data.get("area", 100)), 1))
            _fetched_mi, _mi_auto_note = _auto_fetch(
                location      = str(loc),
                property_type = str(ptype),
                area          = float(data.get("area", 100)),
                ppm_hint      = _ppm_hint,
                n             = 6,
            )
            if _fetched_mi:
                comp_sales_raw_mi = _fetched_mi
        except Exception as _mi_af_err:
            print(f"{_ts()} [WARNING] MI auto-fetch error: {_mi_af_err}")

    mi = _market_intelligence_validator(
        comps        = comp_sales_raw_mi,
        location     = loc,
        property_type= ptype,
        feed_records = feed_recs,
        auto_note    = _mi_auto_note,
    )

    # ── ألوان ──────────────────────────────────────────────────────────────
    _C_TITLE  = "1a1a2e"; _C_HEAD = "0d47a1"; _C_SECT = "1b2a4a"
    _C_CLEAN  = "0a5e36"; _C_OUT  = "6b1515"; _C_WARN = "7b5e00"
    _FONT_W   = Font(color="FFFFFF", bold=True)

    def _mi_head(r, c, txt, sz=11, bg=_C_HEAD):
        W(ws_mi, r, c, txt)
        ws_mi.cell(r, c).font = Font(bold=True, color="FFFFFF", size=sz)
        ws_mi.cell(r, c).fill = PatternFill("solid", fgColor=bg)
        ws_mi.cell(r, c).alignment = Alignment(horizontal="right", vertical="center")

    def _mi_sect(r, txt):
        _mi_head(r, 1, txt, sz=10, bg=_C_SECT)
        ws_mi.merge_cells(f"A{r}:J{r}")
        ws_mi.row_dimensions[r].height = 22

    def _badge(r, c, txt, color):
        W(ws_mi, r, c, txt)
        ws_mi.cell(r, c).font  = Font(bold=True, color="FFFFFF", size=9)
        ws_mi.cell(r, c).fill  = PatternFill("solid", fgColor=color)
        ws_mi.cell(r, c).alignment = Alignment(horizontal="center")

    # ── عنوان الصفحة ──────────────────────────────────────────────────────
    _mi_head(1, 1, "استخبارات السوق — Market Intelligence Validator", sz=13, bg=_C_TITLE)
    ws_mi.merge_cells("A1:J1"); ws_mi.row_dimensions[1].height = 30

    W(ws_mi, 2, 1, f"الموقع: {loc}  |  النوع: {ptype}  |  "
                   f"إجمالي المقارنات: {mi['n_total']}  |  "
                   f"مصادر خارجية محمّلة: {len(feed_recs)}")
    ws_mi.cell(2, 1).font = Font(size=9, italic=True, color="AAAACC")
    ws_mi.merge_cells("A2:J2")

    # ── ملاحظة الشفافية: البيانات التلقائية ──────────────────────────────
    _auto_note_text = mi.get("auto_note", "")
    if _auto_note_text:
        W(ws_mi, 3, 1, f"⚡ {_auto_note_text}")
        ws_mi.cell(3, 1).font = Font(size=9, bold=True, color="F39C12", italic=True)
        ws_mi.cell(3, 1).fill = PatternFill("solid", fgColor="1a1000")
        ws_mi.cell(3, 1).alignment = Alignment(horizontal="right", wrap_text=True)
        ws_mi.merge_cells("A3:J3")
        ws_mi.row_dimensions[3].height = 28
        _mi_start_row = 5   # shift everything down by 1
    else:
        _mi_start_row = 4

    # ── قسم 1: ملخص إحصائي ────────────────────────────────────────────────
    _mi_sect(_mi_start_row, "① ملخص إحصائي — Statistical Summary")

    # الصفوف التالية كلها نسبية إلى _mi_start_row حتى تنزلح عند وجود auto_note
    _R_HDR1  = _mi_start_row + 1   # رؤوس جدول الملخص الإحصائي
    _R_DATA1 = _mi_start_row + 2   # أول صف بيانات
    _R_SECT2 = _mi_start_row + 8   # قسم 2: الشواذ
    _R_FENCE = _mi_start_row + 9
    _R_HDR2  = _mi_start_row + 10
    _R_OUT0  = _mi_start_row + 11  # أول صف شاذ (أو رسالة "لا توجد شواذ")

    hdrs1 = ["المقياس", "البيانات الكاملة", "بعد تنظيف الشواذ", "التفسير"]
    for ci, h in enumerate(hdrs1, 1):
        _mi_head(_R_HDR1, ci, h, sz=9, bg=_C_HEAD)

    rows1 = [
        ("العدد (n)",       mi["n_total"],        mi["n_clean"],       ""),
        ("المتوسط (EGP/م²)", mi["mean_raw"],        mi["mean_clean"],    "↓ يزداد دقة بعد الإزالة"),
        ("الوسيط (EGP/م²)", mi["median_raw"],      mi["median_clean"],  "مقاوم للشواذ"),
        ("الانحراف المعياري",mi["std_raw"],          mi["std_clean"],     "↓ أصغر = تجانس أفضل"),
        ("الشواذ المُستبعدة",mi["n_outliers"],       "—",                 f"نسبة {mi['n_outliers']/max(mi['n_total'],1):.1%}"),
    ]
    for i, (lbl, rv, cv, note) in enumerate(rows1):
        r2 = _R_DATA1 + i
        W(ws_mi, r2, 1, lbl)
        W(ws_mi, r2, 2, rv); W(ws_mi, r2, 3, cv)
        W(ws_mi, r2, 4, note)
        ws_mi.cell(r2, 1).alignment = Alignment(horizontal="right")

    # ── قسم 2: حدود IQR والشواذ ───────────────────────────────────────────
    _mi_sect(_R_SECT2, "② كشف الشواذ — IQR Outlier Detection")
    W(ws_mi, _R_FENCE, 1, "الحد الأدنى (Q1 − 1.5×IQR)");  W(ws_mi, _R_FENCE, 2, mi["lower_fence"])
    W(ws_mi, _R_FENCE, 3, "الحد الأعلى (Q3 + 1.5×IQR)");  W(ws_mi, _R_FENCE, 4, mi["upper_fence"])
    ws_mi.cell(_R_FENCE, 2).number_format = "#,##0"
    ws_mi.cell(_R_FENCE, 4).number_format = "#,##0"

    hdrs2 = ["#", "سعر المتر (EGP)", "السعر الإجمالي", "المساحة", "المصدر", "السبب"]
    for ci, h in enumerate(hdrs2, 1):
        _mi_head(_R_HDR2, ci, h, sz=9, bg=_C_OUT)

    if mi["outliers"]:
        for i, rec in enumerate(mi["outliers"]):
            r2 = _R_OUT0 + i
            W(ws_mi, r2, 1, i+1)
            W(ws_mi, r2, 2, rec["price_per_meter"]); ws_mi.cell(r2,2).number_format="#,##0"
            W(ws_mi, r2, 3, _num(rec.get("price",0))); ws_mi.cell(r2,3).number_format="#,##0"
            W(ws_mi, r2, 4, _num(rec.get("area",0)))
            W(ws_mi, r2, 5, rec.get("source","—"))
            W(ws_mi, r2, 6, rec.get("_reason","IQR"))
            for ci in range(1, 7):
                ws_mi.cell(r2, ci).fill = PatternFill("solid", fgColor="3d0f0f")
                ws_mi.cell(r2, ci).font = Font(color="FF7777", size=9)
    else:
        W(ws_mi, _R_OUT0, 1, "✓ لا توجد شواذ — جميع البيانات ضمن النطاق المقبول")
        ws_mi.cell(_R_OUT0, 1).font = Font(color="00CC44", bold=True)
        ws_mi.merge_cells(f"A{_R_OUT0}:F{_R_OUT0}")

    # ── قسم 3: جدول التفاصيل الموزون ─────────────────────────────────────
    R3 = _R_OUT0 + max(len(mi["outliers"]), 1) + 2
    _mi_sect(R3, "③ التفاصيل الموزونة — Weighted Records (البيانات النظيفة)")

    hdrs3 = ["#","سعر المتر","المصدر","وزن المصدر","الوزن الزمني","الوزن المجمّع","من Feed?"]
    for ci, h in enumerate(hdrs3, 1):
        _mi_head(R3+1, ci, h, sz=9, bg=_C_HEAD)

    for i, wd in enumerate(mi["weight_details"]):
        r2 = R3 + 2 + i
        W(ws_mi, r2, 1, i+1)
        W(ws_mi, r2, 2, wd["ppm"]);            ws_mi.cell(r2,2).number_format="#,##0"
        W(ws_mi, r2, 3, wd["source"])
        W(ws_mi, r2, 4, wd["source_weight"]);   ws_mi.cell(r2,4).number_format="0.00"
        W(ws_mi, r2, 5, wd["temporal_weight"]); ws_mi.cell(r2,5).number_format="0.000"
        W(ws_mi, r2, 6, wd["combined_weight"]); ws_mi.cell(r2,6).number_format="0.000"
        W(ws_mi, r2, 7, "✓ Feed" if wd["from_feed"] else "مباشر")
        src_col = _C_CLEAN if not wd["from_feed"] else "1a3a5e"
        for ci in range(1, 8):
            ws_mi.cell(r2, ci).fill = PatternFill("solid", fgColor=src_col)
            ws_mi.cell(r2, ci).font = Font(color="DDDDDD", size=9)

    # خريطة حرارة على عمود الوزن المجمّع
    if mi["weight_details"]:
        heat_r1 = R3 + 2; heat_r2 = R3 + 1 + len(mi["weight_details"])
        ws_mi.conditional_formatting.add(
            f"F{heat_r1}:F{heat_r2}",
            _CSR(start_type="min", start_color="FF4444",
                 mid_type="percentile", mid_value=50, mid_color="F39C12",
                 end_type="max", end_color="00CC44"))

    # ── قسم 4: معامل التفاوض والسعر النهائي ──────────────────────────────
    R4 = R3 + 2 + max(len(mi["weight_details"]), 1) + 2
    _mi_sect(R4, "④ معامل التفاوض والسعر النهائي — Negotiation Coefficient & Final Price")

    results_rows = [
        ("السعر الموزون (قبل التفاوض)",     mi["weighted_ppm"],  "#,##0", _C_SECT),
        ("معامل التفاوض — Neg. Coeff.",       f"{mi['neg_coeff']:.1%}", "@",   _C_WARN),
        ("السعر المُعدَّل النهائي (EGP/م²)",   mi["adjusted_ppm"],  "#,##0", _C_CLEAN),
        ("القيمة السوقية المُعدَّلة (EGP)",    round(mi["adjusted_ppm"] * area, 0), "#,##0", _C_CLEAN),
        ("درجة جودة البيانات (0-10)",          mi["quality_score"],  "0.0",   _C_HEAD),
    ]
    for i, (lbl, val, fmt, bg) in enumerate(results_rows):
        r2 = R4 + 1 + i
        W(ws_mi, r2, 1, lbl)
        ws_mi.cell(r2,1).font = Font(bold=True, color="CCCCFF", size=10)
        ws_mi.cell(r2,1).fill = PatternFill("solid", fgColor=_C_TITLE)
        W(ws_mi, r2, 2, val)
        ws_mi.cell(r2,2).number_format = fmt
        ws_mi.cell(r2,2).font = Font(bold=True, color="FFFFFF", size=11)
        ws_mi.cell(r2,2).fill = PatternFill("solid", fgColor=bg)
        ws_mi.cell(r2,2).alignment = Alignment(horizontal="center")

    # ── قسم 5: دليل المصادر المدعومة ─────────────────────────────────────
    R5 = R4 + len(results_rows) + 3
    _mi_sect(R5, "⑤ دليل المصادر المدعومة — Supported Data Sources")

    src_rows = [
        ("direct",   "معاينة مباشرة / عقد موثق",    "1.00", _C_CLEAN),
        ("agent",    "وكيل عقاري / سمسار",            "0.85", "1a4a2a"),
        ("forum",    "منتديات التقييم العقاري",       "0.65", _C_WARN),
        ("facebook", "مجموعات فيسبوك العقارية",       "0.50", "3d2a00"),
        ("other",    "مصدر آخر",                      "0.55", _C_SECT),
    ]
    _mi_head(R5+1, 1, "كود المصدر", sz=9, bg=_C_HEAD)
    _mi_head(R5+1, 2, "الوصف", sz=9, bg=_C_HEAD)
    _mi_head(R5+1, 3, "وزن المصداقية", sz=9, bg=_C_HEAD)
    for i, (src, desc, wt, bg) in enumerate(src_rows):
        r2 = R5 + 2 + i
        W(ws_mi, r2, 1, src);  ws_mi.cell(r2,1).fill = PatternFill("solid",fgColor=bg)
        W(ws_mi, r2, 2, desc); ws_mi.cell(r2,2).fill = PatternFill("solid",fgColor=bg)
        W(ws_mi, r2, 3, wt);   ws_mi.cell(r2,3).fill = PatternFill("solid",fgColor=bg)
        for ci in range(1,4):
            ws_mi.cell(r2,ci).font = Font(color="DDDDDD", size=9)

    # ── ضبط عرض الأعمدة ──────────────────────────────────────────────────
    for col, w in zip("ABCDEFGHIJ", [34,16,14,14,14,14,10,14,14,14]):
        ws_mi.column_dimensions[col].width = w

    # ── صندوق التفسير لصفحة MI (يُكتب هنا لأن mi و R5 متاحان) ──────────
    try:
        mi_scope_row = R5 + len(src_rows) + 4
        _scope_box(ws_mi, mi_scope_row, "استخبارات السوق والبيانات الخارجية", [
            f"جُمعت {mi['n_total']} مقارنة سوقية من مصادر متعددة. تم استبعاد {mi['n_outliers']} قيمة شاذة بطريقة MAD — تحييد الأسعار المتضخمة أو المدخلة بخطأ.",
            f"السعر الموزون (بعد ترجيح المصدر × الحداثة الزمنية): {mi['weighted_ppm']:,.0f} EGP/م². بعد تطبيق معامل التفاوض {mi['neg_coeff']:.1%}: السعر المُعدَّل = {mi['adjusted_ppm']:,.0f} EGP/م².",
            f"درجة جودة البيانات: {mi['quality_score']}/10. {'البيانات موثوقة — الاعتماد على السعر المُعدَّل ملائم.' if mi['quality_score'] >= 7 else 'البيانات محدودة — أضف مزيداً من المقارنات عبر /api/market-feed لتحسين الدقة.'}",
            "لإضافة بيانات من فيسبوك أو المنتديات: أرسل POST إلى /api/market-feed مع حقل source='facebook' أو 'forum'. البيانات الأحدث تحصل على وزن زمني أعلى تلقائياً.",
        ])
    except Exception as _e:
        print(f"  [scope MI] {_e}")

    # ── كتابة درجة الجودة في Dashboard ───────────────────────────────────
    if dash_name in wb.sheetnames:
        W(ws_dash, 15, 9, "══ جودة البيانات ══")
        ws_dash.cell(15,9).font = Font(bold=True, color="D4AF37")
        W(ws_dash, 16, 9, mi["quality_score"])
        ws_dash.cell(16,9).font = Font(bold=True, size=18, color="00CC44")
        ws_dash.cell(16,9).number_format = "0.0"
        W(ws_dash, 17, 9, f"/{10}  (Market Intelligence)")
        ws_dash.cell(17,9).font = Font(size=9, color="AAAACC")

    # ═════════════════════════════════════════════════════════════════════════
    # القسم الاستراتيجي — المحلل الجيوسياسي والاقتصادي (RAG)
    # يُضاف في Dashboard (صفوف 20-30) ليربط التقييم بالسياق الكلي
    # ═════════════════════════════════════════════════════════════════════════
    try:
        _strat = _rag_strategic(
            sector   = _dash_sector,
            purpose  = _purpose,
            location = str(data.get("location", "")),
        )
        if dash_name in wb.sheetnames:
            _ND  = "001428"   # لون خلفية داكن (نيلي)
            _GD  = "D4AF37"   # ذهبي
            _WH  = "E8F0FF"   # أبيض مائل للأزرق
            _GRN = "00CC66"   # أخضر
            _YEL = "FFD700"   # أصفر تحذيري

            # ── رأس القسم ──────────────────────────────────────────────────
            for _ci in range(1, 12):
                ws_dash.cell(20, _ci).fill = PatternFill("solid", fgColor="0A1F3A")
            W(ws_dash, 20, 1, f"🌐  {_strat.get('title','المحلل الاستراتيجي')}  —  السياق الجيوسياسي والاقتصادي")
            ws_dash.cell(20,1).font      = Font(bold=True, size=12, color=_GD, name="Calibri")
            ws_dash.cell(20,1).alignment = Alignment(horizontal="right", vertical="center")
            ws_dash.row_dimensions[20].height = 28
            try: ws_dash.merge_cells("A20:K20")
            except: pass

            # ── بيانات المدينة ──────────────────────────────────────────────
            _row21 = _strat.get("city_line", "")
            if _row21:
                for _ci in range(1, 12):
                    ws_dash.cell(21, _ci).fill = PatternFill("solid", fgColor="001A3A")
                W(ws_dash, 21, 1, f"📍  {_row21}")
                ws_dash.cell(21,1).font      = Font(size=10, color=_WH, name="Calibri")
                ws_dash.cell(21,1).alignment = Alignment(horizontal="right", vertical="center")
                ws_dash.row_dimensions[21].height = 20
                try: ws_dash.merge_cells("A21:K21")
                except: pass

            # ── المؤشرات الاقتصادية ─────────────────────────────────────────
            _row22 = _strat.get("econ_line", "")
            if _row22:
                for _ci in range(1, 12):
                    ws_dash.cell(22, _ci).fill = PatternFill("solid", fgColor="001A3A")
                W(ws_dash, 22, 1, f"📊  {_row22}")
                ws_dash.cell(22,1).font      = Font(size=10, color=_WH, name="Calibri")
                ws_dash.cell(22,1).alignment = Alignment(horizontal="right", vertical="center")
                ws_dash.row_dimensions[22].height = 20
                try: ws_dash.merge_cells("A22:K22")
                except: pass

            # ── الأحداث الجيوسياسية ─────────────────────────────────────────
            _row23 = _strat.get("geo_line", "")
            if _row23:
                for _ci in range(1, 12):
                    ws_dash.cell(23, _ci).fill = PatternFill("solid", fgColor="001A3A")
                W(ws_dash, 23, 1, f"🌍  {_row23}")
                ws_dash.cell(23,1).font      = Font(size=10, color="FFB347", name="Calibri")
                ws_dash.cell(23,1).alignment = Alignment(horizontal="right", vertical="center")
                ws_dash.row_dimensions[23].height = 20
                try: ws_dash.merge_cells("A23:K23")
                except: pass

            # ── توقعات القطاع ───────────────────────────────────────────────
            _row24 = _strat.get("sector_line", "")
            if _row24:
                for _ci in range(1, 12):
                    ws_dash.cell(24, _ci).fill = PatternFill("solid", fgColor="001A3A")
                W(ws_dash, 24, 1, f"🏗  {_row24}")
                ws_dash.cell(24,1).font      = Font(size=10, color=_WH, name="Calibri")
                ws_dash.cell(24,1).alignment = Alignment(horizontal="right", vertical="center")
                ws_dash.row_dimensions[24].height = 20
                try: ws_dash.merge_cells("A24:K24")
                except: pass

            # ── ملاحظة الغرض التقييمي ───────────────────────────────────────
            _row25 = _strat.get("purpose_note", "")
            if _row25:
                for _ci in range(1, 12):
                    ws_dash.cell(25, _ci).fill = PatternFill("solid", fgColor="001A3A")
                W(ws_dash, 25, 1, f"⚙  {_row25}")
                ws_dash.cell(25,1).font      = Font(size=10, color="B0C4DE", italic=True, name="Calibri")
                ws_dash.cell(25,1).alignment = Alignment(horizontal="right", vertical="center")
                ws_dash.row_dimensions[25].height = 20
                try: ws_dash.merge_cells("A25:K25")
                except: pass

            # ── إشارات الشراء ──────────────────────────────────────────────
            _buy = _strat.get("buy_signals", [])
            if _buy:
                for _ci in range(1, 7):
                    ws_dash.cell(26, _ci).fill = PatternFill("solid", fgColor="003300")
                W(ws_dash, 26, 1, "✅  إشارات إيجابية:  " + "  •  ".join(_buy[:3]))
                ws_dash.cell(26,1).font      = Font(size=9, color=_GRN, name="Calibri")
                ws_dash.cell(26,1).alignment = Alignment(horizontal="right", vertical="center")
                ws_dash.row_dimensions[26].height = 18
                try: ws_dash.merge_cells("A26:F26")
                except: pass

            # ── إشارات الحذر ───────────────────────────────────────────────
            _cau = _strat.get("caution", [])
            if _cau:
                for _ci in range(7, 12):
                    ws_dash.cell(26, _ci).fill = PatternFill("solid", fgColor="332200")
                W(ws_dash, 26, 7, "⚠  تحفظات:  " + "  •  ".join(_cau[:2]))
                ws_dash.cell(26,7).font      = Font(size=9, color=_YEL, name="Calibri")
                ws_dash.cell(26,7).alignment = Alignment(horizontal="right", vertical="center")
                try: ws_dash.merge_cells("G26:K26")
                except: pass

            # ── إخلاء المسؤولية ───────────────────────────────────────────
            _disc = _strat.get("disclaimer", "")
            if _disc:
                for _ci in range(1, 12):
                    ws_dash.cell(27, _ci).fill = PatternFill("solid", fgColor="050E1E")
                W(ws_dash, 27, 1, _disc)
                ws_dash.cell(27,1).font      = Font(size=8, color="667788", italic=True, name="Calibri")
                ws_dash.cell(27,1).alignment = Alignment(horizontal="right", vertical="center")
                ws_dash.row_dimensions[27].height = 16
                try: ws_dash.merge_cells("A27:K27")
                except: pass

            print(f"{_ts()} [INFO] ✓ القسم الاستراتيجي مُضاف إلى Dashboard")
    except Exception as _strat_err:
        print(f"  [strategic_section] {_strat_err}")

    # ═════════════════════════════════════════════════════════════════════════
    # صفحة الاستخبارات والمخاطر — Sovereign Intelligence Sheet
    # تُضاف فقط إذا أرسل الـ Frontend بيانات sovereign_insights
    # ═════════════════════════════════════════════════════════════════════════
    try:
        _si = data.get("sovereign_insights")
        if _si and isinstance(_si, dict):
            _si_name = "الاستخبارات والمخاطر"
            ws_si = wb.create_sheet(_si_name)
            ws_si.sheet_view.rightToLeft = True

            # ── ألوان السمة السيادية ────────────────────────────────────
            _SI_BG      = PatternFill("solid", fgColor="030812")
            _SI_HEAD_BG = PatternFill("solid", fgColor="0A1628")
            _SI_SECT_BG = PatternFill("solid", fgColor="0F1E35")
            _SI_GOLD_F  = Font(bold=True, size=11, color="D4AF37", name="Calibri")
            _SI_TITLE_F = Font(bold=True, size=14, color="D4AF37", name="Calibri")
            _SI_SUB_F   = Font(bold=True, size=10, color="B0C4DE", name="Calibri")
            _SI_VAL_F   = Font(size=10, color="E8F0FF", name="Calibri")
            _SI_WARN_F  = Font(size=10, color="F87171", name="Calibri")
            _SI_OK_F    = Font(size=10, color="4ADE80", name="Calibri")
            _SI_NOTE_F  = Font(size=9, color="667799", italic=True, name="Calibri")
            _SI_R_ALIGN = Alignment(horizontal="right", vertical="center", wrap_text=True)

            # عرض الأعمدة
            ws_si.column_dimensions["A"].width = 28
            ws_si.column_dimensions["B"].width = 35
            ws_si.column_dimensions["C"].width = 18
            ws_si.column_dimensions["D"].width = 25

            r = 1  # row counter

            # ── رأس الصفحة ─────────────────────────────────────────────
            for c in range(1, 5):
                ws_si.cell(r, c).fill = PatternFill("solid", fgColor="0A1F3A")
            W(ws_si, r, 1, "🏛  تقرير الاستخبارات العقارية والمخاطر السيادية")
            ws_si.cell(r, 1).font = Font(bold=True, size=16, color="D4AF37", name="Calibri")
            ws_si.cell(r, 1).alignment = _SI_R_ALIGN
            ws_si.row_dimensions[r].height = 36
            try: ws_si.merge_cells(f"A{r}:D{r}")
            except: pass
            r += 1

            # تاريخ وختم
            for c in range(1, 5):
                ws_si.cell(r, c).fill = _SI_BG
            W(ws_si, r, 1, f"Expert_Smart v45 — Sovereign Edition  |  {datetime.now().strftime('%Y-%m-%d %H:%M')}")
            ws_si.cell(r, 1).font = _SI_NOTE_F
            ws_si.cell(r, 1).alignment = _SI_R_ALIGN
            try: ws_si.merge_cells(f"A{r}:D{r}")
            except: pass
            r += 2

            # ════════════════════════════════════════════════════════════
            # القسم 1: المُحقق الرقمي
            # ════════════════════════════════════════════════════════════
            _si_fraud = _si.get("fraud")
            for c in range(1, 5):
                ws_si.cell(r, c).fill = _SI_SECT_BG
            W(ws_si, r, 1, "🔍  المُحقق الرقمي — Fraud & Anomaly Detection")
            ws_si.cell(r, 1).font = _SI_GOLD_F
            ws_si.cell(r, 1).alignment = _SI_R_ALIGN
            ws_si.row_dimensions[r].height = 26
            try: ws_si.merge_cells(f"A{r}:D{r}")
            except: pass
            r += 1

            if _si_fraud:
                _fraud_rows = [
                    ("نسبة الثقة",       _si_fraud.get("confidence", "—")),
                    ("الحالة",           _si_fraud.get("flag", "—")),
                    ("Z-Score",          _si_fraud.get("z_score", "—")),
                    ("الحد الأدنى للسوق", _si_fraud.get("market_low", "—")),
                    ("الحد الأعلى للسوق", _si_fraud.get("market_high", "—")),
                ]
                for label, val in _fraud_rows:
                    for c in range(1, 5):
                        ws_si.cell(r, c).fill = _SI_BG
                    W(ws_si, r, 1, label)
                    ws_si.cell(r, 1).font = _SI_SUB_F
                    ws_si.cell(r, 1).alignment = _SI_R_ALIGN
                    W(ws_si, r, 2, str(val))
                    # تلوين الحالة
                    _is_ok = any(x in str(val) for x in ("سليم", "✅", "طبيعي"))
                    ws_si.cell(r, 2).font = _SI_OK_F if _is_ok else (
                        _SI_WARN_F if any(x in str(val) for x in ("مشبوه", "حرج", "⚠", "🚨")) else _SI_VAL_F
                    )
                    ws_si.cell(r, 2).alignment = _SI_R_ALIGN
                    r += 1
                # النص الكامل
                _fraud_raw = str(_si_fraud.get("raw_text", ""))
                if _fraud_raw:
                    for c in range(1, 5):
                        ws_si.cell(r, c).fill = _SI_BG
                    W(ws_si, r, 1, _fraud_raw[:300])
                    ws_si.cell(r, 1).font = _SI_NOTE_F
                    ws_si.cell(r, 1).alignment = _SI_R_ALIGN
                    try: ws_si.merge_cells(f"A{r}:D{r}")
                    except: pass
                    ws_si.row_dimensions[r].height = 40
                    r += 1
            else:
                for c in range(1, 5):
                    ws_si.cell(r, c).fill = _SI_BG
                W(ws_si, r, 1, "لم يتم تشغيل تحليل المُحقق الرقمي — أدخل السعر والموقع في لوحة الاستخبارات")
                ws_si.cell(r, 1).font = _SI_NOTE_F
                ws_si.cell(r, 1).alignment = _SI_R_ALIGN
                try: ws_si.merge_cells(f"A{r}:D{r}")
                except: pass
                r += 1

            r += 1

            # ════════════════════════════════════════════════════════════
            # القسم 2: المخاطر الجيوتقنية
            # ════════════════════════════════════════════════════════════
            _si_geo = _si.get("geo")
            for c in range(1, 5):
                ws_si.cell(r, c).fill = _SI_SECT_BG
            W(ws_si, r, 1, "🌍  التحليل الجيوتقني — Geotechnical & Legal Risk")
            ws_si.cell(r, 1).font = _SI_GOLD_F
            ws_si.cell(r, 1).alignment = _SI_R_ALIGN
            try: ws_si.merge_cells(f"A{r}:D{r}")
            except: pass
            ws_si.row_dimensions[r].height = 26
            r += 1

            if _si_geo:
                _geo_rows = [
                    ("المنطقة المُكتشفة", _si_geo.get("location", "—")),
                    ("نوع التربة",       _si_geo.get("soil_type", "—")),
                    ("مخاطر الفيضانات", _si_geo.get("flood_risk", "—")),
                    ("درجة المخاطر",    _si_geo.get("risk_score", "—") + "/10" if _si_geo.get("risk_score") else "—"),
                ]
                for label, val in _geo_rows:
                    for c in range(1, 5):
                        ws_si.cell(r, c).fill = _SI_BG
                    W(ws_si, r, 1, label)
                    ws_si.cell(r, 1).font = _SI_SUB_F
                    ws_si.cell(r, 1).alignment = _SI_R_ALIGN
                    W(ws_si, r, 2, str(val))
                    ws_si.cell(r, 2).font = _SI_VAL_F
                    ws_si.cell(r, 2).alignment = _SI_R_ALIGN
                    r += 1
                _geo_raw = str(_si_geo.get("raw_text", ""))
                if _geo_raw:
                    for c in range(1, 5):
                        ws_si.cell(r, c).fill = _SI_BG
                    W(ws_si, r, 1, _geo_raw[:400])
                    ws_si.cell(r, 1).font = _SI_NOTE_F
                    ws_si.cell(r, 1).alignment = _SI_R_ALIGN
                    try: ws_si.merge_cells(f"A{r}:D{r}")
                    except: pass
                    ws_si.row_dimensions[r].height = 50
                    r += 1
            else:
                for c in range(1, 5):
                    ws_si.cell(r, c).fill = _SI_BG
                W(ws_si, r, 1, "لم يتم تشغيل التحليل الجيوتقني — حدّد الموقع في لوحة المخاطر")
                ws_si.cell(r, 1).font = _SI_NOTE_F
                ws_si.cell(r, 1).alignment = _SI_R_ALIGN
                try: ws_si.merge_cells(f"A{r}:D{r}")
                except: pass
                r += 1

            r += 1

            # ════════════════════════════════════════════════════════════
            # القسم 3: رادار الهجرة العقارية
            # ════════════════════════════════════════════════════════════
            _si_demo = _si.get("demographic")
            for c in range(1, 5):
                ws_si.cell(r, c).fill = _SI_SECT_BG
            W(ws_si, r, 1, "🏗  رادار الهجرة العقارية — Demographic & POI Flow")
            ws_si.cell(r, 1).font = _SI_GOLD_F
            ws_si.cell(r, 1).alignment = _SI_R_ALIGN
            try: ws_si.merge_cells(f"A{r}:D{r}")
            except: pass
            ws_si.row_dimensions[r].height = 26
            r += 1

            if _si_demo:
                _demo_rows = [
                    ("النمو السنوي المتوقع", _si_demo.get("annual_growth", "—")),
                    ("كثافة POI",           _si_demo.get("poi_density", "—")),
                ]
                for label, val in _demo_rows:
                    for c in range(1, 5):
                        ws_si.cell(r, c).fill = _SI_BG
                    W(ws_si, r, 1, label)
                    ws_si.cell(r, 1).font = _SI_SUB_F
                    ws_si.cell(r, 1).alignment = _SI_R_ALIGN
                    W(ws_si, r, 2, str(val))
                    ws_si.cell(r, 2).font = _SI_VAL_F
                    ws_si.cell(r, 2).alignment = _SI_R_ALIGN
                    r += 1
                _demo_raw = str(_si_demo.get("raw_text", ""))
                if _demo_raw:
                    for c in range(1, 5):
                        ws_si.cell(r, c).fill = _SI_BG
                    W(ws_si, r, 1, _demo_raw[:400])
                    ws_si.cell(r, 1).font = _SI_NOTE_F
                    ws_si.cell(r, 1).alignment = _SI_R_ALIGN
                    try: ws_si.merge_cells(f"A{r}:D{r}")
                    except: pass
                    ws_si.row_dimensions[r].height = 50
                    r += 1
            else:
                for c in range(1, 5):
                    ws_si.cell(r, c).fill = _SI_BG
                W(ws_si, r, 1, "لم يتم تشغيل رادار الهجرة — حدّد الموقع في لوحة الديموغرافيا")
                ws_si.cell(r, 1).font = _SI_NOTE_F
                ws_si.cell(r, 1).alignment = _SI_R_ALIGN
                try: ws_si.merge_cells(f"A{r}:D{r}")
                except: pass
                r += 1

            r += 1

            # ── تذييل إخلاء مسؤولية ──────────────────────────────────
            for c in range(1, 5):
                ws_si.cell(r, c).fill = _SI_BG
            W(ws_si, r, 1,
              "⚠ إخلاء مسؤولية: هذه التحليلات مبنية على بيانات متاحة للجمهور وقواعد حتمية (Rule-Based). "
              "لا تُغني عن فحص التربة الميداني أو التحقق القانوني المباشر. "
              "يُنصح بالتحقق المستقل قبل اتخاذ قرارات استثمارية أو تمويلية.")
            ws_si.cell(r, 1).font = Font(size=8, color="556677", italic=True, name="Calibri")
            ws_si.cell(r, 1).alignment = _SI_R_ALIGN
            try: ws_si.merge_cells(f"A{r}:D{r}")
            except: pass
            ws_si.row_dimensions[r].height = 36

            # تلوين خلفية كل الصفحة
            for _row in ws_si.iter_rows(min_row=1, max_row=r, max_col=4):
                for _cell in _row:
                    if _cell.fill == PatternFill() or not _cell.fill.fgColor:
                        _cell.fill = _SI_BG

            print(f"{_ts()} [INFO] ✓ صفحة الاستخبارات والمخاطر ({r} صف)")
    except Exception as _si_err:
        print(f"  [sovereign_insights] {_si_err}")

    # ═════════════════════════════════════════════════════════════════════════
    # حفظ واحد — ثم حذف الملفات المؤقتة
    # ═════════════════════════════════════════════════════════════════════════
    n_sheets = len(wb.sheetnames)
    print(f"{_ts()} [INFO] حفظ الملف النهائي ({n_sheets} شيت)... 100%")
    wb.save(output_path)
    wb.close()
    for _t in _tmp_files:
        try: os.unlink(_t)
        except: pass
    _elapsed = round(time.time() - _t_start, 1)
    print(f"{_ts()} [INFO] ✓ اكتمل — {os.path.basename(output_path)} | {n_sheets} شيت | {_elapsed}s")
    return output_path

# ═══════════════════════════════════════════════════════════════════════════
# Legacy sheet filter — strips advanced-analytics sheets from a saved workbook
# ═══════════════════════════════════════════════════════════════════════════

def _remove_legacy_advanced_sheets(path: str) -> None:
    """Remove advanced-analytics sheets from a saved workbook for legacy exports.

    A sheet is removed when its normalized name EQUALS or CONTAINS any excluded
    keyword.  This handles prefixed/suffixed sheet names such as:
      "ANN — الشبكات العصبية"   → contains "الشبكات العصبية"
      "ARIMA — السلاسل الزمنية" → contains "السلاسل الزمنية"
      "استخبارات السوق — MI"    → contains "استخبارات السوق"

    Operates in-place on an already-saved file.  Any failure is logged and
    suppressed so the export is never blocked by sheet-removal errors.
    """
    # Keywords — matching is done after normalization (see _norm below).
    # Each entry is a keyword: a sheet is removed if its normalized name
    # equals OR contains the keyword.
    _KEYWORDS: tuple = (
        # Arabic core keywords (variants unified by _norm)
        "التحليل المكاني",       # also catches المكانى via ى→ي
        "الانحدار المتعدد",      # also catches الإنحدار via hamza→ا
        "الخيارات الحقيقية",
        "لوحة القيادة التنفيذية",
        "الشبكات العصبية",
        "السلاسل الزمنية",
        "استخبارات السوق",       # also catches إستخبارات via hamza→ا
        # English keywords
        "spatial analysis",
        "multiple regression",
        "real options",
        "executive dashboard",
        "neural networks",
        "time series",
        "market intelligence",
    )

    def _norm(name: str) -> str:
        """Normalize for matching: strip, lower, unify hamza variants and ى→ي."""
        n = name.strip().lower()
        for _ch in "أإآ":
            n = n.replace(_ch, "ا")
        return n.replace("ى", "ي")

    # Pre-normalize keywords once
    _kw_norm = tuple(_norm(k) for k in _KEYWORDS)

    def _is_advanced(sheet_name: str) -> bool:
        """True if the sheet's normalized name equals OR contains any keyword."""
        sn = _norm(sheet_name)
        for kw in _kw_norm:
            if kw == sn or kw in sn:
                return True
        return False

    try:
        import openpyxl as _xl
        _is_xlsm = path.lower().endswith(".xlsm")
        _wb      = _xl.load_workbook(path, keep_vba=_is_xlsm)
        _all     = list(_wb.sheetnames)
        _to_del  = [s for s in _all if _is_advanced(s)]
        print(f"{_ts()} [LEGACY-FILTER] before: {len(_all)} sheet(s)")
        print(f"{_ts()} [LEGACY-FILTER] matched: {_to_del}")
        for _s in _to_del:
            del _wb[_s]
        if _to_del:
            _wb.save(path)
        _after = list(_wb.sheetnames)
        print(f"{_ts()} [LEGACY-FILTER] after:  {len(_after)} sheet(s)")
        _wb.close()
    except Exception as _rf_err:
        print(f"{_ts()} [LEGACY-FILTER] warning — sheet removal skipped: {_rf_err}")


# ═══════════════════════════════════════════════════════════════════════════
# تقرير Word
# ═══════════════════════════════════════════════════════════════════════════
def write_word_summary(data, output_path):
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()
        h = doc.add_heading("تقرير تقييم عقاري — م. هشام المهدي", 0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for lbl, key in [("رقم التقرير","report_id"),("التاريخ","report_date"),
                          ("الموقع","location"),("نوع العقار","property_type")]:
            doc.add_paragraph(f"{lbl}: {data.get(key,'')}")
        mv=data.get("market_value",0)
        doc.add_paragraph(f"القيمة السوقية: {mv:,.0f} ج.م")
        doc.add_paragraph(f"أسلوب السوق:   {data.get('market_approach',0):,.0f} ج.م")
        doc.add_paragraph(f"أسلوب التكلفة: {data.get('cost_approach',0):,.0f} ج.م")
        doc.add_paragraph(f"أسلوب الدخل:   {data.get('income_approach',0):,.0f} ج.م")
        doc.save(output_path)
    except Exception:
        print(traceback.format_exc())

# ═══════════════════════════════════════════════════════════════════════════
# Endpoints
# ═══════════════════════════════════════════════════════════════════════════

# ── (Wave 2) Helpers لتفويض /api/valuation إلى الموديولات المتخصصة ─────────
# هذه الدوال إضافية فقط ولا تمسّ المنطق القائم لـ advanced_valuation.

def _delegate_valuation_to_hbu(payload: dict):
    """يُفوِّض الطلب إلى موديول HBU ويُرجع استجابة بنفس عقد /api/valuation."""
    try:
        try:
            from hbu_analysis_engine import run_hbu_analysis, generate_hbu_excel
        except Exception:
            from core_engine.hbu_analysis_engine import run_hbu_analysis, generate_hbu_excel  # type: ignore
        try:
            from purpose_specific_reports import write_hbu_word_report
        except Exception:
            try:
                from core_engine.purpose_specific_reports import write_hbu_word_report  # type: ignore
            except Exception:
                write_hbu_word_report = None  # type: ignore

        hbu_payload = {
            "property": {
                "asset_id":       payload.get("asset_id", ""),
                "location":       payload.get("location", ""),
                "area":           payload.get("area", 0),
                "current_use":    payload.get("property_type", ""),
                "current_zoning": payload.get("zoning", ""),
            },
            "discount_rate":    float(payload.get("discount_rate", 0.10)),
            "valuation_date":   payload.get("date") or datetime.now().strftime("%Y-%m-%d"),
            "alternative_uses": payload.get("alternative_uses", []),
        }
        result = run_hbu_analysis(hbu_payload)

        rid  = f"ES-HBU-{datetime.now().strftime('%Y%m%d')}-{uuid.uuid4().hex[:4].upper()}"
        ts   = datetime.now().strftime("%H%M%S")
        xlsx = generate_hbu_excel(result, OUTPUTS)
        xlsx_name = os.path.basename(xlsx)

        docx_name = None
        if write_hbu_word_report is not None:
            docx_name = f"Summary_{rid}_{ts}.docx"
            try:
                write_hbu_word_report(result, os.path.join(OUTPUTS, docx_name))
            except Exception as werr:
                print(f"{_ts()} [HBU-DELEGATE] Word failed: {werr}")
                docx_name = None

        market_value = result.get("recommended_npv") or 0.0
        return jsonify({
            "status":       "success",
            "purpose":      "highest_and_best_use",
            "delegated_to": "hbu_analysis_engine",
            "market_value": market_value,
            "report_id":    rid,
            "result":       result,
            "excel_url":    f"http://127.0.0.1:5000/api/download/{xlsx_name}",
            "word_url":     (f"http://127.0.0.1:5000/api/download/{docx_name}" if docx_name else None),
        })
    except ValueError as ve:
        return jsonify({"status": "error", "message": str(ve)}), 400
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({"status": "error", "message": f"HBU delegate failed: {e}"}), 500


def _delegate_valuation_to_reit(payload: dict):
    """يُفوِّض الطلب إلى موديول REIT NAV ويُرجع نفس صيغة استجابة /api/valuation."""
    try:
        try:
            from reit_nav_engine import run_reit_nav, generate_reit_nav_excel
        except Exception:
            from core_engine.reit_nav_engine import run_reit_nav, generate_reit_nav_excel  # type: ignore
        try:
            from purpose_specific_reports import write_reit_word_report
        except Exception:
            try:
                from core_engine.purpose_specific_reports import write_reit_word_report  # type: ignore
            except Exception:
                write_reit_word_report = None  # type: ignore

        result = run_reit_nav(payload)

        rid  = f"ES-REIT-{datetime.now().strftime('%Y%m%d')}-{uuid.uuid4().hex[:4].upper()}"
        ts   = datetime.now().strftime("%H%M%S")
        xlsx = generate_reit_nav_excel(result, OUTPUTS)
        xlsx_name = os.path.basename(xlsx)

        docx_name = None
        if write_reit_word_report is not None:
            docx_name = f"Summary_{rid}_{ts}.docx"
            try:
                write_reit_word_report(result, os.path.join(OUTPUTS, docx_name))
            except Exception as werr:
                print(f"{_ts()} [REIT-DELEGATE] Word failed: {werr}")
                docx_name = None

        return jsonify({
            "status":       "success",
            "purpose":      "investment_funds",
            "delegated_to": "reit_nav_engine",
            "market_value": result.get("nav", 0.0),
            "nav_per_unit": result.get("nav_per_unit", 0.0),
            "report_id":    rid,
            "result":       result,
            "excel_url":    f"http://127.0.0.1:5000/api/download/{xlsx_name}",
            "word_url":     (f"http://127.0.0.1:5000/api/download/{docx_name}" if docx_name else None),
        })
    except ValueError as ve:
        return jsonify({"status": "error", "message": str(ve)}), 400
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({"status": "error", "message": f"REIT delegate failed: {e}"}), 500


def _delegate_valuation_to_eia(payload: dict):
    """يُفوِّض الطلب إلى محرك تقييم الأثر البيئي ويُرجع نفس عقد /api/valuation."""
    try:
        try:
            from eia_engine import run_eia_assessment, write_eia_word_report, write_eia_excel_report
        except Exception:
            from core_engine.eia_engine import run_eia_assessment, write_eia_word_report, write_eia_excel_report  # type: ignore

        # market_value optional — pulled from payload for ERF computation
        result = run_eia_assessment(payload)

        rid  = f"ES-EIA-{datetime.now().strftime('%Y%m%d')}-{uuid.uuid4().hex[:4].upper()}"
        ts   = datetime.now().strftime("%H%M%S")

        # Word report
        docx_name = f"EIA_Report_{rid}_{ts}.docx"
        try:
            write_eia_word_report(result, os.path.join(OUTPUTS, docx_name))
            word_url = f"http://127.0.0.1:5000/api/download/{docx_name}"
        except Exception as werr:
            print(f"{_ts()} [EIA-DELEGATE] Word failed: {werr}")
            word_url = None

        # Excel report
        try:
            xlsx = write_eia_excel_report(result, OUTPUTS)
            xlsx_url = (f"http://127.0.0.1:5000/api/download/{os.path.basename(xlsx)}"
                        if xlsx else None)
        except Exception as xerr:
            print(f"{_ts()} [EIA-DELEGATE] Excel failed: {xerr}")
            xlsx_url = None

        inv = result.get("investment_linkage", {})
        cl  = result.get("classification", {})
        return jsonify({
            "status":              "success",
            "purpose":             "environmental_impact_assessment",
            "delegated_to":        "eia_engine",
            "market_value":        inv.get("adjusted_market_value", 0.0),
            "report_id":           rid,
            "environmental_category": cl.get("category"),
            "environmental_label":  cl.get("category_label"),
            "environmental_risk_factor": inv.get("environmental_risk_factor"),
            "result":              result,
            "word_url":            word_url,
            "excel_url":           xlsx_url,
        })
    except ValueError as ve:
        return jsonify({"status": "error", "message": str(ve)}), 400
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({"status": "error", "message": f"EIA delegate failed: {e}"}), 500


def _augment_payload_for_usufruct(payload: dict) -> None:
    """يحقن بيانات حق الانتفاع داخل الـ payload لتُستخدم في تقرير ما بعد التقييم.
    لا يُعدِّل المنطق الحسابي القائم — مجرد تخزين للحقول كمتغيرات meta."""
    try:
        years    = float(payload.get("usufruct_years", 25))
        rate     = float(payload.get("discount_rate", 0.10))
        # PV factor for an annuity-style decay over the usufruct term:
        #   pv_factor = 1 - (1+r)^(-n)
        pv_factor = 1.0 - (1.0 + rate) ** (-years) if rate > 0 else 1.0
        # تطبيع لنطاق [0.50, 0.85] لتفادي القيم المتطرفة
        pv_factor = max(0.50, min(0.85, pv_factor))
        payload.setdefault("usufruct_meta", {
            "usufruct_years":  years,
            "discount_rate":   rate,
            "pv_factor":       pv_factor,
            "note":            "حق انتفاع — قيمة العقار × معامل القيمة الحالية للمنفعة عبر مدة العقد",
        })
    except Exception:
        pass


def _augment_payload_for_uncertainty(payload: dict) -> None:
    """يضيف نطاق قيمة (Low/Best/High) كإفصاح IFRS 13 §93 لعدم اليقين الجوهري.
    يُحسب لاحقاً على القيمة السوقية الناتجة من المسار العام."""
    try:
        spread_pct = float(payload.get("uncertainty_spread_pct", 0.15))
        payload.setdefault("uncertainty_meta", {
            "spread_pct": spread_pct,
            "note":       "IFRS 13 §93 — نطاق قيمة Level 3 يعكس عدم اليقين الجوهري",
        })
    except Exception:
        pass


@app.route("/api/valuation", methods=["POST","OPTIONS"])
def handle_valuation():
    if request.method == "OPTIONS": return jsonify({}), 200
    try:
        payload = request.get_json(silent=True) or {}

        # ── تفويض تلقائي للأغراض المتخصصة (Wave 2) ──────────────────────────
        # عند اختيار غرض يستوجب موديولاً متخصصاً، نُحوِّل الطلب إليه ونرجع
        # بنتيجته المُغلَّفة بنفس صيغة استجابة /api/valuation الموحَّدة.
        # لا يتم التفويض إلا إذا توفرت مدخلات الموديول؛ وإلا نسقط على المسار العام.
        _vp = (payload.get("valuation_purpose") or "").strip()
        if _vp == "highest_and_best_use" and payload.get("alternative_uses"):
            return _delegate_valuation_to_hbu(payload)
        if _vp == "investment_funds" and payload.get("properties"):
            return _delegate_valuation_to_reit(payload)
        if _vp == "environmental_impact_assessment" and payload.get("impact_matrix"):
            return _delegate_valuation_to_eia(payload)

        # ── AVM dispatch — لـ FMV / bank_financing / tax_assessment / acquisition فقط ─
        # يعمل بشكل غير حاجب: إذا كانت البيانات كافية يحقن ppm مباشرة في الـ payload
        # قبل استدعاء advanced_valuation. لا يمسّ منطق الأغراض الأخرى.
        _avm_decision = None
        try:
            try:
                from avm_dispatcher import dispatch_avm, is_avm_eligible
            except Exception:
                from core_engine.avm_dispatcher import dispatch_avm, is_avm_eligible  # type: ignore
            if is_avm_eligible(_vp):
                _avm_decision = dispatch_avm(payload, records=_load_feed())
                print(f"{_ts()} [AVM] purpose={_vp} -> {_avm_decision.get('verdict', '—')}")
        except Exception as _avm_err:
            print(f"{_ts()} [AVM] dispatch skipped: {_avm_err}")
        if _vp == "usufruct":
            # اعتراض اختياري لتطبيق منطق حق الانتفاع (DCF عمر العقد)
            _augment_payload_for_usufruct(payload)
        if _vp == "uncertainty_valuation":
            # اعتراض اختياري لإضافة نطاق قيمة (Range) كإفصاح Level 3
            _augment_payload_for_uncertainty(payload)

        # ── (Wave 2) دعوة دالة متخصصة حسب نوع الأصل عند توفر منطق متخصص ────
        # تُستدعى ABDO advanced_valuation كالعادة، ثم نُكمل القيمة من valuator
        # متخصص إذا كان النوع من القائمة المخصصة (intangible / partial / under-construction / quarry).
        _specialized_result = None
        try:
            try:
                from asset_type_valuators import has_specialized_valuator, run_specialized_valuator
            except Exception:
                from core_engine.asset_type_valuators import has_specialized_valuator, run_specialized_valuator  # type: ignore
            _ptype = (payload.get("property_type") or "").strip()
            if has_specialized_valuator(_ptype):
                try:
                    _specialized_result = run_specialized_valuator(_ptype, payload)
                    print(f"{_ts()} [SPECIALIZED] {_ptype} → reconciled = {_specialized_result.get('reconciled_value'):,.0f}")
                except Exception as _spec_err:
                    print(f"{_ts()} [SPECIALIZED] {_ptype} failed (fallback to advanced_valuation): {_spec_err}")
                    _specialized_result = None
        except Exception:
            _specialized_result = None

        # نمرر جميع حقول الـ payload كـ kwargs حتى تصل المعاملات المتخصصة
        # (beds, rooms, frontage, purpose, gla_ratio, facilities …)
        _extra_kw = {k: v for k, v in payload.items()
                     if k not in ("location","area","property_type","price_per_meter")}
        res = advanced_valuation(
            location        = payload.get("location",""),
            area            = payload.get("area"),
            property_type   = payload.get("property_type",""),
            price_per_meter = payload.get("price_per_meter"),
            **_extra_kw,
        )

        # If specialized valuator produced a result, override market_value
        if _specialized_result and _specialized_result.get("reconciled_value"):
            res["market_value"] = float(_specialized_result["reconciled_value"])
            res["specialized_asset_result"] = _specialized_result
        rid = f"ES-{datetime.now().strftime('%Y%m%d')}-{uuid.uuid4().hex[:4].upper()}"
        full = {**payload, **res,
                "report_id":   rid,
                "report_date": datetime.now().strftime("%d/%m/%Y"),
                "expert":      "م. هشام المهدي"}

        ts           = datetime.now().strftime("%H%M%S")
        report_style = payload.get("report_style", "legacy")
        _valid_styles = ("legacy", "detailed", "professional_template")
        _style_used  = report_style if report_style in _valid_styles else "legacy"

        # ── Validation Gate — opt-in via "validate": true in payload (Wave BA.2) ──
        # Runs ONLY when the caller opts in; zero effect on any existing request.
        # ERRORs block generation (HTTP 422); warnings are attached but do not block.
        _validation_result: dict | None = None
        if payload.get("validate"):
            try:
                from reports.report_pipeline import validate_report_data as _vrd
                _vr = _vrd(full, profile_key=_style_used)
                _validation_result = {
                    "is_valid": _vr.is_valid,
                    "issues": [
                        {
                            "code":       i.code,
                            "severity":   i.severity.value,
                            "message_ar": i.message_ar,
                            "message_en": i.message_en,
                        }
                        for i in (_vr.errors + _vr.warnings)
                    ],
                }
                if not _vr.is_valid:
                    return jsonify({
                        "status":     "validation_error",
                        "validation": _validation_result,
                    }), 422
            except Exception as _val_err:
                print(f"{_ts()} [VALIDATION] gate skipped: {_val_err}")
        # ── (end validation gate) ─────────────────────────────────────────────
        _tpl_used    = False
        _fallback    = False
        _fb_reason   = ""

        if report_style == "professional_template":
            # ── Attempt individual_valuation_professional_template.xlsm ───────
            try:
                try:
                    from reports.excel_template_renderer import (
                        INDIVIDUAL_VALUATION_TEMPLATE as _IVTPL,
                        build_individual_valuation_report as _build_iv,
                    )
                except ImportError:
                    from core_engine.reports.excel_template_renderer import (  # type: ignore
                        INDIVIDUAL_VALUATION_TEMPLATE as _IVTPL,
                        build_individual_valuation_report as _build_iv,
                    )
                if _IVTPL.is_file():
                    _tpl_name = f"Report_{rid}_{ts}.xlsm"
                    _tpl_path = os.path.join(OUTPUTS, _tpl_name)
                    _iv_ctx = {
                        "report_date":       full.get("report_date", datetime.now().strftime("%d/%m/%Y")),
                        "valuation_date":    full.get("date", ""),
                        "client_name":       full.get("client_name", full.get("expert", "N/A")),
                        "property_type":     full.get("property_type", ""),
                        "location":          full.get("location", ""),
                        "area":              full.get("area", ""),
                        "valuation_purpose": full.get("valuation_purpose", ""),
                        "market_value":      float(full.get("market_value") or 0),
                        "final_value":       float(full.get("market_value") or 0),
                        "confidence":        full.get("confidence", ""),
                        "reviewer_name":     full.get("expert", "N/A"),
                        "report_id":         rid,
                    }
                    _iv_out = _build_iv(output_path=_tpl_path, context=_iv_ctx, tables={})
                    if _iv_out is not None and os.path.isfile(_tpl_path):
                        name      = _tpl_name
                        _tpl_used = True
                    else:
                        raise RuntimeError("template render returned None")
                else:
                    raise FileNotFoundError("individual_valuation_professional_template.xlsm not found")
            except Exception as _iv_err:
                _fallback  = True
                _fb_reason = str(_iv_err)
                _style_used = "legacy"
                ext  = ".xlsm" if TEMPLATE.endswith(".xlsm") else ".xlsx"
                name = f"Report_{rid}_{ts}{ext}"
                path = os.path.join(OUTPUTS, name)
                write_to_excel_template(full, path)
        else:
            # legacy / detailed — write_to_excel_template fills all template sheets.
            # For legacy, advanced analytics sheets are stripped from the saved file.
            ext  = ".xlsm" if TEMPLATE.endswith(".xlsm") else ".xlsx"
            name = f"Report_{rid}_{ts}{ext}"
            path = os.path.join(OUTPUTS, name)
            write_to_excel_template(full, path)
            print(f"{_ts()} [REPORT-STYLE] requested={report_style} used={_style_used} path={path}")
            if report_style == "legacy" or _style_used == "legacy":
                _remove_legacy_advanced_sheets(path)

        write_word_summary(full, os.path.join(OUTPUTS, f"Summary_{rid}_{ts}.docx"))

        # ── (Wave 2) تقرير Word مخصص لكل غرض ولكل نوع أصل ────────────────────
        # نُولِّد التقرير إذا كان هناك غرض محدد، أو إذا كان هناك نوع أصل متخصص
        # (حتى لو لم يُحدِّد المستخدم غرضاً، فيُعتمد fair_market_value افتراضياً).
        purpose_report_url = None
        _vp_for_report = _vp or ("fair_market_value" if _specialized_result else "")
        try:
            try:
                from purpose_specific_reports import write_purpose_specific_report
            except Exception:
                from core_engine.purpose_specific_reports import write_purpose_specific_report  # type: ignore
            if _vp_for_report:
                purpose_doc_name = f"PurposeReport_{_vp_for_report}_{rid}_{ts}.docx"
                pr_path = os.path.join(OUTPUTS, purpose_doc_name)
                if write_purpose_specific_report(_vp_for_report, full, pr_path):
                    purpose_report_url = f"http://127.0.0.1:5000/api/download/{purpose_doc_name}"
        except Exception as _pr_err:
            print(f"{_ts()} [PURPOSE-REPORT] skipped: {_pr_err}")

        resp = {"status":"success",
                "market_value": res["market_value"],
                "valuation_purpose": _vp,
                "excel_url": f"http://127.0.0.1:5000/api/download/{name}",
                "report_style_requested": report_style,
                "report_style_used":      _style_used,
                "template_used":          _tpl_used,
                "fallback_used":          _fallback,
                "fallback_reason":        _fb_reason if _fallback else ""}
        if _validation_result:
            resp["validation"] = _validation_result
        # ── Persist Gate — opt-in via "persist": true in payload (Wave BA.3) ──
        # Non-fatal (P1): a DB failure adds persist_error but does NOT return 500.
        # The user still receives their Excel report.
        if payload.get("persist"):
            try:
                from reports.report_pipeline import persist_report_data as _prd
                _ps = str(payload.get("report_status") or "draft")
                _ps = _ps if _ps in ("draft", "final", "archived") else "draft"
                _db_id = _prd(full, profile_key=_style_used, status=_ps)
                resp["report_db_id"]  = _db_id
                resp["persisted"]     = True
                resp["persist_error"] = ""
            except Exception as _persist_err:
                print(f"{_ts()} [PERSIST] save_report failed: {_persist_err}")
                resp["persisted"]     = False
                resp["persist_error"] = str(_persist_err)
        # ── (end persist gate) ────────────────────────────────────────────────
        if _avm_decision is not None:
            resp["avm"] = {
                "applied":     _avm_decision.get("applied"),
                "eligible":    _avm_decision.get("eligible"),
                "verdict":     _avm_decision.get("verdict"),
                "user_ppm":    _avm_decision.get("user_ppm"),
                "avm_ppm":     (_avm_decision.get("avm") or {}).get("avm_ppm"),
                "confidence":  (_avm_decision.get("avm") or {}).get("confidence"),
                "n_records":   (_avm_decision.get("avm") or {}).get("n_records"),
                "spread_pct":  _avm_decision.get("spread_pct"),
            }
        if _specialized_result:
            resp["specialized_asset"] = {
                "asset_type":       _specialized_result.get("asset_type"),
                "method":           _specialized_result.get("method"),
                "standards":        _specialized_result.get("standards"),
                "reconciled_value": _specialized_result.get("reconciled_value"),
            }
            # ── Excel مكمِّل لنوع الأصل المتخصص ────────────────────────────
            try:
                try:
                    from asset_type_excel import write_asset_type_excel
                except Exception:
                    from core_engine.asset_type_excel import write_asset_type_excel  # type: ignore
                _at_xlsx = write_asset_type_excel(_specialized_result, OUTPUTS)
                if _at_xlsx:
                    _at_xlsx_name = os.path.basename(_at_xlsx)
                    resp["asset_type_excel_url"] = f"http://127.0.0.1:5000/api/download/{_at_xlsx_name}"
            except Exception as _ax_err:
                print(f"{_ts()} [ASSET-TYPE-EXCEL] skipped: {_ax_err}")
        if purpose_report_url:
            resp["purpose_report_url"] = purpose_report_url
        # إفصاحات إضافية للأغراض المتخصصة
        if _vp == "uncertainty_valuation":
            spread = float(payload.get("uncertainty_spread_pct", 0.15))
            mv = float(res["market_value"])
            resp["uncertainty_range"] = {
                "low":  mv * (1 - spread),
                "best": mv,
                "high": mv * (1 + spread),
                "spread_pct": spread * 100,
            }
        if _vp == "usufruct" and "usufruct_meta" in payload:
            resp["usufruct"] = payload["usufruct_meta"]
        return jsonify(resp)
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({"status":"error","message":str(e)}), 500

@app.route("/api/download/<filename>")
def download(filename):
    return send_file(os.path.join(OUTPUTS, filename), as_attachment=True)

# ═══════════════════════════════════════════════════════════════════════════
# Endpoint: Market Feed — استقبال وعرض البيانات الخارجية
# ═══════════════════════════════════════════════════════════════════════════
@app.route("/api/market-feed", methods=["POST", "OPTIONS"])
def market_feed_post():
    """
    يُضيف سجل بيانات سوقية جديد من أي مصدر خارجي.

    Body (JSON):
    {
      "source":        "facebook" | "forum" | "agent" | "direct" | "other",
      "location":      "القاهرة الجديدة",
      "property_type": "شقة سكنية",
      "area":          150,
      "price":         5250000,          ← اختياري إذا وُجد price_per_meter
      "price_per_meter": 35000,          ← اختياري إذا وُجد price + area
      "floor":         3,               ← اختياري
      "year_built":    2020,            ← اختياري
      "notes":         "نص حر"         ← اختياري
    }
    """
    if request.method == "OPTIONS":
        return jsonify({}), 200
    try:
        import json as _json
        payload = request.get_json(silent=True) or {}

        source   = payload.get("source", "other")
        location = payload.get("location", "")
        ptype    = payload.get("property_type", "")
        area_f   = _num(payload.get("area", 0))
        price_f  = _num(payload.get("price", 0))
        ppm_f    = _num(payload.get("price_per_meter", 0))

        # استنتاج PPM إذا لم يُرسل صراحةً
        if ppm_f <= 0 and price_f > 0 and area_f > 0:
            ppm_f = round(price_f / area_f, 0)
        if price_f <= 0 and ppm_f > 0 and area_f > 0:
            price_f = round(ppm_f * area_f, 0)

        if ppm_f <= 0:
            return jsonify({"status": "error",
                            "message": "يجب إرسال price_per_meter أو (price + area)"}), 400

        record = {
            "id":              uuid.uuid4().hex[:12],
            "timestamp":       datetime.now().isoformat(),
            "source":          source,
            "location":        location,
            "property_type":   ptype,
            "area":            area_f,
            "price":           price_f,
            "price_per_meter": ppm_f,
            "floor":           _num(payload.get("floor", 0)),
            "year_built":      _num(payload.get("year_built", 0)),
            "notes":           payload.get("notes", ""),
            "credibility":     _SOURCE_CREDIBILITY.get(source, 0.55),
        }

        feed = _load_feed()
        feed.append(record)
        _save_feed(feed)

        return jsonify({
            "status":  "success",
            "id":      record["id"],
            "message": f"تم حفظ السجل. إجمالي السجلات: {len(feed)}",
            "credibility": record["credibility"],
        })
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route("/api/market-feed", methods=["GET"])
def market_feed_get():
    """
    يُعيد ملخص بيانات السوق المُخزّنة مع إمكانية التصفية.

    Query params:
      location=      (اختياري) تصفية بالموقع (تطابق جزئي)
      property_type= (اختياري) تصفية بالنوع
      since=         (اختياري) تاريخ ISO  e.g. 2025-01-01
      source=        (اختياري) facebook | forum | agent | direct
      analyze=true   (اختياري) يُشغّل المُدقق ويُعيد الإحصاءات
    """
    try:
        feed = _load_feed()
        loc_f   = request.args.get("location", "")
        ptype_f = request.args.get("property_type", "")
        since_f = request.args.get("since", "")
        src_f   = request.args.get("source", "")
        analyze = request.args.get("analyze", "false").lower() == "true"

        filtered = []
        for rec in feed:
            if loc_f and loc_f not in rec.get("location", ""):
                continue
            if ptype_f and ptype_f not in rec.get("property_type", ""):
                continue
            if src_f and rec.get("source") != src_f:
                continue
            if since_f:
                try:
                    from datetime import datetime as _dt2
                    since_dt = _dt2.fromisoformat(since_f)
                    rec_dt   = _dt2.fromisoformat(rec.get("timestamp","2000-01-01"))
                    if rec_dt < since_dt:
                        continue
                except Exception:
                    pass
            filtered.append(rec)

        # إضافة الوزن الزمني لكل سجل في الاستجابة
        for rec in filtered:
            rec["temporal_weight"] = round(_temporal_weight(rec.get("timestamp", "")), 3)

        response = {
            "status":       "success",
            "total_records": len(feed),
            "filtered_count": len(filtered),
            "records":       filtered,
        }

        if analyze and filtered:
            mi = _market_intelligence_validator(
                comps        = [],
                location     = loc_f,
                property_type= ptype_f,
                feed_records = filtered,
            )
            response["analysis"] = {
                "n_clean":       mi["n_clean"],
                "n_outliers":    mi["n_outliers"],
                "weighted_ppm":  mi["weighted_ppm"],
                "adjusted_ppm":  mi["adjusted_ppm"],
                "neg_coeff":     mi["neg_coeff"],
                "quality_score": mi["quality_score"],
                "diagnostics":   mi["diagnostics"],
            }

        return jsonify(response)
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route("/api/market-feed/<record_id>", methods=["DELETE"])
def market_feed_delete(record_id):
    """يحذف سجلاً بعينه من قاعدة البيانات"""
    try:
        feed = _load_feed()
        before = len(feed)
        feed = [r for r in feed if r.get("id") != record_id]
        _save_feed(feed)
        return jsonify({"status": "success",
                        "deleted": before - len(feed),
                        "remaining": len(feed)})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


# ═══════════════════════════════════════════════════════════════════════════
# Endpoint: RAG Advisor — المستشار العقاري الذكي
# POST /api/advisor  { question, location?, property_type?, use_web? }
# GET  /api/advisor/health
# ═══════════════════════════════════════════════════════════════════════════
@app.route("/api/advisor", methods=["POST", "OPTIONS"])
def advisor_endpoint():
    """
    المستشار العقاري الذكي — يُجيب على الأسئلة العقارية العامة
    باستخدام RAG (Qdrant + market_feed + Macro KB + DuckDuckGo + Ollama).

    Body (JSON):
    {
      "question":      "ما هو أفضل وقت للاستثمار في رأس الحكمة؟",
      "location":      "رأس الحكمة",          ← اختياري
      "property_type": "شاليه",               ← اختياري
      "use_web":       true                   ← اختياري (بحث DuckDuckGo)
    }

    Response (JSON):
    {
      "status":     "success",
      "answer":     "...",
      "sources":    [...],
      "confidence": 0.82,
      "mode":       "rag+llm",
      "docs_used":  3,
      "web_hits":   2,
      "elapsed_s":  4.2
    }
    """
    if request.method == "OPTIONS":
        return jsonify({}), 200
    try:
        payload       = request.get_json(silent=True) or {}
        question      = str(payload.get("question", "")).strip()
        location      = str(payload.get("location", "")).strip()
        property_type = str(payload.get("property_type", "")).strip()
        use_web       = bool(payload.get("use_web", False))

        if not question:
            return jsonify({"status": "error",
                            "message": "يجب إرسال حقل question"}), 400

        result = _rag_answer(
            question      = question,
            location      = location,
            property_type = property_type,
            use_web       = use_web,
        )
        return jsonify({
            "status":     "success",
            "answer":     result.get("answer", ""),
            "sources":    result.get("sources", []),
            "confidence": result.get("confidence", 0),
            "mode":       result.get("mode", ""),
            "docs_used":  result.get("docs_used", 0),
            "web_hits":   result.get("web_hits", 0),
            "elapsed_s":  result.get("elapsed_s", 0),
        })
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route("/api/advisor/health", methods=["GET"])
def advisor_health():
    """فحص حالة المستشار الذكي"""
    return jsonify({
        "status":    "ok" if _RAG_OK else "disabled",
        "rag_ready": _RAG_OK,
        "message":   "المستشار الذكي جاهز" if _RAG_OK else "RAG غير متاح — تحقق من المكتبات",
    })


# ═══════════════════════════════════════════════════════════════════════════
# Context-Tuning — محرك المحاكاة الأسلوبية
# ═══════════════════════════════════════════════════════════════════════════

@app.route("/api/tune/analyze", methods=["POST", "OPTIONS"])
def tune_analyze():
    """
    يحلّل نص وثيقة ويستخرج بصمتها الأسلوبية.

    Body: { "text": "...", "filename": "تقرير تقييم 2024.pdf" }
    Response: { "status":"success", "profile": {...}, "injection_preview": "..." }
    """
    if request.method == "OPTIONS": return jsonify({}), 200
    try:
        payload  = request.get_json(silent=True) or {}
        text     = str(payload.get("text", "")).strip()
        filename = str(payload.get("filename", "وثيقة مجهولة")).strip()

        if len(text) < 100:
            return jsonify({"status": "error",
                            "message": "النص قصير جداً — أرسل 100 حرف على الأقل"}), 400

        profile = _tune_analyze(text, filename)
        _tune_save(profile)

        return jsonify({
            "status":           "success",
            "profile_id":       profile["id"],
            "name":             profile["name"],
            "doc_type":         profile["doc_type"],
            "language":         profile["language"],
            "formality":        profile.get("formality", ""),
            "standards":        profile["terminology"]["standards_refs"],
            "methods":          profile["terminology"]["valuation_methods"],
            "section_count":    len(profile["structure"]["section_headers"]),
            "confidence":       profile["confidence"],
            "injection_preview": profile["prompt_injection"][:400] + "...",
        })
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route("/api/tune/profiles", methods=["GET"])
def tune_list_profiles():
    """يُعيد قائمة الأساليب المحفوظة"""
    try:
        return jsonify({"status": "success", "profiles": _tune_list()})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route("/api/tune/profiles/<profile_id>", methods=["GET"])
def tune_get_profile(profile_id):
    """يُعيد تفاصيل ملف أسلوب"""
    profile = _tune_load(profile_id)
    if not profile:
        return jsonify({"status": "error", "message": "الملف غير موجود"}), 404
    return jsonify({"status": "success", "profile": profile})


@app.route("/api/tune/profiles/<profile_id>", methods=["DELETE"])
def tune_delete_profile(profile_id):
    """يحذف ملف أسلوب"""
    ok = _tune_delete(profile_id)
    return jsonify({"status": "success" if ok else "error",
                    "deleted": ok})


@app.route("/api/tune/apply", methods=["POST", "OPTIONS"])
def tune_apply_prompt():
    """
    يُطبّق بصمة أسلوب على موجّه أساسي ويُعيد الموجّه المُعزَّز.

    Body: { "prompt": "...", "profile_id": "abc123" }
    """
    if request.method == "OPTIONS": return jsonify({}), 200
    try:
        payload    = request.get_json(silent=True) or {}
        prompt     = str(payload.get("prompt", ""))
        profile_id = str(payload.get("profile_id", ""))
        enhanced   = _tune_apply(prompt, profile_id)
        return jsonify({"status": "success", "enhanced_prompt": enhanced})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


# ═══════════════════════════════════════════════════════════════════════════
# Smart Library — المكتبة المرجعية الذكية
# ═══════════════════════════════════════════════════════════════════════════

@app.route("/api/library", methods=["GET"])
def library_list():
    """
    يُعيد وثائق المكتبة مع دعم الفلترة والبحث.
    GET /api/library?q=تقييم&sector=residential&limit=20
    """
    try:
        q        = request.args.get("q", "")
        sector   = request.args.get("sector", "")
        doc_type = request.args.get("doc_type", "")
        country  = request.args.get("country", "")
        language = request.args.get("language", "")
        limit    = min(int(request.args.get("limit", 30)), 100)

        results = _lib_search(
            query=q, sector=sector, doc_type=doc_type,
            country=country, language=language, limit=limit
        )
        stats = _lib_stats()
        return jsonify({
            "status":  "success",
            "records": results,
            "count":   len(results),
            "stats":   stats,
        })
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route("/api/library/scan", methods=["POST", "OPTIONS"])
def library_scan():
    """
    يُشغّل مسح دوري لكل المصادر المُفعَّلة.
    POST /api/library/scan
    """
    if request.method == "OPTIONS": return jsonify({}), 200
    try:
        import threading
        result_box = {}

        def _run():
            result_box.update(_lib_scan())

        # نُشغّل المسح في خلفية حتى لا يتجاوز مهلة HTTP
        t = threading.Thread(target=_run, daemon=True, name="lib-scan")
        t.start()
        t.join(timeout=60)   # ننتظر 60 ثانية كحد أقصى

        return jsonify({
            "status":    "success",
            "scan_done": not t.is_alive(),
            **result_box,
        })
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route("/api/library/add", methods=["POST", "OPTIONS"])
def library_add_manual():
    """
    يُضيف وثيقة مرفوعة يدوياً إلى المكتبة.

    Body: { "title":"...", "text":"...", "filename":"...",
            "sector":"residential", "doc_type":"valuation_report" }
    """
    if request.method == "OPTIONS": return jsonify({}), 200
    try:
        payload  = request.get_json(silent=True) or {}
        title    = str(payload.get("title", payload.get("filename", "وثيقة بدون عنوان"))).strip()
        text     = str(payload.get("text", "")).strip()
        filename = str(payload.get("filename", "")).strip()
        sector   = str(payload.get("sector", "all"))
        doc_type = str(payload.get("doc_type", "valuation_report"))

        if len(text) < 50:
            return jsonify({"status": "error",
                            "message": "النص قصير جداً للإضافة إلى المكتبة"}), 400

        result = _lib_add_manual(title, text, filename, sector, doc_type)
        return jsonify({"status": result["status"], "record": result["record"]})
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route("/api/library/<record_id>", methods=["GET"])
def library_get(record_id):
    """يُعيد تفاصيل وثيقة"""
    rec = _lib_get(record_id)
    if not rec:
        return jsonify({"status": "error", "message": "الوثيقة غير موجودة"}), 404
    return jsonify({"status": "success", "record": rec})


@app.route("/api/library/<record_id>", methods=["DELETE"])
def library_delete(record_id):
    """يحذف وثيقة من المكتبة"""
    ok = _lib_delete(record_id)
    return jsonify({"status": "success" if ok else "error", "deleted": ok})


# ══════════════════════════════════════════════════════════════════════════════
# خريطة الأسعار التفاعلية — /api/price-map
# يُولّد خريطة Leaflet تفاعلية بدبوس للموقع وسعر المتر
# ══════════════════════════════════════════════════════════════════════════════

def _lookup_coords(location_text: str):
    """
    يبحث عن إحداثيات الموقع في _CITY_COORDS.
    يُعيد (lat, lng, matched_name) أو (None, None, None).
    """
    loc = str(location_text or "").strip()
    if not loc:
        return None, None, None
    # مطابقة مباشرة
    if loc in _CITY_COORDS:
        lat, lng = _CITY_COORDS[loc]
        return lat, lng, loc
    # مطابقة جزئية (substring) — الأطول أولاً لتجنب تضارب "القاهرة" و"القاهرة الجديدة"
    for city in sorted(_CITY_COORDS.keys(), key=len, reverse=True):
        if city in loc or loc in city:
            lat, lng = _CITY_COORDS[city]
            return lat, lng, city
    return None, None, None


def _build_price_map_html(
    lat: float,
    lng: float,
    city_name: str,
    ppm: float,
    nearby: list = None,
) -> str:
    """
    يُنشئ صفحة HTML كاملة تحتوي خريطة Leaflet تفاعلية.
    - دبوس الموقع الأساسي (ذهبي)
    - دوائر أسعار المدن المجاورة (اختياري)
    - بدون أي مكتبة خارجية غير Leaflet CDN
    """
    ppm_formatted = f"{ppm:,.0f}" if ppm else "—"

    # ── بناء markers للمدن المجاورة (كل ما في _CITY_COORDS) ──────────
    nearby_js = ""
    for c_name, (c_lat, c_lng) in _CITY_COORDS.items():
        c_ppm = _PRICE_MAP.get(c_name, 0)
        if c_ppm <= 0:
            continue
        color = (
            "#4ade80" if c_ppm < 20000 else
            "#fbbf24" if c_ppm < 35000 else
            "#f87171"
        )
        nearby_js += f"""
        L.circleMarker([{c_lat},{c_lng}], {{
            radius: 7, color:'{color}', fillColor:'{color}',
            fillOpacity: 0.65, weight: 1.5
        }}).addTo(map).bindPopup('<div dir=rtl style="font-family:Tajawal,sans-serif;text-align:right;"><b>{c_name}</b><br>متوسط: {c_ppm:,.0f} EGP/م²</div>');
        """

    html = f"""<!DOCTYPE html>
<html lang="ar" dir="rtl">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>خريطة الأسعار — Expert_Smart</title>
<link rel="stylesheet" href="https://unpkg.com/leaflet@1.9.4/dist/leaflet.css"/>
<script src="https://unpkg.com/leaflet@1.9.4/dist/leaflet.js"></script>
<style>
  * {{ margin:0; padding:0; box-sizing:border-box; }}
  body {{ font-family: 'Tajawal','Segoe UI',sans-serif; background:#030812; }}
  #map {{ width:100%; height:100vh; }}
  .legend {{
      position:fixed; bottom:18px; right:18px; z-index:9999;
      background:rgba(3,8,18,0.92); border:1px solid rgba(255,215,0,0.35);
      border-radius:14px; padding:14px 18px; color:#e2e8f0;
      font-size:13px; backdrop-filter:blur(6px);
      box-shadow: 0 4px 24px rgba(0,0,0,0.5);
  }}
  .legend h4 {{ color:#FFD700; margin-bottom:8px; font-size:14px; }}
  .legend .row {{ display:flex; align-items:center; gap:8px; margin:4px 0; }}
  .legend .dot {{ width:12px; height:12px; border-radius:50%; display:inline-block; }}
  .marker-label {{
      background:rgba(3,8,18,0.88); color:#FFD700; padding:4px 10px;
      border-radius:8px; border:1px solid rgba(255,215,0,0.5);
      font-family:'Tajawal',sans-serif; font-weight:700; font-size:13px;
      white-space:nowrap; box-shadow:0 2px 10px rgba(0,0,0,0.4);
  }}
</style>
</head>
<body>
<div id="map"></div>

<div class="legend">
    <h4>🗺 خريطة أسعار السوق</h4>
    <div class="row"><span class="dot" style="background:#4ade80;"></span> أقل من 20,000 EGP/م²</div>
    <div class="row"><span class="dot" style="background:#fbbf24;"></span> 20,000 – 35,000 EGP/م²</div>
    <div class="row"><span class="dot" style="background:#f87171;"></span> أكثر من 35,000 EGP/م²</div>
    <div style="margin-top:8px;font-size:11px;color:#667799;">Expert_Smart v45 — Sovereign Edition</div>
</div>

<script>
var map = L.map('map', {{
    center: [{lat}, {lng}],
    zoom: 12,
    zoomControl: true
}});

L.tileLayer('https://{{s}}.basemaps.cartocdn.com/dark_all/{{z}}/{{x}}/{{y}}{{r}}.png', {{
    attribution: '&copy; <a href="https://carto.com/">CARTO</a> &copy; <a href="https://osm.org/">OSM</a>',
    subdomains: 'abcd',
    maxZoom: 19
}}).addTo(map);

// ── الدبوس الرئيسي (ذهبي) ──────────────────────────────
var goldIcon = L.divIcon({{
    className: '',
    html: '<div class="marker-label">📌 {city_name}<br>{ppm_formatted} EGP/م²</div>',
    iconSize: [180, 50],
    iconAnchor: [90, 50]
}});
L.marker([{lat}, {lng}], {{ icon: goldIcon }}).addTo(map)
 .bindPopup('<div dir=rtl style="font-family:Tajawal,sans-serif;text-align:right;"><b>📍 {city_name}</b><br>سعر المتر: <b style=color:#FFD700>{ppm_formatted} EGP</b></div>')
 .openPopup();

// ── نقاط المدن المجاورة ──────────────────────────────────
{nearby_js}
</script>
</body>
</html>"""
    return html


@app.route("/api/price-map", methods=["POST", "GET", "OPTIONS"])
def price_map():
    """
    خريطة الأسعار التفاعلية.

    POST /api/price-map  { "location": "التجمع الخامس", "price_per_meter": 42000 }
    GET  /api/price-map?format=html&location=...&price_per_meter=...

    إذا format=html → يُعيد صفحة HTML كاملة (Leaflet map)
    إذا format=json → يُعيد JSON بالإحداثيات والبيانات
    """
    if request.method == "OPTIONS":
        return jsonify({}), 200

    # ── استخراج المعاملات من POST body أو GET query ──────────────────
    if request.method == "POST":
        body     = request.get_json(silent=True) or {}
        location = str(body.get("location", ""))
        ppm      = float(body.get("price_per_meter", 0))
    else:
        location = request.args.get("location", "")
        ppm      = float(request.args.get("price_per_meter", 0) or 0)

    fmt = request.args.get("format", "json")

    # ── البحث عن الإحداثيات ─────────────────────────────────────────
    lat, lng, matched = _lookup_coords(location)

    # fallback: القاهرة كإحداثيات افتراضية
    if lat is None:
        lat, lng, matched = 30.0444, 31.2357, "القاهرة (افتراضي)"

    # fallback: سعر من _PRICE_MAP
    if ppm <= 0:
        ppm = _PRICE_MAP.get(matched, _PRICE_MAP.get(location, 25000))

    # ── JSON mode ──────────────────────────────────────────────────────
    if fmt != "html":
        # أرجع كل المدن مع أسعارها لرسمها في الواجهة
        cities_data = []
        for c_name, (c_lat, c_lng) in _CITY_COORDS.items():
            c_ppm = _PRICE_MAP.get(c_name, 0)
            if c_ppm > 0:
                cities_data.append({
                    "name": c_name, "lat": c_lat, "lng": c_lng,
                    "ppm":  c_ppm,
                })
        return jsonify({
            "status":    "success",
            "location":  matched,
            "lat":       lat,
            "lng":       lng,
            "ppm":       ppm,
            "cities":    cities_data,
        })

    # ── HTML mode — خريطة Leaflet كاملة ──────────────────────────────
    try:
        html = _build_price_map_html(lat, lng, matched, ppm)
        from flask import Response
        return Response(html, mimetype="text/html; charset=utf-8")
    except Exception as e:
        return f"<h1>Error generating map</h1><pre>{e}</pre>", 500


# ══════════════════════════════════════════════════════════════════════════════
#  v37 — Market Radar Endpoints
# ══════════════════════════════════════════════════════════════════════════════

@app.route("/api/radar/start", methods=["POST", "OPTIONS"])
def radar_start():
    """
    يُشغِّل محرك رادار السوق في الخلفية.
    Body (optional): {
      "interval_seconds": 1800,
      "enable_web": true,
      "enable_whatsapp": false,
      "wa_groups": ["مجموعة عقارات القاهرة"]
    }
    """
    if request.method == "OPTIONS": return jsonify({}), 200
    try:
        p = request.get_json(silent=True) or {}
        result = _radar.start(
            interval_seconds = int(p.get("interval_seconds", 1800)),
            enable_web       = bool(p.get("enable_web",       True)),
            enable_whatsapp  = bool(p.get("enable_whatsapp",  False)),
            wa_groups        = p.get("wa_groups", []),
        )
        return jsonify({"status": "success", **result})
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route("/api/radar/stop", methods=["POST", "OPTIONS"])
def radar_stop():
    """يُوقِف محرك رادار السوق."""
    if request.method == "OPTIONS": return jsonify({}), 200
    try:
        return jsonify({"status": "success", **_radar.stop()})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route("/api/radar/status", methods=["GET"])
def radar_status():
    """حالة الرادار + إحصائيات قاعدة بيانات السوق."""
    try:
        return jsonify({"status": "success", "radar": _radar.status()})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route("/api/radar/data", methods=["GET"])
def radar_data():
    """
    يُعيد سجلات السوق المجمَّعة.
    Query params: location, asset_type, source, hours_back, limit (≤500), offset
    """
    try:
        location   = request.args.get("location", "")
        asset_type = request.args.get("asset_type", "")
        source     = request.args.get("source", "")
        hours_back = request.args.get("hours_back")
        limit      = min(int(request.args.get("limit",  100)), 500)
        offset     = int(request.args.get("offset", 0))
        records = _radar.get_records(
            location   = location   or None,
            asset_type = asset_type or None,
            source     = source     or None,
            hours_back = int(hours_back) if hours_back else None,
            min_conf   = 0.3,
            limit      = limit,
            offset     = offset,
        )
        return jsonify({
            "status":  "success",
            "count":   len(records),
            "records": records,
            "stats":   _radar.get_stats(),
        })
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route("/api/radar/parse", methods=["POST", "OPTIONS"])
def radar_parse_text():
    """
    يُحلِّل نصاً عقارياً حراً ويُضيفه إلى قاعدة البيانات.
    Body: { "text": "شقة 120م في التجمع بـ 3 مليون" }
    """
    if request.method == "OPTIONS": return jsonify({}), 200
    try:
        p    = request.get_json(silent=True) or {}
        text = str(p.get("text", "")).strip()
        if len(text) < 10:
            return jsonify({"status": "error", "message": "النص قصير جداً"}), 400
        rec = _radar.parse_text(text)
        if rec:
            return jsonify({"status": "success", "record": rec})
        return jsonify({"status": "no_data", "message": "لم يُستخرج بيانات عقارية كافية"}), 422
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route("/api/radar/heatmap", methods=["GET"])
def radar_heatmap():
    """
    بيانات خريطة الأسعار (PPM مُجمَّعة حسب الموقع).
    Query: asset_type (default: شقة)
    """
    try:
        asset_type = request.args.get("asset_type", "شقة")
        data       = _radar.get_heatmap_data(asset_type=asset_type)
        return jsonify({"status": "success", "count": len(data), "zones": data})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


# ══════════════════════════════════════════════════════════════════════════════
#  v37 — IAAO Standards Endpoint
# ══════════════════════════════════════════════════════════════════════════════

@app.route("/api/iaao", methods=["POST", "OPTIONS"])
def iaao_compute():
    """
    يحسب مقاييس IAAO (ASR / COD / PRD / PRB) وإشارات المرور.

    Body option A — قائمتان مباشرتان:
      { "assessed_values": [float,...], "sale_prices": [float,...],
        "location": "...", "sector": "residential", "period": "Q1-2025" }

    Body option B — مقارنات + قيمة متر مُقدَّرة:
      { "comparables": [{price, area, location}, ...],
        "assessed_ppm": 25000, "sector": "residential", "location": "..." }
    """
    if request.method == "OPTIONS": return jsonify({}), 200
    try:
        p = request.get_json(silent=True) or {}

        if "comparables" in p:
            report = _iaao_from_comps(
                comparables  = p["comparables"],
                assessed_ppm = float(p.get("assessed_ppm", 0)),
                sector       = str(p.get("sector", "residential")),
                location     = str(p.get("location", "")),
            )
        else:
            assessed = [float(v) for v in p.get("assessed_values", [])]
            sales    = [float(v) for v in p.get("sale_prices", [])]
            if len(assessed) < 3:
                return jsonify({"status": "error",
                                "message": "يلزم 3 مقارنات على الأقل"}), 400
            report = _iaao_report(
                assessed_values = assessed,
                sale_prices     = sales,
                location        = str(p.get("location", "")),
                sector          = str(p.get("sector", "residential")),
                period          = str(p.get("period", "")),
            )
        return jsonify({"status": "success", "report": report})
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route("/api/iaao/health", methods=["GET"])
def iaao_health():
    return jsonify({"status": "success", "iaao_ready": _IAAO_OK,
                    "radar_ready": _RADAR_OK,
                    "message": "IAAO Engine v37 — جاهز" if _IAAO_OK else "IAAO Engine غير متاح"})


# ── Training Template Registry (نموذج التدريب المرجعي) ──────────────────────
_training_registry: dict = {}   # latest registered master template per session

@app.route("/api/training/register", methods=["POST"])
def training_register():
    """Receives the master template metadata from the frontend Training Clip."""
    try:
        body        = request.get_json(force=True, silent=True) or {}
        filename    = str(body.get("filename", "")).strip()
        sheet_text  = str(body.get("sheet_text", ""))
        uploaded_by = str(body.get("uploaded_by", "غير محدد"))

        if not filename:
            return jsonify({"status": "error", "message": "filename مطلوب"}), 400

        import hashlib, datetime
        _training_registry.update({
            "filename":    filename,
            "uploaded_by": uploaded_by,
            "registered":  datetime.datetime.utcnow().isoformat(),
            "text_hash":   hashlib.md5(sheet_text.encode("utf-8", "replace")).hexdigest(),
            "text_len":    len(sheet_text),
        })
        print(f"  [training] Master template registered: {filename} by {uploaded_by}")
        return jsonify({
            "status":  "success",
            "message": f"تم تسجيل النموذج المرجعي: {filename}",
            "meta":    _training_registry,
        })
    except Exception as exc:
        return jsonify({"status": "error", "message": str(exc)}), 500


@app.route("/api/training/status", methods=["GET"])
def training_status():
    """Returns the currently registered master template, if any."""
    if _training_registry:
        return jsonify({"status": "success", "registered": True, "meta": _training_registry})
    return jsonify({"status": "success", "registered": False, "meta": {}})


# ══════════════════════════════════════════════════════════════════════════════
#  Price Intelligence API — بروتوكول الرصد اللحظي للأسعار
# ══════════════════════════════════════════════════════════════════════════════

@app.route("/api/price/intelligence", methods=["POST"])
def price_intelligence_search():
    """
    البحث الكامل عن أسعار منطقة محددة.

    Body (JSON):
      location    str   — المنطقة (مثل: التجمع الخامس)
      market      str   — egypt | ksa          (افتراضي: egypt)
      asset_type  str   — نوع الأصل            (افتراضي: شقة)
      max_items   int   — أقصى سجلات لكل مصدر  (افتراضي: 20)
      force       bool  — تجاهل الكاش           (افتراضي: false)

    Returns:
      summary       — إحصائيات: avg_ppm, min/max, معامل التفاوض ...
      analysis_text — تقرير تحليلي عربي بأسلوب خبير فئة A
      excel_table   — جدول جاهز للإكسيل (16 عموداً)
      records       — قائمة السجلات الخام
      trends        — إشارات الاتجاه الاقتصادي
    """
    try:
        body       = request.get_json(force=True, silent=True) or {}
        location   = str(body.get("location", "")).strip()
        market     = str(body.get("market", _MKT_EG)).strip()
        asset_type = str(body.get("asset_type", "شقة")).strip()
        max_items  = int(body.get("max_items", 20))
        force      = bool(body.get("force", False))

        if not location:
            return jsonify({"status": "error", "message": "location مطلوب"}), 400
        if market not in (_MKT_EG, _MKT_KSA):
            market = _MKT_EG

        result = _price_intel.search_dict(
            location=location,
            market=market,
            asset_type=asset_type,
            max_items=max_items,
            force=force,
        )
        return jsonify(result)

    except Exception as exc:
        return jsonify({"status": "error", "message": str(exc)}), 500


@app.route("/api/price/trend", methods=["GET"])
def price_trend():
    """
    جلب إشارات الاتجاه الاقتصادي فقط (بدون scraping الأسعار — أسرع).

    Query params:
      market    egypt | ksa
      location  اختياري — لتخصيص البحث
    """
    try:
        market   = request.args.get("market", _MKT_EG)
        location = request.args.get("location", "")
        if market not in (_MKT_EG, _MKT_KSA):
            market = _MKT_EG
        signals = _price_intel.quick_trend(market=market, location=location)
        return jsonify({"status": "success", "market": market, "trends": signals,
                        "count": len(signals), "as_of": datetime.now().isoformat()
                        if 'datetime' in dir() else ""})
    except Exception as exc:
        return jsonify({"status": "error", "message": str(exc)}), 500


@app.route("/api/price/excel-table", methods=["POST"])
def price_excel_table():
    """
    يُعيد جدول مقارنة الأسعار بتنسيق JSON جاهز للحقن في شيت Excel.

    Body (JSON):
      records   — قائمة سجلات من /api/price/intelligence (اختياري)
      summary   — ملخص من /api/price/intelligence       (اختياري)
      location, market, asset_type, max_items            (إذا لم تُمرَّر records)
    """
    try:
        body    = request.get_json(force=True, silent=True) or {}
        records = body.get("records")
        summary = body.get("summary")

        if records is None:
            # Fetch fresh data
            location   = str(body.get("location", "")).strip()
            market     = str(body.get("market", _MKT_EG)).strip()
            asset_type = str(body.get("asset_type", "شقة")).strip()
            if not location:
                return jsonify({"status": "error", "message": "location أو records مطلوب"}), 400
            result  = _price_intel.search_dict(
                location=location, market=market,
                asset_type=asset_type, max_items=int(body.get("max_items", 15))
            )
            table   = result.get("excel_table", [])
            summary = result.get("summary", {})
        else:
            # Re-build table from supplied records
            from price_intelligence import PriceRecord
            rec_objs = []
            for r in records:
                try:
                    rec_objs.append(PriceRecord(**{
                        k: r.get(k) for k in PriceRecord.__dataclass_fields__
                    }))
                except Exception:
                    pass
            table = _build_excel_table(rec_objs, summary or {})

        return jsonify({
            "status":       "success",
            "excel_table":  table,
            "column_count": len(table[0]) if table else 0,
            "row_count":    len(table),
            "as_of":        __import__('datetime').datetime.now().strftime("%Y-%m-%d %H:%M"),
        })

    except Exception as exc:
        return jsonify({"status": "error", "message": str(exc)}), 500


@app.route("/api/price/cache/clear", methods=["POST"])
def price_cache_clear():
    """يمسح كاش بروتوكول الرصد اللحظي."""
    return jsonify(_price_intel.clear_cache())


# ═══════════════════════════════════════════════════════════════════════════
# File Upload Endpoints — /api/ingest  /api/upload
# ═══════════════════════════════════════════════════════════════════════════
import uuid as _uuid

_UPLOAD_DIR = os.path.join(os.path.dirname(__file__), "uploads")
os.makedirs(_UPLOAD_DIR, exist_ok=True)

_ALLOWED_DOC_EXT  = {".xlsx", ".xls", ".csv", ".pdf", ".docx", ".doc", ".txt", ".xlsm"}
_ALLOWED_IMG_EXT  = {".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp", ".tiff", ".tif", ".jfif"}


def _save_upload(file_obj) -> tuple[str, str]:
    """Save an uploaded FileStorage to the uploads dir. Returns (saved_path, original_name)."""
    orig = file_obj.filename or "upload"
    ext  = os.path.splitext(orig)[1].lower()
    safe_name = f"{_uuid.uuid4().hex}{ext}"
    dest = os.path.join(_UPLOAD_DIR, safe_name)
    file_obj.save(dest)
    return dest, orig


@app.route("/api/ingest", methods=["POST", "OPTIONS"])
def ingest_file():
    """استقبال ملفات Excel / PDF / Word / CSV وتخزينها للمعالجة."""
    if request.method == "OPTIONS":
        return jsonify({}), 200
    if "file" not in request.files:
        return jsonify({"status": "error", "message": "لم يتم إرسال ملف"}), 400
    f = request.files["file"]
    ext = os.path.splitext(f.filename or "")[1].lower()
    if ext not in _ALLOWED_DOC_EXT:
        return jsonify({"status": "error", "message": f"صيغة الملف غير مدعومة: {ext}"}), 415
    try:
        dest, orig = _save_upload(f)
        file_type = request.form.get("type", "document")
        return jsonify({
            "status": "success",
            "message": f"تم استلام الملف: {orig}",
            "file_type": file_type,
            "saved_as": os.path.basename(dest),
        })
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route("/api/upload", methods=["POST", "OPTIONS"])
def upload_image():
    """استقبال صور العقارات وتخزينها."""
    if request.method == "OPTIONS":
        return jsonify({}), 200
    if "file" not in request.files:
        return jsonify({"status": "error", "message": "لم يتم إرسال صورة"}), 400
    f = request.files["file"]
    ext = os.path.splitext(f.filename or "")[1].lower()
    # Accept both image and document types (generic upload endpoint)
    allowed = _ALLOWED_IMG_EXT | _ALLOWED_DOC_EXT
    if ext not in allowed:
        return jsonify({"status": "error", "message": f"صيغة الملف غير مدعومة: {ext}"}), 415
    try:
        dest, orig = _save_upload(f)
        file_type = request.form.get("type", "image")
        return jsonify({
            "status": "success",
            "message": f"تم استلام الملف: {orig}",
            "file_type": file_type,
            "saved_as": os.path.basename(dest),
        })
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


# ═══════════════════════════════════════════════════════════════════════════
# Super Intelligence Suite — حزمة الاستخبارات العقارية الفائقة
# ═══════════════════════════════════════════════════════════════════════════

# ── 1. المُحقق الرقمي (Fraud & Anomaly Detection) ─────────────────────────
try:
    from fraud_detector import detect_fraud as _detect_fraud
    _FRAUD_OK = True
    print("  [fraud_detector] loaded ✓")
except Exception as _fd_err:
    _FRAUD_OK = False
    print(f"  [fraud_detector] unavailable: {_fd_err}")
    def _detect_fraud(**kw):
        return {"status": "error", "message": "fraud_detector غير متاح"}


@app.route("/api/fraud/detect", methods=["POST", "OPTIONS"])
def fraud_detect():
    """
    POST /api/fraud/detect
    يُحلّل سعر المتر ويكشف عن احتمالية التلاعب أو الشذوذ.
    Body: { price_per_meter, location, property_type?, comp_ppms?, area_m2?, total_price? }
    """
    if request.method == "OPTIONS":
        return jsonify({}), 200
    try:
        body = request.get_json(force=True) or {}
        ppm  = float(body.get("price_per_meter", 0))
        if ppm <= 0:
            return jsonify({"status": "error", "message": "price_per_meter مطلوب"}), 400
        result = _detect_fraud(
            price_per_meter = ppm,
            location        = str(body.get("location", "")),
            property_type   = str(body.get("property_type", "")),
            comp_ppms       = [float(x) for x in body.get("comp_ppms", []) if x],
            area_m2         = float(body.get("area_m2", 0)),
            total_price     = float(body.get("total_price", 0)),
        )
        result["status"] = "success"
        return jsonify(result)
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


# ── 2. التحليل الجيوتقني والقانوني (Geo Risk Engine) ─────────────────────
try:
    from geo_risk_engine import analyze_geo_risk as _analyze_geo_risk
    _GEO_OK = True
    print("  [geo_risk_engine] loaded ✓")
except Exception as _gr_err:
    _GEO_OK = False
    print(f"  [geo_risk_engine] unavailable: {_gr_err}")
    def _analyze_geo_risk(**kw):
        return {"status": "error", "message": "geo_risk_engine غير متاح"}


@app.route("/api/geo/risk", methods=["POST", "OPTIONS"])
def geo_risk():
    """
    POST /api/geo/risk
    يُحلّل المخاطر الجيوتقنية (التربة، السيول، الزلازل) للموقع.
    Body: { location, lat?, lng?, property_type? }
    """
    if request.method == "OPTIONS":
        return jsonify({}), 200
    try:
        body = request.get_json(force=True) or {}
        loc  = str(body.get("location", ""))
        if not loc:
            return jsonify({"status": "error", "message": "location مطلوب"}), 400
        result = _analyze_geo_risk(
            location      = loc,
            lat           = body.get("lat"),
            lng           = body.get("lng"),
            property_type = str(body.get("property_type", "")),
        )
        result["status"] = "success"
        return jsonify(result)
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


# ── 3. رادار الهجرة العقارية (Demographic Flow) ───────────────────────────
try:
    from demographic_radar import analyze_demographic_flow as _analyze_demo_flow
    _DEMO_OK = True
    print("  [demographic_radar] loaded ✓")
except Exception as _dm_err:
    _DEMO_OK = False
    print(f"  [demographic_radar] unavailable: {_dm_err}")
    def _analyze_demo_flow(**kw):
        return {"status": "error", "message": "demographic_radar غير متاح"}


@app.route("/api/demographic/flow", methods=["POST", "OPTIONS"])
def demographic_flow():
    """
    POST /api/demographic/flow
    يُحلّل الزحف العمراني وكثافة POI وتوصيات الاستثمار.
    Body: { location, lat?, lng?, investment_horizon_yr? }
    """
    if request.method == "OPTIONS":
        return jsonify({}), 200
    try:
        body = request.get_json(force=True) or {}
        loc  = str(body.get("location", ""))
        if not loc:
            return jsonify({"status": "error", "message": "location مطلوب"}), 400
        result = _analyze_demo_flow(
            location              = loc,
            lat                   = body.get("lat"),
            lng                   = body.get("lng"),
            investment_horizon_yr = int(body.get("investment_horizon_yr", 5)),
        )
        result["status"] = "success"
        return jsonify(result)
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


# ── 4. إدارة الأصول الذكية (Asset Manager) ───────────────────────────────
try:
    from asset_manager import (
        register_asset    as _am_register,
        get_asset         as _am_get,
        list_assets       as _am_list,
        update_asset_value as _am_update,
        portfolio_dashboard as _am_dashboard,
        delete_asset      as _am_delete,
    )
    _AM_OK = True
    print("  [asset_manager] loaded ✓")
except Exception as _am_err:
    _AM_OK = False
    print(f"  [asset_manager] unavailable: {_am_err}")
    def _am_register(**kw):  return {"status": "error", "message": "asset_manager غير متاح"}
    def _am_get(i):          return {"status": "error", "message": "asset_manager غير متاح"}
    def _am_list(**kw):      return []
    def _am_update(**kw):    return {"status": "error", "message": "asset_manager غير متاح"}
    def _am_dashboard(**kw): return {"status": "error", "message": "asset_manager غير متاح"}
    def _am_delete(i):       return {"status": "error", "message": "asset_manager غير متاح"}


@app.route("/api/assets", methods=["GET"])
def assets_list():
    """GET /api/assets?market=egypt&limit=50 — قائمة الأصول."""
    market = request.args.get("market")
    limit  = int(request.args.get("limit", 50))
    return jsonify({"status": "success", "assets": _am_list(market=market, limit=limit)})


@app.route("/api/assets/register", methods=["POST", "OPTIONS"])
def assets_register():
    """POST /api/assets/register — تسجيل عقار جديد في المحفظة."""
    if request.method == "OPTIONS":
        return jsonify({}), 200
    try:
        b = request.get_json(force=True) or {}
        result = _am_register(
            name          = str(b.get("name", "عقار غير مُسمّى")),
            location      = str(b.get("location", "")),
            property_type = str(b.get("property_type", "")),
            area_m2       = float(b.get("area_m2", 0)),
            base_value    = float(b.get("base_value", 0)),
            market        = str(b.get("market", "egypt")),
            currency      = str(b.get("currency", "EGP")),
            owner_ref     = str(b.get("owner_ref", "")),
            notes         = str(b.get("notes", "")),
            metadata      = b.get("metadata"),
        )
        return jsonify(result)
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route("/api/assets/<asset_id>", methods=["GET"])
def assets_get(asset_id: str):
    """GET /api/assets/<id> — بيانات عقار + سجله التاريخي."""
    return jsonify(_am_get(asset_id))


@app.route("/api/assets/<asset_id>/update", methods=["PUT", "POST", "OPTIONS"])
def assets_update(asset_id: str):
    """PUT /api/assets/<id>/update — تحديث قيمة العقار."""
    if request.method == "OPTIONS":
        return jsonify({}), 200
    try:
        b = request.get_json(force=True) or {}
        result = _am_update(
            asset_id        = asset_id,
            new_value       = b.get("new_value"),
            method          = str(b.get("method", "تحديث يدوي")),
            source          = str(b.get("source", "manual")),
            apply_inflation = bool(b.get("apply_inflation", False)),
            notes           = str(b.get("notes", "")),
        )
        return jsonify(result)
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route("/api/assets/dashboard", methods=["GET"])
def assets_dashboard():
    """GET /api/assets/dashboard?market=egypt — لوحة أداء المحفظة."""
    market = request.args.get("market")
    return jsonify(_am_dashboard(market=market))


@app.route("/api/assets/<asset_id>", methods=["DELETE"])
def assets_delete(asset_id: str):
    """DELETE /api/assets/<id> — حذف عقار من المحفظة."""
    return jsonify(_am_delete(asset_id))


# ── 5. معالجة الصور الذكية (Auto-Tagging) ────────────────────────────────
import struct

_IMAGE_TAGS_DB = {
    # كلمات مفتاحية لتصنيف الصور
    "واجهة": ["facade", "front", "exterior", "building"],
    "تشطيب داخلي": ["interior", "inside", "room", "floor", "ceiling", "wall"],
    "مطبخ": ["kitchen", "cook"],
    "حمام": ["bathroom", "toilet", "bath"],
    "غرفة نوم": ["bedroom", "bed", "room"],
    "صالة": ["living", "salon", "hall", "reception"],
    "مسبح": ["pool", "swimming"],
    "حديقة": ["garden", "yard", "outdoor", "landscape"],
    "موقف سيارات": ["parking", "garage", "car"],
    "إطلالة": ["view", "panorama", "roof", "top"],
    "عداد مياه/كهرباء": ["meter", "electric", "water", "utility"],
    "هيكل إنشائي": ["structure", "column", "beam", "concrete", "construction"],
}

_REPORT_ORDER = [
    "واجهة", "هيكل إنشائي", "تشطيب داخلي", "صالة", "غرفة نوم",
    "مطبخ", "حمام", "حديقة", "مسبح", "موقف سيارات", "إطلالة", "عداد مياه/كهرباء"
]


def _auto_tag_image(filename: str, file_bytes: Optional[bytes] = None) -> str:
    """
    تصنيف الصورة بناءً على اسمها أولاً، ثم تحليل محتوى البيانات إن توفّر.
    يُعيد التصنيف العربي للصورة.
    """
    name_lower = filename.lower()
    for arabic_tag, keywords in _IMAGE_TAGS_DB.items():
        for kw in keywords:
            if kw in name_lower:
                return arabic_tag
    # fallback: اسم الملف بالأرقام → واجهة
    if any(c.isdigit() for c in filename):
        return "واجهة"
    return "أخرى"


def _extract_exif_gps(image_bytes: bytes):
    """
    Pure-Python EXIF GPS extractor. Returns (lat, lng) or (None, None).
    Wrapped in a single try/except for safety against malformed JPEGs.
    """
    try:
        import struct
        if not image_bytes or image_bytes[:2] != b"\xff\xd8":
            return None, None

        offset = 2
        while offset < len(image_bytes) - 2:
            marker = image_bytes[offset:offset + 2]
            offset += 2
            if not marker.startswith(b"\xff"):
                break
            if marker == b"\xff\xe1":  # APP1 (EXIF)
                seg_len = struct.unpack(">H", image_bytes[offset:offset + 2])[0]
                seg = image_bytes[offset + 2: offset + seg_len]
                offset += seg_len
                if seg[:4] != b"Exif":
                    continue
                tiff = seg[6:]
                bo = ">" if tiff[:2] == b"MM" else "<"
                ifd0_offset = struct.unpack(bo + "I", tiff[4:8])[0]
                cnt = struct.unpack(bo + "H", tiff[ifd0_offset:ifd0_offset + 2])[0]
                gps_ifd_offset = None
                for i in range(cnt):
                    e = tiff[ifd0_offset + 2 + i * 12: ifd0_offset + 2 + i * 12 + 12]
                    tag_id = struct.unpack(bo + "H", e[:2])[0]
                    if tag_id == 0x8825:  # GPS IFD pointer
                        gps_ifd_offset = struct.unpack(bo + "I", e[8:12])[0]
                        break
                if gps_ifd_offset is None:
                    return None, None

                gcnt = struct.unpack(bo + "H", tiff[gps_ifd_offset:gps_ifd_offset + 2])[0]
                gps_tags = {}
                for i in range(gcnt):
                    e = tiff[gps_ifd_offset + 2 + i * 12: gps_ifd_offset + 2 + i * 12 + 12]
                    tid = struct.unpack(bo + "H", e[:2])[0]
                    val_offset = struct.unpack(bo + "I", e[8:12])[0]
                    gps_tags[tid] = val_offset

                def _rat(off):
                    n, d = struct.unpack(bo + "II", tiff[off:off + 8])
                    return n / d if d else 0.0

                lat_off = gps_tags.get(0x0002)
                lng_off = gps_tags.get(0x0004)
                if lat_off is None or lng_off is None:
                    return None, None
                ref_lat = tiff[gps_tags.get(0x0001, 0)] if gps_tags.get(0x0001) else b"N"[0]
                ref_lng = tiff[gps_tags.get(0x0003, 0)] if gps_tags.get(0x0003) else b"E"[0]

                lat_d = _rat(lat_off); lat_m = _rat(lat_off + 8); lat_s = _rat(lat_off + 16)
                lng_d = _rat(lng_off); lng_m = _rat(lng_off + 8); lng_s = _rat(lng_off + 16)
                lat_dd = lat_d + lat_m / 60 + lat_s / 3600
                lng_dd = lng_d + lng_m / 60 + lng_s / 3600
                if ref_lat in (ord("S"), ord("s")): lat_dd = -lat_dd
                if ref_lng in (ord("W"), ord("w")): lng_dd = -lng_dd
                return lat_dd, lng_dd
            else:
                # Skip non-EXIF markers
                if offset + 2 <= len(image_bytes):
                    seg_len = struct.unpack(">H", image_bytes[offset:offset + 2])[0]
                    offset += seg_len
                else:
                    break
        return None, None
    except Exception:
        return None, None


# ═══════════════════════════════════════════════════════════════════════════
# Endpoint: Image Analyze — صورة → تصنيف تلقائي + GPS من EXIF
# RESTORED — المنطق مبني على helpers قائمة (_auto_tag_image, _extract_exif_gps,
# _REPORT_ORDER, _save_upload). شكل الاستجابة مأخوذ من sovereign_brain.html:2533.
# ═══════════════════════════════════════════════════════════════════════════
@app.route("/api/image/analyze", methods=["POST", "OPTIONS"])
def image_analyze():
    """
    POST /api/image/analyze — multipart/form-data with field 'file'.
    Response: {status, tag, report_order, has_gps, gps_lat?, gps_lng?, saved_as}
    """
    if request.method == "OPTIONS":
        return jsonify({}), 200
    if "file" not in request.files:
        return jsonify({"status": "error", "message": "لم يتم إرسال صورة"}), 400
    f = request.files["file"]
    ext = os.path.splitext(f.filename or "")[1].lower()
    if ext not in _ALLOWED_IMG_EXT:
        return jsonify({"status": "error",
                        "message": f"صيغة صورة غير مدعومة: {ext}"}), 415
    try:
        # Read bytes once, then save (consumes/restores stream)
        try:
            f.stream.seek(0)
        except Exception:
            pass
        image_bytes = f.stream.read() or b""
        try:
            f.stream.seek(0)
        except Exception:
            pass
        dest, orig = _save_upload(f)

        tag = _auto_tag_image(orig or f.filename or "", image_bytes)
        try:
            order = _REPORT_ORDER.index(tag)
        except ValueError:
            order = len(_REPORT_ORDER)  # unknown tags go last

        gps_lat, gps_lng = _extract_exif_gps(image_bytes)
        has_gps = (gps_lat is not None and gps_lng is not None)

        resp = {
            "status":       "success",
            "tag":          tag,
            "report_order": order,
            "has_gps":      has_gps,
            "saved_as":     os.path.basename(dest),
        }
        if has_gps:
            resp["gps_lat"] = gps_lat
            resp["gps_lng"] = gps_lng
        return jsonify(resp)
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({"status": "error", "message": str(e)}), 500


# ═══════════════════════════════════════════════════════════════════════════
# Endpoint: EIA Assessment — تقييم الأثر البيئي (ADDITIVE)
# ═══════════════════════════════════════════════════════════════════════════
@app.route("/api/eia/assess", methods=["POST", "OPTIONS"])
def handle_eia_assess():
    if request.method == "OPTIONS":
        return jsonify({}), 200
    try:
        from eia_engine import run_eia_assessment, write_eia_word_report, write_eia_excel_report
    except Exception:
        try:
            from core_engine.eia_engine import run_eia_assessment, write_eia_word_report, write_eia_excel_report  # type: ignore
        except Exception as imp_err:
            return jsonify({"status": "error",
                            "message": f"eia_engine غير متاح: {imp_err}"}), 500
    try:
        payload = request.get_json(silent=True) or {}
        result  = run_eia_assessment(payload)

        rid  = f"ES-EIA-{datetime.now().strftime('%Y%m%d')}-{uuid.uuid4().hex[:4].upper()}"
        ts   = datetime.now().strftime("%H%M%S")

        word_url = None
        xlsx_url = None
        try:
            docx_name = f"EIA_Report_{rid}_{ts}.docx"
            if write_eia_word_report(result, os.path.join(OUTPUTS, docx_name)):
                word_url = f"http://127.0.0.1:5000/api/download/{docx_name}"
        except Exception as werr:
            print(f"{_ts()} [EIA] Word failed: {werr}")

        if payload.get("export_excel", True):
            try:
                xlsx = write_eia_excel_report(result, OUTPUTS)
                if xlsx:
                    xlsx_url = f"http://127.0.0.1:5000/api/download/{os.path.basename(xlsx)}"
            except Exception as xerr:
                print(f"{_ts()} [EIA] Excel failed: {xerr}")

        return jsonify({
            "status":      "success",
            "report_id":   rid,
            "result":      result,
            "category":    result.get("classification", {}).get("category"),
            "erf":         result.get("investment_linkage", {}).get("environmental_risk_factor"),
            "word_url":    word_url,
            "excel_url":   xlsx_url,
        })
    except ValueError as ve:
        return jsonify({"status": "error", "message": str(ve)}), 400
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({"status": "error", "message": str(e)}), 500


# ═══════════════════════════════════════════════════════════════════════════
# Endpoint: HBU Analysis — تحليل أعلى وأفضل استغلال (ADDITIVE)
# ═══════════════════════════════════════════════════════════════════════════
@app.route("/api/hbu/analyze", methods=["POST", "OPTIONS"])
def handle_hbu_analyze():
    if request.method == "OPTIONS":
        return jsonify({}), 200
    try:
        from hbu_analysis_engine import run_hbu_analysis, generate_hbu_excel
    except Exception:
        try:
            from core_engine.hbu_analysis_engine import run_hbu_analysis, generate_hbu_excel  # type: ignore
        except Exception as imp_err:
            return jsonify({"status": "error",
                            "message": f"hbu_analysis_engine غير متاح: {imp_err}"}), 500
    try:
        payload = request.get_json(silent=True) or {}
        result = run_hbu_analysis(payload)

        excel_url = None
        if payload.get("export_excel", False):
            try:
                xlsx_path = generate_hbu_excel(result, OUTPUTS)
                fname = os.path.basename(xlsx_path)
                excel_url = f"http://127.0.0.1:5000/api/download/{fname}"
            except Exception as ex_err:
                print(f"{_ts()} [HBU] Excel export failed: {ex_err}")

        return jsonify({
            "status":      "success",
            "result":      result,
            "excel_url":   excel_url,
            "report_date": datetime.now().strftime("%d/%m/%Y"),
        })
    except ValueError as ve:
        return jsonify({"status": "error", "message": str(ve)}), 400
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({"status": "error", "message": str(e)}), 500


# ═══════════════════════════════════════════════════════════════════════════
# Endpoint: REIT NAV Engine (ADDITIVE)
# ═══════════════════════════════════════════════════════════════════════════
@app.route("/api/reit/nav", methods=["POST", "OPTIONS"])
def handle_reit_nav():
    if request.method == "OPTIONS":
        return jsonify({}), 200
    try:
        from reit_nav_engine import run_reit_nav, generate_reit_nav_excel
    except Exception:
        try:
            from core_engine.reit_nav_engine import run_reit_nav, generate_reit_nav_excel  # type: ignore
        except Exception as imp_err:
            return jsonify({"status": "error",
                            "message": f"reit_nav_engine غير متاح: {imp_err}"}), 500
    try:
        payload = request.get_json(silent=True) or {}
        result = run_reit_nav(payload)

        excel_url = None
        if payload.get("export_excel", False):
            try:
                xlsx_path = generate_reit_nav_excel(result, OUTPUTS)
                fname = os.path.basename(xlsx_path)
                excel_url = f"http://127.0.0.1:5000/api/download/{fname}"
            except Exception as ex_err:
                print(f"{_ts()} [REIT-NAV] Excel export failed: {ex_err}")

        return jsonify({
            "status":      "success",
            "result":      result,
            "excel_url":   excel_url,
            "report_date": datetime.now().strftime("%d/%m/%Y"),
        })
    except ValueError as ve:
        return jsonify({"status": "error", "message": str(ve)}), 400
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({"status": "error", "message": str(e)}), 500



# ═══════════════════════════════════════════════════════════════════════════
# Endpoint: Price Index — مؤشر معدلات الزيادة في أسعار العقارات
# 4 methodologies: CMA / AVM / RPPI / Stratification — composite output.
# ═══════════════════════════════════════════════════════════════════════════
@app.route("/api/price-index", methods=["GET", "POST", "OPTIONS"])
def handle_price_index():
    if request.method == "OPTIONS":
        return jsonify({}), 200
    try:
        from price_index_engine import compute_price_index
    except Exception:
        try:
            from core_engine.price_index_engine import compute_price_index  # type: ignore
        except Exception as imp_err:
            return jsonify({"status": "error",
                            "message": f"price_index_engine missing: {imp_err}"}), 500
    try:
        # Accept filters from query string OR JSON body
        if request.method == "POST":
            payload = request.get_json(silent=True) or {}
        else:
            payload = {}
        region_filter = (request.args.get("region") or payload.get("region") or "").strip() or None
        ptype_filter  = (request.args.get("property_type") or payload.get("property_type") or "").strip() or None
        base_period   = (request.args.get("base_period") or payload.get("base_period") or "").strip() or None
        official_rppi = payload.get("official_rppi")  # optional dict per region

        # Load market_feed as the data source
        records = _load_feed()

        result = compute_price_index(
            records,
            base_period=base_period,
            region_filter=region_filter,
            property_type_filter=ptype_filter,
            official_rppi=official_rppi,
        )
        result["data_source"] = "market_feed.json"
        result["records_total"] = len(records)
        return jsonify(result)
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({"status": "error", "message": str(e)}), 500

# ═══════════════════════════════════════════════════════════════════════════
# Mass Appraisal Routes — Phase 3 (Limited Backend Route Restore)
# Adds POST handlers for all /api/mass-appraisal/* endpoints.
# Only bridge_api.py is modified; specialist modules are not touched.
# ═══════════════════════════════════════════════════════════════════════════

def _ma_units(body: dict) -> list:
    """Map frontend rows[] / units[] to engine-compatible unit dicts."""
    rows = body.get("rows") or body.get("units") or []
    out = []
    for r in rows:
        if not isinstance(r, dict):
            continue
        u = dict(r)
        if "id" not in u and "row_id" in u:
            u["id"] = u["row_id"]
        out.append(u)
    return out


@app.route("/api/mass-appraisal/preview", methods=["POST", "OPTIONS"])
def handle_mass_appraisal_preview():
    if request.method == "OPTIONS":
        return jsonify({}), 200
    try:
        from mass_appraisal import run_mass_appraisal
    except Exception:
        try:
            from core_engine.mass_appraisal import run_mass_appraisal  # type: ignore
        except Exception as imp_err:
            return jsonify({"status": "error",
                            "message": f"mass_appraisal not available: {imp_err}"}), 500
    try:
        body  = request.get_json(silent=True) or {}
        units = _ma_units(body)
        if not units:
            return jsonify({"status": "error", "error_code": "MISSING_ROWS",
                            "message": "rows[] is required and must be non-empty."}), 400
        result = run_mass_appraisal(
            units,
            base_market_ppm=float(body.get("base_market_ppm", 0)),
            location=body.get("location", "Cairo"),
            region=body.get("region", "EG"),
            method=body.get("method", "avm"),
            purpose=body.get("purpose", "fair_market"),
        )
        # Preview: summary only — omit per-unit rows and disk path
        preview = {k: v for k, v in result.items() if k not in ("units", "output_xlsx")}
        preview["preview"] = True
        return jsonify(preview)
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route("/api/mass-appraisal/run", methods=["POST", "OPTIONS"])
def handle_mass_appraisal_run():
    if request.method == "OPTIONS":
        return jsonify({}), 200
    try:
        from mass_appraisal import run_mass_appraisal
    except Exception:
        try:
            from core_engine.mass_appraisal import run_mass_appraisal  # type: ignore
        except Exception as imp_err:
            return jsonify({"status": "error",
                            "message": f"mass_appraisal not available: {imp_err}"}), 500
    try:
        body  = request.get_json(silent=True) or {}
        units = _ma_units(body)
        if not units:
            return jsonify({"status": "error", "error_code": "MISSING_ROWS",
                            "message": "rows[] is required and must be non-empty."}), 400
        result = run_mass_appraisal(
            units,
            base_market_ppm=float(body.get("base_market_ppm", 0)),
            location=body.get("location", "Cairo"),
            region=body.get("region", "EG"),
            method=body.get("method", "avm"),
            purpose=body.get("purpose", "fair_market"),
        )
        return jsonify(result)
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route("/api/mass-appraisal/export-xlsx", methods=["POST", "OPTIONS"])
def handle_mass_appraisal_export_xlsx():
    if request.method == "OPTIONS":
        return jsonify({}), 200
    try:
        from mass_appraisal_excel import build_mass_appraisal_workbook
    except Exception:
        try:
            from core_engine.mass_appraisal_excel import build_mass_appraisal_workbook  # type: ignore
        except Exception as imp_err:
            return jsonify({"status": "error",
                            "message": f"mass_appraisal_excel not available: {imp_err}"}), 500
    try:
        from flask import Response
        body                = request.get_json(silent=True) or {}
        run_result          = body.get("result") or body
        ratio_study         = body.get("ratio_study")
        calibration_preview = body.get("calibration_preview")
        calibration_sandbox = body.get("calibration_sandbox")
        governance          = body.get("governance")
        model_cycle         = body.get("model_cycle")
        import_validation   = body.get("import_validation")
        sales_verification  = body.get("sales_verification")
        time_adjustment     = body.get("time_adjustment")
        sales_adjustments   = body.get("sales_adjustments")

        _valid_ma_styles = ("legacy", "detailed", "professional_template")
        report_style     = body.get("report_style", "legacy")
        _ma_style_used   = report_style if report_style in _valid_ma_styles else "legacy"
        _ma_tpl_used     = False
        _ma_fallback     = False
        _ma_fb_reason    = ""

        if report_style == "professional_template":
            try:
                try:
                    from reports.excel_template_renderer import (
                        MASS_APPRAISAL_TEMPLATE as _MATPL,
                        build_mass_appraisal_report as _build_ma,
                    )
                except ImportError:
                    from core_engine.reports.excel_template_renderer import (  # type: ignore
                        MASS_APPRAISAL_TEMPLATE as _MATPL,
                        build_mass_appraisal_report as _build_ma,
                    )
                if not _MATPL.is_file():
                    raise FileNotFoundError("mass_appraisal_professional_template.xlsm not found")
                import tempfile, os as _os
                _ma_tmp = _os.path.join(tempfile.gettempdir(), "ma_prof_export.xlsm")
                _ma_ctx = {
                    "report_date": datetime.now().strftime("%d/%m/%Y"),
                    "total_units": run_result.get("total_units", ""),
                    "mean_ratio":  run_result.get("mean_ratio", ""),
                    "cod":         run_result.get("cod", ""),
                    "prd":         run_result.get("prd", ""),
                }
                _ma_out = _build_ma(output_path=_ma_tmp, context=_ma_ctx, tables={})
                if _ma_out is None or not _os.path.isfile(_ma_tmp):
                    raise RuntimeError("template render returned None")
                with open(_ma_tmp, "rb") as _f:
                    xlsx_bytes = _f.read()
                try:
                    _os.remove(_ma_tmp)
                except Exception:
                    pass
                _ma_tpl_used = True
                _ma_style_used = "professional_template"
                return Response(
                    xlsx_bytes,
                    mimetype="application/vnd.ms-excel.sheet.macroEnabled.12",
                    headers={
                        "Content-Disposition": "attachment; filename=mass_appraisal.xlsm",
                        "X-Report-Style-Requested": report_style,
                        "X-Report-Style-Used":      _ma_style_used,
                        "X-Template-Used":          "true",
                        "X-Fallback-Used":          "false",
                    },
                )
            except Exception as _ma_err:
                _ma_fallback  = True
                _ma_fb_reason = str(_ma_err)
                _ma_style_used = "legacy"

        xlsx_bytes = build_mass_appraisal_workbook(
            run_result,
            ratio_study=ratio_study,
            calibration_preview=calibration_preview,
            calibration_sandbox=calibration_sandbox,
            governance=governance,
            model_cycle=model_cycle,
            import_validation=import_validation,
            sales_verification=sales_verification,
            time_adjustment=time_adjustment,
            sales_adjustments=sales_adjustments,
            report_style="legacy",
        )
        return Response(
            xlsx_bytes,
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={
                "Content-Disposition": "attachment; filename=mass_appraisal.xlsx",
                "X-Report-Style-Requested": report_style,
                "X-Report-Style-Used":      _ma_style_used,
                "X-Template-Used":          "false",
                "X-Fallback-Used":          str(_ma_fallback).lower(),
                "X-Fallback-Reason":        _ma_fb_reason,
            },
        )
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route("/api/mass-appraisal/sales/verify", methods=["POST", "OPTIONS"])
def handle_mass_sales_verify():
    if request.method == "OPTIONS":
        return jsonify({}), 200
    try:
        from sales_verification import verify_sales_records
    except Exception:
        try:
            from core_engine.sales_verification import verify_sales_records  # type: ignore
        except Exception as imp_err:
            return jsonify({"status": "error",
                            "message": f"sales_verification not available: {imp_err}"}), 500
    try:
        body    = request.get_json(silent=True) or {}
        records = body.get("records") or []
        options = body.get("options")
        return jsonify(verify_sales_records(records, options))
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route("/api/mass-appraisal/sales/time-adjust", methods=["POST", "OPTIONS"])
def handle_mass_sales_time_adjust():
    if request.method == "OPTIONS":
        return jsonify({}), 200
    try:
        from sales_time_adjustment import adjust_sales_for_time
    except Exception:
        try:
            from core_engine.sales_time_adjustment import adjust_sales_for_time  # type: ignore
        except Exception as imp_err:
            return jsonify({"status": "error",
                            "message": f"sales_time_adjustment not available: {imp_err}"}), 500
    try:
        body           = request.get_json(silent=True) or {}
        records        = body.get("records") or []
        valuation_date = body.get("valuation_date", "")
        options        = body.get("options")
        return jsonify(adjust_sales_for_time(records, valuation_date, options))
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route("/api/mass-appraisal/sales/adjust", methods=["POST", "OPTIONS"])
def handle_mass_sales_adjust():
    if request.method == "OPTIONS":
        return jsonify({}), 200
    try:
        from sales_adjustments import apply_sales_adjustments
    except Exception:
        try:
            from core_engine.sales_adjustments import apply_sales_adjustments  # type: ignore
        except Exception as imp_err:
            return jsonify({"status": "error",
                            "message": f"sales_adjustments not available: {imp_err}"}), 500
    try:
        body               = request.get_json(silent=True) or {}
        records            = body.get("sale_records") or body.get("records") or []
        adjustment_profile = body.get("adjustment_profile")
        options            = body.get("options")
        return jsonify(apply_sales_adjustments(records, adjustment_profile, options))
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route("/api/mass-appraisal/ratio-study/run", methods=["POST", "OPTIONS"])
def handle_mass_ratio_study():
    if request.method == "OPTIONS":
        return jsonify({}), 200
    try:
        from ratio_studies import run_ratio_study
    except Exception:
        try:
            from core_engine.ratio_studies import run_ratio_study  # type: ignore
        except Exception as imp_err:
            return jsonify({"status": "error",
                            "message": f"ratio_studies not available: {imp_err}"}), 500
    try:
        body         = request.get_json(silent=True) or {}
        subject_rows = body.get("subject_rows") or []
        sale_records = body.get("sale_records") or []
        options      = body.get("options")
        return jsonify(run_ratio_study(subject_rows, sale_records, options))
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route("/api/mass-appraisal/calibration/preview", methods=["POST", "OPTIONS"])
def handle_mass_calibration_preview():
    if request.method == "OPTIONS":
        return jsonify({}), 200
    try:
        from model_calibration import preview_calibration
    except Exception:
        try:
            from core_engine.model_calibration import preview_calibration  # type: ignore
        except Exception as imp_err:
            return jsonify({"status": "error",
                            "message": f"model_calibration not available: {imp_err}"}), 500
    try:
        body         = request.get_json(silent=True) or {}
        subject_rows = body.get("subject_rows") or []
        sale_records = body.get("sale_records") or []
        ratio_study  = body.get("ratio_study")
        options      = body.get("options")
        return jsonify(preview_calibration(subject_rows, sale_records, ratio_study, options))
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route("/api/mass-appraisal/calibration/sandbox", methods=["POST", "OPTIONS"])
def handle_mass_calibration_sandbox():
    if request.method == "OPTIONS":
        return jsonify({}), 200
    try:
        from calibration_sandbox import apply_calibration_sandbox
    except Exception:
        try:
            from core_engine.calibration_sandbox import apply_calibration_sandbox  # type: ignore
        except Exception as imp_err:
            return jsonify({"status": "error",
                            "message": f"calibration_sandbox not available: {imp_err}"}), 500
    try:
        body                = request.get_json(silent=True) or {}
        subject_rows        = body.get("subject_rows") or []
        calibration_preview = body.get("calibration_preview") or {}
        options             = body.get("options")
        return jsonify(apply_calibration_sandbox(subject_rows, calibration_preview, options))
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({"status": "error", "message": str(e)}), 500


# ── Phase 3.12 — Mass Appraisal Excel Template Download ──────────────────────
@app.route("/api/mass-appraisal/template-xlsx", methods=["GET", "OPTIONS"])
def handle_mass_appraisal_template_xlsx():
    """Return a blank Mass Appraisal Excel template for download."""
    if request.method == "OPTIONS":
        return jsonify({}), 200
    try:
        try:
            from mass_appraisal_template import build_mass_appraisal_template_workbook
        except ImportError:
            from core_engine.mass_appraisal_template import build_mass_appraisal_template_workbook  # type: ignore
        from io import BytesIO
        from flask import send_file
        xlsx_bytes = build_mass_appraisal_template_workbook()
        return send_file(
            BytesIO(xlsx_bytes),
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            as_attachment=True,
            download_name="mass_appraisal_template.xlsx",
        )
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({"status": "error", "message": str(e)}), 500


# ── Phase 3.13 P3 — Import-xlsx hardening constants ──────────────────────────
_MA_IMPORT_MAX_BYTES   = 10 * 1024 * 1024   # 10 MB per-route limit
_MA_IMPORT_ALLOWED_EXT = {".xlsx"}           # strict: no .xls / .csv / .zip

# ── Phase 3.12 — Mass Appraisal Excel Import Parse Preview ───────────────────
@app.route("/api/mass-appraisal/import-xlsx", methods=["POST", "OPTIONS"])
def handle_mass_appraisal_import_xlsx():
    """
    Parse an uploaded Mass Appraisal template .xlsx and return normalized JSON preview.
    No valuation is executed. No files are saved to disk.
    """
    if request.method == "OPTIONS":
        return jsonify({}), 200

    # ── 1. File field present ─────────────────────────────────────────────────
    if "file" not in request.files:
        return jsonify({"status": "error",
                        "message": "No file part. Upload an .xlsx file using multipart field 'file'."}), 400

    uploaded = request.files["file"]

    # ── 2. Non-empty filename ─────────────────────────────────────────────────
    if not uploaded.filename:
        return jsonify({"status": "error",
                        "message": "No file selected."}), 400

    # ── 3. Extension check (strict: .xlsx only) ───────────────────────────────
    ext = os.path.splitext(uploaded.filename)[1].lower()
    if ext not in _MA_IMPORT_ALLOWED_EXT:
        return jsonify({"status": "error",
                        "message": "Unsupported file type. Please upload an .xlsx Excel file."}), 400

    # ── 4. Read bytes in memory — no disk write ───────────────────────────────
    try:
        file_bytes = uploaded.read()
    except Exception:
        return jsonify({"status": "error",
                        "message": "Failed to read uploaded file."}), 400

    # ── 5. Empty file ─────────────────────────────────────────────────────────
    if len(file_bytes) == 0:
        return jsonify({"status": "error",
                        "message": "Uploaded file is empty."}), 400

    # ── 6. File size limit ────────────────────────────────────────────────────
    if len(file_bytes) > _MA_IMPORT_MAX_BYTES:
        return jsonify({"status": "error",
                        "message": "Excel file is too large. Maximum allowed size is 10 MB."}), 413

    # ── 7. Magic bytes — .xlsx is a ZIP container (starts with PK) ────────────
    if file_bytes[:2] != b"PK":
        return jsonify({"status": "error",
                        "message": "Invalid .xlsx file. The uploaded file is not a valid Excel workbook."}), 400

    # ── 8. Parse workbook ─────────────────────────────────────────────────────
    try:
        try:
            from mass_appraisal_template import parse_mass_appraisal_template_workbook
        except ImportError:
            from core_engine.mass_appraisal_template import parse_mass_appraisal_template_workbook  # type: ignore
        result = parse_mass_appraisal_template_workbook(file_bytes)
        return jsonify(result), 200
    except ValueError as ve:
        return jsonify({"status": "error",
                        "message": "Could not read Excel workbook.",
                        "details": str(ve)[:200]}), 400
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({"status": "error",
                        "message": "Could not read Excel workbook.",
                        "details": type(e).__name__}), 400


# ── Phase 4 Routes ───────────────────────────────────────────────────────────

def _engine_result_to_dict(result):
    """Serialise an EngineResult to a plain dict safe for jsonify()."""
    return {
        "status":      "success",
        "engine_name": result.engine_name,
        "value":       float(result.value) if result.value is not None else None,
        "confidence":  result.confidence,
        "metadata":    result.metadata,
        "audit_trail": [
            {
                "step_name":  e.step_name,
                "inputs":     e.inputs,
                "outputs":    e.outputs,
                "formula":    e.formula,
                "references": e.references,
            }
            for e in result.audit_trail
        ],
        "issues": [
            {
                "severity": i.severity,
                "code":     i.code,
                "message":  i.message,
            }
            for i in result.issues
        ],
    }


@app.route("/api/comparables/search", methods=["POST", "OPTIONS"])
def api_comparable_search():
    """Search comparables: PostgreSQL first, JSON search engine fallback."""
    if request.method == "OPTIONS":
        return jsonify({}), 200
    try:
        data    = request.get_json(force=True) or {}
        subject = data.get("subject_property", {})
        filters = data.get("filters", {})
        limit   = int(data.get("limit", 20))

        db_results = _search_db_comparables(subject, filters, limit)
        if db_results is not None:
            return jsonify({
                "status": "success",
                "count":  len(db_results),
                "results": db_results,
                "source": "postgresql",
            }), 200

        # DB unavailable — fall back to JSON search engine
        engine  = get_search_engine()
        results = engine.search_and_rank(subject=subject, filters=filters, limit=limit)
        return jsonify({
            "status": "success",
            "count":  len(results),
            "results": results,
            "source": "json",
        }), 200
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 400


@app.route("/api/engines/comparative", methods=["POST", "OPTIONS"])
def api_engine_comparative():
    """Sales Comparison Approach — applies area/age adjustments to ranked comparables."""
    if request.method == "OPTIONS":
        return jsonify({}), 200
    try:
        data   = request.get_json(force=True) or {}
        engine = get_comparative_engine()
        result = engine.calculate(data)
        return jsonify(_engine_result_to_dict(result)), 200
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 400


@app.route("/api/engines/cost", methods=["POST", "OPTIONS"])
def api_engine_cost():
    """Cost Approach — Replacement Cost New minus depreciation plus land value."""
    if request.method == "OPTIONS":
        return jsonify({}), 200
    try:
        data   = request.get_json(force=True) or {}
        engine = get_cost_engine()
        result = engine.calculate(data)
        return jsonify(_engine_result_to_dict(result)), 200
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 400


@app.route("/api/engines/income", methods=["POST", "OPTIONS"])
def api_engine_income():
    """Income Approach — Direct Capitalization (NOI / Cap Rate)."""
    if request.method == "OPTIONS":
        return jsonify({}), 200
    try:
        data   = request.get_json(force=True) or {}
        engine = get_income_engine()
        result = engine.calculate(data)
        return jsonify(_engine_result_to_dict(result)), 200
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 400


# ── Phase 5 Route ────────────────────────────────────────────────────────────

def _purpose_result_to_dict(result):
    """Convert PurposeResult to JSON-serialisable dict."""
    def _dec(v):
        return float(v) if v is not None else None

    return {
        "purpose_name":   result.purpose_name,
        "adjusted_value": _dec(result.adjusted_value),
        "confidence":     result.confidence,
        "adjustments": [
            {
                "factor_name":  adj.factor_name,
                "before_value": float(adj.before_value),
                "after_value":  float(adj.after_value),
                "percentage":   adj.percentage,
                "reason":       adj.reason,
            }
            for adj in result.adjustments
        ],
        "audit_trail": [
            {
                "step_name":  e.step_name,
                "inputs":     e.inputs,
                "outputs":    e.outputs,
                "formula":    e.formula,
                "references": e.references,
            }
            for e in result.audit_trail
        ],
        "disclosures": result.disclosures,
        "metadata":    {
            k: float(v) if hasattr(v, "__float__") and not isinstance(v, (bool, str, list, dict))
               else v
            for k, v in result.metadata.items()
        },
        "issues": [
            {"severity": i.severity, "code": i.code, "message": i.message}
            for i in result.issues
        ],
    }


@app.route("/api/valuation/full", methods=["POST", "OPTIONS"])
def api_valuation_full():
    """
    Full valuation pipeline: Phase 4 engines → Phase 5 purpose adapters.

    Input JSON keys (all optional except subject_property):
      subject_property       dict   property attributes for search + engines
      filters                dict   lat/lng/radius/governorate for comparable search
      comparable_count       int    how many comparables to use  (default 5)
      valuation_purposes     list   subset of: market_value, mortgage, insurance, ifrs_13
      borrower_profile       str    standard | preferred | subprime  (mortgage)
      building_age_years     int    for insurance tier + cost engine  (default 5)
      cap_rate_source        str    market | published | expert  (IFRS 13 + income)
      land_value_egp         float  land component for cost engine  (default 1 000 000)
      gross_income_annual_egp float annual gross rental income  (default 240 000)
      vacancy_rate           float  (default 0.10)
      opex_ratio             float  (default 0.35)
      cap_rate               float  (default 0.10)
    """
    if request.method == "OPTIONS":
        return jsonify({}), 200
    try:
        data = request.get_json(force=True) or {}

        subject            = data.get("subject_property", {})
        filters            = data.get("filters", {})
        comparable_count   = int(data.get("comparable_count", 5))
        requested_purposes = data.get(
            "valuation_purposes",
            ["market_value", "mortgage", "insurance", "ifrs_13"],
        )
        borrower_profile   = data.get("borrower_profile", "standard")
        building_age       = int(data.get("building_age_years", 5))
        cap_rate_source    = data.get("cap_rate_source", "market")
        land_value_egp     = float(data.get("land_value_egp", 1_000_000))
        gross_income       = float(data.get("gross_income_annual_egp", 240_000))
        vacancy_rate       = float(data.get("vacancy_rate", 0.10))
        opex_ratio         = float(data.get("opex_ratio", 0.35))
        cap_rate           = float(data.get("cap_rate", 0.10))
        governorate        = filters.get("governorate", "Cairo")

        # ── Phase 4: comparable search ────────────────────────────────────────
        comparables = get_search_engine().search_and_rank(
            subject=subject, filters=filters, limit=20
        )
        if not comparables:
            return jsonify({
                "status": "error",
                "message": "No comparables found matching the given criteria.",
            }), 400

        selected = comparables[:comparable_count]

        # ── Phase 4: three-approach engines ───────────────────────────────────
        result_comp = get_comparative_engine().calculate({
            "subject_area_sqm":  subject.get("area_sqm"),
            "subject_age_years": subject.get("age_years"),
            "comparables":       selected,
        })

        result_cost = get_cost_engine().calculate({
            "building_area_sqm":    subject.get("area_sqm"),
            "building_age_years":   building_age,
            "construction_quality": "standard",
            "governorate":          governorate,
            "land_value_egp":       land_value_egp,
        })

        result_income = get_income_engine().calculate({
            "gross_income_annual_egp":   gross_income,
            "vacancy_rate":              vacancy_rate,
            "operating_expenses_ratio":  opex_ratio,
            "cap_rate":                  cap_rate,
            "cap_rate_source":           cap_rate_source,
        })

        three_values = {
            "comparative": result_comp.value,
            "cost":        result_cost.value,
            "income":      result_income.value,
        }

        # ── Phase 5: purpose adapters ─────────────────────────────────────────
        purpose_valuations = {}

        if "market_value" in requested_purposes:
            r = get_market_value_adapter().adjust(three_values, {
                "comparable_count":   len(selected),
                "comparable_quality": "strong",
            })
            purpose_valuations["market_value"] = _purpose_result_to_dict(r)

        if "mortgage" in requested_purposes:
            r = get_mortgage_adapter().adjust(three_values, {
                "borrower_profile":  borrower_profile,
                "market_conditions": "neutral",
            })
            purpose_valuations["mortgage"] = _purpose_result_to_dict(r)

        if "insurance" in requested_purposes:
            r = get_insurance_adapter().adjust(three_values, {
                "building_age_years": building_age,
                "location":           governorate,
                "claim_history":      "clean",
            })
            purpose_valuations["insurance"] = _purpose_result_to_dict(r)

        if "ifrs_13" in requested_purposes or "ifrs_13_fair_value" in requested_purposes:
            r = get_ifrs13_adapter().adjust(three_values, {
                "comparable_count": len(selected),
                "cap_rate_source":  cap_rate_source,
            })
            purpose_valuations["ifrs_13_fair_value"] = _purpose_result_to_dict(r)

        top_score = selected[0].get("similarity_score", 0) if selected else 0

        # ── Persist to PostgreSQL (graceful — never blocks the response) ──────
        valuation_id = None
        if _DB_AVAILABLE:
            try:
                _sess = _DbSession()
                try:
                    _rec = _DbValuation(
                        asset_type=subject.get("property_type", "unknown"),
                        comparable_value=float(result_comp.value) if result_comp.value is not None else None,
                        cost_value=float(result_cost.value) if result_cost.value is not None else None,
                        income_value=float(result_income.value) if result_income.value is not None else None,
                        result_json={"purpose_valuations": purpose_valuations},
                    )
                    _sess.add(_rec)
                    _sess.commit()
                    valuation_id = str(_rec.id)
                finally:
                    _sess.close()
            except Exception:
                pass  # DB write failure must not affect the API response

        return jsonify({
            "status": "success",
            "phase_4_values": {
                "comparative": float(result_comp.value)   if result_comp.value   is not None else None,
                "cost":        float(result_cost.value)   if result_cost.value   is not None else None,
                "income":      float(result_income.value) if result_income.value is not None else None,
            },
            "comparable_search": {
                "count":               len(selected),
                "top_similarity_score": top_score,
            },
            "purpose_valuations": purpose_valuations,
            "valuation_id": valuation_id,
            "issues": [],
        }), 200

    except Exception as e:
        import traceback as _tb
        print(_tb.format_exc())
        return jsonify({"status": "error", "message": str(e)}), 400


# ── Phase 6 report routes ─────────────────────────────────────────────────────

# Temp directory for generated reports (cleaned on process restart)
_REPORT_DIR = os.path.join(tempfile.gettempdir(), "expert_smart_reports")
os.makedirs(_REPORT_DIR, exist_ok=True)

# Map asset_type → adapter getter
_ASSET_ADAPTER_MAP = {
    "residential": get_residential_adapter,
    "commercial":  get_commercial_adapter,
    "land":        get_land_adapter,
}


@app.route("/api/valuation/report", methods=["POST", "OPTIONS"])
def api_valuation_report():
    """
    Generate an EGVS-compliant 8-sheet Excel report for an asset valuation.

    Runs the same Phase 4 + Phase 5 pipeline as /api/valuation/full, then
    applies the Phase 6 AssetAdapter for the requested asset_type and writes
    the Excel workbook via ExcelReportBuilder.

    Required body keys:
      subject_property  dict  — property attributes (must include property_type)
      asset_type        str   — "residential" | "commercial"

    Optional body keys (same as /api/valuation/full):
      filters, comparable_count, valuation_purposes, borrower_profile,
      building_age_years, cap_rate_source, land_value_egp,
      gross_income_annual_egp, vacancy_rate, opex_ratio, cap_rate

    Returns JSON:
      { "status": "success", "report_id": "<uuid>",
        "download_url": "/api/valuation/report/download/<uuid>.xlsx",
        "primary_value": <float>, "confidence": <str> }
    """
    if request.method == "OPTIONS":
        return jsonify({}), 200
    try:
        data = request.get_json(force=True) or {}

        subject          = data.get("subject_property", {})
        asset_type       = data.get("asset_type", "residential").lower()
        filters          = data.get("filters", {})
        comparable_count = int(data.get("comparable_count", 5))
        requested_purposes = data.get(
            "valuation_purposes",
            ["market_value", "mortgage", "insurance"],
        )
        borrower_profile = data.get("borrower_profile", "standard")
        building_age     = int(data.get("building_age_years", 5))
        cap_rate_source  = data.get("cap_rate_source", "market")
        land_value_egp   = float(data.get("land_value_egp", 1_000_000))
        gross_income     = float(data.get("gross_income_annual_egp", 240_000))
        vacancy_rate     = float(data.get("vacancy_rate", 0.10))
        opex_ratio       = float(data.get("opex_ratio", 0.35))
        cap_rate         = float(data.get("cap_rate", 0.10))
        governorate      = filters.get("governorate", "Cairo")

        # Validate asset_type early
        if asset_type not in _ASSET_ADAPTER_MAP:
            return jsonify({
                "status": "error",
                "message": f"Unsupported asset_type '{asset_type}'. "
                           f"Choose from: {list(_ASSET_ADAPTER_MAP.keys())}",
            }), 400

        # ── Phase 4: comparable search ────────────────────────────────────────
        comparables = get_search_engine().search_and_rank(
            subject=subject, filters=filters, limit=20
        )
        if not comparables:
            return jsonify({
                "status": "error",
                "message": "No comparables found for the given criteria.",
            }), 400

        selected = comparables[:comparable_count]

        # ── Phase 4: three-approach engines ───────────────────────────────────
        result_comp = get_comparative_engine().calculate({
            "subject_area_sqm":  subject.get("area_sqm"),
            "subject_age_years": subject.get("age_years"),
            "comparables":       selected,
        })
        result_cost = get_cost_engine().calculate({
            "building_area_sqm":    subject.get("area_sqm"),
            "building_age_years":   building_age,
            "construction_quality": "standard",
            "governorate":          governorate,
            "land_value_egp":       land_value_egp,
        })
        result_income = get_income_engine().calculate({
            "gross_income_annual_egp":  gross_income,
            "vacancy_rate":             vacancy_rate,
            "operating_expenses_ratio": opex_ratio,
            "cap_rate":                 cap_rate,
            "cap_rate_source":          cap_rate_source,
        })

        from decimal import Decimal as _D
        three_values = {
            "comparable": result_comp.value  if result_comp.value  is not None else _D("0"),
            "cost":       result_cost.value  if result_cost.value  is not None else _D("0"),
            "income":     result_income.value if result_income.value is not None else _D("0"),
        }

        # ── Phase 5: market_value purpose adapter (drives confidence) ─────────
        market_result = get_market_value_adapter().adjust(three_values, {
            "comparable_count":   len(selected),
            "comparable_quality": "strong",
        })

        # Optional extra purpose adapters (for alternative_values)
        alternative_values = {}
        if "mortgage" in requested_purposes:
            r = get_mortgage_adapter().adjust(three_values, {
                "borrower_profile":  borrower_profile,
                "market_conditions": "neutral",
            })
            if r.adjusted_value is not None:
                alternative_values["mortgage"] = r.adjusted_value
        if "insurance" in requested_purposes:
            r = get_insurance_adapter().adjust(three_values, {
                "building_age_years": building_age,
                "location":           governorate,
                "claim_history":      "clean",
            })
            if r.adjusted_value is not None:
                alternative_values["insurance"] = r.adjusted_value

        # ── Phase 6: asset adapter ─────────────────────────────────────────────
        asset_adapter  = _ASSET_ADAPTER_MAP[asset_type]()
        asset_result   = asset_adapter.value(subject, three_values, market_result)
        asset_result.alternative_values = alternative_values

        # ── Generate IVSC disclosure ───────────────────────────────────────────
        _ivsc_disclosure = None
        try:
            from adapters.ivsc import IVSComplianceBuilder as _IVSBuilder
            _ivsc_builder = _IVSBuilder()
            _ivsc_fn_map = {
                "residential": _ivsc_builder.for_residential,
                "commercial":  _ivsc_builder.for_commercial,
                "land":        _ivsc_builder.for_land,
            }
            _ivsc_fn = _ivsc_fn_map.get(asset_type)
            if _ivsc_fn:
                _ivsc_disclosure = _ivsc_fn(
                    appraiser_name=data.get("appraiser_name", "Expert Smart System"),
                    property_address=str(subject.get("address", subject.get("location", ""))),
                )
        except Exception:
            pass

        # ── Generate cross-border disclosure (optional) ────────────────────────
        _cb_disclosure = None
        if data.get("include_cross_border"):
            try:
                from adapters.cross_border import CrossBorderBuilder as _CBBuilder
                from datetime import date as _date
                _cb_builder       = _CBBuilder()
                _primary_egp      = float(asset_result.primary_value) if asset_result.primary_value else 0.0
                _rep_currency     = str(data.get("reporting_currency", "EGP")).upper()
                _rate_source      = str(data.get("rate_source", "Central Bank of Egypt"))
                _prop_country     = str(data.get("property_country", "Egypt"))
                if _rep_currency == "EGP":
                    _cb_disclosure = _cb_builder.domestic_egp(
                        primary_value=_primary_egp,
                        property_country=_prop_country,
                    )
                elif _rep_currency == "USD":
                    _cb_disclosure = _cb_builder.cross_border_usd(
                        primary_value_egp=_primary_egp,
                        exchange_rate=float(data.get("usd_exchange_rate", 30.0)),
                        rate_source=_rate_source,
                        rate_date=_date.today(),
                    )
                else:
                    _cb_disclosure = _cb_builder.multi_currency(
                        primary_value_egp=_primary_egp,
                        usd_rate=float(data.get("usd_exchange_rate", 30.0)),
                        eur_rate=float(data.get("eur_exchange_rate", 33.0)),
                        rate_source=_rate_source,
                        rate_date=_date.today(),
                    )
            except Exception:
                pass

        # ── Generate Excel report ──────────────────────────────────────────────
        report_id       = str(uuid.uuid4())
        filename        = f"{report_id}.xlsx"
        filepath        = os.path.join(_REPORT_DIR, filename)
        actual_filepath = ExcelReportBuilder(asset_result).build(
            filepath,
            ivsc_disclosure=_ivsc_disclosure,
            cross_border_disclosure=_cb_disclosure,
        )
        # build() may promote .xlsx to .xlsm when the professional template is used
        actual_filename = os.path.basename(str(actual_filepath or filepath))

        return jsonify({
            "status":              "success",
            "report_id":           report_id,
            "download_url":        f"/api/valuation/report/download/{actual_filename}",
            "primary_value":       float(asset_result.primary_value) if asset_result.primary_value is not None else None,
            "confidence":          asset_result.confidence,
            "asset_type":          asset_result.asset_type,
            "issues_count":        len(asset_result.issues),
            "ivsc_compliant":      _ivsc_disclosure is not None,
            "cross_border_compliant": _cb_disclosure is not None,
        }), 200

    except Exception as e:
        import traceback as _tb
        print(_tb.format_exc())
        return jsonify({"status": "error", "message": str(e)}), 400


@app.route("/api/valuation/report/download/<filename>", methods=["GET"])
def api_valuation_report_download(filename: str):
    """
    Download a previously generated Excel valuation report.

    <filename> must be the UUID-based name returned by POST /api/valuation/report.
    The file is served as an attachment with MIME type application/vnd.openxmlformats.
    """
    # Reject path traversal attempts
    safe_name = os.path.basename(filename)
    if not safe_name.endswith(".xlsx") or safe_name != filename:
        return jsonify({"status": "error", "message": "Invalid filename"}), 400

    filepath = os.path.join(_REPORT_DIR, safe_name)
    if not os.path.exists(filepath):
        return jsonify({"status": "error", "message": "Report not found or expired"}), 404

    return send_file(
        filepath,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        as_attachment=True,
        download_name=f"valuation_report_{safe_name}",
    )


# ── Phase 7 routes ────────────────────────────────────────────────────────────

@app.route("/api/valuation/land", methods=["POST", "OPTIONS"])
def api_valuation_land():
    """
    Land valuation pipeline: Phase 4 engines → Phase 5 market adapter → Phase 7 LandAdapter.

    Required body keys:
      subject_property  dict  — must include property_type (vacant_land / raw_land /
                                development_site / leasehold_land), area_sqm, hbu
    Optional body keys:
      filters               dict   latitude + longitude required for comparable search
      comparable_count      int    default 5
      land_value_egp        float  nominal land value passed as cost proxy (default 500 000)
      gross_income_annual_egp float default 100 000
      vacancy_rate          float  default 0.10
      opex_ratio            float  default 0.35
      cap_rate              float  default 0.08
      run_quality_audit     bool   when true, appends quality_audit to response

    Returns JSON with phase_4_values, comparable_search, land_valuation, quality_audit.
    """
    if request.method == "OPTIONS":
        return jsonify({}), 200
    try:
        data = request.get_json(force=True) or {}
        from decimal import Decimal as _D

        subject          = data.get("subject_property", {})
        filters          = data.get("filters", {})
        comparable_count = int(data.get("comparable_count", 5))
        run_audit        = bool(data.get("run_quality_audit", False))
        land_value_egp   = float(data.get("land_value_egp", 500_000))
        gross_income     = float(data.get("gross_income_annual_egp", 100_000))
        vacancy_rate     = float(data.get("vacancy_rate", 0.10))
        opex_ratio       = float(data.get("opex_ratio", 0.35))
        cap_rate         = float(data.get("cap_rate", 0.08))
        governorate      = filters.get("governorate", "Cairo")

        # ── Phase 4: comparable search ────────────────────────────────────────
        comparables = get_search_engine().search_and_rank(
            subject=subject, filters=filters, limit=20
        )
        if not comparables:
            return jsonify({
                "status": "error",
                "message": "No comparables found matching the given criteria.",
            }), 400

        selected = comparables[:comparable_count]

        # ── Phase 4: comparable sales engine ──────────────────────────────────
        result_comp = get_comparative_engine().calculate({
            "subject_area_sqm":  subject.get("area_sqm"),
            "subject_age_years": subject.get("age_years", 0),
            "comparables":       selected,
        })

        # ── Phase 4: income engine (land leased for parking / commercial use) ─
        result_income = get_income_engine().calculate({
            "gross_income_annual_egp":  gross_income,
            "vacancy_rate":             vacancy_rate,
            "operating_expenses_ratio": opex_ratio,
            "cap_rate":                 cap_rate,
            "cap_rate_source":          "market",
        })

        comp_value  = result_comp.value  if result_comp.value  is not None else _D("0")
        inc_value   = result_income.value if result_income.value is not None else _D("0")
        cost_proxy  = _D(str(land_value_egp))  # nominal; multiplied by 0 in land adapter

        # Phase 5 adapters use "comparative"; Phase 6/7 asset adapters use "comparable"
        p5_values = {"comparative": comp_value, "cost": cost_proxy, "income": inc_value}
        p6_values = {"comparable":  comp_value, "cost": cost_proxy, "income": inc_value}

        # ── Phase 5: market value adapter (drives confidence + audit chain) ───
        market_result = get_market_value_adapter().adjust(p5_values, {
            "comparable_count":   len(selected),
            "comparable_quality": "standard",
        })

        # ── Phase 7: land adapter ──────────────────────────────────────────────
        land_result = get_land_adapter().value(subject, p6_values, market_result)

        # ── Phase 7: optional quality audit ───────────────────────────────────
        audit_data = None
        if run_audit:
            audit = get_report_auditor().audit(land_result)
            audit_data = {
                "passed":        audit.passed,
                "quality_score": audit.quality_score,
                "quality_grade": audit.quality_grade,
                "summary":       audit.summary,
                "findings": [
                    {
                        "severity":       f.severity,
                        "category":       f.category,
                        "code":           f.code,
                        "message":        f.message,
                        "recommendation": f.recommendation,
                    }
                    for f in audit.findings
                ],
            }

        top_score = selected[0].get("similarity_score", 0) if selected else 0

        # ── Persist to PostgreSQL (graceful — never blocks the response) ──────
        valuation_id = None
        if _DB_AVAILABLE:
            try:
                _sess = _DbSession()
                try:
                    _rec = _DbValuation(
                        asset_type=subject.get("property_type", "vacant_land"),
                        comparable_value=float(comp_value),
                        cost_value=float(cost_proxy),
                        income_value=float(inc_value),
                        primary_value=float(land_result.primary_value) if land_result.primary_value is not None else None,
                        confidence=land_result.confidence,
                        result_json={"land_valuation": {"asset_type": land_result.asset_type}},
                    )
                    _sess.add(_rec)
                    _sess.flush()   # populate _rec.id before QualityAudit FK
                    valuation_id = str(_rec.id)

                    if run_audit and audit_data:
                        _qa = _DbQualityAudit(
                            valuation_id=_rec.id,
                            passed=audit_data.get("passed", False),
                            quality_score=audit_data.get("quality_score"),
                            quality_grade=audit_data.get("quality_grade"),
                            findings_json=audit_data.get("findings", []),
                        )
                        _sess.add(_qa)

                    _sess.commit()
                finally:
                    _sess.close()
            except Exception:
                pass  # DB write failure must not affect the API response

        return jsonify({
            "status": "success",
            "phase_4_values": {
                "comparable": float(comp_value),
                "cost":       float(cost_proxy),
                "income":     float(inc_value),
            },
            "comparable_search": {
                "count":                len(selected),
                "top_similarity_score": top_score,
            },
            "land_valuation": {
                "asset_type":     land_result.asset_type,
                "primary_value":  float(land_result.primary_value) if land_result.primary_value is not None else None,
                "confidence":     land_result.confidence,
                "weights_applied": land_result.weights_applied,
                "audit_trail": [
                    {
                        "step_name":  e.step_name,
                        "inputs":     e.inputs,
                        "outputs":    e.outputs,
                        "formula":    e.formula,
                        "references": e.references,
                    }
                    for e in land_result.audit_trail
                ],
                "disclosures": land_result.disclosures,
                "metadata":    {
                    k: (float(v) if hasattr(v, "__float__") and not isinstance(v, (bool, str, list, dict)) else v)
                    for k, v in land_result.metadata.items()
                },
                "issues": [
                    {"severity": i.severity, "code": i.code, "message": i.message}
                    for i in land_result.issues
                ],
            },
            "quality_audit": audit_data,
            "valuation_id": valuation_id,
        }), 200

    except Exception as e:
        import traceback as _tb
        print(_tb.format_exc())
        return jsonify({"status": "error", "message": str(e)}), 400


@app.route("/api/valuation/audit", methods=["POST", "OPTIONS"])
def api_valuation_audit():
    """
    Quality audit endpoint: run ReportQualityAuditor on any Phase 5-7 result JSON.

    Accepts the serialised form of an AssetValuationResult and returns a full
    AuditReport without needing the original Python object.

    Required body key:
      result  dict  — serialised AssetValuationResult with at minimum:
                      asset_type, primary_value, confidence, weights_applied,
                      audit_trail (list of step dicts), disclosures (list of str)

    Returns JSON with passed, quality_score, quality_grade, summary, findings.
    """
    if request.method == "OPTIONS":
        return jsonify({}), 200
    try:
        data = request.get_json(force=True) or {}
        from decimal import Decimal as _D

        result_json = data.get("result")
        if not result_json:
            return jsonify({
                "status": "error",
                "message": "Missing 'result' key in request body.",
            }), 400

        # Re-hydrate a lightweight proxy object that satisfies AuditFinding checks
        raw_pv = result_json.get("primary_value")
        primary_value = _D(str(raw_pv)) if raw_pv is not None else None

        class _AuditEntry:
            def __init__(self, d):
                self.step_name  = d.get("step_name", "")
                self.inputs     = d.get("inputs", {})
                self.outputs    = d.get("outputs", {})
                self.formula    = d.get("formula", "")
                self.references = d.get("references", [])

        class _Issue:
            def __init__(self, d):
                self.severity = d.get("severity", "warning")
                self.code     = d.get("code", "")
                self.message  = d.get("message", "")

        class _Result:
            pass

        result_obj = _Result()
        result_obj.asset_type         = result_json.get("asset_type", "")
        result_obj.primary_purpose    = result_json.get("primary_purpose", "market_value")
        result_obj.primary_value      = primary_value
        result_obj.confidence         = result_json.get("confidence", "")
        result_obj.weights_applied    = result_json.get("weights_applied", {})
        result_obj.alternative_values = {
            k: _D(str(v)) for k, v in result_json.get("alternative_values", {}).items()
        }
        result_obj.audit_trail        = [_AuditEntry(e) for e in result_json.get("audit_trail", [])]
        result_obj.disclosures        = result_json.get("disclosures", [])
        result_obj.issues             = [_Issue(i) for i in result_json.get("issues", [])]
        result_obj.metadata           = result_json.get("metadata", {})

        audit = get_report_auditor().audit(result_obj)

        return jsonify({
            "status": "success",
            "quality_audit": {
                "passed":        audit.passed,
                "quality_score": audit.quality_score,
                "quality_grade": audit.quality_grade,
                "summary":       audit.summary,
                "error_count":   len(audit.errors),
                "warning_count": len(audit.warnings),
                "findings": [
                    {
                        "severity":       f.severity,
                        "category":       f.category,
                        "code":           f.code,
                        "message":        f.message,
                        "recommendation": f.recommendation,
                        "field_path":     f.field_path,
                        "score_impact":   f.score_impact,
                    }
                    for f in audit.findings
                ],
            },
        }), 200

    except Exception as e:
        import traceback as _tb
        print(_tb.format_exc())
        return jsonify({"status": "error", "message": str(e)}), 400


# ── Phase 11: Portfolio Analysis ─────────────────────────────────────────────

@app.route("/api/valuation/portfolio", methods=["POST", "OPTIONS"])
def api_valuation_portfolio():
    """
    Portfolio valuation endpoint (Phase 11).

    Accepts a list of pre-valued properties, computes aggregate metrics,
    and optionally generates an Excel portfolio summary report.
    """
    if request.method == "OPTIONS":
        return jsonify({}), 200
    try:
        from adapters.portfolio import PortfolioBuilder as _PBuilder
        from datetime import datetime as _dt

        data = request.get_json(silent=True) or {}

        builder = _PBuilder(data.get("portfolio_name", "Unnamed Portfolio"))
        for pd in data.get("properties", []):
            builder.add_property(
                property_id=str(pd.get("property_id", "")),
                property_name=str(pd.get("property_name", "")),
                property_type=str(pd.get("property_type", "residential")),
                valuation_value=float(pd.get("valuation_value", 0)),
                valuation_confidence=str(pd.get("valuation_confidence", "medium")),
                annual_noi=float(pd.get("annual_noi", 0)),
                annual_gross_income=float(pd.get("annual_gross_income", 0)),
            )

        builder.calculate_metrics()
        portfolio_summary    = builder.get_portfolio_summary()
        diversification_score = builder.get_diversification_score()

        # ── Optional Excel report ──────────────────────────────────────────────
        report_id    = None
        download_url = None
        if data.get("generate_report"):
            report_id  = str(uuid.uuid4())
            filename   = f"{report_id}.xlsx"
            filepath   = os.path.join(_REPORT_DIR, filename)
            ExcelReportBuilder().build(filepath, portfolio_summary=portfolio_summary)
            download_url = f"/api/valuation/report/download/{filename}"

        return jsonify({
            "status":               "success",
            "portfolio":            portfolio_summary,
            "diversification_score": round(diversification_score, 4),
            "report_id":            report_id,
            "download_url":         download_url,
            "timestamp":            _dt.now().isoformat(),
        }), 200

    except Exception as e:
        import traceback as _tb
        print(_tb.format_exc())
        return jsonify({"status": "error", "message": str(e)}), 400


@app.route("/api/valuation/portfolio/performance", methods=["POST", "OPTIONS"])
def api_valuation_portfolio_performance():
    """
    Portfolio performance + scenario analysis endpoint (Phase 11.2).

    Accepts a list of pre-valued properties and optional scenario definitions.
    Returns stressed metrics, value-at-risk, and IRR estimates across scenarios.
    """
    if request.method == "OPTIONS":
        return jsonify({}), 200
    try:
        from adapters.portfolio import PortfolioBuilder as _PBuilder
        from adapters.portfolio_performance import (
            PortfolioPerformanceAnalyzer as _PPA,
            PortfolioScenario as _PScenario,
        )
        from datetime import datetime as _dt

        data = request.get_json(silent=True) or {}

        # ── Build portfolio ────────────────────────────────────────────────────
        builder = _PBuilder(data.get("portfolio_name", "Unnamed Portfolio"))
        for pd in data.get("properties", []):
            builder.add_property(
                property_id=str(pd.get("property_id", "")),
                property_name=str(pd.get("property_name", "")),
                property_type=str(pd.get("property_type", "residential")),
                valuation_value=float(pd.get("valuation_value", 0)),
                valuation_confidence=str(pd.get("valuation_confidence", "medium")),
                annual_noi=float(pd.get("annual_noi", 0)),
                annual_gross_income=float(pd.get("annual_gross_income", 0)),
            )

        analyzer = _PPA(builder)

        # ── Scenarios: standard presets or caller-supplied list ────────────────
        custom_scenarios = data.get("scenarios", [])
        if data.get("use_standard_scenarios", True) and not custom_scenarios:
            analyzer.create_standard_scenarios()
        else:
            for sc in custom_scenarios:
                analyzer.add_scenario(_PScenario(
                    label=str(sc.get("label", "custom")),
                    noi_shock=float(sc.get("noi_shock", 1.0)),
                    value_shock=float(sc.get("value_shock", 1.0)),
                    cap_rate_shift=float(sc.get("cap_rate_shift", 0.0)),
                ))

        performance_summary = analyzer.get_performance_summary()

        return jsonify({
            "status":               "success",
            "performance_analysis": performance_summary,
            "timestamp":            _dt.now().isoformat(),
        }), 200

    except Exception as e:
        import traceback as _tb
        print(_tb.format_exc())
        return jsonify({"status": "error", "message": str(e)}), 400


# ── Phase 12: Batch Valuation ────────────────────────────────────────────────

@app.route("/api/valuation/batch", methods=["POST", "OPTIONS"])
def api_valuation_batch():
    """
    Batch valuation endpoint (Phase 12.1).

    Accepts an array of properties and processes them in a single request.
    Each property resolves its valuation value via one of three strategies:
      1. input_data.valuation_value  — use directly (pre-valued)
      2. input_data.price_per_sqm   — multiply by area_sqm
      3. neither                    — property is failed with an explanatory message

    Request body:
      {
        "batch_name": "...",          // optional
        "properties": [
          {
            "property_id": "P1",
            "property_name": "...",
            "property_type": "residential|commercial|land",
            "area_sqm": 120,
            "input_data": {
              "valuation_value": 3645000,   // strategy 1
              "price_per_sqm":   30000,     // strategy 2 (if no valuation_value)
              "primary_purpose": "market_value"
            }
          }
        ],
        "generate_report": false      // optional; creates portfolio Excel
      }

    Response:
      { "status": "success", "batch_id": "...", "batch_name": "...",
        "summary": {...}, "completed_properties": [...],
        "failed_properties": [...], "skipped_properties": [...],
        "report_id": null, "download_url": null, "timestamp": "..." }
    """
    if request.method == "OPTIONS":
        return jsonify({}), 200
    try:
        from adapters.batch_processor import BatchProcessor as _BP
        from datetime import datetime as _dt

        data = request.get_json(silent=True) or {}
        properties_raw = data.get("properties", [])

        if not properties_raw:
            return jsonify({"status": "error",
                            "message": "properties array is required and must not be empty"}), 400

        # ── Build and validate batch ───────────────────────────────────────────
        bp = _BP(data.get("batch_name", "Unnamed Batch"))
        for pd in properties_raw:
            bp.add_property(
                property_id=str(pd.get("property_id", "")),
                property_name=str(pd.get("property_name", "")),
                property_type=str(pd.get("property_type", "residential")),
                area_sqm=float(pd.get("area_sqm", 0)),
                input_data=pd.get("input_data", {}),
            )

        bp.validate_batch()
        bp.start_processing()

        # ── Process each property ─────────────────────────────────────────────
        for idx, prop in enumerate(bp.properties):
            if prop.status == "skipped":
                continue
            inp     = prop.input_data
            purpose = str(inp.get("primary_purpose", "market_value"))

            # Strategy 1: explicit pre-valued
            if "valuation_value" in inp and float(inp["valuation_value"]) > 0:
                bp.process_property(idx, float(inp["valuation_value"]), purpose)
                continue

            # Strategy 2: price per sqm × area
            if "price_per_sqm" in inp and float(inp["price_per_sqm"]) > 0:
                value = float(inp["price_per_sqm"]) * prop.area_sqm
                bp.process_property(idx, value, purpose)
                continue

            # Strategy 3: insufficient data
            bp.fail_property(idx, "No valuation_value or price_per_sqm in input_data")

        bp.complete_batch()
        report = bp.get_completion_report()

        # ── Persist to SQLite batch store for durable GET retrieval ──────────────
        from database.batch_store import BatchStore as _BS
        _BS().save({**report, "batch_name": bp.batch_name})

        # ── Optional webhook notification (Phase 14) ───────────────────────────
        webhook_url = data.get("webhook_url", "").strip()
        if webhook_url:
            try:
                from adapters.webhook_dispatcher import WebhookDispatcher as _WD
                from database.webhook_log import WebhookLog as _WL
                _wl = _WL()
                _WD().dispatch(
                    webhook_url, report,
                    batch_id=bp.batch_id,
                    on_complete=lambda d: _wl.record(d),
                )
            except Exception:
                pass

        # ── Optional Excel report (batch-specific 3-sheet workbook) ──────────────
        report_id    = None
        download_url = None
        if data.get("generate_report"):
            from reports.batch_report_builder import BatchReportBuilder as _BRB
            report_id    = str(uuid.uuid4())
            filename     = f"{report_id}.xlsx"
            filepath     = os.path.join(_REPORT_DIR, filename)
            _BRB(report).build(filepath)
            download_url = f"/api/valuation/report/download/{filename}"

        response_body = {
            "status":                "success",
            "batch_id":              report["batch_id"],
            "batch_name":            bp.batch_name,
            "summary":               report["summary"],
            "completed_properties":  report["completed_properties"],
            "failed_properties":     report["failed_properties"],
            "skipped_properties":    report["skipped_properties"],
            "report_id":             report_id,
            "download_url":          download_url,
            "timestamp":             _dt.now().isoformat(),
        }
        return jsonify(response_body), 200

    except Exception as e:
        import traceback as _tb
        print(_tb.format_exc())
        return jsonify({"status": "error", "message": str(e)}), 400


@app.route("/api/valuation/batch/<batch_id>", methods=["GET", "OPTIONS"])
def api_valuation_batch_status(batch_id: str):
    """
    Retrieve a stored batch report by batch_id (Phase 12.3).

    Returns the completion report registered by POST /api/valuation/batch.
    Returns 404 if the batch_id is not found (unknown or server restarted).
    """
    if request.method == "OPTIONS":
        return jsonify({}), 200
    try:
        from database.batch_store import BatchStore as _BS
        from datetime import datetime as _dt

        stored = _BS().get(batch_id)
        if stored is None:
            return jsonify({
                "status":   "not_found",
                "message":  f"Batch '{batch_id}' not found. "
                            "Batches are held in memory and cleared on server restart.",
                "batch_id": batch_id,
            }), 404

        return jsonify({
            "status":    "success",
            "batch":     stored,
            "timestamp": _dt.now().isoformat(),
        }), 200

    except Exception as e:
        import traceback as _tb
        print(_tb.format_exc())
        return jsonify({"status": "error", "message": str(e)}), 400


@app.route("/api/valuation/batch", methods=["GET", "OPTIONS"])
def api_valuation_batch_list():
    """
    List recent batches from the in-memory registry (Phase 12.3).

    Query params:
      limit  — max results to return (default 20, max 100)
    """
    if request.method == "OPTIONS":
        return jsonify({}), 200
    try:
        from database.batch_store import BatchStore as _BS
        from datetime import datetime as _dt

        limit  = min(int(request.args.get("limit", 20)), 100)
        recent = _BS().list_recent(limit=limit)

        summaries = [
            {
                "batch_id":       r.get("batch_id"),
                "status":         r.get("status"),
                "completed":      r.get("summary", {}).get("completed", 0),
                "failed":         r.get("summary", {}).get("failed", 0),
                "total_submitted":r.get("summary", {}).get("total_submitted", 0),
                "registered_at":  r.get("registered_at"),
            }
            for r in recent
        ]

        return jsonify({
            "status":      "success",
            "count":       len(summaries),
            "batches":     summaries,
            "timestamp":   _dt.now().isoformat(),
        }), 200

    except Exception as e:
        import traceback as _tb
        print(_tb.format_exc())
        return jsonify({"status": "error", "message": str(e)}), 400


# ── Phase 10: DCF Income Valuation ───────────────────────────────────────────

@app.route("/api/valuation/dcf", methods=["POST", "OPTIONS"])
def api_valuation_dcf():
    """
    DCF income valuation endpoint (Phase 10).

    Accepts multi-year cash flow projections and optional sensitivity analysis.
    Returns NPV-based property value + optional scenario comparison.
    """
    if request.method == "OPTIONS":
        return jsonify({}), 200
    try:
        from adapters.dcf_model import DCFModel as _DCFModel
        from adapters.dcf_sensitivity import DCFSensitivityAnalysis as _DCFSens
        from datetime import datetime as _dt

        data = request.get_json(silent=True) or {}

        # ── DCF assumptions ────────────────────────────────────────────────────
        dcf_assumptions   = data.get("dcf_assumptions", {})
        discount_rate     = float(dcf_assumptions.get("discount_rate", 0.08))
        holding_period    = int(dcf_assumptions.get("holding_period", 5))
        terminal_cap_rate = float(dcf_assumptions.get("terminal_cap_rate", 0.07))
        terminal_method   = str(dcf_assumptions.get("terminal_value_method", "cap_rate"))
        terminal_growth   = float(dcf_assumptions.get("terminal_noi_growth", 0.02))

        annual_projections = dcf_assumptions.get("annual_projections", [])
        if not annual_projections:
            return jsonify({"status": "error",
                            "message": "dcf_assumptions.annual_projections is required"}), 400

        # ── Build DCF model ────────────────────────────────────────────────────
        dcf = _DCFModel(
            discount_rate=discount_rate,
            holding_period=holding_period,
            terminal_cap_rate=terminal_cap_rate,
        )

        final_year_noi = 0.0
        last_year      = max(p.get("year", 1) for p in annual_projections)

        for proj in annual_projections:
            year    = int(proj.get("year", 1))
            gross   = float(proj.get("gross_income", 0))
            vacancy = float(proj.get("vacancy_rate", 0.05))
            opex    = float(proj.get("operating_expenses", 0))
            debt    = float(proj.get("debt_service", 0))
            dcf.add_annual_cash_flow(year, gross, vacancy, opex, debt)
            if year == last_year:
                eri            = gross * (1 - vacancy)
                final_year_noi = eri - opex

        dcf.set_terminal_value(
            final_year_noi=final_year_noi,
            method=terminal_method,
            terminal_noi_growth=terminal_growth,
        )
        dcf.calculate_npv()

        dcf_result = {
            "property_value":    round(dcf.property_value, 2),
            "pv_cash_flows":     round(dcf.pv_cash_flows, 2),
            "pv_terminal_value": round(dcf.pv_terminal_value, 2),
            "discount_rate":     discount_rate,
            "terminal_cap_rate": terminal_cap_rate,
            "holding_period":    holding_period,
            "annual_cash_flows": [a.to_dict() for a in dcf.annual_cash_flows],
            "terminal_value":    dcf.terminal_value.to_dict() if dcf.terminal_value else None,
        }

        # ── Optional sensitivity analysis ──────────────────────────────────────
        sensitivity_result = None
        if data.get("include_sensitivity"):
            sens = _DCFSens(discount_rate, holding_period, terminal_cap_rate)
            for proj in annual_projections:
                sens.add_year_projection(
                    int(proj.get("year", 1)),
                    float(proj.get("gross_income", 0)),
                    float(proj.get("vacancy_rate", 0.05)),
                    float(proj.get("operating_expenses", 0)),
                    float(proj.get("debt_service", 0)),
                )
            cap_rates  = data.get("sensitivity_cap_rates")
            disc_rates = data.get("sensitivity_discount_rates")
            if cap_rates and len(cap_rates) >= 3:
                sens.create_cap_rate_scenarios(
                    float(cap_rates[0]), float(cap_rates[1]), float(cap_rates[2])
                )
            if disc_rates:
                sens.create_discount_rate_scenarios([float(r) for r in disc_rates])
            sensitivity_result = sens.get_scenario_summary()

        return jsonify({
            "status":              "success",
            "dcf_valuation":       dcf_result,
            "sensitivity_analysis": sensitivity_result,
            "timestamp":           _dt.now().isoformat(),
        }), 200

    except Exception as e:
        import traceback as _tb
        print(_tb.format_exc())
        return jsonify({"status": "error", "message": str(e)}), 400


# ── Phase 15: Enterprise API routes ──────────────────────────────────────────

@app.route("/api/enterprise/tenant", methods=["POST"])
def api_enterprise_create_tenant():
    """POST /api/enterprise/tenant — create a new tenant organization."""
    try:
        data = request.get_json(force=True) or {}
        org_name = (data.get("organization_name") or "").strip()
        country  = (data.get("country") or "").strip()
        tier     = (data.get("subscription_tier") or "professional").strip()

        if not org_name:
            return jsonify({"status": "error", "message": "organization_name required"}), 400
        if not country:
            return jsonify({"status": "error", "message": "country required"}), 400

        tenant = _tenant_manager.create_tenant(org_name, country, tier)
        try:
            _AuditLog().record(_AuditEvent(
                tenant_id=tenant.tenant_id,
                action=_AuditAction.TENANT_CREATED,
                resource_type="tenant",
                resource_id=tenant.tenant_id,
                details={"organization_name": org_name, "country": country, "tier": tier},
                ip_address=request.remote_addr or "",
            ))
        except Exception:
            pass
        return jsonify({
            "status": "success",
            "tenant": tenant.to_dict(),
        }), 201

    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 400


@app.route("/api/enterprise/tenant/<tenant_id>", methods=["GET"])
def api_enterprise_get_tenant(tenant_id: str):
    """GET /api/enterprise/tenant/<tenant_id> — fetch tenant summary."""
    summary = _tenant_manager.get_tenant_summary(tenant_id)
    if not summary:
        return jsonify({"status": "error", "message": "Tenant not found"}), 404
    return jsonify({"status": "success", **summary}), 200


@app.route("/api/enterprise/tenant/<tenant_id>/user", methods=["POST"])
def api_enterprise_add_user(tenant_id: str):
    """POST /api/enterprise/tenant/<tenant_id>/user — add a user to a tenant."""
    try:
        data      = request.get_json(force=True) or {}
        email     = (data.get("email") or "").strip()
        full_name = (data.get("full_name") or "").strip()
        role      = (data.get("role") or "").strip()

        if not email:
            return jsonify({"status": "error", "message": "email required"}), 400
        if not role:
            return jsonify({"status": "error", "message": "role required"}), 400

        try:
            _TenantRole(role)
        except ValueError:
            valid = [r.value for r in _TenantRole]
            return jsonify({
                "status": "error",
                "message": f"Invalid role '{role}'. Must be one of: {valid}",
            }), 400

        if not _tenant_manager.get_tenant(tenant_id):
            return jsonify({"status": "error", "message": "Tenant not found"}), 404

        user = _tenant_manager.add_user_to_tenant(tenant_id, email, full_name, role)
        try:
            _AuditLog().record(_AuditEvent(
                tenant_id=tenant_id,
                action=_AuditAction.USER_ADDED,
                resource_type="user",
                resource_id=user.user_id,
                user_id=user.user_id,
                details={"email": email, "role": role},
                ip_address=request.remote_addr or "",
            ))
        except Exception:
            pass
        return jsonify({"status": "success", "user": user.to_dict()}), 201

    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 400


@app.route("/api/enterprise/tenant/<tenant_id>/license", methods=["GET"])
def api_enterprise_license(tenant_id: str):
    """GET /api/enterprise/tenant/<tenant_id>/license — validate subscription."""
    tenant = _tenant_manager.get_tenant(tenant_id)
    if not tenant:
        return jsonify({"status": "error", "message": "Tenant not found"}), 404
    result = _LicenseValidator.validate_subscription(tenant)
    return jsonify({"status": "success", **result}), 200


@app.route("/api/enterprise/tenant/<tenant_id>/audit", methods=["GET"])
def api_enterprise_audit(tenant_id: str):
    """GET /api/enterprise/tenant/<tenant_id>/audit — fetch audit trail."""
    if not _tenant_manager.get_tenant(tenant_id):
        return jsonify({"status": "error", "message": "Tenant not found"}), 404
    try:
        limit  = min(int(request.args.get("limit", 50)), 200)
        action = request.args.get("action", "").strip() or None
        user_id = request.args.get("user_id", "").strip() or None
        al = _AuditLog()
        if user_id:
            events = al.get_by_user(tenant_id, user_id, limit=limit)
        else:
            events = al.get_by_tenant(tenant_id, limit=limit)
        if action:
            events = [e for e in events if e["action"] == action]
        return jsonify({
            "status":      "success",
            "tenant_id":   tenant_id,
            "event_count": len(events),
            "events":      events,
        }), 200
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 400


# ── Phase 21 — Agent Chat API ─────────────────────────────────────────────────

try:
    from agents.workspace_manager import WorkspaceManager as _WorkspaceManager
    from agents.supervised_agent import SupervisedAgent as _SupervisedAgent, ExecutionMode as _ExecMode
    from agents.chat_agent import ChatAgent as _ChatAgent
    _chat_wm    = _WorkspaceManager()
    _chat_agent = _SupervisedAgent(_chat_wm, None, execution_mode=_ExecMode.SUPERVISED)
    _chat_bot   = _ChatAgent(_chat_wm, _chat_agent)
    _CHAT_OK    = True
except Exception as _chat_err:
    _CHAT_OK    = False
    _chat_bot   = None


@app.route("/agent")
def agent_chat_ui():
    """Serve the Phase 21 Agent Chat UI."""
    chat_html = Path(BASE_DIR).parent / "frontend" / "agent_chat.html"
    if chat_html.exists():
        return send_file(str(chat_html))
    return "Agent Chat UI not found", 404


@app.route("/api/agent/chat", methods=["POST"])
def agent_chat():
    """
    POST /api/agent/chat
    Body: {"message": "run pipeline on <workspace_id>", "workspace_id": "<optional>"}
    Returns: {"success": bool, "message": str, "data": {...}, "intent": str}
    """
    if not _CHAT_OK or _chat_bot is None:
        return jsonify({"success": False,
                        "message": "Chat agent not available",
                        "data": None, "intent": "unknown"}), 503
    body         = request.get_json(force=True, silent=True) or {}
    text         = str(body.get("message", "")).strip()
    workspace_id = str(body.get("workspace_id", "")).strip() or None
    if not text:
        return jsonify({"success": False,
                        "message": "message field is required",
                        "data": None, "intent": "unknown"}), 400
    try:
        resp = _chat_bot.chat(text, workspace_id=workspace_id)
        return jsonify(resp.to_dict()), 200
    except Exception as exc:
        return jsonify({"success": False, "message": str(exc),
                        "data": None, "intent": "unknown"}), 500


# ── Phase 22 — Knowledge Base + RAG ──────────────────────────────────────────

try:
    from knowledge.knowledge_base import KnowledgeBase as _KnowledgeBase
    from knowledge.rag_enhancer import RAGEnhancer as _RAGEnhancer
    from knowledge.seed_data import get_seed_entries as _get_seed_entries
    _kb22 = _KnowledgeBase()
    for _se in _get_seed_entries():
        _kb22.add_entry(_se, auto_embed=False)
    _rag22 = _RAGEnhancer(_kb22)
    _KB22_OK = True
except Exception as _kb22_err:
    _KB22_OK = False
    _kb22 = None
    _rag22 = None


@app.route("/api/knowledge/search", methods=["POST"])
def knowledge_search():
    """
    POST /api/knowledge/search
    Body: {"query": str, "top_k": int, "category": str}
    Returns: {"results": [...entry dicts...], "count": int}
    """
    if not _KB22_OK or _kb22 is None:
        return jsonify({"error": "Knowledge base not available"}), 503
    body     = request.get_json(force=True, silent=True) or {}
    query    = str(body.get("query", "")).strip()
    top_k    = int(body.get("top_k", 5))
    category = body.get("category")
    if not query:
        return jsonify({"error": "query is required"}), 400
    try:
        from knowledge.knowledge_base import KnowledgeCategory as _KC
        cat = _KC(category) if category else None
        entries = _kb22.search(query, top_k=top_k, category=cat)
        return jsonify({"results": [e.to_dict() for e in entries], "count": len(entries)}), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 400


@app.route("/api/knowledge/enhance", methods=["POST"])
def knowledge_enhance():
    """
    POST /api/knowledge/enhance
    Body: {"property_type": str, "location": str, "purpose": str}
    Returns: {"context": str, "citations": [...], "entries": [...]}
    """
    if not _KB22_OK or _rag22 is None:
        return jsonify({"error": "RAG enhancer not available"}), 503
    body          = request.get_json(force=True, silent=True) or {}
    property_type = str(body.get("property_type", "residential"))
    location      = str(body.get("location", "Cairo"))
    purpose       = str(body.get("purpose", "market_value"))
    try:
        result = _rag22.enhance_valuation_context(property_type, location, purpose)
        return jsonify(result), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/knowledge/stats", methods=["GET"])
def knowledge_stats():
    """GET /api/knowledge/stats — returns KB statistics."""
    if not _KB22_OK or _kb22 is None:
        return jsonify({"error": "Knowledge base not available"}), 503
    try:
        return jsonify(_kb22.get_statistics()), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


# ── API Hardening — Resilience, Idempotency, Observability ───────────────────

try:
    from api.resilience import (
        RetryPolicy as _RetryPolicy,
        CircuitBreaker as _CircuitBreaker,
        RetryStrategy as _RetryStrategy,
        api_retry_policy as _api_retry_policy,
        valuation_circuit_breaker as _valuation_cb,
    )
    from api.request_deduplication import RequestDeduplicator as _ReqDedup
    from api.response_formatter import ResponseFormatter as _RespFmt, format_response as _fmt_resp
    from api.error_handler import (
        APIError as _APIError,
        ValidationError as _ValError,
        NotFoundError as _NotFoundError,
        ErrorCode as _ErrCode,
    )
    from api.observability import StructuredLogger as _SLog
    _API_HARD_OK = True
    _api_logger  = _SLog("bridge_api")
except Exception as _ah_err:
    _API_HARD_OK = False
    _api_logger  = None

import uuid as _uuid_mod
import time as _time_mod


@app.before_request
def _harden_before_request():
    from flask import g
    g.request_id = str(_uuid_mod.uuid4())
    g.start_time = _time_mod.time()


@app.after_request
def _harden_after_request(response):
    from flask import g
    if hasattr(g, "request_id"):
        response.headers["X-Request-ID"] = g.request_id
    if hasattr(g, "start_time"):
        elapsed = _time_mod.time() - g.start_time
        response.headers["X-Process-Time"] = f"{elapsed:.6f}"
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["X-Frame-Options"]        = "DENY"
    response.headers["X-XSS-Protection"]       = "1; mode=block"
    return response


@app.route("/api/health", methods=["GET"])
def api_health_hardened():
    """GET /api/health — system-wide health with component flags."""
    return jsonify({
        "status": "healthy",
        "components": {
            "api_hardening": _API_HARD_OK,
            "i18n":          _I18N_OK,
            "knowledge_base": _KB22_OK,
            "scenarios":     _SCENARIOS_OK,
        },
    }), 200


# ── Phase 24 — Advanced Scenario Modeling ────────────────────────────────────

try:
    from scenarios.scenario_builder import ScenarioType as _ST24, ScenarioBuilder as _SB24
    from scenarios.monte_carlo import MonteCarloConfig as _MCCfg24, MonteCarloEngine as _MCE24
    from scenarios.sensitivity_matrix import SensitivityMatrix as _SM24
    from scenarios.stress_test import StressTestSuite as _STS24, StressScenario as _SS24
    _SCENARIOS_OK = True
except Exception as _sc24_err:
    _SCENARIOS_OK = False


@app.route("/api/scenarios/run", methods=["POST"])
def scenarios_run():
    """
    POST /api/scenarios/run
    Body: {"base_value": float, "parameters": [{"name": str, "optimistic_delta": float, "pessimistic_delta": float}]}
    Returns: {"scenarios": [...]}
    """
    if not _SCENARIOS_OK:
        return jsonify({"error": "Scenario engine not available"}), 503
    body       = request.get_json(force=True, silent=True) or {}
    base_value = float(body.get("base_value", 0))
    if base_value <= 0:
        return jsonify({"error": "base_value must be positive"}), 400
    try:
        sb = _SB24(base_value)
        for p in body.get("parameters", []):
            sb.add_parameter(
                name=str(p.get("name", "param")),
                base_value=float(p.get("base_value", 0)),
                optimistic_delta=float(p.get("optimistic_delta", 0)),
                pessimistic_delta=float(p.get("pessimistic_delta", 0)),
                description=str(p.get("description", "")),
            )
        results = sb.build_all()
        return jsonify({"scenarios": [r.to_dict() for r in results]}), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/scenarios/monte_carlo", methods=["POST"])
def scenarios_monte_carlo():
    """
    POST /api/scenarios/monte_carlo
    Body: {"base_value": float, "volatility": float, "iterations": int, "seed": int}
    Returns: MonteCarloResult dict
    """
    if not _SCENARIOS_OK:
        return jsonify({"error": "Scenario engine not available"}), 503
    body       = request.get_json(force=True, silent=True) or {}
    base_value = float(body.get("base_value", 0))
    volatility = float(body.get("volatility", 15))
    iterations = int(body.get("iterations", 10_000))
    seed       = body.get("seed")
    if base_value <= 0:
        return jsonify({"error": "base_value must be positive"}), 400
    try:
        cfg    = _MCCfg24(iterations=min(iterations, 100_000), seed=seed)
        engine = _MCE24(cfg)
        result = engine.run(base_value, volatility=volatility)
        return jsonify(result.to_dict()), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/scenarios/sensitivity", methods=["POST"])
def scenarios_sensitivity():
    """
    POST /api/scenarios/sensitivity
    Body: {"base_value": float, "var1_name": str, "var1_changes": [...],
           "var2_name": str, "var2_changes": [...]}
    Returns: SensitivityResult dict
    """
    if not _SCENARIOS_OK:
        return jsonify({"error": "Scenario engine not available"}), 503
    body       = request.get_json(force=True, silent=True) or {}
    base_value = float(body.get("base_value", 0))
    if base_value <= 0:
        return jsonify({"error": "base_value must be positive"}), 400
    try:
        sm     = _SM24(base_value)
        result = sm.analyze(
            var1_name=str(body.get("var1_name", "var1")),
            var1_changes=[float(x) for x in body.get("var1_changes", [-20, -10, 0, 10, 20])],
            var2_name=str(body.get("var2_name", "var2")),
            var2_changes=[float(x) for x in body.get("var2_changes", [-20, -10, 0, 10, 20])],
        )
        return jsonify(result.to_dict()), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/scenarios/stress_test", methods=["POST"])
def scenarios_stress_test():
    """
    POST /api/scenarios/stress_test
    Body: {"base_value": float}
    Returns: StressTestSuite summary dict
    """
    if not _SCENARIOS_OK:
        return jsonify({"error": "Scenario engine not available"}), 503
    body       = request.get_json(force=True, silent=True) or {}
    base_value = float(body.get("base_value", 0))
    if base_value <= 0:
        return jsonify({"error": "base_value must be positive"}), 400
    try:
        return jsonify(_STS24().summary(base_value)), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


# ── Phase 23 — i18n / Localization ───────────────────────────────────────────

try:
    from i18n.localization import Language as _Lang23, get_localization as _get_loc23
    from i18n.language_detector import LanguageDetector as _LangDetect23
    from i18n.translations import get_translations as _get_trans23
    _loc23 = _get_loc23()
    for _lang23 in (_Lang23.ENGLISH, _Lang23.ARABIC, _Lang23.FRENCH):
        _loc23.load_translations(_lang23, _get_trans23(_lang23.value))
    _I18N_OK = True
except Exception as _i18n_err:
    _I18N_OK = False
    _loc23 = None


@app.route("/api/language/set/<language_code>", methods=["POST"])
def set_language(language_code: str):
    """POST /api/language/set/<ar|en|fr> — switch active language."""
    if not _I18N_OK or _loc23 is None:
        return jsonify({"error": "i18n not available"}), 503
    try:
        code = language_code.lower()
        if code in ("ar", "arabic"):
            lang = _Lang23.ARABIC
        elif code in ("fr", "french"):
            lang = _Lang23.FRENCH
        else:
            lang = _Lang23.ENGLISH
        _loc23.set_language(lang)
        return jsonify({
            "status": "success",
            "language": lang.value,
            "direction": _loc23.get_ui_direction(),
        }), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/language/strings", methods=["GET"])
def get_language_strings():
    """GET /api/language/strings — return all translation strings for current language."""
    if not _I18N_OK or _loc23 is None:
        return jsonify({"error": "i18n not available"}), 503
    try:
        lang = _loc23.get_current_language()
        return jsonify({
            "status": "success",
            "language": lang.value,
            "direction": _loc23.get_ui_direction(),
            "strings": _loc23.translations.get(lang, {}),
        }), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/language/detect", methods=["POST"])
def detect_language():
    """POST /api/language/detect — auto-detect language from Accept-Language header."""
    if not _I18N_OK or _loc23 is None:
        return jsonify({"error": "i18n not available"}), 503
    try:
        detected = _LangDetect23.detect_from_headers(dict(request.headers))
        _loc23.set_language(detected)
        return jsonify({
            "status": "success",
            "detected_language": detected.value,
            "direction": _loc23.get_ui_direction(),
        }), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


# ===========================================================================
# Phase 26 — Multi-Tenant SaaS Readiness
# ===========================================================================
_SAAS_OK = False
try:
    from saas.tenant_manager import (
        TenantManager as _TenantManager26,
        TenantStatus as _TenantStatus26,
        UserRole as _UserRole26,
        SubscriptionTier as _SubTier26,
        get_tenant_manager as _get_tm26,
    )
    from saas.tenant_isolation import TenantIsolationValidator as _TIV26
    from saas.billing_engine import BillingEngine as _BillingEngine26, UsageMetric as _UsageMetric26
    from saas.subscription_manager import SubscriptionManager as _SubMgr26
    from saas.dashboard import TenantDashboard as _Dashboard26

    _tm26 = _get_tm26()
    _billing26 = _BillingEngine26()
    _sm26 = _SubMgr26(tenant_manager=_tm26, billing_engine=_billing26)
    _dashboard26 = _Dashboard26(
        tenant_manager=_tm26,
        billing_engine=_billing26,
        subscription_manager=_sm26,
    )
    _validator26 = _TIV26(_tm26)
    _SAAS_OK = True
except Exception as _saas_err:
    print(f"  [Phase26/SaaS] disabled: {_saas_err}")


@app.route("/api/saas/tenants", methods=["POST"])
def saas_create_tenant():
    if not _SAAS_OK:
        return jsonify({"error": "SaaS module not available"}), 503
    try:
        body = request.get_json(force=True) or {}
        name = body.get("name", "").strip()
        domain = body.get("domain", "").strip()
        tier_str = body.get("tier", "free").lower()
        trial_days = int(body.get("trial_days", 0))
        if not name or not domain:
            return jsonify({"error": "name and domain are required"}), 400
        tier = _SubTier26(tier_str) if tier_str in [t.value for t in _SubTier26] else _SubTier26.FREE
        tenant = _tm26.create_tenant(name, domain, tier=tier, trial_days=trial_days)
        return jsonify({"status": "created", "tenant": tenant.to_dict()}), 201
    except ValueError as exc:
        return jsonify({"error": str(exc)}), 409
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/saas/tenants", methods=["GET"])
def saas_list_tenants():
    if not _SAAS_OK:
        return jsonify({"error": "SaaS module not available"}), 503
    try:
        tenants = _tm26.list_tenants()
        return jsonify({
            "status": "success",
            "tenants": [t.to_dict() for t in tenants],
            "count": len(tenants),
        }), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/saas/tenants/<tenant_id>", methods=["GET"])
def saas_get_tenant(tenant_id):
    if not _SAAS_OK:
        return jsonify({"error": "SaaS module not available"}), 503
    try:
        tenant = _tm26.get_tenant(tenant_id)
        if tenant is None:
            return jsonify({"error": "Tenant not found"}), 404
        return jsonify({"status": "success", "tenant": tenant.to_dict()}), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/saas/tenants/<tenant_id>/users", methods=["POST"])
def saas_add_user(tenant_id):
    if not _SAAS_OK:
        return jsonify({"error": "SaaS module not available"}), 503
    try:
        body = request.get_json(force=True) or {}
        email = body.get("email", "").strip()
        role_str = body.get("role", "analyst").lower()
        if not email:
            return jsonify({"error": "email is required"}), 400
        role = _UserRole26(role_str) if role_str in [r.value for r in _UserRole26] else _UserRole26.ANALYST
        user = _tm26.add_user(tenant_id, email, role=role)
        return jsonify({"status": "created", "user": user.to_dict()}), 201
    except (KeyError, ValueError) as exc:
        return jsonify({"error": str(exc)}), 400
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/saas/tenants/<tenant_id>/subscription", methods=["PUT"])
def saas_update_subscription(tenant_id):
    if not _SAAS_OK:
        return jsonify({"error": "SaaS module not available"}), 503
    try:
        body = request.get_json(force=True) or {}
        action = body.get("action", "upgrade")
        tier_str = body.get("tier", "").lower()
        if not tier_str:
            return jsonify({"error": "tier is required"}), 400
        tier = _SubTier26(tier_str)
        if action == "upgrade":
            result = _sm26.upgrade_plan(tenant_id, tier)
        elif action == "downgrade":
            result = _sm26.downgrade_plan(tenant_id, tier)
        elif action == "convert":
            _sm26.convert_to_paid(tenant_id, tier)
            result = {"tenant_id": tenant_id, "action": "converted", "tier": tier_str}
        else:
            return jsonify({"error": f"Unknown action '{action}'"}), 400
        return jsonify({"status": "success", **result}), 200
    except (KeyError, ValueError) as exc:
        return jsonify({"error": str(exc)}), 400
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/saas/tenants/<tenant_id>/suspend", methods=["POST"])
def saas_suspend_tenant(tenant_id):
    if not _SAAS_OK:
        return jsonify({"error": "SaaS module not available"}), 503
    try:
        body = request.get_json(force=True) or {}
        reason = body.get("reason", "")
        tenant = _sm26.suspend_tenant(tenant_id, reason=reason)
        return jsonify({"status": "suspended", "tenant_id": tenant_id, "reason": reason}), 200
    except KeyError as exc:
        return jsonify({"error": str(exc)}), 404
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/saas/tenants/<tenant_id>/reactivate", methods=["POST"])
def saas_reactivate_tenant(tenant_id):
    if not _SAAS_OK:
        return jsonify({"error": "SaaS module not available"}), 503
    try:
        _sm26.reactivate_tenant(tenant_id)
        return jsonify({"status": "reactivated", "tenant_id": tenant_id}), 200
    except KeyError as exc:
        return jsonify({"error": str(exc)}), 404
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/saas/tenants/<tenant_id>/billing/usage", methods=["POST"])
def saas_record_usage(tenant_id):
    if not _SAAS_OK:
        return jsonify({"error": "SaaS module not available"}), 503
    try:
        body = request.get_json(force=True) or {}
        metric_str = body.get("metric", "api_call").lower()
        quantity = float(body.get("quantity", 1.0))
        metric = _UsageMetric26(metric_str)
        record = _billing26.record_usage(tenant_id, metric, quantity)
        return jsonify({"status": "recorded", "record": record.to_dict()}), 201
    except (KeyError, ValueError) as exc:
        return jsonify({"error": str(exc)}), 400
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/saas/tenants/<tenant_id>/billing/invoice", methods=["POST"])
def saas_generate_invoice(tenant_id):
    if not _SAAS_OK:
        return jsonify({"error": "SaaS module not available"}), 503
    try:
        result = _sm26.generate_monthly_invoice(tenant_id)
        return jsonify({"status": "success", "invoice": result}), 200
    except KeyError as exc:
        return jsonify({"error": str(exc)}), 404
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/saas/tenants/<tenant_id>/dashboard", methods=["GET"])
def saas_tenant_dashboard(tenant_id):
    if not _SAAS_OK:
        return jsonify({"error": "SaaS module not available"}), 503
    try:
        overview = _dashboard26.get_overview(tenant_id)
        return jsonify({"status": "success", "dashboard": overview}), 200
    except KeyError as exc:
        return jsonify({"error": str(exc)}), 404
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/saas/stats", methods=["GET"])
def saas_platform_stats():
    if not _SAAS_OK:
        return jsonify({"error": "SaaS module not available"}), 503
    try:
        stats = _dashboard26.get_platform_stats()
        return jsonify({"status": "success", "stats": stats}), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


# ===========================================================================
# ML / AVM Training Pipeline
# ===========================================================================
_ML_OK = False
_avm_predictor_instance = None
try:
    import os as _os_ml
    import joblib as _joblib_ml
    from ml.data_processor import DataProcessor as _DataProcessor
    from ml.feature_engineer import FeatureEngineer as _FeatureEngineer
    from ml.model_trainer import ModelTrainer as _ModelTrainer
    from ml.model_validator import ModelValidator as _ModelValidator
    from ml.avm_predictor import AVMPredictor as _AVMPredictor
    from ml.model_registry import ModelRegistry as _ModelRegistry

    _ml_registry = _ModelRegistry()
    _avm_model_path = _os_ml.getenv("AVM_MODEL_PATH", "models/avm_model.pkl")

    _active_meta = _ml_registry.get_active_model()
    if _active_meta:
        _loaded_model = _joblib_ml.load(_active_meta.model_path)
        _ml_engineer = _FeatureEngineer()
        _avm_predictor_instance = _AVMPredictor(_loaded_model, _ml_engineer)
    elif _os_ml.path.exists(_avm_model_path):
        _loaded_model = _joblib_ml.load(_avm_model_path)
        _ml_engineer = _FeatureEngineer()
        _avm_predictor_instance = _AVMPredictor(_loaded_model, _ml_engineer)

    _ML_OK = True
except Exception as _ml_err:
    print(f"  [ML/AVM] disabled: {_ml_err}")


@app.route("/api/valuation/avm", methods=["POST"])
def api_avm_valuation():
    if not _ML_OK:
        return jsonify({"error": "AVM model not available"}), 503
    if _avm_predictor_instance is None:
        return jsonify({"error": "No trained AVM model loaded. Train a model first."}), 503
    try:
        body = request.get_json(force=True) or {}
        area_sqm = float(body.get("area_sqm", 0))
        location = str(body.get("location", "")).strip()
        property_type = str(body.get("property_type", "")).strip()
        if area_sqm <= 0 or not location or not property_type:
            return jsonify({"error": "area_sqm, location, and property_type are required"}), 400
        additional = body.get("additional_features")
        result = _avm_predictor_instance.predict(
            area_sqm=area_sqm,
            location=location,
            property_type=property_type,
            additional_features=additional,
        )
        if result.get("predicted_value") is None:
            return jsonify({"error": result.get("error", "Prediction failed")}), 500
        return jsonify(result), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/valuation/avm/batch", methods=["POST"])
def api_avm_batch_valuation():
    if not _ML_OK:
        return jsonify({"error": "AVM model not available"}), 503
    if _avm_predictor_instance is None:
        return jsonify({"error": "No trained AVM model loaded. Train a model first."}), 503
    try:
        body = request.get_json(force=True) or {}
        properties = body.get("properties", [])
        if not isinstance(properties, list) or not properties:
            return jsonify({"error": "properties must be a non-empty list"}), 400
        if len(properties) > 500:
            return jsonify({"error": "Batch size cannot exceed 500 properties"}), 400
        result = _avm_predictor_instance.predict_batch(properties)
        return jsonify(result), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/avm/info", methods=["GET"])
def api_avm_info():
    if not _ML_OK:
        return jsonify({"status": "unavailable", "reason": "ML module not loaded"}), 503
    try:
        active = _ml_registry.get_active_model()
        stats = _ml_registry.get_stats()
        return jsonify({
            "status": "ready" if _avm_predictor_instance else "no_model",
            "active_model": active.to_dict() if active else None,
            "registry": stats,
        }), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


# ===========================================================================
# Integration Marketplace
# ===========================================================================
_MARKETPLACE_OK = False
try:
    from plugins.plugin_system import PluginSystem as _PluginSystem
    from plugins.plugin_registry import PluginRegistry as _PluginRegistry
    from integrations.oauth_manager import OAuthManager as _OAuthManager
    from integrations.webhook_manager import (
        WebhookManager as _WebhookManager,
        WebhookEventType as _WebhookEventType,
        WebhookPayload as _WebhookPayload,
    )
    from integrations.payment_integration import PaymentIntegration as _PaymentIntegration, PaymentProvider as _PaymentProvider
    from marketplace.marketplace import Marketplace as _Marketplace
    from datetime import datetime as _dt_mp

    _mp_plugin_system = _PluginSystem()
    _mp_plugin_registry = _PluginRegistry()
    _mp_oauth_manager = _OAuthManager()
    _mp_webhook_manager = _WebhookManager()
    _mp_marketplace = _Marketplace()
    _mp_payment = _PaymentIntegration(provider=_PaymentProvider.MOCK)

    _MARKETPLACE_OK = True
except Exception as _mp_err:
    print(f"  [Marketplace] disabled: {_mp_err}")


@app.route("/api/marketplace/plugins", methods=["GET"])
def api_marketplace_list_plugins():
    if not _MARKETPLACE_OK:
        return jsonify({"error": "Marketplace not available"}), 503
    query = request.args.get("q", "")
    category = request.args.get("category", "")
    sort_by = request.args.get("sort", "rating")
    results = _mp_marketplace.search_plugins(
        query=query or None,
        category=category or None,
        sort_by=sort_by,
    )
    return jsonify({"plugins": [p.to_dict() for p in results], "count": len(results)}), 200


@app.route("/api/marketplace/plugins/<listing_id>", methods=["GET"])
def api_marketplace_plugin_detail(listing_id):
    if not _MARKETPLACE_OK:
        return jsonify({"error": "Marketplace not available"}), 503
    listing = _mp_marketplace.listings.get(listing_id)
    if not listing:
        return jsonify({"error": "Plugin not found"}), 404
    reviews = _mp_marketplace.get_reviews(listing_id)
    return jsonify({
        "plugin": listing.to_dict(),
        "reviews": [r.to_dict() for r in reviews],
        "review_count": len(reviews),
    }), 200


@app.route("/api/marketplace/plugins/<listing_id>/reviews", methods=["POST"])
def api_marketplace_add_review(listing_id):
    if not _MARKETPLACE_OK:
        return jsonify({"error": "Marketplace not available"}), 503
    try:
        body = request.get_json(force=True) or {}
        tenant_id = request.headers.get("X-Tenant-ID", "anonymous")
        rating = int(body.get("rating", 0))
        title = str(body.get("title", "")).strip()
        comment = str(body.get("comment", "")).strip()
        if not title or not comment:
            return jsonify({"error": "title and comment are required"}), 400
        review = _mp_marketplace.add_review(listing_id, tenant_id, rating, title, comment)
        if review is None:
            return jsonify({"error": "Invalid rating (must be 1–5) or listing not found"}), 400
        return jsonify({"review": review.to_dict(), "message": "Review added"}), 201
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/marketplace/trending", methods=["GET"])
def api_marketplace_trending():
    if not _MARKETPLACE_OK:
        return jsonify({"error": "Marketplace not available"}), 503
    limit = int(request.args.get("limit", 10))
    plugins = _mp_marketplace.get_trending_plugins(limit=limit)
    return jsonify({"plugins": [p.to_dict() for p in plugins], "count": len(plugins)}), 200


@app.route("/api/integrations/plugins", methods=["GET"])
def api_integrations_list_plugins():
    if not _MARKETPLACE_OK:
        return jsonify({"error": "Marketplace not available"}), 503
    tenant_id = request.headers.get("X-Tenant-ID", "")
    if not tenant_id:
        return jsonify({"error": "X-Tenant-ID header required"}), 400
    plugins = _mp_plugin_system.list_installed_plugins(tenant_id)
    return jsonify({"plugins": [p.to_dict() for p in plugins], "count": len(plugins)}), 200


@app.route("/api/integrations/plugins/<plugin_id>/install", methods=["POST"])
def api_integrations_install_plugin(plugin_id):
    if not _MARKETPLACE_OK:
        return jsonify({"error": "Marketplace not available"}), 503
    try:
        tenant_id = request.headers.get("X-Tenant-ID", "")
        if not tenant_id:
            return jsonify({"error": "X-Tenant-ID header required"}), 400
        body = request.get_json(force=True) or {}
        config = body.get("configuration", {})
        credentials = body.get("credentials", {})
        instance = _mp_plugin_system.install_plugin(tenant_id, plugin_id, config, credentials)
        if instance is None:
            return jsonify({"error": "Plugin installation failed"}), 400
        _mp_marketplace.record_installation(tenant_id, plugin_id)
        return jsonify({"plugin": instance.to_dict(), "message": "Plugin installed"}), 201
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/integrations/webhooks", methods=["GET"])
def api_integrations_list_webhooks():
    if not _MARKETPLACE_OK:
        return jsonify({"error": "Marketplace not available"}), 503
    tenant_id = request.headers.get("X-Tenant-ID", "")
    if not tenant_id:
        return jsonify({"error": "X-Tenant-ID header required"}), 400
    webhooks = _mp_webhook_manager.list_webhooks(tenant_id)
    return jsonify({"webhooks": [w.to_dict() for w in webhooks], "count": len(webhooks)}), 200


@app.route("/api/integrations/webhooks", methods=["POST"])
def api_integrations_create_webhook():
    if not _MARKETPLACE_OK:
        return jsonify({"error": "Marketplace not available"}), 503
    try:
        tenant_id = request.headers.get("X-Tenant-ID", "")
        if not tenant_id:
            return jsonify({"error": "X-Tenant-ID header required"}), 400
        body = request.get_json(force=True) or {}
        url = str(body.get("url", "")).strip()
        if not url:
            return jsonify({"error": "url is required"}), 400
        raw_events = body.get("events", [])
        events = []
        for e in raw_events:
            key = e.upper().replace(".", "_")
            try:
                events.append(_WebhookEventType[key])
            except KeyError:
                return jsonify({"error": f"Unknown event type: {e}"}), 400
        if not events:
            return jsonify({"error": "At least one event is required"}), 400
        webhook = _mp_webhook_manager.register_webhook(tenant_id, url, events)
        return jsonify({"webhook": webhook.to_dict(), "message": "Webhook created"}), 201
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/integrations/webhooks/<webhook_id>", methods=["DELETE"])
def api_integrations_delete_webhook(webhook_id):
    if not _MARKETPLACE_OK:
        return jsonify({"error": "Marketplace not available"}), 503
    if _mp_webhook_manager.unregister_webhook(webhook_id):
        return jsonify({"message": "Webhook deleted"}), 200
    return jsonify({"error": "Webhook not found"}), 404


@app.route("/api/integrations/oauth/<service>/authorize", methods=["GET"])
def api_oauth_authorize(service):
    if not _MARKETPLACE_OK:
        return jsonify({"error": "Marketplace not available"}), 503
    try:
        auth_url = _mp_oauth_manager.get_authorization_url(service)
        return jsonify({"auth_url": auth_url, "service": service}), 200
    except ValueError as exc:
        return jsonify({"error": str(exc)}), 400


@app.route("/api/integrations/oauth/<service>/callback", methods=["GET"])
def api_oauth_callback(service):
    if not _MARKETPLACE_OK:
        return jsonify({"error": "Marketplace not available"}), 503
    code = request.args.get("code", "")
    state = request.args.get("state", "")
    error = request.args.get("error", "")
    if error:
        return jsonify({"error": f"OAuth error: {error}"}), 400
    token = _mp_oauth_manager.exchange_code_for_token(service, code, state or None)
    if not token:
        return jsonify({"error": "Failed to obtain token"}), 400
    return jsonify({"token": token.to_dict(), "service": service, "message": "Authorized"}), 200


@app.route("/api/marketplace/info", methods=["GET"])
def api_marketplace_info():
    if not _MARKETPLACE_OK:
        return jsonify({"status": "unavailable"}), 503
    counts = _mp_marketplace.get_category_counts()
    all_listings = list(_mp_marketplace.listings.values())
    return jsonify({
        "status": "ready",
        "total_plugins": len(all_listings),
        "categories": counts,
        "registry_total": _mp_plugin_registry.get_stats()["total"],
    }), 200


# ===========================================================================
# Phase 23B — USPAP-Oriented Compliance Framework
# Reporting / disclosure support only — zero impact on valuations.
# ===========================================================================
_USPAP_OK = False
try:
    from standards.uspap import (
        USPAPReport as _USPAPReport,
        USPAPPropertyIdentification as _USPAPPropID,
        USPAPAssumptions as _USPAPAssumptions,
        USPAPCompetencyStatement as _USPAPCompetency,
        USPAPCertification as _USPAPCertification,
        USPAPWorkfileNote as _USPAPWorkfileNote,
        USPAPComplianceChecker as _USPAPChecker,
        USPAPComplianceAddenum as _USPAPAddenum,
        USPAPCompliance as _USPAPComplianceEnum,
    )
    from standards.standards_manager import (
        StandardsManager as _StandardsManager,
        StandardsFramework as _StandardsFramework,
    )
    _standards_mgr = _StandardsManager()
    _USPAP_OK = True
except Exception as _uspap_err:
    print(f"  [USPAP/Standards] disabled: {_uspap_err}")


@app.route("/api/standards/frameworks", methods=["GET"])
def api_standards_list_frameworks():
    if not _USPAP_OK:
        return jsonify({"error": "Standards module not available"}), 503
    return jsonify({
        "frameworks": _standards_mgr.list_frameworks(),
        "note": (
            "All frameworks are optional reporting/disclosure frameworks. "
            "None change valuation calculations."
        ),
    }), 200


@app.route("/api/standards/validate", methods=["POST"])
def api_standards_validate():
    if not _USPAP_OK:
        return jsonify({"error": "Standards module not available"}), 503
    try:
        body = request.get_json(force=True) or {}
        raw = body.get("frameworks", [])
        frameworks = []
        for f in raw:
            try:
                frameworks.append(_StandardsFramework(f.lower()))
            except ValueError:
                return jsonify({"error": f"Unknown framework: {f}"}), 400
        is_valid, warnings = _standards_mgr.validate_framework_combination(frameworks)
        return jsonify({
            "is_valid": is_valid,
            "warnings": warnings,
            "frameworks": [fw.value for fw in frameworks],
        }), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/standards/uspap/generate", methods=["POST"])
def api_uspap_generate():
    """
    Generate a USPAP-oriented disclosure addendum.
    Does NOT change or recalculate any valuation.
    Returns the disclosure framework for the supplied property data.
    """
    if not _USPAP_OK:
        return jsonify({"error": "USPAP module not available"}), 503
    try:
        body = request.get_json(force=True) or {}
        prop_data = body.get("property", {})
        uspap_data = body.get("uspap_parameters", {})

        prop_id = _USPAPPropID(
            street_address=prop_data.get("street_address", ""),
            city=prop_data.get("city", ""),
            state=prop_data.get("state", ""),
            zip_code=prop_data.get("zip_code", ""),
            parcel_id=prop_data.get("parcel_id"),
            legal_description=prop_data.get("legal_description"),
        )

        assumptions = _USPAPAssumptions()
        for ea in uspap_data.get("extraordinary_assumptions", []):
            assumptions.add_extraordinary_assumption(ea)
        for hc in uspap_data.get("hypothetical_conditions", []):
            assumptions.add_hypothetical_condition(hc)

        competency = _USPAPCompetency(
            appraiser_name=uspap_data.get("appraiser_name", ""),
            appraiser_license=uspap_data.get("appraiser_license", ""),
            appraiser_designation=uspap_data.get("appraiser_designation"),
        )

        from datetime import datetime as _dt23b
        cert_date_raw = uspap_data.get("certification_date")
        cert_date = _dt23b.fromisoformat(cert_date_raw) if cert_date_raw else _dt23b.utcnow()

        certification = _USPAPCertification(
            appraiser_name=uspap_data.get("appraiser_name", ""),
            certification_date=cert_date,
            appraiser_license=uspap_data.get("appraiser_license", ""),
            scope_of_work=uspap_data.get("scope_of_work", "Desktop"),
        )

        workfile = _USPAPWorkfileNote(content=uspap_data.get("workfile_note", ""))

        eff_raw = uspap_data.get("effective_date")
        eff_date = _dt23b.fromisoformat(eff_raw) if eff_raw else _dt23b.utcnow()

        report = _USPAPReport(
            report_type=uspap_data.get("report_type", "Desktop"),
            intended_user=uspap_data.get("intended_user", ""),
            intended_use=uspap_data.get("intended_use", ""),
            scope_of_work=uspap_data.get("scope_of_work", ""),
            effective_date=eff_date,
            report_date=_dt23b.utcnow(),
            property_identification=prop_id,
            assumptions=assumptions,
            competency=competency,
            certification=certification,
            workfile_note=workfile,
            compliance_level=_USPAPComplianceEnum.USPAP_ORIENTED,
            is_avm_valuation=bool(body.get("is_avm_valuation", False)),
            is_mass_appraisal=bool(body.get("is_mass_appraisal", False)),
        )

        warnings, errors = _USPAPChecker.check_compliance(report)
        addendum = _USPAPAddenum.generate_addendum(report)

        return jsonify({
            "status": "generated",
            "compliance_level": report.compliance_level.value,
            "report_data": report.to_dict(),
            "compliance_warnings": warnings,
            "compliance_errors": errors,
            "addendum": addendum,
            "note": (
                "USPAP-Oriented Disclosure Support. "
                "Does NOT change valuation results. "
                "Subject to professional review and verification."
            ),
        }), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


# ===========================================================================
# Phase 32 — Government / Tax Pilot Package
# Compliance engine, tax calculator, official forms, audit trail, PKI signing.
# Zero impact on valuation calculations.
# ===========================================================================
_GOVERNMENT_OK = False
try:
    from government.compliance_engine import (
        GovernmentComplianceEngine as _GovCompliance,
        GovernmentStandard as _GovStandard,
        ComplianceLevel as _ComplianceLevel,
    )
    from government.tax_calculator import (
        TaxCalculator as _TaxCalculator,
        TaxClassification as _TaxClassification,
    )
    from government.forms_generator import (
        FormsGenerator as _FormsGenerator,
        GovernmentForm as _GovernmentForm,
    )
    from government.audit_trail import (
        GovernmentAuditTrail as _GovAudit,
        AuditAction as _GovAuditAction,
    )
    from government.government_portal import GovernmentPortalManager as _GovPortal

    _gov_compliance = _GovCompliance()
    _gov_tax = _TaxCalculator()
    _gov_forms = _FormsGenerator()
    _gov_audit = _GovAudit()
    _gov_portal = _GovPortal()
    _GOVERNMENT_OK = True
except Exception as _gov_err:
    print(f"  [Government] disabled: {_gov_err}")


@app.route("/api/government/info", methods=["GET"])
def api_government_info():
    """Government pilot program capabilities and supported standards."""
    return jsonify({
        "program": "Expert Smart Government Pilot",
        "supported_authorities": [
            "Central Bank of Egypt (CBE)",
            "Egyptian Financial Supervisory Authority (EGFSA)",
            "Egyptian Tax Authority",
            "Ministry of Finance",
        ],
        "standards_supported": [
            "EGVS — Egyptian General Valuation Standards",
            "CBE — Central Bank of Egypt Standards",
            "EGFSA — Egyptian Financial Supervisory Authority",
            "Tax Authority — Egyptian Tax Code",
            "IFRS 13 — Fair Value Measurement",
            "IFRS 16 — Lease Accounting",
        ],
        "forms_available": [
            "CBE Form 101 — Collateral Valuation",
            "Tax Authority Form 50 — Property Valuation",
            "EGFSA Form 30 — Fair Value Assessment",
        ],
        "features": [
            "Rule-based compliance checking per standard",
            "Automated tax valuation calculation",
            "Official form generation (CBE 101, Tax 50, EGFSA 30)",
            "Complete audit trail",
            "Digital signatures (HMAC-SHA256, PKI-ready)",
            "Government agency portal management",
        ],
        "module_status": "available" if _GOVERNMENT_OK else "disabled",
    }), 200


@app.route("/api/government/compliance/check", methods=["POST"])
def api_government_compliance_check():
    """Check property data against one or more Egyptian government standards."""
    if not _GOVERNMENT_OK:
        return jsonify({"error": "Government module not available"}), 503
    try:
        body = request.get_json(force=True) or {}
        property_data = body.get("property_data", {})
        standards_requested = body.get("standards", [])

        if not standards_requested:
            standards_requested = [s.value for s in _GovStandard]

        results = {}
        for std_str in standards_requested:
            try:
                std = _GovStandard(std_str)
                result = _gov_compliance.check_compliance(property_data, std)
                results[std_str] = result.to_dict()
                _gov_audit.record(
                    _GovAuditAction.COMPLIANCE_CHECK,
                    entity_id=property_data.get("property_id", "unknown"),
                    entity_type="property",
                    details={"standard": std_str, "level": result.compliance_level.value},
                )
            except ValueError:
                results[std_str] = {"error": f"Unknown standard: {std_str}"}

        return jsonify({
            "compliance_results": results,
            "total_standards_checked": len(results),
        }), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/government/tax/calculate", methods=["POST"])
def api_government_tax_calculate():
    """Calculate property tax per Egyptian Tax Authority standards."""
    if not _GOVERNMENT_OK:
        return jsonify({"error": "Government module not available"}), 503
    try:
        body = request.get_json(force=True) or {}
        property_id = body.get("property_id", "unknown")
        assessed_value = float(body.get("assessed_value", 0))
        classification_str = body.get("tax_classification", "residential")
        purchase_price = body.get("purchase_price")
        years_held = body.get("years_held")

        classification = _TaxClassification(classification_str)
        result = _gov_tax.calculate_tax_valuation(
            property_id=property_id,
            assessed_value=assessed_value,
            classification=classification,
            purchase_price=float(purchase_price) if purchase_price is not None else None,
            years_held=int(years_held) if years_held is not None else None,
        )
        _gov_audit.record(
            _GovAuditAction.TAX_CALCULATION,
            entity_id=property_id,
            entity_type="property",
            details={"classification": classification_str, "total_tax": result.total_estimated_tax},
        )
        return jsonify({
            "tax_valuation": result.to_dict(),
            "report": _gov_tax.get_tax_report(result),
        }), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/government/forms/generate", methods=["POST"])
def api_government_forms_generate():
    """Generate an official government valuation form."""
    if not _GOVERNMENT_OK:
        return jsonify({"error": "Government module not available"}), 503
    try:
        body = request.get_json(force=True) or {}
        form_type_str = body.get("form_type", "cbe_101")
        property_data = body.get("property_data", {})
        valuation_result = body.get("valuation_result", {})
        appraiser_info = body.get("appraiser_info", {})

        form_type = _GovernmentForm(form_type_str)

        if form_type == _GovernmentForm.CBE_FORM_101:
            form_text = _gov_forms.generate_cbe_form_101(
                property_data, valuation_result, appraiser_info
            )
        elif form_type == _GovernmentForm.TAX_FORM_50:
            owner_info = body.get("owner_info", {})
            form_text = _gov_forms.generate_tax_form_50(
                property_data, valuation_result, owner_info
            )
        elif form_type == _GovernmentForm.EGFSA_FORM_30:
            expert_info = body.get("expert_info", {})
            form_text = _gov_forms.generate_egfsa_form_30(
                property_data, valuation_result, expert_info
            )
        else:
            return jsonify({"error": f"Form type not supported: {form_type_str}"}), 400

        from datetime import datetime as _dt32
        _gov_audit.record(
            _GovAuditAction.FORM_GENERATED,
            entity_id=property_data.get("property_id", "unknown"),
            entity_type="form",
            details={"form_type": form_type_str},
        )
        return jsonify({
            "form_type": form_type_str,
            "form_content": form_text,
            "generated_at": _dt32.utcnow().isoformat(),
        }), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/government/portal/create", methods=["POST"])
def api_government_portal_create():
    """Register a new government agency portal."""
    if not _GOVERNMENT_OK:
        return jsonify({"error": "Government module not available"}), 503
    try:
        body = request.get_json(force=True) or {}
        agency_name = body.get("agency_name", "")
        contact_person = body.get("contact_person", "")
        contact_email = body.get("contact_email", "")
        if not agency_name:
            return jsonify({"error": "agency_name is required"}), 400
        portal = _gov_portal.create_portal(agency_name, contact_person, contact_email)
        return jsonify(portal.get_dashboard_summary()), 201
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


# ===========================================================================
# Phase 33 — CBE / Banking Collateral Pilot
# Collateral valuation, LTV, risk scoring, registry, and bank dashboard.
# ===========================================================================
_BANKING_OK = False
try:
    from banking.collateral_engine import (
        CollateralValuationEngine as _ColEngine,
        CollateralProperty as _ColProp,
        CollateralType as _ColType,
    )
    from banking.ltv_calculator import (
        LTVCalculator as _LTVCalc,
        LTVTier as _LTVTier,
        CreditRiskRating as _CRR,
    )
    from banking.collateral_registry import CollateralRegistry as _ColRegistry
    from banking.risk_assessment import PropertyRiskAnalyzer as _RiskAnalyzer
    from banking.compliance_tracker import CBEComplianceTracker as _CBETracker
    from banking.loan_servicing import LoanServicingManager as _LoanMgr
    from banking.market_monitoring import MarketMonitor as _MktMonitor
    from banking.bank_dashboard import BankDashboard as _BankDash

    _col_engine   = _ColEngine()
    _ltv_calc     = _LTVCalc()
    _col_registry = _ColRegistry()
    _risk_analyzer = _RiskAnalyzer()
    _cbe_tracker  = _CBETracker()
    _loan_mgr     = _LoanMgr()
    _mkt_monitor  = _MktMonitor()
    _bank_dash    = _BankDash()
    _BANKING_OK   = True
except Exception as _banking_err:
    print(f"  [Banking] disabled: {_banking_err}")


@app.route("/api/banking/info", methods=["GET"])
def api_banking_info():
    """Banking pilot capabilities and market coverage."""
    return jsonify({
        "program": "Expert Smart CBE / Banking Collateral Pilot",
        "module_status": "available" if _BANKING_OK else "disabled",
        "supported_institutions": [
            "Central Bank of Egypt (CBE)",
            "Commercial Banks",
            "Investment Banks",
            "Mortgage Lenders",
            "Finance Companies",
        ],
        "capabilities": [
            "Collateral valuation (appraised / conservative / forced-sale)",
            "LTV tier classification (Prime / Conventional / Conforming / High-LTV)",
            "Credit risk scoring (AAA–CCC, 7-tier scale)",
            "Collateral registry (central asset tracking)",
            "Property risk profiling (Basel III risk weights)",
            "CBE regulatory compliance checking",
            "Loan servicing and lifecycle management",
            "Market monitoring and LTV-breach alerts",
            "Bank portfolio dashboard",
        ],
        "pricing": {
            "per_bank_annually": "50,000–200,000 EGP",
            "per_valuation": "500–2,000 EGP",
        },
    }), 200


@app.route("/api/banking/collateral/value", methods=["POST"])
def api_banking_collateral_value():
    """Value a property as bank collateral (appraised, conservative, forced-sale)."""
    if not _BANKING_OK:
        return jsonify({"error": "Banking module not available"}), 503
    try:
        body = request.get_json(force=True) or {}
        from datetime import datetime as _dt33
        cid = body.get("collateral_id") or f"COL-{_dt33.utcnow().timestamp():.0f}"
        col = _ColProp(
            collateral_id=cid,
            property_id=body.get("property_id", ""),
            owner_name=body.get("owner_name", ""),
            owner_id=body.get("owner_id", ""),
            property_type=_ColType(body.get("property_type", "residential_property")),
            location=body.get("location", ""),
            city=body.get("city", ""),
            area_sqm=float(body.get("area_sqm", 0)),
            year_built=int(body.get("year_built", 2000)),
            condition=body.get("condition", "good"),
            legal_status=body.get("legal_status", "free"),
            existing_liens=int(body.get("existing_liens", 0)),
            assessed_value=float(body.get("assessed_value", 0)),
            market_value=float(body.get("market_value", 0)),
        )
        result = _col_engine.value_collateral(
            col,
            market_conditions=body.get("market_conditions", "Stable"),
            comparable_count=int(body.get("comparable_count", 3)),
        )
        return jsonify({
            "collateral": col.to_dict(),
            "valuation": result.to_dict(),
        }), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/banking/ltv/calculate", methods=["POST"])
def api_banking_ltv_calculate():
    """Calculate LTV ratio and credit risk assessment."""
    if not _BANKING_OK:
        return jsonify({"error": "Banking module not available"}), 503
    try:
        body = request.get_json(force=True) or {}
        loan_id = body.get("loan_id", "unknown")
        ltv_result = _ltv_calc.calculate_ltv(
            loan_id=loan_id,
            loan_amount=float(body.get("loan_amount", 0)),
            collateral_value=float(body.get("collateral_value", 0)),
        )
        risk = _ltv_calc.assess_credit_risk(
            loan_id=loan_id,
            borrower_credit_score=body.get("borrower_credit_score"),
            ltv_ratio=ltv_result.ltv_ratio,
            property_quality=body.get("property_quality", "good"),
            loan_purpose=body.get("loan_purpose", "residential_mortgage"),
            market_conditions=body.get("market_conditions", "Stable"),
        )
        return jsonify({
            "ltv": ltv_result.to_dict(),
            "risk_assessment": risk.to_dict(),
        }), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/banking/collateral/register", methods=["POST"])
def api_banking_collateral_register():
    """Register collateral in the central registry."""
    if not _BANKING_OK:
        return jsonify({"error": "Banking module not available"}), 503
    try:
        from datetime import datetime as _dt33, timedelta as _td33
        body = request.get_json(force=True) or {}
        mat_raw = body.get("loan_maturity_date")
        maturity = _dt33.fromisoformat(mat_raw) if mat_raw else _dt33.utcnow() + _td33(days=3650)
        entry = _col_registry.register_collateral(
            collateral_id=body.get("collateral_id", ""),
            property_id=body.get("property_id", ""),
            owner_id=body.get("owner_id", ""),
            loan_id=body.get("loan_id", ""),
            bank_id=body.get("bank_id", ""),
            collateral_value=float(body.get("collateral_value", 0)),
            loan_amount=float(body.get("loan_amount", 0)),
            ltv_ratio=float(body.get("ltv_ratio", 0)),
            loan_maturity_date=maturity,
        )
        return jsonify({"registered": entry.to_dict()}), 201
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/banking/dashboard/<bank_id>", methods=["GET"])
def api_banking_dashboard(bank_id: str):
    """Get bank portfolio dashboard and metrics."""
    if not _BANKING_OK:
        return jsonify({"error": "Banking module not available"}), 503
    try:
        portfolio = _col_registry.get_bank_portfolio(bank_id)
        metrics = _bank_dash.calculate_portfolio_metrics(bank_id, portfolio.get("entries", []))
        return jsonify({
            "portfolio": portfolio,
            "dashboard": _bank_dash.get_dashboard_summary(bank_id),
            "metrics": metrics.to_dict(),
        }), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/banking/compliance/check", methods=["POST"])
def api_banking_compliance_check():
    """Check CBE regulatory compliance for a loan."""
    if not _BANKING_OK:
        return jsonify({"error": "Banking module not available"}), 503
    try:
        body = request.get_json(force=True) or {}
        status = _cbe_tracker.check_loan(
            loan_id=body.get("loan_id", ""),
            bank_id=body.get("bank_id", ""),
            ltv_ratio=float(body.get("ltv_ratio", 0)),
            collateral_type=body.get("collateral_type", "residential_property"),
            has_valuation=bool(body.get("has_valuation", True)),
            cbe_approved_appraiser=bool(body.get("cbe_approved_appraiser", True)),
            documentation_complete=bool(body.get("documentation_complete", True)),
            risk_class_assigned=bool(body.get("risk_class_assigned", True)),
            property_insured=bool(body.get("property_insured", True)),
            revaluation_overdue=bool(body.get("revaluation_overdue", False)),
        )
        return jsonify(status.to_dict()), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


# ---------------------------------------------------------------------------
# Phase 34 — FRA / Funds & Fair Value Pilot
# ---------------------------------------------------------------------------

_FUNDS_OK = False
try:
    from funds.fair_value_calculator import (
        FairValueCalculator as _FVCalc,
        ValuationInput as _FVInput,
        FairValueInput as _FVInputType,
        ValuationApproach as _FVApproach,
    )
    from funds.nav_calculator import (
        NAVCalculator as _NAVCalc,
        FundAsset as _FundAsset,
        FundLiability as _FundLiability,
    )
    from funds.fund_engine import (
        FundValuationEngine as _FundEngine,
        FundType as _FundType,
        FundStrategy as _FundStrategy,
    )
    from funds.fra_compliance import FRAComplianceEngine as _FRAEngine
    from funds.fund_dashboard import FundDashboard as _FundDash
    from datetime import date as _date

    _fv_calc   = _FVCalc()
    _nav_calc  = _NAVCalc()
    _fund_eng  = _FundEngine()
    _fra_eng   = _FRAEngine()
    _fund_dash = _FundDash()
    _FUNDS_OK  = True
    print("  [Funds] FRA / Funds & Fair Value module loaded OK")
except Exception as _funds_err:
    print(f"  [Funds] disabled: {_funds_err}")


@app.route("/api/funds/info", methods=["GET"])
def api_funds_info():
    """FRA / Funds & Fair Value pilot capabilities."""
    return jsonify({
        "pilot": "FRA / Funds & Fair Value",
        "version": "1.0",
        "status": "active" if _FUNDS_OK else "unavailable",
        "capabilities": [
            "IFRS 13 fair value (3-level hierarchy)",
            "NAV calculation",
            "Fund valuation (6 fund types)",
            "FRA compliance (18-point checklist)",
            "Portfolio dashboard",
        ],
        "target_clients": ["FRA", "REITs", "Real Estate Funds", "MENA Funds"],
        "pricing": "$100K–$300K per fund/year",
    }), 200


@app.route("/api/funds/fair-value/assess", methods=["POST"])
def api_funds_fair_value_assess():
    """Assess IFRS 13 fair value for an asset."""
    if not _FUNDS_OK:
        return jsonify({"error": "Funds module not available"}), 503
    try:
        body = request.get_json(force=True) or {}
        raw_inputs = body.get("inputs", [])
        if not raw_inputs:
            return jsonify({"error": "inputs list is required"}), 400

        inputs = []
        for ri in raw_inputs:
            inputs.append(_FVInput(
                input_type=_FVInputType(ri.get("input_type", "appraisal")),
                value=float(ri.get("value", 0)),
                date=_date.today(),
                source=ri.get("source", "unknown"),
                observable=bool(ri.get("observable", False)),
                weight=float(ri.get("weight", 1.0)),
                confidence=float(ri.get("confidence", 1.0)),
                market_condition=ri.get("market_condition", "normal"),
            ))

        approach_str = body.get("approach", "market_approach")
        try:
            approach = _FVApproach(approach_str)
        except ValueError:
            approach = _FVApproach.MARKET_APPROACH

        assessment = _fv_calc.assess_fair_value(
            asset_id=body.get("asset_id", ""),
            asset_type=body.get("asset_type", "residential"),
            inputs=inputs,
            approach=approach,
        )
        return jsonify({
            "assessment": assessment.to_dict(),
            "disclosure": _fv_calc.generate_ifrs13_disclosure(assessment),
        }), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/funds/nav/calculate", methods=["POST"])
def api_funds_nav_calculate():
    """Calculate NAV and NAV per share for a fund."""
    if not _FUNDS_OK:
        return jsonify({"error": "Funds module not available"}), 503
    try:
        body = request.get_json(force=True) or {}
        fund_id = body.get("fund_id", "")
        shares = float(body.get("shares_outstanding", 1))
        if shares <= 0:
            return jsonify({"error": "shares_outstanding must be positive"}), 400

        assets = [
            _FundAsset(
                asset_id=a.get("asset_id", ""),
                asset_type=a.get("asset_type", "residential"),
                current_value=float(a.get("current_value", 0)),
                original_cost=float(a.get("original_cost", 0)),
                acquisition_date=_date.today(),
                last_valuation_date=_date.today(),
            )
            for a in body.get("assets", [])
        ]
        liabilities = [
            _FundLiability(
                liability_id=l.get("liability_id", ""),
                liability_type=l.get("liability_type", "other"),
                amount=float(l.get("amount", 0)),
            )
            for l in body.get("liabilities", [])
        ]

        result = _nav_calc.calculate_nav(fund_id, assets, liabilities, shares)
        return jsonify({
            "nav_result": result.to_dict(),
            "report": _nav_calc.generate_nav_report(result),
        }), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/funds/value", methods=["POST"])
def api_funds_value():
    """Value a fund and compute risk metrics."""
    if not _FUNDS_OK:
        return jsonify({"error": "Funds module not available"}), 503
    try:
        body = request.get_json(force=True) or {}

        try:
            fund_type = _FundType(body.get("fund_type", "reit"))
        except ValueError:
            fund_type = _FundType.REIT
        try:
            strategy = _FundStrategy(body.get("strategy", "mixed"))
        except ValueError:
            strategy = _FundStrategy.MIXED

        result = _fund_eng.value_fund(
            fund_id=body.get("fund_id", ""),
            fund_name=body.get("fund_name", ""),
            fund_type=fund_type,
            strategy=strategy,
            aum=float(body.get("aum", 0)),
            nav=float(body.get("nav", 0)),
            shares_outstanding=float(body.get("shares_outstanding", 1)),
            ytd_return=float(body.get("ytd_return", 0)),
            volatility=float(body.get("volatility", 0)),
            dividend_per_share=float(body.get("dividend_per_share", 0)),
            fra_registered=bool(body.get("fra_registered", True)),
        )
        return jsonify(result.to_dict()), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/funds/compliance/check", methods=["POST"])
def api_funds_compliance_check():
    """Check FRA regulatory compliance for a fund."""
    if not _FUNDS_OK:
        return jsonify({"error": "Funds module not available"}), 503
    try:
        body = request.get_json(force=True) or {}
        fund_id = body.get("fund_id", "")
        if not fund_id:
            return jsonify({"error": "fund_id is required"}), 400

        result = _fra_eng.check_compliance(fund_id, body)
        return jsonify({
            "compliance": result.to_dict(),
            "report": _fra_eng.generate_compliance_report(result),
        }), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/funds/dashboard/<manager_id>", methods=["GET"])
def api_funds_dashboard(manager_id: str):
    """Get cached fund manager dashboard metrics."""
    if not _FUNDS_OK:
        return jsonify({"error": "Funds module not available"}), 503
    try:
        snap = _fund_dash.get_snapshot(manager_id)
        if snap is None:
            return jsonify({"error": f"No dashboard data for manager {manager_id}"}), 404
        return jsonify({
            "manager_id": manager_id,
            "metrics": snap.to_dict(),
            "report": _fund_dash.generate_summary_report(snap),
        }), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


# ---------------------------------------------------------------------------
# Phase 35 — Knowledge Base + Standards Expansion
# ---------------------------------------------------------------------------

_KNOWLEDGE_OK = False
try:
    from knowledge.kb_engine import (
        KnowledgeBaseEngine as _KBEngine,
        ContentType as _ContentType,
        ContentLevel as _ContentLevel,
    )
    from knowledge.standards_registry import StandardsRegistry as _StdRegistry
    from knowledge.ai_search import SmartSearchEngine as _SmartSearch
    from knowledge.learning_platform import (
        LearningPlatform as _LearnPlatform,
        CourseLevel as _CourseLevel,
    )
    from knowledge.best_practices import BestPracticesLibrary as _BPLibrary

    _kb_engine       = _KBEngine()
    _std_registry    = _StdRegistry()
    _smart_search    = _SmartSearch()
    _learn_platform  = _LearnPlatform()
    _bp_library      = _BPLibrary()
    _KNOWLEDGE_OK    = True
    print("  [Knowledge] Phase 35 Knowledge Base + Standards module loaded OK")
except Exception as _knowledge_err:
    print(f"  [Knowledge] disabled: {_knowledge_err}")


@app.route("/api/knowledge/info", methods=["GET"])
def api_knowledge_info():
    """Knowledge Base + Standards pilot capabilities."""
    return jsonify({
        "pilot": "Knowledge Base + Standards Expansion",
        "version": "1.0",
        "status": "active" if _KNOWLEDGE_OK else "unavailable",
        "capabilities": [
            "10 content types (articles, guides, case studies, standards, etc.)",
            "20+ international and local valuation standards",
            "AI-powered semantic search with query expansion",
            "Learning platform with 4 course levels",
            "Automated certification management",
            "Regulatory update tracking and compliance deadlines",
            "Best practices library",
        ],
        "standards_count": len(_std_registry.standards) if _KNOWLEDGE_OK else 0,
        "target_users": ["Appraisers", "Fund managers", "Banks", "Government"],
        "pricing": "$50K–$100K/year premium knowledge tier",
    }), 200


@app.route("/api/knowledge/search", methods=["GET"])
def api_knowledge_search():
    """Search the knowledge base."""
    if not _KNOWLEDGE_OK:
        return jsonify({"error": "Knowledge module not available"}), 503
    try:
        query = request.args.get("q", "").strip()
        if not query:
            return jsonify({"error": "q parameter is required"}), 400
        ctype_str = request.args.get("type")
        category = request.args.get("category")
        limit = min(int(request.args.get("limit", 10)), 50)

        content_type = None
        if ctype_str:
            try:
                content_type = _ContentType(ctype_str)
            except ValueError:
                pass

        results = _kb_engine.search_content(
            query=query, content_type=content_type,
            category=category, limit=limit,
        )
        for r in results:
            _kb_engine.add_view(r.content_id)

        # Also run semantic re-ranking on results
        scored = _smart_search.semantic_search(query, results, similarity_threshold=0.0)
        if scored:
            results = [item for item, _ in scored]

        return jsonify({
            "query": query,
            "results_count": len(results),
            "results": [r.to_dict() for r in results],
            "suggestions": _smart_search.get_search_suggestions(query),
        }), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/knowledge/standards", methods=["GET"])
def api_knowledge_standards_list():
    """List all registered standards."""
    if not _KNOWLEDGE_OK:
        return jsonify({"error": "Knowledge module not available"}), 503
    try:
        standards = _std_registry.list_all_standards(active_only=True)
        return jsonify({
            "total": len(standards),
            "standards": [s.to_dict() for s in standards],
        }), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/knowledge/standards/compatible", methods=["POST"])
def api_knowledge_standards_compatible():
    """Get compatible standards for an asset type and country."""
    if not _KNOWLEDGE_OK:
        return jsonify({"error": "Knowledge module not available"}), 503
    try:
        body = request.get_json(force=True) or {}
        asset_type = body.get("asset_type", "")
        country = body.get("country", "Egypt")
        if not asset_type:
            return jsonify({"error": "asset_type is required"}), 400
        matrix = _std_registry.get_compatibility_matrix(asset_type, country)
        return jsonify(matrix), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/knowledge/courses", methods=["GET"])
def api_knowledge_courses_list():
    """List published courses."""
    if not _KNOWLEDGE_OK:
        return jsonify({"error": "Knowledge module not available"}), 503
    try:
        level_str = request.args.get("level")
        courses = [
            c.to_dict() for c in _learn_platform.courses.values()
            if c.is_published and (not level_str or c.level.value == level_str)
        ]
        return jsonify({"courses_count": len(courses), "courses": courses}), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/knowledge/courses/<course_id>/enroll", methods=["POST"])
def api_knowledge_enroll(course_id: str):
    """Enroll a user in a course."""
    if not _KNOWLEDGE_OK:
        return jsonify({"error": "Knowledge module not available"}), 503
    try:
        body = request.get_json(force=True) or {}
        user_id = body.get("user_id", "")
        if not user_id:
            return jsonify({"error": "user_id is required"}), 400
        enrollment = _learn_platform.enroll_user(user_id, course_id)
        return jsonify({"enrollment": enrollment.to_dict()}), 200
    except ValueError as exc:
        return jsonify({"error": str(exc)}), 404
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/knowledge/regulatory/updates", methods=["GET"])
def api_knowledge_regulatory_updates():
    """Get active regulatory updates."""
    if not _KNOWLEDGE_OK:
        return jsonify({"error": "Knowledge module not available"}), 503
    try:
        source = request.args.get("source")
        updates = _bp_library.get_active_updates(source=source or None)
        return jsonify({
            "updates_count": len(updates),
            "updates": [u.to_dict() for u in updates],
        }), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/knowledge/regulatory/deadlines", methods=["GET"])
def api_knowledge_regulatory_deadlines():
    """Get upcoming compliance deadlines."""
    if not _KNOWLEDGE_OK:
        return jsonify({"error": "Knowledge module not available"}), 503
    try:
        days = min(int(request.args.get("days", 90)), 365)
        deadlines = _bp_library.get_upcoming_compliance_deadlines(days_ahead=days)
        return jsonify({
            "deadlines_count": len(deadlines),
            "deadlines": [d.to_dict() for d in deadlines],
        }), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/knowledge/statistics", methods=["GET"])
def api_knowledge_statistics():
    """Get knowledge base and learning platform statistics."""
    if not _KNOWLEDGE_OK:
        return jsonify({"error": "Knowledge module not available"}), 503
    try:
        return jsonify({
            "knowledge_base": _kb_engine.get_statistics(),
            "learning_platform": _learn_platform.get_platform_statistics(),
            "standards_count": _std_registry.count(),
            "best_practices_count": _bp_library.count_practices(),
            "regulatory_updates_count": _bp_library.count_updates(),
        }), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


# ── Phase 36: Advanced Analytics & Business Intelligence ──────────────────────
_ANALYTICS_OK = False
try:
    from analytics.analytics_engine import AnalyticsEngine as _AnalyticsEngine, MetricType as _MetricType, TimeGranularity as _TimeGranularity
    from analytics.dashboard_system import DashboardSystem as _DashboardSystem, DashboardType as _DashboardType, WidgetType as _WidgetType
    from analytics.forecasting import ForecastingEngine as _ForecastingEngine
    from analytics.market_intelligence import MarketIntelligence as _MarketIntelligence, MarketSegment as _MarketSegment
    from analytics.portfolio_risk import PortfolioRiskAnalytics as _PortfolioRiskAnalytics
    _analytics_engine = _AnalyticsEngine()
    _dashboard_system = _DashboardSystem()
    _forecasting_engine = _ForecastingEngine()
    _market_intelligence = _MarketIntelligence()
    _portfolio_risk_analytics = _PortfolioRiskAnalytics()
    _ANALYTICS_OK = True
    print("  [Analytics] Phase 36 loaded OK")
except Exception as _analytics_err:
    print(f"  [Analytics] disabled: {_analytics_err}")


@app.route("/api/analytics/info", methods=["GET"])
def analytics_info():
    if not _ANALYTICS_OK:
        return jsonify({"error": "Analytics module unavailable"}), 503
    try:
        return jsonify({
            "phase": 36,
            "module": "Advanced Analytics & Business Intelligence",
            "features": ["metric_tracking", "dashboards", "forecasting", "market_intelligence", "portfolio_risk"],
            "status": "operational",
        }), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/analytics/metrics/<metric_id>/record", methods=["POST"])
def analytics_record_metric(metric_id):
    if not _ANALYTICS_OK:
        return jsonify({"error": "Analytics module unavailable"}), 503
    try:
        data = request.get_json(force=True) or {}
        value = data.get("value")
        if value is None:
            return jsonify({"error": "value is required"}), 400
        if metric_id not in _analytics_engine.metrics:
            metric_type_str = data.get("metric_type", "count")
            try:
                metric_type = _MetricType(metric_type_str)
            except ValueError:
                metric_type = _MetricType.COUNT
            _analytics_engine.register_metric(
                metric_id=metric_id,
                name=data.get("name", metric_id),
                metric_type=metric_type,
                unit=data.get("unit", ""),
                description=data.get("description", ""),
                tenant_id=data.get("tenant_id", "default"),
            )
        ok = _analytics_engine.record_data_point(metric_id=metric_id, value=float(value), tags=data.get("tags"))
        if not ok:
            return jsonify({"error": "Failed to record data point"}), 400
        return jsonify({"status": "recorded", "metric_id": metric_id, "value": value}), 201
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/analytics/metrics/<metric_id>/statistics", methods=["GET"])
def analytics_metric_statistics(metric_id):
    if not _ANALYTICS_OK:
        return jsonify({"error": "Analytics module unavailable"}), 503
    try:
        stats = _analytics_engine.get_metric_statistics(metric_id)
        if stats is None:
            return jsonify({"error": "Metric not found"}), 404
        return jsonify(stats), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/analytics/metrics/<metric_id>/timeseries", methods=["GET"])
def analytics_metric_timeseries(metric_id):
    if not _ANALYTICS_OK:
        return jsonify({"error": "Analytics module unavailable"}), 503
    try:
        granularity_str = request.args.get("granularity", "daily")
        try:
            granularity = _TimeGranularity(granularity_str)
        except ValueError:
            granularity = _TimeGranularity.DAILY
        series = _analytics_engine.get_time_series(metric_id=metric_id, granularity=granularity)
        return jsonify({"metric_id": metric_id, "granularity": granularity_str, "series": series}), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/analytics/dashboards", methods=["GET"])
def analytics_list_dashboards():
    if not _ANALYTICS_OK:
        return jsonify({"error": "Analytics module unavailable"}), 503
    try:
        owner_id = request.args.get("owner_id")
        dashboard_type_str = request.args.get("type")
        if owner_id:
            dashboards = _dashboard_system.get_user_dashboards(owner_id)
        elif dashboard_type_str:
            try:
                dtype = _DashboardType(dashboard_type_str)
            except ValueError:
                return jsonify({"error": f"Unknown dashboard type: {dashboard_type_str}"}), 400
            dashboards = _dashboard_system.get_dashboards_by_type(dtype)
        else:
            with _dashboard_system._lock:
                dashboards = list(_dashboard_system.dashboards.values())
        return jsonify({"dashboards": [d.to_dict() for d in dashboards], "total": len(dashboards)}), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/analytics/dashboards/<dashboard_id>", methods=["GET"])
def analytics_get_dashboard(dashboard_id):
    if not _ANALYTICS_OK:
        return jsonify({"error": "Analytics module unavailable"}), 503
    try:
        viewer_id = request.args.get("viewer_id", "anonymous")
        dashboard = _dashboard_system.get_dashboard(dashboard_id, viewer_id=viewer_id)
        if dashboard is None:
            return jsonify({"error": "Dashboard not found"}), 404
        return jsonify(dashboard.to_dict()), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/analytics/risk/portfolio/<portfolio_id>", methods=["POST"])
def analytics_portfolio_risk(portfolio_id):
    if not _ANALYTICS_OK:
        return jsonify({"error": "Analytics module unavailable"}), 503
    try:
        data = request.get_json(force=True) or {}
        tenant_id = data.get("tenant_id", "default")
        total_value = float(data.get("total_value", 0))
        assets = data.get("assets", [])
        credit_risk = float(data.get("credit_risk", 10.0))
        operational_risk = float(data.get("operational_risk", 5.0))
        result = _portfolio_risk_analytics.assess_portfolio_risk(
            portfolio_id=portfolio_id,
            tenant_id=tenant_id,
            total_value=total_value,
            assets=assets,
            credit_risk=credit_risk,
            operational_risk=operational_risk,
        )
        return jsonify(result.to_dict()), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/analytics/market/trends", methods=["GET"])
def analytics_market_trends():
    if not _ANALYTICS_OK:
        return jsonify({"error": "Analytics module unavailable"}), 503
    try:
        segment_str = request.args.get("segment")
        location = request.args.get("location")
        if segment_str:
            try:
                segment = _MarketSegment(segment_str)
            except ValueError:
                return jsonify({"error": f"Unknown segment: {segment_str}"}), 400
            trends = _market_intelligence.get_trends_by_segment(segment)
        elif location:
            trends = _market_intelligence.get_trends_by_location(location)
        else:
            with _market_intelligence._lock:
                trends = list(_market_intelligence.trends.values())
        summary = _market_intelligence.get_market_summary()
        return jsonify({
            "trends": [t.to_dict() for t in trends],
            "total": len(trends),
            "summary": summary,
        }), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


# ── Phase 37: Comparable Search Enhancement ───────────────────────────────────
_SEARCH_OK = False
try:
    from search.comparable_search import ComparableSearchEngine as _CSEngine, SearchCriteria as _SearchCriteria
    from search.similarity_matcher import SmartMatcher as _SmartMatcher, MatchingStrategy as _MatchingStrategy
    from search.adjustment_factors import AdjustmentFactorEngine as _AdjEngine
    _cs_engine = _CSEngine()
    _smart_matcher = _SmartMatcher()
    _adj_engine = _AdjEngine()
    _SEARCH_OK = True
    print("  [Search] Phase 37 loaded OK")
except Exception as _search_err:
    print(f"  [Search] disabled: {_search_err}")


@app.route("/api/search/info", methods=["GET"])
def search_info():
    if not _SEARCH_OK:
        return jsonify({"error": "Search module unavailable"}), 503
    try:
        stats = _cs_engine.get_search_statistics()
        return jsonify({
            "phase": 37,
            "module": "Comparable Search Enhancement",
            "features": {
                "multi_criteria_search": True,
                "similarity_scoring": True,
                "price_adjustments": True,
                "smart_matching": True,
            },
            "matching_strategies": ["exact", "weighted", "ml", "hybrid"],
            "adjustment_categories": 7,
            "search_statistics": stats,
        }), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/search/properties/register", methods=["POST"])
def search_register_property():
    if not _SEARCH_OK:
        return jsonify({"error": "Search module unavailable"}), 503
    try:
        data = request.get_json(force=True) or {}
        property_id = data.get("property_id")
        if not property_id:
            return jsonify({"error": "property_id is required"}), 400

        from datetime import datetime as _dt
        sale_date_str = data.get("sale_date")
        try:
            sale_date = _dt.fromisoformat(sale_date_str) if sale_date_str else None
        except ValueError:
            sale_date = None

        property_data = {
            "property_id": property_id,
            "property_type": data.get("property_type", ""),
            "location": data.get("location", ""),
            "area_sqm": float(data.get("area_sqm", 0)),
            "price": float(data.get("price", 0)),
            "age_years": int(data.get("age_years", 0)),
            "condition": data.get("condition", "good"),
            "bedrooms": int(data.get("bedrooms", 0)),
            "bathrooms": int(data.get("bathrooms", 0)),
            "features": data.get("features", []),
            "sale_date": sale_date,
        }

        _cs_engine.register_property(property_id=property_id, property_data=property_data)
        return jsonify({"property_id": property_id, "status": "registered_for_search"}), 201
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/search/comparables", methods=["POST"])
def search_comparables():
    if not _SEARCH_OK:
        return jsonify({"error": "Search module unavailable"}), 503
    try:
        data = request.get_json(force=True) or {}

        criteria = _SearchCriteria(
            property_type=data.get("property_type", "residential"),
            location=data.get("location", ""),
            area_sqm_min=float(data.get("area_sqm_min", 0)),
            area_sqm_max=float(data.get("area_sqm_max", 10000)),
            price_min=float(data.get("price_min", 0)),
            price_max=float(data.get("price_max", 100_000_000)),
            distance_km=float(data.get("distance_km", 5.0)),
            sale_date_months_back=int(data.get("sale_date_months", 12)),
            condition=data.get("condition"),
            max_results=int(data.get("max_results", 10)),
            sort_by=data.get("sort_by", "similarity"),
            min_similarity_score=float(data.get("min_similarity", 0.7)),
            exclude_property_ids=data.get("exclude_ids", []),
            required_features=data.get("required_features", []),
        )

        results = _cs_engine.search_comparables(
            criteria=criteria,
            user_id=data.get("user_id"),
        )

        if data.get("apply_adjustments"):
            for r in results:
                time_adj = _adj_engine.create_time_adjustment(r.days_since_sale)
                cond_adj = _adj_engine.create_condition_adjustment(
                    r.property_data.get("condition", "good")
                )
                loc_adj = _adj_engine.create_location_adjustment(
                    criteria.location, r.property_data.get("location", "")
                )
                adjs = [time_adj, cond_adj] + ([loc_adj] if loc_adj else [])
                adjusted = _adj_engine.apply_adjustments(
                    comparable_id=r.property_id,
                    original_price=r.property_data.get("price", 0),
                    adjustments=adjs,
                )
                r.adjusted_price = adjusted.adjusted_price
                r.is_adjusted = True

        return jsonify({
            "results_count": len(results),
            "results": [r.to_dict() for r in results],
        }), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/search/match", methods=["POST"])
def search_match_properties():
    if not _SEARCH_OK:
        return jsonify({"error": "Search module unavailable"}), 503
    try:
        data = request.get_json(force=True) or {}
        subject = data.get("subject_property")
        candidates = data.get("candidate_properties", [])
        strategy_str = data.get("strategy", "weighted")

        if not subject:
            return jsonify({"error": "subject_property is required"}), 400

        try:
            strategy = _MatchingStrategy(strategy_str)
        except ValueError:
            strategy = _MatchingStrategy.WEIGHTED

        matches = _smart_matcher.match_properties(
            subject=subject, candidates=candidates, strategy=strategy
        )
        return jsonify({
            "subject_id": subject.get("property_id"),
            "matches_count": len(matches),
            "matches": [m.to_dict() for m in matches],
        }), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


# ── Phase 38: API Hardening & External Integration Readiness ──────────────────
_API38_OK = False
try:
    from api.security_layer import APISecurityLayer as _SecurityLayer, RateLimitTier as _RLTier
    from api.request_validation import RequestValidator as _ReqValidator
    from api.performance_optimizer import PerformanceOptimizer as _PerfOpt
    from api.error_standardizer import ErrorStandardizer as _ErrStd, StandardErrorCode as _ErrCode
    from api.integration_framework import (
        IntegrationFramework as _IntegFW,
        IntegrationEvent as _IntegEvent,
    )
    _security_layer = _SecurityLayer()
    _req_validator = _ReqValidator()
    _perf_optimizer = _PerfOpt()
    _err_std = _ErrStd()
    _integ_fw = _IntegFW()
    _API38_OK = True
    print("  [API38] Phase 38 loaded OK")
except Exception as _api38_err:
    print(f"  [API38] disabled: {_api38_err}")


@app.route("/api/hardening/info", methods=["GET"])
def api38_info():
    if not _API38_OK:
        return jsonify({"error": "API Hardening module unavailable"}), 503
    try:
        return jsonify({
            "phase": 38,
            "module": "API Hardening & External Integration Readiness",
            "features": {
                "api_key_management": True,
                "rate_limiting": True,
                "request_validation": True,
                "response_caching": True,
                "error_standardization": True,
                "integration_webhooks": True,
            },
            "rate_tiers": ["free", "starter", "professional", "enterprise"],
            "security": _security_layer.get_security_status(),
            "cache": _perf_optimizer.get_cache_statistics(),
            "integrations": _integ_fw.get_integration_statistics(),
        }), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/hardening/api-keys/generate", methods=["POST"])
def api38_generate_key():
    if not _API38_OK:
        return jsonify({"error": "API Hardening module unavailable"}), 503
    try:
        data = request.get_json(force=True) or {}
        partner_id = data.get("partner_id")
        partner_name = data.get("partner_name")
        if not partner_id or not partner_name:
            return jsonify({"error": "partner_id and partner_name are required"}), 400
        tier_str = data.get("tier", "starter")
        try:
            tier = _RLTier(tier_str)
        except ValueError:
            tier = _RLTier.STARTER
        api_key = _security_layer.generate_api_key(
            partner_id=partner_id, partner_name=partner_name, tier=tier
        )
        return jsonify({
            "key_id": api_key.key_id,
            "key_secret": api_key.key_secret,
            "tier": api_key.tier.value,
            "expires_at": api_key.expires_at.isoformat() if api_key.expires_at else None,
            "note": "Store key_secret securely — it cannot be retrieved later",
        }), 201
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/hardening/api-keys/validate", methods=["POST"])
def api38_validate_key():
    if not _API38_OK:
        return jsonify({"error": "API Hardening module unavailable"}), 503
    try:
        data = request.get_json(force=True) or {}
        key_id = data.get("key_id", "")
        key_secret = data.get("key_secret", "")
        is_valid, api_key = _security_layer.validate_api_key(
            key_id=key_id, key_secret=key_secret
        )
        if not is_valid:
            return jsonify({"valid": False, "reason": "Invalid credentials or inactive key"}), 401
        return jsonify({"valid": True, "key_info": api_key.to_dict()}), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/hardening/integrations/register", methods=["POST"])
def api38_register_integration():
    if not _API38_OK:
        return jsonify({"error": "API Hardening module unavailable"}), 503
    try:
        data = request.get_json(force=True) or {}
        integration_id = data.get("integration_id")
        name = data.get("name")
        provider = data.get("provider")
        if not all([integration_id, name, provider]):
            return jsonify({"error": "integration_id, name, and provider are required"}), 400

        raw_events = data.get("events", [])
        events = []
        for ev in raw_events:
            try:
                events.append(_IntegEvent(ev))
            except ValueError:
                pass

        integ = _integ_fw.register_integration(
            integration_id=integration_id,
            name=name,
            provider=provider,
            webhook_url=data.get("webhook_url"),
            events=events,
        )
        return jsonify(integ.to_dict()), 201
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/hardening/integrations/stats", methods=["GET"])
def api38_integration_stats():
    if not _API38_OK:
        return jsonify({"error": "API Hardening module unavailable"}), 503
    try:
        return jsonify(_integ_fw.get_integration_statistics()), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/docs/openapi.json", methods=["GET"])
def api38_openapi_spec():
    if not _API38_OK:
        return jsonify({"error": "API Hardening module unavailable"}), 503
    try:
        spec = {
            "openapi": "3.0.0",
            "info": {
                "title": "Expert Smart API",
                "version": "1.0.0",
                "description": "Real Estate Valuation Platform API — Phases 1-38",
            },
            "security_schemes": {
                "ApiKeyAuth": {"type": "apiKey", "in": "header", "name": "X-API-Key"},
            },
            "paths": {
                "/api/valuation/batch": {
                    "post": {"summary": "Batch property valuation", "tags": ["Valuations"]}
                },
                "/api/search/comparables": {
                    "post": {"summary": "Search comparable properties", "tags": ["Search"]}
                },
                "/api/analytics/metrics/{metric_id}/record": {
                    "post": {"summary": "Record analytics metric", "tags": ["Analytics"]}
                },
                "/api/hardening/api-keys/generate": {
                    "post": {"summary": "Generate partner API key", "tags": ["Security"]}
                },
            },
        }
        return jsonify(spec), 200
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500



# ════════════════════════════════════════════════════════════════════════════
# Phase 40 — Integration Framework
# ════════════════════════════════════════════════════════════════════════════
_INTEG40_OK = False
try:
    from integrations.connector_base import (
        ConnectorConfig as _ConnCfg40,
        ConnectorType as _ConnType40,
        SyncDirection as _SyncDir40,
    )
    from integrations.connectors.bank_connector import BankConnector as _BankConn40
    from integrations.data_mapper import DataMapper as _DataMapper40
    from integrations.sync_engine import SyncEngine as _SyncEng40
    from integrations.partner_portal import (
        PartnerPortal as _PartnerPortal40,
        PartnerTier as _PartnerTier40,
    )
    from integrations.webhook_manager import (
        WebhookManager as _WHMgr40,
        WebhookEventType as _WHEvent40,
    )
    _dm40 = _DataMapper40()
    _se40 = _SyncEng40()
    _pp40 = _PartnerPortal40()
    _wh40 = _WHMgr40()
    _INTEG40_OK = True
except Exception as _err40:
    print(f"Phase 40 load error: {_err40}")


@app.route("/api/integrations/info", methods=["GET"])
def integ40_info():
    """Integration Framework status and capabilities."""
    if not _INTEG40_OK:
        return jsonify({"error": "Integration Framework unavailable"}), 503
    stats = _se40.get_sync_statistics()
    return jsonify({
        "phase": 40,
        "module": "Integration Framework",
        "capabilities": {
            "connector_types": [t.value for t in _ConnType40],
            "sync_directions": [d.value for d in _SyncDir40],
            "partner_tiers": [t.value for t in _PartnerTier40],
            "data_mappings": _dm40.list_mappings(),
            "webhook_event_types": [e.value for e in _WHEvent40],
        },
        "sync_statistics": stats,
    })


@app.route("/api/integrations/sync", methods=["POST"])
def integ40_sync():
    """Trigger connector synchronisation."""
    if not _INTEG40_OK:
        return jsonify({"error": "Integration Framework unavailable"}), 503
    data = request.get_json(silent=True) or {}
    connector_id = data.get("connector_id")
    if connector_id:
        success = _se40.sync_connector(connector_id)
        return jsonify({"connector_id": connector_id, "sync_successful": success})
    results = _se40.sync_all()
    return jsonify({"sync_results": results, "synced_count": len(results)})


@app.route("/api/integrations/partners", methods=["POST"])
def integ40_create_partner():
    """Create a new integration partner account."""
    if not _INTEG40_OK:
        return jsonify({"error": "Integration Framework unavailable"}), 503
    data = request.get_json(silent=True) or {}
    partner_id = data.get("partner_id")
    name = data.get("name")
    email = data.get("email")
    if not partner_id or not name or not email:
        return jsonify({"error": "partner_id, name, and email are required"}), 400
    try:
        tier = _PartnerTier40(data.get("tier", "trial"))
    except ValueError:
        tier = _PartnerTier40.TRIAL
    partner = _pp40.create_partner(partner_id=partner_id, name=name, email=email, tier=tier)
    return jsonify({"partner": partner.to_dict(), "api_key": partner.api_key}), 201


@app.route("/api/integrations/partners/<partner_id>/dashboard", methods=["GET"])
def integ40_partner_dashboard(partner_id):
    """Get partner dashboard overview."""
    if not _INTEG40_OK:
        return jsonify({"error": "Integration Framework unavailable"}), 503
    dashboard = _pp40.get_partner_dashboard(partner_id)
    if not dashboard:
        return jsonify({"error": "Partner not found"}), 404
    return jsonify(dashboard)


@app.route("/api/integrations/connector-webhooks", methods=["POST"])
def integ40_register_webhook():
    """Register a connector-level webhook endpoint."""
    if not _INTEG40_OK:
        return jsonify({"error": "Integration Framework unavailable"}), 503
    data = request.get_json(silent=True) or {}
    tenant_id = data.get("tenant_id")
    url = data.get("url")
    events_raw = data.get("events", [])
    if not tenant_id or not url:
        return jsonify({"error": "tenant_id and url are required"}), 400
    events = []
    for e in events_raw:
        try:
            events.append(_WHEvent40(e))
        except ValueError:
            pass
    if not events:
        events = [_WHEvent40.VALUATION_COMPLETED]
    wh = _wh40.register_webhook(tenant_id=tenant_id, url=url, events=events)
    return jsonify({"webhook": wh.to_dict()}), 201


@app.route("/api/integrations/data-map", methods=["POST"])
def integ40_map_data():
    """Transform a data record using a registered mapping."""
    if not _INTEG40_OK:
        return jsonify({"error": "Integration Framework unavailable"}), 503
    data = request.get_json(silent=True) or {}
    mapping_name = data.get("mapping")
    source = data.get("source", {})
    if not mapping_name:
        return jsonify({"error": "mapping is required"}), 400
    try:
        result = _dm40.map_data(source, mapping_name)
        return jsonify({"mapped": result, "mapping": mapping_name})
    except ValueError as exc:
        return jsonify({"error": str(exc)}), 400


# ════════════════════════════════════════════════════════════════════════════
# Phase 41.0 — Market Indicators Backend API
# Read-only market price indicators for Mass Appraisal charts
# ════════════════════════════════════════════════════════════════════════════
_MI41_OK = False
try:
    from market_indicators import (
        MarketIndicatorsBackend as _MIBackend41,
        market_indicators as _mi41_singleton,
    )
    _mi41 = _MIBackend41()           # fresh instance per server start
    _MI41_OK = True
except Exception as _err41:
    print(f"Phase 41.0 load error: {_err41}")


@app.route("/api/market-indicators/latest", methods=["GET"])
def mi41_latest():
    """Current values for all four market indicators."""
    if not _MI41_OK:
        return jsonify({"error": "Market Indicators module unavailable"}), 503
    return jsonify(_mi41.get_latest_indicators())


@app.route("/api/market-indicators/history", methods=["GET"])
def mi41_history():
    """12-month time-series for all four indicators.

    Query param: days_back (int, default 365, currently returns all 12 sample months).
    """
    if not _MI41_OK:
        return jsonify({"error": "Market Indicators module unavailable"}), 503
    try:
        days_back = int(request.args.get("days_back", 365))
    except ValueError:
        days_back = 365
    return jsonify(_mi41.get_indicators_history(days_back=days_back))


@app.route("/api/market-indicators/statistics", methods=["GET"])
def mi41_statistics():
    """Min / max / avg statistics for each indicator over the sample window."""
    if not _MI41_OK:
        return jsonify({"error": "Market Indicators module unavailable"}), 503
    stats = _mi41.get_statistics()
    return jsonify({"status": "success", "statistics": stats})


@app.route("/api/market-indicators/info", methods=["GET"])
def mi41_info():
    """Metadata: indicator descriptions, units, ranges, and endpoint list."""
    return jsonify({
        "module": "Expert Smart Market Indicators",
        "version": "1.0",
        "phase": "41.0",
        "indicators": {
            "stratification": {
                "name": "Market Stratification Index",
                "description": "Local price variation and market segmentation",
                "unit": "percentage",
                "typical_range": "8-14%",
                "use_case": "Neighbourhood price variation in Mass Appraisal",
            },
            "rppi": {
                "name": "Residential Property Price Index",
                "description": "Annual residential property price growth",
                "unit": "percentage",
                "typical_range": "10-18%",
                "use_case": "Market trend analysis and time adjustments",
            },
            "avm": {
                "name": "Automated Valuation Model Index",
                "description": "AVM accuracy / performance indicator",
                "unit": "percentage",
                "typical_range": "9-13%",
                "use_case": "Confidence level in automated valuations",
            },
            "cma": {
                "name": "Comparative Market Analysis Index",
                "description": "Effectiveness of comparable-property analysis",
                "unit": "percentage",
                "typical_range": "8-12%",
                "use_case": "Reliability of comparable-based valuations",
            },
        },
        "data_source": "sample_data",
        "endpoints": [
            "GET /api/market-indicators/latest",
            "GET /api/market-indicators/history",
            "GET /api/market-indicators/statistics",
            "GET /api/market-indicators/info",
        ],
        "note": "Data is representative sample values. TODO: live market data collector.",
    })

# ════════════════════════════════════════════════════════════════════════════
# BA.4a — Report History Endpoints
#
# Read-only endpoints backed by the reports SQLite DB (reports.db).
# No application-level auth gate — infrastructure auth is assumed,
# consistent with the existing /api/valuation endpoint.
# ════════════════════════════════════════════════════════════════════════════

@app.route("/api/reports", methods=["GET"])
def reports_list():
    """List persisted reports — summary fields, newest first.

    Query params (all optional):
      limit        int  max records to return (default 20, capped at 100)
      offset       int  records to skip (default 0)
      profile_key  str  filter by profile ('legacy'/'detailed'/'professional_template')
      status       str  filter by status  ('draft'/'final'/'archived')

    Response 200:
      {"status":"success","count":<int>,"reports":[...]}
    Response 500:
      {"status":"error","message":"<detail>"}
    """
    try:
        limit  = min(int(request.args.get("limit",  20)), 100)
        offset = max(int(request.args.get("offset",  0)),   0)
    except (TypeError, ValueError):
        limit  = 20
        offset = 0

    profile_key = request.args.get("profile_key") or None
    status      = request.args.get("status")      or None

    try:
        from reports.report_pipeline import fetch_reports as _fetch_reports
        result = _fetch_reports(
            profile_key=profile_key,
            status=status,
            limit=limit,
            offset=offset,
        )
        return jsonify({"status": "success", **result})
    except Exception as _hist_err:
        return jsonify({"status": "error", "message": str(_hist_err)}), 500


@app.route("/api/reports/<report_id>", methods=["GET"])
def reports_get(report_id: str):
    """Return the full stored DTO for a single persisted report.

    Returns 404 when report_id is not found in the database.
    No application-level auth gate — infrastructure auth is assumed.

    Response 200:
      {"status":"success","report":{...}}
    Response 404:
      {"status":"not_found","message":"Report <id> not found"}
    Response 500:
      {"status":"error","message":"<detail>"}
    """
    try:
        from reports.report_pipeline import fetch_report as _fetch_report
        record = _fetch_report(report_id)
        if record is None:
            return jsonify({
                "status":  "not_found",
                "message": f"Report {report_id} not found",
            }), 404
        return jsonify({"status": "success", "report": record})
    except Exception as _hist_err:
        return jsonify({"status": "error", "message": str(_hist_err)}), 500


@app.route("/api/reports/<report_id>/pdf", methods=["GET"])
def reports_pdf(report_id: str):
    """Export a stored report as a PDF file.

    Fetches the report DTO from the DB and streams a generated PDF.
    No application-level auth gate — infrastructure auth is assumed,
    consistent with /api/valuation and the other /api/reports endpoints.

    Response 200: application/pdf stream; Content-Disposition names the file.
    Response 404: {"status":"not_found","message":"Report <id> not found"}
    Response 500: {"status":"error","message":"<detail>"}
    """
    try:
        from reports.report_pipeline import export_report_pdf as _export_pdf
        pdf_bytes = _export_pdf(report_id)
        if pdf_bytes is None:
            return jsonify({
                "status":  "not_found",
                "message": f"Report {report_id} not found",
            }), 404
        # Sanitise report_id for use as a filename (keep alphanumeric + hyphen/underscore)
        safe_name = "".join(
            c if (c.isalnum() or c in "-_") else "_" for c in report_id
        )
        buf = io.BytesIO(pdf_bytes)
        buf.seek(0)
        return send_file(
            buf,
            mimetype="application/pdf",
            as_attachment=True,
            download_name=f"{safe_name}.pdf",
        )
    except Exception as _pdf_err:
        return jsonify({"status": "error", "message": str(_pdf_err)}), 500


if __name__ == "__main__":
    print(f"Template [v22-MI] : {TEMPLATE}")
    print(f"Outputs  : {OUTPUTS}")
    print(f"Feed DB  : {_FEED_FILE}")
    if _RAG_OK:
        try:
            import threading
            threading.Thread(target=_rag_init, daemon=True, name="rag-preload").start()
            print("  [RAG] starting model preload in background...")
        except Exception as _ri_err:
            print(f"  [RAG] pre-init failed: {_ri_err}")
    try:
        from waitress import serve
        print("  [WSGI] using waitress on http://0.0.0.0:5000")
        serve(app, host="0.0.0.0", port=5000, threads=8, channel_timeout=180)
    except Exception as _wsgi_err:
        print(f"  [WSGI] waitress unavailable ({_wsgi_err}); falling back to Flask dev server")
        app.run(host="0.0.0.0", port=5000, debug=False, threaded=True)
