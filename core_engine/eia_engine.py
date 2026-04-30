"""
eia_engine.py
==============
Environmental Impact Assessment (EIA) engine for the Expert_Smart platform.

Implements:
    1. وصف المشروع                      Project description
    2. تحليل الموقع GIS                  Site analysis (proximity scores)
    3. خط الأساس البيئي                 Environmental baseline
    4. الانبعاثات والمخلفات             Emissions & waste quantification
    5. IMPACT MATRIX                     Activity × element → Severity × Probability
    6. تقييم الأثر                       Aggregate impact assessment
    7. التخفيف والالتزام (ISO 14001)    Mitigation & compliance
    8. الرصد والمتابعة                  Monitoring KPIs
    9. التصنيف البيئي A/B/C             Environmental classification
    10. الربط الاستثماري — ERF           Investment linkage (Environmental Risk Factor)

Standards referenced:
    - ISO 14001:2015 (Environmental Management Systems)
    - Egyptian Environment Law No. 4/1994 (and amendments by Law 105/2015)
    - World Bank Environmental & Social Framework (ESF) — Categorization A/B/C
    - IFC Performance Standards on Environmental and Social Sustainability
    - IVS 104 §60 — Investment Value (environmental discount adjustment)

Public API:
    run_eia_assessment(payload)              -> dict
    write_eia_word_report(result, path)      -> bool
    write_eia_excel_report(result, dir)      -> Optional[str]
"""

from __future__ import annotations
import os
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple


_BASE_DIR = os.path.dirname(os.path.abspath(__file__))
_DEF_OUT = os.path.join(_BASE_DIR, "outputs", "reports")


# ════════════════════════════════════════════════════════════════════════════
#  CONSTANTS
# ════════════════════════════════════════════════════════════════════════════

# Encoding: low / medium / high → 1 / 2 / 3
_LEVEL_TO_SCORE = {
    "low": 1, "منخفض": 1, "low ": 1, "1": 1, 1: 1,
    "medium": 2, "med": 2, "متوسط": 2, "moderate": 2, "2": 2, 2: 2,
    "high": 3, "عالي": 3, "high ": 3, "3": 3, 3: 3,
}

# Score → label
_SCORE_TO_LABEL_AR = {1: "منخفض", 2: "متوسط", 3: "عالي"}

# Category thresholds on a 1–9 average impact-score scale
#   Avg < 3.0  → A (Low)
#   Avg ∈ [3,6) → B (Medium)
#   Avg ≥ 6.0  → C (High)
_CATEGORY_THRESHOLDS = (3.0, 6.0)

# Environmental Risk Factor (ERF) per category (multiplier on market value)
_CATEGORY_ERF = {
    "A": 1.00,  # No discount — well-managed
    "B": 0.95,  # 5% discount — moderate residual risk
    "C": 0.85,  # 15% discount — significant residual risk
}

# Proximity sensitivity thresholds in metres
#   < 500m  : very high sensitivity
#   500–1500: high
#   1500–3000: medium
#   > 3000  : low
_PROX_THRESHOLDS_M = (500.0, 1500.0, 3000.0)


# ════════════════════════════════════════════════════════════════════════════
#  HELPERS
# ════════════════════════════════════════════════════════════════════════════

def _norm_level(v: Any) -> int:
    """Normalize a low/med/high value (Arabic or English or numeric) to 1..3."""
    if v is None:
        return 1
    if isinstance(v, (int, float)):
        return max(1, min(3, int(round(v))))
    key = str(v).strip().lower()
    return _LEVEL_TO_SCORE.get(key, 1)


def _proximity_sensitivity(distance_m: Optional[float]) -> Tuple[int, str]:
    """Given a distance to a sensitive receptor, return (score 1-3, label)."""
    if distance_m is None or distance_m <= 0:
        return 3, "غير محدد — يُفترض حساسية عالية"
    if distance_m < _PROX_THRESHOLDS_M[0]:
        return 3, "حرج — أقل من 500م"
    if distance_m < _PROX_THRESHOLDS_M[1]:
        return 3, "عالٍ — 500م إلى 1.5كم"
    if distance_m < _PROX_THRESHOLDS_M[2]:
        return 2, "متوسط — 1.5 إلى 3كم"
    return 1, "منخفض — أكثر من 3كم"


def _classify(avg_score: float) -> Tuple[str, str]:
    """Map an average 1–9 impact score to ESF category A/B/C."""
    lo, hi = _CATEGORY_THRESHOLDS
    if avg_score < lo:
        return "A", "فئة (أ) — منخفض الأثر"
    if avg_score < hi:
        return "B", "فئة (ب) — متوسط الأثر"
    return "C", "فئة (ج) — مرتفع الأثر"


# ════════════════════════════════════════════════════════════════════════════
#  CORE EIA ASSESSMENT
# ════════════════════════════════════════════════════════════════════════════

def run_eia_assessment(payload: Dict[str, Any]) -> Dict[str, Any]:
    """
    Runs a full EIA assessment from a structured payload.

    Required-ish payload structure:
        project: {
            activity_type, location, gis_lat, gis_lon, area_m2,
            production_capacity, raw_materials,
            water_consumption_m3_yr, energy_consumption_kwh_yr, fuel_consumption_l_yr
        }
        site_analysis: {
            distance_to_residential_m, distance_to_water_m, distance_to_infrastructure_m
        }
        baseline: {
            air_quality, water_quality, soil_quality, noise_level_db, social_status
        }
        emissions: {
            air_emissions_tons_yr, industrial_wastewater_m3_yr,
            ordinary_waste_tons_yr, hazardous_waste_tons_yr
        }
        impact_matrix: list[{activity, element, severity, duration, probability}]
        mitigation: list[str]
        monitoring: {frequency, kpis: list[str], reporting_cycle}
        market_value: float                   (optional — for investment linkage)

    Returns dict with all 10 sections + classification + ERF + adjusted value.
    """
    if not isinstance(payload, dict):
        raise ValueError("payload يجب أن يكون كائناً.")

    project   = payload.get("project") or {}
    site      = payload.get("site_analysis") or {}
    baseline  = payload.get("baseline") or {}
    emissions = payload.get("emissions") or {}
    matrix    = list(payload.get("impact_matrix") or [])
    mitigation= list(payload.get("mitigation") or [])
    monitoring= payload.get("monitoring") or {}

    if not matrix:
        raise ValueError("impact_matrix فارغة — يجب توفير سجل واحد على الأقل.")

    # ── 1. PROJECT DESCRIPTION ───────────────────────────────────────────────
    project_summary = {
        "activity_type":         project.get("activity_type", "—"),
        "location":              project.get("location", "—"),
        "gis_coords":            (f"{project.get('gis_lat')}, {project.get('gis_lon')}"
                                   if project.get("gis_lat") and project.get("gis_lon") else "—"),
        "area_m2":               float(project.get("area_m2", 0) or 0),
        "production_capacity":   project.get("production_capacity", "—"),
        "raw_materials":         project.get("raw_materials", "—"),
        "water_m3_yr":           float(project.get("water_consumption_m3_yr", 0) or 0),
        "energy_kwh_yr":         float(project.get("energy_consumption_kwh_yr", 0) or 0),
        "fuel_l_yr":             float(project.get("fuel_consumption_l_yr", 0) or 0),
    }

    # ── 2. SITE ANALYSIS — proximity sensitivity ─────────────────────────────
    d_res   = site.get("distance_to_residential_m")
    d_water = site.get("distance_to_water_m")
    d_infra = site.get("distance_to_infrastructure_m")
    res_score,   res_label   = _proximity_sensitivity(float(d_res)   if d_res   is not None else None)
    water_score, water_label = _proximity_sensitivity(float(d_water) if d_water is not None else None)
    infra_score, infra_label = _proximity_sensitivity(float(d_infra) if d_infra is not None else None)
    site_sensitivity_avg = (res_score + water_score + infra_score) / 3.0

    site_analysis = {
        "distance_to_residential_m": d_res,
        "residential_score":         res_score,
        "residential_label":         res_label,
        "distance_to_water_m":       d_water,
        "water_score":               water_score,
        "water_label":               water_label,
        "distance_to_infrastructure_m": d_infra,
        "infra_score":               infra_score,
        "infra_label":               infra_label,
        "site_sensitivity_avg":      round(site_sensitivity_avg, 2),
    }

    # ── 3. ENVIRONMENTAL BASELINE ────────────────────────────────────────────
    env_baseline = {
        "air_quality":     baseline.get("air_quality", "—"),
        "water_quality":   baseline.get("water_quality", "—"),
        "soil_quality":    baseline.get("soil_quality", "—"),
        "noise_level_db":  baseline.get("noise_level_db", "—"),
        "social_status":   baseline.get("social_status", "—"),
    }

    # ── 4. EMISSIONS & WASTE ─────────────────────────────────────────────────
    air      = float(emissions.get("air_emissions_tons_yr", 0) or 0)
    waste_w  = float(emissions.get("industrial_wastewater_m3_yr", 0) or 0)
    ord_w    = float(emissions.get("ordinary_waste_tons_yr", 0) or 0)
    haz_w    = float(emissions.get("hazardous_waste_tons_yr", 0) or 0)
    total_waste_intensity = (air * 1.5) + (waste_w * 0.001) + (ord_w * 0.5) + (haz_w * 3.0)
    emissions_block = {
        "air_emissions_tons_yr":         air,
        "industrial_wastewater_m3_yr":   waste_w,
        "ordinary_waste_tons_yr":        ord_w,
        "hazardous_waste_tons_yr":       haz_w,
        "total_waste_intensity_index":   round(total_waste_intensity, 2),
    }

    # ── 5. IMPACT MATRIX — Severity × Probability ────────────────────────────
    matrix_evaluated: List[Dict[str, Any]] = []
    critical_impacts: List[str] = []
    total_score = 0
    for row in matrix:
        sev_s = _norm_level(row.get("severity"))
        prob_s = _norm_level(row.get("probability"))
        dur_s = _norm_level(row.get("duration"))
        score = sev_s * prob_s
        is_critical = (score >= 6) or (sev_s == 3 and dur_s >= 2)
        verdict = ("حرج — يستوجب تدخلاً فورياً" if is_critical
                   else "متوسط — تخفيف موصى به" if score >= 3
                   else "مقبول — رصد دوري كافٍ")
        evaluated = {
            "activity":          row.get("activity", "—"),
            "element":           row.get("element", "—"),
            "severity":          _SCORE_TO_LABEL_AR[sev_s],
            "severity_score":    sev_s,
            "duration":          ("طويل المدى" if dur_s >= 2 else "قصير المدى"),
            "duration_score":    dur_s,
            "probability":       _SCORE_TO_LABEL_AR[prob_s],
            "probability_score": prob_s,
            "impact_score":      score,
            "is_critical":       is_critical,
            "verdict":           verdict,
        }
        matrix_evaluated.append(evaluated)
        total_score += score
        if is_critical:
            critical_impacts.append(f"{evaluated['activity']} → {evaluated['element']}")

    avg_impact_score = total_score / max(len(matrix_evaluated), 1)

    # ── 6. IMPACT ASSESSMENT (aggregate) ─────────────────────────────────────
    short_term_impacts = [e for e in matrix_evaluated if e["duration_score"] == 1]
    long_term_impacts  = [e for e in matrix_evaluated if e["duration_score"] >= 2]
    impact_assessment = {
        "avg_impact_score":    round(avg_impact_score, 2),
        "max_impact_score":    max((e["impact_score"] for e in matrix_evaluated), default=0),
        "n_critical":          len(critical_impacts),
        "critical_impacts":    critical_impacts,
        "n_short_term":        len(short_term_impacts),
        "n_long_term":         len(long_term_impacts),
    }

    # ── 7. MITIGATION & COMPLIANCE ───────────────────────────────────────────
    mitigation_block = {
        "measures":         mitigation if mitigation else ["لم يتم توفير تدابير تخفيف — يُوصى بإلزامية وضع EMP"],
        "iso_14001":        bool(payload.get("iso_14001_compliant", False)),
        "egyptian_law":     "قانون البيئة المصري 4/1994 (المعدَّل بالقانون 105/2015)",
        "international":    "World Bank ESF Categorization + IFC Performance Standards",
    }

    # ── 8. MONITORING ────────────────────────────────────────────────────────
    monitoring_block = {
        "frequency":       monitoring.get("frequency", "ربع سنوي"),
        "kpis":            list(monitoring.get("kpis") or [
                                 "تركيز PM₁₀ في الهواء",
                                 "تركيز COD/BOD₅ في صرف الصناعي",
                                 "مستوى الضوضاء dB(A)",
                                 "كفاءة الفلاتر / المعالجة",
                                 "كميات النفايات الخطرة المُولَّدة",
                              ]),
        "reporting_cycle": monitoring.get("reporting_cycle", "تقرير ربع سنوي + تقرير سنوي شامل"),
    }

    # ── 9. ENVIRONMENTAL CLASSIFICATION ──────────────────────────────────────
    # Classification considers BOTH impact matrix AND site sensitivity.
    composite_score = (avg_impact_score * 0.7) + (site_sensitivity_avg * 1.5 * 0.3)
    category, category_label = _classify(composite_score)

    # Bump category up if any single critical impact and site is sensitive
    if critical_impacts and site_sensitivity_avg >= 2.5 and category == "A":
        category, category_label = "B", "فئة (ب) — متوسط الأثر (مرفوع بسبب حساسية الموقع)"
    if len(critical_impacts) >= 3 and category == "B":
        category, category_label = "C", "فئة (ج) — مرتفع الأثر (مرفوع بسبب تعدد الآثار الحرجة)"

    classification = {
        "category":          category,
        "category_label":    category_label,
        "composite_score":   round(composite_score, 2),
        "rationale":         (f"متوسط درجة الأثر = {avg_impact_score:.2f}، "
                              f"حساسية الموقع = {site_sensitivity_avg:.2f}، "
                              f"عدد الآثار الحرجة = {len(critical_impacts)}"),
    }

    # ── 10. INVESTMENT LINKAGE — Environmental Risk Factor (ERF) ─────────────
    erf = _CATEGORY_ERF[category]
    market_value = float(payload.get("market_value", 0) or 0)
    adjusted_value = market_value * erf if market_value > 0 else 0.0
    operational_impact = {
        "A": "تشغيل طبيعي بدون قيود إضافية",
        "B": "تشغيل مع متطلبات تخفيف منتظمة، تكلفة امتثال 1–3% من OPEX",
        "C": "تشغيل تحت ضغط رقابي، تكلفة امتثال 5–10% من OPEX، احتمال إيقاف",
    }[category]

    investment_linkage = {
        "market_value":            market_value,
        "environmental_risk_factor": erf,
        "erf_pct":                 round((1 - erf) * 100, 2),  # نسبة الخصم البيئي
        "adjusted_market_value":   round(adjusted_value, 0),
        "value_at_risk":           round(market_value - adjusted_value, 0),
        "operational_impact":      operational_impact,
    }

    # ── ASSEMBLE ─────────────────────────────────────────────────────────────
    return {
        "logic":             "EIA",
        "valuation_date":    payload.get("valuation_date") or datetime.now().strftime("%Y-%m-%d"),
        "standards":         "ISO 14001:2015 / Law 4-1994 / World Bank ESF / IFC PS",
        "project":           project_summary,
        "site_analysis":     site_analysis,
        "baseline":          env_baseline,
        "emissions":         emissions_block,
        "impact_matrix":     matrix_evaluated,
        "impact_assessment": impact_assessment,
        "mitigation":        mitigation_block,
        "monitoring":        monitoring_block,
        "classification":    classification,
        "investment_linkage": investment_linkage,
    }


# ════════════════════════════════════════════════════════════════════════════
#  WORD REPORT
# ════════════════════════════════════════════════════════════════════════════

def write_eia_word_report(result: Dict[str, Any], output_path: str) -> bool:
    """Generates the full 10-section EIA report as a Word document."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except Exception:
        return False

    NAVY  = RGBColor(0x1F, 0x38, 0x64)
    GOLD  = RGBColor(0xC9, 0xA2, 0x27)
    GREEN = RGBColor(0x1E, 0x84, 0x49)
    RED   = RGBColor(0xC0, 0x39, 0x2B)
    GREY  = RGBColor(0x55, 0x55, 0x55)
    AMBER = RGBColor(0xD6, 0x8B, 0x18)

    def _para_rtl(p):
        pPr = p._p.get_or_add_pPr()
        bidi = OxmlElement("w:bidi"); bidi.set(qn("w:val"), "1"); pPr.append(bidi)

    def _h1(doc, text, color=None):
        p = doc.add_paragraph(); _para_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); r.font.size = Pt(20); r.font.bold = True
        r.font.color.rgb = color or NAVY

    def _h2(doc, text, color=None):
        p = doc.add_paragraph(); _para_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(text); r.font.size = Pt(14); r.font.bold = True
        r.font.color.rgb = color or GOLD

    def _para(doc, text, *, bold=False, color=None, size=12):
        p = doc.add_paragraph(); _para_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(text); r.font.size = Pt(size); r.font.bold = bold
        if color is not None: r.font.color.rgb = color

    def _disclosure(doc, text):
        p = doc.add_paragraph(); _para_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(text); r.font.italic = True; r.font.size = Pt(10); r.font.color.rgb = GREY

    def _kv(doc, rows):
        if not rows: return
        tbl = doc.add_table(rows=len(rows), cols=2); tbl.style = "Light Grid Accent 1"
        bidi = OxmlElement("w:bidiVisual"); tbl._tbl.tblPr.append(bidi)
        for i, (k, v) in enumerate(rows):
            rc = tbl.rows[i].cells; rc[0].text = str(k); rc[1].text = str(v)
            for c in rc:
                for p in c.paragraphs:
                    _para_rtl(p); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for r in p.runs: r.font.name = "Arial"; r.font.size = Pt(11)
            rc[0].paragraphs[0].runs[0].font.bold = True

    def _grid(doc, headers, rows, *, highlight_critical=False):
        tbl = doc.add_table(rows=1+len(rows), cols=len(headers)); tbl.style = "Light Grid Accent 1"
        bidi = OxmlElement("w:bidiVisual"); tbl._tbl.tblPr.append(bidi)
        for j, h in enumerate(headers):
            tbl.rows[0].cells[j].text = h
            for p in tbl.rows[0].cells[j].paragraphs:
                _para_rtl(p); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: r.font.bold = True; r.font.color.rgb = NAVY; r.font.size = Pt(11)
        for i, row in enumerate(rows):
            crit = highlight_critical and len(row) > 0 and "حرج" in str(row[-1])
            for j, val in enumerate(row):
                cell = tbl.rows[i+1].cells[j]
                cell.text = str(val)
                for p in cell.paragraphs:
                    _para_rtl(p); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for r in p.runs:
                        r.font.size = Pt(10)
                        if crit:
                            r.font.bold = True; r.font.color.rgb = RED

    # ── BUILD DOCUMENT ────────────────────────────────────────────────────────
    doc = Document()
    style = doc.styles["Normal"]; style.font.name = "Arial"; style.font.size = Pt(12)
    rpr = style.element.get_or_add_rPr()
    rpr.append(OxmlElement("w:bidi")); rpr.append(OxmlElement("w:rtl"))

    _h1(doc, "تقرير تقييم الأثر البيئي (EIA)")
    _h2(doc, f"المعيار المُطبَّق: {result.get('standards', '—')}", color=GOLD)
    _para(doc, f"تاريخ التقييم: {result.get('valuation_date', '—')}", bold=True, color=NAVY)
    doc.add_paragraph()

    # 1. وصف المشروع
    _h2(doc, "1. وصف المشروع")
    p = result.get("project", {})
    _kv(doc, [
        ("نوع النشاط",                       p.get("activity_type", "—")),
        ("الموقع",                           p.get("location", "—")),
        ("إحداثيات GIS",                     p.get("gis_coords", "—")),
        ("المساحة (م²)",                     f"{p.get('area_m2', 0):,.0f}"),
        ("الطاقة الإنتاجية",                 p.get("production_capacity", "—")),
        ("المواد الخام",                     p.get("raw_materials", "—")),
        ("استهلاك المياه (م³/سنة)",          f"{p.get('water_m3_yr', 0):,.0f}"),
        ("استهلاك الطاقة (ك.و.س/سنة)",       f"{p.get('energy_kwh_yr', 0):,.0f}"),
        ("استهلاك الوقود (لتر/سنة)",         f"{p.get('fuel_l_yr', 0):,.0f}"),
    ])

    # 2. تحليل الموقع GIS
    _h2(doc, "2. تحليل الموقع (GIS)")
    s = result.get("site_analysis", {})
    _kv(doc, [
        ("المسافة إلى أقرب منطقة سكنية (م)",         s.get("distance_to_residential_m", "—")),
        ("حساسية القرب السكني",                       s.get("residential_label", "—")),
        ("المسافة إلى أقرب مورد مائي (م)",           s.get("distance_to_water_m", "—")),
        ("حساسية القرب المائي",                       s.get("water_label", "—")),
        ("المسافة إلى البنية التحتية الحساسة (م)",   s.get("distance_to_infrastructure_m", "—")),
        ("حساسية القرب من البنية التحتية",            s.get("infra_label", "—")),
        ("متوسط حساسية الموقع (1–3)",                s.get("site_sensitivity_avg", "—")),
    ])

    # 3. خط الأساس البيئي
    _h2(doc, "3. خط الأساس البيئي")
    b = result.get("baseline", {})
    _kv(doc, [
        ("جودة الهواء",      b.get("air_quality", "—")),
        ("جودة المياه",      b.get("water_quality", "—")),
        ("جودة التربة",       b.get("soil_quality", "—")),
        ("مستوى الضوضاء (dB)", b.get("noise_level_db", "—")),
        ("الوضع الاجتماعي",   b.get("social_status", "—")),
    ])

    # 4. الانبعاثات والمخلفات
    _h2(doc, "4. الانبعاثات والمخلفات")
    e = result.get("emissions", {})
    _kv(doc, [
        ("الانبعاثات الهوائية (طن/سنة)",    f"{e.get('air_emissions_tons_yr', 0):,.2f}"),
        ("الصرف الصناعي (م³/سنة)",          f"{e.get('industrial_wastewater_m3_yr', 0):,.0f}"),
        ("النفايات العادية (طن/سنة)",        f"{e.get('ordinary_waste_tons_yr', 0):,.2f}"),
        ("النفايات الخطرة (طن/سنة)",         f"{e.get('hazardous_waste_tons_yr', 0):,.2f}"),
        ("مؤشر شدة المخلفات",                f"{e.get('total_waste_intensity_index', 0):,.2f}"),
    ])

    # 5. IMPACT MATRIX
    _h2(doc, "5. مصفوفة الآثار (Impact Matrix)")
    matrix = result.get("impact_matrix", [])
    _grid(doc,
          headers=["النشاط", "العنصر البيئي", "الشدة", "المدة", "الاحتمالية", "Impact Score", "التقييم"],
          rows=[[m["activity"], m["element"], m["severity"], m["duration"],
                 m["probability"], m["impact_score"], m["verdict"]] for m in matrix],
          highlight_critical=True)
    _para(doc, "Impact Score = Severity × Probability  |  حرج إذا Score ≥ 6 أو شدة عالية + مدة طويلة",
          bold=True, color=NAVY)

    # 6. تقييم الأثر
    _h2(doc, "6. تقييم الأثر الكلي")
    ia = result.get("impact_assessment", {})
    _kv(doc, [
        ("متوسط درجة الأثر",         f"{ia.get('avg_impact_score', 0):.2f}"),
        ("أعلى درجة أثر",            ia.get("max_impact_score", 0)),
        ("عدد الآثار قصيرة المدى",   ia.get("n_short_term", 0)),
        ("عدد الآثار طويلة المدى",   ia.get("n_long_term", 0)),
        ("عدد الآثار الحرجة",        ia.get("n_critical", 0)),
    ])
    if ia.get("critical_impacts"):
        _para(doc, "الآثار الحرجة المُحدَّدة:", bold=True, color=RED)
        for c in ia["critical_impacts"]:
            _para(doc, f"• {c}", color=RED)

    # 7. التخفيف والالتزام
    _h2(doc, "7. التخفيف والالتزام")
    m = result.get("mitigation", {})
    _para(doc, f"الالتزام بـ ISO 14001:2015 — {'مطابق' if m.get('iso_14001') else 'غير مطابق / يستوجب اعتماد'}",
          bold=True, color=GREEN if m.get('iso_14001') else AMBER)
    _para(doc, "الإطار القانوني المحلي: " + m.get("egyptian_law", "—"))
    _para(doc, "الإطار الدولي: " + m.get("international", "—"))
    _para(doc, "تدابير التخفيف المُوصى بها:", bold=True)
    for measure in m.get("measures", []):
        _para(doc, f"• {measure}")

    # 8. الرصد والمتابعة
    _h2(doc, "8. الرصد والمتابعة")
    mon = result.get("monitoring", {})
    _kv(doc, [
        ("تكرار الرصد",          mon.get("frequency", "—")),
        ("دورة التقارير",        mon.get("reporting_cycle", "—")),
    ])
    _para(doc, "مؤشرات الأداء البيئية الرئيسية (KPIs):", bold=True)
    for kpi in mon.get("kpis", []):
        _para(doc, f"• {kpi}")

    # 9. التصنيف البيئي
    _h2(doc, "9. التصنيف البيئي")
    cl = result.get("classification", {})
    cat = cl.get("category", "—")
    cat_color = GREEN if cat == "A" else (AMBER if cat == "B" else RED)
    _para(doc, f"التصنيف: {cl.get('category_label', '—')}", bold=True, color=cat_color, size=14)
    _kv(doc, [
        ("درجة الأثر المركَّبة",   f"{cl.get('composite_score', 0):.2f}"),
        ("مبرر التصنيف",          cl.get("rationale", "—")),
    ])
    _para(doc, "تفسير الفئات:", bold=True)
    _para(doc, "• فئة (أ) منخفض: متوسط درجة الأثر < 3.0 — قبول مباشر مع رصد دوري.")
    _para(doc, "• فئة (ب) متوسط: 3.0 ≤ درجة الأثر < 6.0 — يستوجب EMP وتدابير تخفيف موثقة.")
    _para(doc, "• فئة (ج) مرتفع: درجة الأثر ≥ 6.0 أو ≥3 آثار حرجة — يستوجب EIA كامل ومراجعة دورية.")

    # 10. الربط الاستثماري
    _h2(doc, "10. الربط الاستثماري — Environmental Risk Factor")
    inv = result.get("investment_linkage", {})
    _kv(doc, [
        ("القيمة السوقية قبل التعديل البيئي",    f"{inv.get('market_value', 0):,.0f} ج.م"),
        ("Environmental Risk Factor (ERF)",      f"{inv.get('environmental_risk_factor', 0):.4f}"),
        ("نسبة الخصم البيئي",                   f"{inv.get('erf_pct', 0):.2f}%"),
        ("القيمة بعد التعديل البيئي",            f"{inv.get('adjusted_market_value', 0):,.0f} ج.م"),
        ("القيمة المعرَّضة للمخاطر البيئية",      f"{inv.get('value_at_risk', 0):,.0f} ج.م"),
        ("الأثر التشغيلي المتوقع",               inv.get("operational_impact", "—")),
    ])

    _disclosure(doc,
        "هذا التقرير وفق ISO 14001:2015 وقانون البيئة المصري 4/1994 (المعدَّل بالقانون 105/2015) "
        "وإطار البنك الدولي ESF ومعايير IFC Performance Standards. ERF يُطبَّق كمعامل ضربي على "
        "القيمة السوقية لتعكس المخاطر البيئية الكامنة، ويُستخدم في التقييم الاستثماري لتحويل "
        "الآثار البيئية إلى قيمة كمية واضحة.")

    doc.save(output_path)
    return True


# ════════════════════════════════════════════════════════════════════════════
#  EXCEL REPORT (companion)
# ════════════════════════════════════════════════════════════════════════════

def write_eia_excel_report(result: Dict[str, Any], output_dir: str = "") -> Optional[str]:
    """Renders an RTL Arabic Excel companion to the EIA Word report."""
    try:
        import xlsxwriter  # type: ignore
    except Exception:
        return None

    out_dir = output_dir.strip() if output_dir.strip() else _DEF_OUT
    os.makedirs(out_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filepath = os.path.join(out_dir, f"eia_report_{timestamp}.xlsx")

    NAVY, GOLD, WHITE, LIGHT, GREEN, RED, AMBER, GREY = (
        "#1F3864", "#C9A227", "#FFFFFF", "#EAF0FB", "#1E8449",
        "#C0392B", "#D68B18", "#F2F2F2"
    )

    wb = xlsxwriter.Workbook(filepath)

    title_fmt = wb.add_format({"bold": True, "font_size": 16, "bg_color": NAVY,
                                "font_color": WHITE, "align": "center", "valign": "vcenter"})
    section_fmt = wb.add_format({"bold": True, "font_size": 12, "bg_color": GOLD,
                                  "font_color": NAVY, "align": "right"})
    label_fmt = wb.add_format({"bold": True, "bg_color": LIGHT, "align": "right", "border": 1})
    val_fmt   = wb.add_format({"align": "right", "border": 1})
    num_fmt   = wb.add_format({"num_format": "#,##0.00", "align": "right", "border": 1})
    pct_fmt   = wb.add_format({"num_format": "0.00%", "align": "right", "border": 1})
    head_fmt  = wb.add_format({"bold": True, "bg_color": NAVY, "font_color": WHITE,
                                "align": "center", "border": 1, "text_wrap": True})
    crit_fmt  = wb.add_format({"bg_color": RED, "font_color": WHITE, "align": "right", "border": 1})
    cat_a_fmt = wb.add_format({"bg_color": GREEN, "font_color": WHITE, "align": "center",
                                "border": 1, "bold": True, "font_size": 14})
    cat_b_fmt = wb.add_format({"bg_color": AMBER, "font_color": WHITE, "align": "center",
                                "border": 1, "bold": True, "font_size": 14})
    cat_c_fmt = wb.add_format({"bg_color": RED, "font_color": WHITE, "align": "center",
                                "border": 1, "bold": True, "font_size": 14})

    # Sheet 1: Summary
    ws1 = wb.add_worksheet("ملخص EIA")
    ws1.right_to_left()
    ws1.set_column(0, 0, 32); ws1.set_column(1, 1, 28)
    ws1.merge_range(0, 0, 0, 1, "تقرير تقييم الأثر البيئي (EIA)", title_fmt)
    ws1.set_row(0, 32)

    cl  = result.get("classification", {})
    inv = result.get("investment_linkage", {})
    ia  = result.get("impact_assessment", {})

    ws1.write(2, 0, "التصنيف البيئي", section_fmt)
    cat = cl.get("category", "A")
    cat_format = {"A": cat_a_fmt, "B": cat_b_fmt, "C": cat_c_fmt}.get(cat, val_fmt)
    ws1.merge_range(3, 0, 3, 1, cl.get("category_label", "—"), cat_format)
    ws1.set_row(3, 28)

    r = 5
    ws1.write(r, 0, "متوسط درجة الأثر", label_fmt); ws1.write_number(r, 1, ia.get("avg_impact_score", 0), num_fmt); r += 1
    ws1.write(r, 0, "عدد الآثار الحرجة", label_fmt); ws1.write(r, 1, ia.get("n_critical", 0), val_fmt); r += 1
    ws1.write(r, 0, "Environmental Risk Factor", label_fmt); ws1.write_number(r, 1, inv.get("environmental_risk_factor", 0), num_fmt); r += 1
    ws1.write(r, 0, "نسبة الخصم البيئي", label_fmt); ws1.write_number(r, 1, inv.get("erf_pct", 0) / 100.0, pct_fmt); r += 1
    ws1.write(r, 0, "القيمة قبل التعديل (ج.م)", label_fmt); ws1.write_number(r, 1, inv.get("market_value", 0), num_fmt); r += 1
    ws1.write(r, 0, "القيمة بعد التعديل (ج.م)", label_fmt); ws1.write_number(r, 1, inv.get("adjusted_market_value", 0), num_fmt); r += 1
    ws1.write(r, 0, "القيمة المعرَّضة للمخاطر (ج.م)", label_fmt); ws1.write_number(r, 1, inv.get("value_at_risk", 0), num_fmt); r += 1

    # Sheet 2: Impact Matrix
    ws2 = wb.add_worksheet("Impact Matrix")
    ws2.right_to_left()
    ws2.set_column(0, 6, 18)
    headers = ["النشاط", "العنصر البيئي", "الشدة", "المدة", "الاحتمالية", "Score", "التقييم"]
    for j, h in enumerate(headers):
        ws2.write(0, j, h, head_fmt)
    ws2.set_row(0, 36)
    for i, m in enumerate(result.get("impact_matrix", []), start=1):
        is_crit = m.get("is_critical", False)
        row_fmt = crit_fmt if is_crit else val_fmt
        ws2.write(i, 0, m["activity"], row_fmt)
        ws2.write(i, 1, m["element"], row_fmt)
        ws2.write(i, 2, m["severity"], row_fmt)
        ws2.write(i, 3, m["duration"], row_fmt)
        ws2.write(i, 4, m["probability"], row_fmt)
        ws2.write_number(i, 5, m["impact_score"], row_fmt)
        ws2.write(i, 6, m["verdict"], row_fmt)

    wb.close()
    return filepath
