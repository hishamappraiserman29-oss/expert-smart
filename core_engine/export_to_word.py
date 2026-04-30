import os
from datetime import datetime

def generate_full_valuation_report_docx(data: dict) -> str:
    """
    يقوم بإنشاء تقرير تقييم عقاري رسمي ورصين بصيغة Word.
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError:
        print("❌ مكتبة python-docx غير مثبتة. لا يمكن تصدير ملف Word.")
        return ""

    doc = Document()
    
    # تحديد المتغيرات حسب المنطقة
    region = data.get("region", "SA")
    currency = "ريال سعودي" if region == "SA" else "جنيه مصري"
    standard_name = "الهيئة السعودية للمقيمين المعتمدين (تقييم)" if region == "SA" else "المعايير المصرية للتقييم العقاري"
    
    area = data.get("area_sqm", 0)
    price = data.get("price_per_meter", 0)
    total = data.get("total_value", 0)
    location = data.get("location", "موقع غير محدد")
    query = data.get("query", "")
    trend = data.get("market_trend", "مستقر")
    
    # تنسيق للغة رسمية
    doc.add_heading('تقرير تقييم عقاري احترافي', 0)
    
    # 1. المقدمة
    doc.add_heading('أولاً: مقدّمة التقييم والغرض منه', level=1)
    doc.add_paragraph(
        f"تم إعداد هذا التقرير بتاريخ {datetime.now().strftime('%Y-%m-%d')} بناءً على طلبكم لتقدير القيمة السوقية العادلة للعقار. "
        f"يهدف هذا التقييم إلى تقديم رأي مهني ومستقل يحاكي ظروف السوق الحالية لدعم قراراتكم، "
        f"بناءً على طلبكم المرجعي المُستلَم: '{query}'."
    )
    
    # 2. وصف العقار
    doc.add_heading('ثانياً: وصف العقار وتحليل الموقع', level=1)
    doc.add_paragraph(
        f"يقع العقار محل التقييم في إطار منطقة البحث ({location}). "
        f"يبلغ إجمالي المساحة التقديرية للعقار ({area}) متر مربع. "
        f"وتشير البيانات التحليلية المستخرجة من قراءات السوق الخاصة بنا لتلك المنطقة بصفة عامة، إلى أن اتجاه الأسعار في المرحلة الحالية يعتبر ({trend})."
    )
    
    # 3. المنهجية
    doc.add_heading('ثالثاً: المنهجية الحسابية وأسس التقييم', level=1)
    doc.add_paragraph(
        f"تم تطبيق المنهجيات العلمية المحددة في أدلة وإرشادات {standard_name}. "
        f"حيث اُعتمدت منهجية أسلوب المقارنة والمقاربات السوقية في تطوير القيمة، وذلك بربط العوامل والخصائص الفنية للعقار محل التقييم مع المعروض في السوق. "
        f"وبناءً عليه، تم استنتاج أن متوسط السعر التقديري للمتر المربع يبلغ نحو {price:,.2f} {currency}."
    )
    
    # 4. الخلاصة
    doc.add_heading('رابعاً: الخلاصة الإجمالية والتوصيات', level=1)
    doc.add_paragraph(
        f"استناداً إلى إجراءات التقييم المطبقة، والرأي المهني المبني على دراسة وقراءة ظروف السوق الحالية، "
        f"فإن القيمة السوقية التقديرية الإجمالية للعقار المُقيَّم تبلغ:\n"
        f"*** ({total:,.2f} {currency}) ***.\n\n"
        f"التوصيات: يُوصى الخبير بإجراء مطابقة الفحص الفني والعناية الواجبة (Due Diligence) على مستندات الملكية قبل اتخاذ أي قرار تنفيذي للاستثمار، لضمان صحة وخلو العقار من أي التزامات قانونية تعوق الاستخدام المعتزم."
    )
    
    # المحاذاة لليمين ودعم الخطوط للغة العربية
    for p in doc.paragraphs:
        try:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        except:
            pass

    # حفظ المستند
    output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'outputs', 'reports')
    os.makedirs(output_dir, exist_ok=True)
    report_id = datetime.now().strftime('%Y%m%d_%H%M%S')
    file_path = os.path.join(output_dir, f'Valuation_Report_RAG_{report_id}.docx')
    
    doc.save(file_path)
    return file_path
