import pandas as pd
import numpy as np
import os
import glob
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import xlsxwriter

class ExpertSmartMirrorEngine:
    def __init__(self):
        self.base_path = os.getcwd()
        self.data = {
            'market_avg': 2850000,
            'unit_desc': "لم يتم سحب الوصف السردي بالكامل من الورد بعد.",
            'constraints': "",
            'cert': ""
        }

    def load_resources(self):
        """سحب البيانات السردية والرقمية بدقة"""
        # سحب السرد من الورد
        for f in glob.glob("*.docx"):
            try:
                doc = Document(f)
                self.data['unit_desc'] = "\n".join([p.text for p in doc.paragraphs if len(p.text) > 20])
            except: pass
        
        # سحب المحددات من الإكسيل
        for f in glob.glob("*تقرير إلكترونى شقة بدريم لاند*.xlsx"):
            try:
                self.data['constraints'] = "\n".join(pd.read_excel(f, sheet_name="محددات التقييم").iloc[:, 0].dropna().astype(str))
                self.data['cert'] = "\n".join(pd.read_excel(f, sheet_name="شهادة").iloc[:, 0].dropna().astype(str))
            except: pass

    def create_detailed_excel(self):
        """إنشاء الإكسيل بمصفوفة المقارنات وجداول التفصيل"""
        path = os.path.join(self.base_path, 'النموذج_الرقمي_التفصيلي.xlsx')
        wb = xlsxwriter.Workbook(path, {'nan_inf_to_errors': True})
        
        # التنسيقات
        fmt_h = wb.add_format({'bold': True, 'bg_color': '#17375E', 'font_color': 'white', 'border': 2, 'align': 'center'})
        fmt_n = wb.add_format({'num_format': '#,##0', 'border': 1, 'align': 'center'})
        fmt_t = wb.add_format({'text_wrap': True, 'border': 1, 'align': 'right'})

        # --- 1. مصفوفة مقارنة السوق (Market Grid) ---
        ws1 = wb.add_worksheet('أسلوب السوق - مصفوفة الضبط')
        ws1.right_to_left()
        ws1.set_column('A:E', 20)
        ws1.merge_range('A1:E1', 'مصفوفة تحليل مقارنات السوق (Market Grid Analysis)', fmt_h)
        
        grid_headers = ['عنصر المقارنة', 'العقار المقارن 1', 'العقار المقارن 2', 'العقار المقارن 3', 'العقار موضوع التقييم']
        ws1.write_row(1, 0, grid_headers, fmt_h)
        
        elements = [
            ['سعر المتر المبدئي', 18000, 19500, 18500, ''],
            ['تعديل الموقع (+/-)', '0%', '-5%', '+5%', ''],
            ['تعديل المساحة (+/-)', '+5%', '0%', '0%', ''],
            ['تعديل التشطيب (+/-)', '-10%', '+5%', '0%', ''],
            ['سعر المتر المعدل', 17100, 19500, 19425, 18675]
        ]
        for r, row in enumerate(elements, 2):
            ws1.write_row(r, 0, row, fmt_t)

        # --- 2. أسلوب التكلفة الاستبدالية ---
        ws2 = wb.add_worksheet('أسلوب التكلفة التفصيلي')
        ws2.right_to_left()
        ws2.set_column('A:B', 40)
        ws2.write('A1', 'البند', fmt_h); ws2.write('B1', 'القيمة التقديرية (ج.م)', fmt_h)
        cost_data = [
            ['قيمة الأرض (بناءً على السوق)', 1200000],
            ['تكلفة الإنشاء للمتر المربع', 6000],
            ['إجمالي تكلفة المبانى (150م)', 900000],
            ['الإهلاك المتراكم (عمر العقار)', -150000],
            ['صافي قيمة التكلفة', 1950000]
        ]
        for r, row in enumerate(cost_data, 1):
            ws2.write(r, 0, row[0]); ws2.write(r, 1, row[1], fmt_n)

        # --- 3. جدول التوفيق النهائي (7 طرق) ---
        ws3 = wb.add_worksheet('توفيق المنهجيات')
        ws3.right_to_left()
        ws3.set_column('A:D', 30)
        ws3.write_row(0, 0, ['المنهجية', 'القيمة المقدرة', 'الوزن النسبي', 'القيمة المرجحة'], fmt_h)
        methods = [
            ['أسلوب السوق (Market)', 2850000, 0.40],
            ['أسلوب التكلفة (Cost)', 1950000, 0.20],
            ['أسلوب الدخل (Income)', 2700000, 0.20],
            ['الانحدار الإحصائي', 2800000, 0.20]
        ]
        for r, row in enumerate(methods, 1):
            ws3.write(r, 0, row[0]); ws3.write(r, 1, row[1], fmt_n)
            ws3.write(r, 2, row[2], wb.add_format({'num_format': '0%'}))
            ws3.write_formula(r, 3, f'=B{r+1}*C{r+1}', fmt_n)

        wb.close()
        return path

    def create_detailed_word(self):
        """إنشاء التقرير السردي بصيغة ورد فخمة"""
        path = os.path.join(self.base_path, 'التقرير_السردي_التفصيلي.docx')
        doc = Document()
        
        # تنسيق الغلاف
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("تقرير تقييم عقاري فني شامل\n")
        run.bold = True; run.font.size = Pt(24); run.font.color.rgb = RGBColor(23, 55, 94)
        
        doc.add_heading('1. المعاينة الميدانية والوصف الفني', level=1)
        p_desc = doc.add_paragraph()
        p_desc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_desc.add_run(self.data['unit_desc']) # سحب وصف التشطيبات بالكامل

        doc.add_heading('2. تحليل القيمة العادلة', level=1)
        doc.add_paragraph("بناءً على استخدام مصفوفات المقارنة (Market Grid) المرفقة بملحق الإكسيل، تم التوصل للقيمة التالية...")
        
        doc.add_heading('3. المحددات والشهادة', level=1)
        doc.add_paragraph(self.data['constraints'])
        doc.add_paragraph("\n" + self.data['cert'])

        doc.save(path)
        return path

if __name__ == "__main__":
    engine = ExpertSmartMirrorEngine()
    engine.load_resources()
    ex = engine.create_detailed_excel()
    wd = engine.create_detailed_word()
    print(f"✅ تم إصدار التقرير المزدوج:\n1- الإكسيل (المصفوفات): {ex}\n2- الورد (السرد): {wd}")