"""
fund_valuation_engine.py
========================
REITs / Investment Fund Valuation Engine
نظام تقييم صناديق الاستثمار العقاري

Standards: IFRS 13 (Fair Value Measurement) & IVS 103 (Valuation Approaches)

Functions:
    run_fund_valuation(payload)            -> dict
    generate_fund_excel(result, out_dir)   -> str   (path to .xlsx)
    generate_fund_word_section(doc, result)         (appends to python-docx Document)
"""

from __future__ import annotations
import os
from datetime import datetime
from typing import Dict

_BASE_DIR = os.path.dirname(os.path.abspath(__file__))
_OUT_DIR = os.path.join(_BASE_DIR, "outputs", "reports")

# ─── Non-real-estate asset types that are out of scope ────────────────────────
_EXCLUDED_TYPES = {"machinery", "business", "equipment", "vehicle", "aircraft"}


# ══════════════════════════════════════════════════════════════════════════════
#  CORE VALUATION
# ══════════════════════════════════════════════════════════════════════════════

def run_fund_valuation(payload: dict) -> dict:
    """
    Runs a full REIT / Investment Fund valuation under IFRS 13 and IVS 103.

    Parameters
    ----------
    payload : dict
        property_type       str
        location            str
        area                float   (sq m)
        market_value        float   (EGP)
        annual_rent         float   (EGP, total annual)
        vacancy_rate        float   (e.g. 0.10 for 10%)
        operating_expenses  float   (EGP, annual)
        loan_amount         float   optional — for NAV
        total_units         int     optional — portfolio size, default 1
        fund_name           str     optional — default "صندوق الاستثمار العقاري"
        ifrs_level          int     optional — 1/2/3, default 3
        client_name         str     optional

    Returns
    -------
    dict with all computed valuation fields.

    Raises
    ------
    ValueError  if property_type is a non-real-estate asset.
    ValueError  if required numeric fields are missing or non-positive.
    """
    # ── Validate asset type ───────────────────────────────────────────────────
    property_type: str = str(payload.get("property_type", "")).strip().lower()
    if property_type in _EXCLUDED_TYPES:
        raise ValueError(
            f"نوع الأصل '{property_type}' خارج نطاق تقييم الصناديق العقارية. "
            f"هذه الوحدة مخصصة للعقارات فقط وفق IFRS 13 / IVS 103. "
            f"الأنواع المستثناة: {sorted(_EXCLUDED_TYPES)}"
        )

    # ── Extract & validate inputs ─────────────────────────────────────────────
    location: str           = str(payload.get("location", "غير محدد")).strip()
    area: float             = float(payload.get("area", 0))
    market_value: float     = float(payload.get("market_value", 0))
    annual_rent: float      = float(payload.get("annual_rent", 0))
    vacancy_rate: float     = float(payload.get("vacancy_rate", 0))
    operating_expenses: float = float(payload.get("operating_expenses", 0))
    loan_amount: float      = float(payload.get("loan_amount", 0))
    total_units: int        = int(payload.get("total_units", 1)) or 1
    fund_name: str          = str(payload.get("fund_name", "صندوق الاستثمار العقاري")).strip()
    ifrs_level: int         = int(payload.get("ifrs_level", 3))
    client_name: str        = str(payload.get("client_name", "")).strip()

    if market_value <= 0:
        raise ValueError("market_value يجب أن يكون أكبر من صفر.")
    if area < 0:
        raise ValueError("area يجب أن تكون قيمة موجبة.")
    if not 0.0 <= vacancy_rate <= 1.0:
        raise ValueError("vacancy_rate يجب أن يكون بين 0 و 1.")
    if ifrs_level not in (1, 2, 3):
        raise ValueError("ifrs_level يجب أن يكون 1 أو 2 أو 3.")

    # ── 1. NOI & Cap Rate ─────────────────────────────────────────────────────
    gross_income: float     = annual_rent
    vacancy_loss: float     = gross_income * vacancy_rate
    effective_income: float = gross_income - vacancy_loss
    noi: float              = effective_income - operating_expenses
    cap_rate: float         = noi / market_value          # decimal
    cap_rate_pct: float     = cap_rate * 100

    # ── 2. IFRS 13 Fair Value Classification ──────────────────────────────────
    ifrs_level_labels = {
        1: "Level 1 — أسعار معلنة في أسواق نشطة (لا تنطبق على العقار المباشر)",
        2: "Level 2 — مدخلات قابلة للملاحظة (معاملات مقارنة)",
        3: "Level 3 — مدخلات غير قابلة للملاحظة (نهج الدخل / DCF)",
    }

    if ifrs_level == 3:
        fair_value_method = "Income Capitalisation (IVS 103 / IFRS 13 Level 3)"
    elif ifrs_level == 2:
        fair_value_method = "Market Comparison Approach (IFRS 13 Level 2)"
    else:
        fair_value_method = "Quoted Market Price (IFRS 13 Level 1)"

    # Direct Capitalisation: FV = NOI / Cap Rate
    fair_value: float = (noi / cap_rate) if cap_rate > 0 else market_value
    fair_value_variance_pct: float = (fair_value - market_value) / market_value * 100

    # ── 3. Dividend Yield & FFO ───────────────────────────────────────────────
    dividend_yield_pct: float = (noi / market_value) * 100   # same as cap rate for direct RE
    ffo: float                = noi                           # simplified (no D&A add-back for land)
    ffo_yield_pct: float      = ffo / market_value * 100
    price_to_ffo: float       = (1.0 / (ffo_yield_pct / 100)) if ffo_yield_pct > 0 else 0.0

    # ── 4. Portfolio Health Score ─────────────────────────────────────────────
    health_score: float = 100.0

    if vacancy_rate > 0.15:
        health_score -= 20
    if cap_rate_pct < 5.0:
        health_score -= 15
    if fair_value_variance_pct < -10:
        health_score -= 15
    if gross_income > 0 and (operating_expenses / gross_income) > 0.40:
        health_score -= 10
    if cap_rate_pct > 8.0:
        health_score += 10

    health_score = max(0.0, min(100.0, health_score))

    if health_score >= 85:
        health_label = "ممتاز — استثمار مؤسسي آمن"
    elif health_score >= 70:
        health_label = "جيد — مقبول للصناديق"
    elif health_score >= 55:
        health_label = "متوسط — يحتاج مراجعة"
    else:
        health_label = "ضعيف — توصية بإعادة الهيكلة"

    # ── 5. NAV (Net Asset Value) ──────────────────────────────────────────────
    nav: float          = fair_value - loan_amount
    nav_per_unit: float = nav / total_units

    # ── 6. Audit Timestamp (IFRS 13 §93 disclosure) ───────────────────────────
    today_str: str      = datetime.today().strftime("%Y-%m-%d")
    valuation_date: str = today_str
    ifrs_disclosure: str = (
        f"وفقاً لـ IFRS 13 §93، تم الإفصاح عن القيمة العادلة كـ Level {ifrs_level} "
        f"بتاريخ {today_str}"
    )

    # ── 7. Assemble & return result dict ─────────────────────────────────────
    return {
        # Metadata
        "fund_name":               fund_name,
        "client_name":             client_name,
        "property_type":           property_type,
        "location":                location,
        "area":                    area,
        "total_units":             total_units,
        "valuation_date":          valuation_date,
        # Inputs
        "market_value":            market_value,
        "annual_rent":             annual_rent,
        "vacancy_rate":            vacancy_rate,
        "vacancy_rate_pct":        vacancy_rate * 100,
        "operating_expenses":      operating_expenses,
        "loan_amount":             loan_amount,
        # NOI
        "gross_income":            gross_income,
        "vacancy_loss":            vacancy_loss,
        "effective_income":        effective_income,
        "noi":                     noi,
        # Cap Rate
        "cap_rate":                cap_rate,
        "cap_rate_pct":            cap_rate_pct,
        # IFRS 13
        "ifrs_level":              ifrs_level,
        "ifrs_level_label":        ifrs_level_labels[ifrs_level],
        "fair_value_method":       fair_value_method,
        "fair_value":              fair_value,
        "fair_value_variance_pct": fair_value_variance_pct,
        # Yield / FFO
        "dividend_yield_pct":      dividend_yield_pct,
        "ffo":                     ffo,
        "ffo_yield_pct":           ffo_yield_pct,
        "price_to_ffo":            price_to_ffo,
        # Health
        "health_score":            health_score,
        "health_label":            health_label,
        # NAV
        "nav":                     nav,
        "nav_per_unit":            nav_per_unit,
        # IFRS Disclosure
        "ifrs_disclosure":         ifrs_disclosure,
    }


# ══════════════════════════════════════════════════════════════════════════════
#  EXCEL REPORT
# ══════════════════════════════════════════════════════════════════════════════

def generate_fund_excel(result: dict, output_dir: str = "") -> str:
    """
    Generates an RTL Arabic IFRS 13 fund valuation Excel report using xlsxwriter.

    Parameters
    ----------
    result     : dict returned by run_fund_valuation()
    output_dir : override output directory (default: _OUT_DIR)

    Returns
    -------
    str — absolute path to the saved .xlsx file.
    """
    import xlsxwriter  # type: ignore

    out_dir = output_dir.strip() if output_dir.strip() else _OUT_DIR
    os.makedirs(out_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filepath = os.path.join(out_dir, f"fund_valuation_{timestamp}.xlsx")

    wb = xlsxwriter.Workbook(filepath)
    ws = wb.add_worksheet("تقرير الصندوق العقاري — IFRS 13")

    # ── Right-to-left sheet ───────────────────────────────────────────────────
    ws.right_to_left()

    # ── Colour palette ────────────────────────────────────────────────────────
    NAVY   = "#1F3864"
    GOLD   = "#C9A227"
    WHITE  = "#FFFFFF"
    LIGHT  = "#EAF0FB"
    GREEN  = "#1E8449"
    YELLOW = "#D4AC0D"
    RED    = "#C0392B"
    GREY   = "#F2F2F2"
    BLACK  = "#000000"

    # ── Base formats ──────────────────────────────────────────────────────────
    def _fmt(**kw):
        base = {
            "font_name": "Simplified Arabic",
            "font_size": 11,
            "text_wrap": True,
            "valign": "vcenter",
            "align": "right",
            "border": 1,
        }
        base.update(kw)
        return wb.add_format(base)

    fmt_title = _fmt(
        bold=True, font_size=16, font_color=WHITE,
        bg_color=NAVY, align="center", border=0
    )
    fmt_subtitle = _fmt(
        bold=True, font_size=12, font_color=GOLD,
        bg_color=NAVY, align="center", border=0
    )
    fmt_header = _fmt(
        bold=True, font_size=11, font_color=WHITE,
        bg_color=NAVY
    )
    fmt_label = _fmt(bold=True, bg_color=LIGHT)
    fmt_value = _fmt(num_format="#,##0.00")
    fmt_pct   = _fmt(num_format="0.00%")
    fmt_int   = _fmt(num_format="#,##0")
    fmt_plain = _fmt()
    fmt_gold_label = _fmt(bold=True, font_color=GOLD, bg_color=NAVY)
    fmt_ifrs_box   = _fmt(bg_color=LIGHT, italic=True, font_size=10)

    # Health gauge colours
    hs = result["health_score"]
    if hs >= 85:
        gauge_color = GREEN
    elif hs >= 70:
        gauge_color = YELLOW
    else:
        gauge_color = RED
    fmt_health = _fmt(bold=True, font_color=WHITE, bg_color=gauge_color,
                      align="center", font_size=13)

    # ── Column widths ─────────────────────────────────────────────────────────
    ws.set_column("A:A", 38)
    ws.set_column("B:B", 28)
    ws.set_column("C:C", 28)
    ws.set_column("D:D", 18)

    row = 0

    # ── Title block ───────────────────────────────────────────────────────────
    ws.merge_range(row, 0, row, 3,
                   "تقرير تقييم الصندوق العقاري — IFRS 13 / IVS 103",
                   fmt_title)
    row += 1
    ws.merge_range(row, 0, row, 3, result["fund_name"], fmt_subtitle)
    row += 1
    ws.merge_range(row, 0, row, 3,
                   f"تاريخ التقييم: {result['valuation_date']}   |   "
                   f"العميل: {result.get('client_name', '—')}   |   "
                   f"الموقع: {result['location']}",
                   fmt_subtitle)
    row += 2

    # ── Helper: write a two-column KPI table ──────────────────────────────────
    def _section_header(title: str):
        nonlocal row
        ws.merge_range(row, 0, row, 3, title, fmt_gold_label)
        row += 1

    def _kpi_row(label: str, value, fmt=None):
        nonlocal row
        ws.write(row, 0, label, fmt_label)
        ws.write(row, 1, value, fmt or fmt_value)
        ws.write(row, 2, "", fmt_plain)
        ws.write(row, 3, "", fmt_plain)
        row += 1

    def _kpi_row2(label1, val1, label2, val2, fmt1=None, fmt2=None):
        nonlocal row
        ws.write(row, 0, label1, fmt_label)
        ws.write(row, 1, val1, fmt1 or fmt_value)
        ws.write(row, 2, label2, fmt_label)
        ws.write(row, 3, val2, fmt2 or fmt_value)
        row += 1

    # ── Section A: Property & Fund Info ───────────────────────────────────────
    _section_header("أ — معلومات الصندوق والعقار")
    _kpi_row2("نوع العقار",    result["property_type"],   "عدد الوحدات",  result["total_units"], fmt_plain, fmt_int)
    _kpi_row2("المساحة (م²)",  result["area"],            "القيمة السوقية (ج.م.)", result["market_value"])
    _kpi_row2("الإيجار السنوي (ج.م.)", result["annual_rent"],
              "نسبة الشغور",   result["vacancy_rate_pct"] / 100, fmt_value, fmt_pct)
    _kpi_row2("المصاريف التشغيلية (ج.م.)", result["operating_expenses"],
              "مبلغ القرض (ج.م.)",         result["loan_amount"])
    row += 1

    # ── Section B: NOI ────────────────────────────────────────────────────────
    _section_header("ب — صافي الدخل التشغيلي (NOI)")
    ws.write(row, 0, "البند", fmt_header)
    ws.write(row, 1, "القيمة (ج.م.)", fmt_header)
    ws.write(row, 2, "", fmt_plain)
    ws.write(row, 3, "", fmt_plain)
    row += 1
    for label, key in [
        ("إجمالي الدخل (Gross Income)",   "gross_income"),
        ("خسارة الشغور (Vacancy Loss)",   "vacancy_loss"),
        ("الدخل الفعلي (Effective Income)", "effective_income"),
        ("المصاريف التشغيلية",             "operating_expenses"),
        ("صافي الدخل التشغيلي (NOI)",      "noi"),
    ]:
        ws.write(row, 0, label, fmt_label)
        ws.write(row, 1, result[key], fmt_value)
        ws.write(row, 2, "", fmt_plain)
        ws.write(row, 3, "", fmt_plain)
        row += 1
    row += 1

    # ── Section C: Key KPIs ───────────────────────────────────────────────────
    _section_header("ج — المؤشرات الرئيسية للصندوق")
    ws.write(row, 0, "المؤشر",        fmt_header)
    ws.write(row, 1, "القيمة",        fmt_header)
    ws.write(row, 2, "المؤشر",        fmt_header)
    ws.write(row, 3, "القيمة",        fmt_header)
    row += 1
    _kpi_row2("Cap Rate (معدل الرسملة)",  result["cap_rate"],
              "Cap Rate %",               result["cap_rate_pct"] / 100,
              fmt_pct, fmt_pct)
    _kpi_row2("القيمة العادلة IFRS 13 (ج.م.)", result["fair_value"],
              "فارق القيمة العادلة %",    result["fair_value_variance_pct"] / 100,
              fmt_value, fmt_pct)
    _kpi_row2("عائد التوزيعات (Dividend Yield)", result["dividend_yield_pct"] / 100,
              "FFO Yield",                result["ffo_yield_pct"] / 100,
              fmt_pct, fmt_pct)
    _kpi_row2("صافي الدخل التشغيلي FFO (ج.م.)", result["ffo"],
              "P/FFO (السعر / FFO)",       result["price_to_ffo"],
              fmt_value, fmt_value)
    row += 1

    # ── Section D: IFRS 13 Classification Box ────────────────────────────────
    _section_header("د — تصنيف IFRS 13 للقيمة العادلة")

    levels_info = [
        ("Level 1", "أسعار معلنة في أسواق نشطة",   "لا تنطبق على العقار المباشر"),
        ("Level 2", "مدخلات قابلة للملاحظة",        "المعاملات المقارنة في السوق"),
        ("Level 3", "مدخلات غير قابلة للملاحظة",    "نهج الدخل والتدفقات النقدية المخصومة"),
    ]
    for lvl, ar_name, desc in levels_info:
        active = (int(lvl.split()[1]) == result["ifrs_level"])
        cell_fmt = _fmt(
            bold=active,
            bg_color=NAVY if active else GREY,
            font_color=GOLD if active else BLACK,
        )
        ws.write(row, 0, f"{'◄ ' if active else ''}{lvl} — {ar_name}", cell_fmt)
        ws.merge_range(row, 1, row, 3, desc, fmt_ifrs_box)
        row += 1

    ws.merge_range(row, 0, row, 3,
                   f"أسلوب التقييم المعتمد: {result['fair_value_method']}",
                   _fmt(bold=True, bg_color=LIGHT, font_color=NAVY))
    row += 1
    ws.merge_range(row, 0, row, 3, result["ifrs_disclosure"],
                   fmt_ifrs_box)
    row += 2

    # ── Section E: Portfolio Health Gauge ─────────────────────────────────────
    _section_header("هـ — مؤشر صحة المحفظة (Portfolio Health)")
    ws.merge_range(row, 0, row, 1,
                   f"درجة الصحة: {result['health_score']:.1f} / 100",
                   fmt_health)
    ws.merge_range(row, 2, row, 3, result["health_label"], fmt_health)
    row += 1

    # Gauge breakdown legend
    for threshold, label, color in [
        ("≥ 85", "ممتاز — استثمار مؤسسي آمن", GREEN),
        ("≥ 70", "جيد — مقبول للصناديق",       YELLOW),
        ("≥ 55", "متوسط — يحتاج مراجعة",       "#E67E22"),
        ("< 55",  "ضعيف — توصية بإعادة الهيكلة", RED),
    ]:
        legend_fmt = _fmt(font_color=WHITE, bg_color=color, font_size=10, border=0)
        ws.write(row, 0, threshold,  legend_fmt)
        ws.merge_range(row, 1, row, 3, label, legend_fmt)
        row += 1
    row += 1

    # ── Section F: NAV ────────────────────────────────────────────────────────
    _section_header("و — صافي قيمة الأصول (NAV)")
    _kpi_row2("القيمة العادلة (ج.م.)",         result["fair_value"],
              "القرض (ج.م.)",                   result["loan_amount"])
    _kpi_row2("NAV الإجمالي (ج.م.)",            result["nav"],
              "NAV لكل وحدة (ج.م.)",            result["nav_per_unit"])
    row += 2

    # ── Footer ────────────────────────────────────────────────────────────────
    footer_fmt = _fmt(italic=True, font_size=9, font_color=NAVY,
                      bg_color=LIGHT, border=0, align="center")
    ws.merge_range(row, 0, row, 3,
                   "تم إعداد هذا التقرير وفق معايير IFRS 13 و IVS 103 | "
                   "نظام Expert_Smart للتقييم العقاري",
                   footer_fmt)
    row += 1
    ws.merge_range(row, 0, row, 3,
                   "هذا التقرير سري ومُعدّ للاستخدام المؤسسي فقط",
                   footer_fmt)

    wb.close()
    return filepath


# ══════════════════════════════════════════════════════════════════════════════
#  WORD SECTION
# ══════════════════════════════════════════════════════════════════════════════

def generate_fund_word_section(doc, result: dict) -> None:
    """
    Appends an IFRS 13 / IVS 103 fund valuation section to a python-docx Document.

    Parameters
    ----------
    doc    : docx.Document  — existing document to append to
    result : dict           — output of run_fund_valuation()
    """
    from docx.shared import Pt, RGBColor, Inches  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.enum.table import WD_TABLE_ALIGNMENT  # type: ignore
    from docx.oxml.ns import qn                     # type: ignore
    from docx.oxml import OxmlElement               # type: ignore

    NAVY  = RGBColor(0x1F, 0x38, 0x64)
    GOLD  = RGBColor(0xC9, 0xA2, 0x27)
    GREEN = RGBColor(0x1E, 0x84, 0x49)
    AMBER = RGBColor(0xD4, 0xAC, 0x0D)
    RED   = RGBColor(0xC0, 0x39, 0x2B)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)

    FONT_AR = "Simplified Arabic"

    def _set_rtl(paragraph):
        """Force RTL direction on a paragraph."""
        pPr = paragraph._p.get_or_add_pPr()
        bidi = OxmlElement("w:bidi")
        bidi.set(qn("w:val"), "1")
        pPr.append(bidi)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _heading(text: str, level: int = 1):
        p = doc.add_paragraph()
        _set_rtl(p)
        run = p.add_run(text)
        run.font.name     = FONT_AR
        run.font.size     = Pt(16 if level == 1 else 13)
        run.font.bold     = True
        run.font.color.rgb = NAVY if level == 1 else GOLD
        return p

    def _body(text: str, color: RGBColor = None, bold: bool = False):
        p = doc.add_paragraph()
        _set_rtl(p)
        run = p.add_run(text)
        run.font.name  = FONT_AR
        run.font.size  = Pt(11)
        run.font.bold  = bold
        if color:
            run.font.color.rgb = color
        return p

    def _add_table(headers: list, rows: list):
        """Add a simple styled table."""
        col_count = len(headers)
        tbl = doc.add_table(rows=1 + len(rows), cols=col_count)
        tbl.alignment = WD_TABLE_ALIGNMENT.RIGHT
        tbl.style = "Table Grid"

        # Header row
        hdr_cells = tbl.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            run = hdr_cells[i].paragraphs[0].runs[0]
            run.font.name      = FONT_AR
            run.font.size      = Pt(11)
            run.font.bold      = True
            run.font.color.rgb = WHITE
            # Navy background
            tc_pr = hdr_cells[i]._tc.get_or_add_tcPr()
            shd   = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  "1F3864")
            tc_pr.append(shd)
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Data rows
        for r_idx, row_data in enumerate(rows):
            row_cells = tbl.rows[r_idx + 1].cells
            for c_idx, cell_val in enumerate(row_data):
                row_cells[c_idx].text = str(cell_val)
                run = row_cells[c_idx].paragraphs[0].runs[0]
                run.font.name = FONT_AR
                run.font.size = Pt(10)
                row_cells[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        return tbl

    def _fmt_egp(v: float) -> str:
        return f"{v:,.2f} ج.م."

    def _fmt_pct(v: float) -> str:
        return f"{v:.2f}%"

    # ── Section Title ─────────────────────────────────────────────────────────
    doc.add_page_break()
    _heading("تقرير الصندوق العقاري — IFRS 13 / IVS 103")
    _body(
        f"الصندوق: {result['fund_name']}   |   "
        f"العميل: {result.get('client_name', '—')}   |   "
        f"الموقع: {result['location']}   |   "
        f"تاريخ التقييم: {result['valuation_date']}",
        color=NAVY
    )
    doc.add_paragraph()

    # ── Fair Value Table ──────────────────────────────────────────────────────
    _heading("أ — القيمة العادلة وفق IFRS 13", level=2)
    _add_table(
        headers=["البند", "القيمة"],
        rows=[
            ["القيمة السوقية",               _fmt_egp(result["market_value"])],
            ["القيمة العادلة (IFRS 13)",      _fmt_egp(result["fair_value"])],
            ["فارق القيمة العادلة",           _fmt_pct(result["fair_value_variance_pct"])],
            ["أسلوب التقييم",                result["fair_value_method"]],
            ["تصنيف IFRS 13",               result["ifrs_level_label"]],
        ]
    )
    doc.add_paragraph()

    # ── Cap Rate & Yield Table ────────────────────────────────────────────────
    _heading("ب — معدل الرسملة والعوائد", level=2)
    _add_table(
        headers=["المؤشر", "القيمة"],
        rows=[
            ["صافي الدخل التشغيلي (NOI)",  _fmt_egp(result["noi"])],
            ["معدل الرسملة (Cap Rate)",      _fmt_pct(result["cap_rate_pct"])],
            ["عائد التوزيعات (Dividend Yield)", _fmt_pct(result["dividend_yield_pct"])],
            ["FFO",                           _fmt_egp(result["ffo"])],
            ["FFO Yield",                     _fmt_pct(result["ffo_yield_pct"])],
            ["P/FFO",                         f"{result['price_to_ffo']:.2f}x"],
        ]
    )
    doc.add_paragraph()

    # ── NAV Section ───────────────────────────────────────────────────────────
    _heading("ج — صافي قيمة الأصول (NAV)", level=2)
    _add_table(
        headers=["البند", "القيمة"],
        rows=[
            ["القيمة العادلة",       _fmt_egp(result["fair_value"])],
            ["القرض المستحق",        _fmt_egp(result["loan_amount"])],
            ["NAV الإجمالي",         _fmt_egp(result["nav"])],
            ["NAV لكل وحدة",         _fmt_egp(result["nav_per_unit"])],
            ["عدد الوحدات",          str(result["total_units"])],
        ]
    )
    doc.add_paragraph()

    # ── Portfolio Health ──────────────────────────────────────────────────────
    _heading("د — مؤشر صحة المحفظة", level=2)
    hs = result["health_score"]
    if hs >= 85:
        health_color = GREEN
    elif hs >= 70:
        health_color = AMBER
    else:
        health_color = RED

    _body(
        f"درجة الصحة: {hs:.1f} / 100   —   {result['health_label']}",
        color=health_color,
        bold=True
    )
    doc.add_paragraph()

    # ── IFRS 13 Level Disclosure ──────────────────────────────────────────────
    _heading("هـ — إفصاح IFRS 13 §93", level=2)
    _body(result["ifrs_disclosure"], color=NAVY)
    doc.add_paragraph()

    # ── Footer note ───────────────────────────────────────────────────────────
    footer = doc.add_paragraph()
    _set_rtl(footer)
    run = footer.add_run(
        "تم إعداد هذا التقرير وفق معايير IFRS 13 و IVS 103 | "
        "نظام Expert_Smart للتقييم العقاري | سري — للاستخدام المؤسسي فقط"
    )
    run.font.name      = FONT_AR
    run.font.size      = Pt(9)
    run.font.italic    = True
    run.font.color.rgb = NAVY
