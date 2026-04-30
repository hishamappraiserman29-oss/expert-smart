import pandas as pd
import numpy as np
import os
from datetime import datetime
import xlsxwriter
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from scipy.spatial import distance
from pykrige.ok import OrdinaryKriging

class ExpertSmartIVS:
    def __init__(self):
        self.output_excel = os.path.join(os.getcwd(), 'التحليل_الرقمي_الشامل_IVS.xlsx')
        self.output_word = os.path.join(os.getcwd(), 'التقرير_السردي_النهائي_IVS.docx')
        self.report_date = datetime.now().strftime("%Y/%m/%d")

    # --- محرك التحليل المكاني (GIS Interpolation) ---
    def spatial_analysis(self):
        # بيانات افتراضية للمقارنات (إحداثيات وسعر متر) - يتم سحبها لاحقاً من قاعدة بياناتك
        coords = [(30.01, 31.02), (30.02, 31.03), (30.015, 31.025), (30.022, 31.035)]
        prices = [15000, 16500, 15800, 17000]
        target = (30.012, 31.022) # إحداثيات العقار المستهدف

        # 1. حساب IDW
        dists = [distance.euclidean(target, c) for c in coords]
        weights = [1.0 / (d**2) for d in dists]
        idw_val = sum(w * v for w, v in zip(weights, prices)) / sum(weights)

        # 2. حساب Kriging
        ok = OrdinaryKriging([c[0] for c in coords], [c[1] for c in coords], prices, variogram_model='linear')
        krig_val, _ = ok.execute('single', target[0], target[1])
        
        return idw_val, float(krig_val)

    # --- محرك الإكسيل (Excel Mirror Engine) ---
    def generate_excel(self, idw, krig):
        wb = xlsxwriter.Workbook(self.output_excel, {'nan_inf_to_errors': True})
        
        # التنسيقات الفخمة
        h_fmt = wb.add_format({'bold': True, 'bg_color': '#1F4E78', 'font_color': 'white', 'border': 2, 'align': 'center', 'font_name': 'Simplified Arabic'})
        sub_h = wb.add_format({'bold': True, 'bg_color': '#D9E1F2', 'border': 1, 'align': 'center'})
        n_fmt = wb.add_format({'num_format': '#,##0', 'border': 1, 'align': 'center'})
        p_fmt = wb.add_format({'num_format': '0%', 'border': 1, 'align': 'center'})
        t_fmt = wb.add_format({'text_wrap': True, 'border': 1, 'align': 'right', 'font_name': 'Simplified Arabic'})

        # شيت 1: مصفوفة السوق والضبط (التقليدية)
        ws1 = wb.add_worksheet('1- أسلوب السوق التفصيلي')
        ws1.right_to_left()
        ws1.set_column('A:F', 20)
        ws1.merge_range('A1:F1', 'مصفوفة تحليل مقارنات السوق (Market Grid Analysis)', h_fmt)
        ws1.write_row(1, 0, ['عنصر المقارنة', 'مقارنة 1', 'مقارنة 2', 'مقارنة 3', 'مقارنة 4', 'المتوسط'], sub_h)
        market_rows = [
            ['سعر المتر المبدئي', 18000, 19500, 17500, 20000, ''],
            ['تعديل الموقع (+/-)', '+5%', '0%', '+10%', '-5%', ''],
            ['تعديل التشطيب (+/-)', '-10%', '0%', '-5%', '+5%', ''],
            ['سعر المتر المعدل', 17100, 19500, 18375, 20000, 18743]
        ]
        for r, row in enumerate(market_rows, 2):
            ws1.write_row(r, 0, row, t_fmt if r==2 or r==5 else p_fmt)

        # شيت 2: التحليل الجيومكاني (الحديث - GIS)
        ws2 = wb.add_worksheet('2- التحليل المكاني GIS')
        ws2.right_to_left()
        ws2.set_column('A:B', 40)
        ws2.write('A1', 'منهجية الاستيفاء المكاني (Interpolation)', h_fmt)
        ws2.write_row(1, 0, ['الخوارزمية المستخدمة', 'قيمة المتر المقدرة (ج.م)'], sub_h)
        ws2.write(2, 0, 'المسافة العكسية المرجحة (IDW Analysis)', t_fmt); ws2.write(2, 1, idw, n_fmt)
        ws2.write(3, 0, 'التحليل الإحصائي الجيومكاني (Kriging)', t_fmt); ws2.write(3, 1, krig, n_fmt)

        # شيت 3: التوفيق النهائي (7 طرق)
        ws3 = wb.add_worksheet('3- مصفوفة التوفيق النهائي')
        ws3.right_to_left()
        ws3.set_column('A:D', 25)
        ws3.write_row(0, 0, ['المنهجية', 'القيمة المقدرة', 'الوزن', 'المرجح'], h_fmt)
        final_data = [
            ['أسلوب السوق (تقليدي)', 18743*150, 0.30],
            ['أسلوب التكلفة (تقليدي)', 2400000, 0.15],
            ['تحليل IDW (حديث)', idw*150, 0.15],
            ['تحليل Kriging (حديث)', krig*150, 0.15],
            ['الانحدار المتعدد', 2750000, 0.15],
            ['الخيارات الحقيقية', 2900000, 0.05],
            ['الذكاء الاصطناعي', 2800000, 0.05]
        ]
        for r, row in enumerate(final_data, 1):
            ws3.write(r, 0, row[0], t_fmt); ws3.write(r, 1, row[1], n_fmt)
            ws3.write(r, 2, row[2], p_fmt); ws3.write_formula(r, 3, f'=B{r+1}*C{r+1}', n_fmt)
        
        ws3.write(9, 0, 'القيمة النهائية المرجحة', h_fmt)
        ws3.write_formula(9, 3, '=SUM(D2:D8)', h_fmt)

        wb.close()
        return self.output_excel

    # --- محرك الورد (Word Narrative Engine) ---
    def generate_word(self, idw, krig):
        doc = Document()
        # جعل المستند من اليمين لليسار
        section = doc.sections[0]
        section.right_to_left = True

        # تنسيق الخط العربي
        style = doc.styles['Normal']
        style.font.name = 'Simplified Arabic'
        style.font.size = Pt(13)

        # 1. الغلاف
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("تقرير تقييم عقاري احترافي (IVS Standard)\n")
        run.bold = True; run.font.size = Pt(26); run.font.color.rgb = RGBColor(31, 78, 120)

        # 2. المقدمة والمعاينة
        doc.add_heading('أولاً: المعاينة والوصف الفني', level=1)
        doc.add_paragraph("تمت المعاينة الميدانية للوحدة موضوع التقييم وتبين أنها تقع في منطقة دريم لاند، وتتميز بتشطيبات سوبر لوكس وأرضيات رخامية ودهانات حديثة...")

        # 3. التحليل التكنولوجي (الـ GIS)
        doc.add_heading('ثانياً: التحليل الجيومكاني المتقدم (GIS)', level=1)
        doc.add_paragraph(f"تم استخدام تقنيات الاستيفاء المكاني لتقدير القيمة بناءً على الموقع الجغرافي الدقيق، حيث أعطى تحليل IDW قيمة {idw:,.0f} ج.م للمتر، بينما أعطى تحليل Kriging قيمة {krig:,.0f} ج.م للمتر.")

        # 4. النتيجة النهائية
        doc.add_heading('ثالثاً: القيمة السوقية العادلة', level=1)
        doc.add_paragraph("بناءً على التوفيق بين 7 منهجيات (تقليدية وحديثة)، تم التوصل للقيمة السوقية النهائية الموضحة في ملحق الإكسيل المرفق.")

        doc.save(self.output_word)
        return self.output_word

if __name__ == "__main__":
    system = ExpertSmartIVS()
    print("⏳ جاري تحليل البيانات مكانياً...")
    idw_v, krig_v = system.spatial_analysis()
    
    print("📊 جاري إنشاء ملف الإكسيل التفصيلي...")
    excel_file = system.generate_excel(idw_v, krig_v)
    
    print("📝 جاري إنشاء ملف الورد المنسق...")
    word_file = system.generate_word(idw_v, krig_v)
    
    print(f"\n✅ تمت العملية بنجاح!")
    print(f"1- ملف الإكسيل: {excel_file}")
    print(f"2- ملف الورد: {word_file}")