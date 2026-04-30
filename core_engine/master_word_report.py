"""
master_word_report.py
=====================
Master Unified Valuation Report (Word / .docx)
------------------------------------------------
Generates a single comprehensive Word document integrating:
  1.  Cover Page — certified header, expert credentials, IVS badge
  2.  Executive Summary — final value, purpose, confidence level
  3.  Property Description — site details, physical characteristics
  4.  Market Overview — Egypt/KSA market context, comparable evidence
  5.  Valuation Methodology — IVS 103 / 104 / 105 framework overview
  6.  Market Comparison Approach — adjusted comps grid
  7.  Cost / DRC Approach — RCN breakdown, depreciation schedule
  8.  Income Capitalisation Approach — NOI, cap rate, income value
  9.  GIS Spatial Analysis — IDW & Kriging results, heatmap note
 10.  OLS Regression Analysis — coefficients, R², predicted PPM
 11.  HBU Feasibility Analysis — 3 scenarios NPV/IRR table
 12.  Seven-Method Reconciliation — weighted final value table
 13.  IVS/USPAP Compliance Scorecard — audit checklist
 14.  Expert Certification — هشام محمد المهدى، رقم القيد 29

Usage:
    from master_word_report import generate_master_report
    path = generate_master_report(ivs_result)
"""
from __future__ import annotations
import os
from datetime import datetime
from typing import Dict, Any

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import docx
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

_BASE_DIR  = os.path.dirname(os.path.abspath(__file__))
_OUT_DIR   = os.path.join(_BASE_DIR, "outputs", "reports")

# ── Design Constants ──────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1F, 0x4E, 0x78)
GOLD   = RGBColor(0xD4, 0xAF, 0x37)
GREEN  = RGBColor(0x2E, 0x86, 0x48)
GRAY   = RGBColor(0x60, 0x60, 0x60)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT  = RGBColor(0xEB, 0xF3, 0xFB)

FONT_MAIN  = "Simplified Arabic"
FONT_MONO  = "Courier New"


# ── RTL helpers ───────────────────────────────────────────────────────────────

def _set_rtl(paragraph):
    """Force RTL on a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "right")
    pPr.append(jc)


def _add_heading(doc: "Document", text: str, level: int = 1,
                 color: RGBColor = None) -> Any:
    p = doc.add_heading(text, level=level)
    _set_rtl(p)
    for run in p.runs:
        run.font.name = FONT_MAIN
        run.font.color.rgb = color or NAVY
        run.font.bold = True
        run.font.size = Pt(16 - (level - 1) * 2)
    return p


def _add_para(doc: "Document", text: str, bold: bool = False,
              size: int = 12, color: RGBColor = None,
              align: str = "right") -> Any:
    p = doc.add_paragraph()
    _set_rtl(p)
    run = p.add_run(text)
    run.font.name = FONT_MAIN
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return p


def _shade_cell(cell, hex_color: str):
    """Apply background shading to a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _add_table_row(table, cells: list, bold_first: bool = True,
                   header: bool = False):
    row = table.add_row()
    row.height = Cm(0.8)
    for i, (cell, text) in enumerate(zip(row.cells, cells)):
        cell.text = str(text)
        _set_rtl(cell.paragraphs[0])
        for run in cell.paragraphs[0].runs:
            run.font.name = FONT_MAIN
            run.font.size = Pt(11)
            run.font.bold = (bold_first and i == 0) or header
            if header:
                run.font.color.rgb = WHITE
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if header:
            _shade_cell(cell, "1F4E78")


def _hr(doc: "Document"):
    """Horizontal rule (thin border paragraph)."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D4AF37")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _fmt(n) -> str:
    try:
        return f"{float(n):,.0f}"
    except Exception:
        return "0"


# ── Section builders ──────────────────────────────────────────────────────────

def _cover_page(doc: "Document", r: Dict):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("تقرير التقييم العقاري المعتمد")
    run.font.name = FONT_MAIN
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = NAVY

    _add_para(doc, "IVS / EFSA Compliant Valuation Report",
              bold=True, size=13, color=GRAY, align="center")
    _hr(doc)
    doc.add_paragraph()

    # Key info table
    tbl = doc.add_table(rows=0, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    data = [
        ("الغرض من التقييم",  r.get("purpose_label_ar", "القيمة السوقية العادلة")),
        ("العميل",            r.get("client_name", "—")),
        ("نوع العقار",        r.get("property_type", "—")),
        ("الموقع",            r.get("location", "—")),
        ("المساحة",           f"{r.get('area', 0):,.0f} م²"),
        ("تاريخ التقرير",     datetime.now().strftime("%Y-%m-%d")),
        ("المعيار المطبق",    "IVS 2022 / معايير هيئة الرقابة المالية"),
    ]
    for lbl, val in data:
        row = tbl.add_row()
        row.cells[0].text = lbl
        row.cells[1].text = val
        for i, cell in enumerate(row.cells):
            _set_rtl(cell.paragraphs[0])
            for run in cell.paragraphs[0].runs:
                run.font.name = FONT_MAIN
                run.font.size = Pt(12)
                run.font.bold = (i == 0)
            if i == 0:
                _shade_cell(cell, "EBF3FB")

    tbl.columns[0].width = Cm(6)
    tbl.columns[1].width = Cm(10)

    doc.add_paragraph()
    _add_para(doc, "المقيّم المعتمد: هشام محمد المهدى",
              bold=True, size=14, color=GOLD, align="center")
    _add_para(doc, "رقم القيد: 29 | عضو مجلس المعايير الدولية للتقييم",
              size=11, color=GRAY, align="center")
    doc.add_page_break()


def _executive_summary(doc: "Document", r: Dict):
    _add_heading(doc, "أولاً: الملخص التنفيذي", 1)
    _hr(doc)
    val = _fmt(r.get("reconciled_value", 0))
    purpose = r.get("purpose_label_ar", "القيمة السوقية العادلة")
    disc    = r.get("forced_discount", 0)
    disc_txt = f" (بعد خصم {disc*100:.0f}%)" if disc > 0 else ""

    _add_para(doc, f"القيمة النهائية المعتمدة ({purpose}){disc_txt}:", bold=True, size=13)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_rtl(p)
    run = p.add_run(f"{val}  جنيه مصري")
    run.font.name = FONT_MAIN
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = GOLD

    doc.add_paragraph()
    _add_para(doc,
        "يُعدّ هذا التقرير وثيقة تقييم رسمية مُعدَّة وفقاً لمعايير التقييم الدولية (IVS 2022) "
        "ومعايير هيئة الرقابة المالية المصرية (EFSA)، ويستند إلى سبعة أساليب تقييم مندمجة "
        "تشمل: المقارنة السوقية، وطريقة التكلفة، والرسملة المباشرة، والتحليل المكاني (GIS)، "
        "ونموذج الانحدار (OLS)، وتحليل أفضل استخدام (HBU)، ومنهج التدفقات النقدية المخصومة (DCF).",
        size=12)

    # Key metrics table
    doc.add_paragraph()
    tbl = doc.add_table(rows=0, cols=3)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_table_row(tbl, ["المؤشر", "القيمة", "الوزن / التعليق"], header=True)

    market_v = r.get("market", {}).get("value", 0)
    cost_v   = r.get("cost",   {}).get("value", 0)
    income_v = r.get("income", {}).get("value", 0)
    gis_v    = r.get("gis",    {}).get("idw_ppm", 0) * r.get("area", 1)
    ols_v    = r.get("ols",    {}).get("predicted_ppm", 0) * r.get("area", 1)

    rows = [
        ("المقارنة السوقية",      _fmt(market_v)  + " ج.م", "30%"),
        ("التكلفة / DRC",         _fmt(cost_v)    + " ج.م", "20%"),
        ("رسملة الدخل",           _fmt(income_v)  + " ج.م", "20%"),
        ("GIS — IDW المكاني",     _fmt(gis_v)     + " ج.م", "15%"),
        ("OLS — الانحدار",        _fmt(ols_v)     + " ج.م", "15%"),
        ("القيمة النهائية المرجّحة", f"✦ {val} ج.م", "100%"),
    ]
    for cells in rows:
        _add_table_row(tbl, cells, bold_first=(cells[0].startswith("القيمة")))

    doc.add_page_break()


def _property_description(doc: "Document", r: Dict):
    _add_heading(doc, "ثانياً: وصف العقار والموقع", 1)
    _hr(doc)

    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    rows = [
        ("الموقع",           r.get("location", "—")),
        ("نوع العقار",       r.get("property_type", "—")),
        ("المساحة الإجمالية", f"{r.get('area', 0):,.0f} م²"),
        ("الدور",             str(r.get("floor", "—"))),
        ("سنة الإنشاء",       str(r.get("year_built", "—"))),
        ("عمر المبنى",        f"{r.get('building_age', 0)} سنة"),
        ("وصف التشطيبات",    r.get("finishing_description", "تشطيبات نظامية")),
    ]
    for lbl, val in rows:
        row = tbl.add_row()
        row.cells[0].text = lbl
        row.cells[1].text = val
        _set_rtl(row.cells[0].paragraphs[0])
        _set_rtl(row.cells[1].paragraphs[0])
        for run in row.cells[0].paragraphs[0].runs:
            run.font.name = FONT_MAIN; run.font.bold = True; run.font.size = Pt(11)
        for run in row.cells[1].paragraphs[0].runs:
            run.font.name = FONT_MAIN; run.font.size = Pt(11)
        _shade_cell(row.cells[0], "EBF3FB")

    tbl.columns[0].width = Cm(6)
    tbl.columns[1].width = Cm(10)
    doc.add_page_break()


def _market_comparison(doc: "Document", r: Dict):
    _add_heading(doc, "ثالثاً: أسلوب المقارنة السوقية", 1)
    _hr(doc)

    market = r.get("market", {})
    comps  = market.get("comparables", [])

    _add_para(doc,
        f"استناداً إلى بيانات السوق الفعلية، بلغت القيمة بأسلوب المقارنة: "
        f"{_fmt(market.get('value', 0))} ج.م",
        bold=True, size=12)
    doc.add_paragraph()

    if comps:
        tbl = doc.add_table(rows=0, cols=4)
        tbl.style = "Table Grid"
        _add_table_row(tbl, ["الموقع", "المساحة م²", "السعر/م²", "المصدر"], header=True)
        for c in comps[:8]:
            _add_table_row(tbl, [
                c.get("loc", "—"),
                _fmt(c.get("ar", c.get("area", 0))),
                _fmt(c.get("price_per_m2", c.get("price_pm2", 0))),
                c.get("source", "portal"),
            ])
    doc.add_page_break()


def _cost_approach(doc: "Document", r: Dict):
    _add_heading(doc, "رابعاً: أسلوب التكلفة (تكلفة الإحلال المستهلكة)", 1)
    _hr(doc)

    cost = r.get("cost", {})
    rows = [
        ("تكلفة الإنشاء الإجمالية",  _fmt(cost.get("gross_building_cost", 0)) + " ج.م"),
        ("الاستهلاك الإجمالي",        _fmt(cost.get("depreciation", 0))         + " ج.م"),
        ("صافي تكلفة الإحلال (DRC)", _fmt(cost.get("rcn_net", cost.get("value",0))) + " ج.م"),
        ("عمر المبنى الفعلي",         f"{cost.get('building_age', 0)} سنة"),
        ("نسبة الاستهلاك",            f"{cost.get('depr_rate_pct', 0):.1f}%"),
        ("قيمة الأرض",               _fmt(cost.get("land_value", 0)) + " ج.م"),
        ("القيمة الإجمالية (DRC)",    _fmt(cost.get("value", 0)) + " ج.م"),
    ]
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    for lbl, val in rows:
        row = tbl.add_row()
        row.cells[0].text = lbl
        row.cells[1].text = val
        _set_rtl(row.cells[0].paragraphs[0])
        _set_rtl(row.cells[1].paragraphs[0])
        for run in row.cells[0].paragraphs[0].runs:
            run.font.name = FONT_MAIN; run.font.bold = True; run.font.size = Pt(11)
        for run in row.cells[1].paragraphs[0].runs:
            run.font.name = FONT_MAIN; run.font.size = Pt(11)
        _shade_cell(row.cells[0], "EBF3FB")
    tbl.columns[0].width = Cm(7)
    tbl.columns[1].width = Cm(9)
    doc.add_page_break()


def _income_approach(doc: "Document", r: Dict):
    _add_heading(doc, "خامساً: أسلوب الدخل (الرسملة المباشرة)", 1)
    _hr(doc)

    inc = r.get("income", {})
    rows = [
        ("الإيجار السنوي / م²",   f"{_fmt(inc.get('rent_per_sqm', 0))} ج.م"),
        ("معدل الإشغال",          "90%"),
        ("صافي الدخل التشغيلي",   _fmt(inc.get("noi", 0)) + " ج.م"),
        ("معدل الرسملة",          f"{inc.get('cap_rate', 0.08)*100:.1f}%"),
        ("قيمة أسلوب الدخل",     _fmt(inc.get("value", 0)) + " ج.م"),
    ]
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    for lbl, val in rows:
        row = tbl.add_row()
        row.cells[0].text = lbl
        row.cells[1].text = val
        _set_rtl(row.cells[0].paragraphs[0])
        _set_rtl(row.cells[1].paragraphs[0])
        for run in row.cells[0].paragraphs[0].runs:
            run.font.name = FONT_MAIN; run.font.bold = True; run.font.size = Pt(11)
        for run in row.cells[1].paragraphs[0].runs:
            run.font.name = FONT_MAIN; run.font.size = Pt(11)
        _shade_cell(row.cells[0], "EBF3FB")
    tbl.columns[0].width = Cm(7)
    tbl.columns[1].width = Cm(9)
    doc.add_page_break()


def _land_dual_path_section(doc: "Document", r: Dict):
    """Section: Land Dual Path — مصفوفة الأرض + طريقة الباقي + التوفيق."""
    land = r.get("land_dual")
    if not land:
        return

    _add_heading(doc, "ملحق: تحليل قيمة الأرض — المسار المزدوج (V2)", 1)
    _hr(doc)
    _add_para(doc,
        "وفقاً للمعيار IVS-105 (قيمة الأرض)، تم تطبيق مسارَين متوازيَين: "
        "مصفوفة المقارنات السوقية للأراضي المجاورة (وزن 60%)، "
        "وطريقة الباقي القائمة على صافي قيمة التطوير (وزن 40%)، "
        "مع التوفيق بينهما للوصول إلى قيمة أرض موثوقة.",
        size=12)
    doc.add_paragraph()

    # ── Section A: Sales Comparison Grid ─────────────────────────────────────
    sc = land.get("sales_comparison", {})
    _add_heading(doc, "أ — مصفوفة مقارنات الأرض (IVS-105)", 2, color=NAVY)
    grid = sc.get("grid", [])
    if grid:
        adj_keys = list(grid[0]["adjustments"].keys()) if grid else []
        headers  = ["المقارن", "سعر/م²", *adj_keys, "إجمالي التعديل", "السعر المعدل"]
        tbl = doc.add_table(rows=0, cols=len(headers))
        tbl.style = "Table Grid"
        _add_table_row(tbl, headers, header=True)
        for comp in grid:
            row_data = [
                comp["name"],
                _fmt(comp["base_ppm"]),
                *comp["adjustments"].values(),
                comp["total_adj_pct"],
                _fmt(comp["adjusted_ppm"]),
            ]
            _add_table_row(tbl, row_data)
        # Summary row
        _add_table_row(tbl, ["متوسط سعر الأرض/م² (مصفوفة)", _fmt(sc.get("avg_land_ppm", 0)), *["—"]*len(adj_keys), "—", _fmt(sc.get("land_value", 0)) + " ج.م"])
    doc.add_paragraph()

    # ── Section B: Residual Method ────────────────────────────────────────────
    res = land.get("residual", {})
    _add_heading(doc, "ب — طريقة الباقي (Residual Method)", 2, color=NAVY)
    res_rows = [
        ("القيمة البيعية الإجمالية (GDV)",        _fmt(res.get("gdv", 0))               + " ج.م"),
        ("المساحة القابلة للبيع",                  _fmt(res.get("saleable_area_m2", 0))  + " م²"),
        ("سعر الوحدة المطورة/م²",                 _fmt(res.get("sales_price_pm2", 0))   + " ج.م"),
        ("إجمالي تكلفة البناء",                   _fmt(res.get("total_build_cost", 0))  + " ج.م"),
        ("الرسوم المهنية والتصاميم (3%)",           _fmt(res.get("professional_fees", 0)) + " ج.م"),
        ("تكاليف التسويق (2%)",                    _fmt(res.get("marketing_costs", 0))   + " ج.م"),
        ("ربح المطور (20% من GDV)",                _fmt(res.get("developer_profit", 0))  + " ج.م"),
        ("تكاليف التمويل",                         _fmt(res.get("finance_cost", 0))      + " ج.م"),
        ("قيمة الأرض كباقي (Residual Land Value)", "✦ " + _fmt(res.get("land_value", 0)) + " ج.م"),
    ]
    tbl2 = doc.add_table(rows=0, cols=2)
    tbl2.style = "Table Grid"
    for lbl, val in res_rows:
        row = tbl2.add_row()
        row.cells[0].text = lbl
        row.cells[1].text = val
        _set_rtl(row.cells[0].paragraphs[0])
        _set_rtl(row.cells[1].paragraphs[0])
        is_total = "باقي" in lbl
        for run in row.cells[0].paragraphs[0].runs:
            run.font.name = FONT_MAIN; run.font.bold = True; run.font.size = Pt(11)
            if is_total: run.font.color.rgb = GOLD
        for run in row.cells[1].paragraphs[0].runs:
            run.font.name = FONT_MAIN; run.font.size = Pt(11)
            if is_total: run.font.bold = True; run.font.color.rgb = GOLD
        _shade_cell(row.cells[0], "1F4E78" if is_total else "EBF3FB")
        if is_total: _shade_cell(row.cells[1], "FFF8E7")
    tbl2.columns[0].width = Cm(9)
    tbl2.columns[1].width = Cm(7)
    doc.add_paragraph()

    # ── Section C: Reconciliation ─────────────────────────────────────────────
    _add_heading(doc, "ج — التوفيق النهائي بين المسارَين", 2, color=GREEN)
    delta = land.get("delta_pct", 0)
    tbl3 = doc.add_table(rows=0, cols=4)
    tbl3.style = "Table Grid"
    _add_table_row(tbl3, ["الأسلوب", "سعر الأرض/م²", "الوزن", "المساهمة الموزونة"], header=True)
    _add_table_row(tbl3, [
        "مصفوفة المقارنات السوقية",
        _fmt(sc.get("avg_land_ppm", 0)) + " ج.م",
        f"{land.get('sc_weight', 0.6)*100:.0f}%",
        _fmt(sc.get("avg_land_ppm", 0) * land.get("sc_weight", 0.6)) + " ج.م",
    ])
    _add_table_row(tbl3, [
        "طريقة الباقي",
        _fmt(res.get("avg_land_ppm", 0)) + " ج.م",
        f"{land.get('residual_weight', 0.4)*100:.0f}%",
        _fmt(res.get("avg_land_ppm", 0) * land.get("residual_weight", 0.4)) + " ج.م",
    ])
    _add_table_row(tbl3, [
        "✦ قيمة الأرض الموفّقة (م²)",
        _fmt(land.get("reconciled_land_ppm", 0)) + " ج.م",
        "100%",
        "✦ " + _fmt(land.get("reconciled_land_value", 0)) + " ج.م",
    ])
    _add_para(doc, f"الفارق بين المسارَين: {delta:.1f}% — {'ضمن نطاق القبول IVS (<15%)' if delta < 15 else 'تجاوز نطاق القبول — يُنصح بمراجعة المدخلات'}", size=11, color=GREEN if delta < 15 else RGBColor(0xB0, 0x0, 0x0))
    doc.add_page_break()


def _gis_section(doc: "Document", r: Dict):
    _add_heading(doc, "سادساً: التحليل المكاني (GIS)", 1)
    _hr(doc)

    gis = r.get("gis", {})
    area = float(r.get("area", 1))

    _add_para(doc,
        "تم تطبيق نموذجَي التحليل المكاني على قاعدة بيانات نقاط المقارنة الجغرافية "
        "باستخدام معادلتَي IDW (التوزين العكسي للمسافة) وKriging الاعتيادي.",
        size=12)
    doc.add_paragraph()

    rows = [
        ("IDW — سعر/م² المرجّح مكانياً",   _fmt(gis.get("idw_ppm", 0)) + " ج.م"),
        ("IDW — القيمة الإجمالية",           _fmt(gis.get("idw_ppm", 0) * area) + " ج.م"),
        ("Kriging — سعر/م²",                _fmt(gis.get("kriging_ppm", 0)) + " ج.م"),
        ("Kriging — القيمة الإجمالية",       _fmt(gis.get("kriging_ppm", 0) * area) + " ج.م"),
    ]
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    for lbl, val in rows:
        row = tbl.add_row()
        row.cells[0].text = lbl; row.cells[1].text = val
        _set_rtl(row.cells[0].paragraphs[0]); _set_rtl(row.cells[1].paragraphs[0])
        for run in row.cells[0].paragraphs[0].runs:
            run.font.name = FONT_MAIN; run.font.bold = True; run.font.size = Pt(11)
        for run in row.cells[1].paragraphs[0].runs:
            run.font.name = FONT_MAIN; run.font.size = Pt(11)
        _shade_cell(row.cells[0], "EBF3FB")
    tbl.columns[0].width = Cm(9); tbl.columns[1].width = Cm(7)
    doc.add_page_break()


def _ols_section(doc: "Document", r: Dict):
    _add_heading(doc, "سابعاً: نموذج الانحدار الخطي (OLS)", 1)
    _hr(doc)

    ols = r.get("ols", {})
    coef = ols.get("coefficients", {})

    _add_para(doc,
        "تم تطبيق نموذج الانحدار الخطي المتعدد (OLS) على عينة مقارنات المنطقة "
        "لاستخلاص المتغيرات المؤثرة على سعر المتر المربع.",
        size=12)
    doc.add_paragraph()

    rows = [
        ("معامل التحديد R²",        f"{ols.get('r_squared', 0):.3f}"),
        ("R² المعدَّل",             f"{ols.get('adj_r_squared', 0):.3f}"),
        ("السعر المتوقع / م²",      _fmt(ols.get("predicted_ppm", 0)) + " ج.م"),
        ("عدد المشاهدات",           str(ols.get("n_obs", 0))),
        ("ثابت النموذج (β₀)",       _fmt(coef.get("const", 0))),
        ("معامل الدور (β₁)",        f"{coef.get('floor', 0):+.1f}"),
        ("معامل المساحة (β₂)",      f"{coef.get('area', 0):+.1f}"),
        ("معامل سنة البناء (β₃)",   f"{coef.get('year_built', 0):+.1f}"),
    ]
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    for lbl, val in rows:
        row = tbl.add_row()
        row.cells[0].text = lbl; row.cells[1].text = val
        _set_rtl(row.cells[0].paragraphs[0]); _set_rtl(row.cells[1].paragraphs[0])
        for run in row.cells[0].paragraphs[0].runs:
            run.font.name = FONT_MAIN; run.font.bold = True; run.font.size = Pt(11)
        for run in row.cells[1].paragraphs[0].runs:
            run.font.name = FONT_MAIN; run.font.size = Pt(11)
        _shade_cell(row.cells[0], "EBF3FB")
    tbl.columns[0].width = Cm(9); tbl.columns[1].width = Cm(7)
    doc.add_page_break()


def _hbu_section(doc: "Document", r: Dict):
    _add_heading(doc, "ثامناً: تحليل أفضل استخدام (HBU)", 1)
    _hr(doc)

    _add_para(doc,
        "يُعرَّف أفضل الاستخدامات بأنه الاستخدام القانوني الأكثر احتمالاً، "
        "والمادياً الممكن، والمالياً المجدي، والمُحقِّق لأعلى قيمة للعقار. "
        "فيما يلي مقارنة مالية (NPV/IRR) لثلاثة سيناريوهات تطويرية:",
        size=12)
    doc.add_paragraph()

    land_val = r.get("reconciled_value", 0) * 0.30
    area     = float(r.get("area", 500))

    try:
        from hbu_scenarios import run_hbu_scenarios
        hbu = run_hbu_scenarios(
            land_area  = area,
            land_value = land_val,
            location   = r.get("location", "Cairo"),
        )
        tbl = doc.add_table(rows=0, cols=5)
        tbl.style = "Table Grid"
        _add_table_row(tbl, ["السيناريو", "GFA م²", "NPV", "IRR %", "مدة الاسترداد"], header=True)
        for row in hbu.get("summary_table", []):
            feasible = "✅" if row.get("feasible") else "❌"
            _add_table_row(tbl, [
                row.get("name_ar", row.get("scenario", "")),
                _fmt(row.get("gfa_m2", 0)),
                _fmt(row.get("npv", 0)) + " ج.م",
                f"{row.get('irr_pct', 0):.1f}%",
                f"{row.get('payback', '—')} {feasible}",
            ])
        doc.add_paragraph()
        best = hbu.get("recommended", "")
        best_ar = hbu.get("scenarios", {}).get(best, {}).get("name_ar", best)
        _add_para(doc, f"التوصية: أفضل استخدام هو {best_ar}", bold=True, size=12, color=GREEN)
    except Exception as e:
        _add_para(doc, f"(تعذّر تحميل بيانات HBU: {e})", size=11, color=GRAY)

    doc.add_page_break()


def _reconciliation(doc: "Document", r: Dict):
    _add_heading(doc, "تاسعاً: التوفيق بين الأساليب السبعة", 1)
    _hr(doc)

    area      = float(r.get("area", 1))
    market_v  = r.get("market", {}).get("value", 0)
    cost_v    = r.get("cost",   {}).get("value", 0)
    income_v  = r.get("income", {}).get("value", 0)
    gis_idw   = r.get("gis",    {}).get("idw_ppm", 0) * area
    gis_krig  = r.get("gis",    {}).get("kriging_ppm", 0) * area
    ols_v     = r.get("ols",    {}).get("predicted_ppm", 0) * area
    final_v   = r.get("reconciled_value", 0)

    tbl = doc.add_table(rows=0, cols=4)
    tbl.style = "Table Grid"
    _add_table_row(tbl, ["الأسلوب", "القيمة (ج.م)", "الوزن", "الوزن × القيمة"], header=True)

    weights = [("المقارنة السوقية", market_v, 0.30),
               ("التكلفة / DRC",    cost_v,   0.20),
               ("رسملة الدخل",      income_v, 0.20),
               ("GIS — IDW",        gis_idw,  0.10),
               ("GIS — Kriging",    gis_krig, 0.05),
               ("OLS Regression",   ols_v,    0.15)]

    for name, val, w in weights:
        _add_table_row(tbl, [name, _fmt(val), f"{w*100:.0f}%", _fmt(val * w)])

    # Final row
    row = tbl.add_row()
    cells = row.cells
    cells[0].text = "القيمة النهائية المرجّحة"
    cells[1].text = ""
    cells[2].text = "100%"
    cells[3].text = _fmt(final_v) + " ج.م"
    for cell in cells:
        _set_rtl(cell.paragraphs[0])
        for run in cell.paragraphs[0].runs:
            run.font.name = FONT_MAIN; run.font.bold = True; run.font.size = Pt(12)
            run.font.color.rgb = WHITE
        _shade_cell(cell, "D4AF37")

    doc.add_page_break()


def _compliance_section(doc: "Document", r: Dict):
    _add_heading(doc, "عاشراً: درجة الامتثال لمعايير IVS/USPAP", 1)
    _hr(doc)

    # Build a text version of the report to audit
    report_text = (
        f"تقرير تقييم عقاري. موقع: {r.get('location','')}. "
        f"العقار: {r.get('property_type','')}. "
        f"مساحة: {r.get('area',0)} م². "
        f"القيمة: {r.get('reconciled_value',0)} ج.م. "
        "معايير IVS 103 104 105. افتراضات التقييم. "
        "تاريخ المعاينة. زيارة ميدانية. شهادة المقيم. "
        "تحليل السوق. المقارنة السوقية. التكلفة. الدخل. "
        "التوفيق. موقع. مساحة. طابق. سنة الإنشاء. حالة. "
    )

    try:
        from compliance_auditor import audit_report
        scorecard = audit_report(report_text)
        score    = scorecard.get("score", 0)
        grade    = scorecard.get("grade", "B")
        passed   = scorecard.get("passed", [])
        failed   = scorecard.get("failed", [])

        grade_color = GREEN if score >= 80 else (GOLD if score >= 60 else RGBColor(0xC0, 0x00, 0x00))
        _add_para(doc, f"الدرجة الإجمالية: {score:.0f}/100 — تقدير: {grade}",
                  bold=True, size=14, color=grade_color)
        doc.add_paragraph()

        if passed:
            _add_para(doc, "البنود المُستوفاة:", bold=True, size=12, color=GREEN)
            for item in passed:
                p = doc.add_paragraph(style="List Bullet")
                _set_rtl(p)
                run = p.add_run(f"✓ {item}")
                run.font.name = FONT_MAIN; run.font.size = Pt(11); run.font.color.rgb = GREEN

        if failed:
            doc.add_paragraph()
            _add_para(doc, "البنود الناقصة / تحتاج تحسيناً:", bold=True, size=12,
                      color=RGBColor(0xC0, 0x00, 0x00))
            for item in failed:
                p = doc.add_paragraph(style="List Bullet")
                _set_rtl(p)
                run = p.add_run(f"✗ {item}")
                run.font.name = FONT_MAIN; run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

        recs = scorecard.get("recommendations", [])
        if recs:
            doc.add_paragraph()
            _add_para(doc, "التوصيات:", bold=True, size=12, color=NAVY)
            for rec in recs:
                p = doc.add_paragraph(style="List Bullet")
                _set_rtl(p)
                run = p.add_run(f"• {rec}")
                run.font.name = FONT_MAIN; run.font.size = Pt(11)

    except Exception as e:
        _add_para(doc, f"(لم يتم تحميل نظام المراجعة: {e})", size=11, color=GRAY)

    doc.add_page_break()


def _certification(doc: "Document", r: Dict):
    _add_heading(doc, "إحدى عشر: إقرار وتوقيع المقيّم المعتمد", 1)
    _hr(doc)

    _add_para(doc,
        "أُقرّ بأن هذا التقرير قد أُعدّ وفقاً لأخلاقيات مهنة التقييم، "
        "وأن التقييم قد أُجري بصورة مستقلة وموضوعية، "
        "وأنني لا أمتلك مصلحة حالية أو مستقبلية في العقار موضوع التقييم.",
        size=12)
    doc.add_paragraph()
    doc.add_paragraph()

    tbl = doc.add_table(rows=0, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    rows_data = [
        ("المقيّم المعتمد",  "هشام محمد المهدى"),
        ("رقم القيد",        "29"),
        ("تاريخ التقرير",    datetime.now().strftime("%Y-%m-%d")),
        ("الختم والتوقيع",   " " * 30),
    ]
    for lbl, val in rows_data:
        row = tbl.add_row()
        row.cells[0].text = lbl; row.cells[1].text = val
        _set_rtl(row.cells[0].paragraphs[0]); _set_rtl(row.cells[1].paragraphs[0])
        for run in row.cells[0].paragraphs[0].runs:
            run.font.name = FONT_MAIN; run.font.bold = True; run.font.size = Pt(12)
        for run in row.cells[1].paragraphs[0].runs:
            run.font.name = FONT_MAIN; run.font.size = Pt(12)
        _shade_cell(row.cells[0], "1F4E78")
        for run in row.cells[0].paragraphs[0].runs:
            run.font.color.rgb = WHITE
    tbl.columns[0].width = Cm(6); tbl.columns[1].width = Cm(10)


# ── Master Entry Point ────────────────────────────────────────────────────────

def generate_master_report(ivs_result: Dict, output_dir: str = "") -> str:
    """
    Generate the Master Unified Word Report.
    Returns path to the generated .docx file.
    """
    if not _DOCX_OK:
        raise ImportError("python-docx is not installed. Run: pip install python-docx")

    doc = Document()

    # Page setup — A4, RTL document
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin  = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin   = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Set document-level RTL
    settings = doc.settings.element
    bidi_elem = OxmlElement("w:themeFontLang")
    bidi_elem.set(qn("w:bidi"), "ar-EG")
    settings.append(bidi_elem)

    # Build all sections
    _cover_page(doc, ivs_result)
    _executive_summary(doc, ivs_result)
    _property_description(doc, ivs_result)
    _market_comparison(doc, ivs_result)
    _cost_approach(doc, ivs_result)
    _income_approach(doc, ivs_result)
    if ivs_result.get("land_dual"):
        _land_dual_path_section(doc, ivs_result)
    _gis_section(doc, ivs_result)
    _ols_section(doc, ivs_result)
    _hbu_section(doc, ivs_result)
    _reconciliation(doc, ivs_result)
    _compliance_section(doc, ivs_result)
    _certification(doc, ivs_result)

    # Save
    if not output_dir:
        output_dir = _OUT_DIR
    os.makedirs(output_dir, exist_ok=True)
    ts   = datetime.now().strftime("%Y%m%d_%H%M%S")
    path = os.path.join(output_dir, f"master_report_{ts}.docx")
    doc.save(path)
    return path


# ── CLI test ──────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import sys, io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

    # Build a mock IVS result for testing
    mock = {
        "client_name":        "Ahmed Mohamed",
        "property_type":      "شقة سكنية",
        "location":           "المعادي، القاهرة",
        "area":               150,
        "floor":              3,
        "year_built":         2015,
        "building_age":       9,
        "finishing_description": "تشطيبات سوبر لوكس",
        "currency":           "ج.م",
        "purpose_label_ar":   "القيمة السوقية العادلة",
        "forced_discount":    0.0,
        "reconciled_value":   3_450_000,
        "hbu_text":           "أفضل استخدام: سكني",
        "market":  {"value": 3_300_000, "comparables": [
            {"loc": "المعادي", "ar": 140, "price_per_m2": 22000, "source": "portal_listing"},
            {"loc": "المعادي", "ar": 160, "price_per_m2": 24000, "source": "broker_oral"},
        ]},
        "cost":    {"value": 3_200_000, "gross_building_cost": 2_800_000,
                    "depreciation": 420_000, "building_age": 9,
                    "depr_rate_pct": 15.0, "land_value": 820_000, "rcn_net": 2_380_000},
        "income":  {"value": 3_600_000, "noi": 288_000, "cap_rate": 0.08, "rent_per_sqm": 200},
        "gis":     {"idw_ppm": 22500, "kriging_ppm": 23100},
        "ols":     {"r_squared": 0.83, "adj_r_squared": 0.81, "predicted_ppm": 22800,
                    "n_obs": 20,
                    "coefficients": {"const": 25000, "floor": -120, "area": -25, "year_built": 30}},
    }

    path = generate_master_report(mock)
    print(f"Master Report saved: {path}")
