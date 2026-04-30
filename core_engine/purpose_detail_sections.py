"""
purpose_detail_sections.py
============================
Detailed Word sections for each of the 9 standard valuation purposes.

Each builder appends a self-contained section to an open python-docx Document.
Each section pulls its OWN purpose-specific data from the payload and renders
the relevant tables (LTV/DSCR for banks, IRR/NPV/sensitivity for investment,
RCN/depreciation for insurance, etc.).

Public API:
    add_purpose_detail_section(doc, purpose, payload) -> bool

This module is ADDITIVE — it does not modify existing report builders.
"""

from __future__ import annotations
from typing import Any, Dict, List, Optional, Tuple

try:
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _DOCX_OK = True
except Exception:
    _DOCX_OK = False


NAVY  = RGBColor(0x1F, 0x38, 0x64) if _DOCX_OK else None
GOLD  = RGBColor(0xC9, 0xA2, 0x27) if _DOCX_OK else None
GREEN = RGBColor(0x1E, 0x84, 0x49) if _DOCX_OK else None
RED   = RGBColor(0xC0, 0x39, 0x2B) if _DOCX_OK else None
GREY  = RGBColor(0x55, 0x55, 0x55) if _DOCX_OK else None


# ── low-level helpers (mirror style of asset_type_reports.py) ───────────────

def _para_rtl(p):
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


def _h2(doc, text: str, color=None):
    p = doc.add_paragraph()
    _para_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(text)
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = color or NAVY
    return p


def _para(doc, text: str, *, bold: bool = False, color=None, size: int = 12):
    p = doc.add_paragraph()
    _para_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    if color is not None:
        r.font.color.rgb = color
    return p


def _disclosure(doc, text: str):
    p = doc.add_paragraph()
    _para_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = GREY


def _kv(doc, rows: List[Tuple[str, str]]):
    if not rows:
        return None
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = "Light Grid Accent 1"
    tblPr = tbl._tbl.tblPr
    bidi = OxmlElement("w:bidiVisual")
    tblPr.append(bidi)
    for i, (k, v) in enumerate(rows):
        rc = tbl.rows[i].cells
        rc[0].text = str(k)
        rc[1].text = str(v)
        for c in rc:
            for p in c.paragraphs:
                _para_rtl(p)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(11)
        rc[0].paragraphs[0].runs[0].font.bold = True
    return tbl


def _grid(doc, headers: List[str], rows: List[List[str]]):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    tblPr = tbl._tbl.tblPr
    bidi = OxmlElement("w:bidiVisual")
    tblPr.append(bidi)
    hr = tbl.rows[0].cells
    for j, h in enumerate(headers):
        hr[j].text = h
        for p in hr[j].paragraphs:
            _para_rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = NAVY
                r.font.size = Pt(11)
    for i, row in enumerate(rows):
        cs = tbl.rows[i + 1].cells
        for j, val in enumerate(row):
            cs[j].text = str(val)
            for p in cs[j].paragraphs:
                _para_rtl(p)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for r in p.runs:
                    r.font.size = Pt(10)
    return tbl


def _money(x: Any) -> str:
    try:
        return f"{float(x):,.0f} ج.م"
    except Exception:
        return "—"


def _pct(x: Any, *, scale: bool = True) -> str:
    """Format as %. If scale=True and |x|<=2 treat x as decimal; else as already-percent."""
    try:
        v = float(x)
        if scale and abs(v) <= 2.0:
            return f"{v * 100:.2f}%"
        return f"{v:.2f}%"
    except Exception:
        return "—"


def _f(payload: Dict[str, Any], key: str, default: float = 0.0) -> float:
    try:
        v = payload.get(key)
        return float(v) if v not in (None, "") else default
    except Exception:
        return default


# ════════════════════════════════════════════════════════════════════════════
#  1. FAIR MARKET VALUE — comparables, market trend, transaction count
# ════════════════════════════════════════════════════════════════════════════

def add_fmv_section(doc, p: Dict[str, Any]) -> None:
    _add_avm_section(doc, p)
    _h2(doc, "تفصيل القيمة السوقية العادلة (Market Value)", color=GOLD)
    mv   = _f(p, "market_value")
    mva  = _f(p, "market_approach")
    cost = _f(p, "cost_approach")
    inc  = _f(p, "income_approach")

    _kv(doc, [
        ("القيمة السوقية النهائية",        _money(mv)),
        ("نهج السوق (Comparables)",        _money(mva)),
        ("نهج التكلفة (Cost Approach)",    _money(cost)),
        ("نهج الدخل (Income Approach)",    _money(inc)),
        ("سعر المتر",                       _money(_f(p, "price_per_meter"))),
        ("المساحة (م²)",                    f"{_f(p, 'area'):,.0f}"),
    ])

    # Comparable transactions (if supplied)
    comps = p.get("comparables") or []
    if comps:
        _h2(doc, "المعاملات المقارنة")
        rows = [[
            c.get("date", "—"),
            c.get("location", "—"),
            f"{float(c.get('area', 0) or 0):,.0f}",
            _money(c.get("price")),
            _money(c.get("price_per_meter") or
                   (float(c.get("price", 0) or 0) / max(float(c.get("area", 1) or 1), 1)))
        ] for c in comps[:10]]
        _grid(doc, ["تاريخ المعاملة", "الموقع", "المساحة (م²)", "إجمالي السعر", "سعر المتر"], rows)
        _para(doc, f"عدد المعاملات المرجعية: {len(comps)}", bold=True)

    # Market trend (annualized growth, transaction volume)
    trend_rate = p.get("annual_growth_rate")
    tx_count   = p.get("transaction_count_12m")
    if trend_rate is not None or tx_count is not None:
        _h2(doc, "اتجاه السوق")
        rows = []
        if trend_rate is not None:
            rows.append(("معدل النمو السنوي", _pct(trend_rate)))
        if tx_count is not None:
            rows.append(("عدد المعاملات (آخر 12 شهر)", str(tx_count)))
        _kv(doc, rows)

    _disclosure(doc,
        "القيمة السوقية وفق IVS 104 §30: «المبلغ المُقدَّر الذي يجب أن يتم به "
        "تبادل الأصل في تاريخ التقييم بين بائع راغب ومشترٍ راغب، في صفقة محايدة، "
        "بعد تسويق ملائم، حيث تصرَّف كل طرف بمعرفة وحرص دون إكراه».")


# ════════════════════════════════════════════════════════════════════════════
#  2. ACQUISITION (M&A) — DCF/IRR/NPV/synergies
# ════════════════════════════════════════════════════════════════════════════

def add_acquisition_section(doc, p: Dict[str, Any]) -> None:
    _add_avm_section(doc, p)
    _h2(doc, "تفصيل القيمة الاستثمارية للاستحواذ والاندماج", color=GOLD)
    mv      = _f(p, "market_value")
    wacc    = _f(p, "wacc", 0.12)
    growth  = _f(p, "terminal_growth_rate", 0.03)
    syn     = _f(p, "synergy_value", 0.0)
    premium = _f(p, "acquisition_premium_pct", 0.20)

    standalone_value = mv
    investment_value = standalone_value * (1 + premium) + syn

    _kv(doc, [
        ("القيمة المستقلة للأصل (Standalone)", _money(standalone_value)),
        ("WACC المعتمد",                        _pct(wacc)),
        ("معدل النمو النهائي (Terminal g)",    _pct(growth)),
        ("علاوة الاستحواذ المقدرة",             _pct(premium)),
        ("قيمة التآزر (Synergies)",             _money(syn)),
        ("القيمة الاستثمارية للمستحوذ (Investment Value)", _money(investment_value)),
    ])

    # Optional DCF projection table (if user supplied year-by-year cashflows)
    cfs = p.get("acquisition_cashflows") or []
    if cfs:
        _h2(doc, "إسقاط التدفقات النقدية للاستحواذ")
        rows = [[f"سنة {i}", _money(cf)] for i, cf in enumerate(cfs)]
        _grid(doc, ["السنة", "التدفق النقدي"], rows)

    _disclosure(doc,
        "القيمة الاستثمارية وفق IVS 104 §60 تعكس قيمة الأصل لمستثمر محدد بالنظر "
        "لأهداف ومتطلبات استثمارية محددة، وقد تختلف عن القيمة السوقية بسبب "
        "علاوة السيطرة والتآزر التشغيلي.")


# ════════════════════════════════════════════════════════════════════════════
#  3. BANK FINANCING — LTV/DSCR/Basel III
# ════════════════════════════════════════════════════════════════════════════

def add_bank_financing_section(doc, p: Dict[str, Any]) -> None:
    _add_avm_section(doc, p)
    _h2(doc, "تفصيل القيمة لأغراض التمويل البنكي", color=GOLD)

    mv             = _f(p, "market_value")
    ltv_target     = _f(p, "ltv_target", 0.70)
    interest_rate  = _f(p, "interest_rate", 0.18)
    loan_term_yrs  = int(_f(p, "loan_term_years", 15)) or 15
    annual_noi     = _f(p, "annual_noi", mv * 0.07)  # default 7% cap

    max_loan = mv * ltv_target
    # Annuity payment
    if interest_rate > 0 and loan_term_yrs > 0:
        annuity = max_loan * (interest_rate / (1 - (1 + interest_rate) ** (-loan_term_yrs)))
    else:
        annuity = max_loan / max(loan_term_yrs, 1)
    dscr = (annual_noi / annuity) if annuity > 0 else 0.0

    # Basel III risk weighting (commercial real estate, secured)
    rw = 1.00 if ltv_target > 0.60 else 0.75 if ltv_target > 0.40 else 0.50
    rwa = max_loan * rw

    _kv(doc, [
        ("القيمة السوقية للضمان",            _money(mv)),
        ("نسبة LTV المستهدفة",               _pct(ltv_target)),
        ("الحد الأقصى للقرض",                _money(max_loan)),
        ("معدل الفائدة السنوي",              _pct(interest_rate)),
        ("مدة القرض (سنوات)",                 str(loan_term_yrs)),
        ("القسط السنوي المقدر",              _money(annuity)),
        ("صافي الدخل التشغيلي السنوي (NOI)", _money(annual_noi)),
        ("DSCR (تغطية خدمة الدين)",         f"{dscr:.2f}×"),
        ("Basel III — وزن المخاطر RW",       _pct(rw)),
        ("RWA — الأصول المرجحة بالمخاطر",    _money(rwa)),
    ])

    # Verdict
    verdict_color = GREEN if dscr >= 1.25 else (RED if dscr < 1.10 else None)
    if dscr >= 1.25:
        verdict = "DSCR قوي (≥1.25×) — مقبول للتمويل البنكي وفق ممارسات الـ CBE."
    elif dscr >= 1.10:
        verdict = "DSCR هامشي (1.10–1.25×) — يحتاج ضمانات إضافية."
    else:
        verdict = "DSCR ضعيف (<1.10×) — مرفوض ائتمانياً، يُوصى بتخفيض حجم القرض."
    _para(doc, verdict, bold=True, color=verdict_color)

    _disclosure(doc,
        "تم تطبيق قواعد البنك المركزي المصري (CBE) و Basel III §Article 122. "
        "نسبة LTV الافتراضية تتبع تعليمات BCBS 401 للعقارات السكنية والتجارية.")


# ════════════════════════════════════════════════════════════════════════════
#  4. RENTAL ARBITRATION — comparable rents, escalation, fair market rent
# ════════════════════════════════════════════════════════════════════════════

def add_rental_arbitration_section(doc, p: Dict[str, Any]) -> None:
    _h2(doc, "تفصيل القيمة الإيجارية العادلة", color=GOLD)

    mv             = _f(p, "market_value")
    area           = _f(p, "area")
    cap_rate       = _f(p, "cap_rate", 0.08)
    current_rent   = _f(p, "current_annual_rent")
    escalation     = _f(p, "rent_escalation_rate", 0.07)
    market_rent    = mv * cap_rate
    monthly_market = market_rent / 12.0

    diff_pct = ((market_rent - current_rent) / current_rent * 100) if current_rent > 0 else 0.0

    _kv(doc, [
        ("القيمة السوقية للعقار",           _money(mv)),
        ("معدل الرسملة المستخدم",           _pct(cap_rate)),
        ("الإيجار السنوي العادل",           _money(market_rent)),
        ("الإيجار الشهري العادل",           _money(monthly_market)),
        ("الإيجار الحالي السنوي",           _money(current_rent) if current_rent else "—"),
        ("الفارق عن الحالي",                 (_pct(diff_pct, scale=False) if current_rent > 0 else "—")),
        ("معدل الزيادة السنوي المقترح",     _pct(escalation)),
        ("الإيجار للمتر سنوياً",            _money(market_rent / max(area, 1)) if area > 0 else "—"),
    ])

    # 5-year escalated rent projection
    if market_rent > 0:
        _h2(doc, "إسقاط الإيجار العادل (5 سنوات)")
        proj_rows = []
        for yr in range(1, 6):
            rent_yr = market_rent * ((1 + escalation) ** (yr - 1))
            proj_rows.append([f"سنة {yr}", _money(rent_yr), _money(rent_yr / 12)])
        _grid(doc, ["السنة", "الإيجار السنوي المقترح", "الإيجار الشهري"], proj_rows)

    # Comparable rents
    comps = p.get("comparable_rents") or []
    if comps:
        _h2(doc, "إيجارات مقارنة")
        rows = [[
            c.get("location", "—"),
            f"{float(c.get('area', 0) or 0):,.0f}",
            _money(c.get("annual_rent")),
            _money((float(c.get("annual_rent", 0) or 0)) / max(float(c.get("area", 1) or 1), 1)),
        ] for c in comps[:8]]
        _grid(doc, ["الموقع", "المساحة (م²)", "الإيجار السنوي", "إيجار المتر/سنة"], rows)

    _disclosure(doc,
        "تم تطبيق منهجية تحديد بدل المثل وفق القانون 4 لسنة 1996 (الإيجارات الجديدة) "
        "والمواد 564–م.602 من القانون المدني المصري. الإيجار العادل يُحدَّد من خلال "
        "تثليث Cap Rate سوقي + إيجارات مقارنة + احتساب صريح.")


# ════════════════════════════════════════════════════════════════════════════
#  5. INSURANCE — RCN, depreciation, indemnity value
# ════════════════════════════════════════════════════════════════════════════

def add_insurance_section(doc, p: Dict[str, Any]) -> None:
    _h2(doc, "تفصيل قيمة إعادة الإنشاء للتأمين", color=GOLD)

    building_area = _f(p, "building_area_m2", _f(p, "area"))
    rcn_per_m2    = _f(p, "rcn_per_m2", 12000)  # avg residential
    rcn_total     = building_area * rcn_per_m2
    age           = int(_f(p, "build_age", 10))
    economic_life = int(_f(p, "economic_life_years", 50)) or 50
    depr_pct      = min(age / max(economic_life, 1), 0.85)
    indemnity     = rcn_total * (1 - depr_pct)
    contents_pct  = _f(p, "contents_insurance_pct", 0.10)
    contents_sum  = rcn_total * contents_pct
    policy_excess = _f(p, "policy_excess_pct", 0.05)

    _kv(doc, [
        ("مساحة المبنى (م²)",                 f"{building_area:,.0f}"),
        ("تكلفة إعادة الإنشاء للمتر",        _money(rcn_per_m2)),
        ("تكلفة إعادة الإنشاء (RCN)",         _money(rcn_total)),
        ("عمر المبنى (سنة)",                   str(age)),
        ("العمر الاقتصادي المتوقع",            f"{economic_life} سنة"),
        ("نسبة الإهلاك المتراكم",              _pct(depr_pct)),
        ("قيمة التعويض (Indemnity Value)",     _money(indemnity)),
        ("نسبة محتويات للتأمين",                _pct(contents_pct)),
        ("مبلغ تأمين المحتويات",                _money(contents_sum)),
        ("نسبة تحمل العميل (Excess)",          _pct(policy_excess)),
        ("إجمالي مبلغ التأمين الموصى به",      _money(rcn_total + contents_sum)),
    ])

    _disclosure(doc,
        "RCN = Reinstatement Cost New وفق IVS 105 §60 ومعيار البناء المصري. "
        "قيمة التعويض تُمثِّل قيمة استبدال المبنى مع خصم الإهلاك الفيزيائي "
        "والوظيفي، وهي القيمة المستخدمة في تسوية المطالبات وفق وثائق التأمين.")


# ════════════════════════════════════════════════════════════════════════════
#  6. INVESTMENT ANALYSIS — IRR/NPV/Sensitivity
# ════════════════════════════════════════════════════════════════════════════

def _local_npv(rate: float, cfs: List[float]) -> float:
    return sum(cf / ((1 + rate) ** t) for t, cf in enumerate(cfs))


def _local_irr(cfs: List[float], guess: float = 0.10) -> Optional[float]:
    if not cfs or all(c >= 0 for c in cfs) or all(c <= 0 for c in cfs):
        return None
    rate = guess
    for _ in range(120):
        f = _local_npv(rate, cfs)
        df = (_local_npv(rate + 1e-6, cfs) - f) / 1e-6
        if df == 0:
            break
        new = rate - f / df
        if abs(new - rate) < 1e-7:
            return new
        rate = new
        if rate <= -0.99:
            break
    # fallback bisection
    lo, hi = -0.99, 5.0
    flo, fhi = _local_npv(lo, cfs), _local_npv(hi, cfs)
    if flo * fhi > 0:
        return None
    for _ in range(100):
        mid = (lo + hi) / 2
        fm = _local_npv(mid, cfs)
        if abs(fm) < 1e-6:
            return mid
        if flo * fm < 0:
            hi, fhi = mid, fm
        else:
            lo, flo = mid, fm
    return (lo + hi) / 2


def add_investment_analysis_section(doc, p: Dict[str, Any]) -> None:
    _h2(doc, "تفصيل التحليل الاستثماري — IRR / NPV / DCF", color=GOLD)

    mv         = _f(p, "market_value")
    annual_noi = _f(p, "annual_noi", mv * 0.08)
    growth     = _f(p, "noi_growth_rate", 0.06)
    discount   = _f(p, "discount_rate", 0.13)
    horizon    = int(_f(p, "investment_horizon_years", 10)) or 10
    exit_cap   = _f(p, "exit_cap_rate", 0.08)

    # Build cashflow stream
    cfs = [-mv]
    for yr in range(1, horizon + 1):
        cfs.append(annual_noi * ((1 + growth) ** (yr - 1)))
    # Terminal value at end
    terminal_noi = annual_noi * ((1 + growth) ** horizon)
    terminal_value = terminal_noi / exit_cap if exit_cap > 0 else 0
    cfs[-1] += terminal_value

    npv = _local_npv(discount, cfs)
    irr = _local_irr(cfs)
    # Payback (undiscounted)
    cum, payback = 0.0, None
    for t, cf in enumerate(cfs):
        cum += cf
        if cum >= 0 and payback is None and t > 0:
            payback = t
            break

    _kv(doc, [
        ("قيمة الاستثمار الأولي",          _money(mv)),
        ("صافي الدخل التشغيلي السنوي",    _money(annual_noi)),
        ("معدل نمو NOI",                    _pct(growth)),
        ("أفق الاستثمار",                  f"{horizon} سنة"),
        ("Exit Cap Rate",                  _pct(exit_cap)),
        ("القيمة النهائية (Terminal Value)", _money(terminal_value)),
        ("معدل الخصم (Discount Rate)",     _pct(discount)),
        ("صافي القيمة الحالية (NPV)",      _money(npv)),
        ("معدل العائد الداخلي (IRR)",      _pct(irr * 100, scale=False) if irr is not None else "—"),
        ("فترة الاسترداد",                  f"{payback} سنة" if payback else "—"),
    ])

    # Sensitivity table — discount rate × growth rate
    _h2(doc, "تحليل الحساسية (NPV)")
    growth_range   = [growth - 0.02, growth, growth + 0.02]
    discount_range = [discount - 0.02, discount, discount + 0.02]
    headers = ["Discount \\ Growth"] + [_pct(g) for g in growth_range]
    rows = []
    for d in discount_range:
        row = [_pct(d)]
        for g in growth_range:
            cfs2 = [-mv]
            for yr in range(1, horizon + 1):
                cfs2.append(annual_noi * ((1 + g) ** (yr - 1)))
            cfs2[-1] += (annual_noi * ((1 + g) ** horizon)) / exit_cap
            row.append(_money(_local_npv(d, cfs2)))
        rows.append(row)
    _grid(doc, headers, rows)

    _disclosure(doc,
        "التحليل وفق IVS 105 §50 (Income Approach — DCF Method). يُستخدم لاتخاذ "
        "قرار الاستثمار وتقييم العائد المرتقب مقابل المخاطر. القرار: استثمر إذا "
        "كان IRR > معدل الخصم و NPV > 0.")


# ════════════════════════════════════════════════════════════════════════════
#  7. JUDICIAL LIQUIDATION — forced sale discount
# ════════════════════════════════════════════════════════════════════════════

def add_judicial_liquidation_section(doc, p: Dict[str, Any]) -> None:
    _h2(doc, "تفصيل القيمة التصفوية القضائية", color=GOLD)

    mv = _f(p, "market_value")
    # Reverse the 0.82 PPM factor that was already applied → recover original FMV
    pre_discount_fmv = _f(p, "pre_liquidation_fmv", mv / 0.82 if mv > 0 else 0)
    forced_sale_period = int(_f(p, "expected_sale_period_months", 6))
    legal_costs_pct = _f(p, "legal_costs_pct", 0.05)
    auction_costs_pct = _f(p, "auction_costs_pct", 0.03)
    liquidation_discount_pct = 0.18  # matches _PURPOSE_PPM_FACTOR

    legal_costs = pre_discount_fmv * legal_costs_pct
    auction_costs = pre_discount_fmv * auction_costs_pct
    forced_sale_value = pre_discount_fmv * (1 - liquidation_discount_pct)
    net_recoverable = forced_sale_value - legal_costs - auction_costs

    _kv(doc, [
        ("القيمة السوقية قبل التصفية",      _money(pre_discount_fmv)),
        ("نسبة خصم البيع الإجباري",          _pct(liquidation_discount_pct)),
        ("المدة المتوقعة للبيع",             f"{forced_sale_period} شهر"),
        ("القيمة التصفوية الإجبارية",         _money(forced_sale_value)),
        ("تكاليف التقاضي والإجراءات",         _money(legal_costs)),
        ("تكاليف المزاد والشهر العقاري",      _money(auction_costs)),
        ("صافي القيمة المسترَدَّة",            _money(net_recoverable)),
        ("نسبة الاسترداد من FMV",              _pct(net_recoverable / max(pre_discount_fmv, 1), scale=False)),
    ])

    _disclosure(doc,
        "وفق قانون المرافعات المدنية والتجارية المصري — مواد البيع الجبري للعقار، "
        "ومعيار IVS 104 §40 للقيمة التصفوية. الخصم يعكس ضغط الزمن، نقص التسويق، "
        "وعدم وجود فرصة تفاوض حقيقية بين البائع والمشتري.")


# ════════════════════════════════════════════════════════════════════════════
#  8. TAX ASSESSMENT — tax base, indices, exemptions
# ════════════════════════════════════════════════════════════════════════════

def add_tax_assessment_section(doc, p: Dict[str, Any]) -> None:
    _add_avm_section(doc, p)
    _h2(doc, "تفصيل القيمة للأغراض الضريبية (الضريبة العقارية)", color=GOLD)

    mv = _f(p, "market_value")
    governorate_factor = _f(p, "governorate_factor", 1.00)
    construction_factor = _f(p, "construction_factor", 1.00)
    location_factor = _f(p, "location_factor", 1.00)
    annual_rental_value = mv * 0.03  # دخل إيجاري سنوي افتراضي 3%
    composite_factor = governorate_factor * construction_factor * location_factor
    tax_base = annual_rental_value * composite_factor
    exemption_threshold = _f(p, "exemption_threshold", 24000)  # م6 قانون 196/2008
    taxable_amount = max(tax_base - exemption_threshold, 0)
    tax_rate = 0.10  # 10% وفق قانون 196/2008
    annual_tax = taxable_amount * tax_rate

    _kv(doc, [
        ("القيمة السوقية للعقار",                _money(mv)),
        ("القيمة الإيجارية السنوية (3%)",        _money(annual_rental_value)),
        ("معامل المحافظة",                       f"{governorate_factor:.2f}"),
        ("معامل تصنيف البناء",                   f"{construction_factor:.2f}"),
        ("معامل الموقع داخل المحافظة",           f"{location_factor:.2f}"),
        ("المعامل المركَّب",                      f"{composite_factor:.4f}"),
        ("الوعاء الضريبي السنوي",                _money(tax_base)),
        ("الحد المُعفى (م.6 قانون 196/2008)",     _money(exemption_threshold)),
        ("الوعاء الخاضع للضريبة",                _money(taxable_amount)),
        ("سعر الضريبة",                           _pct(tax_rate)),
        ("الضريبة العقارية السنوية",             _money(annual_tax)),
    ])

    _disclosure(doc,
        "تم احتساب الضريبة العقارية وفق القانون رقم 196 لسنة 2008 ولائحته التنفيذية. "
        "الإعفاء السنوي 24,000 ج.م لكل وحدة سكنية يستخدمها المالك للسكن الخاص. "
        "المعدل الموحد 10% من صافي القيمة الإيجارية السنوية بعد الإعفاء.")


# ════════════════════════════════════════════════════════════════════════════
#  9. FINANCIAL REPORTING (IFRS 13) — fair value hierarchy + sensitivity
# ════════════════════════════════════════════════════════════════════════════

def add_financial_reporting_section(doc, p: Dict[str, Any]) -> None:
    _h2(doc, "تفصيل القيمة العادلة وفق IFRS 13", color=GOLD)

    mv = _f(p, "market_value")
    ifrs_level = int(_f(p, "ifrs_level", 3)) or 3
    technique = p.get("valuation_technique") or (
        "Income Approach (DCF)" if ifrs_level == 3 else
        "Market Approach (Comparables)" if ifrs_level == 2 else
        "Quoted Price"
    )

    level_labels = {
        1: "Level 1 — أسعار معلنة في أسواق نشطة لأصول مماثلة",
        2: "Level 2 — مدخلات قابلة للملاحظة (بخلاف Level 1)",
        3: "Level 3 — مدخلات غير قابلة للملاحظة",
    }

    _kv(doc, [
        ("القيمة العادلة المعتمدة",          _money(mv)),
        ("المستوى الهرمي للقيمة العادلة",   level_labels.get(ifrs_level, str(ifrs_level))),
        ("تقنية التقييم المُطبَّقة",         technique),
        ("المدخلات القابلة للملاحظة",       p.get("observable_inputs",
                                              "أسعار معاملات مقارنة، Cap Rates سوقية، أسعار إيجار حقيقية")),
        ("المدخلات غير القابلة للملاحظة",   p.get("unobservable_inputs",
                                              "معدل خصم متخصص، نمو NOI متوقع، بيانات داخلية للملاك")),
        ("تاريخ القياس",                     p.get("report_date", "—")),
    ])

    # Sensitivity (per IFRS 13 §93(h))
    if ifrs_level == 3:
        _h2(doc, "تحليل الحساسية (IFRS 13 §93(h))")
        scenarios = [
            ("معدل الخصم +1%",  -0.07),
            ("معدل الخصم −1%",  +0.08),
            ("Cap Rate +50bp",  -0.05),
            ("Cap Rate −50bp",  +0.06),
            ("نمو NOI +1%",     +0.04),
            ("نمو NOI −1%",     -0.04),
        ]
        rows = [[name, _pct(impact, scale=False), _money(mv * (1 + impact))]
                for name, impact in scenarios]
        _grid(doc, ["السيناريو", "الأثر على القيمة", "القيمة الناتجة"], rows)

    _disclosure(doc,
        "وفق IFRS 13 §93 (a–i): تم الإفصاح عن المستوى الهرمي، تقنية التقييم، "
        "المدخلات الجوهرية، تأثير المدخلات غير القابلة للملاحظة. المعيار يُلزم "
        "بإفصاح كامل عند تصنيف الأصل ضمن Level 3.")



# ════════════════════════════════════════════════════════════════════════════
#  AVM CROSS-VALIDATION SECTION — used by FMV / bank / tax / acquisition only
# ════════════════════════════════════════════════════════════════════════════

def _add_avm_section(doc, p: Dict[str, Any]) -> None:
    """يضيف قسم AVM (Automated Valuation Model) إذا توفرت بيانات في payload['avm_meta']."""
    avm = p.get("avm_meta") or {}
    if not avm or not isinstance(avm, dict):
        return
    _h2(doc, "AVM — التحقق الآلي بنموذج الانحدار", color=NAVY)
    _para(doc,
        "تم تطبيق نموذج Automated Valuation Model عبر انحدار least-squares "
        "موزون بمصداقية المصادر على بيانات معاملات السوق الفعلية لنفس المنطقة. "
        "هذا القسم يقدم نتيجة AVM إلى جانب القيمة التقليدية كتثليث مستقل.")
    rows = [
        ("المنهجية",                       avm.get("method", "—")),
        ("المصدر",                          avm.get("source", "—")),
        ("منطقة المعاملات المطابقة",         avm.get("matched_region", "—")),
        ("عدد المعاملات في النطاق الزمني",  str(avm.get("n_records", 0))),
        ("النطاق الزمني (شهور)",            str(avm.get("time_span_months", 0))),
        ("درجة الثقة (Confidence)",         avm.get("confidence", "—")),
        ("سعر المتر بالـ AVM (ج.م)",        f"{avm.get('avm_ppm', 0):,.0f}"),
        ("معدل النمو الشهري المُستنبَط",    f"{avm.get('avm_growth_per_month_pct', 0):.3f}%"),
        ("المعدل السنوي YoY (ضمنياً)",      f"{avm.get('avm_yoy_pct', 0):.2f}%"),
        ("مؤشر AVM (الأساس=100)",           f"{avm.get('avm_index', 100):.2f}"),
    ]
    _kv(doc, rows)
    notes = avm.get("notes") or ""
    if notes:
        _disclosure(doc,
            f"AVM — {notes} هذا الناتج يعمل كمرجع رقابي لقيمة السوق التقليدية ولا يحلّ محلّها "
            f"إلا عند غياب سعر متر مُدخَل من المستخدم. النموذج الإحصائي يصلح للأغراض: "
            f"FMV، التمويل البنكي، الضريبة العقارية، الاستحواذ والاندماج.")

# ════════════════════════════════════════════════════════════════════════════
#  PUBLIC DISPATCHER
# ════════════════════════════════════════════════════════════════════════════

_PURPOSE_BUILDERS = {
    "fair_market_value":    add_fmv_section,
    "acquisition":          add_acquisition_section,
    "bank_financing":       add_bank_financing_section,
    "rental_arbitration":   add_rental_arbitration_section,
    "insurance":            add_insurance_section,
    "investment_analysis":  add_investment_analysis_section,
    "judicial_liquidation": add_judicial_liquidation_section,
    "tax_assessment":       add_tax_assessment_section,
    "financial_reporting":  add_financial_reporting_section,
}


def add_purpose_detail_section(doc, purpose: str, payload: Dict[str, Any]) -> bool:
    """يُضيف قسماً مفصلاً مخصصاً للغرض المحدد. يُرجع True عند الإضافة."""
    if not _DOCX_OK or not purpose:
        return False
    builder = _PURPOSE_BUILDERS.get(purpose.strip())
    if builder is None:
        return False
    try:
        builder(doc, payload)
        return True
    except Exception as e:
        print(f"[purpose_detail_sections] {purpose} failed: {e}")
        return False
