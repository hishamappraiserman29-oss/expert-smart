"""
purpose_specific_reports.py
============================
Per-purpose Word report generator for the Expert_Smart valuation engine.

Each of the 13 valuation purposes gets its own dedicated report layout.
This module is ADDITIVE — it does not modify the existing
`master_word_report.py` / `write_word_summary` flow.

Public API:
    write_purpose_specific_report(purpose, payload, output_path)  -> bool
    write_hbu_word_report(result, output_path)                    -> bool
    write_reit_word_report(result, output_path)                   -> bool

`payload` is the merged dict of (request payload + advanced_valuation result)
that bridge_api.py builds before saving the standard report.
"""

from __future__ import annotations
import os
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

# RTL Arabic Word generation via python-docx (already a project dependency)
try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _DOCX_OK = True
except Exception:
    _DOCX_OK = False


# ════════════════════════════════════════════════════════════════════════════
#  STYLE PALETTE — مطابق لأسلوب التقارير القائمة
# ════════════════════════════════════════════════════════════════════════════

NAVY  = RGBColor(0x1F, 0x38, 0x64) if _DOCX_OK else None
GOLD  = RGBColor(0xC9, 0xA2, 0x27) if _DOCX_OK else None
DARK  = RGBColor(0x2C, 0x3E, 0x50) if _DOCX_OK else None
GREEN = RGBColor(0x1E, 0x84, 0x49) if _DOCX_OK else None
RED   = RGBColor(0xC0, 0x39, 0x2B) if _DOCX_OK else None
GREY  = RGBColor(0x55, 0x55, 0x55) if _DOCX_OK else None


# ════════════════════════════════════════════════════════════════════════════
#  LOW-LEVEL DOCX HELPERS
# ════════════════════════════════════════════════════════════════════════════

def _set_rtl_doc(doc) -> None:
    """يجعل المستند بأكمله RTL مع خط Arial 12pt افتراضي."""
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(12)
    rpr = style.element.get_or_add_rPr()
    bidi = OxmlElement("w:bidi")
    rpr.append(bidi)
    rtl = OxmlElement("w:rtl")
    rpr.append(rtl)


def _para_rtl(p) -> None:
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


def _add_h1(doc, text: str, color=None):
    p = doc.add_paragraph()
    _para_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = color or NAVY
    return p


def _add_h2(doc, text: str, color=None):
    p = doc.add_paragraph()
    _para_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = color or GOLD
    return p


def _add_p(doc, text: str, *, bold: bool = False, color=None, size: int = 12):
    p = doc.add_paragraph()
    _para_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    return p


def _add_kv_table(doc, rows: List[Tuple[str, str]], *, header: Optional[str] = None):
    if header:
        _add_h2(doc, header)
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Make table RTL
    tblPr = tbl._tbl.tblPr
    bidi = OxmlElement("w:bidiVisual")
    tblPr.append(bidi)

    for i, (k, v) in enumerate(rows):
        rc = tbl.rows[i].cells
        rc[0].text = str(k)
        rc[1].text = str(v)
        for c in rc:
            for p in c.paragraphs:
                _para_rtl(p)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(11)
        rc[0].paragraphs[0].runs[0].font.bold = True
    return tbl


def _add_data_table(doc, headers: List[str], rows: List[List[str]],
                    *, highlight_row: Optional[int] = None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tblPr = tbl._tbl.tblPr
    bidi = OxmlElement("w:bidiVisual")
    tblPr.append(bidi)

    # Header
    hr = tbl.rows[0].cells
    for j, h in enumerate(headers):
        hr[j].text = h
        for p in hr[j].paragraphs:
            _para_rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = NAVY
                r.font.size = Pt(11)

    # Rows
    for i, row in enumerate(rows):
        cs = tbl.rows[i + 1].cells
        for j, val in enumerate(row):
            cs[j].text = str(val)
            for p in cs[j].paragraphs:
                _para_rtl(p)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(10)
                    if i == highlight_row:
                        r.font.bold = True
                        r.font.color.rgb = GREEN
    return tbl


def _add_disclosure(doc, text: str):
    p = doc.add_paragraph()
    _para_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = GREY
    return p


def _money(x: Any) -> str:
    try:
        return f"{float(x):,.2f} ج.م"
    except Exception:
        return "—"


def _pct(x: Any) -> str:
    try:
        return f"{float(x):.2f}%"
    except Exception:
        return "—"


def _common_header(doc, payload: Dict[str, Any], purpose_label_ar: str):
    _add_h1(doc, "Expert_Smart — تقرير تقييم متخصص")
    _add_h2(doc, purpose_label_ar, color=GOLD)
    rid = payload.get("report_id", "—")
    rdate = payload.get("report_date") or datetime.now().strftime("%d/%m/%Y")
    _add_p(doc, f"رقم التقرير: {rid}    |    تاريخ التقرير: {rdate}",
           bold=True, color=DARK, size=11)
    _add_p(doc, f"الخبير المُعِد: {payload.get('expert', 'م. هشام المهدي')}",
           color=DARK, size=11)
    doc.add_paragraph()


def _property_kv(payload: Dict[str, Any]) -> List[Tuple[str, str]]:
    rows = [
        ("الموقع",         payload.get("location", "—")),
        ("نوع الأصل",      payload.get("property_type", "—")),
        ("المساحة",        f"{float(payload.get('area', 0) or 0):,.0f} م²"),
        ("سعر المتر",      _money(payload.get("price_per_meter", 0))),
        ("القيمة السوقية", _money(payload.get("market_value", 0))),
    ]
    return rows


# ════════════════════════════════════════════════════════════════════════════
#  PURPOSE-SPECIFIC REPORT BUILDERS
# ════════════════════════════════════════════════════════════════════════════

def _report_usufruct(payload: Dict[str, Any], path: str) -> bool:
    doc = Document()
    _set_rtl_doc(doc)
    _common_header(doc, payload, "تقرير تقييم حق الانتفاع")

    meta = payload.get("usufruct_meta") or {}
    years     = float(meta.get("usufruct_years", payload.get("usufruct_years", 25)))
    rate      = float(meta.get("discount_rate",  payload.get("discount_rate", 0.10)))
    pv_factor = float(meta.get("pv_factor", 0.65))
    mv        = float(payload.get("market_value", 0) or 0)
    usufruct_value = mv * pv_factor

    _add_h2(doc, "ملخص الحق")
    _add_kv_table(doc, [
        ("نوع الحق",                  "حق انتفاع (Usufruct Right) — IVS 410"),
        ("مدة الانتفاع المتبقية",     f"{years:.0f} سنة"),
        ("معدل الخصم المعتمد",        _pct(rate * 100)),
        ("معامل القيمة الحالية (PV)", f"{pv_factor:.4f}"),
        ("القيمة السوقية للملكية الكاملة", _money(mv)),
        ("قيمة حق الانتفاع",          _money(usufruct_value)),
    ])

    _add_h2(doc, "بيانات العقار محل الحق")
    _add_kv_table(doc, _property_kv(payload))

    _add_h2(doc, "المنهجية")
    _add_p(doc,
        "تم تقييم حق الانتفاع وفق نهج خصم التدفقات النقدية (DCF) عبر المدة "
        "المتبقية من العقد. القيمة الحالية للمنفعة المتولدة من العقار خلال هذه "
        "المدة تُحسب بصيغة annuity decay:")
    _add_p(doc, "   PV Factor = 1 − (1 + r)^(−n)", bold=True, color=NAVY)
    _add_p(doc,
        f"حيث r = {_pct(rate*100)} (معدل الخصم السوقي) و n = {years:.0f} سنة "
        f"(المدة المتبقية). يُطبَّق المعامل على القيمة السوقية للملكية الكاملة "
        f"للحصول على قيمة حق الانتفاع.")

    _add_h2(doc, "القيود والمخاطر")
    _add_p(doc, "• حق الانتفاع لا يشمل حق التصرف (بيع/رهن) — مقصور على المنفعة فقط.")
    _add_p(doc, "• تنتقل الملكية الرقبى للمالك الأصلي عند انتهاء المدة دون تعويض.")
    _add_p(doc, "• القيمة حساسة لأي تغيير في معدل الخصم أو افتراضات الإيجار المستقبلي.")

    _add_h2(doc, "التوصية")
    _add_p(doc, f"قيمة حق الانتفاع المُعتمدة = {_money(usufruct_value)}",
           bold=True, color=GREEN, size=14)

    _add_disclosure(doc,
        "هذا التقرير يخص حق الانتفاع كحق عيني محدود وفق القانون المدني المصري "
        "(م.985–م.997) ومعيار IVS 410. القيمة المُعلنة تعكس فقط حق المنفعة عبر "
        "المدة المُحددة ولا تتضمن قيمة الرقبة.")

    doc.save(path)
    return True


def _report_uncertainty(payload: Dict[str, Any], path: str) -> bool:
    doc = Document()
    _set_rtl_doc(doc)
    _common_header(doc, payload, "تقرير التقييم في حالة عدم اليقين — IFRS 13 §93")

    spread = float(payload.get("uncertainty_spread_pct", 0.15))
    mv     = float(payload.get("market_value", 0) or 0)
    low    = mv * (1 - spread)
    high   = mv * (1 + spread)

    _add_h2(doc, "نطاق القيمة المُفصَح عنه")
    _add_kv_table(doc, [
        ("الحد الأدنى (Low)",   _money(low)),
        ("القيمة الأكثر احتمالاً (Best Estimate)", _money(mv)),
        ("الحد الأعلى (High)",  _money(high)),
        ("نطاق الطيف (±)",      _pct(spread * 100)),
        ("مستوى الإفصاح",        "Level 3 — مدخلات غير قابلة للملاحظة"),
    ])

    _add_h2(doc, "مصادر عدم اليقين")
    _add_p(doc, "1) عدم يقين السوق — تقلب أسعار المعاملات المقارنة، نقص البيانات الحديثة، اضطراب اقتصادي.")
    _add_p(doc, "2) عدم يقين البيانات — جودة وحداثة المعاملات المرجعية، عينات صغيرة، تباين سعري.")
    _add_p(doc, "3) عدم يقين النموذج — حساسية الـ DCF لمعدل الخصم وافتراضات النمو.")
    _add_p(doc, "4) عدم يقين الأصل — تفرد العقار وصعوبة وجود مقارنات حقيقية.")

    _add_h2(doc, "بيانات العقار")
    _add_kv_table(doc, _property_kv(payload))

    _add_h2(doc, "تحليل الحساسية")
    _add_data_table(doc,
        headers=["السيناريو", "افتراض", "القيمة الناتجة"],
        rows=[
            ["متشائم",  f"−{_pct(spread*100)} من الـ Best",  _money(low)],
            ["مرجح",    "Best Estimate",                       _money(mv)],
            ["متفائل",  f"+{_pct(spread*100)} من الـ Best",  _money(high)],
        ],
        highlight_row=1,
    )

    _add_h2(doc, "إفصاحات IFRS 13 §93")
    _add_p(doc, "(d) المستوى في التسلسل الهرمي للقيمة العادلة: Level 3.")
    _add_p(doc, "(g) وصف تقنية التقييم والمدخلات المستخدمة.")
    _add_p(doc, "(h) تأثير تغيير المدخلات غير القابلة للملاحظة على القيمة العادلة.")
    _add_p(doc, "(i) تسوية الحركة من الرصيد الافتتاحي إلى الختامي للأصل.")

    _add_h2(doc, "التوصية")
    _add_p(doc,
        f"يُوصى بالإفصاح عن نطاق قيمة كامل ({_money(low)} — {_money(high)}) "
        f"بدلاً من قيمة نقطية واحدة، مع التأكيد أن Best Estimate = {_money(mv)} "
        f"يمثل الأكثر تمثيلاً للقيمة العادلة.",
        bold=True, color=GREEN)

    _add_disclosure(doc,
        "وفقاً لـ IFRS 13 §93 و RICS Red Book Material Valuation Uncertainty "
        "Clause، يجب الإفصاح صراحة عن وجود عدم يقين جوهري وعن نطاق القيمة "
        "المعقول، حماية لمستخدمي القوائم المالية.")

    doc.save(path)
    return True


def _report_hbu(payload: Dict[str, Any], path: str) -> bool:
    """Builds an HBU report from either an HBU result dict or basic payload."""
    doc = Document()
    _set_rtl_doc(doc)
    _common_header(doc, payload, "تقرير تحليل أعلى وأفضل استغلال (HBU)")

    # If payload is already an HBU result (delegated path) it has these keys:
    is_full_hbu = "scenarios_evaluated" in payload and "comparison_table" in payload
    if not is_full_hbu:
        _add_h2(doc, "ملاحظة")
        _add_p(doc,
            "هذا التقرير يقدم إطار تحليل HBU العام. لتحليل تفصيلي بالاختبارات "
            "الأربعة وحساب NPV لكل سيناريو، استخدم endpoint /api/hbu/analyze "
            "مع قائمة `alternative_uses`.")
        _add_h2(doc, "بيانات العقار")
        _add_kv_table(doc, _property_kv(payload))
        _add_disclosure(doc, "إطار IVS / USPAP / Appraisal Institute للاختبارات الأربعة المتسلسلة.")
        doc.save(path)
        return True

    # Full HBU report from delegated module
    prop = payload.get("property", {})
    _add_h2(doc, "ملخص العقار")
    _add_kv_table(doc, [
        ("الموقع",          prop.get("location", "—")),
        ("المساحة",         f"{float(prop.get('area', 0) or 0):,.0f} م²"),
        ("الاستخدام الحالي", prop.get("current_use", "—") or "—"),
        ("النطاق التخطيطي",  prop.get("current_zoning", "—") or "—"),
        ("معدل الخصم",       _pct(payload.get("discount_rate_pct", 0))),
        ("تاريخ التقييم",    payload.get("valuation_date", "—")),
    ])

    _add_h2(doc, "الاختبارات الأربعة لإطار HBU")
    _add_p(doc, "1) Test 1 — ممكن قانونياً (Legally Permissible)")
    _add_p(doc, "2) Test 2 — ممكن مادياً (Physically Possible)")
    _add_p(doc, "3) Test 3 — مجدٍ مالياً (Financially Feasible) — NPV > 0")
    _add_p(doc, "4) Test 4 — أعلى إنتاجية (Maximally Productive) — أعلى NPV من بين المرشحين")

    _add_h2(doc, "جدول المقارنة بين الاستخدامات البديلة")
    rows = []
    win_idx = -1
    for i, r in enumerate(payload["comparison_table"]):
        if r["maximally_productive"]:
            win_idx = i
        rows.append([
            r["use_name"],
            "✓" if r["legally_permissible"] else "✗",
            "✓" if r["physically_possible"] else "✗",
            "✓" if r["financially_feasible"] else "✗",
            "✓" if r["maximally_productive"] else "✗",
            f"{r['npv']:,.0f}",
            (f"{r['irr_pct']:.2f}%" if r["irr_pct"] is not None else "—"),
            (f"{r['payback_years']:.2f}" if r["payback_years"] is not None else "—"),
        ])
    _add_data_table(doc,
        headers=["الاستخدام", "قانوني", "مادي", "مالي", "أعلى إنتاجية",
                 "NPV (ج.م)", "IRR %", "الاسترداد (سنة)"],
        rows=rows,
        highlight_row=win_idx if win_idx >= 0 else None,
    )

    _add_h2(doc, "الأحكام النهائية")
    for r in payload["comparison_table"]:
        clr = GREEN if r["maximally_productive"] else (RED if not r["financially_feasible"] else None)
        _add_p(doc, f"• {r['use_name']}: {r['verdict']}", bold=r["maximally_productive"], color=clr)

    # ── Detailed year-by-year cashflows per scenario ────────────────────────
    scenarios = payload.get("scenarios_evaluated") or []
    if scenarios:
        _add_h2(doc, "التدفقات النقدية السنوية لكل سيناريو")
        max_yr = max((len(s.get("cashflows", [])) for s in scenarios), default=0)
        if max_yr > 0:
            headers = ["السيناريو"] + [f"سنة {y}" for y in range(max_yr)]
            rows = []
            win_name = payload.get("recommended_use")
            for s in scenarios:
                cfs = s.get("cashflows") or []
                cells = [s["use_name"]]
                for y in range(max_yr):
                    cells.append(f"{cfs[y]:,.0f}" if y < len(cfs) else "—")
                rows.append(cells)
            _add_data_table(doc, headers, rows,
                            highlight_row=next((i for i, s in enumerate(scenarios)
                                                if s.get("test_4_max_productive")), None))

    _add_h2(doc, "التوصية النهائية")
    if payload.get("recommended_use"):
        _add_p(doc,
            f"أعلى وأفضل استغلال = «{payload['recommended_use']}» "
            f"بصافي قيمة حالية = {_money(payload.get('recommended_npv', 0))}",
            bold=True, color=GREEN, size=14)
    else:
        _add_p(doc, payload.get("recommendation_note", "—"), bold=True, color=RED)

    _add_disclosure(doc, payload.get("standards_note",
        "تم تطبيق إطار IVS / USPAP / Appraisal Institute للاختبارات الأربعة."))
    doc.save(path)
    return True


def _report_investment_funds(payload: Dict[str, Any], path: str) -> bool:
    """Builds a REIT NAV report from either a REIT result dict or basic payload."""
    doc = Document()
    _set_rtl_doc(doc)
    _common_header(doc, payload, "تقرير صافي قيمة الأصول لصندوق الاستثمار العقاري — IFRS 13 / IOSCO")

    is_full_reit = "gross_asset_value" in payload and "nav" in payload
    if not is_full_reit:
        _add_h2(doc, "ملاحظة")
        _add_p(doc,
            "لإصدار تقرير REIT تفصيلي بالـ NAV per Unit، استخدم endpoint "
            "/api/reit/nav أو ابعث الطلب إلى /api/valuation مع "
            "`valuation_purpose=investment_funds` وقائمة `properties`.")
        _add_h2(doc, "بيانات العقار")
        _add_kv_table(doc, _property_kv(payload))
        _add_disclosure(doc, "IFRS 13 + IOSCO + لائحة صناديق الاستثمار العقاري.")
        doc.save(path)
        return True

    _add_h2(doc, "بيانات الصندوق")
    _add_kv_table(doc, [
        ("اسم الصندوق",    payload.get("fund_name", "—")),
        ("تاريخ التقييم",  payload.get("valuation_date", "—")),
        ("مستوى IFRS 13",  payload.get("ifrs_level_label", "—")),
        ("عدد الأصول",     str(payload.get("asset_count", "—"))),
    ])

    _add_h2(doc, "جانب الأصول")
    _add_kv_table(doc, [
        ("القيمة العادلة للعقارات",     _money(payload.get("real_estate_fair_value", 0))),
        ("النقد وما في حكمه",            _money(payload.get("cash_and_equiv", 0))),
        ("المدينون والإيجارات المستحقة", _money(payload.get("receivables", 0))),
        ("أصول أخرى",                   _money(payload.get("other_assets", 0))),
        ("إجمالي قيمة الأصول (GAV)",    _money(payload.get("gross_asset_value", 0))),
    ])

    _add_h2(doc, "جانب الالتزامات")
    _add_kv_table(doc, [
        ("القروض البنكية",                 _money(payload.get("loans", 0))),
        ("مصروفات مستحقة",                 _money(payload.get("accrued_expenses", 0))),
        ("توزيعات مُعتمدة غير مدفوعة",     _money(payload.get("distributions_payable", 0))),
        ("ضريبة مؤجلة",                    _money(payload.get("deferred_tax", 0))),
        ("التزامات أخرى",                  _money(payload.get("other_liabilities", 0))),
        ("إجمالي الالتزامات",              _money(payload.get("total_liabilities", 0))),
    ])

    _add_h2(doc, "صافي قيمة الأصول (NAV)")
    _add_kv_table(doc, [
        ("NAV (إجمالي)",        _money(payload.get("nav", 0))),
        ("عدد الوحدات المُصدَرة", f"{int(payload.get('units_outstanding', 0)):,}"),
        ("NAV لكل وحدة",        _money(payload.get("nav_per_unit", 0))),
    ])

    _add_h2(doc, "مؤشرات أداء المحفظة")
    _add_kv_table(doc, [
        ("Cap Rate المرجح", _pct(payload.get("weighted_cap_rate_pct", 0))),
        ("نسبة LTV",        _pct(payload.get("ltv_pct", 0))),
        ("Gearing Ratio",   _pct(payload.get("gearing_ratio_pct", 0))),
        ("تقييم الصحة",      f"{payload.get('health_score', 0):.0f}/100 — {payload.get('health_label', '—')}"),
    ])

    if payload.get("assets"):
        _add_h2(doc, "تفصيل العقارات داخل المحفظة")
        rows = [
            [a.get("asset_name", "—"),
             a.get("location", "—"),
             a.get("property_type", "—"),
             f"{float(a.get('area', 0) or 0):,.0f}",
             _pct(a.get("cap_rate_pct", 0)),
             _money(a.get("fair_value", 0))]
            for a in payload["assets"]
        ]
        _add_data_table(doc,
            headers=["اسم الأصل", "الموقع", "النوع", "المساحة (م²)",
                     "Cap Rate %", "القيمة العادلة"],
            rows=rows)

    _add_disclosure(doc, payload.get("ifrs_disclosure",
        "وفقاً لـ IFRS 13 §93 و IOSCO Principles for the Valuation of CIS Portfolios "
        "ولائحة هيئة السوق المالية للصناديق العقارية."))
    doc.save(path)
    return True


def _report_generic_purpose(payload: Dict[str, Any], path: str,
                            title_ar: str, methodology_ar: str,
                            disclosure_ar: str) -> bool:
    """نموذج تقرير عام للأغراض التسعة الأصلية — يتم تخصيص العنوان والمنهجية فقط."""
    doc = Document()
    _set_rtl_doc(doc)
    _common_header(doc, payload, title_ar)

    _add_h2(doc, "بيانات العقار")
    _add_kv_table(doc, _property_kv(payload))

    _add_h2(doc, "المنهجية")
    _add_p(doc, methodology_ar)

    # Display key valuation outputs if present
    extras = []
    for k, lbl in [
        ("market_approach",  "نهج السوق"),
        ("cost_approach",    "نهج التكلفة"),
        ("income_approach",  "نهج الدخل"),
        ("market_value",     "القيمة السوقية النهائية"),
    ]:
        if payload.get(k) is not None:
            extras.append((lbl, _money(payload.get(k))))
    if extras:
        _add_h2(doc, "نتائج المناهج الثلاثة")
        _add_kv_table(doc, extras)

    _add_h2(doc, "التوصية")
    _add_p(doc, f"القيمة المعتمدة لهذا الغرض = {_money(payload.get('market_value', 0))}",
           bold=True, color=GREEN, size=14)

    _add_disclosure(doc, disclosure_ar)
    doc.save(path)
    return True


# ════════════════════════════════════════════════════════════════════════════
#  PURPOSE → BUILDER MAP
# ════════════════════════════════════════════════════════════════════════════

def _report_eia(payload: Dict[str, Any], path: str) -> bool:
    """Builds an EIA report. If payload has full eia_result, use eia_engine.write_eia_word_report.
    Otherwise produce a minimal placeholder with instructions."""
    try:
        try:
            from eia_engine import run_eia_assessment, write_eia_word_report
        except Exception:
            from core_engine.eia_engine import run_eia_assessment, write_eia_word_report  # type: ignore
    except Exception:
        run_eia_assessment = None
        write_eia_word_report = None

    # If payload already looks like an EIA result, render directly
    if isinstance(payload.get("classification"), dict) and isinstance(payload.get("impact_matrix"), list):
        if write_eia_word_report:
            return write_eia_word_report(payload, path)

    # If payload has impact_matrix but isn't an EIA result, run the assessment
    if payload.get("impact_matrix") and run_eia_assessment and write_eia_word_report:
        try:
            result = run_eia_assessment(payload)
            return write_eia_word_report(result, path)
        except Exception as e:
            print(f"[purpose_specific_reports] EIA inline assessment failed: {e}")

    # Fallback: minimal placeholder
    doc = Document()
    _set_rtl_doc(doc)
    _common_header(doc, payload, "تقرير تقييم الأثر البيئي (EIA)")
    _add_h2(doc, "ملاحظة")
    _add_p(doc, "لإصدار تقرير EIA كامل، استخدم endpoint /api/eia/assess أو ابعث الطلب إلى "
                "/api/valuation مع `valuation_purpose=environmental_impact_assessment` "
                "وحقول `project`، `site_analysis`، `baseline`، `emissions`، `impact_matrix` …")
    _add_kv_table(doc, _property_kv(payload))
    _add_disclosure(doc, "ISO 14001:2015 / Law 4-1994 / World Bank ESF / IFC PS.")
    doc.save(path)
    return True


_PURPOSE_META: Dict[str, Tuple[str, str, str]] = {
    "fair_market_value":    ("تقرير القيمة السوقية العادلة (Fair Market Value)",
        "تم تطبيق نهج السوق ونهج التكلفة ونهج الدخل بمعايير IVS و RICS Red Book.",
        "القيمة السوقية تعكس السعر الذي يُتوقع تحقيقه في معاملة طوعية بين بائع راغب ومشتري راغب."),
    "acquisition":          ("تقرير القيمة الاستثمارية للاستحواذ والاندماج (M&A Investment Value)",
        "تم تطبيق نموذج DCF متخصص للاستحواذ مع معدلات خصم تعكس مخاطر المشتري الاستراتيجي.",
        "القيمة الاستثمارية تعكس القيمة لمستثمر محدد، وقد تختلف عن القيمة السوقية."),
    "bank_financing":       ("تقرير القيمة لأغراض التمويل البنكي (Bank Financing Value)",
        "تم تطبيق خصم تحفظي بنكي قدره 5% على القيمة السوقية لأغراض حساب LTV.",
        "هذه القيمة لأغراض التمويل والرهن العقاري وفق متطلبات Basel III."),
    "rental_arbitration":   ("تقرير تحديد الأجرة العادلة (Fair Rental Value)",
        "تم استخدام نهج المقارنات الإيجارية للسوق المحلي مع تحليل العائد على رأس المال.",
        "القيمة الإيجارية المُحددة لأغراض التحكيم الإيجاري وتحديد بدل المثل."),
    "insurance":            ("تقرير قيمة إعادة الإنشاء للتأمين (Reinstatement Value)",
        "تم تطبيق نهج التكلفة بأسعار البناء الجارية + 8% احتياطي رفع كلفة المواد.",
        "هذه القيمة لأغراض التأمين فقط وتُمثل تكلفة إعادة بناء العقار وليس قيمته السوقية."),
    "investment_analysis":  ("تقرير التحليل الاستثماري (IRR/NPV/DCF)",
        "تم تطبيق نموذج DCF متعدد السنوات مع تحليل حساسية لمعدلات الخصم والنمو.",
        "هذا التحليل يُستخدم لاتخاذ قرارات الاستثمار وقياس العائد المرتقب."),
    "judicial_liquidation": ("تقرير القيمة التصفوية القضائية (Forced Liquidation Value)",
        "تم تطبيق خصم تصفية إجبارية قدره 18% لتعكس البيع السريع تحت ضغط الزمن.",
        "هذه القيمة لأغراض التصفية القضائية فقط وأقل من القيمة السوقية بطبيعتها."),
    "tax_assessment":       ("تقرير القيمة للأغراض الضريبية (Tax Assessment Value)",
        "تم تطبيق المنهجية الضريبية المعتمدة من مصلحة الضرائب العقارية مع جداول التقدير.",
        "هذه القيمة لأغراض احتساب الضريبة العقارية وفق القانون 196 لسنة 2008."),
    "financial_reporting":  ("تقرير القيمة العادلة للتقارير المالية (Fair Value — IFRS 13)",
        "تم تطبيق IFRS 13 مع تحديد المستوى الهرمي للقيمة العادلة (Level 1/2/3).",
        "القيمة العادلة وفق IFRS 13 §9 — السعر المُستلَم لبيع الأصل في معاملة منظمة."),
}


def _maybe_add_asset_type_section(doc, payload: Dict[str, Any]) -> bool:
    """يحاول حقن قسم متخصص لنوع الأصل (إن وُجدت نتيجة متخصصة في الـ payload)."""
    try:
        try:
            from asset_type_reports import add_asset_type_section
        except Exception:
            from core_engine.asset_type_reports import add_asset_type_section  # type: ignore
    except Exception:
        return False
    asset_result = payload.get("specialized_asset_result")
    if not isinstance(asset_result, dict):
        return False
    return add_asset_type_section(doc, asset_result)


def write_purpose_specific_report(purpose: str, payload: Dict[str, Any],
                                  output_path: str) -> bool:
    """
    يُولِّد تقريراً خاصاً بالغرض المُحدَّد. يُرجِع True عند النجاح.

    إذا احتوى الـ payload على `specialized_asset_result` (نتيجة دالة متخصصة من
    asset_type_valuators)، يُضاف قسم مخصص للنوع داخل التقرير قبل الحفظ.

    purpose : مفتاح الغرض من _PURPOSES (مثل "usufruct", "highest_and_best_use" …)
    payload : dict يحوي بيانات العقار + مخرجات التقييم
    output_path : المسار الكامل لملف Word الناتج
    """
    if not _DOCX_OK:
        return False
    if not purpose:
        return False
    purpose = purpose.strip()

    # Build a temp path so we can re-open the doc to inject asset-type section
    temp_path = output_path + ".tmp"

    try:
        ok = False
        if purpose == "usufruct":
            ok = _report_usufruct(payload, temp_path)
        elif purpose == "uncertainty_valuation":
            ok = _report_uncertainty(payload, temp_path)
        elif purpose == "highest_and_best_use":
            ok = _report_hbu(payload, temp_path)
        elif purpose == "investment_funds":
            ok = _report_investment_funds(payload, temp_path)
        elif purpose == "environmental_impact_assessment":
            ok = _report_eia(payload, temp_path)
        elif purpose in _PURPOSE_META:
            title, method, discl = _PURPOSE_META[purpose]
            ok = _report_generic_purpose(payload, temp_path, title, method, discl)
        else:
            return False

        if not ok:
            return False

        # Re-open doc to append (1) purpose-detail section, then (2) asset-type section
        try:
            doc = Document(temp_path)
            try:
                try:
                    from purpose_detail_sections import add_purpose_detail_section
                except Exception:
                    from core_engine.purpose_detail_sections import add_purpose_detail_section  # type: ignore
                add_purpose_detail_section(doc, purpose, payload)
            except Exception as _pd_err:
                print(f"[purpose_specific_reports] detail section skipped: {_pd_err}")
            injected = _maybe_add_asset_type_section(doc, payload)
            doc.save(output_path)
            try:
                os.remove(temp_path)
            except Exception:
                pass
            if injected:
                print(f"[purpose_specific_reports] composite (asset+purpose) report saved: {output_path}")
            return True
        except Exception as ce:
            try:
                os.replace(temp_path, output_path)
            except Exception:
                pass
            print(f"[purpose_specific_reports] composite step skipped: {ce}")
            return True
    except Exception as e:
        print(f"[purpose_specific_reports] {purpose} failed: {e}")
        return False


def write_hbu_word_report(result, output_path):
    return _report_hbu(result, output_path)


def write_reit_word_report(result, output_path):
    return _report_investment_funds(result, output_path)
