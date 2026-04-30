"""
run_final_valuation.py
======================
تنفيذ التقييم الكامل لوحدة التجمع الخامس وفق البرومبيت:

  المدخلات:
    الموقع  : التجمع الخامس (New Cairo) — بالقرب من بيفرلي هيلز / الرحاب
    المساحة : 200 م²
    النوع   : شقة سكنية — تشطيب سوبر لوكس
    العمر   : 3 سنوات

  المخرجات:
    Final_Valuation_Report.xlsx  — تقرير Excel متعدد الأوراق (RTL)
    Final_Valuation_Report.docx  — تقرير Word سردي بالعربية (RTL)

التشغيل:
    python run_final_valuation.py
"""

import os
import sys
import io
import math
import warnings
from datetime import datetime

# ── مسارات المشروع ────────────────────────────────────────────────────────────
_CORE_DIR = os.path.dirname(os.path.abspath(__file__))
_ROOT_DIR = os.path.join(_CORE_DIR, "..")
sys.path.insert(0, _CORE_DIR)
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
warnings.filterwarnings("ignore")

OUTPUT_DIR  = os.path.join(_CORE_DIR, "outputs", "reports")
EXCEL_PATH  = os.path.join(OUTPUT_DIR, "Final_Valuation_Report.xlsx")
WORD_PATH   = os.path.join(OUTPUT_DIR, "Final_Valuation_Report.docx")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ── بيانات العقار الهدف ───────────────────────────────────────────────────────
PROPERTY = {
    "client_name":   "لجنة الاستثمار العقاري",
    "property_type": "شقة سكنية سوبر لوكس",
    "location":      "التجمع الخامس — New Cairo — بالقرب من بيفرلي هيلز والرحاب",
    "area":          200.0,
    "floor":         5,
    "rooms":         4,
    "year_built":    2022,
    "building_age":  3,
    "price_per_m2":  27000.0,    # سعر المتر المرجعي (سوبر لوكس — التجمع الخامس)
    "rent_per_sqm":  550.0,      # إيجار سنوي EGP/م² (تقدير سوبر لوكس)
    "cap_rate":      0.075,      # معدل رسملة عقارات راقية (7.5%)
    "target_x":      31.4450,    # إحداثي خطي (longitude — التجمع الخامس)
    "target_y":      30.0139,    # إحداثي عرضي (latitude)
    "report_number": f"VAL-FINAL-{datetime.now().strftime('%Y%m%d-%H%M')}",
}


# ══════════════════════════════════════════════════════════════════════════════
# STEP 1 — Agentic RAG: جلب مقارنات حقيقية من Qdrant
# ══════════════════════════════════════════════════════════════════════════════
def step1_rag():
    print("[ 1/5 ] تشغيل RAG — البحث عن مقارنات في قاعدة البيانات المتجهية...")
    from valuation_logic import _rag_search
    comps = _rag_search(PROPERTY["location"], top_k=5)
    if comps:
        print(f"        ✔ تم جلب {len(comps)} مقارنة حقيقية من Qdrant")
        for c in comps:
            print(f"          — {c['loc']}: {c['price_per_m2']:,.0f} EGP/م²  ({c['ar']:.0f} م²)")
    else:
        print("        ⚠ RAG لم يجد نتائج — سيُستخدم النظام الاصطناعي")
    return comps


# ══════════════════════════════════════════════════════════════════════════════
# STEP 2 — IVS Valuation: تقييم كامل (6 طرق)
# ══════════════════════════════════════════════════════════════════════════════
def step2_ivs_valuation():
    print("[ 2/5 ] تشغيل نماذج IVS (السوق / التكلفة / الدخل / GIS / الانحدار / الخيارات)...")
    from valuation_logic import advanced_valuation
    data = {
        "area":          PROPERTY["area"],
        "price_per_meter": PROPERTY["price_per_m2"],
        "location":      PROPERTY["location"],
        "building_age":  PROPERTY["building_age"],
        "rent_per_sqm":  PROPERTY["rent_per_sqm"],
        "cap_rate":      PROPERTY["cap_rate"],
        "property_type": PROPERTY["property_type"],
    }
    result = advanced_valuation(data)
    rv = result["reconciled_value"]
    print(f"        ✔ القيمة التوفيقية (3 طرق): {rv:,.0f} EGP  ({rv/PROPERTY['area']:,.0f} EGP/م²)")
    print(f"        — HBU: {result['hbu_text'][:80]}...")
    return result


# ══════════════════════════════════════════════════════════════════════════════
# STEP 3 — DCF + Market Forecast
# ══════════════════════════════════════════════════════════════════════════════
def step3_dcf_forecast(ivs_result):
    print("[ 3/5 ] تشغيل DCF + توقعات السوق (Prophet / ARIMA)...")
    from investment_engine import run_market_forecast, run_dcf, run_hbu_advanced

    final_value  = ivs_result["reconciled_value"]
    current_ppm  = final_value / PROPERTY["area"]

    forecast = run_market_forecast(current_ppm, PROPERTY["location"])
    print(f"        ✔ اتجاه السوق: {forecast['trend']}  (12 شهر: {forecast['signals'].get('12m',0):,.0f} EGP/م²)")

    dcf = run_dcf(
        total_value      = final_value,
        area             = PROPERTY["area"],
        rent_per_sqm     = PROPERTY["rent_per_sqm"],
        cap_rate         = PROPERTY["cap_rate"],
        building_age     = PROPERTY["building_age"],
        forecast_signals = forecast["signals"],
    )
    print(f"        ✔ WACC: {dcf['wacc']*100:.1f}%  |  IRR: {dcf['irr']*100:.1f}%  |  NPV: {dcf['npv']:,.0f} EGP")

    land_ppm = current_ppm * 0.30
    hbu_adv = run_hbu_advanced(
        area         = PROPERTY["area"],
        land_ppm     = land_ppm,
        location     = PROPERTY["location"],
        wacc         = dcf["wacc"],
        forecast_pct = forecast["trend_pct"],
    )
    best = hbu_adv["best"]["name"]
    print(f"        ✔ أفضل استخدام (HBU): {best}  |  NPV = {hbu_adv['best']['npv']:,.0f} EGP")

    return forecast, dcf, hbu_adv


# ══════════════════════════════════════════════════════════════════════════════
# STEP 4 — Excel Report: تقرير Excel كامل
# ══════════════════════════════════════════════════════════════════════════════
def step4_excel(ivs_result, forecast, dcf, hbu_adv):
    print("[ 4/5 ] إنشاء تقرير Excel (IVS + DCF + GIS + توقعات) ...")

    # --- ورقة IVS الكاملة (11 ورقة) ---
    from master_report_generator import generate_report
    ivs_path = generate_report(
        client_name   = PROPERTY["client_name"],
        property_type = PROPERTY["property_type"],
        location      = PROPERTY["location"],
        area          = PROPERTY["area"],
        floor         = PROPERTY["floor"],
        rooms         = PROPERTY["rooms"],
        year_built    = PROPERTY["year_built"],
        price_per_m2  = PROPERTY["price_per_m2"],
        rent_per_sqm  = PROPERTY["rent_per_sqm"],
        cap_rate      = PROPERTY["cap_rate"],
        target_x      = PROPERTY["target_x"],
        target_y      = PROPERTY["target_y"],
        report_number = PROPERTY["report_number"],
        output_dir    = OUTPUT_DIR,
    )

    # --- ورقة DCF الاستثمارية (4 أوراق) ---
    from investment_engine import build_excel_report, _generate_narrative
    narrative = _generate_narrative(
        property_type = PROPERTY["property_type"],
        location      = PROPERTY["location"],
        area          = PROPERTY["area"],
        dcf           = dcf,
        hbu           = hbu_adv,
        forecast      = forecast,
        final_value   = ivs_result["reconciled_value"],
    )
    invest_path = build_excel_report(
        client_name   = PROPERTY["client_name"],
        property_type = PROPERTY["property_type"],
        location      = PROPERTY["location"],
        area          = PROPERTY["area"],
        final_value   = ivs_result["reconciled_value"],
        forecast      = forecast,
        dcf           = dcf,
        hbu           = hbu_adv,
        narrative     = narrative,
        output_dir    = OUTPUT_DIR,
    )

    # --- دمج الملفين في Final_Valuation_Report.xlsx ---
    import xlsxwriter, openpyxl, shutil

    try:
        _merge_workbooks(ivs_path, invest_path, EXCEL_PATH)
        print(f"        ✔ تم دمج الأوراق في: {EXCEL_PATH}")
    except Exception as e:
        # إذا فشل الدمج، احفظ التقرير الكامل (IVS) باسم Final
        shutil.copy(ivs_path, EXCEL_PATH)
        print(f"        ⚠ تعذّر الدمج ({e}) — تم حفظ تقرير IVS كـ Final")

    return narrative, ivs_path


def _merge_workbooks(ivs_path: str, invest_path: str, out_path: str):
    """
    يدمج ورقتَي IVS والاستثمارية في ملف واحد باستخدام openpyxl (cell-by-cell copy).
    """
    import openpyxl
    from copy import copy as _copy

    wb_main   = openpyxl.load_workbook(ivs_path)
    wb_invest = openpyxl.load_workbook(invest_path)

    for sh_name in wb_invest.sheetnames:
        src = wb_invest[sh_name]
        # avoid name clash
        dst_name = sh_name[:31]
        if dst_name in wb_main.sheetnames:
            dst_name = (sh_name[:28] + "_I")[:31]
        dst = wb_main.create_sheet(title=dst_name)

        # copy column widths
        for col_letter, col_dim in src.column_dimensions.items():
            dst.column_dimensions[col_letter].width = col_dim.width

        # copy row heights
        for row_num, row_dim in src.row_dimensions.items():
            dst.row_dimensions[row_num].height = row_dim.height

        # RTL
        dst.sheet_view.rightToLeft = True

        # copy merged cells first
        for merge in src.merged_cells.ranges:
            dst.merge_cells(str(merge))

        # copy cell values and basic styles (skip non-primary merged cells)
        for row in src.iter_rows():
            for cell in row:
                # MergedCell objects (non-top-left cells of a merge) are read-only
                if type(cell).__name__ == "MergedCell":
                    continue
                dst_cell = dst.cell(row=cell.row, column=cell.column)
                dst_cell.value = cell.value
                try:
                    if cell.has_style:
                        dst_cell.font      = _copy(cell.font)
                        dst_cell.fill      = _copy(cell.fill)
                        dst_cell.border    = _copy(cell.border)
                        dst_cell.alignment = _copy(cell.alignment)
                        if cell.number_format:
                            dst_cell.number_format = cell.number_format
                except Exception:
                    pass

    wb_main.save(out_path)


# ══════════════════════════════════════════════════════════════════════════════
# STEP 5 — Word Report: تقرير Word سردي
# ══════════════════════════════════════════════════════════════════════════════
def step5_word(ivs_result, forecast, dcf, hbu_adv, narrative):
    print("[ 5/5 ] إنشاء تقرير Word (Arabic RTL Narrative)...")
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import lxml.etree as etree

    doc = Document()

    # ── إعداد الصفحة (A4 RTL) ──────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = Cm(2.5)
    section.top_margin  = section.bottom_margin = Cm(2.0)

    def set_rtl_body(doc):
        """جعل النص الافتراضي RTL في الجسم."""
        body = doc.element.body
        body_pr = body.get_or_add_pPr()
        bidi = OxmlElement("w:bidi")
        body_pr.append(bidi)

    def add_para(doc, text, bold=False, size=12, color=None,
                 align=WD_ALIGN_PARAGRAPH.RIGHT, rtl=True,
                 space_before=0, space_after=6, italic=False):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        # RTL paragraph
        pPr = p._p.get_or_add_pPr()
        bidi_el = OxmlElement("w:bidi")
        pPr.append(bidi_el)
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.name = "Simplified Arabic"
        if color:
            run.font.color.rgb = RGBColor(*color)
        # RTL run
        rPr = run._r.get_or_add_rPr()
        rtl_el = OxmlElement("w:rtl")
        rPr.append(rtl_el)
        return p

    def add_heading(doc, text, level=1):
        sizes  = {1: 20, 2: 16, 3: 13}
        colors = {1: (31, 78, 120), 2: (46, 116, 181), 3: (0, 0, 0)}
        bold   = True
        add_para(doc, text, bold=bold, size=sizes.get(level, 12),
                 color=colors.get(level), space_before=10, space_after=4)

    def add_kv(doc, key, value):
        """سطر مفتاح: قيمة"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pPr = p._p.get_or_add_pPr()
        pPr.append(OxmlElement("w:bidi"))
        r_key = p.add_run(f"{key}:  ")
        r_key.bold      = True
        r_key.font.size = Pt(11)
        r_key.font.name = "Simplified Arabic"
        r_val = p.add_run(str(value))
        r_val.font.size = Pt(11)
        r_val.font.name = "Simplified Arabic"
        for r in [r_key, r_val]:
            rPr = r._r.get_or_add_rPr()
            rPr.append(OxmlElement("w:rtl"))

    def add_table_row(table, cells_data, header=False):
        row = table.add_row()
        for i, (txt, width) in enumerate(cells_data):
            cell = row.cells[i]
            cell.text = str(txt)
            cell.width = Cm(width)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.name = "Simplified Arabic"
                    run.font.size = Pt(10)
                    run.bold = header
                    if header:
                        run.font.color.rgb = RGBColor(255, 255, 255)
            if header:
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "1F4E78")
                tc_pr.append(shd)

    today     = datetime.now().strftime("%Y/%m/%d")
    final_val = ivs_result["reconciled_value"]
    ppm       = final_val / PROPERTY["area"]

    # ══ غلاف التقرير ══════════════════════════════════════════════════════════
    add_para(doc, "تقرير التقييم العقاري الاستثماري", bold=True, size=24,
             color=(31, 78, 120), align=WD_ALIGN_PARAGRAPH.CENTER, space_before=40)
    add_para(doc, "Final Investment Valuation Report", bold=False, size=14,
             color=(128, 128, 128), align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "─" * 60, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=2)
    add_para(doc, f"التاريخ: {today}  |  رقم التقرير: {PROPERTY['report_number']}",
             size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=(80, 80, 80))
    add_para(doc, f"المُعِدّ: خبير تقييم معتمد  |  الجهة: الهيئة العامة للرقابة المالية",
             size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=(80, 80, 80))
    doc.add_page_break()

    # ══ القسم الأول: بيانات العقار ════════════════════════════════════════════
    add_heading(doc, "أولاً: بيانات العقار الهدف", level=1)
    add_kv(doc, "الموقع",         PROPERTY["location"])
    add_kv(doc, "المساحة",        f"{PROPERTY['area']:.0f} م²")
    add_kv(doc, "النوع",          PROPERTY["property_type"])
    add_kv(doc, "الطابق",         f"الطابق {PROPERTY['floor']}")
    add_kv(doc, "عدد الغرف",      f"{PROPERTY['rooms']} غرف")
    add_kv(doc, "سنة الإنشاء",    str(PROPERTY["year_built"]))
    add_kv(doc, "عمر المبنى",     f"{PROPERTY['building_age']} سنوات")
    add_kv(doc, "سعر المتر المرجعي", f"{PROPERTY['price_per_m2']:,.0f} EGP/م²")
    add_kv(doc, "الإيجار السنوي",  f"{PROPERTY['rent_per_sqm']:,.0f} EGP/م²/سنة")
    add_kv(doc, "معدل الرسملة",   f"{PROPERTY['cap_rate']*100:.1f}%")

    # ══ القسم الثاني: نتائج RAG ════════════════════════════════════════════════
    doc.add_paragraph()
    add_heading(doc, "ثانياً: نتائج البحث التوليدي المعزز (Agentic RAG)", level=1)
    add_para(doc, (
        "تم تشغيل نظام RAG (Retrieval-Augmented Generation) باستخدام قاعدة بيانات متجهية "
        "(Qdrant) تحتوي على عقارات السوق المصري. تم استرداد أقرب 5 مقارنات سوقية للموقع "
        "المستهدف وفق نموذج التضمين (intfloat/multilingual-e5-large)."
    ), size=11)

    comps = ivs_result["market"].get("comparables", [])
    if comps:
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        add_table_row(tbl, [("العقار المقارن", 6), ("السعر الأساسي (EGP/م²)", 4.5), ("السعر المعدل (EGP/م²)", 4.5)], header=True)
        for c in comps:
            add_table_row(tbl, [
                (c.get("name", c.get("location", "—")), 6),
                (f"{c.get('base_price', 0):,.0f}", 4.5),
                (f"{c.get('adj_price', 0):,.0f}", 4.5),
            ])
        doc.add_paragraph()

    rag_flag = "✔ RAG نشط — مقارنات حقيقية" if ivs_result["market"].get("rag_used") else "⚠ RAG غير نشط — مقارنات اصطناعية"
    add_para(doc, rag_flag, size=11, italic=True, color=(0, 112, 0) if ivs_result["market"].get("rag_used") else (180, 0, 0))

    # ══ القسم الثالث: نتائج التقييم (IVS) ═════════════════════════════════════
    doc.add_paragraph()
    add_heading(doc, "ثالثاً: نتائج التقييم وفق معايير IVS", level=1)
    add_para(doc, (
        "تم تطبيق المنهجية الثلاثية وفق معايير التقييم الدولية (IVS 105) باستخدام "
        "ثلاثة أساليب تقييم رئيسية يُوزَن بينها في مرحلة التوفيق."
    ), size=11)

    add_heading(doc, "1. أسلوب مقارنة السوق (Sales Comparison Approach)", level=2)
    market_val = ivs_result["market"]["value"]
    add_kv(doc, "القيمة السوقية",  f"{market_val:,.0f} EGP")
    add_kv(doc, "سعر المتر المعدَّل", f"{ivs_result['market']['ppm']:,.0f} EGP/م²")
    add_para(doc, (
        "يعتمد هذا الأسلوب على مقارنة العقار بصفقات مماثلة منقضية في نفس المنطقة الجغرافية، "
        "مع تطبيق تعديلات نسبية على الموقع والمساحة والتشطيب وزمن البيع."
    ), size=11, color=(80, 80, 80))

    add_heading(doc, "2. أسلوب التكلفة (Depreciated Replacement Cost)", level=2)
    cost_val = ivs_result["cost"]["value"]
    add_kv(doc, "القيمة الاستبدالية المستهلكة", f"{cost_val:,.0f} EGP")
    add_kv(doc, "تكلفة البناء الإجمالية", f"{ivs_result['cost']['gross_building_cost']:,.0f} EGP")
    add_kv(doc, "الإهلاك المتراكم", f"{ivs_result['cost']['depreciation']:,.0f} EGP")
    add_para(doc, (
        "يحسب هذا الأسلوب تكلفة إعادة بناء العقار بمواصفاته الحالية (5,500 EGP/م²) "
        "بإضافة قيمة الأرض وخصم الإهلاك المادي والوظيفي وفق العمر الاقتصادي للمبنى."
    ), size=11, color=(80, 80, 80))

    add_heading(doc, "3. أسلوب رأسمالة الدخل (Income Capitalization)", level=2)
    income_val = ivs_result["income"]["value"]
    add_kv(doc, "القيمة بالرسملة المباشرة", f"{income_val:,.0f} EGP")
    add_kv(doc, "صافي الدخل التشغيلي (NOI)", f"{ivs_result['income']['noi']:,.0f} EGP/سنة")
    add_kv(doc, "معدل الرسملة المُطبَّق", f"{ivs_result['income']['cap_rate']*100:.1f}%")
    add_para(doc, (
        "يُرسمل هذا الأسلوب صافي الدخل التشغيلي المتوقع باستخدام معدل رسملة مستمد من "
        "السوق المحلي، مع تعديل معدل الرسملة الفعّال بإضافة استرداد رأس المال وفق العمر المتبقي."
    ), size=11, color=(80, 80, 80))

    # توفيق
    add_heading(doc, "4. التوفيق بين الطرق (Reconciliation)", level=2)
    add_kv(doc, "القيمة التوفيقية النهائية", f"{final_val:,.0f} EGP")
    add_kv(doc, "سعر المتر المُوحَّد",       f"{ppm:,.0f} EGP/م²")
    add_para(doc, "أوزان التوفيق: 50% سوق + 30% تكلفة + 20% دخل", size=11, italic=True)

    # ══ القسم الرابع: تحليل HBU ════════════════════════════════════════════════
    doc.add_paragraph()
    add_heading(doc, "رابعاً: تحليل أعلى وأفضل استخدام (HBU Analysis)", level=1)
    add_para(doc, ivs_result.get("hbu_text", "—"), size=11)

    # ══ القسم الخامس: توقعات السوق ════════════════════════════════════════════
    doc.add_paragraph()
    add_heading(doc, "خامساً: توقعات السوق (Market Forecast — Prophet Model)", level=1)

    signals  = forecast["signals"]
    cur_ppm  = forecast["current_ppm"]
    add_kv(doc, "السعر الحالي",       f"{cur_ppm:,.0f} EGP/م²")
    add_kv(doc, "التوقع بعد 3 أشهر",  f"{signals.get('3m', cur_ppm):,.0f} EGP/م²")
    add_kv(doc, "التوقع بعد 6 أشهر",  f"{signals.get('6m', cur_ppm):,.0f} EGP/م²")
    add_kv(doc, "التوقع بعد 9 أشهر",  f"{signals.get('9m', cur_ppm):,.0f} EGP/م²")
    add_kv(doc, "التوقع بعد 12 شهرًا", f"{signals.get('12m', cur_ppm):,.0f} EGP/م²")
    add_kv(doc, "الاتجاه العام",       forecast["trend"])
    add_kv(doc, "نسبة التغير (12 شهر)", f"{forecast['trend_pct']:.1f}%")
    add_para(doc, (
        "يعتمد نموذج التنبؤ على Prophet (Meta) مع بيانات تاريخية مُقدَّرة لـ24 شهراً. "
        "تُمثّل النتائج توقعات سعر المتر المربع بحدود الثقة."
    ), size=11, color=(80, 80, 80))

    # ══ القسم السادس: DCF والعوائد الاستثمارية ════════════════════════════════
    doc.add_paragraph()
    add_heading(doc, "سادساً: التدفقات النقدية المخصومة (DCF) ومصفوفة المخاطر", level=1)

    add_heading(doc, "أ. تحليل WACC (Weighted Average Cost of Capital)", level=2)
    add_kv(doc, "معدل الفائدة الخالي من المخاطر", f"{dcf['risk_free']*100:.1f}%  (أذون خزانة مصرية)")
    add_kv(doc, "علاوة مخاطر السوق × Beta",       f"{dcf['equity_risk']*100:.2f}%  (β = {dcf['beta']})")
    add_kv(doc, "علاوة مخاطر التطوير",             f"{dcf['dev_risk']*100:.1f}%")
    add_kv(doc, "علاوة مخاطر السيولة",             f"{dcf['liquidity_risk']*100:.1f}%")
    add_kv(doc, "معدل الخصم الإجمالي (WACC)",      f"{dcf['wacc']*100:.2f}%")

    add_heading(doc, "ب. جدول التدفقات النقدية (5 سنوات)", level=2)
    tbl2 = doc.add_table(rows=1, cols=3)
    tbl2.style = "Table Grid"
    add_table_row(tbl2, [("السنة", 3), ("NOI السنوي (EGP)", 5), ("القيمة الحالية (EGP)", 5)], header=True)
    for cf in dcf["cash_flows"]:
        add_table_row(tbl2, [
            (f"السنة {cf['year']}", 3),
            (f"{cf['noi']:,.0f}", 5),
            (f"{cf['pv']:,.0f}", 5),
        ])
    doc.add_paragraph()

    add_heading(doc, "ج. القيمة النهائية (Terminal Value)", level=2)
    add_kv(doc, "Gordon Growth Model",    f"{dcf['tv_gordon']:,.0f} EGP")
    add_kv(doc, "Exit Cap Rate Method",   f"{dcf['tv_exit_cap']:,.0f} EGP")
    add_kv(doc, "Exit Multiple Method",   f"{dcf['tv_exit_mult']:,.0f} EGP")
    add_kv(doc, "متوسط القيمة النهائية", f"{dcf['tv_average']:,.0f} EGP")

    add_heading(doc, "د. مؤشرات الجدوى", level=2)
    add_kv(doc, "صافي القيمة الحالية (NPV)",    f"{dcf['npv']:,.0f} EGP")
    add_kv(doc, "معدل العائد الداخلي (IRR)",    f"{dcf['irr']*100:.1f}%")
    add_kv(doc, "القيمة الحالية للتدفقات",      f"{dcf['pv_cashflows']:,.0f} EGP")
    add_kv(doc, "القيمة الحالية للمحطة النهائية", f"{dcf['pv_terminal']:,.0f} EGP")

    feasibility = "مُجدٍ استثمارياً" if dcf["npv"] > 0 and dcf["irr"] > dcf["wacc"] else "يحتاج مراجعة هيكل التمويل"
    add_para(doc, f"الحكم الاستثماري: {feasibility}", bold=True, size=12,
             color=(0, 112, 0) if dcf["npv"] > 0 else (180, 0, 0))

    # ══ القسم السابع: HBU المتقدم ═══════════════════════════════════════════
    doc.add_paragraph()
    add_heading(doc, "سابعاً: تحليل HBU المتقدم — مقارنة السيناريوهات", level=1)

    tbl3 = doc.add_table(rows=1, cols=5)
    tbl3.style = "Table Grid"
    add_table_row(tbl3, [
        ("السيناريو", 4), ("NOI السنوي", 3.5), ("NPV (EGP)", 3.5),
        ("IRR التقريبي", 3), ("التوصية", 2)
    ], header=True)
    for s in hbu_adv["scenarios"]:
        is_best = s["name"] == hbu_adv["best"]["name"]
        add_table_row(tbl3, [
            (s["name"], 4),
            (f"{s['annual_noi']:,.0f}", 3.5),
            (f"{s['npv']:,.0f}", 3.5),
            (f"{s['irr']*100:.1f}%", 3),
            ("★ الأفضل" if is_best else "—", 2),
        ])
    doc.add_paragraph()
    add_para(doc, f"الاستخدام الأعلى إنتاجية: {hbu_adv['best']['name']}", bold=True, size=12)

    # ══ القسم الثامن: الملخص الاستثماري (AI Narrative) ═══════════════════════
    doc.add_paragraph()
    add_heading(doc, "ثامناً: الملخص الاستثماري (Investment Committee Summary)", level=1)
    add_para(doc, narrative, size=11)

    # ══ القسم التاسع: القيمة السوقية النهائية ══════════════════════════════════
    doc.add_paragraph()
    add_heading(doc, "تاسعاً: القيمة السوقية النهائية المُعتمَدة", level=1)
    add_para(doc, f"{final_val:,.0f}  جنيه مصري", bold=True, size=20,
             color=(31, 78, 120), align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10)
    add_para(doc, f"سعر المتر المربع: {ppm:,.0f} EGP/م²", bold=True, size=14,
             align=WD_ALIGN_PARAGRAPH.CENTER, color=(212, 175, 55))
    add_para(doc, "تاريخ صلاحية التقييم: 3 أشهر من تاريخ الإصدار", size=10,
             align=WD_ALIGN_PARAGRAPH.CENTER, color=(128, 128, 128))

    doc.add_page_break()

    # ══ شهادة المُقيِّم ═════════════════════════════════════════════════════
    add_heading(doc, "شهادة المُقيِّم", level=1)
    add_para(doc, (
        "أُقرّ أنا المُقيِّم الموقّع أدناه أن هذا التقييم قد أُعِدَّ وفق معايير التقييم الدولية (IVS) "
        "الصادرة عن مجلس معايير التقييم الدولية (IVSC)، وأن القيمة المُقدَّرة تعكس رأيي المهني "
        "المستقل المبني على تحليل السوق والبيانات المتاحة في تاريخ التقييم."
    ), size=11)
    doc.add_paragraph()
    add_kv(doc, "الخبير المُقيِّم", "هشام محمد محمد المهدى")
    add_kv(doc, "رقم القيد",        "29 — الهيئة العامة للرقابة المالية")
    add_kv(doc, "تاريخ التقرير",    today)
    add_para(doc, "\n\n_" * 30 + "  توقيع المُقيِّم", size=11)

    doc.save(WORD_PATH)
    print(f"        ✔ تم حفظ: {WORD_PATH}")


# ══════════════════════════════════════════════════════════════════════════════
# MAIN — تشغيل كل الخطوات
# ══════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("=" * 65)
    print("  Expert Smart — Final Valuation Engine")
    print(f"  العقار: {PROPERTY['property_type']} — {PROPERTY['location']}")
    print(f"  المساحة: {PROPERTY['area']:.0f} م²")
    print("=" * 65)
    print()

    # تسلسل الخطوات
    rag_comps         = step1_rag()
    ivs_result        = step2_ivs_valuation()
    forecast, dcf, hbu = step3_dcf_forecast(ivs_result)
    narrative, _      = step4_excel(ivs_result, forecast, dcf, hbu)
    step5_word(ivs_result, forecast, dcf, hbu, narrative)

    print()
    print("=" * 65)
    print("  التقارير جاهزة في مجلد outputs/reports:")
    print(f"  Excel: {EXCEL_PATH}")
    print(f"  Word:  {WORD_PATH}")
    print("=" * 65)
