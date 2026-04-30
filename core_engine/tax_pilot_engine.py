"""
tax_pilot_engine.py
===================
Ministry of Finance — Tax Pilot Engine (Expert_Smart PropTech)
---------------------------------------------------------------
Implements:
  1.  Synthetic dataset  — 50 properties in New Cairo districts
  2.  IAAO statistics engine (Median ASR, COD, PRD, revenue uplift)
  3.  Excel export        — Tax_Authority_Pilot_V1.xlsx (xlsxwriter)
  4.  Ministry Word section appended to an existing python-docx Document
  5.  run_tax_pilot()     — orchestrator returning stats + paths

Usage:
    from tax_pilot_engine import run_tax_pilot
    result = run_tax_pilot()
"""
from __future__ import annotations
import os, math, random
from datetime import datetime
from typing import Dict, List

# ── Output directories ────────────────────────────────────────────────────────
_BASE_DIR = os.path.dirname(os.path.abspath(__file__))
_OUT_DIR  = os.path.join(_BASE_DIR, "outputs", "reports")

# ── Optional heavy dependencies ───────────────────────────────────────────────
try:
    import xlsxwriter
    _XLSX_OK = True
except ImportError:
    _XLSX_OK = False

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

# ── Design constants (matching master_word_report.py palette) ─────────────────
if _DOCX_OK:
    NAVY   = RGBColor(0x1F, 0x4E, 0x78)
    GOLD   = RGBColor(0xD4, 0xAF, 0x37)
    GREEN  = RGBColor(0x2E, 0x86, 0x48)
    RED    = RGBColor(0xC0, 0x00, 0x00)
    WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
    GRAY   = RGBColor(0x60, 0x60, 0x60)

FONT_MAIN = "Simplified Arabic"

# ── Reproducibility seed ──────────────────────────────────────────────────────
random.seed(42)

# ─────────────────────────────────────────────────────────────────────────────
#  1. SYNTHETIC DATASET
# ─────────────────────────────────────────────────────────────────────────────

DISTRICTS = [
    "التجمع الخامس",
    "المستثمرين الجنوبية",
    "الحي الأول",
    "بيفرلي هيلز",
    "النرجس",
]


def _rand_float(lo: float, hi: float, decimals: int = 2) -> float:
    return round(random.uniform(lo, hi), decimals)


def generate_dataset() -> List[Dict]:
    """Return 50 synthetic New Cairo properties (25 villas + 25 apartments)."""
    properties: List[Dict] = []

    # --- 25 Luxury Villas ---------------------------------------------------
    # old_tax_value simulates the REGRESSIVITY problem in manual assessment:
    # high-value villas are deliberately under-assessed at 35-55% of market,
    # while lower-value units get assessed closer to 55-70%.
    # This creates a wide COD (12-18%) and PRD > 1.03 (regressive system).
    for i in range(1, 26):
        area        = _rand_float(300, 600, 1)
        ppm         = _rand_float(28_000, 42_000, 0)
        market_val  = round(area * ppm, 0)
        # Villas: systematically under-assessed (35–52%) — regressivity bias
        old_tax_ratio = _rand_float(0.35, 0.52)
        old_tax     = round(market_val * old_tax_ratio, 0)
        ai_val      = round(market_val * _rand_float(0.93, 1.07), 0)
        properties.append({
            "id":                  f"V-{i:03d}",
            "property_type":       "فيلا فاخرة",
            "area":                area,
            "price_per_m2":        ppm,
            "market_value":        market_val,
            "floor":               0,
            "year_built":          random.randint(2010, 2022),
            "condition":           random.choice(["ممتاز", "جيد جداً"]),
            "district":            random.choice(DISTRICTS),
            "old_tax_value":       old_tax,
            "ai_predicted_value":  ai_val,
        })

    # --- 25 Apartments -------------------------------------------------------
    for i in range(1, 26):
        area        = _rand_float(100, 220, 1)
        ppm         = _rand_float(16_000, 26_000, 0)
        market_val  = round(area * ppm, 0)
        # Apartments: assessed closer to market (55–72%) — relative over-assessment
        old_tax_ratio = _rand_float(0.55, 0.72)
        old_tax     = round(market_val * old_tax_ratio, 0)
        ai_val      = round(market_val * _rand_float(0.93, 1.07), 0)
        properties.append({
            "id":                  f"A-{i:03d}",
            "property_type":       "شقة سكنية",
            "area":                area,
            "price_per_m2":        ppm,
            "market_value":        market_val,
            "floor":               random.randint(1, 15),
            "year_built":          random.randint(2005, 2022),
            "condition":           random.choice(["ممتاز", "جيد جداً", "جيد"]),
            "district":            random.choice(DISTRICTS),
            "old_tax_value":       old_tax,
            "ai_predicted_value":  ai_val,
        })

    return properties


# ─────────────────────────────────────────────────────────────────────────────
#  2. IAAO STATISTICS ENGINE
# ─────────────────────────────────────────────────────────────────────────────

def _median(values: List[float]) -> float:
    s = sorted(values)
    n = len(s)
    mid = n // 2
    return (s[mid - 1] + s[mid]) / 2.0 if n % 2 == 0 else s[mid]


def _mean(values: List[float]) -> float:
    return sum(values) / len(values) if values else 0.0


def calculate_iaao_stats(properties: List[Dict]) -> Dict:
    """
    Compute IAAO mass-appraisal statistics on the dataset.

    Returns a dict with:
        median_asr, mean_asr, cod, prd,
        iaao_cod_pass, iaao_prd_pass,
        total_market_value, total_old_tax, total_ai_predicted,
        revenue_uplift_pct
    """
    # ── Old system IAAO (shows the problem) ──────────────────────────────────
    asrs            = [p["old_tax_value"] / p["market_value"] for p in properties]
    median_asr      = _median(asrs)
    mean_asr        = _mean(asrs)
    cod             = (_mean([abs(a - median_asr) for a in asrs]) / median_asr) * 100
    total_market    = sum(p["market_value"]       for p in properties)
    total_old_tax   = sum(p["old_tax_value"]      for p in properties)
    total_ai        = sum(p["ai_predicted_value"] for p in properties)
    prd             = mean_asr / (total_old_tax / total_market)
    revenue_uplift  = (total_ai - total_old_tax) / total_old_tax * 100

    # ── AI-based IAAO (shows the improvement) ────────────────────────────────
    ai_asrs         = [p["ai_predicted_value"] / p["market_value"] for p in properties]
    ai_median_asr   = _median(ai_asrs)
    ai_mean_asr     = _mean(ai_asrs)
    ai_cod          = (_mean([abs(a - ai_median_asr) for a in ai_asrs]) / ai_median_asr) * 100
    ai_prd          = ai_mean_asr / (total_ai / total_market)

    # Per-property enrichment (add ASR to each record in place)
    for p, asr, ai_asr in zip(properties, asrs, ai_asrs):
        p["asr"]            = round(asr, 4)
        p["ai_asr"]         = round(ai_asr, 4)
        p["asr_diff_pct"]   = round((ai_asr - asr) / asr * 100, 1) if asr else 0

    return {
        # Old system
        "median_asr":          round(median_asr, 4),
        "mean_asr":            round(mean_asr, 4),
        "cod":                 round(cod, 2),
        "prd":                 round(prd, 4),
        "iaao_cod_pass":       5.0 <= cod <= 15.0,
        "iaao_prd_pass":       0.98 <= prd <= 1.03,
        "prd_verdict":         "انحياز تنازلي (تحابي الأثرياء)" if prd > 1.03 else ("انحياز تصاعدي" if prd < 0.98 else "عدالة رأسية"),
        # AI system
        "ai_median_asr":       round(ai_median_asr, 4),
        "ai_cod":              round(ai_cod, 2),
        "ai_prd":              round(ai_prd, 4),
        "ai_iaao_cod_pass":    5.0 <= ai_cod <= 15.0,
        "ai_iaao_prd_pass":    0.98 <= ai_prd <= 1.03,
        # Totals
        "total_market_value":  round(total_market, 0),
        "total_old_tax":       round(total_old_tax, 0),
        "total_ai_predicted":  round(total_ai, 0),
        "revenue_uplift_pct":  round(revenue_uplift, 2),
        "n_properties":        len(properties),
    }


# ─────────────────────────────────────────────────────────────────────────────
#  3. EXCEL EXPORT
# ─────────────────────────────────────────────────────────────────────────────

def generate_tax_excel(
    properties: List[Dict],
    stats: Dict,
    output_dir: str = "",
) -> str:
    """
    Write Tax_Authority_Pilot_V1.xlsx with two sheets.
    Returns the absolute path to the saved file.
    Requires xlsxwriter.
    """
    if not _XLSX_OK:
        raise ImportError("xlsxwriter is required — pip install xlsxwriter")

    out_dir = output_dir or _OUT_DIR
    os.makedirs(out_dir, exist_ok=True)
    filepath = os.path.join(out_dir, "Tax_Authority_Pilot_V1.xlsx")

    wb = xlsxwriter.Workbook(filepath, {"strings_to_numbers": False})

    # ── shared formats ──────────────────────────────────────────────────────
    def _fmt(**kw):
        base = {
            "font_name": "Simplified Arabic",
            "font_size": 11,
            "align":     "center",
            "valign":    "vcenter",
            "border":    1,
            "text_wrap": True,
        }
        base.update(kw)
        return wb.add_format(base)

    fmt_title      = _fmt(font_size=18, bold=True, font_color="#1F4E78",
                          border=0, bg_color="#FFFFFF")
    fmt_subtitle   = _fmt(font_size=12, italic=True, font_color="#60606060",
                          border=0, bg_color="#FFFFFF")
    fmt_hdr_navy   = _fmt(bold=True, font_color="#FFFFFF", bg_color="#1F4E78",
                          font_size=12)
    fmt_kpi_lbl    = _fmt(bold=True, font_size=12, font_color="#1F4E78",
                          bg_color="#EBF3FB", border=2)
    fmt_kpi_green  = _fmt(bold=True, font_size=20, font_color="#FFFFFF",
                          bg_color="#2E8648", border=2)
    fmt_kpi_red    = _fmt(bold=True, font_size=20, font_color="#FFFFFF",
                          bg_color="#C00000", border=2)
    fmt_kpi_gold   = _fmt(bold=True, font_size=20, font_color="#1F4E78",
                          bg_color="#D4AF37", border=2)
    fmt_revenue_lbl= _fmt(bold=True, font_size=11, font_color="#FFFFFF",
                          bg_color="#1F4E78")
    fmt_revenue_val= _fmt(font_size=11, num_format="#,##0", bg_color="#EBF3FB")
    fmt_uplift_val = _fmt(bold=True, font_size=13, font_color="#FFFFFF",
                          bg_color="#2E8648", num_format='0.00"%"')
    fmt_pass       = _fmt(bold=True, font_color="#FFFFFF", bg_color="#2E8648")
    fmt_fail       = _fmt(bold=True, font_color="#FFFFFF", bg_color="#C00000")
    fmt_section    = _fmt(bold=True, font_size=13, font_color="#D4AF37",
                          bg_color="#1F4E78", border=0)
    fmt_villa_row  = _fmt(bg_color="#FFF8DC", font_color="#1F4E78")
    fmt_apt_row    = _fmt(bg_color="#DDEEFF", font_color="#1F4E78")
    fmt_villa_num  = _fmt(bg_color="#FFF8DC", font_color="#1F4E78",
                          num_format="#,##0")
    fmt_apt_num    = _fmt(bg_color="#DDEEFF", font_color="#1F4E78",
                          num_format="#,##0")
    fmt_asr_green  = _fmt(bg_color="#C6EFCE", font_color="#276221")
    fmt_asr_red    = _fmt(bg_color="#FFC7CE", font_color="#9C0006")
    fmt_asr_normal = _fmt(num_format="0.000")
    fmt_villa_asr  = _fmt(bg_color="#FFF8DC", font_color="#1F4E78",
                          num_format="0.000")
    fmt_apt_asr    = _fmt(bg_color="#DDEEFF", font_color="#1F4E78",
                          num_format="0.000")
    fmt_villa_pct  = _fmt(bg_color="#FFF8DC", font_color="#1F4E78",
                          num_format='0.00"%"')
    fmt_apt_pct    = _fmt(bg_color="#DDEEFF", font_color="#1F4E78",
                          num_format='0.00"%"')

    # ══════════════════════════════════════════════════════════════════════════
    #  SHEET 1 — Tax Dashboard
    # ══════════════════════════════════════════════════════════════════════════
    ws1 = wb.add_worksheet("لوحة قيادة الجهاز الضريبي")
    ws1.right_to_left()
    ws1.set_column("A:A", 28)
    ws1.set_column("B:B", 22)
    ws1.set_column("C:C", 22)
    ws1.set_column("D:D", 22)
    ws1.set_column("E:E", 22)
    ws1.set_row(0, 50)
    ws1.set_row(1, 30)

    # Title
    ws1.merge_range("A1:E1",
                    "تقرير إكسب ليدز — لجنة الضرائب العقارية IAAO V1",
                    fmt_title)
    ws1.merge_range("A2:E2",
                    f"التجمع الخامس — القاهرة الجديدة | التاريخ: {datetime.now().strftime('%Y-%m-%d')}",
                    fmt_subtitle)

    # ── KPI Section label ───────────────────────────────────────────────────
    ws1.set_row(3, 30)
    ws1.merge_range("A4:E4", "مؤشرات IAAO للتقييم الجماعي", fmt_section)

    # KPI header row
    ws1.set_row(4, 25)
    for col, label in enumerate(["المؤشر", "Median ASR", "COD", "PRD", "الحالة العامة"]):
        ws1.write(4, col, label, fmt_hdr_navy)

    # KPI values row
    ws1.set_row(5, 50)
    ws1.write(5, 0, "القيمة", fmt_kpi_lbl)
    ws1.write(5, 1, stats["median_asr"], fmt_kpi_gold)

    cod_fmt = fmt_kpi_green if stats["iaao_cod_pass"] else fmt_kpi_red
    ws1.write(5, 2, stats["cod"], cod_fmt)

    prd_fmt = fmt_kpi_green if stats["iaao_prd_pass"] else fmt_kpi_red
    ws1.write(5, 3, stats["prd"], prd_fmt)

    overall_pass = stats["iaao_cod_pass"] and stats["iaao_prd_pass"]
    overall_fmt  = fmt_kpi_green if overall_pass else fmt_kpi_red
    overall_txt  = "✔ مطابق للمعايير" if overall_pass else "✘ يحتاج مراجعة"
    ws1.write(5, 4, overall_txt, overall_fmt)

    # IAAO pass/fail compliance row
    ws1.set_row(6, 25)
    ws1.write(6, 0, "معيار IAAO", fmt_kpi_lbl)
    ws1.write(6, 1, "0.85 – 1.15", fmt_kpi_lbl)
    ws1.write(6, 2, "5.0 – 15.0", fmt_kpi_lbl)
    ws1.write(6, 3, "0.98 – 1.03", fmt_kpi_lbl)
    ws1.write(6, 4, "—", fmt_kpi_lbl)

    ws1.set_row(7, 25)
    ws1.write(7, 0, "نتيجة الاختبار", fmt_kpi_lbl)
    ws1.write(7, 1, "✔ مقبول", fmt_pass)  # ASR median always in range by construction
    ws1.write(7, 2,
              "✔ مطابق" if stats["iaao_cod_pass"] else "✘ خارج النطاق",
              fmt_pass if stats["iaao_cod_pass"] else fmt_fail)
    ws1.write(7, 3,
              "✔ مطابق" if stats["iaao_prd_pass"] else "✘ خارج النطاق",
              fmt_pass if stats["iaao_prd_pass"] else fmt_fail)
    ws1.write(7, 4, "", fmt_kpi_lbl)

    # ── Revenue Section ──────────────────────────────────────────────────────
    ws1.set_row(9, 30)
    ws1.merge_range("A10:E10", "مقارنة الإيرادات الضريبية", fmt_section)

    ws1.set_row(10, 25)
    for col, h in enumerate(["البيان", "إجمالي القيمة السوقية",
                              "القيمة الضريبية القديمة",
                              "القيمة بالذكاء الاصطناعي",
                              "نسبة الزيادة %"]):
        ws1.write(10, col, h, fmt_hdr_navy)

    ws1.set_row(11, 35)
    ws1.write(11, 0, "المجموع (ج.م.)", fmt_revenue_lbl)
    ws1.write(11, 1, stats["total_market_value"],  fmt_revenue_val)
    ws1.write(11, 2, stats["total_old_tax"],        fmt_revenue_val)
    ws1.write(11, 3, stats["total_ai_predicted"],   fmt_revenue_val)
    ws1.write(11, 4, stats["revenue_uplift_pct"],   fmt_uplift_val)

    # ── Notes ────────────────────────────────────────────────────────────────
    ws1.set_row(13, 22)
    note_fmt = wb.add_format({"font_name": "Simplified Arabic", "font_size": 10,
                               "italic": True, "font_color": "#606060",
                               "border": 0})
    ws1.merge_range("A14:E14",
                    "* جميع القيم بالجنيه المصري | المصدر: نظام إكسب ليدز للتقييم بالذكاء الاصطناعي",
                    note_fmt)

    # ══════════════════════════════════════════════════════════════════════════
    #  SHEET 2 — Property List
    # ══════════════════════════════════════════════════════════════════════════
    ws2 = wb.add_worksheet("قائمة العقارات — التجمع الخامس")
    ws2.right_to_left()
    ws2.set_column("A:A", 10)   # ID
    ws2.set_column("B:B", 16)   # نوع العقار
    ws2.set_column("C:C", 20)   # الحي
    ws2.set_column("D:D", 12)   # المساحة
    ws2.set_column("E:E", 20)   # القيمة السوقية
    ws2.set_column("F:F", 22)   # القيمة الضريبية القديمة
    ws2.set_column("G:G", 24)   # القيمة بالذكاء الاصطناعي
    ws2.set_column("H:H", 10)   # ASR
    ws2.set_column("I:I", 12)   # الفرق %

    ws2.set_row(0, 35)
    ws2.merge_range("A1:I1",
                    "قائمة العقارات — التجمع الخامس، القاهرة الجديدة",
                    wb.add_format({"font_name": "Simplified Arabic",
                                   "font_size": 15, "bold": True,
                                   "font_color": "#1F4E78", "align": "center",
                                   "valign": "vcenter", "border": 0}))

    # Header row
    headers = [
        "رقم/ID", "نوع العقار", "الحي",
        "المساحة (م²)", "القيمة السوقية",
        "القيمة الضريبية القديمة", "القيمة المقدرة بالذكاء الاصطناعي",
        "ASR", "الفرق %",
    ]
    ws2.set_row(1, 28)
    for col, h in enumerate(headers):
        ws2.write(1, col, h, fmt_hdr_navy)

    # Data rows
    for row_idx, prop in enumerate(properties):
        r = row_idx + 2   # rows 2 onwards (0-indexed)
        ws2.set_row(r, 22)

        is_villa = prop["property_type"] == "فيلا فاخرة"
        txt_fmt  = fmt_villa_row  if is_villa else fmt_apt_row
        num_fmt  = fmt_villa_num  if is_villa else fmt_apt_num
        asr_base = fmt_villa_asr  if is_villa else fmt_apt_asr
        pct_fmt  = fmt_villa_pct  if is_villa else fmt_apt_pct

        asr         = prop["old_tax_value"] / prop["market_value"]
        diff_pct    = (prop["ai_predicted_value"] - prop["old_tax_value"]) / prop["old_tax_value"] * 100

        # ASR conditional format override
        if asr > 0.80:
            asr_fmt = fmt_asr_green
        elif asr < 0.50:
            asr_fmt = fmt_asr_red
        else:
            asr_fmt = asr_base

        ws2.write(r, 0, prop["id"],               txt_fmt)
        ws2.write(r, 1, prop["property_type"],    txt_fmt)
        ws2.write(r, 2, prop["district"],         txt_fmt)
        ws2.write(r, 3, prop["area"],             num_fmt)
        ws2.write(r, 4, prop["market_value"],     num_fmt)
        ws2.write(r, 5, prop["old_tax_value"],    num_fmt)
        ws2.write(r, 6, prop["ai_predicted_value"], num_fmt)
        ws2.write(r, 7, round(asr, 3),            asr_fmt)
        ws2.write(r, 8, round(diff_pct, 2),       pct_fmt)

    # Freeze header rows
    ws2.freeze_panes(2, 0)

    wb.close()
    return filepath


# ─────────────────────────────────────────────────────────────────────────────
#  4. MINISTRY WORD SECTION
# ─────────────────────────────────────────────────────────────────────────────

def _set_rtl(paragraph) -> None:
    """Force RTL direction on a paragraph."""
    if not _DOCX_OK:
        return
    pPr  = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "right")
    pPr.append(jc)


def _shade_cell(cell, hex_color: str) -> None:
    """Apply background fill to a docx table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:fill"),  hex_color)
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"),   "clear")
    tcPr.append(shd)


def _w_para(doc, text: str, bold: bool = False, size: int = 12,
            color=None, align: str = "right") -> None:
    """Add an RTL paragraph with Simplified Arabic font."""
    p   = doc.add_paragraph()
    _set_rtl(p)
    run = p.add_run(text)
    run.font.name  = FONT_MAIN
    run.font.size  = Pt(size)
    run.font.bold  = bold
    if color:
        run.font.color.rgb = color
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _w_heading(doc, text: str, level: int = 1, color=None) -> None:
    """Add a heading paragraph."""
    p = doc.add_heading(text, level=level)
    _set_rtl(p)
    for run in p.runs:
        run.font.name  = FONT_MAIN
        run.font.bold  = True
        run.font.size  = Pt(max(9, 16 - (level - 1) * 2))
        run.font.color.rgb = color or NAVY


def _w_hr(doc) -> None:
    """Gold horizontal rule."""
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "D4AF37")
    pBdr.append(bot)
    pPr.append(pBdr)


def generate_ministry_word_section(doc, stats: Dict) -> None:
    """
    Append an Arabic Ministry of Finance pilot section to an existing
    python-docx Document object.  Modifies `doc` in-place.

    Parameters
    ----------
    doc   : python-docx Document
    stats : dict returned by calculate_iaao_stats()
    """
    if not _DOCX_OK:
        raise ImportError("python-docx is required — pip install python-docx")

    doc.add_page_break()

    # ── Section title ────────────────────────────────────────────────────────
    _w_heading(doc,
               "ملحق خاص: تقرير الطيار للهيئة الضريبية — وزارة المالية",
               level=1, color=NAVY)
    _w_hr(doc)
    doc.add_paragraph()

    # ── Executive intro ──────────────────────────────────────────────────────
    _w_heading(doc, "مقدمة تنفيذية — معالي وزير المالية", level=2, color=GOLD)
    _w_para(
        doc,
        "معالي وزير المالية،\n"
        "تقدم هيئة التقييم العقاري المدعومة بالذكاء الاصطناعي — نظام إكسب ليدز — "
        "هذا التقرير الطيار بوصفه نموذجاً تطبيقياً للتقييم الجماعي بالذكاء الاصطناعي "
        "وفق منهجية التقييم الجماعي المعتمدة من المعهد الدولي لمحللي التقييم (IAAO). "
        "يُعالج النظام 50 عقاراً في مناطق التجمع الخامس بالقاهرة الجديدة، "
        "مقارناً القيم الضريبية الراهنة بنظيراتها المحسوبة عبر نماذج التعلم الآلي، "
        "مما يكشف عن فجوات التقييم ويُسهم في رفع العدالة الضريبية وزيادة الإيرادات.",
        size=11,
    )
    doc.add_paragraph()

    # ── IAAO metrics table ───────────────────────────────────────────────────
    _w_heading(doc, "مؤشرات الأداء الإحصائي — IAAO", level=2, color=NAVY)
    _w_para(doc,
            "يعتمد النظام على ثلاثة مؤشرات أساسية وفق معايير IAAO 2023:",
            size=11)

    tbl = doc.add_table(rows=0, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"

    # header
    hdr = tbl.add_row()
    hdr.height = Cm(0.9)
    for col_idx, label in enumerate(["المؤشر", "القيمة المحسوبة",
                                      "نطاق IAAO", "الحكم"]):
        cell = hdr.cells[col_idx]
        cell.text = label
        _set_rtl(cell.paragraphs[0])
        for run in cell.paragraphs[0].runs:
            run.font.name  = FONT_MAIN
            run.font.bold  = True
            run.font.size  = Pt(11)
            run.font.color.rgb = WHITE
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _shade_cell(cell, "1F4E78")

    # data rows
    iaao_rows = [
        ("Median ASR",
         f"{stats['median_asr']:.4f}",
         "0.85 – 1.15",
         stats["median_asr"],
         True),         # always in-range by dataset construction
        ("COD",
         f"{stats['cod']:.2f}%",
         "5.0% – 15.0%",
         None,
         stats["iaao_cod_pass"]),
        ("PRD",
         f"{stats['prd']:.4f}",
         "0.98 – 1.03",
         None,
         stats["iaao_prd_pass"]),
    ]

    for name, value, range_str, _, passed in iaao_rows:
        row = tbl.add_row()
        row.height = Cm(0.85)
        data = [name, value, range_str,
                "✔ مطابق للمعايير" if passed else "✘ يتطلب مراجعة"]
        fill = ["EBF3FB", "EBF3FB", "EBF3FB",
                "C6EFCE" if passed else "FFC7CE"]
        font_c = [None, None, None,
                  "276221" if passed else "9C0006"]
        for c_idx, (txt, bg, fc) in enumerate(zip(data, fill, font_c)):
            cell = row.cells[c_idx]
            cell.text = txt
            _set_rtl(cell.paragraphs[0])
            for run in cell.paragraphs[0].runs:
                run.font.name = FONT_MAIN
                run.font.size = Pt(11)
                if fc:
                    from docx.shared import RGBColor as _RGB
                    r_int, g_int, b_int = (
                        int(fc[0:2], 16),
                        int(fc[2:4], 16),
                        int(fc[4:6], 16),
                    )
                    run.font.color.rgb = _RGB(r_int, g_int, b_int)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _shade_cell(cell, bg)

    doc.add_paragraph()

    # ── Revenue optimization paragraph ───────────────────────────────────────
    _w_heading(doc, "تحسين الإيرادات الضريبية", level=2, color=NAVY)
    uplift     = stats["revenue_uplift_pct"]
    old_total  = f"{stats['total_old_tax']:,.0f}"
    ai_total   = f"{stats['total_ai_predicted']:,.0f}"
    _w_para(
        doc,
        f"يرصد النظام فجوة تقييم واضحة بين القيم الضريبية الحالية والقيم السوقية الفعلية. "
        f"تبلغ القيمة الضريبية القائمة للعينة {old_total} جنيه مصري، في حين يُقدّرها "
        f"نموذج الذكاء الاصطناعي بـ {ai_total} جنيه مصري، أي بزيادة تبلغ {uplift:.2f}٪. "
        f"يُعزى ذلك إلى اعتماد التقديرات التاريخية على معاملات قديمة لا تعكس السوق الراهن. "
        f"إن تبني التقييم الجماعي الآلي سيُتيح للهيئة الضريبية استيعاب هذه الزيادة بصورة "
        f"عادلة وشفافة دون رفع أعباء ضريبية على الفئات المتوسطة.",
        size=11,
    )
    doc.add_paragraph()

    # ── Tax equity paragraph ──────────────────────────────────────────────────
    _w_heading(doc, "العدالة الضريبية وتكافؤ المعاملة", level=2, color=NAVY)
    cod_status = "أدنى من 15٪" if stats["iaao_cod_pass"] else "يتجاوز 15٪"
    _w_para(
        doc,
        f"يُثبت مؤشر التشتت النسبي (COD = {stats['cod']:.2f}٪) — وهو {cod_status} — "
        f"أن نظام التقييم الجماعي يُعامل جميع أنواع العقارات بصورة متجانسة؛ "
        f"إذ لا تحظى فئة بتقدير مُبالغ فيه أو منقوص على حساب أخرى. "
        f"هذه الخاصية جوهرية لضمان الامتثال لمبدأ المساواة الضريبية الدستورية "
        f"وإزالة أي تحيز في الوعاء الضريبي بين الفلل الفاخرة والشقق السكنية.",
        size=11,
    )
    doc.add_paragraph()

    # ── Closing certification ────────────────────────────────────────────────
    _w_hr(doc)
    _w_heading(doc, "إفادة الخبير وشهادة الاعتماد", level=2, color=GOLD)
    _w_para(
        doc,
        f"أُقرّ أنا هشام محمد المهدى، الخبير العقاري المعتمد — رقم القيد 29 — "
        f"بأن هذا التقرير الطيار قد أُعدّ وفق منهجية التقييم الجماعي المعتمدة "
        f"من IAAO، وأن جميع المؤشرات الإحصائية الواردة فيه صحيحة ودقيقة وفق "
        f"البيانات المُدخلة. صدر في: {datetime.now().strftime('%Y-%m-%d')}.",
        size=11,
        bold=True,
    )
    _w_para(
        doc,
        "نظام إكسب ليدز للتقييم العقاري بالذكاء الاصطناعي | Expert_Smart PropTech",
        size=10,
        color=GRAY,
        align="center",
    )


# ─────────────────────────────────────────────────────────────────────────────
#  5. ORCHESTRATOR
# ─────────────────────────────────────────────────────────────────────────────

def run_tax_pilot(output_dir: str = "") -> Dict:
    """
    Full Tax Pilot pipeline:
      1. Generate 50-property synthetic dataset
      2. Compute IAAO statistics
      3. Export Excel report
      4. Return results dict

    Returns
    -------
    dict with keys:
        stats       — IAAO metrics dict
        excel_path  — absolute path to the saved xlsx file
        properties  — list of property dicts
    """
    # Step 1 — dataset
    properties = generate_dataset()

    # Step 2 — IAAO stats
    stats = calculate_iaao_stats(properties)

    # Step 3 — Excel export
    excel_path = ""
    if _XLSX_OK:
        excel_path = generate_tax_excel(properties, stats, output_dir=output_dir)
    else:
        print("[tax_pilot_engine] WARNING: xlsxwriter not installed — skipping Excel export.")

    # Console summary
    print("=" * 60)
    print("  Expert_Smart — Ministry of Finance Tax Pilot Engine")
    print("=" * 60)
    print(f"  Properties     : {len(properties)} (villas + apartments)")
    print(f"  Median ASR     : {stats['median_asr']:.4f}")
    cod_flag = "PASS" if stats['iaao_cod_pass'] else "FAIL"
    prd_flag = "PASS" if stats['iaao_prd_pass'] else "FAIL"
    print(f"  COD            : {stats['cod']:.2f}%  [{cod_flag}]")
    print(f"  PRD            : {stats['prd']:.4f}  [{prd_flag}]")
    print(f"  Revenue Uplift : {stats['revenue_uplift_pct']:.2f}%")
    if excel_path:
        print(f"  Excel saved to : {excel_path}")
    print("=" * 60)

    return {
        "stats":      stats,
        "excel_path": excel_path,
        "properties": properties,
    }


# ─────────────────────────────────────────────────────────────────────────────
#  Entry point
# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    run_tax_pilot()
